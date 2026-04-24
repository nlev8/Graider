"""Tests for flashcard generation and export."""

import json
import os
import tempfile
import pytest
from unittest.mock import patch, MagicMock
from flask import Flask, g


def _make_app():
    """Create a minimal Flask app with planner routes."""
    app = Flask(__name__)
    app.config['TESTING'] = True
    app.config['SECRET_KEY'] = 'test'
    app.config['RATELIMIT_ENABLED'] = False

    from backend.extensions import limiter
    limiter.init_app(app)

    from backend.routes.planner_routes import planner_bp
    app.register_blueprint(planner_bp)

    @app.before_request
    def set_test_user():
        from flask import request as req
        teacher_id = req.headers.get("X-Test-Teacher-Id")
        if teacher_id:
            g.user_id = teacher_id

    return app


def _mock_genai_response(text):
    """Create a mock response compatible with GeminiAdapter's genai.Client."""
    mock_resp = MagicMock()
    mock_resp.text = text
    # NOTE: MagicMock(name=...) sets the mock's repr name, NOT a `.name` attribute.
    # Gemini's finish_reason is accessed as `fr.name`, so set it explicitly via attribute assignment.
    fr = MagicMock()
    fr.name = "STOP"
    mock_resp.candidates = [MagicMock(finish_reason=fr)]
    mock_resp.usage_metadata = MagicMock(prompt_token_count=100, candidates_token_count=300)
    return mock_resp


SAMPLE_FLASHCARDS = json.dumps({
    "title": "Unit 3: The Constitution",
    "cards": [
        {"term": "Federalism", "definition": "Division of power between national and state governments."},
        {"term": "Ratification", "definition": "Formal approval of the Constitution by the states."},
        {"term": "Amendment", "definition": "A change or addition to the Constitution."},
        {"term": "Bill of Rights", "definition": "The first ten amendments to the Constitution, guaranteeing individual rights."},
        {"term": "Separation of Powers", "definition": "Division of government into three branches: legislative, executive, and judicial."},
        {"term": "Checks and Balances", "definition": "System where each branch can limit the powers of the other branches."},
    ]
})


class TestGenerateFlashcards:
    def test_returns_flashcards_json(self):
        """Should return structured flashcards from Gemini."""
        app = _make_app()
        with app.test_client() as client:
            with patch('backend.api_keys.get_api_key', return_value='fake-key'), \
                 patch('backend.services.llm_adapter.gemini_adapter.genai') as mock_genai:
                mock_client = MagicMock()
                mock_client.models.generate_content.return_value = _mock_genai_response(SAMPLE_FLASHCARDS)
                mock_genai.Client.return_value = mock_client

                resp = client.post('/api/generate-flashcards', json={
                    "title": "Constitution Flashcards",
                    "content": "Federalism is the division of power...",
                    "subject": "US History",
                    "grade": "8",
                }, headers={"X-Test-Teacher-Id": "teacher-1"})

            assert resp.status_code == 200
            data = resp.get_json()
            assert "flashcards" in data
            assert len(data["flashcards"]["cards"]) == 6
            assert data["flashcards"]["cards"][0]["term"] == "Federalism"

    def test_returns_400_without_content(self):
        """Should reject requests without content."""
        app = _make_app()
        with app.test_client() as client:
            resp = client.post('/api/generate-flashcards', json={
                "title": "Empty",
            }, headers={"X-Test-Teacher-Id": "teacher-1"})
        assert resp.status_code == 400

    def test_handles_gemini_error(self):
        """Should return 500 on Gemini failure."""
        app = _make_app()
        with app.test_client() as client:
            with patch('backend.api_keys.get_api_key', return_value='fake-key'), \
                 patch('backend.services.llm_adapter.gemini_adapter.genai') as mock_genai:
                mock_client = MagicMock()
                mock_client.models.generate_content.side_effect = Exception("API error")
                mock_genai.Client.return_value = mock_client

                resp = client.post('/api/generate-flashcards', json={
                    "title": "Broken",
                    "content": "Some content",
                    "subject": "Math",
                    "grade": "7",
                }, headers={"X-Test-Teacher-Id": "teacher-1"})

            assert resp.status_code == 500

    def test_includes_lesson_plan_vocab(self):
        """Should incorporate lesson plan vocabulary into the prompt."""
        app = _make_app()
        captured_prompt = []

        def capture(model=None, contents=None, config=None, **kwargs):
            prompt_text = contents[0]["parts"][0]["text"] if contents else ""
            captured_prompt.append(prompt_text)
            return _mock_genai_response(SAMPLE_FLASHCARDS)

        with app.test_client() as client:
            with patch('backend.api_keys.get_api_key', return_value='fake-key'), \
                 patch('backend.services.llm_adapter.gemini_adapter.genai') as mock_genai:
                mock_client = MagicMock()
                mock_client.models.generate_content.side_effect = capture
                mock_genai.Client.return_value = mock_client

                resp = client.post('/api/generate-flashcards', json={
                    "title": "Vocab Cards",
                    "content": "Day 1: Introduction to the Constitution",
                    "lessonPlan": {
                        "title": "Constitution Unit",
                        "vocabulary": ["federalism", "ratification", "amendment"],
                    },
                    "subject": "US History",
                    "grade": "8",
                }, headers={"X-Test-Teacher-Id": "teacher-1"})

            assert resp.status_code == 200
            assert captured_prompt, "generate_content was not called"
            assert "federalism" in captured_prompt[0]

    def test_respects_global_ai_notes(self):
        """Should include globalAINotes in the prompt."""
        app = _make_app()
        captured_prompt = []

        def capture(model=None, contents=None, config=None, **kwargs):
            prompt_text = contents[0]["parts"][0]["text"] if contents else ""
            captured_prompt.append(prompt_text)
            return _mock_genai_response(SAMPLE_FLASHCARDS)

        with app.test_client() as client:
            with patch('backend.api_keys.get_api_key', return_value='fake-key'), \
                 patch('backend.services.llm_adapter.gemini_adapter.genai') as mock_genai:
                mock_client = MagicMock()
                mock_client.models.generate_content.side_effect = capture
                mock_genai.Client.return_value = mock_client

                resp = client.post('/api/generate-flashcards', json={
                    "title": "Cards",
                    "content": "Some content",
                    "subject": "Science",
                    "grade": "6",
                    "globalAINotes": "Use simplified vocabulary for grade 6",
                }, headers={"X-Test-Teacher-Id": "teacher-1"})

            assert resp.status_code == 200
            assert captured_prompt, "generate_content was not called"
            assert "simplified vocabulary" in captured_prompt[0]

    def test_custom_card_count(self):
        """Should pass card count to the prompt."""
        app = _make_app()
        captured_prompt = []

        def capture(model=None, contents=None, config=None, **kwargs):
            prompt_text = contents[0]["parts"][0]["text"] if contents else ""
            captured_prompt.append(prompt_text)
            return _mock_genai_response(SAMPLE_FLASHCARDS)

        with app.test_client() as client:
            with patch('backend.api_keys.get_api_key', return_value='fake-key'), \
                 patch('backend.services.llm_adapter.gemini_adapter.genai') as mock_genai:
                mock_client = MagicMock()
                mock_client.models.generate_content.side_effect = capture
                mock_genai.Client.return_value = mock_client

                resp = client.post('/api/generate-flashcards', json={
                    "title": "Cards",
                    "content": "Some content",
                    "subject": "Math",
                    "grade": "7",
                    "cardCount": 20,
                }, headers={"X-Test-Teacher-Id": "teacher-1"})

            assert resp.status_code == 200
            assert captured_prompt, "generate_content was not called"
            assert "20" in captured_prompt[0]


class TestExportFlashcards:
    def test_exports_pdf(self):
        """Should export flashcards as PDF with card grid."""
        app = _make_app()
        cards = json.loads(SAMPLE_FLASHCARDS)
        with app.test_client() as client:
            with patch('backend.routes.planner_routes._get_export_dir',
                       return_value=tempfile.mkdtemp()):
                resp = client.post('/api/export-flashcards', json={
                    "flashcards": cards,
                    "format": "pdf",
                }, headers={"X-Test-Teacher-Id": "teacher-1"})

        assert resp.status_code == 200
        assert resp.content_type == 'application/pdf'
        assert resp.data[:5] == b'%PDF-'

    def test_exports_docx(self):
        """Should export flashcards as DOCX with term/definition table."""
        app = _make_app()
        cards = json.loads(SAMPLE_FLASHCARDS)
        with app.test_client() as client:
            with patch('backend.routes.planner_routes._get_export_dir',
                       return_value=tempfile.mkdtemp()):
                resp = client.post('/api/export-flashcards', json={
                    "flashcards": cards,
                    "format": "docx",
                }, headers={"X-Test-Teacher-Id": "teacher-1"})

        assert resp.status_code == 200
        assert 'wordprocessingml' in resp.content_type

    def test_docx_contains_terms(self):
        """Exported DOCX should contain the flashcard terms."""
        from docx import Document
        from io import BytesIO
        app = _make_app()
        cards = json.loads(SAMPLE_FLASHCARDS)
        with app.test_client() as client:
            with patch('backend.routes.planner_routes._get_export_dir',
                       return_value=tempfile.mkdtemp()):
                resp = client.post('/api/export-flashcards', json={
                    "flashcards": cards,
                    "format": "docx",
                }, headers={"X-Test-Teacher-Id": "teacher-1"})

        doc = Document(BytesIO(resp.data))
        text = "\n".join(p.text for p in doc.paragraphs)
        tables_text = "\n".join(
            cell.text for table in doc.tables for row in table.rows for cell in row.cells
        )
        all_text = text + tables_text
        assert "Federalism" in all_text
        assert "Ratification" in all_text

    def test_returns_400_without_flashcards(self):
        """Should reject requests without flashcard data."""
        app = _make_app()
        with app.test_client() as client:
            resp = client.post('/api/export-flashcards', json={
                "format": "pdf",
            }, headers={"X-Test-Teacher-Id": "teacher-1"})
        assert resp.status_code == 400

    def test_defaults_to_pdf(self):
        """Should default to PDF when no format specified."""
        app = _make_app()
        cards = json.loads(SAMPLE_FLASHCARDS)
        with app.test_client() as client:
            with patch('backend.routes.planner_routes._get_export_dir',
                       return_value=tempfile.mkdtemp()):
                resp = client.post('/api/export-flashcards', json={
                    "flashcards": cards,
                }, headers={"X-Test-Teacher-Id": "teacher-1"})

        assert resp.status_code == 200
        assert resp.content_type == 'application/pdf'

    def test_empty_cards_does_not_crash(self):
        """Should handle empty card list gracefully."""
        app = _make_app()
        cards = {"title": "Empty Set", "cards": []}
        with app.test_client() as client:
            with patch('backend.routes.planner_routes._get_export_dir',
                       return_value=tempfile.mkdtemp()):
                resp = client.post('/api/export-flashcards', json={
                    "flashcards": cards,
                    "format": "pdf",
                }, headers={"X-Test-Teacher-Id": "teacher-1"})

        assert resp.status_code == 200
