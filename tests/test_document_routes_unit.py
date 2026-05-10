"""Unit tests for backend/routes/document_routes.py.

Audit MAJOR #4 sprint follow-up to PR #285. Targets the 88 uncovered LOC
(28% baseline). Single endpoint /api/parse-document + 3 internal
helpers (_parse_docx with mammoth + python-docx fallback, _parse_pdf
with PyMuPDF, _parse_txt).

Strategy
--------
Flask test_client + extensive library mocks for `mammoth`,
`docx.Document`, and `fitz` (PyMuPDF). Each helper path covered.
File uploads via `test_client.post(..., data={'file': (BytesIO, 'name.ext')})`.
"""
from __future__ import annotations

import io
from unittest.mock import patch, MagicMock

import pytest


@pytest.fixture
def client():
    from backend.app import app
    from backend.extensions import limiter
    try:
        limiter.reset()
    except Exception:
        pass
    with app.test_client() as c:
        yield c


@pytest.fixture
def auth_headers():
    return {"X-Test-Teacher-Id": "teach-1"}


@pytest.fixture(autouse=True)
def dev_env(monkeypatch):
    monkeypatch.setenv("FLASK_ENV", "development")


# ──────────────────────────────────────────────────────────────────
# /api/parse-document
# ──────────────────────────────────────────────────────────────────


class TestParseDocumentDispatch:
    def test_no_file_returns_400(self, client, auth_headers):
        resp = client.post("/api/parse-document", headers=auth_headers)
        assert resp.status_code == 400
        assert "No file uploaded" in resp.get_json()["error"]

    def test_unsupported_extension_returns_400(self, client, auth_headers):
        resp = client.post(
            "/api/parse-document",
            data={"file": (io.BytesIO(b"data"), "document.xyz")},
            headers=auth_headers,
        )
        assert resp.status_code == 400
        body = resp.get_json()
        assert "Unsupported file type" in body["error"]
        assert ".xyz" in body["error"]

    def test_no_extension_returns_400(self, client, auth_headers):
        # Bare filename without an extension at all
        resp = client.post(
            "/api/parse-document",
            data={"file": (io.BytesIO(b"data"), "noextension")},
            headers=auth_headers,
        )
        assert resp.status_code == 400

    def test_bare_dotfile_uses_filename_as_extension(self, client, auth_headers):
        # Filename like ".txt" hits the `not ext and filename.startswith('.')`
        # branch (line 33). os.path.splitext('.txt') → ('.txt','') so ext=''
        # then ext gets reassigned to '.txt'. filename.endswith('.txt') is
        # True → dispatches to _parse_txt and decodes successfully.
        resp = client.post(
            "/api/parse-document",
            data={"file": (io.BytesIO(b"hello world"), ".txt")},
            headers=auth_headers,
        )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["text"] == "hello world"
        assert body["type"] == "html"


# ──────────────────────────────────────────────────────────────────
# .docx parsing — mammoth path
# ──────────────────────────────────────────────────────────────────


class TestParseDocx:
    def test_mammoth_path_with_metadata_title(self, client, auth_headers):
        # Mock mammoth.convert_to_html
        mock_mammoth_result = MagicMock()
        mock_mammoth_result.value = "<p>Hello world</p>"

        # Mock python-docx Document
        mock_para = MagicMock()
        mock_para.text = "Hello world"
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para]
        mock_doc.tables = []
        mock_doc.core_properties.title = "My Custom Title"

        mock_mammoth = MagicMock()
        mock_mammoth.convert_to_html.return_value = mock_mammoth_result

        with patch.dict("sys.modules", {"mammoth": mock_mammoth}), \
             patch("docx.Document", return_value=mock_doc):
            resp = client.post(
                "/api/parse-document",
                data={"file": (io.BytesIO(b"fake docx"), "test.docx")},
                headers=auth_headers,
            )

        assert resp.status_code == 200
        body = resp.get_json()
        assert body["type"] == "html"
        assert body["filename"] == "test.docx"
        assert "Hello world" in body["html"]
        assert body["text"] == "Hello world"
        # Metadata title wins
        assert body["doc_title"] == "My Custom Title"

    def test_mammoth_path_fallback_to_first_para_when_no_metadata_title(
        self, client, auth_headers,
    ):
        mock_mammoth_result = MagicMock()
        mock_mammoth_result.value = "<p>Worksheet</p>"

        mock_para1 = MagicMock()
        mock_para1.text = "Worksheet Title"
        mock_para2 = MagicMock()
        mock_para2.text = "Question 1"
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para1, mock_para2]
        mock_doc.tables = []
        mock_doc.core_properties.title = ""

        mock_mammoth = MagicMock()
        mock_mammoth.convert_to_html.return_value = mock_mammoth_result

        with patch.dict("sys.modules", {"mammoth": mock_mammoth}), \
             patch("docx.Document", return_value=mock_doc):
            resp = client.post(
                "/api/parse-document",
                data={"file": (io.BytesIO(b"fake"), "x.docx")},
                headers=auth_headers,
            )

        body = resp.get_json()
        # Fallback: first non-empty line is the title
        assert body["doc_title"] == "Worksheet Title"

    def test_mammoth_path_with_table(self, client, auth_headers):
        # Document with a table — exercises the table-row joining
        mock_mammoth_result = MagicMock()
        mock_mammoth_result.value = "<table><tr><td>A</td></tr></table>"

        # Cells in a row
        cell_a = MagicMock(); cell_a.text = "Header A"
        cell_b = MagicMock(); cell_b.text = "Header B"
        cell_data = MagicMock(); cell_data.text = "Data"
        cell_empty = MagicMock(); cell_empty.text = ""  # filtered out

        row1 = MagicMock(); row1.cells = [cell_a, cell_b]
        row2 = MagicMock(); row2.cells = [cell_data, cell_empty]

        mock_table = MagicMock()
        mock_table.rows = [row1, row2]

        mock_doc = MagicMock()
        mock_doc.paragraphs = []
        mock_doc.tables = [mock_table]
        mock_doc.core_properties.title = "Tab Doc"

        mock_mammoth = MagicMock()
        mock_mammoth.convert_to_html.return_value = mock_mammoth_result

        with patch.dict("sys.modules", {"mammoth": mock_mammoth}), \
             patch("docx.Document", return_value=mock_doc):
            resp = client.post(
                "/api/parse-document",
                data={"file": (io.BytesIO(b"x"), "tabdoc.docx")},
                headers=auth_headers,
            )

        body = resp.get_json()
        # Both header cells joined; data row has only 1 non-empty cell
        assert "Header A | Header B" in body["text"]
        assert "Data" in body["text"]

    def test_mammoth_path_metadata_exception_falls_back(
        self, client, auth_headers,
    ):
        # If accessing core_properties.title raises, the except sentry
        # branch fires and falls back to first paragraph.
        mock_result = MagicMock()
        mock_result.value = "<p>x</p>"

        mock_para = MagicMock(); mock_para.text = "Fallback Title"
        mock_doc = MagicMock()
        mock_doc.paragraphs = [mock_para]
        mock_doc.tables = []
        # Make accessing core_properties.title raise
        type(mock_doc).core_properties = property(
            lambda self: (_ for _ in ()).throw(RuntimeError("metadata error"))
        )

        mock_mammoth = MagicMock()
        mock_mammoth.convert_to_html.return_value = mock_result

        with patch.dict("sys.modules", {"mammoth": mock_mammoth}), \
             patch("docx.Document", return_value=mock_doc), \
             patch("backend.routes.document_routes.sentry_sdk.capture_exception"):
            resp = client.post(
                "/api/parse-document",
                data={"file": (io.BytesIO(b"x"), "x.docx")},
                headers=auth_headers,
            )

        body = resp.get_json()
        # Metadata exception swallowed; falls back to first paragraph
        assert body["doc_title"] == "Fallback Title"

    def test_mammoth_path_skips_empty_paragraphs(self, client, auth_headers):
        # Empty paragraphs are filtered out of the plain_text result
        mock_result = MagicMock(); mock_result.value = "<p>x</p>"

        empty_para = MagicMock(); empty_para.text = "   "  # whitespace only
        good_para = MagicMock(); good_para.text = "Real content"

        mock_doc = MagicMock()
        mock_doc.paragraphs = [empty_para, good_para]
        mock_doc.tables = []
        mock_doc.core_properties.title = ""

        mock_mammoth = MagicMock()
        mock_mammoth.convert_to_html.return_value = mock_result

        with patch.dict("sys.modules", {"mammoth": mock_mammoth}), \
             patch("docx.Document", return_value=mock_doc):
            resp = client.post(
                "/api/parse-document",
                data={"file": (io.BytesIO(b"x"), "x.docx")},
                headers=auth_headers,
            )

        body = resp.get_json()
        # Empty paragraph excluded from text
        assert body["text"] == "Real content"
        # First-paragraph fallback skips the empty one
        assert body["doc_title"] == "Real content"


# ──────────────────────────────────────────────────────────────────
# .docx parsing — python-docx fallback path
# ──────────────────────────────────────────────────────────────────


class TestParseDocxFallback:
    def test_fallback_when_mammoth_missing(self, client, auth_headers):
        # Simulate mammoth ImportError → falls into the python-docx-only
        # branch at line 108-161.
        from docx.oxml.text.paragraph import CT_P
        from docx.oxml.table import CT_Tbl

        # Mock paragraphs covering all four style branches:
        # Heading 1, Heading 2, generic "Heading", Normal
        mock_para_h1 = MagicMock(spec=CT_P)
        mock_para_h2 = MagicMock(spec=CT_P)
        mock_para_h_generic = MagicMock(spec=CT_P)
        mock_para_p = MagicMock(spec=CT_P)
        mock_table_elem = MagicMock(spec=CT_Tbl)

        # Fake doc structure
        mock_doc = MagicMock()
        mock_doc.element.body.iterchildren.return_value = [
            mock_para_h1, mock_para_h2, mock_para_h_generic,
            mock_para_p, mock_table_elem,
        ]

        # Mock Paragraph wrappers
        h1_paragraph = MagicMock()
        h1_paragraph.text = "Big Heading"
        h1_paragraph.style.name = "Heading 1"

        h2_paragraph = MagicMock()
        h2_paragraph.text = "Sub Heading"
        h2_paragraph.style.name = "Heading 2"

        h_generic_paragraph = MagicMock()
        h_generic_paragraph.text = "Heading Three"
        h_generic_paragraph.style.name = "Heading 3"  # generic 'Heading' branch

        normal_paragraph = MagicMock()
        normal_paragraph.text = "Regular text"
        normal_paragraph.style.name = "Normal"

        # Mock Table wrapper for the iterchildren loop
        cell = MagicMock(); cell.text = "cell"
        row = MagicMock(); row.cells = [cell]
        table = MagicMock(); table.rows = [row]

        # paragraphs/tables for the second loop (plain_text extraction).
        # Include a table with a non-empty cell so lines 151-154 fire.
        plain_para = MagicMock(); plain_para.text = "Extracted"
        plain_cell = MagicMock(); plain_cell.text = "TableCellText"
        plain_row = MagicMock(); plain_row.cells = [plain_cell]
        plain_table = MagicMock(); plain_table.rows = [plain_row]
        mock_doc.paragraphs = [plain_para]
        mock_doc.tables = [plain_table]

        # Build a side_effect for `Paragraph(child, doc)` calls
        def paragraph_factory(child, doc):
            if child is mock_para_h1:
                return h1_paragraph
            if child is mock_para_h2:
                return h2_paragraph
            if child is mock_para_h_generic:
                return h_generic_paragraph
            if child is mock_para_p:
                return normal_paragraph
            return MagicMock(text="", style=MagicMock(name=""))

        # Also need a Table factory
        def table_factory(child, doc):
            return table

        # Force ImportError specifically on `import mammoth` so the
        # python-docx-only fallback branch executes. Patch builtins.__import__
        # because sys.modules tricks let Python re-resolve the package from
        # disk (mammoth is installed in this venv).
        import builtins
        real_import = builtins.__import__

        def fake_import(name, globals=None, locals=None, fromlist=(), level=0):
            if name == "mammoth":
                raise ImportError("mocked mammoth import failure")
            return real_import(name, globals, locals, fromlist, level)

        with patch("builtins.__import__", side_effect=fake_import), \
             patch("docx.Document", return_value=mock_doc), \
             patch("docx.text.paragraph.Paragraph",
                   side_effect=paragraph_factory), \
             patch("docx.table.Table", side_effect=table_factory):
            resp = client.post(
                "/api/parse-document",
                data={"file": (io.BytesIO(b"docx data"), "fb.docx")},
                headers=auth_headers,
            )

        body = resp.get_json()
        assert body["type"] == "html"
        assert body["filename"] == "fb.docx"
        # Heading 1 → <h1>, Heading 2 → <h2>, generic Heading → <h3>
        assert "<h1>Big Heading</h1>" in body["html"]
        assert "<h2>Sub Heading</h2>" in body["html"]
        assert "<h3>Heading Three</h3>" in body["html"]
        # Normal paragraph as <p>
        assert "<p>Regular text</p>" in body["html"]
        # Plain text from the second loop includes para + table-cell row text
        assert "Extracted" in body["text"]
        assert "TableCellText" in body["text"]


# ──────────────────────────────────────────────────────────────────
# .pdf parsing
# ──────────────────────────────────────────────────────────────────


class TestParsePdf:
    def test_pdf_happy_path(self, client, auth_headers):
        # Mock fitz (PyMuPDF)
        mock_pixmap = MagicMock()
        mock_pixmap.tobytes.return_value = b"\x89PNG\r\n\x1a\nfake-png-data"

        mock_page = MagicMock()
        mock_page.get_pixmap.return_value = mock_pixmap
        mock_page.get_text.return_value = "Page 1 text content"

        mock_doc = MagicMock()
        mock_doc.__iter__.return_value = iter([mock_page])
        mock_doc.close = MagicMock()

        mock_fitz = MagicMock()
        mock_fitz.open.return_value = mock_doc
        mock_fitz.Matrix.return_value = "matrix-1.5"

        with patch.dict("sys.modules", {"fitz": mock_fitz}):
            resp = client.post(
                "/api/parse-document",
                data={"file": (io.BytesIO(b"%PDF-1.4 fake"), "test.pdf")},
                headers=auth_headers,
            )

        assert resp.status_code == 200
        body = resp.get_json()
        assert body["type"] == "html"
        assert body["filename"] == "test.pdf"
        # PNG image data URL embedded
        assert "data:image/png;base64," in body["html"]
        assert "Page 1" in body["html"]
        assert body["text"] == "Page 1 text content"
        # Doc was closed (resource cleanup)
        mock_doc.close.assert_called_once()

    def test_pdf_multiple_pages(self, client, auth_headers):
        # Multi-page PDF — each page produces its own image div
        mock_pages = []
        for i in range(3):
            mp = MagicMock()
            mp.tobytes.return_value = b"\x89PNG fake"
            mock_page = MagicMock()
            mock_page.get_pixmap.return_value = mp
            mock_page.get_text.return_value = f"Page {i + 1} text"
            mock_pages.append(mock_page)

        mock_doc = MagicMock()
        mock_doc.__iter__.return_value = iter(mock_pages)

        mock_fitz = MagicMock()
        mock_fitz.open.return_value = mock_doc
        mock_fitz.Matrix.return_value = "m"

        with patch.dict("sys.modules", {"fitz": mock_fitz}):
            resp = client.post(
                "/api/parse-document",
                data={"file": (io.BytesIO(b"%PDF"), "multi.pdf")},
                headers=auth_headers,
            )

        body = resp.get_json()
        # All 3 pages text joined
        for i in range(1, 4):
            assert f"Page {i} text" in body["text"]

    def test_pdf_pymupdf_missing_returns_helpful_error(
        self, client, auth_headers,
    ):
        # When PyMuPDF (fitz) isn't installed → ImportError → helpful msg
        import sys as _sys
        original_fitz = _sys.modules.pop("fitz", None)
        # Make the import statement inside _parse_pdf raise
        try:
            with patch.dict("sys.modules", {"fitz": None}):
                resp = client.post(
                    "/api/parse-document",
                    data={"file": (io.BytesIO(b"%PDF"), "x.pdf")},
                    headers=auth_headers,
                )
        finally:
            if original_fitz is not None:
                _sys.modules["fitz"] = original_fitz

        body = resp.get_json()
        assert "PyMuPDF" in body["error"]
        assert "pip3 install pymupdf" in body["error"]


# ──────────────────────────────────────────────────────────────────
# .txt parsing
# ──────────────────────────────────────────────────────────────────


class TestParseTxt:
    def test_txt_happy_path(self, client, auth_headers):
        text_data = b"Line 1\nLine 2\nLine 3"
        resp = client.post(
            "/api/parse-document",
            data={"file": (io.BytesIO(text_data), "notes.txt")},
            headers=auth_headers,
        )
        assert resp.status_code == 200
        body = resp.get_json()
        assert body["type"] == "html"
        assert body["filename"] == "notes.txt"
        assert "Line 1" in body["text"]
        assert "<pre" in body["html"]

    def test_txt_handles_encoding_errors_gracefully(
        self, client, auth_headers,
    ):
        # Production uses `errors='ignore'` in decode → invalid bytes stripped
        bad_bytes = b"valid\xff\xfeinvalid"
        resp = client.post(
            "/api/parse-document",
            data={"file": (io.BytesIO(bad_bytes), "weird.txt")},
            headers=auth_headers,
        )
        assert resp.status_code == 200
        body = resp.get_json()
        # Valid portion preserved, invalid bytes silently dropped
        assert "valid" in body["text"]
        assert "invalid" in body["text"]
