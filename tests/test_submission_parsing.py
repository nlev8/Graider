"""Characterization tests for parse_filename (Wave 7 Slice 6 — grader decomposition).

Pins the student-info-from-filename parsing BEFORE moving parse_filename into a new
backend/services/submission_parsing.py (the home for the upcoming file-reader cluster).
Pure (pathlib + string ops — no file I/O despite the domain). Imported via
`assignment_grader` (re-export shim).
"""
import sys
import os

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

import base64
import tempfile

from assignment_grader import parse_filename, read_image_file, read_docx_file


def test_standard_underscore_format():
    assert parse_filename("Eli_Long_Hamilton_Jefferson_Graphic_Organizer.docx") == {
        "first_name": "Eli", "last_name": "Long",
        "assignment_part": "Hamilton_Jefferson_Graphic_Organizer", "lookup_key": "eli long"}


def test_comma_last_first_format():
    assert parse_filename("Deloach, Rylee M._Washington_Stations_Handout.docx") == {
        "first_name": "Rylee", "last_name": "Deloach",
        "assignment_part": "Washington_Stations_Handout", "lookup_key": "rylee deloach"}


def test_apostrophe_stripped_in_lookup_key():
    assert parse_filename("A'kareah_West_Cornell Notes_ Political Parties.docx") == {
        "first_name": "A'kareah", "last_name": "West",
        "assignment_part": "Cornell Notes_ Political Parties", "lookup_key": "akareah west"}


def test_two_part_no_assignment():
    assert parse_filename("Jane_Doe.pdf") == {
        "first_name": "Jane", "last_name": "Doe", "assignment_part": "", "lookup_key": "jane doe"}


def test_unparseable_single_word_fallback():
    assert parse_filename("singleword.txt") == {
        "first_name": "singleword", "last_name": "", "assignment_part": "",
        "lookup_key": "singleword"}


# ── read_image_file (return-value golden; the print→logger conversion preserves these) ──


def test_read_image_returns_base64_dict():
    raw = b"\x89PNG\r\n\x1a\nFAKE-PNG-BYTES"
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as t:
        t.write(raw)
        path = t.name
    try:
        assert read_image_file(path) == {
            "type": "image",
            "data": base64.b64encode(raw).decode("utf-8"),
            "media_type": "image/png"}
    finally:
        os.unlink(path)


def test_read_image_maps_jpg_to_jpeg_mime():
    with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as t:
        t.write(b"\xff\xd8\xff")
        path = t.name
    try:
        out = read_image_file(path)
        assert out["media_type"] == "image/jpeg"
    finally:
        os.unlink(path)


def test_read_image_unsupported_ext_returns_none():
    assert read_image_file("photo.bmp") is None  # .bmp not in mime_types


def test_read_image_nonexistent_returns_none():
    assert read_image_file("/nonexistent/dir/x.png") is None  # open() fails -> None


def test_read_image_unsupported_logs_warning(caplog):
    # lock the print→logger contract (diagnostics go to the module logger, not stdout)
    with caplog.at_level("WARNING", logger="backend.services.submission_parsing"):
        assert read_image_file("photo.bmp") is None
    assert "Unsupported image type" in caplog.text


# ── read_docx_file (return-value golden; print→logger conversion preserves these) ──


def test_read_docx_interleaves_paragraphs_and_tables():
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("")  # blank — skipped
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Name"
    t.rows[0].cells[1].text = "Score"
    t.rows[1].cells[0].text = "Alice"
    t.rows[1].cells[1].text = "90"
    doc.add_paragraph("Closing paragraph.")
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        path = f.name
    try:
        # document order preserved; table rows joined with ' | '; blank para dropped
        assert read_docx_file(path) == (
            "First paragraph.\nName | Score\nAlice | 90\nClosing paragraph.")
    finally:
        os.unlink(path)


def test_read_docx_missing_file_returns_none():
    assert read_docx_file("/nonexistent/x.docx") is None


# ── read_assignment_file (dispatcher: .docx/.txt/image → typed content dict) ──


def test_read_assignment_graider_docx_returns_text_with_tables():
    from docx import Document
    from assignment_grader import read_assignment_file
    doc = Document()
    doc.add_paragraph("GRAIDER_TABLE_V1")
    t = doc.add_table(rows=2, cols=1)
    t.rows[0].cells[0].text = "[GRAIDER:VOCAB:Photosynthesis] Photosynthesis (5 pts)"
    t.rows[1].cells[0].text = "Plants make food from sunlight"
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        path = f.name
    try:
        out = read_assignment_file(path)
        assert out["type"] == "text"
        assert len(out["graider_tables"]) == 1
        assert out["graider_tables"][0]["tag_id"] == "Photosynthesis"
    finally:
        os.unlink(path)


def test_read_assignment_txt_returns_text():
    from assignment_grader import read_assignment_file
    with tempfile.NamedTemporaryFile(suffix=".txt", delete=False, mode="w") as f:
        f.write("My essay answer.")
        path = f.name
    try:
        assert read_assignment_file(path) == {"type": "text", "content": "My essay answer."}
    finally:
        os.unlink(path)


def test_read_assignment_image_returns_image():
    from assignment_grader import read_assignment_file
    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
        f.write(b"\x89PNGdata")
        path = f.name
    try:
        out = read_assignment_file(path)
        assert out["type"] == "image"
        assert out["media_type"] == "image/png"
        assert base64.b64decode(out["content"]) == b"\x89PNGdata"
    finally:
        os.unlink(path)


def test_read_assignment_unsupported_returns_none():
    from assignment_grader import read_assignment_file
    assert read_assignment_file("/tmp/nonexistent.xyz") is None
