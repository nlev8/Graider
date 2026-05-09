"""Unit tests for backend/services/visualization.py.

Strategy: drive matplotlib via the real Agg backend (already configured by
`_get_plt`) and assert that every public function returns a `data:image/png;base64,...`
data URL whose decoded bytes start with the PNG signature `\\x89PNG\\r\\n\\x1a\\n`.

Each function gets at least:
  * one populated/normal-path call
  * one `blank=True` (or empty-data) call where supported

For high-value chart functions, the `captured_figures` fixture monkeypatches
`figure_to_base64` to snapshot artist counts BEFORE the figure is closed —
so we can assert e.g. `ax.scatter()` actually added a `PathCollection` rather
than the function silently returning a blank PNG. (PR #266 Codex round-1
MAJOR fold: PNG-signature checks alone were smoke tests.)

Per project audit MAJOR #4 coverage sprint — picks the biggest single-PR gain
remaining (visualization.py was at 5%, 502 uncovered LOC).
"""

import base64
import io

import pytest

from backend.services import visualization as viz


PNG_SIGNATURE = b"\x89PNG\r\n\x1a\n"
DATA_URL_PREFIX = "data:image/png;base64,"


def _decode_data_url(data_url: str) -> bytes:
    """Strip `data:image/png;base64,` prefix and base64-decode."""
    assert isinstance(data_url, str), f"expected str, got {type(data_url)}"
    assert data_url.startswith(DATA_URL_PREFIX), f"missing data-URL prefix: {data_url[:40]!r}"
    payload = data_url.split(",", 1)[1]
    return base64.b64decode(payload)


def _assert_png_data_url(data_url: str, *, min_size: int = 100) -> bytes:
    """Assert the value is a PNG data URL and return the decoded bytes."""
    decoded = _decode_data_url(data_url)
    assert decoded[:8] == PNG_SIGNATURE, f"PNG signature missing: {decoded[:8]!r}"
    assert len(decoded) >= min_size, f"PNG suspiciously small: {len(decoded)} bytes"
    return decoded


@pytest.fixture
def captured_figures(monkeypatch):
    """Intercept `figure_to_base64` to snapshot artist counts before close.

    The wrapped function records the figure's first-axes artist collections
    (lines, collections, patches, texts, legend, title) and then forwards to
    the real implementation, which closes the figure as usual.

    Returned value is a list of per-call snapshot dicts in call order. Use
    when you need to assert the chart function actually invoked the right
    matplotlib drawing primitive — not just that *some* PNG came out.
    """
    captures: list[dict] = []
    real = viz.figure_to_base64

    def _intercept(fig):
        ax = fig.axes[0] if fig.axes else None
        snap = {
            "lines": list(ax.lines) if ax is not None else [],
            "collections": list(ax.collections) if ax is not None else [],
            "patches": list(ax.patches) if ax is not None else [],
            "texts": list(ax.texts) if ax is not None else [],
            "legend": ax.get_legend() if ax is not None else None,
            "title": ax.get_title() if ax is not None else "",
            "n_axes": len(fig.axes),
        }
        captures.append(snap)
        return real(fig)

    monkeypatch.setattr(viz, "figure_to_base64", _intercept)
    return captures


class TestLazyLoaders:
    """Cover _get_plt and _get_np singleton-style lazy loaders."""

    def test_get_plt_returns_pyplot_module(self):
        plt = viz._get_plt()
        assert plt is not None
        assert hasattr(plt, "subplots")
        # Cached on second call
        assert viz._get_plt() is plt

    def test_get_np_returns_numpy_module(self):
        np = viz._get_np()
        assert np is not None
        assert hasattr(np, "linspace")
        assert viz._get_np() is np


class TestFigureToBase64:
    def test_returns_data_url_with_png_payload(self):
        plt = viz._get_plt()
        fig, ax = plt.subplots(figsize=(2, 2))
        ax.plot([0, 1], [0, 1])
        try:
            result = viz.figure_to_base64(fig)
            _assert_png_data_url(result)
        finally:
            plt.close(fig)


class TestSaveFigure:
    def test_writes_png_file(self, tmp_path):
        plt = viz._get_plt()
        fig, ax = plt.subplots(figsize=(2, 2))
        ax.plot([0, 1], [0, 1])
        out = tmp_path / "out.png"
        try:
            viz.save_figure(fig, str(out))
        finally:
            plt.close(fig)
        assert out.exists()
        body = out.read_bytes()
        assert body[:8] == PNG_SIGNATURE


class TestRenderLatex:
    def test_renders_simple_expression(self):
        result = viz.render_latex(r"\frac{1}{2}")
        _assert_png_data_url(result)

    def test_respects_font_size(self):
        small = viz.render_latex("x^2", font_size=10)
        large = viz.render_latex("x^2", font_size=40)
        # Different font sizes should yield distinct images.
        assert small != large
        _assert_png_data_url(small)
        _assert_png_data_url(large)

    def test_handles_complex_latex_with_greek(self):
        result = viz.render_latex(r"\sum_{i=1}^{n} \alpha_i \beta_i")
        _assert_png_data_url(result)


class TestNumberLine:
    def test_default_call_renders(self):
        result = viz.create_number_line()
        _assert_png_data_url(result)

    def test_with_points_and_labels(self):
        result = viz.create_number_line(
            min_val=-5, max_val=5,
            points=[-3, 0.5, 2.75],
            labels=["A", "B", "C"],
            title="Plot the rational numbers",
        )
        _assert_png_data_url(result)

    def test_blank_template_skips_points(self):
        result = viz.create_number_line(
            min_val=0, max_val=10,
            points=[2, 5, 8],
            labels=["X", "Y", "Z"],
            blank=True,
        )
        _assert_png_data_url(result)

    def test_show_integers_false_hides_ticks(self):
        result = viz.create_number_line(min_val=0, max_val=4, show_integers=False)
        _assert_png_data_url(result)

    def test_more_points_than_labels(self):
        # zip stops at shortest — exercises labels-shorter-than-points branch
        result = viz.create_number_line(points=[1, 2, 3], labels=["only-one"])
        _assert_png_data_url(result)


class TestCoordinatePlane:
    def test_default_call_renders(self):
        result = viz.create_coordinate_plane()
        _assert_png_data_url(result)

    def test_with_points_and_labels(self):
        result = viz.create_coordinate_plane(
            x_range=(-5, 5), y_range=(-5, 5),
            points=[(3, 4), (-2, 5), (-4, -3), (2, -5)],
            labels=["A", "B", "C", "D"],
            title="Plot the points",
        )
        _assert_png_data_url(result)

    def test_with_points_no_labels_uses_coord_text(self):
        # Hits `f"({pt[0]}, {pt[1]})"` fallback path
        result = viz.create_coordinate_plane(
            points=[(1, 2), (-3, -4)],
        )
        _assert_png_data_url(result)

    def test_grid_off(self):
        result = viz.create_coordinate_plane(grid=False)
        _assert_png_data_url(result)

    def test_blank_template_skips_points(self):
        result = viz.create_coordinate_plane(
            points=[(1, 2)], labels=["A"], blank=True,
        )
        _assert_png_data_url(result)


class TestBoxPlot:
    def test_single_dataset_default_labels(self):
        result = viz.create_box_plot(
            data=[[1, 2, 3, 4, 5, 6, 7, 8, 9, 10]],
            title="Test Scores",
        )
        _assert_png_data_url(result)

    def test_two_datasets_with_labels(self):
        result = viz.create_box_plot(
            data=[
                [23, 45, 67, 32, 44, 55, 66, 77, 34, 45],
                [34, 56, 78, 43, 55, 66, 77, 88, 45, 56],
            ],
            labels=["Class A", "Class B"],
            title="Test Scores by Class",
            show_values=True,
        )
        _assert_png_data_url(result)

    def test_show_values_off(self):
        result = viz.create_box_plot(
            data=[[1, 2, 3, 4, 5]],
            show_values=False,
        )
        _assert_png_data_url(result)

    def test_blank_template(self):
        result = viz.create_box_plot(data=[], blank=True)
        _assert_png_data_url(result)


class TestBarChart:
    def test_with_data(self):
        result = viz.create_bar_chart(
            categories=["Math", "Science", "English", "History"],
            values=[85, 92, 78, 88],
            title="Average Grades by Subject",
            x_label="Subject",
            y_label="Score",
            color="seagreen",
        )
        _assert_png_data_url(result)

    def test_blank_template(self):
        result = viz.create_bar_chart(
            categories=[], values=[], blank=True,
            title="Draw a bar chart",
        )
        _assert_png_data_url(result)


class TestLineGraph:
    def test_with_data_show_points(self):
        result = viz.create_line_graph(
            x_data=[1, 2, 3, 4, 5],
            y_data=[2, 4, 6, 8, 10],
            title="y = 2x",
            x_label="x",
            y_label="y",
            show_points=True,
        )
        _assert_png_data_url(result)

    def test_with_data_no_points(self):
        result = viz.create_line_graph(
            x_data=[1, 2, 3], y_data=[1, 4, 9],
            show_points=False,
        )
        _assert_png_data_url(result)

    def test_blank_template(self):
        result = viz.create_line_graph(x_data=[], y_data=[], blank=True)
        _assert_png_data_url(result)


class TestScatterPlot:
    def test_with_data_no_trend(self):
        result = viz.create_scatter_plot(
            x_data=[1, 2, 3, 4, 5],
            y_data=[2, 5, 7, 6, 9],
            title="Scatter",
            x_label="x", y_label="y",
            show_trend=False,
        )
        _assert_png_data_url(result)

    def test_with_data_show_trend(self):
        result = viz.create_scatter_plot(
            x_data=[1, 2, 3, 4, 5],
            y_data=[2, 4, 5, 4, 5],
            show_trend=True,
        )
        _assert_png_data_url(result)

    def test_blank_template(self):
        result = viz.create_scatter_plot(x_data=[], y_data=[], blank=True)
        _assert_png_data_url(result)


class TestTriangle:
    def test_default(self):
        result = viz.create_triangle()
        _assert_png_data_url(result)

    def test_full_labels_and_dimensions(self):
        result = viz.create_triangle(
            base=8, height=6,
            show_labels=True, show_dimensions=True,
            title="Find the Area",
        )
        _assert_png_data_url(result)

    def test_no_labels_no_dimensions(self):
        result = viz.create_triangle(
            base=4, height=3,
            show_labels=False, show_dimensions=False,
        )
        _assert_png_data_url(result)

    def test_blank_template(self):
        result = viz.create_triangle(blank=True)
        _assert_png_data_url(result)


class TestRectangle:
    def test_default(self):
        result = viz.create_rectangle()
        _assert_png_data_url(result)

    def test_full_labels_and_dimensions(self):
        result = viz.create_rectangle(
            width=10, height=5,
            show_labels=True, show_dimensions=True,
            title="Area",
        )
        _assert_png_data_url(result)

    def test_no_labels_no_dimensions(self):
        result = viz.create_rectangle(
            width=2, height=3,
            show_labels=False, show_dimensions=False,
        )
        _assert_png_data_url(result)

    def test_blank_template(self):
        result = viz.create_rectangle(blank=True)
        _assert_png_data_url(result)


class TestFunctionGraph:
    def test_single_linear(self):
        result = viz.create_function_graph(
            expressions=["2*x + 1"],
            x_range=(-5, 5),
            title="y = 2x + 1",
        )
        _assert_png_data_url(result)

    def test_multiple_expressions_with_labels(self):
        result = viz.create_function_graph(
            expressions=["x**2", "x + 1", "y = 3*x - 2"],
            x_range=(-3, 3),
            labels=["parabola", "line", "another"],
        )
        _assert_png_data_url(result)

    def test_caret_notation_handled(self):
        # Hits replace('^', '**')
        result = viz.create_function_graph(expressions=["x^2 + 1"])
        _assert_png_data_url(result)

    def test_invalid_expression_continues(self):
        # `not_a_function(x)` raises during sympify or lambdify — continue branch
        result = viz.create_function_graph(
            expressions=["definitely-invalid-???", "x"],
            x_range=(-2, 2),
        )
        _assert_png_data_url(result)

    def test_all_invalid_expressions_keeps_all_y_empty(self):
        # All expressions fail → `all_y` stays empty → both `if y_range is None
        # and all_y` and the `elif y_range:` branches fall through.
        # (PR #266 Codex round-1 MINOR fold: prior test paired invalid with
        # a valid `"x"`, never exercising the empty-`all_y` path.)
        result = viz.create_function_graph(
            expressions=["definitely-invalid-???", "also @ broken"],
            x_range=(-2, 2),
        )
        _assert_png_data_url(result)

    def test_constant_expression(self):
        # Hits not hasattr(y_vals, '__len__') — full_like branch
        result = viz.create_function_graph(expressions=["3"])
        _assert_png_data_url(result)

    def test_explicit_y_range(self):
        result = viz.create_function_graph(
            expressions=["x"], y_range=(-5, 5),
        )
        _assert_png_data_url(result)

    def test_blank_with_y_range(self):
        result = viz.create_function_graph(
            expressions=[], blank=True, y_range=(-2, 2),
        )
        _assert_png_data_url(result)

    def test_blank_no_y_range(self):
        result = viz.create_function_graph(expressions=[], blank=True)
        _assert_png_data_url(result)

    def test_grid_off(self):
        result = viz.create_function_graph(
            expressions=["x"], show_grid=False,
        )
        _assert_png_data_url(result)


class TestCircle:
    def test_default(self):
        result = viz.create_circle()
        _assert_png_data_url(result)

    def test_show_radius_diameter_area(self):
        result = viz.create_circle(
            radius=4, center=(1, 1),
            show_radius=True, show_diameter=True, show_area=True,
            title="Circle",
        )
        _assert_png_data_url(result)

    def test_no_decorations(self):
        result = viz.create_circle(
            show_radius=False, show_diameter=False, show_area=False,
        )
        _assert_png_data_url(result)

    def test_blank_unlabeled(self):
        result = viz.create_circle(blank=True, show_radius=True, show_area=True)
        _assert_png_data_url(result)


class TestPolygon:
    @pytest.mark.parametrize("sides", [3, 4, 5, 6, 8])
    def test_various_sides_render(self, sides):
        result = viz.create_polygon(sides=sides, side_length=3)
        _assert_png_data_url(result)

    def test_with_labels_and_dimensions(self):
        result = viz.create_polygon(
            sides=5, side_length=2,
            show_labels=True, show_dimensions=True,
            title="Pentagon",
        )
        _assert_png_data_url(result)

    def test_no_labels_no_dimensions(self):
        result = viz.create_polygon(
            sides=6, show_labels=False, show_dimensions=False,
        )
        _assert_png_data_url(result)

    def test_blank(self):
        result = viz.create_polygon(sides=4, blank=True)
        _assert_png_data_url(result)


class TestHistogram:
    def test_with_data(self):
        result = viz.create_histogram(
            data=[1, 1, 2, 3, 3, 3, 4, 5, 5, 5, 5, 6, 7, 8, 9],
            bins=5,
            title="Frequency",
            x_label="Value",
            show_values=True,
        )
        _assert_png_data_url(result)

    def test_show_values_off(self):
        result = viz.create_histogram(
            data=[1, 2, 3, 4, 5], bins=3, show_values=False,
        )
        _assert_png_data_url(result)

    def test_blank_template(self):
        result = viz.create_histogram(data=[], blank=True)
        _assert_png_data_url(result)


class TestPieChart:
    def test_with_data(self):
        result = viz.create_pie_chart(
            categories=["A", "B", "C"],
            values=[30, 40, 30],
            title="Distribution",
            show_percentages=True,
        )
        _assert_png_data_url(result)

    def test_with_explode(self):
        result = viz.create_pie_chart(
            categories=["X", "Y", "Z"],
            values=[1, 2, 3],
            explode=[0, 0.1, 0],
            show_percentages=False,
        )
        _assert_png_data_url(result)

    def test_blank_template(self):
        result = viz.create_pie_chart(categories=[], values=[], blank=True)
        _assert_png_data_url(result)


class TestDotPlot:
    def test_with_categories_and_dots(self):
        result = viz.create_dot_plot(
            categories=["A", "B", "C"],
            dots={"A": 2, "B": 3, "C": 1},
            title="Dot Plot",
        )
        _assert_png_data_url(result)

    def test_auto_generate_categories(self):
        # Hits `items = ... range` branch when categories is None
        result = viz.create_dot_plot(
            min_val=0, max_val=4, step=1,
            dots={"0": 1, "2": 2, "4": 3},
        )
        _assert_png_data_url(result)

    def test_blank_template(self):
        result = viz.create_dot_plot(
            categories=["A", "B"], dots={"A": 1}, blank=True,
        )
        _assert_png_data_url(result)

    def test_no_dots_default_max(self):
        # Hits `default=3` for max_count
        result = viz.create_dot_plot(categories=["X", "Y"])
        _assert_png_data_url(result)


class TestStemAndLeaf:
    def test_with_data(self):
        result = viz.create_stem_and_leaf(
            data=[12, 15, 22, 25, 28, 31, 33, 47],
            title="Test Scores",
        )
        _assert_png_data_url(result)

    def test_blank_with_data_hides_leaves(self):
        result = viz.create_stem_and_leaf(
            data=[10, 20, 30, 40], blank=True,
        )
        _assert_png_data_url(result)

    def test_no_data_placeholder(self):
        result = viz.create_stem_and_leaf(data=None)
        _assert_png_data_url(result)


class TestVennDiagram:
    def test_two_sets_with_regions(self):
        result = viz.create_venn_diagram(
            sets=2,
            labels=["Cats", "Dogs"],
            regions={"only_a": "5", "a_and_b": "2", "only_b": "8"},
            title="Pets",
        )
        _assert_png_data_url(result)

    def test_three_sets_with_regions(self):
        result = viz.create_venn_diagram(
            sets=3,
            regions={
                "only_a": "1", "only_b": "2", "only_c": "3",
                "a_and_b": "4", "a_and_c": "5", "b_and_c": "6",
                "all": "7",
            },
        )
        _assert_png_data_url(result)

    def test_default_labels(self):
        # No labels passed — hits `Set A`, `Set B` default
        result = viz.create_venn_diagram(sets=2)
        _assert_png_data_url(result)

    def test_blank_no_regions(self):
        result = viz.create_venn_diagram(sets=2, blank=True)
        _assert_png_data_url(result)


class TestProtractor:
    def test_default(self):
        result = viz.create_protractor()
        _assert_png_data_url(result)

    @pytest.mark.parametrize("angle", [0, 30, 45, 60, 90, 135, 180])
    def test_various_angles(self, angle):
        result = viz.create_protractor(given_angle=angle)
        _assert_png_data_url(result)

    def test_show_answer_false(self):
        result = viz.create_protractor(given_angle=72, show_answer=False, title="Measure")
        _assert_png_data_url(result)


# =============================================================================
# Semantic figure-artist inspection (PR #266 Codex round-1 MAJOR fold).
#
# PNG-signature checks alone are smoke tests: they pass even if a chart
# function silently stops calling its drawing primitive and returns a blank
# canvas. These tests assert the *expected matplotlib artists* exist on the
# axes before the figure is closed, so a regression that drops e.g.
# `ax.scatter()` would actually fail.
# =============================================================================

class TestChartActuallyDraws:
    def test_scatter_plot_calls_ax_scatter(self, captured_figures):
        from matplotlib.collections import PathCollection

        viz.create_scatter_plot(x_data=[1, 2, 3], y_data=[4, 5, 6])

        snap = captured_figures[0]
        # ax.scatter() adds a PathCollection to ax.collections
        assert any(isinstance(c, PathCollection) for c in snap["collections"]), (
            f"Expected PathCollection from ax.scatter(); got: {snap['collections']}"
        )
        # No trend line → no legend, no extra ax.plot lines
        assert snap["legend"] is None
        assert len(snap["lines"]) == 0

    def test_scatter_plot_with_trend_adds_line_and_legend(self, captured_figures):
        from matplotlib.collections import PathCollection

        viz.create_scatter_plot(
            x_data=[1, 2, 3, 4, 5],
            y_data=[2, 4, 5, 4, 5],
            show_trend=True,
        )
        snap = captured_figures[0]
        assert any(isinstance(c, PathCollection) for c in snap["collections"])
        # Trend line is drawn via ax.plot — adds a Line2D to ax.lines
        assert len(snap["lines"]) >= 1, (
            f"Expected ≥1 trend line; got {len(snap['lines'])}"
        )
        assert snap["legend"] is not None, "Trend line should produce a legend"

    def test_bar_chart_creates_rectangle_patches(self, captured_figures):
        from matplotlib.patches import Rectangle

        viz.create_bar_chart(categories=["A", "B", "C"], values=[1, 2, 3])

        snap = captured_figures[0]
        rects = [p for p in snap["patches"] if isinstance(p, Rectangle)]
        # ax.bar() produces one Rectangle per bar
        assert len(rects) >= 3, f"Expected ≥3 bar rectangles; got {len(rects)}"

    def test_line_graph_plots_line_and_points(self, captured_figures):
        viz.create_line_graph(
            x_data=[1, 2, 3], y_data=[1, 4, 9], show_points=True,
        )
        snap = captured_figures[0]
        # show_points=True → two ax.plot calls → two Line2D entries
        assert len(snap["lines"]) >= 2, (
            f"Expected ≥2 lines (line + points); got {len(snap['lines'])}"
        )

    def test_line_graph_no_points_only_one_line(self, captured_figures):
        viz.create_line_graph(
            x_data=[1, 2, 3], y_data=[1, 4, 9], show_points=False,
        )
        snap = captured_figures[0]
        assert len(snap["lines"]) == 1, (
            f"Expected exactly 1 line (no points); got {len(snap['lines'])}"
        )

    def test_box_plot_creates_box_patches(self, captured_figures):
        viz.create_box_plot(data=[[1, 2, 3, 4, 5], [6, 7, 8, 9, 10]])

        snap = captured_figures[0]
        # boxplot's `patch_artist=True` produces filled box patches; one per data set
        assert len(snap["patches"]) >= 2, (
            f"Expected ≥2 box patches; got {len(snap['patches'])}"
        )

    def test_histogram_creates_bar_patches(self, captured_figures):
        from matplotlib.patches import Rectangle

        viz.create_histogram(data=[1, 1, 2, 3, 3, 3, 4, 5], bins=3)

        snap = captured_figures[0]
        rects = [p for p in snap["patches"] if isinstance(p, Rectangle)]
        # bins=3 → 3 histogram bars
        assert len(rects) >= 3, (
            f"Expected ≥3 histogram bar rectangles; got {len(rects)}"
        )

    def test_pie_chart_creates_wedges(self, captured_figures):
        from matplotlib.patches import Wedge

        viz.create_pie_chart(categories=["A", "B", "C"], values=[1, 2, 3])

        snap = captured_figures[0]
        wedges = [p for p in snap["patches"] if isinstance(p, Wedge)]
        assert len(wedges) >= 3, (
            f"Expected ≥3 pie wedges; got {len(wedges)}"
        )

    def test_triangle_creates_polygon(self, captured_figures):
        from matplotlib.patches import Polygon

        viz.create_triangle(base=4, height=3)
        snap = captured_figures[0]
        polys = [p for p in snap["patches"] if isinstance(p, Polygon)]
        assert len(polys) >= 1, "Expected ≥1 Polygon patch for triangle body"

    def test_rectangle_creates_rectangle_patch(self, captured_figures):
        from matplotlib.patches import Rectangle

        viz.create_rectangle(width=4, height=3)
        snap = captured_figures[0]
        rects = [p for p in snap["patches"] if isinstance(p, Rectangle)]
        assert len(rects) >= 1, "Expected ≥1 Rectangle patch"

    def test_circle_creates_circle_patch(self, captured_figures):
        from matplotlib.patches import Circle

        viz.create_circle(radius=4)
        snap = captured_figures[0]
        circles = [p for p in snap["patches"] if isinstance(p, Circle)]
        assert len(circles) >= 1, "Expected ≥1 Circle patch"

    def test_polygon_creates_polygon_patch(self, captured_figures):
        from matplotlib.patches import Polygon

        viz.create_polygon(sides=6, side_length=2)
        snap = captured_figures[0]
        polys = [p for p in snap["patches"] if isinstance(p, Polygon)]
        assert len(polys) >= 1, "Expected ≥1 Polygon patch for hexagon body"

    def test_function_graph_plots_each_expression(self, captured_figures):
        viz.create_function_graph(
            expressions=["x", "x**2"], x_range=(-2, 2),
        )
        snap = captured_figures[0]
        # The function adds 2 origin lines (axhline + axvline) + 1 line per
        # successfully plotted expression = 4 lines total. The legend, however,
        # contains *only* the expression labels — so it's a more direct claim
        # that both expressions actually plotted.
        assert snap["legend"] is not None, "Multi-expression call should produce a legend"
        legend_texts = [t.get_text() for t in snap["legend"].get_texts()]
        assert len(legend_texts) == 2, (
            f"Expected exactly 2 legend entries (one per expression); got {legend_texts}"
        )

    def test_number_line_plots_each_point(self, captured_figures):
        # PR #266 Codex round-2 MAJOR fold: prior assertion `len(lines) >= 3`
        # was satisfied by the 11 integer tick lines alone (range(-5, 6))
        # plus the axhline — even if the point-marker draw at viz.py:152 were
        # removed. Distinguish by marker style: tick lines use `linestyle='-'`
        # with no marker; point markers use `marker='o'` with `linestyle='None'`.
        viz.create_number_line(
            min_val=-5, max_val=5,
            points=[-3, 0, 2.5],
            labels=["A", "B", "C"],
        )
        snap = captured_figures[0]
        point_markers = [
            ln for ln in snap["lines"]
            if ln.get_marker() == "o" and ln.get_linestyle() == "None"
        ]
        assert len(point_markers) == 3, (
            f"Expected exactly 3 'o' markers (one per point); got {len(point_markers)}"
        )

    def test_number_line_blank_omits_point_markers(self, captured_figures):
        # Regression guard for the `if points and not blank` branch
        viz.create_number_line(
            min_val=-3, max_val=3,
            points=[1, 2], labels=["A", "B"], blank=True,
        )
        snap = captured_figures[0]
        point_markers = [
            ln for ln in snap["lines"]
            if ln.get_marker() == "o" and ln.get_linestyle() == "None"
        ]
        assert len(point_markers) == 0, (
            "blank=True should suppress all point markers"
        )

    def test_coordinate_plane_plots_each_point(self, captured_figures):
        # Same marker-style filter as number-line so that removing
        # `ax.plot(pt[0], pt[1], 'o', ...)` would actually fail this test.
        viz.create_coordinate_plane(
            points=[(1, 2), (-3, -4), (4, -1)], labels=["A", "B", "C"],
        )
        snap = captured_figures[0]
        point_markers = [
            ln for ln in snap["lines"]
            if ln.get_marker() == "o" and ln.get_linestyle() == "None"
        ]
        assert len(point_markers) == 3, (
            f"Expected exactly 3 'o' markers (one per point); got {len(point_markers)}"
        )

    def test_coordinate_plane_blank_omits_point_markers(self, captured_figures):
        viz.create_coordinate_plane(
            points=[(1, 2), (3, 4)], blank=True,
        )
        snap = captured_figures[0]
        point_markers = [
            ln for ln in snap["lines"]
            if ln.get_marker() == "o" and ln.get_linestyle() == "None"
        ]
        assert len(point_markers) == 0, (
            "blank=True should suppress all point markers"
        )

    def test_dot_plot_draws_each_dot(self, captured_figures):
        # PR #266 Codex round-2 MAJOR fold: TestDotPlot was smoke-only;
        # this asserts the actual data dots at viz.py:945 are drawn —
        # one 'o' marker per count value across all categories.
        viz.create_dot_plot(
            categories=["A", "B", "C"],
            dots={"A": 2, "B": 3, "C": 1},
        )
        snap = captured_figures[0]
        dot_markers = [
            ln for ln in snap["lines"]
            if ln.get_marker() == "o" and ln.get_linestyle() == "None"
        ]
        # Total dots = 2 + 3 + 1 = 6
        assert len(dot_markers) == 6, (
            f"Expected exactly 6 dot markers (sum of dots dict); "
            f"got {len(dot_markers)}"
        )

    def test_dot_plot_blank_omits_dots(self, captured_figures):
        viz.create_dot_plot(
            categories=["A", "B"], dots={"A": 1, "B": 2}, blank=True,
        )
        snap = captured_figures[0]
        dot_markers = [
            ln for ln in snap["lines"]
            if ln.get_marker() == "o" and ln.get_linestyle() == "None"
        ]
        assert len(dot_markers) == 0, (
            "blank=True should omit dot markers"
        )

    def test_dot_plot_zero_count_categories_draw_nothing(self, captured_figures):
        # If a category has count 0, `range(count)` is empty → no dot drawn
        viz.create_dot_plot(
            categories=["X", "Y", "Z"],
            dots={"X": 1, "Z": 2},  # Y missing → defaults to 0
        )
        snap = captured_figures[0]
        dot_markers = [
            ln for ln in snap["lines"]
            if ln.get_marker() == "o" and ln.get_linestyle() == "None"
        ]
        # 1 + 0 + 2 = 3
        assert len(dot_markers) == 3, (
            f"Expected exactly 3 dot markers (X=1, Y=0, Z=2); "
            f"got {len(dot_markers)}"
        )


class TestAddImageToDocx:
    def test_strips_data_url_prefix_and_adds_picture(self):
        # Render a real PNG to base64, then exercise both prefix-stripping and
        # `doc.add_picture` invocation via a stub.
        from docx.shared import Inches

        png = viz.create_rectangle(width=2, height=2, title="t")

        recorded: dict = {}

        class _StubDoc:
            def add_picture(self, stream, width=None):
                recorded["bytes"] = stream.read()
                recorded["width"] = width

        doc = _StubDoc()
        viz.add_image_to_docx(doc, png, width_inches=3)

        assert recorded["bytes"][:8] == PNG_SIGNATURE
        # PR #266 Codex round-1 MINOR fold: confirm width_inches is actually
        # threaded into `Inches(width_inches)` — equal to the EMU value
        # produced by docx.shared.Inches(3) (== 3 * 914400).
        assert recorded["width"] == Inches(3)

    def test_default_width_inches(self):
        from docx.shared import Inches

        png = viz.create_rectangle(width=2, height=2)

        recorded: dict = {}

        class _StubDoc:
            def add_picture(self, stream, width=None):
                recorded["width"] = width

        viz.add_image_to_docx(_StubDoc(), png)  # uses default width_inches=5
        assert recorded["width"] == Inches(5)

    def test_accepts_raw_base64_without_prefix(self):
        # PNG signature in raw base64
        from docx.shared import Inches

        raw_png = base64.b64encode(PNG_SIGNATURE + b"\x00" * 32).decode("utf-8")

        captured: dict = {}

        class _StubDoc:
            def add_picture(self, stream, width=None):
                captured["bytes"] = stream.read()
                captured["width"] = width

        viz.add_image_to_docx(_StubDoc(), raw_png, width_inches=1)
        assert captured["bytes"][:8] == PNG_SIGNATURE
        assert captured["width"] == Inches(1)
