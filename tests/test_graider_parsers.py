"""Characterization tests for the Graider marker/table parsers
(Wave 7 — assignment_grader.py decomposition, complex file-reader parsers).

Pins the EXACT extraction output of extract_from_graider_text (+ extract_from_tables)
on a REAL Graider-formatted submission BEFORE moving them into
backend/services/submission_parsing.py. These are print-heavy → the prints become
_logger calls on extraction (the validated print→logger pattern); RETURN VALUES are
unchanged (what these tests pin). Imported via `assignment_grader` (re-export shim).
"""
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))

from assignment_grader import extract_from_graider_text, extract_from_tables

FIXTURE = os.path.join(os.path.dirname(__file__), "fixtures", "grading",
                       "submission_social_studies.txt")


def _content():
    with open(FIXTURE) as f:
        return f.read()


def test_graider_text_empty_or_no_markers_returns_none():
    assert extract_from_graider_text("") is None
    assert extract_from_graider_text("just plain text, no GRAIDER markers") is None


def test_graider_text_extracts_all_responses_from_real_fixture():
    out = extract_from_graider_text(_content())
    assert out["total_questions"] == 10
    assert out["answered_questions"] == 10
    assert len(out["extracted_responses"]) == 10
    assert out["extraction_summary"] == "Graider text fallback: Found 10 responses out of 10 sections."
    assert set(out.keys()) == {
        "extracted_responses", "blank_questions", "total_questions",
        "answered_questions", "extraction_summary", "excluded_sections", "missing_sections"}


def test_graider_text_first_response_is_vocab_term():
    out = extract_from_graider_text(_content())
    assert out["extracted_responses"][0] == {
        "question": "Louisiana Purchase",
        "answer": "When the US bought a huge piece of land from France in 1803 for $15 million",
        "type": "vocab_term", "section": "VOCAB", "tag_id": "Louisiana Purchase"}


def test_graider_text_last_response_is_summary():
    out = extract_from_graider_text(_content())
    last = out["extracted_responses"][-1]
    assert last["type"] == "summary"
    assert last["section"] == "SUMMARY"
    assert last["tag_id"] == "main"
    assert last["question"] == "Summary"
    assert last["answer"].startswith("The Louisiana Purchase of 1803 was one of the most important")


def test_graider_text_numbered_question_parsed():
    out = extract_from_graider_text(_content())
    numbered = [r for r in out["extracted_responses"] if r["type"] == "numbered_question"]
    assert len(numbered) == 4  # the fixture has 4 numbered questions
    assert numbered[0]["section"] == "QUESTION"
    assert numbered[0]["tag_id"] == "1"


# ── extract_from_tables (structured Graider table data → same dict shape) ──

TABLE_DATA = [
    {"tag_type": "VOCAB", "tag_id": "Photosynthesis", "header_text": "Photosynthesis (5 pts)",
     "response": "The process plants use to make food from sunlight"},
    {"tag_type": "QUESTION", "tag_id": "1",
     "header_text": "1) What is the powerhouse of the cell? (10 pts)", "response": "The mitochondria"},
    {"tag_type": "SUMMARY", "tag_id": "main", "header_text": "Summary (20 pts)",
     "response": "Cells are the basic unit of life and all living things are made of them."},
    {"tag_type": "VOCAB", "tag_id": "Osmosis", "header_text": "Osmosis (5 pts)", "response": "___"},
]


def test_tables_extracts_responses_and_blanks():
    out = extract_from_tables(TABLE_DATA)
    assert out["total_questions"] == 4
    assert out["answered_questions"] == 3
    assert out["blank_questions"] == ["Osmosis"]
    assert out["extraction_summary"] == (
        "Table extraction: Found 3 responses out of 4 sections. 1 left blank.")
    assert [r["type"] for r in out["extracted_responses"]] == [
        "vocab_term", "numbered_question", "summary"]


def test_tables_vocab_and_summary_question_labels():
    out = extract_from_tables(TABLE_DATA)
    r = out["extracted_responses"]
    assert r[0] == {"question": "Photosynthesis",
                    "answer": "The process plants use to make food from sunlight",
                    "type": "vocab_term", "section": "VOCAB", "tag_id": "Photosynthesis"}
    assert r[2]["question"] == "Summary"  # SUMMARY label is fixed
    assert r[1]["question"] == "1) What is the powerhouse of the cell? (10 pts)"  # QUESTION uses header


def test_tables_exclude_markers_skips_section():
    out = extract_from_tables(TABLE_DATA, exclude_markers=["Summary"])
    assert out["excluded_sections"] == ["Summary (20 pts)"]


def test_tables_empty_input_quirk():
    # empty table_data reports total_questions: 1 (preserved quirk)
    assert extract_from_tables([]) == {
        "extracted_responses": [], "blank_questions": [], "total_questions": 1,
        "answered_questions": 0,
        "extraction_summary": "Table extraction: Found 0 responses out of 0 sections.",
        "excluded_sections": [], "missing_sections": []}


# ── read_docx_file_structured (detects [GRAIDER:..] 2-row tables in a .docx) ──


def _make_graider_docx():
    import tempfile
    from docx import Document
    doc = Document()
    doc.add_paragraph("GRAIDER_TABLE_V1")
    t1 = doc.add_table(rows=2, cols=1)
    t1.rows[0].cells[0].text = "[GRAIDER:VOCAB:Photosynthesis] Photosynthesis (5 pts)"
    t1.rows[1].cells[0].text = "The process plants use to make food from sunlight"
    t2 = doc.add_table(rows=2, cols=1)
    t2.rows[0].cells[0].text = "[GRAIDER:QUESTION:1] 1) What is the powerhouse of the cell? (10 pts)"
    t2.rows[1].cells[0].text = "The mitochondria"
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(f.name)
    return f.name


def test_structured_detects_graider_tables():
    from assignment_grader import read_docx_file_structured
    import os
    path = _make_graider_docx()
    try:
        out = read_docx_file_structured(path)
        assert out["is_graider_table"] is True
        assert out["tables"] == [
            {"tag_type": "VOCAB", "tag_id": "Photosynthesis", "header_text": "Photosynthesis (5 pts)",
             "response": "The process plants use to make food from sunlight"},
            {"tag_type": "QUESTION", "tag_id": "1",
             "header_text": "1) What is the powerhouse of the cell? (10 pts)",
             "response": "The mitochondria"}]
        assert out["plain_text"] == (
            "GRAIDER_TABLE_V1\nPhotosynthesis (5 pts)\n"
            "The process plants use to make food from sunlight\n"
            "1) What is the powerhouse of the cell? (10 pts)\nThe mitochondria")
    finally:
        os.unlink(path)


def test_structured_non_graider_docx_returns_false():
    from assignment_grader import read_docx_file_structured
    import os, tempfile
    from docx import Document
    d = Document()
    d.add_paragraph("Just a normal essay with no markers.")
    f = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    d.save(f.name)
    try:
        out = read_docx_file_structured(f.name)
        assert out["is_graider_table"] is False
        assert out["tables"] == []
    finally:
        os.unlink(f.name)
