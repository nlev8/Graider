"""Direct-import tests for the render helpers moved into planner_export (Wave 6 Slice 1).

Fixtures use the REAL data shapes the helpers read (study-guide sections carry a
`content` LIST of point strings + `terms` [{term, definition}]; flashcards carry
`cards` [{term, definition}]). The docx tests read the file back and assert the
content actually rendered — a real regression guard, not just "file non-empty".
"""
import os
import tempfile

from docx import Document

STUDY_GUIDE = {
    "title": "Photosynthesis",
    "sections": [{
        "heading": "Overview",
        "content": ["Plants convert light to energy.", "Chlorophyll absorbs light."],
        "terms": [{"term": "Chlorophyll", "definition": "Green pigment in plants."}],
    }],
}
FLASHCARDS = {"title": "Bio Vocab", "cards": [{"term": "cat", "definition": "a feline"}]}


def _docx_text(path):
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(
        c.text for tbl in doc.tables for row in tbl.rows for c in row.cells
    )


def test_export_study_guide_docx_renders_content():
    from backend.services.planner_export import _export_study_guide_docx
    path = os.path.join(tempfile.mkdtemp(), "sg.docx")
    _export_study_guide_docx(STUDY_GUIDE, path)
    text = _docx_text(path)
    assert "Photosynthesis" in text
    assert "Chlorophyll absorbs light." in text   # a content list item rendered
    assert "Green pigment in plants." in text       # a term definition rendered


def test_export_study_guide_pdf_writes_nonempty_file():
    from backend.services.planner_export import _export_study_guide_pdf
    path = os.path.join(tempfile.mkdtemp(), "sg.pdf")
    _export_study_guide_pdf(STUDY_GUIDE, path)
    assert os.path.exists(path) and os.path.getsize(path) > 0


def test_export_flashcards_docx_renders_term_and_definition():
    from backend.services.planner_export import _export_flashcards_docx
    path = os.path.join(tempfile.mkdtemp(), "fc.docx")
    _export_flashcards_docx(FLASHCARDS, path)
    text = _docx_text(path)
    assert "cat" in text and "a feline" in text


def test_export_flashcards_pdf_writes_nonempty_file():
    from backend.services.planner_export import _export_flashcards_pdf
    path = os.path.join(tempfile.mkdtemp(), "fc.pdf")
    _export_flashcards_pdf(FLASHCARDS, path)
    assert os.path.exists(path) and os.path.getsize(path) > 0


def test_export_flashcards_pdf_empty_cards_does_not_crash():
    # The helper has an explicit `if not cards:` early-return branch.
    from backend.services.planner_export import _export_flashcards_pdf
    path = os.path.join(tempfile.mkdtemp(), "empty.pdf")
    _export_flashcards_pdf({"title": "Empty", "cards": []}, path)
    assert os.path.exists(path) and os.path.getsize(path) > 0
