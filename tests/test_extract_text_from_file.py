"""Characterization tests for /api/extract-text (Wave 6 Slice 10).

Written BEFORE extracting the docx/pdf/txt/image dispatch into
planner_standards.extract_text_from_upload. These pin the OBSERVABLE contract —
the route's 400 branches (no file, unsupported ext, in-ALLOWED-but-unhandled
.doc/.rtf else, image-without-key) and a few happy paths — so the extraction is
constrained regardless of how the route ends up resolving the OpenAI key
(eager vs lazy) internally. The internal dispatch is moved verbatim and
separately byte-verified via AST compare.

Multipart uploads use a teacher-id-only header (NOT application/json) so Flask
sets the multipart boundary. get_api_key is mocked in non-key-path tests so the
suite is robust to eager key resolution.
"""
import io
import os
import sys
from unittest.mock import patch, MagicMock

import pytest

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'backend'))


@pytest.fixture
def app():
    os.environ['FLASK_ENV'] = 'development'
    os.environ['DEV_USER_ID'] = 'test-teacher-001'
    from backend.app import app as flask_app
    flask_app.config['TESTING'] = True
    return flask_app


@pytest.fixture
def client(app):
    return app.test_client()


@pytest.fixture
def mp_headers():
    # multipart: only the teacher id — let Flask set the boundary Content-Type
    return {'X-Test-Teacher-Id': 'test-teacher-001'}


def _completion(text):
    c = MagicMock()
    c.content_parts = [MagicMock(text=text)]
    c.usage = None
    return c


def test_extract_no_file_returns_400(client, mp_headers):
    resp = client.post('/api/extract-text', data={}, content_type='multipart/form-data',
                      headers=mp_headers)
    assert resp.status_code == 400
    assert resp.get_json()["error"] == "No file uploaded"


def test_extract_unsupported_ext_returns_400(client, mp_headers):
    data = {'file': (io.BytesIO(b'junk'), 'mystery.xyz')}
    resp = client.post('/api/extract-text', data=data, content_type='multipart/form-data',
                      headers=mp_headers)
    assert resp.status_code == 400
    assert resp.get_json()["error"] == "Unsupported file type: .xyz"


def test_extract_doc_rtf_in_allowed_but_unhandled_returns_else_400(client, mp_headers):
    # .doc is in ALLOWED_DOC_EXTENSIONS but has no parse branch -> the else 400
    data = {'file': (io.BytesIO(b'\xd0\xcf legacy doc bytes'), 'legacy.doc')}
    resp = client.post('/api/extract-text', data=data, content_type='multipart/form-data',
                      headers=mp_headers)
    assert resp.status_code == 400
    assert resp.get_json()["error"] == \
        "Unsupported file type. Use .docx, .pdf, .txt, .png, .jpg, or .jpeg"


def test_extract_txt_happy_path(client, mp_headers):
    data = {'file': (io.BytesIO('Line one\nLíne two'.encode('utf-8')), 'notes.txt')}
    with patch('backend.api_keys.get_api_key', return_value='x'):
        resp = client.post('/api/extract-text', data=data, content_type='multipart/form-data',
                          headers=mp_headers)
    assert resp.status_code == 200
    assert resp.get_json()["text"] == 'Line one\nLíne two'


def test_extract_docx_happy_path(client, mp_headers):
    from docx import Document
    doc = Document()
    doc.add_paragraph("First paragraph.")
    doc.add_paragraph("")  # blank — should be skipped
    doc.add_paragraph("Second paragraph.")
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    data = {'file': (buf, 'essay.docx')}
    with patch('backend.api_keys.get_api_key', return_value='x'):
        resp = client.post('/api/extract-text', data=data, content_type='multipart/form-data',
                          headers=mp_headers)
    assert resp.status_code == 200
    assert resp.get_json()["text"] == "First paragraph.\nSecond paragraph."


def test_extract_image_without_key_returns_400(client, mp_headers):
    data = {'file': (io.BytesIO(b'\x89PNG fake'), 'scan.png')}
    with patch('backend.api_keys.get_api_key', return_value=None):
        resp = client.post('/api/extract-text', data=data, content_type='multipart/form-data',
                          headers=mp_headers)
    assert resp.status_code == 400
    assert resp.get_json()["error"] == "OpenAI API key required for image text extraction"


def test_extract_image_happy_path_uses_vision(client, mp_headers):
    fake_adapter = MagicMock()
    fake_adapter.chat.return_value = _completion("  Extracted words from the scan.  ")
    data = {'file': (io.BytesIO(b'\x89PNG fake'), 'scan.png')}
    with patch('backend.api_keys.get_api_key', return_value='fake-key'), \
         patch('backend.services.llm_adapter.OpenAIAdapter', return_value=fake_adapter):
        resp = client.post('/api/extract-text', data=data, content_type='multipart/form-data',
                          headers=mp_headers)
    assert resp.status_code == 200
    assert resp.get_json()["text"] == "Extracted words from the scan."  # stripped


# ── Direct service-level tests (pin extract_text_from_upload contract) ──


def test_service_txt_returns_decoded_text():
    from backend.services.planner_standards import extract_text_from_upload
    out = extract_text_from_upload(file_data='héllo'.encode('utf-8'), filename='a.txt', api_key=None)
    assert out == 'héllo'  # raw string, not stripped/wrapped


def test_service_docx_returns_joined_paragraphs():
    from docx import Document
    from backend.services.planner_standards import extract_text_from_upload
    doc = Document()
    doc.add_paragraph("Alpha")
    doc.add_paragraph("   ")  # whitespace-only — skipped
    doc.add_paragraph("Beta")
    buf = io.BytesIO()
    doc.save(buf)
    out = extract_text_from_upload(file_data=buf.getvalue(), filename='x.docx', api_key=None)
    assert out == "Alpha\nBeta"


def test_service_unhandled_type_raises_text_extraction_error():
    from backend.services.planner_standards import extract_text_from_upload, TextExtractionError
    with pytest.raises(TextExtractionError) as exc:
        extract_text_from_upload(file_data=b'legacy', filename='old.doc', api_key=None)
    assert str(exc.value) == "Unsupported file type. Use .docx, .pdf, .txt, .png, .jpg, or .jpeg"


def test_image_extensions_single_source_of_truth():
    # Route and service must reference the SAME tuple object so the two image
    # branches can never diverge (a divergence would let an image reach the
    # service with api_key=None and fail in the 500 catch-all, not the 400).
    import backend.routes.planner_routes as pr
    from backend.services.planner_standards import IMAGE_EXTENSIONS
    assert pr.IMAGE_EXTENSIONS is IMAGE_EXTENSIONS
    assert IMAGE_EXTENSIONS == ('.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp')


def test_service_image_uses_vision_with_key():
    from backend.services.planner_standards import extract_text_from_upload
    fake_adapter = MagicMock()
    fake_adapter.chat.return_value = _completion("  page text  ")
    with patch('backend.services.llm_adapter.OpenAIAdapter', return_value=fake_adapter):
        out = extract_text_from_upload(file_data=b'\x89PNG', filename='p.png', api_key='fake-key')
    assert out == "page text"  # stripped
    fake_adapter.chat.assert_called_once()
