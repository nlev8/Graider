"""
Unit tests for backend/services/document_generator.py.

Audit MAJOR #4 sprint follow-up to PR #251. Targets the 253 uncovered LOC.

Strategy:
- Pure helpers: _hex_to_rgb, _build_document_text — direct assertions.
- File-IO helpers: load_style, save_style, list_styles — HOME redirected
  to tmp_path so no real ~/.graider_data/doc_styles writes (PR #250 R1
  lesson).
- generate_document: entry point — DOCUMENTS_DIR + ASSIGNMENTS_DIR
  redirected to tmp_path; verifies docx file is created and
  save_to_builder writes a valid config.
- create_document_docx: smoke test — round-trip render + read with
  python-docx to confirm core blocks (heading, paragraph, list, table)
  produce non-empty output.

Pattern matches tests/test_assistant_tools_grading_unit.py (PR #251).
"""
from __future__ import annotations

import json
import os
from unittest.mock import patch

import pytest


# ──────────────────────────────────────────────────────────────────
# Fixtures — HOME redirect + module path patches
# ──────────────────────────────────────────────────────────────────


@pytest.fixture
def isolated_dirs(tmp_path, monkeypatch):
    """Redirect HOME + module-level dir constants to tmp_path so file-IO
    functions don't touch the user's real ~/.graider_data or ~/Downloads.

    PR #250 R1 lesson: HOME-redirect from the fixture instead of relying
    on ad-hoc os.path.expanduser patches in each test.
    """
    import backend.services.document_generator as dg

    monkeypatch.setenv("HOME", str(tmp_path))
    monkeypatch.setenv("GRAIDER_EXPORT_DIR", str(tmp_path))
    styles_dir = str(tmp_path / ".graider_data" / "doc_styles")
    assignments_dir = str(tmp_path / ".graider_assignments")

    monkeypatch.setattr(dg, "STYLES_DIR", styles_dir)
    monkeypatch.setattr(dg, "ASSIGNMENTS_DIR", assignments_dir)

    return tmp_path, dg


# ──────────────────────────────────────────────────────────────────
# _hex_to_rgb (pure helper)
# ──────────────────────────────────────────────────────────────────


class TestHexToRgb:
    def test_valid_hex_returns_rgbcolor(self):
        from docx.shared import RGBColor
        from backend.services.document_generator import _hex_to_rgb
        result = _hex_to_rgb("#FF0080")
        assert result is not None
        assert isinstance(result, RGBColor)

    def test_without_hash_prefix(self):
        from backend.services.document_generator import _hex_to_rgb
        assert _hex_to_rgb("FF0080") is not None

    def test_lowercase_valid(self):
        from backend.services.document_generator import _hex_to_rgb
        assert _hex_to_rgb("#abcdef") is not None

    def test_invalid_length_returns_none(self):
        from backend.services.document_generator import _hex_to_rgb
        assert _hex_to_rgb("#FF") is None
        assert _hex_to_rgb("#FF008") is None
        assert _hex_to_rgb("#FF008011") is None

    def test_non_hex_chars_returns_none(self):
        from backend.services.document_generator import _hex_to_rgb
        assert _hex_to_rgb("#GGHHII") is None
        assert _hex_to_rgb("#XYZ123") is None

    def test_none_returns_none(self):
        from backend.services.document_generator import _hex_to_rgb
        assert _hex_to_rgb(None) is None

    def test_empty_string_returns_none(self):
        from backend.services.document_generator import _hex_to_rgb
        assert _hex_to_rgb("") is None

    def test_non_string_returns_none(self):
        from backend.services.document_generator import _hex_to_rgb
        assert _hex_to_rgb(123) is None
        assert _hex_to_rgb([0xFF, 0x00, 0x80]) is None


# ──────────────────────────────────────────────────────────────────
# load_style — defaults + saved style merge + type coercion
# ──────────────────────────────────────────────────────────────────


class TestLoadStyle:
    def test_no_style_name_returns_default(self, isolated_dirs):
        _, dg = isolated_dirs
        result = dg.load_style(None)
        assert result == dg.DEFAULT_STYLE

    def test_empty_string_returns_default(self, isolated_dirs):
        _, dg = isolated_dirs
        result = dg.load_style("")
        assert result == dg.DEFAULT_STYLE

    def test_missing_file_returns_default(self, isolated_dirs):
        _, dg = isolated_dirs
        result = dg.load_style("nonexistent")
        assert result == dg.DEFAULT_STYLE

    def test_saved_style_merged_with_defaults(self, isolated_dirs):
        tmp, dg = isolated_dirs
        os.makedirs(dg.STYLES_DIR, exist_ok=True)
        with open(os.path.join(dg.STYLES_DIR, "custom.json"), 'w') as f:
            json.dump({
                "title_font_name": "Helvetica",
                "title_font_size": 30,
            }, f)

        result = dg.load_style("custom")
        # Override applied
        assert result["title_font_name"] == "Helvetica"
        assert result["title_font_size"] == 30
        # Defaults preserved for non-overridden keys
        assert result["body_font_name"] == dg.DEFAULT_STYLE["body_font_name"]
        assert result["heading_color"] == dg.DEFAULT_STYLE["heading_color"]

    def test_string_int_coerced_to_int(self, isolated_dirs):
        """JSON-from-AI may pass strings; load_style coerces numeric fields."""
        tmp, dg = isolated_dirs
        os.makedirs(dg.STYLES_DIR, exist_ok=True)
        with open(os.path.join(dg.STYLES_DIR, "stringy.json"), 'w') as f:
            json.dump({"title_font_size": "24", "body_font_size": "12"}, f)

        result = dg.load_style("stringy")
        assert result["title_font_size"] == 24
        assert isinstance(result["title_font_size"], int)
        assert result["body_font_size"] == 12

    def test_invalid_int_falls_back_to_default(self, isolated_dirs):
        tmp, dg = isolated_dirs
        os.makedirs(dg.STYLES_DIR, exist_ok=True)
        with open(os.path.join(dg.STYLES_DIR, "bad.json"), 'w') as f:
            json.dump({"title_font_size": "not-a-number"}, f)

        result = dg.load_style("bad")
        assert result["title_font_size"] == dg.DEFAULT_STYLE["title_font_size"]

    def test_string_float_coerced_for_line_spacing(self, isolated_dirs):
        tmp, dg = isolated_dirs
        os.makedirs(dg.STYLES_DIR, exist_ok=True)
        with open(os.path.join(dg.STYLES_DIR, "spacing.json"), 'w') as f:
            json.dump({"line_spacing": "1.5"}, f)

        result = dg.load_style("spacing")
        assert result["line_spacing"] == 1.5
        assert isinstance(result["line_spacing"], float)

    def test_heading_sizes_dict_coerced(self, isolated_dirs):
        tmp, dg = isolated_dirs
        os.makedirs(dg.STYLES_DIR, exist_ok=True)
        with open(os.path.join(dg.STYLES_DIR, "headings.json"), 'w') as f:
            json.dump({
                "heading_sizes": {"1": "20", "2": "16", "3": "bad"},
            }, f)

        result = dg.load_style("headings")
        assert result["heading_sizes"]["1"] == 20
        assert result["heading_sizes"]["2"] == 16
        # Invalid value falls back to default
        assert result["heading_sizes"]["3"] == dg.DEFAULT_STYLE["heading_sizes"]["3"]

    def test_corrupt_json_falls_back_to_default(self, isolated_dirs):
        tmp, dg = isolated_dirs
        os.makedirs(dg.STYLES_DIR, exist_ok=True)
        with open(os.path.join(dg.STYLES_DIR, "corrupt.json"), 'w') as f:
            f.write("not valid json{{{")

        result = dg.load_style("corrupt")
        # Falls back to defaults silently (sentry capture, no raise)
        assert result == dg.DEFAULT_STYLE


# ──────────────────────────────────────────────────────────────────
# save_style + list_styles
# ──────────────────────────────────────────────────────────────────


class TestSaveStyle:
    def test_creates_styles_dir_and_writes_file(self, isolated_dirs):
        tmp, dg = isolated_dirs
        # Dir doesn't exist yet
        assert not os.path.exists(dg.STYLES_DIR)

        result = dg.save_style("my-style", {"title_font_name": "Arial"})

        assert result["status"] == "saved"
        assert result["style_name"] == "my-style"
        assert os.path.exists(result["filepath"])

        with open(result["filepath"]) as f:
            saved = json.load(f)
        assert saved["title_font_name"] == "Arial"
        assert saved["name"] == "my-style"

    def test_filters_unrecognized_keys(self, isolated_dirs):
        tmp, dg = isolated_dirs
        dg.save_style("filtered", {
            "title_font_name": "Arial",
            "unknown_key": "should be dropped",
            "evil_path": "../../etc/passwd",
        })

        with open(os.path.join(dg.STYLES_DIR, "filtered.json")) as f:
            saved = json.load(f)
        assert "title_font_name" in saved
        assert "unknown_key" not in saved
        assert "evil_path" not in saved


class TestListStyles:
    def test_returns_empty_when_no_dir(self, isolated_dirs):
        _, dg = isolated_dirs
        # STYLES_DIR doesn't exist yet
        result = dg.list_styles()
        assert result == {"styles": []}

    def test_returns_sorted_style_names(self, isolated_dirs):
        tmp, dg = isolated_dirs
        os.makedirs(dg.STYLES_DIR)
        for name in ["zebra", "alpha", "middle"]:
            with open(os.path.join(dg.STYLES_DIR, name + ".json"), 'w') as f:
                json.dump({}, f)
        # Non-json files ignored
        with open(os.path.join(dg.STYLES_DIR, "ignore.txt"), 'w') as f:
            f.write("nope")

        result = dg.list_styles()
        assert result["styles"] == ["alpha", "middle", "zebra"]


# ──────────────────────────────────────────────────────────────────
# _build_document_text (pure)
# ──────────────────────────────────────────────────────────────────


class TestBuildDocumentText:
    def test_title_and_paragraphs(self):
        from backend.services.document_generator import _build_document_text
        result = _build_document_text("My Doc", [
            {"type": "paragraph", "text": "Para 1"},
            {"type": "paragraph", "text": "Para 2"},
        ])
        assert "My Doc" in result
        assert "Para 1" in result
        assert "Para 2" in result

    def test_heading_block(self):
        from backend.services.document_generator import _build_document_text
        result = _build_document_text("Title", [
            {"type": "heading", "text": "Section 1", "level": 1},
        ])
        assert "Section 1" in result

    def test_bullet_list(self):
        from backend.services.document_generator import _build_document_text
        result = _build_document_text("Title", [
            {"type": "bullet_list", "items": ["First", "Second", "Third"]},
        ])
        assert "- First" in result
        assert "- Second" in result
        assert "- Third" in result

    def test_numbered_list(self):
        from backend.services.document_generator import _build_document_text
        result = _build_document_text("Title", [
            {"type": "numbered_list", "items": ["A", "B"]},
        ])
        assert "1. A" in result
        assert "2. B" in result

    def test_table(self):
        from backend.services.document_generator import _build_document_text
        result = _build_document_text("Title", [
            {"type": "table", "rows": [["H1", "H2"], ["v1", "v2"]]},
        ])
        assert "H1 | H2" in result
        assert "v1 | v2" in result

    def test_mixed_blocks(self):
        from backend.services.document_generator import _build_document_text
        blocks = [
            {"type": "heading", "text": "Heading", "level": 1},
            {"type": "paragraph", "text": "Body."},
            {"type": "bullet_list", "items": ["Bullet A"]},
        ]
        result = _build_document_text("T", blocks)
        # Lines should be in document order
        idx_heading = result.find("Heading")
        idx_body = result.find("Body.")
        idx_bullet = result.find("- Bullet A")
        assert 0 <= idx_heading < idx_body < idx_bullet


# ──────────────────────────────────────────────────────────────────
# generate_document — entry point
# ──────────────────────────────────────────────────────────────────


class TestGenerateDocument:
    def test_creates_docx_and_returns_metadata(self, isolated_dirs):
        tmp, dg = isolated_dirs
        result = dg.generate_document(
            "Hello World",
            content=[{"type": "paragraph", "text": "Body"}],
        )
        assert result["status"] == "created"
        assert result["filename"] == "Hello World.docx"
        assert os.path.exists(result["filepath"])
        # Download URL is URL-encoded
        assert result["download_url"] == "/api/download-document/Hello%20World.docx"
        assert result["style_used"] == "default"

    def test_unsafe_title_chars_stripped(self, isolated_dirs):
        tmp, dg = isolated_dirs
        result = dg.generate_document(
            "Bad/Title:With*Chars?",
            content=[{"type": "paragraph", "text": "x"}],
        )
        # Only alphanum + space + - + _ kept, then strip
        assert result["filename"] == "BadTitleWithChars.docx"

    def test_custom_style_loaded_and_applied_to_docx(self, isolated_dirs):
        """Codex round-1 MINOR: previous version only checked metadata.
        Now reads back the docx and confirms the saved style's
        distinctive font landed in the title run."""
        from docx import Document
        tmp, dg = isolated_dirs
        os.makedirs(dg.STYLES_DIR, exist_ok=True)
        with open(os.path.join(dg.STYLES_DIR, "custom.json"), 'w') as f:
            json.dump({
                "title_font_name": "Verdana",  # distinctive — not in DEFAULT_STYLE
                "body_font_name": "Times New Roman",
            }, f)

        result = dg.generate_document(
            "Doc", content=[{"type": "paragraph", "text": "Hello"}],
            style_name="custom",
        )
        assert result["style_used"] == "custom"
        assert os.path.exists(result["filepath"])

        # Read back the docx and verify the title run uses Verdana
        doc = Document(result["filepath"])
        title_para = doc.paragraphs[0]
        assert title_para.text == "Doc"
        title_fonts = [r.font.name for r in title_para.runs if r.font.name]
        assert "Verdana" in title_fonts, (
            f"Custom style title_font_name='Verdana' was not applied; "
            f"title runs use {title_fonts!r}"
        )

    def test_save_to_builder_writes_assignment_config(self, isolated_dirs):
        tmp, dg = isolated_dirs
        result = dg.generate_document(
            "Quiz Title",
            content=[
                {"type": "heading", "text": "Section 1"},
                {"type": "paragraph", "text": "Body"},
            ],
            save_to_builder=True,
        )
        assert result.get("saved_to_builder") is True
        assert result["config_name"] == "Quiz Title"
        config_path = os.path.join(dg.ASSIGNMENTS_DIR, "Quiz Title.json")
        assert os.path.exists(config_path)
        with open(config_path) as f:
            config = json.load(f)
        # Schema sanity
        assert config["title"] == "Quiz Title"
        assert config["totalPoints"] == 100
        assert "Section 1" in config["importedDoc"]["text"]
        assert config["worksheetDownloadUrl"].startswith("/api/download-document/")

    def test_no_save_to_builder_means_no_config(self, isolated_dirs):
        tmp, dg = isolated_dirs
        result = dg.generate_document(
            "NotSaved", content=[],
        )
        assert "saved_to_builder" not in result
        assert not os.path.exists(os.path.join(dg.ASSIGNMENTS_DIR, "NotSaved.json"))


# ──────────────────────────────────────────────────────────────────
# create_document_docx — round-trip smoke
# ──────────────────────────────────────────────────────────────────


class TestCreateDocumentDocx:
    def test_round_trip_paragraph_and_heading(self, isolated_dirs):
        from docx import Document
        from backend.services.document_generator import create_document_docx, DEFAULT_STYLE

        tmp, dg = isolated_dirs
        path = str(tmp / "out.docx")
        os.makedirs(os.path.dirname(path), exist_ok=True)

        create_document_docx(
            path,
            "My Title",
            [
                {"type": "heading", "text": "Section A", "level": 1},
                {"type": "paragraph", "text": "Para body."},
            ],
            DEFAULT_STYLE,
        )

        # Read back to confirm content was written
        doc = Document(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "My Title" in all_text
        assert "Section A" in all_text
        assert "Para body." in all_text

    def test_round_trip_bullet_list(self, isolated_dirs):
        from docx import Document
        from backend.services.document_generator import create_document_docx, DEFAULT_STYLE

        tmp, dg = isolated_dirs
        path = str(tmp / "bullets.docx")

        create_document_docx(
            path, "List Doc",
            [{"type": "bullet_list", "items": ["A", "B", "C"]}],
            DEFAULT_STYLE,
        )

        doc = Document(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "A" in all_text
        assert "B" in all_text
        assert "C" in all_text

    def test_round_trip_table(self, isolated_dirs):
        from docx import Document
        from backend.services.document_generator import create_document_docx, DEFAULT_STYLE

        tmp, dg = isolated_dirs
        path = str(tmp / "tbl.docx")

        create_document_docx(
            path, "Table Doc",
            [{"type": "table", "rows": [["Header A", "Header B"], ["v1", "v2"]]}],
            DEFAULT_STYLE,
        )

        doc = Document(path)
        assert len(doc.tables) == 1
        table = doc.tables[0]
        assert table.cell(0, 0).text == "Header A"
        assert table.cell(0, 1).text == "Header B"
        assert table.cell(1, 0).text == "v1"
        assert table.cell(1, 1).text == "v2"

    def test_empty_blocks_does_not_crash(self, isolated_dirs):
        from docx import Document
        from backend.services.document_generator import create_document_docx, DEFAULT_STYLE

        tmp, dg = isolated_dirs
        path = str(tmp / "empty.docx")

        create_document_docx(path, "Just a Title", [], DEFAULT_STYLE)
        doc = Document(path)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Just a Title" in all_text


# ──────────────────────────────────────────────────────────────────
# _parse_markdown_runs — bold/italic markers split into runs
# ──────────────────────────────────────────────────────────────────


class TestParseMarkdownRuns:
    def _make_paragraph(self):
        from docx import Document
        doc = Document()
        return doc.add_paragraph(), doc

    def test_plain_text_single_run(self):
        from backend.services.document_generator import _parse_markdown_runs
        para, _ = self._make_paragraph()
        _parse_markdown_runs(para, "Just plain text")
        # All runs concatenated equals input
        joined = "".join(r.text for r in para.runs)
        assert joined == "Just plain text"

    def test_bold_marker_creates_bold_run(self):
        from backend.services.document_generator import _parse_markdown_runs
        para, _ = self._make_paragraph()
        _parse_markdown_runs(para, "Hello **bold** world")
        # Find the run containing 'bold'
        bold_runs = [r for r in para.runs if "bold" in r.text]
        assert len(bold_runs) >= 1
        assert any(r.bold for r in bold_runs), "Expected a bold run"
        # Joined text drops the markers
        joined = "".join(r.text for r in para.runs)
        assert joined == "Hello bold world"

    def test_italic_marker_creates_italic_run(self):
        from backend.services.document_generator import _parse_markdown_runs
        para, _ = self._make_paragraph()
        _parse_markdown_runs(para, "Hello *italic* world")
        italic_runs = [r for r in para.runs if "italic" in r.text]
        assert len(italic_runs) >= 1
        assert any(r.italic for r in italic_runs), "Expected an italic run"
        joined = "".join(r.text for r in para.runs)
        assert joined == "Hello italic world"

    def test_bold_italic_marker(self):
        from backend.services.document_generator import _parse_markdown_runs
        para, _ = self._make_paragraph()
        _parse_markdown_runs(para, "***bold-italic***")
        joined = "".join(r.text for r in para.runs)
        assert joined == "bold-italic"
        # At least one run should be both bold and italic
        bi_runs = [r for r in para.runs if r.bold and r.italic]
        assert len(bi_runs) >= 1
