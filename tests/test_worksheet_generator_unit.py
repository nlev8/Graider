"""
Unit tests for backend/services/worksheet_generator.py.

Audit MAJOR #4 sprint follow-up to PR #259. Targets the 250 uncovered
LOC in worksheet_generator.py (29% baseline). Pattern matches PR #252
(document_generator).

Strategy:
- Pure logic helpers (`_normalize_correct_answer_to_letter`,
  `_build_document_text`, `_build_assignment_config`) via direct assertions.
- DOCX generators (`generate_worksheet`, `_create_answer_key_doc`,
  `_add_options_with_bubbles`) via round-trip docx.Document inspection.
- HOME redirect from `isolated_dirs` fixture (PR #250 R1 lesson) so
  no real `~/Downloads/Graider/` or `~/.graider_assignments/` writes.
"""
from __future__ import annotations

import json
import os

import pytest


@pytest.fixture
def isolated_dirs(tmp_path, monkeypatch):
    """Redirect HOME + module path constants to tmp_path."""
    import backend.services.worksheet_generator as wg

    monkeypatch.setenv("HOME", str(tmp_path))
    monkeypatch.setenv("GRAIDER_EXPORT_DIR", str(tmp_path / "Downloads" / "Graider"))
    monkeypatch.setattr(wg, "ASSIGNMENTS_DIR", str(tmp_path / "assignments"))
    return tmp_path, wg


# ──────────────────────────────────────────────────────────────────
# _normalize_correct_answer_to_letter
# ──────────────────────────────────────────────────────────────────


class TestNormalizeCorrectAnswerToLetter:
    def test_single_letter(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        opts = ["A) Foo", "B) Bar", "C) Baz"]
        assert _normalize_correct_answer_to_letter("B", opts) == "B"
        # Lowercase is normalized
        assert _normalize_correct_answer_to_letter("c", opts) == "C"

    def test_letter_with_paren(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        opts = ["A) Foo", "B) Bar"]
        assert _normalize_correct_answer_to_letter("B)", opts) == "B"
        assert _normalize_correct_answer_to_letter("A.", opts) == "A"

    def test_zero_index_returns_none_due_to_top_guard(self):
        """Pinned production behavior: the early `if not correct_answer`
        guard treats 0 as falsy and short-circuits to None — even though
        0 is a valid zero-based index. Documented quirk; pinning so a
        future change to the guard doesn't accidentally drop this branch."""
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        assert _normalize_correct_answer_to_letter(0, ["X", "Y", "Z"]) is None

    def test_nonzero_zero_based_integer_index(self):
        """Non-zero indices: 0-based works."""
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        opts = ["X", "Y", "Z"]
        assert _normalize_correct_answer_to_letter(2, opts) == "C"

    def test_one_based_integer_index_fallback(self):
        """When index ≥ len(opts) (out of 0-based range) but fits 1-based."""
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        opts = ["X", "Y"]
        # 0-based: idx=2 → out of range. 1-based: idx=2 → "Y" (B)
        assert _normalize_correct_answer_to_letter(2, opts) == "B"

    def test_full_option_text_match(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        opts = ["A) Constitution", "B) Declaration of Independence", "C) Bill of Rights"]
        assert _normalize_correct_answer_to_letter(
            "Declaration of Independence", opts
        ) == "B"

    def test_full_match_with_letter_prefix_stripped(self):
        """Match works even when option_text has letter prefix and answer doesn't."""
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        opts = ["A) Constitution", "B) Bill of Rights"]
        assert _normalize_correct_answer_to_letter("Bill of Rights", opts) == "B"

    def test_no_match_returns_none(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        opts = ["A) Foo", "B) Bar"]
        assert _normalize_correct_answer_to_letter("Random", opts) is None

    def test_empty_inputs(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        assert _normalize_correct_answer_to_letter("", ["A"]) is None
        assert _normalize_correct_answer_to_letter("A", []) is None
        assert _normalize_correct_answer_to_letter(None, ["A"]) is None


# ──────────────────────────────────────────────────────────────────
# _build_document_text (pure)
# ──────────────────────────────────────────────────────────────────


class TestBuildDocumentText:
    def test_minimal_includes_name_date_period_lines(self):
        from backend.services.worksheet_generator import _build_document_text
        result = _build_document_text("My Worksheet", [], [], None, [], 100)
        assert "My Worksheet" in result
        assert "Name:" in result
        assert "Date:" in result
        assert "Period:" in result

    def test_includes_vocabulary_section(self):
        from backend.services.worksheet_generator import _build_document_text
        terms = [{"term": "ratio", "definition": "comparison of two numbers"}]
        result = _build_document_text("WS", terms, [], None, [], 100)
        assert "VOCABULARY" in result
        assert "ratio:" in result

    def test_includes_questions_section_with_points(self):
        from backend.services.worksheet_generator import _build_document_text
        questions = [
            {"question": "What is 2+2?", "points": 5},
            {"question": "What is 3*3?", "points": 10},
        ]
        result = _build_document_text("WS", [], questions, None, [], 100)
        assert "QUESTIONS" in result
        assert "1) What is 2+2?" in result
        assert "(5 pts)" in result
        assert "2) What is 3*3?" in result
        assert "(10 pts)" in result

    def test_includes_summary_when_prompt_set(self):
        from backend.services.worksheet_generator import _build_document_text
        result = _build_document_text(
            "WS", [], [], "Summarize the chapter.", [], 100,
        )
        assert "SUMMARY" in result
        assert "Summarize the chapter." in result

    def test_no_summary_section_when_prompt_empty(self):
        from backend.services.worksheet_generator import _build_document_text
        result = _build_document_text("WS", [], [], None, [], 100)
        assert "SUMMARY" not in result


# ──────────────────────────────────────────────────────────────────
# _build_assignment_config (pure)
# ──────────────────────────────────────────────────────────────────


class TestBuildAssignmentConfig:
    def test_basic_shape(self):
        from backend.services.worksheet_generator import _build_assignment_config
        config = _build_assignment_config(
            "Quiz 1", "short-answer", [], [{"question": "x?", "expected_answer": "y"}],
            None, [], 100,
        )
        assert config["title"] == "Quiz 1"
        assert config["totalPoints"] == 100
        assert config["effortPoints"] == 15
        assert config["countsTowardsGrade"] is True
        assert config["completionOnly"] is False

    def test_rubric_type_mapping(self):
        from backend.services.worksheet_generator import _build_assignment_config
        # cornell-notes maps to cornell-notes
        c = _build_assignment_config("X", "cornell-notes", [], [], None, [], 100)
        assert c["rubricType"] == "cornell-notes"
        # fill-in-blank
        c = _build_assignment_config("X", "fill-in-blank", [], [], None, [], 100)
        assert c["rubricType"] == "fill-in-blank"
        # short-answer → standard
        c = _build_assignment_config("X", "short-answer", [], [], None, [], 100)
        assert c["rubricType"] == "standard"
        # vocabulary → fill-in-blank
        c = _build_assignment_config("X", "vocabulary", [], [], None, [], 100)
        assert c["rubricType"] == "fill-in-blank"
        # Unknown type → standard fallback
        c = _build_assignment_config("X", "weird-type", [], [], None, [], 100)
        assert c["rubricType"] == "standard"

    def test_vocab_marker_emitted(self):
        from backend.services.worksheet_generator import _build_assignment_config
        terms = [{"term": "ratio", "definition": "x"}, {"term": "fraction", "definition": "y"}]
        config = _build_assignment_config(
            "X", "vocabulary", terms, [], None, [], 100,
        )
        markers = config["customMarkers"]
        vocab_marker = next((m for m in markers if m["start"] == "VOCABULARY"), None)
        assert vocab_marker is not None
        assert vocab_marker["type"] == "vocab_term"
        # 2 vocab questions in question_objects
        vocab_qs = [q for q in config["questions"] if q["type"] == "vocab_term"]
        assert len(vocab_qs) == 2

    def test_questions_marker_emitted(self):
        from backend.services.worksheet_generator import _build_assignment_config
        qs = [{"question": "a", "expected_answer": "1"}]
        config = _build_assignment_config(
            "X", "short-answer", [], qs, None, [], 100,
        )
        q_marker = next(
            (m for m in config["customMarkers"] if m["start"] == "QUESTIONS"), None
        )
        assert q_marker is not None
        assert q_marker["type"] == "short_answer"

    def test_summary_marker_emitted(self):
        from backend.services.worksheet_generator import _build_assignment_config
        config = _build_assignment_config(
            "X", "short-answer", [], [], "Summarize.", [], 100,
        )
        s_marker = next(
            (m for m in config["customMarkers"] if m["start"] == "SUMMARY"), None
        )
        assert s_marker is not None
        assert s_marker["type"] == "written"

    def test_grading_notes_includes_expected_answers(self):
        from backend.services.worksheet_generator import _build_assignment_config
        terms = [{"term": "ratio", "definition": "comparison"}]
        qs = [{"question": "What is X?", "expected_answer": "42"}]
        config = _build_assignment_config(
            "X", "short-answer", terms, qs, None, [], 100,
        )
        notes = config["gradingNotes"]
        assert "VOCABULARY EXPECTED DEFINITIONS:" in notes
        assert "ratio: comparison" in notes
        assert "EXPECTED ANSWERS:" in notes
        assert "Q1: 42" in notes

    def test_summary_key_points_included_in_grading_notes(self):
        from backend.services.worksheet_generator import _build_assignment_config
        config = _build_assignment_config(
            "X", "short-answer", [], [], "Summarize.",
            ["Industrial Revolution caused urban growth", "Steam engine impact"], 100,
        )
        notes = config["gradingNotes"]
        assert "SUMMARY SHOULD INCLUDE:" in notes
        assert "Industrial Revolution" in notes
        assert "Steam engine impact" in notes

    def test_subject_passes_through(self):
        from backend.services.worksheet_generator import _build_assignment_config
        config = _build_assignment_config(
            "X", "short-answer", [], [], None, [], 100, subject="Math",
        )
        assert config["subject"] == "Math"

    def test_point_distribution_across_items(self):
        """content_points (total - 15 effort) split across (vocab + questions + summary)."""
        from backend.services.worksheet_generator import _build_assignment_config
        terms = [{"term": "x", "definition": "y"}, {"term": "a", "definition": "b"}]
        qs = [{"question": "?", "expected_answer": "!"}]
        # Total 100, 15 for effort, 85 for content, 3 items → ~28 each
        config = _build_assignment_config(
            "X", "short-answer", terms, qs, "summary", [], 100,
        )
        # Sum of marker points should approximately equal content_points (85)
        total_marker_pts = sum(m["points"] for m in config["customMarkers"])
        assert 80 <= total_marker_pts <= 90  # approximate due to rounding


# ──────────────────────────────────────────────────────────────────
# _create_answer_key_doc
# ──────────────────────────────────────────────────────────────────


class TestCreateAnswerKeyDoc:
    def test_empty_questions_returns_none(self):
        from backend.services.worksheet_generator import _create_answer_key_doc
        assert _create_answer_key_doc("Quiz", []) is None

    def test_creates_doc_with_title_and_warning(self):
        from backend.services.worksheet_generator import _create_answer_key_doc
        questions = [{
            "number": 1, "question_text": "What is X?",
            "option_texts": ["A) Foo", "B) Bar"], "correct_answer": "B",
            "is_tf": False,
        }]
        doc = _create_answer_key_doc("Quiz 1", questions)
        assert doc is not None
        all_text = "\n".join(p.text for p in doc.paragraphs)
        # Header includes title + ANSWER KEY
        assert "Quiz 1" in all_text
        assert "ANSWER KEY" in all_text
        # Teacher warning
        assert "teacher use only" in all_text.lower()

    def test_includes_question_text_and_options(self):
        from backend.services.worksheet_generator import _create_answer_key_doc
        questions = [{
            "number": 5, "question_text": "Capital of France?",
            "option_texts": ["A) Berlin", "B) Paris", "C) Madrid"],
            "correct_answer": "B", "is_tf": False,
        }]
        doc = _create_answer_key_doc("Geo Quiz", questions)
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Capital of France?" in all_text
        # All option texts present
        assert "Berlin" in all_text
        assert "Paris" in all_text
        assert "Madrid" in all_text

    def test_correct_answer_filled_bubble_on_right_option_only(self):
        """The correct option's paragraph contains the filled bubble (●);
        the wrong option's paragraph contains the empty bubble (○).
        Codex round-1 MINOR: previous version had a duplicate-assertion
        weakness (`'●' in s or '●' in s`) — strengthened to assert the
        filled bubble lives on the correct paragraph specifically and
        the empty bubble lives on the wrong paragraph."""
        from backend.services.worksheet_generator import _create_answer_key_doc
        questions = [{
            "number": 1, "question_text": "Q?",
            "option_texts": ["A) Wrong", "B) Right"],
            "correct_answer": "B", "is_tf": False,
        }]
        doc = _create_answer_key_doc("Quiz", questions)
        # Find paragraphs containing each option text
        wrong_para = next(
            (p for p in doc.paragraphs if "A) Wrong" in p.text), None
        )
        right_para = next(
            (p for p in doc.paragraphs if "B) Right" in p.text), None
        )
        assert wrong_para is not None, "A) Wrong paragraph not rendered"
        assert right_para is not None, "B) Right paragraph not rendered"
        # Correct option has filled bubble (●), wrong option has empty (○)
        assert "●" in right_para.text, (
            f"Correct option missing filled bubble: {right_para.text!r}"
        )
        assert "○" in wrong_para.text, (
            f"Wrong option missing empty bubble: {wrong_para.text!r}"
        )
        # And the wrong option does NOT have a filled bubble
        assert "●" not in wrong_para.text, (
            f"Wrong option has filled bubble (should be empty): {wrong_para.text!r}"
        )


# ──────────────────────────────────────────────────────────────────
# generate_worksheet (entry point with file I/O)
# ──────────────────────────────────────────────────────────────────


class TestGenerateWorksheet:
    def test_creates_docx_and_config(self, isolated_dirs):
        tmp, wg = isolated_dirs
        result = wg.generate_worksheet(
            title="My Worksheet",
            worksheet_type="short-answer",
            questions=[{"question": "What is 2+2?", "expected_answer": "4"}],
            total_points=100,
        )
        assert result["status"] == "created"
        assert result["filename"] == "My Worksheet.docx"
        assert result["saved_to_builder"] is True
        # File written
        assert os.path.exists(result["filepath"])
        # Config written
        config_path = os.path.join(wg.ASSIGNMENTS_DIR, "My Worksheet.json")
        assert os.path.exists(config_path)
        with open(config_path) as f:
            config = json.load(f)
        # Schema sanity
        assert config["title"] == "My Worksheet"
        assert config["totalPoints"] == 100
        # importedDoc populated
        assert "What is 2+2?" in config["importedDoc"]["text"]
        assert config["worksheetDownloadUrl"].startswith("/api/download-worksheet/")

    def test_unsafe_title_chars_stripped(self, isolated_dirs):
        tmp, wg = isolated_dirs
        result = wg.generate_worksheet(
            title="Bad/Title:With*Chars?",
            worksheet_type="short-answer",
            questions=[{"question": "Q", "expected_answer": "A"}],
        )
        assert result["filename"] == "BadTitleWithChars.docx"
        assert os.path.exists(result["filepath"])

    def test_default_style_when_no_style_name(self, isolated_dirs):
        tmp, wg = isolated_dirs
        result = wg.generate_worksheet(
            title="Default Style Test",
            worksheet_type="cornell-notes",
            vocab_terms=[{"term": "x", "definition": "y"}],
        )
        assert result["style_used"] == "default"

    def test_round_trip_docx_includes_title(self, isolated_dirs):
        from docx import Document
        tmp, wg = isolated_dirs
        result = wg.generate_worksheet(
            title="Round Trip Test",
            worksheet_type="short-answer",
            questions=[{"question": "What is X?", "expected_answer": "1"}],
        )
        doc = Document(result["filepath"])
        all_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Round Trip Test" in all_text
