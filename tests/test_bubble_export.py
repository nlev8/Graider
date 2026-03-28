"""Unit tests for bubble answer sheet export."""
import pytest


class TestNormalizeAnswer:
    def test_letter(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        assert _normalize_correct_answer_to_letter('B', ['A) X', 'B) Y']) == 'B'

    def test_letter_lowercase(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        assert _normalize_correct_answer_to_letter('c', ['A) X', 'B) Y', 'C) Z']) == 'C'

    def test_letter_paren(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        assert _normalize_correct_answer_to_letter('C)', ['A) X', 'B) Y', 'C) Z']) == 'C'

    def test_full_text(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        result = _normalize_correct_answer_to_letter(
            'Paris', ['A) London', 'B) Berlin', 'C) Paris', 'D) Madrid'])
        assert result == 'C'

    def test_no_match(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        assert _normalize_correct_answer_to_letter('XYZ', ['A) One', 'B) Two']) is None

    def test_index_zero_based(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        assert _normalize_correct_answer_to_letter('1', ['A) X', 'B) Y', 'C) Z']) == 'B'

    def test_empty_answer(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        assert _normalize_correct_answer_to_letter('', ['A) X']) is None

    def test_empty_options(self):
        from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
        assert _normalize_correct_answer_to_letter('B', []) is None


class TestAddOptionWithBubble:
    def test_empty_bubble(self):
        from docx import Document
        from backend.services.worksheet_generator import _add_option_with_bubble
        doc = Document()
        para = _add_option_with_bubble(doc, 'A) Constitution', is_filled=False)
        text = para.text
        assert 'A) Constitution' in text
        assert '\u25CB' in text  # empty circle
        assert '\u25CF' not in text  # no filled

    def test_filled_bubble(self):
        from docx import Document
        from backend.services.worksheet_generator import _add_option_with_bubble
        doc = Document()
        para = _add_option_with_bubble(doc, 'B) Declaration of Independence', is_filled=True)
        text = para.text
        assert 'B) Declaration of Independence' in text
        assert '\u25CF' in text  # filled circle

    def test_no_question_number(self):
        """Bubble rows should NOT contain question numbers."""
        from docx import Document
        from backend.services.worksheet_generator import _add_option_with_bubble
        doc = Document()
        para = _add_option_with_bubble(doc, 'A) Test', is_filled=False)
        text = para.text
        # Should not start with a number
        assert not text.strip().startswith('1')
        assert not text.strip().startswith('2')


class TestAddOptionsWithBubbles:
    def test_empty_bubbles_mc(self):
        from docx import Document
        from backend.services.worksheet_generator import _add_options_with_bubbles
        doc = Document()
        _add_options_with_bubbles(doc, ['A) One', 'B) Two', 'C) Three', 'D) Four'])
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert 'A) One' in full_text
        assert 'D) Four' in full_text
        assert '\u25CB' in full_text  # empty circles
        assert '\u25CF' not in full_text  # no filled

    def test_filled_correct_mc(self):
        from docx import Document
        from backend.services.worksheet_generator import _add_options_with_bubbles
        doc = Document()
        _add_options_with_bubbles(doc, ['A) One', 'B) Two', 'C) Three'], correct_answer='B')
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert full_text.count('\u25CF') == 1  # one filled

    def test_filled_correct_tf(self):
        from docx import Document
        from backend.services.worksheet_generator import _add_options_with_bubbles
        doc = Document()
        _add_options_with_bubbles(doc, ['True', 'False'], correct_answer='True', is_tf=True)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert full_text.count('\u25CF') == 1

    def test_full_text_answer(self):
        from docx import Document
        from backend.services.worksheet_generator import _add_options_with_bubbles
        doc = Document()
        _add_options_with_bubbles(doc,
                                  ['A) London', 'B) Paris', 'C) Berlin'],
                                  correct_answer='Paris')
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert full_text.count('\u25CF') == 1

    def test_no_match_no_fill(self):
        from docx import Document
        from backend.services.worksheet_generator import _add_options_with_bubbles
        doc = Document()
        _add_options_with_bubbles(doc, ['A) X', 'B) Y'], correct_answer='Nonsense')
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert '\u25CF' not in full_text


class TestCreateAnswerKeyDoc:
    def test_creates_separate_document(self):
        from backend.services.worksheet_generator import _create_answer_key_doc
        questions = [
            {"number": 1, "option_texts": ["A) X", "B) Y", "C) Z", "D) W"],
             "correct_answer": "B", "is_tf": False, "question_text": "What is Y?"},
        ]
        doc = _create_answer_key_doc("Test Quiz", questions)
        assert doc is not None
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "ANSWER KEY" in full_text
        assert "Test Quiz" in full_text
        assert "teacher use only" in full_text.lower()

    def test_has_filled_bubbles(self):
        from backend.services.worksheet_generator import _create_answer_key_doc
        questions = [
            {"number": 1, "option_texts": ["A) X", "B) Y"], "correct_answer": "B",
             "is_tf": False, "question_text": "Pick Y"},
        ]
        doc = _create_answer_key_doc("Quiz", questions)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert '\u25CF' in full_text

    def test_empty_list_returns_none(self):
        from backend.services.worksheet_generator import _create_answer_key_doc
        assert _create_answer_key_doc("Quiz", []) is None

    def test_includes_question_text(self):
        from backend.services.worksheet_generator import _create_answer_key_doc
        questions = [
            {"number": 1, "option_texts": ["True", "False"], "correct_answer": "True",
             "is_tf": True, "question_text": "Water boils at 100C"},
        ]
        doc = _create_answer_key_doc("Science Quiz", questions)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        assert "Water boils" in full_text
