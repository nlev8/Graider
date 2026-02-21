"""
Assignment Grader for 6th Grade Social Studies
===============================================
Tailored for Southwestern Middle School

FILE NAMING EXPECTED: FirstName_LastName_AssignmentName.docx
ROSTER FORMAT: "LastName; FirstName" with Student ID and Email columns

FERPA COMPLIANCE:
- Student names are NOT sent to OpenAI's API
- Only assignment content is analyzed for grading
- All student identification stays local on your computer
- OpenAI API data is not used to train models (per their policy)
- Consult your district's policies for AI tool usage

SETUP:
1. pip install openai python-docx openpyxl python-dotenv
2. Create .env file with: OPENAI_API_KEY=your-key-here
3. Update the folder paths below
4. Update the ASSIGNMENT_INSTRUCTIONS for each assignment
"""

import os
import csv
import json
import re
import random
import threading
import concurrent.futures
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv
from typing import List, Optional

# Structured output models for reliable JSON responses from OpenAI
from pydantic import BaseModel


class GradingBreakdown(BaseModel):
    content_accuracy: int
    completeness: int
    writing_quality: int
    effort_engagement: int


class SkillsDemonstrated(BaseModel):
    strengths: List[str]
    developing: List[str]


class AiDetectionResult(BaseModel):
    flag: str  # "none", "unlikely", "possible", "likely"
    confidence: int
    reason: str


class PlagiarismDetectionResult(BaseModel):
    flag: str  # "none", "possible", "likely"
    reason: str


class GradingResponse(BaseModel):
    score: int
    letter_grade: str
    breakdown: GradingBreakdown
    student_responses: List[str]
    unanswered_questions: List[str]
    excellent_answers: List[str]
    needs_improvement: List[str]
    skills_demonstrated: SkillsDemonstrated
    ai_detection: AiDetectionResult
    plagiarism_detection: PlagiarismDetectionResult
    feedback: str


class DetectionResponse(BaseModel):
    ai_detection: AiDetectionResult
    plagiarism_detection: PlagiarismDetectionResult

# Import student history for personalized feedback
try:
    from backend.student_history import build_history_context
except ImportError:
    # Fallback if running standalone or module not available
    def build_history_context(student_id):
        return ""

# Import accommodations for IEP/504 support (FERPA compliant)
try:
    from backend.accommodations import build_accommodation_prompt
except ImportError:
    # Fallback if running standalone or module not available
    def build_accommodation_prompt(student_id):
        return ""

# Load environment variables from .env file (override system env vars)
import os
app_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(app_dir, '.env'), override=True)

# =============================================================================
# TOKEN / COST TRACKING
# =============================================================================

MODEL_PRICING = {
    # OpenAI — price per 1M tokens
    "gpt-4o-mini":    {"input": 0.15,  "output": 0.60},
    "gpt-4o":         {"input": 2.50,  "output": 10.00},
    # Claude
    "claude-3-5-haiku-latest":    {"input": 0.80,  "output": 4.00},
    "claude-haiku-4-5-20251001":  {"input": 0.80,  "output": 4.00},
    "claude-sonnet-4-20250514":   {"input": 3.00,  "output": 15.00},
    "claude-opus-4-20250514":     {"input": 15.00, "output": 75.00},
    # Gemini
    "gemini-2.0-flash":    {"input": 0.10,  "output": 0.40},
    "gemini-2.0-pro-exp":  {"input": 1.25,  "output": 5.00},
}

class TokenTracker:
    """Accumulates token usage across multiple API calls for a single student grading."""

    def __init__(self):
        self._lock = threading.Lock()
        self.total_input_tokens = 0
        self.total_output_tokens = 0
        self.calls = []

    def record_openai(self, response, model: str):
        if not response or not hasattr(response, 'usage') or not response.usage:
            return
        inp = response.usage.prompt_tokens or 0
        out = response.usage.completion_tokens or 0
        self._add(model, inp, out)

    def record_anthropic(self, response, model: str):
        if not response or not hasattr(response, 'usage') or not response.usage:
            return
        inp = response.usage.input_tokens or 0
        out = response.usage.output_tokens or 0
        self._add(model, inp, out)

    def record_gemini(self, response, model: str):
        if not response or not hasattr(response, 'usage_metadata') or not response.usage_metadata:
            return
        inp = getattr(response.usage_metadata, 'prompt_token_count', 0) or 0
        out = getattr(response.usage_metadata, 'candidates_token_count', 0) or 0
        self._add(model, inp, out)

    def _add(self, model: str, input_tokens: int, output_tokens: int):
        pricing = MODEL_PRICING.get(model, {"input": 0, "output": 0})
        cost = (input_tokens * pricing["input"] + output_tokens * pricing["output"]) / 1_000_000
        with self._lock:
            self.total_input_tokens += input_tokens
            self.total_output_tokens += output_tokens
            self.calls.append({
                "model": model,
                "input_tokens": input_tokens,
                "output_tokens": output_tokens,
                "cost": round(cost, 6)
            })

    def summary(self) -> dict:
        total_cost = sum(c["cost"] for c in self.calls)
        return {
            "total_input_tokens": self.total_input_tokens,
            "total_output_tokens": self.total_output_tokens,
            "total_tokens": self.total_input_tokens + self.total_output_tokens,
            "total_cost": round(total_cost, 6),
            "total_cost_display": f"${total_cost:.4f}",
            "api_calls": len(self.calls),
            "calls": self.calls
        }


# =============================================================================
# PRE-EXTRACTION - Extract student responses BEFORE AI grading (prevents hallucination)
# =============================================================================

def is_question_or_prompt(text: str) -> bool:
    """
    Check if text looks like a question or instruction prompt rather than a student response.
    Returns True if the text should NOT be treated as a student response.

    Uses a SCORING approach with positive signals (instruction-like) and negative signals
    (student-content-like) so it generalizes to any assignment without needing specific keywords.
    """
    if not text or not text.strip():
        return True  # Empty is not a valid response

    text = text.strip()
    text_lower = text.lower()

    # === QUICK CHECKS ===

    # Ends with question mark → question
    if text.rstrip().endswith('?'):
        # But if there's an answer after a question mark on the same line, it's a response
        # e.g., "What year? 1803"
        if '?' in text:
            after_q = text.split('?', 1)[1].strip()
            if after_q and len(after_q) > 2:
                return False
        return True

    # Just a number with period/paren (bare question number)
    if re.match(r'^\d+[\.\)]\s*$', text_lower):
        return True

    # === SCORING-BASED DETECTION ===
    # Positive score = instruction/prompt, Negative score = student response
    score = 0

    # --- POSITIVE SIGNALS (instruction-like) ---

    # Signal 1: Starts with imperative verb (the most reliable signal)
    # These are verbs used to give instructions, in base form without a subject
    imperative_starts = re.match(
        r'^(summarize|include|write|explain|describe|identify|compare|analyze|'
        r'discuss|list|define|name|state|create|draw|use|provide|give|think|'
        r'consider|refer|remember|note|evaluate|assess|determine|select|choose|'
        r'complete|answer|fill|circle|underline|highlight|review|read|examine|'
        r'outline|support|justify|predict|infer|contrast|classify|organize|'
        r'paraphrase|restate|cite|mention|address|respond|demonstrate|illustrate|'
        r'show|prove|argue|defend|critique|interpret|apply|connect|relate|'
        r'be sure|make sure|don\'t forget|do not forget)\b',
        text_lower
    )
    if imperative_starts:
        score += 3

    # Signal 2: Starts with question word (what/how/why/etc.)
    question_start = re.match(
        r'^(what|how|why|when|where|who|which)\b', text_lower
    )
    if question_start:
        # Question words WITHOUT ? and long text are often student statements
        # "How they resolved the issue was by..." is a student response
        if '?' not in text and len(text) > 40:
            score -= 1  # Likely a student statement
        else:
            score += 3  # Likely a question prompt

    # Signal 3: Contains quantifier/requirement language
    # "in 3-4 sentences", "at least 2 examples", "a 5-paragraph essay"
    if re.search(r'\d+[\-–]\d+\s*(sentences?|words?|paragraphs?|examples?|reasons?|points?)', text_lower):
        score += 3
    if re.search(r'at least\s+\d+', text_lower):
        score += 2

    # Signal 4: Contains assignment meta-language (addresses the student about the task)
    meta_phrases = (
        'your answer', 'your response', 'your own words', 'your summary',
        'the text', 'the reading', 'the passage', 'the chapter', 'the article',
        'the space below', 'the lines below', 'the box below',
        'from the reading', 'from the text', 'from the passage',
        'using evidence', 'using details', 'using examples', 'using information',
        'based on the', 'refer to the', 'according to the',
        'key reasons', 'key points', 'key events', 'key figures', 'key details',
        'main ideas', 'main causes', 'main points', 'main events',
        'supporting details', 'supporting evidence', 'textual evidence',
    )
    meta_count = sum(1 for phrase in meta_phrases if phrase in text_lower)
    score += min(meta_count * 2, 4)  # Cap at +4

    # Signal 5: Short text (< 100 chars) starting with imperative = very likely instruction
    if imperative_starts and len(text) < 100:
        score += 1

    # Signal 6: Contains "you" / "your" addressing the student
    if re.search(r'\byou(r)?\b', text_lower):
        score += 1

    # Signal 7: Instruction structure patterns
    if re.match(r'^(in\s+\d+|after\s+reading|before\s+reading|based\s+on|using\s+the)\b', text_lower):
        score += 3

    # --- NEGATIVE SIGNALS (student-response-like) ---

    # Counter 1: Contains specific historical dates (1700s-2000s)
    year_matches = re.findall(r'\b(1[5-9]\d{2}|20\d{2})\b', text)
    if year_matches:
        score -= 2

    # Counter 2: Past tense narrative verbs (student writing about events)
    past_verbs = re.findall(
        r'\b(was|were|had|did|led|fought|signed|began|ended|caused|resulted|'
        r'believed|wanted|decided|became|created|established|passed|declared|'
        r'invaded|conquered|defeated|surrendered|negotiated|agreed|refused|'
        r'moved|settled|built|discovered|explored|traveled|arrived|died|killed|'
        r'escaped|captured|freed|enslaved|governed|ruled|elected|appointed)\b',
        text_lower
    )
    if len(past_verbs) >= 2:
        score -= 3
    elif len(past_verbs) == 1:
        score -= 1

    # Counter 3: Proper nouns mid-sentence (names of people, places, events)
    # Only a strong signal when combined with past tense verbs (narrative writing)
    # Instructions can also mention proper nouns ("Describe the impact on Native Americans")
    words = text.split()
    if len(words) > 3:
        proper_nouns = sum(1 for w in words[1:] if w[0:1].isupper() and not w.isupper()
                          and w not in ('I', 'I\'m', 'I\'ll', 'I\'ve', 'I\'d'))
        if proper_nouns >= 2 and len(past_verbs) >= 1:
            score -= 2  # Proper nouns + past tense = strong student response signal
        elif proper_nouns >= 3 and not imperative_starts:
            score -= 1  # Many proper nouns without imperative start

    # Counter 4: Contains because/since/therefore (reasoning = student work)
    if re.search(r'\b(because|since|therefore|however|although|furthermore|moreover|consequently)\b', text_lower):
        score -= 2

    # Counter 5: Long text (>150 chars) with no meta-language is likely student content
    if len(text) > 150 and meta_count == 0 and not imperative_starts:
        score -= 2

    # Counter 6: Definition patterns ("X means...", "X is defined as...", "X refers to...")
    if re.search(r'\b(means|is defined as|refers to|is when|is the|is a |are the|are a )\b', text_lower):
        score -= 2

    # === DECISION ===
    # score >= 3 means instruction/prompt
    return score >= 3


def filter_questions_from_response(response_text: str) -> str:
    """
    Filter out question/prompt text from a response, keeping only actual answers.

    For example:
    Input: "What year was the Louisiana Purchase? 1803"
    Output: "1803"

    Input: "How did slavery affect daily life?"
    Output: "" (blank - student didn't answer)
    """
    if not response_text:
        return ""

    lines = response_text.strip().split('\n')
    filtered_lines = []

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # If line contains ?, check if there's an answer after it
        if '?' in line:
            parts = line.split('?')
            # Get everything after the last question mark
            after_question = parts[-1].strip() if len(parts) > 1 else ''
            if after_question and len(after_question) > 2 and not is_question_or_prompt(after_question):
                filtered_lines.append(after_question)
        elif not is_question_or_prompt(line):
            filtered_lines.append(line)
        else:
            # Line IS a prompt/instruction — but check if student content follows
            # the instruction on the same line (after a period)
            line_lower = line.lower()
            period_pos = line.find('.')
            if period_pos != -1 and period_pos < len(line) - 5:
                after_period = line[period_pos + 1:].strip()
                if after_period and len(after_period) > 10 and not is_question_or_prompt(after_period):
                    filtered_lines.append(after_period)

    return '\n'.join(filtered_lines).strip()


def _strip_template_lines(response: str, marker_text: str, template_text: str, is_short_answer: bool = False) -> str:
    """Remove template/prompt lines from extracted response by comparing with template.

    After the marker heading (e.g., 'SUMMARY'), the template may contain instruction
    text like 'Summarize the key events in 4-5 sentences.' which appears in the student
    doc but is NOT student work. This function strips those template lines.
    """
    if not template_text or not response:
        return response

    # Find the marker section in the template
    marker_lower = marker_text.lower().strip()
    template_lower = template_text.lower()
    marker_pos = template_lower.find(marker_lower)
    if marker_pos == -1:
        return response

    # Extract template content after this marker
    template_after = template_text[marker_pos + len(marker_text):]
    # Stop at the next likely section marker (all-caps word on its own line)
    template_section_lines = []
    for line in template_after.split('\n'):
        stripped = line.strip()
        if not stripped:
            continue
        # Stop if we hit another section marker (all-caps word, 3+ chars)
        if stripped.isupper() and len(stripped) >= 3 and stripped.isalpha():
            break
        template_section_lines.append(stripped)

    if not template_section_lines:
        return response

    # Build set of normalized template lines (lowered, stripped of underscores/whitespace)
    template_line_set = set()
    template_word_sets = []  # For fuzzy matching
    for tl in template_section_lines:
        cleaned = re.sub(r'[_\s]+', ' ', tl).strip().lower()
        if len(cleaned) >= 10:  # Only match substantial lines
            template_line_set.add(cleaned)
            # Also store word sets for fuzzy overlap matching
            words = set(re.findall(r'[a-z]{3,}', cleaned))
            if len(words) >= 3:
                template_word_sets.append(words)

    # Filter response lines: remove template text, preserving student answers
    # that may appear on the same line after a template prompt
    response_lines = response.split('\n')
    filtered = []
    for line in response_lines:
        line_cleaned = re.sub(r'[_\s]+', ' ', line).strip().lower()
        is_template = False
        partial_strip = None  # Student content remaining after template prefix

        # Exact / substring match
        for tl in template_line_set:
            if line_cleaned == tl or (len(line_cleaned) >= 15 and line_cleaned in tl):
                # Whole line is template or line is inside a longer template line
                is_template = True
                break
            if len(tl) >= 15 and tl in line_cleaned:
                # Template text is a prefix/substring of this line — the student
                # may have typed their answer on the same line after the prompt.
                # Strip the template portion and keep the rest.
                tl_pos = line_cleaned.find(tl)
                remainder = line_cleaned[tl_pos + len(tl):].strip()
                # Only keep if there's meaningful student content after the template
                min_remainder = 2 if is_short_answer else 10
                if len(remainder) >= min_remainder:
                    # Find the same position in the original (non-lowered) line
                    orig_cleaned = re.sub(r'[_\s]+', ' ', line).strip()
                    partial_strip = orig_cleaned[tl_pos + len(tl):].strip()
                else:
                    is_template = True
                break

        # Fuzzy match: high word overlap with any template line (catches rephrased prompts)
        if not is_template and partial_strip is None and len(line_cleaned) >= 15:
            line_words = set(re.findall(r'[a-z]{3,}', line_cleaned))
            if len(line_words) >= 3:
                for tw_set in template_word_sets:
                    overlap = len(line_words & tw_set)
                    # If 60%+ of the response line's words appear in a template line,
                    # it's likely template text (possibly reworded)
                    if overlap >= max(3, len(line_words) * 0.6):
                        is_template = True
                        break

        if partial_strip is not None:
            filtered.append(partial_strip)
        elif not is_template:
            filtered.append(line)

    result = '\n'.join(filtered).strip()
    if result != response.strip():
        stripped_count = len(response_lines) - len(filtered)
        print(f"      🧹 Stripped {stripped_count} template line(s) from {marker_text[:30]} response")
    return result


def strip_emojis(text: str) -> str:
    """Remove emojis and special unicode characters from text."""
    import re
    # Remove emojis (unicode ranges for emojis)
    emoji_pattern = re.compile("["
        u"\U0001F600-\U0001F64F"  # emoticons
        u"\U0001F300-\U0001F5FF"  # symbols & pictographs
        u"\U0001F680-\U0001F6FF"  # transport & map symbols
        u"\U0001F700-\U0001F77F"  # alchemical symbols
        u"\U0001F780-\U0001F7FF"  # Geometric Shapes
        u"\U0001F800-\U0001F8FF"  # Supplemental Arrows-C
        u"\U0001F900-\U0001F9FF"  # Supplemental Symbols and Pictographs
        u"\U0001FA00-\U0001FA6F"  # Chess Symbols
        u"\U0001FA70-\U0001FAFF"  # Symbols and Pictographs Extended-A
        u"\U00002702-\U000027B0"  # Dingbats
        u"\U000024C2-\U0001F251"
        "]+", flags=re.UNICODE)
    return emoji_pattern.sub('', text).strip()


def fuzzy_find_marker(doc_text: str, marker: str, threshold: float = 0.7) -> int:
    """
    Find a marker in document using fuzzy matching.
    Handles cases where students slightly edit formatting.

    Returns position if found with confidence >= threshold, else -1.
    Processing time: ~1-2ms per marker.
    """
    doc_lower = doc_text.lower()
    marker_lower = marker.lower().strip()

    # First try exact match
    exact_pos = doc_lower.find(marker_lower)
    if exact_pos != -1:
        return exact_pos

    # Try matching with emojis stripped from both sides
    marker_no_emoji = strip_emojis(marker_lower)
    doc_no_emoji = strip_emojis(doc_lower)
    if marker_no_emoji:
        emoji_pos = doc_no_emoji.find(marker_no_emoji)
        if emoji_pos != -1:
            # Find the actual position in original doc by searching near this location
            # Account for removed emojis by doing a fuzzy position find
            search_start = max(0, emoji_pos - 20)
            search_chunk = doc_lower[search_start:search_start + len(marker_lower) + 50]
            actual_pos = search_chunk.find(marker_no_emoji)
            if actual_pos != -1:
                return search_start + actual_pos
            return emoji_pos  # Fallback to stripped position

    # Try matching without extra whitespace/punctuation
    marker_normalized = ' '.join(marker_lower.split())  # Collapse whitespace
    marker_words = marker_normalized.split()

    # For single-word markers (like "Summary", "Vocabulary"), try word boundary search
    if len(marker_words) == 1:
        # Look for the word with word boundaries
        pattern = r'\b' + re.escape(marker_words[0]) + r'\b'
        match = re.search(pattern, doc_lower)
        if match:
            return match.start()
        # Also try with emoji-stripped doc
        match = re.search(pattern, doc_no_emoji)
        if match:
            return match.start()
        return -1

    # Try finding first 3-4 significant words (skip common words)
    skip_words = {'the', 'a', 'an', 'of', 'to', 'in', 'for', 'on', 'with', 'is', 'are'}
    sig_words = [w for w in marker_words if w not in skip_words and len(w) > 2][:4]

    if len(sig_words) >= 2:
        # Build a pattern: first word.*second word.*third word (with flexibility)
        pattern_parts = [re.escape(w) for w in sig_words]
        pattern = r'.{0,20}'.join(pattern_parts)

        try:
            match = re.search(pattern, doc_lower)
            if match:
                return match.start()
        except re.error:
            pass

    # Try first significant phrase (first 5 words)
    first_phrase = ' '.join(marker_words[:5])
    # Allow for inserted spaces/chars
    flexible_pattern = r'\s*'.join([re.escape(c) for c in first_phrase if c.strip()])
    try:
        match = re.search(flexible_pattern[:100], doc_lower)  # Limit pattern length
        if match:
            return match.start()
    except re.error:
        pass

    return -1


def extract_fitb_by_template_comparison(student_text: str, template_text: str) -> list:
    """
    Extract fill-in-the-blank answers from student submission.

    For FITB assignments, the student's completed text IS the answer.
    This extracts content flexibly - works with:
    - Numbered lines (1., 2., 3.)
    - Lettered lines (a., b., c.)
    - Timestamped lines (1. (0:00) ...)
    - Plain paragraphs
    - Any format where blanks were filled inline

    Returns list of {"question": context, "answer": filled_text, "type": "fill_in_blank"}
    """
    if not student_text:
        return []

    extracted = []
    student_lines = student_text.strip().split('\n')

    # Skip patterns - these are template/header lines, not student answers
    skip_patterns = [
        r'^name\s*[:_]', r'^date\s*[:_]', r'^period\s*[:_]', r'^class\s*[:_]',
        r'^directions?\s*:', r'^instructions?\s*:', r'^fill in', r'^watch the video',
        r'^answer each', r'^complete the', r'^use evidence', r'^_{5,}$'
    ]

    item_number = 0
    for student_line in student_lines:
        student_line = student_line.strip()
        if not student_line or len(student_line) < 10:
            continue

        # Skip header/template lines
        line_lower = student_line.lower()
        if any(re.match(p, line_lower) for p in skip_patterns):
            continue

        # Try to match various formats:
        # 1. Numbered: "1.", "1)", "1:"
        # 2. Lettered: "a.", "a)", "A."
        # 3. With timestamp: "1. (0:00-0:25)"
        # 4. Plain content lines

        content = None
        label = None

        # Numbered with optional timestamp (handle bullet prefixes like •, *, -)
        num_match = re.match(r'^[•\*\-\u2022\u2023\u25E6\u2043\u2219\s]*(\d+)[\.\)\:]\s*(\([^)]+\))?\s*(.+)', student_line)
        if num_match:
            label = f"Item {num_match.group(1)}"
            if num_match.group(2):
                label += f" {num_match.group(2)}"
            content = num_match.group(3).strip()

        # Lettered (handle bullet prefixes)
        if not content:
            letter_match = re.match(r'^[•\*\-\u2022\u2023\u25E6\u2043\u2219\s]*([a-zA-Z])[\.\)\:]\s*(.+)', student_line)
            if letter_match:
                label = f"Item {letter_match.group(1).upper()}"
                content = letter_match.group(2).strip()

        # Plain line with substance (not just a title)
        if not content and len(student_line) > 20:
            # Check if it looks like content (has filled blanks or is a complete sentence)
            if '_' in student_line or student_line.endswith('.') or ',' in student_line:
                item_number += 1
                label = f"Response {item_number}"
                content = student_line

        if content and len(content) >= 10:
            # Don't add if it's just the assignment title
            if 'fill-in-the-blank' in content.lower() or 'activity' in content.lower()[:20]:
                continue

            extracted.append({
                "question": label or f"Response {len(extracted) + 1}",
                "answer": content,
                "type": "fill_in_blank_sentence"
            })

    # Also extract any text wrapped in underscores _answer_ or __answer__
    underscore_matches = re.findall(r'_+([^_\n]{1,80})_+', student_text)
    for match in underscore_matches:
        answer = match.strip()
        if answer and len(answer) > 0:
            # Skip common template words
            if answer.lower() not in ['name', 'date', 'period', 'class', 'blank', 'answer', 'your answer']:
                # Avoid duplicates
                already_exists = any(answer.lower() in e.get('answer', '').lower() for e in extracted)
                if not already_exists:
                    extracted.append({
                        "question": "Fill-in-blank",
                        "answer": answer,
                        "type": "fill_in_blank"
                    })

    return extracted


def parse_numbered_questions(text: str) -> list:
    """
    Parse numbered questions and their answers from a text block.

    Handles formats like:
    - "1. Question text? Answer text"
    - "1. Question text?\n   Answer text"
    - "1) Question text? Answer text"

    Returns list of {"question": "...", "answer": "..."} dicts.
    Returns empty list if no numbered questions found.
    """
    if not text or len(text) < 20:
        return []

    # Pattern to find numbered questions: "1." or "1)" at start of line or after newline
    # Captures: number, question text (up to ?), and everything after until next number
    number_pattern = r'(?:^|\n)\s*(\d+)[.\)]\s*'

    # Find all numbered items
    matches = list(re.finditer(number_pattern, text))

    if len(matches) < 2:
        # Not enough numbered items to indicate a numbered list
        return []

    # Check if numbers are sequential (1, 2, 3, ...)
    numbers = [int(m.group(1)) for m in matches]
    if numbers[0] != 1:
        return []  # Doesn't start with 1

    # Allow some gaps but should be roughly sequential
    is_sequential = all(numbers[i] <= numbers[i-1] + 2 for i in range(1, len(numbers)))
    if not is_sequential:
        return []

    results = []

    for i, match in enumerate(matches):
        q_num = match.group(1)
        q_start = match.end()  # Start of question text

        # Find end of this question's content (start of next number or end of text)
        if i + 1 < len(matches):
            q_end = matches[i + 1].start()
        else:
            q_end = len(text)

        content = text[q_start:q_end].strip()

        if not content:
            results.append({
                "question": f"Question {q_num}",
                "answer": "",
                "is_blank": True
            })
            continue

        # Split into question and answer
        # The question ends at '?' and answer is everything after
        question_mark_pos = content.find('?')

        if question_mark_pos != -1:
            question_text = content[:question_mark_pos + 1].strip()
            answer_text = content[question_mark_pos + 1:].strip()
        else:
            # No question mark - might be a statement or instruction
            # Look for newline as separator
            newline_pos = content.find('\n')
            if newline_pos != -1:
                question_text = content[:newline_pos].strip()
                answer_text = content[newline_pos:].strip()
            else:
                # Single line with no clear separator
                question_text = content
                answer_text = ""

        # Clean up answer - remove leading/trailing whitespace and blank lines
        answer_lines = [line.strip() for line in answer_text.split('\n') if line.strip()]

        answer_text = '\n'.join(answer_lines)

        # If there's any non-empty text after the question, it's an answer.
        # Don't try to second-guess whether it "looks like" a question —
        # content between two numbered items after the ? is the student's response.
        if answer_text and len(answer_text.strip()) > 2:
            results.append({
                "question": f"{q_num}. {question_text}",
                "answer": answer_text,
                "is_blank": False
            })
        else:
            results.append({
                "question": f"{q_num}. {question_text}",
                "answer": "",
                "is_blank": True
            })

    return results


def parse_vocab_terms(text: str) -> list:
    """
    Parse vocabulary Term: definition pairs from a text block.

    Handles formats like:
    - "Assimilate: to adopt the customs of another group"
    - "Indian Removal Act: 1830 law signed by Jackson"
    - "Term: _______________" (blank)
    - "Term:" followed by definition on next line

    Returns list of {"term": "...", "answer": "...", "is_blank": bool} dicts.
    Returns empty list if no vocab terms detected.
    """
    if not text or len(text) < 10:
        return []

    lines = text.strip().split('\n')
    results = []

    # Skip header keywords that aren't vocab terms
    skip_keywords = {'vocabulary', 'questions', 'summary', 'notes', 'section',
                     'directions', 'instructions', 'response', 'name', 'date',
                     'period', 'write your', 'use bullets', 'define the',
                     'in your own words', 'total', 'points'}

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1

        if not line or line.replace('_', '').replace(' ', '').replace('-', '') == '':
            continue

        # Check for "Term: definition" or "Term:" pattern
        if ':' not in line:
            continue

        parts = line.split(':', 1)
        term = parts[0].strip()
        defn = parts[1].strip() if len(parts) > 1 else ''

        # Validate it looks like a vocab term (not a header/instruction)
        if not term or len(term) < 2 or len(term) > 60:
            continue
        if len(term.split()) > 5:
            continue

        term_lower = term.lower()
        if any(kw in term_lower for kw in skip_keywords):
            continue

        # Skip if term starts with a number (likely a numbered question, not vocab)
        if re.match(r'^\d+[\.\)]\s', term):
            continue

        # Check if definition is blank (underscores, empty, dashes)
        defn_clean = re.sub(r'[_\-\s\.]+', '', defn)
        if len(defn_clean) < 3:
            # Check next line(s) for definition
            found_defn = False
            for look in range(i, min(i + 2, len(lines))):
                next_line = lines[look].strip()
                next_clean = re.sub(r'[_\-\s\.]+', '', next_line)
                if next_clean and len(next_clean) >= 3 and ':' not in next_line:
                    # Next line has content without a colon — it's the definition
                    defn = next_line
                    defn_clean = next_clean
                    found_defn = True
                    i = look + 1  # Skip past the definition line
                    break
                elif next_clean and ':' in next_line:
                    break  # Next line is another term — this one is blank

            if not found_defn or len(defn_clean) < 3:
                results.append({"term": term, "answer": "", "is_blank": True})
                continue

        results.append({"term": term, "answer": defn, "is_blank": False})

    # Only return results if we found at least 2 vocab terms (avoid false positives)
    if len(results) < 2:
        return []

    return results


def extract_student_responses(document_text: str, custom_markers: list = None, exclude_markers: list = None, template_text: str = None) -> dict:
    """
    Extract student responses from document using customMarkers from Builder.

    Args:
        document_text: The student's submitted document text
        custom_markers: List of markers to look for
        exclude_markers: List of markers to exclude from grading
        template_text: Original assignment template (for fill-in-the-blank comparison)

    APPROACH (with fallbacks):
    1. TEMPLATE COMPARISON: If template provided, compare to find filled blanks
    2. EXACT MATCH: Find markers exactly in document
    2. FUZZY MATCH: If exact fails, try fuzzy matching (~1-2ms)
    3. PATTERN MATCH: If all else fails, use regex patterns

    Args:
        document_text: The full document text
        custom_markers: List of marker strings from Builder (the template/prompt text)
        exclude_markers: List of section names to NOT grade (e.g., "Notes Section")

    Returns:
        {
            "extracted_responses": [{"question": "...", "answer": "...", "type": "..."}, ...],
            "blank_questions": ["question text", ...],
            "total_questions": int,
            "answered_questions": int,
            "extraction_summary": "string describing what was found",
            "excluded_sections": ["section name", ...] - sections that were skipped
        }
    """
    import re

    extracted = []
    blank_questions = []
    doc_lower = document_text.lower()

    # Pre-compute normalized doc text (en-dash → hyphen, smart quotes → regular)
    def normalize_for_search(text):
        return text.replace('\u2013', '-').replace('\u2014', '-').replace('\u2018', "'").replace('\u2019', "'").replace('\u201c', '"').replace('\u201d', '"')
    doc_lower_normalized = normalize_for_search(doc_lower)

    # PRIORITY 0: Fill-in-the-blank detection
    # Only trigger FITB if the document/title explicitly mentions fill-in-the-blank
    # Don't trigger just because there are underscores (Cornell Notes has blank lines too)
    has_fitb_keyword = 'fill-in' in doc_lower or 'fill in the blank' in doc_lower or 'fillintheblank' in doc_lower.replace(' ', '').replace('-', '')
    has_timestamps = bool(re.search(r'\d+\.\s*\(\d+:\d+', document_text))  # Has timestamps like "1. (0:00)"

    # Only use underscores as indicator if FITB keyword is also present
    is_fitb = has_fitb_keyword or has_timestamps

    if is_fitb:
        print(f"  📝 Detected fill-in-the-blank format - using FITB extraction")
        fitb_results = extract_fitb_by_template_comparison(document_text, template_text)
        if fitb_results:
            # Filter out excluded sections and blank/template content
            exclude_lower = [em.lower().strip() for em in exclude_markers] if exclude_markers else []
            filtered_fitb = []
            for resp in fitb_results:
                question = resp.get('question', '').lower()
                answer = resp.get('answer', '')
                answer_lower = answer.lower()

                # Skip blank underscore lines (e.g., "• _______________")
                answer_stripped = re.sub(r'[_•\-\s]', '', answer)
                if len(answer_stripped) < 5:
                    continue

                # Skip template instruction text
                if any(kw in answer_lower for kw in ['use bullets', 'write important', 'add more notes', 'drawing']):
                    continue

                # Check if this response is from an excluded section
                # Check both directions: exclude text in answer, AND answer in exclude text
                is_excluded = False
                for em in exclude_lower:
                    # Original: exclude marker contained in answer (short marker, long answer)
                    if em in answer_lower:
                        is_excluded = True
                        break
                    # Reverse: answer contained in exclude marker (long marker, short answer)
                    # Only for answers with enough length to avoid false positives
                    if len(answer_lower) >= 15 and answer_lower in em:
                        is_excluded = True
                        break
                if not is_excluded:
                    filtered_fitb.append(resp)
            fitb_results = filtered_fitb
            extracted.extend(fitb_results)
            print(f"      Found {len(fitb_results)} responses via FITB extraction")

    # PRIORITY 1: Use customMarkers from Builder (most reliable)
    # Markers can be:
    #   - String: "Summary (Bottom Section)" - extracts until next marker
    #   - Object: {"start": "Summary", "end": "📖"} - extracts until end marker
    if custom_markers and len(custom_markers) > 0:
        # Find positions of all markers in the document (exact + fuzzy)
        marker_positions = []
        exact_matches = 0
        fuzzy_matches = 0

        for marker in custom_markers:
            # Handle both string and object markers
            if isinstance(marker, dict):
                marker_clean = marker.get('start', '').strip()
                end_marker = marker.get('end', '').strip()
            else:
                marker_clean = str(marker).strip()
                end_marker = None

            if not marker_clean:
                continue

            # Split marker into lines for multi-line matching strategies
            marker_lines = [l.strip() for l in marker_clean.split('\n') if l.strip()]
            first_line = marker_lines[0] if marker_lines else marker_clean
            marker_lower = marker_clean.lower()

            pos = -1
            match_type = None
            content_end_pos = -1  # Where content actually starts (after all marker/prompt lines)

            # Strategy 1: Full exact match (ideal — whole multi-line marker found as-is)
            # Prefer line-start matches to avoid substring hits (e.g., "NOTES" in "CORNELL NOTES")
            # A "line-start" match means the marker is at position 0, or preceded by a newline
            # (possibly with emoji/whitespace between the newline and the marker text).
            pos = -1
            search_start = 0
            while True:
                candidate = doc_lower.find(marker_lower, search_start)
                if candidate == -1:
                    break
                # Check if this is a standalone section header (at line start)
                is_line_start = (candidate == 0)
                if not is_line_start and candidate > 0:
                    # Look back from candidate to the previous newline — if only
                    # whitespace/emoji/special chars between newline and marker, it's line-start
                    lookback = doc_lower[max(0, candidate - 10):candidate]
                    newline_pos = lookback.rfind('\n')
                    if newline_pos != -1:
                        between = lookback[newline_pos + 1:]
                        # Only whitespace, emoji, or common prefix chars (star, bullet, dash)
                        is_line_start = all(
                            c in ' \t\u2b50\U0001f31f\U0001f4d3\U0001f4dd\U0001f331\U0001f4d6\U0001f4da\U0001f50d\u2022\u2013\u2014-*'
                            or ord(c) > 0x2600  # emoji/symbol unicode ranges
                            for c in between
                        )
                if is_line_start:
                    pos = candidate
                    break
                # Otherwise keep searching for a line-start match
                if pos == -1:
                    pos = candidate  # fallback to first occurrence if no line-start match
                search_start = candidate + 1
            if pos != -1:
                match_type = 'exact'
                content_end_pos = pos + len(marker_clean)

            # Strategy 2: First-line exact match + forward scan for subsequent lines
            if pos == -1:
                first_line_lower = first_line.lower()
                matched_len = len(first_line)  # Track actual matched text length

                # Try with original text
                pos = doc_lower.find(first_line_lower)

                # Try with normalized dashes/quotes
                if pos == -1:
                    first_normalized = normalize_for_search(first_line_lower)
                    pos = doc_lower_normalized.find(first_normalized)

                # Try without emojis (matched_len changes since emoji chars aren't in doc)
                if pos == -1:
                    first_no_emoji = strip_emojis(first_line_lower).strip()
                    if first_no_emoji and len(first_no_emoji) >= 3:
                        pos = doc_lower.find(first_no_emoji)
                        if pos == -1:
                            # Try normalized + emoji-stripped
                            pos = doc_lower_normalized.find(normalize_for_search(first_no_emoji))
                        if pos != -1:
                            matched_len = len(first_no_emoji)  # Use stripped length

                if pos != -1:
                    match_type = 'first_line'
                    # Find the actual end of the heading line in the document
                    heading_newline = doc_lower.find('\n', pos)
                    if heading_newline != -1 and heading_newline < pos + matched_len + 50:
                        content_end_pos = heading_newline + 1  # Start of next line
                    else:
                        content_end_pos = pos + matched_len

                    # Scan forward for instruction/prompt lines that come RIGHT AFTER
                    # the heading — skip past them so content starts at student answers.
                    # STOP scanning at numbered questions (1., 2.) since those contain
                    # individual Q&A pairs that parse_numbered_questions will handle.
                    if len(marker_lines) > 1:
                        last_advance = content_end_pos
                        search_to = min(last_advance + len(marker_clean) + 500, len(doc_lower))

                        for mline in marker_lines[1:]:
                            mline_lower = mline.lower().strip()

                            # Stop at numbered questions — don't skip past them
                            if re.match(r'^\d+[\.\)]\s', mline_lower):
                                break

                            mline_normalized = normalize_for_search(mline_lower)

                            # Try finding this line near the current position
                            line_pos = doc_lower.find(mline_lower, last_advance, search_to)
                            if line_pos == -1:
                                line_pos = doc_lower_normalized.find(mline_normalized, last_advance, search_to)
                            if line_pos == -1:
                                mline_no_emoji = strip_emojis(mline_lower).strip()
                                if mline_no_emoji and len(mline_no_emoji) >= 5:
                                    line_pos = doc_lower.find(mline_no_emoji, last_advance, search_to)
                            if line_pos == -1:
                                # Try first few words (handles table | format)
                                mline_words = mline_lower.split()[:3]
                                if len(mline_words) >= 2:
                                    short_search = ' '.join(mline_words)
                                    if len(short_search) >= 5:
                                        line_pos = doc_lower.find(short_search, last_advance, search_to)

                            if line_pos != -1:
                                # Only advance if this line is CLOSE (no big gap with student content)
                                if line_pos - last_advance > 200:
                                    break  # Gap too large — student content in between
                                newline_after = doc_lower.find('\n', line_pos)
                                if newline_after != -1 and newline_after < search_to:
                                    content_end_pos = max(content_end_pos, newline_after + 1)
                                    last_advance = newline_after + 1
                                else:
                                    content_end_pos = max(content_end_pos, line_pos + len(mline_lower))
                                    last_advance = content_end_pos
                            else:
                                break  # Line not found — stop scanning

            # Strategy 3: Fuzzy match (fallback)
            if pos == -1:
                pos = fuzzy_find_marker(document_text, marker_clean)
                if pos != -1:
                    match_type = 'fuzzy'
                    content_end_pos = pos + min(len(marker_clean), 100)

            if pos != -1:
                if match_type == 'exact':
                    exact_matches += 1
                elif match_type == 'first_line':
                    exact_matches += 1  # First-line is reliable enough to count as exact
                else:
                    fuzzy_matches += 1

                marker_positions.append({
                    'marker': marker_clean,
                    'start': pos,
                    'end': content_end_pos,
                    'end_marker': end_marker,
                    'match_type': match_type,
                    'section_type': marker.get('type', 'written') if isinstance(marker, dict) else 'written',
                })

        # Sort by position in document
        marker_positions.sort(key=lambda x: x['start'])

        # Track excluded sections
        excluded_sections = []

        # Normalize exclude markers for comparison
        exclude_markers_normalized = []
        if exclude_markers:
            for em in exclude_markers:
                exclude_markers_normalized.append(em.lower().strip())

        # Track markers that were NOT found in the document at all
        # Build a set of found marker names using multiple representations
        # (full text, first line, emoji-stripped) so multi-line or emoji markers
        # are correctly matched against the config
        found_marker_names = set()
        for mp in marker_positions:
            m = mp['marker'].lower()
            found_marker_names.add(m)
            # Also add first line (for multi-line markers found by first-line strategy)
            first_line_found = m.split('\n')[0].strip()
            if first_line_found:
                found_marker_names.add(first_line_found)
            # Also add emoji-stripped versions
            m_no_emoji = strip_emojis(m).strip()
            if m_no_emoji:
                found_marker_names.add(m_no_emoji)
            first_no_emoji = strip_emojis(first_line_found).strip()
            if first_no_emoji:
                found_marker_names.add(first_no_emoji)

        missing_sections = []
        for marker in custom_markers:
            if isinstance(marker, dict):
                marker_name = marker.get('start', '').strip()
            else:
                marker_name = str(marker).strip()
            if not marker_name:
                continue
            # Check multiple representations of the marker name
            marker_lower = marker_name.lower()
            marker_first_line = marker_lower.split('\n')[0].strip()
            marker_no_emoji = strip_emojis(marker_lower).strip()
            marker_first_no_emoji = strip_emojis(marker_first_line).strip()
            is_found = (marker_lower in found_marker_names
                       or marker_first_line in found_marker_names
                       or marker_no_emoji in found_marker_names
                       or marker_first_no_emoji in found_marker_names)
            if not is_found:
                # Don't flag excluded sections as missing
                is_excluded = any(em in marker_name.lower() or marker_name.lower() in em
                                for em in exclude_markers_normalized)
                if not is_excluded:
                    missing_sections.append(marker_name)

        # Extract response after each marker
        for i, mp in enumerate(marker_positions):
            marker_text = mp['marker']
            content_start = mp['end']
            end_marker = mp.get('end_marker')
            section_type = mp.get('section_type', 'written')
            # Detect short-answer from explicit type OR from marker name containing FITB keywords
            marker_name_hint = mp['marker'].lower()
            is_short_answer = (
                section_type in ('fill-blank', 'fill_in_blank', 'vocabulary', 'matching')
                or 'fill-in' in marker_name_hint
                or 'fill in' in marker_name_hint
                or 'fitb' in marker_name_hint
                or 'blanks' in marker_name_hint
            )

            # Check if this marker should be excluded from grading
            marker_lower = marker_text.lower().strip()
            marker_first_line = marker_lower.split('\n')[0].strip()
            marker_first_no_emoji = strip_emojis(marker_first_line).strip()
            is_excluded = False
            for em in exclude_markers_normalized:
                em_first_line = em.split('\n')[0].strip()
                em_first_no_emoji = strip_emojis(em_first_line).strip()
                # Compare full text, first lines, and emoji-stripped first lines
                if (em in marker_lower or marker_lower in em
                    or em_first_line in marker_first_line or marker_first_line in em_first_line
                    or (em_first_no_emoji and marker_first_no_emoji
                        and (em_first_no_emoji in marker_first_no_emoji or marker_first_no_emoji in em_first_no_emoji))):
                    is_excluded = True
                    excluded_sections.append(marker_text[:80])
                    break

            if is_excluded:
                continue  # Skip this section - don't grade it

            # Determine where content ends (in priority order):
            # 1. Explicit end marker (if defined)
            # 2. Next start marker
            # 3. Section delimiters
            # 4. Document end (capped at 1500 chars)

            if end_marker:
                # Use explicit end marker with fallback for suspiciously short content
                MIN_CONTENT_LENGTH = 50  # If content is shorter, end marker may have been found too early

                end_pos = doc_lower.find(end_marker.lower(), content_start)
                if end_pos != -1:
                    # Check if content is suspiciously short
                    potential_content = document_text[content_start:end_pos].strip()

                    # If too short, look for NEXT occurrence of end marker
                    search_pos = end_pos + 1
                    while len(potential_content) < MIN_CONTENT_LENGTH and search_pos < len(document_text):
                        next_end_pos = doc_lower.find(end_marker.lower(), search_pos)
                        if next_end_pos == -1:
                            break  # No more occurrences
                        potential_content = document_text[content_start:next_end_pos].strip()
                        if len(potential_content) >= MIN_CONTENT_LENGTH:
                            end_pos = next_end_pos
                            break
                        search_pos = next_end_pos + 1

                    content_end = end_pos
                else:
                    # End marker not found, fall back to next marker or cap
                    if i + 1 < len(marker_positions):
                        content_end = marker_positions[i + 1]['start']
                    else:
                        content_end = min(content_start + 1500, len(document_text))
            elif i + 1 < len(marker_positions):
                # End at next marker
                content_end = marker_positions[i + 1]['start']
            else:
                # Last marker - cap at 1500 chars
                content_end = min(content_start + 1500, len(document_text))

            # Also stop at known section delimiters (reading material, etc.)
            # BUT skip this if we already have an explicit end marker (to avoid re-truncating)
            # NOTE: Removed '___' from delimiters - it conflicts with fill-in-the-blank notation
            if not end_marker:
                section_delimiters = ['📖', '📚', '🔍', '--- ', '***', '===']
                for delim in section_delimiters:
                    delim_pos = document_text.find(delim, content_start, content_end)
                    if delim_pos != -1 and delim_pos > content_start:
                        content_end = delim_pos

            # Also stop at exclude marker boundaries — don't capture excluded
            # section content as part of an adjacent graded section's response
            if exclude_markers_normalized:
                for em in exclude_markers_normalized:
                    # Try multiple search strategies (emoji/dash/encoding differences)
                    em_lines = [l.strip() for l in em.split('\n') if l.strip()]
                    search_terms = set()
                    for el in em_lines[:3]:  # First 3 lines
                        search_terms.add(el)
                        # Without leading emoji/special chars
                        stripped = re.sub(r'^[^a-zA-Z]*', '', el).strip()
                        if stripped and len(stripped) >= 5:
                            search_terms.add(stripped)
                        # Normalized dashes/quotes
                        normalized = el.replace('\u2013', '-').replace('\u2014', '-').replace('\u2018', "'").replace('\u2019', "'")
                        if normalized != el:
                            search_terms.add(normalized)

                    found_boundary = False
                    for em_search in search_terms:
                        if len(em_search) >= 5 and not found_boundary:
                            em_pos = doc_lower.find(em_search, content_start, content_end)
                            if em_pos == -1:
                                # Also try in normalized doc text
                                em_pos = doc_lower_normalized.find(em_search, content_start, content_end)
                            if em_pos != -1 and em_pos > content_start:
                                content_end = min(content_end, em_pos)
                                found_boundary = True

            # Extract the response
            response = document_text[content_start:content_end].strip()

            # Clean up: remove leading colons, newlines
            response = re.sub(r'^[:\s\n]+', '', response).strip()

            # Check if response contains numbered questions (1., 2., 3., etc.)
            # If so, parse them individually instead of as one blob
            numbered_items = parse_numbered_questions(response)
            if numbered_items:
                print(f"      📝 Found {len(numbered_items)} numbered questions in section")
                for item in numbered_items:
                    answer = item.get("answer", "")

                    # Clean template artifacts from numbered question answers:
                    # "(25 pts)", "Response: ___", "_____" lines are template text, not student content
                    if answer:
                        cleaned_lines = []
                        for aline in answer.split('\n'):
                            stripped = aline.strip()
                            # Skip point value markers like "(20 pts)" or "(25 pts)"
                            if re.match(r'^\(\s*\d+\s*(?:pts?|points?)\s*\)\s*$', stripped, re.IGNORECASE):
                                continue
                            # Skip "Response:" with only underscores after it
                            if re.match(r'^Response\s*:\s*[_\s]*$', stripped, re.IGNORECASE):
                                continue
                            # Skip lines that are only underscores/dashes/spaces
                            if re.match(r'^[_\-\s]+$', stripped) or not stripped:
                                continue
                            # Skip "Response:" prefix but keep any actual text after it
                            resp_match = re.match(r'^Response\s*:\s*(.+)', stripped, re.IGNORECASE)
                            if resp_match:
                                actual = resp_match.group(1).strip()
                                if actual and not re.match(r'^[_\-\s]+$', actual):
                                    cleaned_lines.append(actual)
                                continue
                            cleaned_lines.append(stripped)
                        answer = '\n'.join(cleaned_lines).strip()

                    if not answer or len(answer) < 3:
                        blank_questions.append(item.get("question", "Unknown"))
                    else:
                        extracted.append({
                            "question": item["question"],
                            "answer": answer[:10000],
                            "type": "numbered_question"
                        })
                continue  # Skip normal processing for this marker

            # Check if this is a VOCABULARY section with Term: definition pairs
            # Parse each vocab term individually for better grading granularity
            marker_name_lower = marker_text.lower().split('\n')[0].strip()
            is_vocab_section = 'vocab' in marker_name_lower
            if is_vocab_section:
                vocab_items = parse_vocab_terms(response)
                if vocab_items:
                    print(f"      📖 Found {len(vocab_items)} vocab terms in section")
                    for vitem in vocab_items:
                        if vitem.get("is_blank"):
                            blank_questions.append(vitem["term"] + " (no definition)")
                        else:
                            extracted.append({
                                "question": vitem["term"],
                                "answer": vitem["answer"][:10000],
                                "type": "vocab_term"
                            })
                    continue  # Skip normal blob processing

            # Get a short label for the question (first 50 chars of marker)
            question_label = marker_text[:80] + '...' if len(marker_text) > 80 else marker_text
            # Clean up newlines in label
            question_label = ' '.join(question_label.split())

            # Check if response is actually blank (just whitespace, or only template/boilerplate)
            is_blank = False
            if not response or len(response) <= 5:
                is_blank = True
            else:
                # Filter out common template patterns that aren't student answers
                response_clean = response.strip()

                # Remove lines that start with instructions/prompts
                lines = [l.strip() for l in response_clean.split('\n') if l.strip()]
                # Filter out blank placeholder lines, track blank vocab terms
                student_lines = []
                for line in lines:
                    # Skip lines that are just underscores/spaces (blank placeholders)
                    if line.replace('_', '').replace(' ', '') == '':
                        continue
                    # Check if line has 3+ consecutive underscores (fill-in-the-blank format)
                    if re.search(r'_{3,}', line):
                        # Strip all underscores — if meaningful text remains, keep it
                        text_only = re.sub(r'_+', ' ', line).strip()
                        # Remove the term/label before colon too for checking
                        if ':' in text_only:
                            after_colon = text_only.split(':', 1)[1].strip()
                        else:
                            after_colon = text_only
                        min_answer_len = 1 if is_short_answer else 3
                        if len(after_colon) < min_answer_len:
                            # Truly blank — if it looks like a vocab term (Term: ___), track it
                            if ':' in text_only:
                                term = text_only.split(':', 1)[0].strip()
                                if term and len(term.split()) <= 4:
                                    blank_questions.append(f"{term} (no definition)")
                            continue
                        # Otherwise student wrote something between/around the underscores — keep it
                    student_lines.append(line)

                # If no student content remains, it's blank
                min_content_len = 3 if is_short_answer else 10
                if not student_lines or sum(len(l) for l in student_lines) < min_content_len:
                    is_blank = True
                else:
                    # Use filtered response
                    response = '\n'.join(student_lines)

            # CRITICAL: Strip template lines from response using original template
            # This removes instruction/prompt text that appears between the marker
            # heading and the student's actual response (e.g., "Summarize the key events...")
            # SKIP for FITB sections: the template lines ARE the questions — the student's
            # answer is the filled-in version, so stripping would remove their work.
            if not is_blank and template_text and not is_short_answer:
                response = _strip_template_lines(response, marker_text, template_text, is_short_answer=is_short_answer)

            # CRITICAL: Filter out questions/prompts from response
            # This prevents detecting "What was the cause?" as a student answer
            # SKIP for FITB: the prompts are the fill-in sentences themselves
            if not is_blank and not is_short_answer:
                response = filter_questions_from_response(response)
                if not response or len(response.strip()) < 3:
                    is_blank = True

            # Additional blank check: if remaining response is very short and looks like
            # template fragments (sub-prompts, topic lists, page refs), mark as blank.
            # Real student answers for questions are typically 20+ characters.
            if not is_blank and response:
                resp_clean = response.strip()
                # Remove common template artifacts: page refs, point values, "Response:" labels
                resp_stripped = re.sub(r'\((?:pp?\.?\s*\d+[\-–]\d+|\d+\s*(?:pts?|points?))\)', '', resp_clean)
                resp_stripped = re.sub(r'(?i)^response\s*:\s*', '', resp_stripped).strip()
                resp_stripped = re.sub(r'[_\-\s]+$', '', resp_stripped).strip()
                min_response_len = 3 if is_short_answer else 15
                if len(resp_stripped) < min_response_len:
                    is_blank = True

            if not is_blank:
                # Check for blank vocab terms within this section (Term: with no definition)
                resp_lines = response.split('\n')
                for line_idx, line in enumerate(resp_lines):
                    line = line.strip()
                    if ':' in line and not line.startswith('http'):
                        parts = line.split(':', 1)
                        term = parts[0].strip()
                        defn = parts[1].strip() if len(parts) > 1 else ''
                        # If term is short (vocab-like) and definition is empty on this line,
                        # check the NEXT line(s) — student may have put the definition there
                        if len(term.split()) <= 3 and not defn:
                            # Look ahead up to 2 lines for a definition
                            has_next_line_content = False
                            for look_ahead in range(1, 3):
                                if line_idx + look_ahead < len(resp_lines):
                                    next_line = resp_lines[line_idx + look_ahead].strip()
                                    # Skip empty lines and underscore-only lines
                                    if next_line and not re.match(r'^[_\s\-\.]+$', next_line) and len(next_line) > 3:
                                        has_next_line_content = True
                                        break
                            if has_next_line_content:
                                continue  # Definition is on the next line — not blank
                            # Don't flag terms that match excluded section keywords
                            term_lower = term.lower()
                            is_from_excluded = any(
                                term_lower in em or em.split('\n')[0].strip().endswith(term_lower)
                                for em in exclude_markers_normalized
                            ) if exclude_markers_normalized else False
                            if not is_from_excluded:
                                blank_questions.append(f"{term} (no definition)")

                extracted.append({
                    "question": question_label,
                    "answer": response[:10000],
                    "type": "marker_response"
                })
            else:
                blank_questions.append(question_label)

        # Post-processing: strip any excluded marker text that leaked into responses
        if exclude_markers_normalized and extracted:
            for resp in extracted:
                answer = resp.get("answer", "")
                if not answer:
                    continue
                answer_lower = answer.lower()
                for em in exclude_markers_normalized:
                    em_lines = [l.strip() for l in em.split('\n') if l.strip()]
                    for em_line in em_lines[:3]:
                        stripped_em = re.sub(r'^[^a-zA-Z]*', '', em_line).strip()
                        # If the exclude marker line appears at the start of the answer, strip it and everything after
                        for search_term in [em_line, stripped_em]:
                            if search_term and len(search_term) >= 5:
                                pos = answer_lower.find(search_term)
                                if pos != -1:
                                    answer = answer[:pos].strip()
                                    answer_lower = answer.lower()
                # Strip trailing emoji/whitespace/special chars left over
                answer = answer.rstrip()
                answer = re.sub(r'[\s\U00010000-\U0010ffff]+$', '', answer).strip()
                if answer != resp["answer"]:
                    resp["answer"] = answer
                    if not answer or len(answer.strip()) < 5:
                        # Answer was entirely excluded content — mark as blank
                        extracted.remove(resp)
                        blank_questions.append(resp.get("question", "Unknown"))

        # If we found markers, return results (skip pattern matching)
        if marker_positions:
            total_q = len(extracted) + len(blank_questions)
            match_summary = f"{exact_matches} exact"
            if fuzzy_matches > 0:
                match_summary += f", {fuzzy_matches} fuzzy"
            summary = f"Extracted {len(extracted)} responses using {len(marker_positions)} markers ({match_summary})."
            if excluded_sections:
                summary += f" Excluded {len(excluded_sections)} section(s) from grading."
            if missing_sections:
                summary += f" {len(missing_sections)} required section(s) MISSING from submission."
            return {
                "extracted_responses": extracted,
                "blank_questions": blank_questions,
                "missing_sections": missing_sections,
                "total_questions": total_q,
                "answered_questions": len(extracted),
                "extraction_summary": summary,
                "excluded_sections": excluded_sections
            }

    # FALLBACK: If no markers provided, use simple pattern matching
    lines = document_text.split('\n')

    # Look for numbered questions with answers
    current_question = None
    current_answer_lines = []

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Check for numbered question (1. or 1))
        num_match = re.match(r'^(\d+)[\.\)]\s*(.+)', line_stripped)
        if num_match:
            # Save previous Q&A if exists
            if current_question and current_answer_lines:
                answer = ' '.join(current_answer_lines).strip()
                # Filter out questions/prompts from answer
                answer = filter_questions_from_response(answer)
                if answer and len(answer) > 3:
                    extracted.append({
                        "question": current_question,
                        "answer": answer[:10000],
                        "type": "numbered_qa"
                    })

            # Start new question
            question_text = num_match.group(2).strip()
            # Check if answer is on same line (after ? or :)
            if '?' in question_text:
                parts = question_text.split('?', 1)
                current_question = parts[0].strip() + '?'
                if len(parts) > 1 and parts[1].strip():
                    current_answer_lines = [parts[1].strip()]
                else:
                    current_answer_lines = []
            else:
                current_question = question_text
                current_answer_lines = []
        elif current_question:
            # This might be part of the answer
            current_answer_lines.append(line_stripped)

    # Don't forget last Q&A
    if current_question and current_answer_lines:
        answer = ' '.join(current_answer_lines).strip()
        # Filter out questions/prompts from answer
        answer = filter_questions_from_response(answer)
        if answer and len(answer) > 3:
            extracted.append({
                "question": current_question,
                "answer": answer[:10000],
                "type": "numbered_qa"
            })

    # Fill-in-the-blank pattern 1a: ___answer___ (wrapped in 2+ underscores)
    blank_matches = re.findall(r'_{2,}([^_\n]{1,100})_{2,}', document_text)
    for match in blank_matches:
        answer = match.strip()
        if answer and len(answer) > 0:
            extracted.append({
                "question": "Fill-in-blank",
                "answer": answer,
                "type": "fill_in_blank"
            })

    # Fill-in-the-blank pattern 1b: _answer_ (wrapped in single underscores)
    # Common format: "_Meriwether_ Lewis" or "named _Pomp_"
    single_blank_matches = re.findall(r'(?<![_\w])_([^_\n]{1,80})_(?![_\w])', document_text)
    for match in single_blank_matches:
        answer = match.strip()
        # Filter out likely template placeholders vs actual answers
        if answer and len(answer) > 0 and len(answer) < 60:
            # Skip common template words
            if answer.lower() not in ['name', 'date', 'period', 'class', 'blank', 'answer']:
                extracted.append({
                    "question": "Fill-in-blank",
                    "answer": answer,
                    "type": "fill_in_blank"
                })

    # Fill-in-the-blank pattern 2: Numbered blanks where student replaced underscores
    # e.g., "1. Clark" or "5. Louisiana Purchase" (student replaced _____ with answer)
    for line in lines:
        line_stripped = line.strip()
        # Match numbered items that have short answers (likely FITB)
        num_short = re.match(r'^(\d+)[\.\)]\s*([A-Za-z][^?]{1,60})$', line_stripped)
        if num_short:
            q_num = num_short.group(1)
            answer = num_short.group(2).strip()
            # Skip if it's a question (ends with ?) or instruction
            if answer and not answer.endswith('?') and not any(kw in answer.lower() for kw in ['write', 'explain', 'describe', 'answer', 'complete', 'fill in']):
                # Check if this looks like a student answer (not a question prompt)
                if len(answer) < 50 and not re.match(r'^[_\s\-\.]+$', answer):
                    extracted.append({
                        "question": f"Question {q_num}",
                        "answer": answer,
                        "type": "fill_in_blank"
                    })

    # Fill-in-the-blank pattern 3: Lines with blanks followed by text
    # e.g., "_____ led the expedition" where student wrote "Clark led the expedition"
    # ONLY triggers when underscores are present — prevents false positives on reading passages
    for line in lines:
        line_stripped = line.strip()
        if not line_stripped or line_stripped.endswith('?'):
            continue
        # Only match if the line actually contains fill-in-the-blank underscores
        if '___' not in line_stripped:
            continue
        leading_match = re.match(r'^([A-Z][a-z]+(?:\s+[A-Za-z]+){0,2})\s+(.{10,})$', line_stripped)
        if leading_match:
            potential_answer = leading_match.group(1)
            context = leading_match.group(2)
            if '___' in context:
                if not any(kw in potential_answer.lower() for kw in ['name', 'date', 'period', 'class', 'write', 'answer']):
                    extracted.append({
                        "question": f"Fill-in: {context[:40]}...",
                        "answer": potential_answer,
                        "type": "fill_in_blank"
                    })

    total_q = len(extracted) + len(blank_questions)
    return {
        "extracted_responses": extracted,
        "blank_questions": blank_questions,
        "total_questions": max(total_q, 1),
        "answered_questions": len(extracted),
        "extraction_summary": f"Found {len(extracted)} responses via pattern matching."
    }


def extract_student_responses_legacy(document_text: str, custom_markers: list = None) -> dict:
    """LEGACY: Old complex extraction - kept for reference but not used."""
    import re

    lines = document_text.split('\n')
    extracted = []
    blank_questions = []

    if response_sections:
        for section in response_sections:
            start_marker = section.get('start', '').strip()
            end_marker = section.get('end', '').strip() if section.get('end') else None

            if not start_marker:
                continue

            start_lower = start_marker.lower()
            doc_lower = document_text.lower()
            start_pos = doc_lower.find(start_lower)

            if start_pos == -1:
                blank_questions.append(start_marker)
                continue

            content_start = start_pos + len(start_marker)

            if end_marker:
                end_pos = doc_lower.find(end_marker.lower(), content_start)
                if end_pos == -1:
                    end_pos = min(content_start + 1000, len(document_text))  # Cap at 1000 chars
            else:
                # No end marker - find next likely section or cap at 800 chars
                # Look for common section boundaries
                next_section = len(document_text)
                for boundary in ['\n\n\n', '📝', '✏️', '🌱', '\nVocabulary', '\nQuestions', '\nReflection', '\nMain Idea']:
                    pos = doc_lower.find(boundary.lower(), content_start + 20)  # Skip a bit to avoid finding ourselves
                    if pos != -1 and pos < next_section:
                        next_section = pos
                end_pos = min(next_section, content_start + 800)

            # Extract the content
            content = document_text[content_start:end_pos].strip()

            # Clean up: remove leading colons, newlines, and prompt fragments
            content = re.sub(r'^[:\s\n]+', '', content)
            content = re.sub(r'^\s*of\s+(?:today.?s?\s+)?(?:the\s+)?(?:reading|lesson|section|article)[:\s]*\n*', '', content, flags=re.IGNORECASE).strip()
            content = re.sub(r'^\s*(?:today.?s?\s+)?(?:the\s+)?(?:reading|lesson|article)[:\s]*\n*', '', content, flags=re.IGNORECASE).strip()

            # Check if there's actual student content (not just template text)
            if content and len(content) > 10:
                # Filter out template text
                is_template = any(kw in content.lower() for kw in [
                    'write your', 'answer each', 'explain how', 'describe how',
                    'complete the', 'using the reading', 'type your answer',
                    '[your answer here]', '[type here]'
                ])

                if not is_template:
                    extracted.append({
                        "question": start_marker,
                        "answer": content[:10000],
                        "type": "highlighted_section"
                    })
                else:
                    blank_questions.append(start_marker)
            else:
                blank_questions.append(start_marker)

    # Patterns for questions/prompts
    question_patterns = [
        r'^\d+[\.\)]\s*(.+)',  # "1. Question" or "1) Question"
        r'^[a-zA-Z][\.\)]\s*(.+)',  # "a. Question" or "a) Question"
    ]

    # Track current question for multi-line answers
    current_question = None
    current_question_idx = -1

    for i, line in enumerate(lines):
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Pattern 1: Fill-in-the-blank with answer (text between underscores)
        # e.g., "The year was ___1803___"
        blank_matches = re.findall(r'_{2,}([^_\n]{1,50})_{2,}', line)
        for match in blank_matches:
            answer = match.strip()
            if answer and len(answer) > 0:
                # Try to find the question context
                question_context = line_stripped[:50] + "..." if len(line_stripped) > 50 else line_stripped
                extracted.append({
                    "question": question_context,
                    "answer": answer,
                    "type": "fill_in_blank"
                })

        # Pattern 2: Vocabulary definitions (Term: answer)
        # e.g., "Rebellion: going again"
        if ':' in line_stripped:
            parts = line_stripped.split(':', 1)
            term = parts[0].strip()
            answer = parts[1].strip() if len(parts) > 1 else ""

            # Check if this looks like a vocabulary term (not a section header)
            is_vocab = (
                len(term) > 2 and
                len(term) < 50 and
                not any(kw in term.lower() for kw in ['question', 'task', 'summary', 'reflection', 'main idea', 'vocabulary', 'answer', 'notes', 'reading', 'section'])
            )

            if is_vocab:
                # Check if answer is blank or just underscores
                if answer and not re.match(r'^[_\s\-\.]+$', answer) and len(answer) > 1:
                    extracted.append({
                        "question": f"Define: {term}",
                        "answer": answer,
                        "type": "vocabulary"
                    })
                else:
                    blank_questions.append(f"Define: {term}")

        # Pattern 3: Numbered questions with answers on same line
        # e.g., "1. Why was this important? It helped the country grow"
        num_match = re.match(r'^(\d+)[\.\)]\s*(.+?)[\?:]\s*(.+)$', line_stripped)
        if num_match:
            q_num = num_match.group(1)
            question = num_match.group(2).strip() + "?"
            answer = num_match.group(3).strip()

            if answer and not re.match(r'^[_\s\-\.]+$', answer) and len(answer) > 1:
                extracted.append({
                    "question": f"{q_num}. {question}",
                    "answer": answer,
                    "type": "short_answer"
                })
            else:
                blank_questions.append(f"{q_num}. {question}")
            continue

        # Pattern 4: Questions ending with ? (track for next line answer)
        if line_stripped.endswith('?'):
            current_question = line_stripped
            current_question_idx = i
            continue

        # Pattern 5: Numbered question without answer on same line
        num_only = re.match(r'^(\d+)[\.\)]\s*(.+)$', line_stripped)
        if num_only and '?' in line_stripped:
            current_question = line_stripped
            current_question_idx = i
            continue

        # Pattern 6: Check if this line is an answer to the previous question
        if current_question and i == current_question_idx + 1:
            # This might be an answer
            potential_answer = line_stripped

            # Filter out template text, headers, and blanks
            is_template = any(kw in potential_answer.lower() for kw in [
                'write your', 'explain', 'describe', 'answer', 'complete',
                'section', 'chapter', 'reading', 'vocabulary', 'summary',
                'reflection', 'task', 'question'
            ])
            is_blank = re.match(r'^[_\s\-\.]*$', potential_answer) or len(potential_answer) < 3

            if not is_template and not is_blank:
                extracted.append({
                    "question": current_question,
                    "answer": potential_answer,
                    "type": "short_answer"
                })
            else:
                blank_questions.append(current_question)

            current_question = None

    # Pattern 7: Check for section headers that expect written responses
    # These are headers like "Summary:", "📝 SUMMARY", "Reflection:", etc. followed by student paragraphs
    # Using (?:📝|✏️|✍️|🖊️|🌟)? to optionally match common emojis
    section_headers = [
        # Summary patterns - be specific to avoid matching "Questions / Summary" headers
        (r'write\s+(?:a\s+)?(?:\d+[–-]\d+\s+)?(?:sentence\s+)?summary[:\s]+', 'Summary'),  # "Write a 2-3 sentence summary:"
        (r'(?:📝|✏️|✍️|🖊️)\s*summary\s*(?:\([^)]*\))?\s*[:\.\n]', 'Summary'),  # 📝 SUMMARY (3-4 sentences):
        (r'\bsummary\s*of\s+(?:today\'?s?\s+)?(?:reading|lesson|section)[:\s]+', 'Summary'),  # "Summary of today's reading:"
        (r'(?:^|\n)\s*summary[:\s]*(?=\n)', 'Summary'),  # "Summary:" on its own line
        # Reflection patterns
        (r'(?:📝|✏️|✍️)?\s*reflection\s*[^\.]*[:\.]', 'Reflection'),
        (r'(?:^|\n)\s*reflection[:\s]+', 'Reflection'),
        (r'final reflection[:\s]*', 'Final Reflection'),
        # Task patterns
        (r'(?:📝|✏️|✍️)?\s*student task\s*[^\.]*[:\.]', 'Student Task'),
        (r'student task[:\s]+', 'Student Task'),
        # Other patterns
        (r'main idea[:\s]+', 'Main Idea'),
        (r'explain in your own words[:\s]*', 'Explanation'),
        (r'write your answer[:\s]*', 'Written Response'),
        (r'your response[:\s]*', 'Response'),
    ]

    full_text_lower = document_text.lower()
    full_text = document_text  # Keep original case for answer extraction

    for pattern, section_name in section_headers:
        match = re.search(pattern, full_text_lower)
        if match:
            # Find the position of the header
            header_end = match.end()

            # Look for paragraph content after the header
            # Skip blank lines and find the next substantial text
            remaining_text = full_text[header_end:]

            # Split into lines/paragraphs to find student response
            # Skip instruction lines that start with prompt words
            instruction_starters = [
                'explain how', 'describe how', 'write a', 'write your', 'answer each',
                'complete the', 'using the reading', 'in a few sentences', 'in your own words',
                'what is', 'what are', 'what was', 'what were', 'how did', 'how do', 'why did', 'why do',
                'summarize', 'identify', 'list the', 'define the'
            ]

            # Find paragraphs (split by double newline or single newline)
            paragraphs = re.split(r'\n\s*\n|\n', remaining_text)
            student_response = None

            for para in paragraphs:
                para = para.strip()
                if not para or len(para) < 15:
                    continue

                # Check if this looks like an instruction line
                para_lower = para.lower()
                is_instruction = any(para_lower.startswith(starter) for starter in instruction_starters)

                # Also check for question marks at the end (likely a prompt, not answer)
                if para.rstrip().endswith('?'):
                    is_instruction = True

                if not is_instruction:
                    # This looks like a student response
                    # Clean up any remaining prompt fragments
                    content = para
                    content = re.sub(r'^\s*of\s+(?:today.?s?\s+)?(?:the\s+)?(?:reading|lesson|section|article)[:\s]*', '', content, flags=re.IGNORECASE).strip()
                    content = re.sub(r'^\s*(?:today.?s?\s+)?(?:the\s+)?(?:reading|lesson|article)[:\s]*', '', content, flags=re.IGNORECASE).strip()

                    # Check it has substance
                    if len(content) > 20 and any(c.isalpha() for c in content):
                        student_response = content
                        break

            if student_response:
                # Check if we already captured this (avoid duplicates)
                already_captured = any(student_response[:50] in r.get('answer', '')[:50] for r in extracted)
                if not already_captured:
                    extracted.append({
                        "question": section_name,
                        "answer": student_response[:10000],
                        "type": "written_response"
                    })
            else:
                # Section appears to be blank or only has instructions
                blank_questions.append(section_name)

    # Build summary
    total_q = len(extracted) + len(blank_questions)
    answered_q = len(extracted)

    summary = f"Found {answered_q} responses out of {total_q} detected questions/prompts."
    if blank_questions:
        summary += f" {len(blank_questions)} questions left blank."

    return {
        "extracted_responses": extracted,
        "blank_questions": blank_questions,
        "total_questions": total_q,
        "answered_questions": answered_q,
        "extraction_summary": summary
    }


def format_extracted_for_grading(extraction_result: dict, marker_config: list = None, extraction_mode: str = 'structured') -> str:
    """
    Format pre-extracted responses for the grading prompt.
    Includes section point values if provided.

    Args:
        extraction_result: Dict with extracted_responses, blank_questions, etc.
        marker_config: List of marker configs with points, e.g. [{"start": "Summary", "points": 20}, ...]
        extraction_mode: "structured" (parse with rules) or "ai" (send raw, let AI figure it out)
    """
    if not extraction_result or not extraction_result.get("extracted_responses"):
        return "NO STUDENT RESPONSES FOUND - Document appears to be blank or unfinished."

    # Build marker points lookup
    marker_points = {}
    if marker_config:
        for m in marker_config:
            if isinstance(m, dict):
                marker_points[m.get('start', '').lower()] = m.get('points', 10)
            elif isinstance(m, str):
                marker_points[m.lower()] = 10

    output = []
    output.append("=" * 50)
    if extraction_mode == 'ai':
        output.append("RAW SECTION CONTENT (AI will identify prompts vs student responses)")
    else:
        output.append("VERIFIED STUDENT RESPONSES (extracted from document)")
    output.append("=" * 50)
    output.append("")

    for i, item in enumerate(extraction_result["extracted_responses"], 1):
        q_type = item.get("type", "unknown")
        question = item.get("question", "Unknown question")
        answer = item.get("answer", "")

        if extraction_mode == 'ai':
            # AI mode: send raw content, let AI identify what's prompt vs response
            cleaned_answer = answer
        else:
            # Structured mode: clean answer by stripping out question text
            cleaned_lines = []
            for line in answer.split('\n'):
                line = line.strip()
                if not line:
                    continue
                # Skip common prompt-only lines (no student content)
                if line.lower() in ('write your answer:', 'your answer:', 'answer:'):
                    continue
                # If line has a question mark, only keep text AFTER the ?
                if '?' in line:
                    parts = line.split('?')
                    after_q = parts[-1].strip() if len(parts) > 1 else ''
                    if after_q and len(after_q) > 1:
                        cleaned_lines.append(after_q)
                    # Skip lines that are ONLY questions (nothing after ?)
                # If line has colon (vocab format), only keep text AFTER the :
                elif ':' in line and not line.startswith('http'):
                    parts = line.split(':', 1)
                    term = parts[0].strip()
                    defn = parts[1].strip() if len(parts) > 1 else ''
                    # Only treat as vocab if term is short (1-4 words) AND has a definition
                    if len(term.split()) <= 4 and defn and len(defn) > 1:
                        cleaned_lines.append(defn)
                    elif defn and len(defn) > 1:
                        # Longer term but has definition - keep whole line
                        cleaned_lines.append(line)
                    # Skip lines that are just "Term:" with no definition (blank vocab)
                else:
                    cleaned_lines.append(line)

            cleaned_answer = '\n'.join(cleaned_lines) if cleaned_lines else answer

        # Look up points for this section
        points_str = ""
        for marker_key, pts in marker_points.items():
            if marker_key in question.lower():
                points_str = f" [{pts} points]"
                break

        output.append(f"[{i}] {question}{points_str}")
        output.append(f"    STUDENT ANSWER: \"{cleaned_answer}\"")
        output.append(f"    (Type: {q_type})")
        output.append("")

    # Add blank questions with their point values
    if extraction_result.get("blank_questions"):
        output.append("-" * 50)
        output.append("UNANSWERED QUESTIONS (left blank by student):")
        for q in extraction_result["blank_questions"]:
            points_str = ""
            for marker_key, pts in marker_points.items():
                if marker_key in q.lower():
                    points_str = f" [LOSES {pts} points]"
                    break
            output.append(f"  • {q}{points_str}")

    # Add missing sections (required by assignment but entirely absent from submission)
    if extraction_result.get("missing_sections"):
        output.append("-" * 50)
        output.append("MISSING SECTIONS (required by assignment but not found in student submission):")
        for section in extraction_result["missing_sections"]:
            points_str = ""
            for marker_key, pts in marker_points.items():
                if marker_key in section.lower():
                    points_str = f" [LOSES {pts} points]"
                    break
            output.append(f"  ✗ {section}{points_str} — ENTIRELY OMITTED")

    output.append("")
    output.append(f"SUMMARY: {extraction_result.get('extraction_summary', '')}")

    return "\n".join(output)


# =============================================================================
# WRITING STYLE ANALYSIS - For AI Detection
# =============================================================================

def analyze_writing_style(text: str) -> dict:
    """
    Analyze writing style metrics from student text.
    Used to build a profile and detect AI-generated content.
    """
    if not text or len(text.strip()) < 20:
        return None

    # Clean text
    clean_text = text.strip()

    # Split into sentences (basic sentence detection)
    sentences = re.split(r'[.!?]+', clean_text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 3]

    # Split into words
    words = re.findall(r'\b[a-zA-Z]+\b', clean_text)
    if len(words) < 5:
        return None

    # Calculate metrics
    avg_word_length = sum(len(w) for w in words) / len(words)
    avg_sentence_length = len(words) / max(len(sentences), 1)

    # Vocabulary complexity (based on word length distribution)
    long_words = [w for w in words if len(w) > 7]
    complex_word_ratio = len(long_words) / len(words)

    # Detect common misspellings (lowercase words that might be proper nouns)
    potential_misspellings = []
    common_misspelled = re.findall(r'\b[a-z]+[A-Z][a-z]*\b|\b[a-z]{2,}\b', clean_text)

    # Check for specific patterns
    uses_contractions = bool(re.search(r"\b(don't|can't|won't|isn't|aren't|doesn't|didn't|wouldn't|couldn't|shouldn't|I'm|you're|they're|we're|it's|that's|what's|there's|here's)\b", clean_text, re.IGNORECASE))

    # Capitalization habits
    proper_caps = len(re.findall(r'\b[A-Z][a-z]+\b', clean_text))
    all_caps = len(re.findall(r'\b[A-Z]{2,}\b', clean_text))

    # Simple vs complex vocabulary indicators
    simple_words = ['the', 'a', 'an', 'is', 'was', 'are', 'were', 'it', 'they', 'he', 'she', 'we', 'you', 'i', 'and', 'but', 'or', 'so', 'because', 'like', 'just', 'really', 'very', 'good', 'bad', 'big', 'small']
    simple_count = sum(1 for w in words if w.lower() in simple_words)
    simple_ratio = simple_count / len(words)

    # Academic/AI indicator words
    academic_words = ['furthermore', 'therefore', 'consequently', 'however', 'nevertheless', 'moreover', 'subsequently', 'fundamental', 'significant', 'essentially', 'particularly', 'specifically', 'transforming', 'establishing', 'securing', 'trajectory', 'precedent', 'constitutional', 'acquisition', 'vital', 'expansion']
    academic_count = sum(1 for w in words if w.lower() in academic_words)

    # Calculate complexity score (1-10 scale)
    complexity_score = min(10, max(1,
        (avg_word_length - 3) * 1.5 +  # Word length contribution
        (avg_sentence_length / 5) +     # Sentence length contribution
        (complex_word_ratio * 10) +     # Complex words contribution
        (academic_count * 2) -          # Academic words add complexity
        (simple_ratio * 3)              # Simple words reduce complexity
    ))

    return {
        "avg_word_length": round(avg_word_length, 2),
        "avg_sentence_length": round(avg_sentence_length, 2),
        "word_count": len(words),
        "sentence_count": len(sentences),
        "complex_word_ratio": round(complex_word_ratio, 3),
        "simple_word_ratio": round(simple_ratio, 3),
        "academic_word_count": academic_count,
        "uses_contractions": uses_contractions,
        "complexity_score": round(complexity_score, 2)
    }


def compare_writing_styles(current_style: dict, historical_profile: dict) -> dict:
    """
    Compare current submission's writing style against student's historical profile.
    Returns deviation analysis and AI likelihood.
    """
    if not current_style or not historical_profile:
        return {"deviation": "unknown", "ai_likelihood": "unknown", "reason": "Insufficient data"}

    deviations = []

    # Check complexity score deviation
    hist_complexity = historical_profile.get("avg_complexity_score", 3.0)
    curr_complexity = current_style.get("complexity_score", 3.0)
    complexity_diff = curr_complexity - hist_complexity

    if complexity_diff > 3:
        deviations.append(f"Complexity jumped from {hist_complexity:.1f} to {curr_complexity:.1f}")

    # Check sentence length deviation
    hist_sent_len = historical_profile.get("avg_sentence_length", 8.0)
    curr_sent_len = current_style.get("avg_sentence_length", 8.0)
    sent_len_diff = curr_sent_len - hist_sent_len

    if sent_len_diff > 10:
        deviations.append(f"Sentence length jumped from {hist_sent_len:.1f} to {curr_sent_len:.1f} words")

    # Check for sudden academic vocabulary
    hist_academic = historical_profile.get("avg_academic_words", 0)
    curr_academic = current_style.get("academic_word_count", 0)

    if curr_academic > hist_academic + 2:
        deviations.append(f"Academic vocabulary increased significantly ({curr_academic} vs typical {hist_academic})")

    # Check word length deviation
    hist_word_len = historical_profile.get("avg_word_length", 4.0)
    curr_word_len = current_style.get("avg_word_length", 4.0)

    if curr_word_len - hist_word_len > 1.5:
        deviations.append(f"Word length increased from {hist_word_len:.1f} to {curr_word_len:.1f}")

    # Determine AI likelihood based on deviations
    if len(deviations) >= 3:
        ai_likelihood = "likely"
    elif len(deviations) >= 2:
        ai_likelihood = "possible"
    elif len(deviations) == 1 and complexity_diff > 4:
        ai_likelihood = "possible"
    else:
        ai_likelihood = "none"

    return {
        "deviation": "significant" if len(deviations) >= 2 else "minor" if len(deviations) == 1 else "none",
        "ai_likelihood": ai_likelihood,
        "deviations": deviations,
        "reason": "; ".join(deviations) if deviations else "Writing style consistent with history"
    }


def update_writing_profile(student_id: str, current_style: dict, student_name: str = None):
    """
    Update student's writing profile with new submission data.
    Maintains running averages across assignments.
    """
    if not current_style or not student_id:
        return

    history_dir = os.path.expanduser("~/.graider_data/student_history")
    history_file = os.path.join(history_dir, f"{student_id}.json")

    try:
        if os.path.exists(history_file):
            with open(history_file, 'r') as f:
                history = json.load(f)
        else:
            history = {"student_id": student_id, "assignments": []}

        # Always update the student name if provided
        if student_name:
            history["name"] = student_name

        # Get or initialize writing profile
        profile = history.get("writing_profile", {
            "avg_word_length": 0,
            "avg_sentence_length": 0,
            "avg_complexity_score": 0,
            "avg_academic_words": 0,
            "uses_contractions": False,
            "sample_count": 0
        })

        # Update running averages
        n = profile.get("sample_count", 0)
        if n > 0:
            profile["avg_word_length"] = (profile["avg_word_length"] * n + current_style["avg_word_length"]) / (n + 1)
            profile["avg_sentence_length"] = (profile["avg_sentence_length"] * n + current_style["avg_sentence_length"]) / (n + 1)
            profile["avg_complexity_score"] = (profile["avg_complexity_score"] * n + current_style["complexity_score"]) / (n + 1)
            profile["avg_academic_words"] = (profile["avg_academic_words"] * n + current_style["academic_word_count"]) / (n + 1)
        else:
            profile["avg_word_length"] = current_style["avg_word_length"]
            profile["avg_sentence_length"] = current_style["avg_sentence_length"]
            profile["avg_complexity_score"] = current_style["complexity_score"]
            profile["avg_academic_words"] = current_style["academic_word_count"]

        profile["uses_contractions"] = profile.get("uses_contractions", False) or current_style["uses_contractions"]
        profile["sample_count"] = n + 1

        # Round values
        for key in ["avg_word_length", "avg_sentence_length", "avg_complexity_score", "avg_academic_words"]:
            if key in profile:
                profile[key] = round(profile[key], 2)

        history["writing_profile"] = profile
        history["last_updated"] = datetime.now().isoformat()

        # Save updated history
        os.makedirs(history_dir, exist_ok=True)
        with open(history_file, 'w') as f:
            json.dump(history, f, indent=2)

    except Exception as e:
        print(f"  ⚠️ Could not update writing profile: {e}")


def get_writing_profile(student_id: str) -> dict:
    """
    Retrieve student's historical writing profile.
    """
    if not student_id:
        return None

    history_dir = os.path.expanduser("~/.graider_data/student_history")
    history_file = os.path.join(history_dir, f"{student_id}.json")

    try:
        if os.path.exists(history_file):
            with open(history_file, 'r') as f:
                history = json.load(f)
                return history.get("writing_profile")
    except Exception:
        pass

    return None

# =============================================================================
# CONFIGURATION - UPDATE THESE FOR EACH GRADING SESSION
# =============================================================================

# Folder containing student assignment files (.docx)
ASSIGNMENT_FOLDER = "/Users/alexc/Downloads/Assignments"

# Output folder for CSV and email files
OUTPUT_FOLDER = "/Users/alexc/Downloads/Assignment Grader/Results"

# Path to your student roster Excel file
ROSTER_FILE = "/Users/alexc/Downloads/Assignment Grader/all_students_updated.xlsx"

# Your OpenAI API key (set in .env file or paste here)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "your-api-key-here")

# Anthropic API key for Claude models
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")

# Google Gemini API key
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

# Assignment name (used in output files and emails)
ASSIGNMENT_NAME = "Cornell Notes - Political Parties"  # UPDATE FOR EACH ASSIGNMENT

# Marker phrase(s) that indicate where student work begins
# Only content within the section (until next header) will be graded
STUDENT_WORK_MARKERS = [
    # Direct task indicators
    "student task",
    "your turn",
    "now you try",
    "student responses below",
    
    # Question indicators
    "answer the question",
    "questions to check understanding",
    "questions / summary",
    
    # Activity types
    "answer the",
    "match",
    "fill-in-the-blank",
    "fill in the blank",
    "write",
    "explain",
    "complete",
    "describe",
    "summarize",
    
    # Analysis & thinking verbs
    "analyze",
    "compare",
    "contrast",
    "evaluate",
    "identify",
    "list",
    "define",
    
    # Application verbs
    "apply",
    "demonstrate",
    "illustrate",
    "predict",
    "solve",
    
    # Creation verbs
    "create",
    "design",
    "develop",
    "construct",
    
    # Reflection indicators
    "final reflection",
    "final reflection question",
    "reflect",
    "think about",
    
    # Learning check indicators
    "let's see what you've learned",
    "lets see what you've learned",
    "let's see what you've learned",  # curly apostrophe version
    "check your understanding",
    "show what you know",
    "practice",
]

# Section headers that indicate a NEW section (stops extraction from previous marker)
SECTION_HEADERS = [
    "vocabulary",
    "vocabulary mini-lesson",
    "key vocabulary",
    "notes",
    "guided notes",
    "reading",
    "directions",
    "instructions",
    "primary source",
    "background",
    "overview",
    "introduction",
]


# =============================================================================
# GRADING RUBRIC - Customize point values and criteria as needed
# =============================================================================

GRADING_RUBRIC = """
You are grading 6th grade Social Studies assignments. Be ENCOURAGING and GENEROUS.
These are 11-12 year old students - grade with appropriate expectations for their age.

IMPORTANT GRADING GUIDELINES:
- For FILL-IN-THE-BLANK exercises: Accept any answer that is factually correct or very close. Spelling mistakes should NOT reduce the grade if the intent is clear.
- Accept reasonable synonyms, alternate phrasings, and partial answers that demonstrate understanding.
- If a student gets the main idea right but uses slightly different wording, give FULL CREDIT.
- Minor spelling errors (like "piler" instead of "pillar") should NOT be penalized.
- Be GENEROUS - when in doubt, give the student the benefit of the doubt.

GRADING SCALE (out of 100 points):

1. CONTENT ACCURACY (40 points)
   - Are the answers factually correct or demonstrate understanding?
   - For fill-in-the-blank: Is the completed statement essentially true?
   - 40 pts: Most answers correct (90%+)
   - 35 pts: Good understanding (80-89% correct)
   - 30 pts: Solid effort (70-79% correct)
   - 25 pts: Some understanding (60-69% correct)
   - 20 pts: Partial understanding (50-59% correct)
   - Below 20: Less than half correct

2. COMPLETENESS (25 points)
   - Did the student attempt ALL questions/blanks?
   - 25 pts: All questions attempted
   - 20 pts: Nearly all attempted (90%+)
   - 15 pts: Most attempted (75%+)
   - 10 pts: About half attempted
   - 5 pts: Less than half attempted

3. WRITING QUALITY (20 points)
   - Is the writing legible and understandable?
   - For fill-in-blank, this is less important - be generous
   - 20 pts: Clear and readable
   - 15 pts: Minor issues but understandable
   - 10 pts: Some difficulty but can figure out meaning
   - 5 pts: Hard to understand

4. EFFORT & ENGAGEMENT (15 points)
   - Did the student put in genuine effort?
   - Are answers thoughtful (not random guesses)?
   - 15 pts: Clear effort shown
   - 10 pts: Good effort
   - 5 pts: Minimal effort
   - 0 pts: No real effort

GRADE RANGES:
- A: 90-100 (Great job!)
- B: 80-89 (Good work!)
- C: 70-79 (Solid effort)
- D: 60-69 (Needs improvement)
- F: Below 60 (Significant concerns)

REMEMBER: These are 6th graders. Be kind, encouraging, and generous with grading.
A student who attempts all questions and gets most right should get an A or B.
"""

# DEFAULT ASSIGNMENT INSTRUCTIONS - Only used if no assignment-specific config provided
# This should be empty or generic - assignment-specific instructions come from Builder configs
ASSIGNMENT_INSTRUCTIONS = """
Grade the student's work based on what they were asked to do.
Focus on the content they provided, not on sections that may not apply to this assignment type.
"""


# =============================================================================
# ROSTER LOADING - Reads your Excel student list
# =============================================================================

def load_roster(roster_path: str) -> dict:
    """
    Load student roster from Excel or CSV file.

    Excel format: Student, Student ID, Local ID, Email, Grade, Team
    CSV format: FirstName, LastName, StudentID, Email, Period

    Returns dict mapping "firstname lastname" (lowercase) -> student info
    """
    roster = {}

    if not Path(roster_path).exists():
        print(f"⚠️  Roster file not found: {roster_path}")
        return {}

    roster_path_lower = roster_path.lower()

    # Handle CSV files
    if roster_path_lower.endswith('.csv'):
        import csv
        with open(roster_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                # Try various column name formats
                first_name = row.get('FirstName') or row.get('First Name') or row.get('first_name') or row.get('First') or ''
                last_name = row.get('LastName') or row.get('Last Name') or row.get('last_name') or row.get('Last') or ''
                student_id = row.get('StudentID') or row.get('Student ID') or row.get('student_id') or row.get('ID') or ''
                email = row.get('Email') or row.get('email') or ''
                period = row.get('Period') or row.get('period') or row.get('Class') or ''

                # If Student column exists with "Last; First" format
                if not first_name and not last_name:
                    student_col = row.get('Student') or row.get('Name') or ''
                    if ';' in student_col:
                        parts = student_col.split(';')
                        last_name = parts[0].strip()
                        first_name = parts[1].strip() if len(parts) > 1 else ''
                    elif ',' in student_col:
                        parts = student_col.split(',')
                        last_name = parts[0].strip()
                        first_name = parts[1].strip() if len(parts) > 1 else ''

                if not first_name and not last_name:
                    continue

                first_name_simple = first_name.split()[0] if first_name else ''
                lookup_key = f"{first_name_simple} {last_name}".lower().strip()

                entry = {
                    "student_id": str(student_id),
                    "student_name": f"{first_name} {last_name}".strip(),
                    "first_name": first_name,
                    "last_name": last_name,
                    "email": email or "",
                    "period": str(period) if period else ""
                }
                roster[lookup_key] = entry

                reverse_key = f"{last_name} {first_name_simple}".lower().strip()
                roster[reverse_key] = entry

                # For compound last names ("Wilkins Reels"), also add key with just first part
                # so "Dicen_Wilkins" filename matches "Dicen Macheil Wilkins Reels"
                last_parts = last_name.split()
                if len(last_parts) > 1:
                    short_key = f"{first_name_simple} {last_parts[0]}".lower().strip()
                    if short_key not in roster:
                        roster[short_key] = entry
                    reverse_short = f"{last_parts[0]} {first_name_simple}".lower().strip()
                    if reverse_short not in roster:
                        roster[reverse_short] = entry

        print(f"📋 Loaded {len(set(id(v) for v in roster.values()))} students from CSV roster")
        return roster

    # Handle Excel files
    try:
        import openpyxl
    except ImportError:
        print("❌ openpyxl not installed. Run: pip install openpyxl")
        return {}

    wb = openpyxl.load_workbook(roster_path)
    sheet = wb.active

    # Get header row to find column indices
    headers = [str(cell.value).lower() if cell.value else '' for cell in sheet[1]]

    # Try to find period column
    period_col = None
    for i, h in enumerate(headers):
        if 'period' in h or 'class' in h or 'team' in h:
            period_col = i
            break

    # Skip header row
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not row[0]:
            continue

        student_name = row[0]
        student_id = row[1]
        email = row[3] if len(row) > 3 else ''
        period = row[period_col] if period_col is not None and len(row) > period_col else ''

        if ';' in str(student_name):
            parts = str(student_name).split(';')
            last_name = parts[0].strip()
            first_name = parts[1].strip() if len(parts) > 1 else ""
            first_name_simple = first_name.split()[0] if first_name else ""
        else:
            last_name = str(student_name)
            first_name = ""
            first_name_simple = ""

        lookup_key = f"{first_name_simple} {last_name}".lower().strip()

        entry = {
            "student_id": str(student_id),
            "student_name": f"{first_name} {last_name}".strip(),
            "first_name": first_name,
            "last_name": last_name,
            "email": email or "",
            "period": str(period) if period else ""
        }
        roster[lookup_key] = entry

        reverse_key = f"{last_name} {first_name_simple}".lower().strip()
        roster[reverse_key] = entry

        # For compound last names, add short key with just first part of last name
        last_parts = last_name.split()
        if len(last_parts) > 1:
            short_key = f"{first_name_simple} {last_parts[0]}".lower().strip()
            if short_key not in roster:
                roster[short_key] = entry
            reverse_short = f"{last_parts[0]} {first_name_simple}".lower().strip()
            if reverse_short not in roster:
                roster[reverse_short] = entry

    print(f"📋 Loaded {len(set(id(v) for v in roster.values()))} students from Excel roster")

    # Supplement with period CSVs from Focus Import (adds students not in Excel roster)
    periods_dir = os.path.expanduser("~/.graider_data/periods")
    if os.path.exists(periods_dir):
        import csv as csv_mod
        added = 0
        for period_file in sorted(os.listdir(periods_dir)):
            if not period_file.endswith('.csv'):
                continue
            period_name = period_file.replace('.csv', '').replace('_', ' ')
            # Try to get period name from metadata
            meta_path = os.path.join(periods_dir, f"{period_file}.meta.json")
            if os.path.exists(meta_path):
                try:
                    with open(meta_path, 'r') as mf:
                        meta = json.load(mf)
                        period_name = meta.get('period_name', period_name)
                except Exception:
                    pass
            filepath = os.path.join(periods_dir, period_file)
            try:
                with open(filepath, 'r', encoding='utf-8') as fh:
                    reader = csv_mod.DictReader(fh)
                    for row in reader:
                        student_col = row.get('Student', row.get('Name', '')).strip().strip('"')
                        student_id = row.get('Student ID', '').strip().strip('"')
                        local_id = row.get('Local ID', '').strip().strip('"')
                        grade = row.get('Grade', '').strip().strip('"')

                        # Parse "Last, First" or "Last; First" format
                        first_name = ''
                        last_name = ''
                        for sep in [';', ',']:
                            if sep in student_col:
                                parts = student_col.split(sep, 1)
                                last_name = parts[0].strip()
                                first_name = parts[1].strip() if len(parts) > 1 else ''
                                break
                        if not first_name and not last_name:
                            continue

                        first_name_simple = first_name.split()[0] if first_name else ''
                        lookup_key = f"{first_name_simple} {last_name}".lower().strip()

                        # Only add if not already in roster (don't overwrite Excel data)
                        if lookup_key in roster:
                            continue

                        email = f"{local_id}@vcs2go.net" if local_id else ""
                        entry = {
                            "student_id": str(student_id),
                            "student_name": f"{first_name} {last_name}".strip(),
                            "first_name": first_name,
                            "last_name": last_name,
                            "email": email,
                            "period": period_name
                        }
                        roster[lookup_key] = entry
                        reverse_key = f"{last_name} {first_name_simple}".lower().strip()
                        if reverse_key not in roster:
                            roster[reverse_key] = entry
                        # Compound last name short keys
                        last_parts = last_name.split()
                        if len(last_parts) > 1:
                            short_key = f"{first_name_simple} {last_parts[0]}".lower().strip()
                            if short_key not in roster:
                                roster[short_key] = entry
                        added += 1
            except Exception:
                pass
        if added:
            print(f"📋 Supplemented with {added} students from period CSVs")

    return roster


# =============================================================================
# FILE PARSING - Extract student name from filename
# =============================================================================

def parse_filename(filename: str) -> dict:
    """
    Parse student info from filename.

    Expected formats:
        FirstName_LastName_AssignmentName.docx
        Last, First M._AssignmentName.docx
    Examples:
        A'kareah_West_Cornell Notes_ Political Parties.docx
        Eli_Long_Hamilton_Jefferson_Graphic_Organizer.docx
        Deloach, Rylee M._Washington_Stations_Handout.docx

    Returns: {"first_name": ..., "last_name": ..., "assignment_part": ...}
    """
    # Remove extension
    name = Path(filename).stem

    # Handle "Last, First M._Assignment" format (comma before first underscore)
    first_underscore = name.find('_')
    if first_underscore > 0 and ',' in name[:first_underscore]:
        name_part = name[:first_underscore]
        assignment_part = name[first_underscore + 1:] if first_underscore < len(name) - 1 else ""
        comma_parts = name_part.split(',')
        last_name = comma_parts[0].strip()
        # First name may include middle initial like "Rylee M."
        first_full = comma_parts[1].strip() if len(comma_parts) > 1 else ""
        first_name = first_full.split()[0] if first_full else ""

        return {
            "first_name": first_name,
            "last_name": last_name,
            "assignment_part": assignment_part,
            "lookup_key": f"{first_name} {last_name}".lower()
        }

    # Standard format: FirstName_LastName_AssignmentName
    parts = name.split('_')

    if len(parts) >= 2:
        first_name = parts[0].strip()
        last_name = parts[1].strip()
        assignment_part = '_'.join(parts[2:]) if len(parts) > 2 else ""

        return {
            "first_name": first_name,
            "last_name": last_name,
            "assignment_part": assignment_part,
            "lookup_key": f"{first_name} {last_name}".lower()
        }
    else:
        # Can't parse - return filename as-is
        return {
            "first_name": name,
            "last_name": "",
            "assignment_part": "",
            "lookup_key": name.lower()
        }


def read_docx_file(filepath: str) -> str:
    """
    Read text content from a Word document (.docx) in document order.
    This properly interleaves paragraphs and tables as they appear.
    """
    try:
        from docx import Document
        from docx.document import Document as DocType
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        print("❌ python-docx not installed. Run: pip install python-docx")
        return None

    try:
        doc = Document(filepath)
        full_text = []

        # Iterate through document body elements in order
        # This ensures tables and paragraphs appear in their actual document order
        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                para = Paragraph(element, doc)
                if para.text.strip():
                    full_text.append(para.text)
            # Check if it's a table
            elif element.tag.endswith('tbl'):
                table = Table(element, doc)
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))

        return '\n'.join(full_text)
    except Exception as e:
        print(f"  ⚠️  Error reading file: {e}")
        return None


def read_docx_file_structured(filepath: str) -> dict:
    """Read a .docx file and detect Graider structured tables.

    Iterates through doc.element.body looking for 2-row tables whose first cell
    contains a [GRAIDER:TYPE:ID] tag. Also checks for the GRAIDER_TABLE_V1 marker.

    Returns:
        {
            "is_graider_table": bool,
            "plain_text": str (full document text for fallback),
            "tables": [
                {"tag_type": "VOCAB"|"QUESTION"|"SUMMARY",
                 "tag_id": str, "header_text": str, "response": str},
                ...
            ]
        }
    """
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.shared import Pt
    except ImportError:
        return {"is_graider_table": False, "plain_text": None, "tables": []}

    try:
        doc = Document(filepath)
        tables_found = []
        full_text = []
        has_marker = False
        tag_pattern = re.compile(r'\[GRAIDER:(VOCAB|QUESTION|SUMMARY):([^\]]+)\]')

        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = Paragraph(element, doc)
                text = para.text.strip()
                if text:
                    full_text.append(text)
                    if text == "GRAIDER_TABLE_V1":
                        has_marker = True

            elif element.tag.endswith('tbl'):
                table = Table(element, doc)
                rows = table.rows
                if len(rows) >= 2:
                    header_cell_text = rows[0].cells[0].text
                    match = tag_pattern.search(header_cell_text)
                    if match:
                        tag_type = match.group(1)
                        tag_id = match.group(2)
                        # Strip the hidden tag from visible text
                        visible_header = tag_pattern.sub('', header_cell_text).strip()
                        # Collect response text from ALL rows after the header
                        # (handles cases where Enter key or word processors add extra rows)
                        response_parts = []
                        for r_idx in range(1, len(rows)):
                            cell_text = rows[r_idx].cells[0].text.strip()
                            if cell_text:
                                response_parts.append(cell_text)
                        response = '\n'.join(response_parts)
                        # Strip placeholder text from response cell
                        response = response.replace("Type your answer here...", "").replace("Your Answer:", "").strip()

                        # Fix: Detect student text typed in header cell (row 0)
                        # Students sometimes type in the blue header cell instead of the white response cell
                        response_cleaned = re.sub(r'[_\s]', '', response)
                        if len(response_cleaned) < 2:
                            header_cell_obj = rows[0].cells[0]
                            header_paras = [p.text.strip() for p in header_cell_obj.paragraphs]

                            # Case 1: Student pressed Enter in header cell — extra paragraphs
                            if len(header_paras) > 1:
                                extra_text = '\n'.join(p for p in header_paras[1:] if p)
                                if extra_text and len(re.sub(r'[_\s]', '', extra_text)) >= 2:
                                    response = extra_text
                                    visible_header = tag_pattern.sub('', header_paras[0]).strip()

                            # Case 2: Student typed on same line after (N pts) — no Enter
                            if len(re.sub(r'[_\s]', '', response)) < 2:
                                pts_match = re.search(r'\(\d+\s*pts?\)', header_cell_text)
                                if pts_match:
                                    after_pts = header_cell_text[pts_match.end():].strip()
                                    if after_pts and len(re.sub(r'[_\s]', '', after_pts)) >= 2:
                                        response = after_pts
                                        visible_header = tag_pattern.sub('', header_cell_text[:pts_match.end()]).strip()

                            # Case 3: SUMMARY (no pts marker) — check for non-bold student runs
                            if len(re.sub(r'[_\s]', '', response)) < 2 and tag_type == "SUMMARY":
                                student_text_parts = []
                                for para in header_cell_obj.paragraphs:
                                    for run in para.runs:
                                        # Header runs are bold or tiny (1pt tag) or 9pt (pts)
                                        is_tag_run = run.font.size is not None and run.font.size <= Pt(2)
                                        is_pts_run = run.font.size is not None and run.font.size == Pt(9)
                                        is_header_run = run.bold is True
                                        if not is_tag_run and not is_pts_run and not is_header_run:
                                            txt = run.text.strip()
                                            if txt:
                                                student_text_parts.append(txt)
                                if student_text_parts:
                                    candidate = ' '.join(student_text_parts).strip()
                                    if len(re.sub(r'[_\s]', '', candidate)) >= 2:
                                        response = candidate

                        # Deduplicate: skip if we already have this tag (duplicate tables in doc)
                        existing = next(
                            (t for t in tables_found
                             if t["tag_type"] == tag_type and t["tag_id"] == tag_id),
                            None
                        )
                        if existing:
                            # Keep the version with the longer response (more likely to be correct)
                            if len(response) > len(existing.get("response", "")):
                                existing["response"] = response
                                existing["header_text"] = visible_header
                        else:
                            tables_found.append({
                                "tag_type": tag_type,
                                "tag_id": tag_id,
                                "header_text": visible_header,
                                "response": response
                            })
                        # Also add to plain text for content reference
                        full_text.append(visible_header)
                        if response:
                            full_text.append(response)
                        continue

                # Non-Graider table — add as plain text
                for row in rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))

        is_graider = has_marker or len(tables_found) > 0
        plain_text = '\n'.join(full_text)

        if is_graider:
            print(f"  📊 Detected Graider table format: {len(tables_found)} structured tables")

        return {
            "is_graider_table": is_graider,
            "plain_text": plain_text,
            "tables": tables_found
        }
    except Exception as e:
        print(f"  ⚠️  Error in structured read: {e}")
        return {"is_graider_table": False, "plain_text": None, "tables": []}


def extract_from_tables(table_data, exclude_markers=None):
    """Extract student responses from Graider structured table data.

    Maps table entries to the same dict format as extract_student_responses().

    Args:
        table_data: List of dicts from read_docx_file_structured()["tables"]
        exclude_markers: List of section names to skip

    Returns:
        Same shape as extract_student_responses():
        {
            "extracted_responses": [{"question": ..., "answer": ..., "type": ...}],
            "blank_questions": [...],
            "total_questions": int,
            "answered_questions": int,
            "extraction_summary": str
        }
    """
    extracted = []
    blank_questions = []
    excluded_sections = []
    exclude_lower = [em.lower().strip() for em in exclude_markers] if exclude_markers else []

    type_map = {
        "VOCAB": "vocab_term",
        "QUESTION": "numbered_question",
        "SUMMARY": "summary"
    }

    for entry in table_data:
        tag_type = entry.get("tag_type", "")
        tag_id = entry.get("tag_id", "")
        header = entry.get("header_text", "")
        response = entry.get("response", "")
        # Strip placeholder text from response
        response = response.replace("Type your answer here...", "").replace("Your Answer:", "").strip()

        # Check exclusion
        header_lower = header.lower()
        is_excluded = any(em in header_lower for em in exclude_lower)
        if is_excluded:
            excluded_sections.append(header)
            continue

        # Build question label
        if tag_type == "VOCAB":
            question = tag_id  # The term name
        elif tag_type == "QUESTION":
            question = header  # e.g. "1) What is photosynthesis?"
        elif tag_type == "SUMMARY":
            question = "Summary"
        else:
            question = header

        # Check if blank
        response_cleaned = re.sub(r'[_\s]', '', response)
        if len(response_cleaned) < 2:
            # Backup: Student may have typed answer in the header cell after (N pts)
            # read_docx_file_structured has Cases 1-3 for this, but they can fail
            # when Word/Google Docs alters the cell structure
            pts_match = re.search(r'\(\d+\s*pts?\)', header)
            if pts_match:
                after_pts = header[pts_match.end():].strip()
                if after_pts and len(re.sub(r'[_\s]', '', after_pts)) >= 2:
                    response = after_pts
                    header = header[:pts_match.end()].strip()
                    # Rebuild question label with trimmed header
                    if tag_type == "QUESTION":
                        question = header
                    print(f"    🔧 Recovered same-line answer from header for {tag_type}:{tag_id}")

            # Summary fallback: no (N pts) marker — check if prompt is followed by student text
            # Look for sentence-ending punctuation followed by student text
            if len(re.sub(r'[_\s]', '', response)) < 2 and tag_type == "SUMMARY":
                # Try splitting after the prompt (usually ends with a period or question mark)
                # Match the last sentence-ending punctuation that's part of the prompt
                prompt_end = re.search(r'[.?!]\s{2,}', header)
                if prompt_end:
                    after_prompt = header[prompt_end.end():].strip()
                    if after_prompt and len(re.sub(r'[_\s]', '', after_prompt)) >= 10:
                        response = after_prompt
                        header = header[:prompt_end.start() + 1].strip()
                        question = "Summary"
                        print(f"    🔧 Recovered same-line answer from summary header")

        # After recovery attempts, final blank check
        response_cleaned = re.sub(r'[_\s]', '', response)
        if len(response_cleaned) < 2:
            # Deduplicate: don't add if this question is already in blank_questions
            if question not in blank_questions:
                blank_questions.append(question)
            continue

        # Deduplicate: don't add if we already have an extracted response for this tag
        already_extracted = any(
            e.get("tag_id") == tag_id and e.get("section") == tag_type
            for e in extracted
        )
        if already_extracted:
            continue

        extracted.append({
            "question": question,
            "answer": response,
            "type": type_map.get(tag_type, "numbered_question"),
            "section": tag_type,
            "tag_id": tag_id
        })

    # Remove from blank_questions any items that were successfully extracted
    extracted_tags = {(e.get("section"), e.get("tag_id")) for e in extracted}
    # Also build a set of extracted question labels for cross-reference
    extracted_questions = {e.get("question", "").lower().strip() for e in extracted}
    blank_questions = [
        bq for bq in blank_questions
        if bq.lower().strip() not in extracted_questions
    ]

    total_q = len(extracted) + len(blank_questions)
    answered_q = len(extracted)
    summary = "Table extraction: Found " + str(answered_q) + " responses out of " + str(total_q) + " sections."
    if blank_questions:
        summary += " " + str(len(blank_questions)) + " left blank."
    if excluded_sections:
        summary += " Excluded " + str(len(excluded_sections)) + " section(s)."

    print(f"  📊 Table extraction: {answered_q}/{total_q} answered")

    return {
        "extracted_responses": extracted,
        "blank_questions": blank_questions,
        "total_questions": max(total_q, 1),
        "answered_questions": answered_q,
        "extraction_summary": summary,
        "excluded_sections": excluded_sections,
        "missing_sections": []  # Table format has no separate missing sections (avoids double-counting with blank_questions)
    }


def extract_from_graider_text(document_text, exclude_markers=None):
    """Extract student responses from plain text containing [GRAIDER:TYPE:ID] markers.

    Fallback for when structured table reading fails (e.g., tables were flattened
    by Google Docs, copy-paste, or format conversion).  Parses the text between
    consecutive GRAIDER markers to capture student answers.

    Args:
        document_text: Plain text that may contain [GRAIDER:...] tags.
        exclude_markers: List of section names to skip.

    Returns:
        Same shape as extract_from_tables(), or None if no GRAIDER tags found.
    """
    tag_pattern = re.compile(r'\[GRAIDER:(VOCAB|QUESTION|SUMMARY):([^\]]+)\]')
    matches = list(tag_pattern.finditer(document_text))

    if not matches:
        return None

    print(f"  📝 Graider text fallback: Found {len(matches)} GRAIDER markers in plain text")

    extracted = []
    blank_questions = []
    excluded_sections = []
    exclude_lower = [em.lower().strip() for em in exclude_markers] if exclude_markers else []

    type_map = {
        "VOCAB": "vocab_term",
        "QUESTION": "numbered_question",
        "SUMMARY": "summary"
    }

    for i, match in enumerate(matches):
        tag_type = match.group(1)
        tag_id = match.group(2)

        # Text after the tag up to the next tag (or GRAIDER_TABLE_V1 marker or end of doc)
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(document_text)
        raw_block = document_text[start:end]
        # Truncate at GRAIDER_TABLE_V1 marker if present (end-of-worksheet sentinel)
        marker_pos = raw_block.find('GRAIDER_TABLE_V1')
        if marker_pos != -1:
            raw_block = raw_block[:marker_pos]
        raw_block = raw_block.strip()

        # The block contains: "  visible_header  (N pts)\nstudent answer\n..."
        # Split into lines, skip the header line (contains the term/question + pts),
        # and grab everything else as the student response.
        lines = [ln.strip() for ln in raw_block.split('\n') if ln.strip()]

        # First non-empty line is typically the visible header (term + pts)
        header = lines[0] if lines else ""
        # Student response is everything after the header line,
        # filtering out section headers and metadata that appear between markers
        section_headers = {'vocabulary', 'questions', 'summary', 'question', 'vocab'}
        response_lines = []
        for ln in lines[1:]:
            # Skip section headers and metadata
            if ln.lower() in section_headers:
                continue
            if 'GRAIDER_TABLE_V1' in ln:
                continue
            response_lines.append(ln)
        response = '\n'.join(response_lines).strip()
        # Strip placeholder text from response cell
        response = response.replace("Type your answer here...", "").replace("Your Answer:", "").strip()

        # Fix: Student typed answer on same line as header (no newline separation)
        # Common when student types in the header cell of a Graider table
        if not response or len(re.sub(r'[_\s]', '', response)) < 2:
            pts_match = re.search(r'\(\d+\s*pts?\)', header)
            if pts_match:
                after_pts = header[pts_match.end():].strip()
                if after_pts and len(re.sub(r'[_\s]', '', after_pts)) >= 2:
                    response = after_pts
                    header = header[:pts_match.end()].strip()
                    print(f"    🔧 Text fallback: recovered same-line answer for {tag_type}:{tag_id}")

            # Summary fallback: no (N pts) marker — check if prompt is followed by student text
            if (not response or len(re.sub(r'[_\s]', '', response)) < 2) and tag_type == "SUMMARY":
                prompt_end = re.search(r'[.?!]\s{2,}', header)
                if prompt_end:
                    after_prompt = header[prompt_end.end():].strip()
                    if after_prompt and len(re.sub(r'[_\s]', '', after_prompt)) >= 10:
                        response = after_prompt
                        header = header[:prompt_end.start() + 1].strip()
                        print(f"    🔧 Text fallback: recovered same-line answer for SUMMARY")

        # Check exclusion
        header_lower = header.lower()
        is_excluded = any(em in header_lower for em in exclude_lower)
        if is_excluded:
            excluded_sections.append(header)
            continue

        # Build question label
        if tag_type == "VOCAB":
            question = tag_id
        elif tag_type == "QUESTION":
            question = header
        elif tag_type == "SUMMARY":
            question = "Summary"
        else:
            question = header

        # Check if blank
        response_cleaned = re.sub(r'[_\s]', '', response)
        if len(response_cleaned) < 2:
            # Deduplicate
            if question not in blank_questions:
                blank_questions.append(question)
            continue

        # Deduplicate: skip if already extracted for this tag
        already_extracted = any(
            e.get("tag_id") == tag_id and e.get("section") == tag_type
            for e in extracted
        )
        if already_extracted:
            continue

        extracted.append({
            "question": question,
            "answer": response,
            "type": type_map.get(tag_type, "numbered_question"),
            "section": tag_type,
            "tag_id": tag_id
        })

    # Cross-reference: remove blank entries that were successfully extracted
    extracted_questions = {e.get("question", "").lower().strip() for e in extracted}
    blank_questions = [
        bq for bq in blank_questions
        if bq.lower().strip() not in extracted_questions
    ]

    total_q = len(extracted) + len(blank_questions)
    answered_q = len(extracted)
    summary = f"Graider text fallback: Found {answered_q} responses out of {total_q} sections."
    if blank_questions:
        summary += f" {len(blank_questions)} left blank."

    print(f"  📝 Graider text extraction: {answered_q}/{total_q} answered")

    return {
        "extracted_responses": extracted,
        "blank_questions": blank_questions,
        "total_questions": max(total_q, 1),
        "answered_questions": answered_q,
        "extraction_summary": summary,
        "excluded_sections": excluded_sections,
        "missing_sections": []  # Text fallback has no separate missing sections (avoids double-counting with blank_questions)
    }


def read_image_file(filepath: str) -> dict:
    """
    Read an image file and return it as base64 for GPT-4o vision.
    
    Returns dict with:
    - type: "image"
    - data: base64 encoded image
    - media_type: image MIME type
    """
    import base64
    
    filepath = Path(filepath)
    extension = filepath.suffix.lower()
    
    # Map extensions to MIME types
    mime_types = {
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.webp': 'image/webp'
    }
    
    if extension not in mime_types:
        print(f"  ⚠️  Unsupported image type: {extension}")
        return None
    
    try:
        with open(filepath, 'rb') as f:
            image_data = base64.b64encode(f.read()).decode('utf-8')
        
        return {
            "type": "image",
            "data": image_data,
            "media_type": mime_types[extension]
        }
    except Exception as e:
        print(f"  ⚠️  Error reading image: {e}")
        return None


def read_assignment_file(filepath: str) -> dict:
    """
    Read assignment file based on its extension.
    Supports: .docx, .txt, .jpg, .jpeg, .png, .gif, .webp
    
    Returns dict with:
    - type: "text" or "image"
    - content: text content or base64 image data
    """
    filepath = Path(filepath)
    extension = filepath.suffix.lower()
    
    # Text-based files
    if extension == '.docx':
        # Try structured table reading first (Graider-generated worksheets)
        structured = read_docx_file_structured(filepath)
        if structured.get("is_graider_table") and structured.get("tables"):
            content = structured.get("plain_text", "")
            if "GRAIDER_ANSWER_KEY_START" in content:
                content = content.split("GRAIDER_ANSWER_KEY_START")[0].rstrip().rstrip('-')
            return {
                "type": "text",
                "content": content,
                "graider_tables": structured["tables"]
            }

        # Fallback to standard text reading
        content = read_docx_file(filepath)
        if content:
            # Strip embedded answer key from generated worksheets (handles -- and --- variants)
            if "GRAIDER_ANSWER_KEY_START" in content:
                content = content.split("GRAIDER_ANSWER_KEY_START")[0].rstrip().rstrip('-')
                print(f"  🧹 Stripped embedded answer key at file read")
            return {"type": "text", "content": content}
        return None
    
    elif extension == '.txt':
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return {"type": "text", "content": f.read()}
        except Exception as e:
            print(f"  ⚠️  Error reading text file: {e}")
            return None
    
    # Image files
    elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
        image_data = read_image_file(filepath)
        if image_data:
            return {
                "type": "image",
                "content": image_data["data"],
                "media_type": image_data["media_type"]
            }
        return None
    
    else:
        print(f"  ⚠️  Unsupported file type: {extension}")
        return None


def extract_student_work(content: str) -> tuple:
    """
    Extract only the student work portions of the document.

    Finds the first student work marker and returns everything after it.
    This ensures we capture student responses without duplicating content.

    Returns: (student_work, markers_found)
    - student_work: The extracted student content (or full content if no marker)
    - markers_found: List of markers that were found
    """
    content_lower = content.lower()

    # Find the earliest marker position
    found_markers = []
    earliest_pos = len(content)
    earliest_marker = None

    for marker in STUDENT_WORK_MARKERS:
        marker_lower = marker.lower()
        pos = content_lower.find(marker_lower)
        if pos != -1:
            if marker not in found_markers:
                found_markers.append(marker)
            if pos < earliest_pos:
                earliest_pos = pos
                earliest_marker = marker

    if not earliest_marker:
        # No markers found - return full content
        return content, []

    # Return everything from the first marker onward
    # Find the line containing the marker and start from there
    line_start = content.rfind('\n', 0, earliest_pos)
    if line_start == -1:
        line_start = 0
    else:
        line_start += 1  # Skip the newline character

    student_work = content[line_start:].strip()
    return student_work, found_markers


# =============================================================================
# FERPA COMPLIANCE - PII SANITIZATION
# =============================================================================

import hashlib

def sanitize_pii_for_ai(student_name: str, content: str) -> tuple:
    """
    FERPA Compliance: Remove all Personally Identifiable Information (PII)
    before sending student work to external AI services.

    Returns:
        tuple: (anonymous_id, sanitized_content)
    """
    if not content:
        return "Student_0000", ""

    # Create consistent anonymous identifier from student name
    if student_name:
        hash_val = int(hashlib.md5(student_name.encode()).hexdigest(), 16) % 10000
        anon_id = f"Student_{hash_val:04d}"
    else:
        anon_id = "Student_0000"

    sanitized = content

    # Remove student name variations (first name, last name, full name)
    if student_name:
        name_parts = student_name.split()
        for part in name_parts:
            if len(part) > 2:  # Avoid removing short words like "I" or "A"
                sanitized = re.sub(
                    rf'\b{re.escape(part)}\b',
                    '[STUDENT]',
                    sanitized,
                    flags=re.IGNORECASE
                )

    # Remove Social Security Numbers (XXX-XX-XXXX)
    sanitized = re.sub(r'\b\d{3}-\d{2}-\d{4}\b', '[SSN-REMOVED]', sanitized)

    # Remove Student ID numbers (7-10 digit numbers that look like IDs)
    sanitized = re.sub(r'\b\d{7,10}\b', '[ID-REMOVED]', sanitized)

    # Remove email addresses
    sanitized = re.sub(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', '[EMAIL-REMOVED]', sanitized)

    # Remove phone numbers (various formats)
    sanitized = re.sub(r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b', '[PHONE-REMOVED]', sanitized)
    sanitized = re.sub(r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}', '[PHONE-REMOVED]', sanitized)

    # Remove dates that might be birthdates (MM/DD/YYYY, MM-DD-YYYY, etc.)
    # But preserve historical dates (years before 2000 are likely historical)
    sanitized = re.sub(r'\b\d{1,2}[/-]\d{1,2}[/-](20\d{2}|19[89]\d)\b', '[DATE-REMOVED]', sanitized)

    # Remove street addresses (basic pattern)
    sanitized = re.sub(
        r'\b\d+\s+[\w\s]+(?:street|st|avenue|ave|road|rd|drive|dr|lane|ln|way|court|ct|boulevard|blvd|circle|cir|place|pl)\.?\b',
        '[ADDRESS-REMOVED]',
        sanitized,
        flags=re.IGNORECASE
    )

    # Remove zip codes (5 digit or 5+4 format)
    sanitized = re.sub(r'\b\d{5}(-\d{4})?\b', '[ZIP-REMOVED]', sanitized)

    return anon_id, sanitized


def log_pii_sanitization(student_name: str, original_len: int, sanitized_len: int, removals: dict):
    """
    Log PII sanitization actions for audit purposes.
    Does not log actual PII - only counts and types of removals.
    """
    # This could be extended to write to an audit log file
    if any(removals.values()):
        print(f"  🔒 PII sanitized for student submission (removed: {', '.join(k for k, v in removals.items() if v > 0)})")


# =============================================================================
# AI/PLAGIARISM DETECTION (Parallel Agent using GPT-4o-mini)
# =============================================================================

def preprocess_for_ai_detection(text: str) -> str:
    """
    Preprocess extracted text to focus on student-written content for AI detection.

    - Extracts fill-in-the-blank answers (text between underscores)
    - Removes template/instructional text
    - Focuses on longer written responses, not factual answers

    Returns cleaned text suitable for AI detection analysis.
    Short fill-in-blank answers (dates, names, single words) are excluded since
    they're factual answers, not writing that can be AI-detected.
    """
    import re

    lines = text.split('\n')
    student_written = []
    fill_in_answers = []

    # Patterns that indicate template/instructional text (not student writing)
    template_patterns = [
        r'^Q:\s*',  # Question prefix
        r'^A:\s*$',  # Empty answer prefix
        r'^The .+ happened in the year',
        r'^The U\.S\. bought',
        r'^The purchase doubled',
        r'^The price paid',
        r'^This purchase helped',
        r'^Explain in your own words',
        r'^Write a \d+',
        r'^How do you think .+\?$',  # Questions ending with ?
        r'^Why was the .+\?$',
        r'^Who was the .+\?$',
        r'^What role did .+\?$',
        r'^How did .+\?$',
        r'^\d+\.\s*(Why|What|Who|How|Where|When)\s+.+\?$',  # Numbered questions
        r'^Primary Source Quote',
        r'^Quote from',
        r'^"The acquisition of',  # Known quote text
        r'^Student Task:',
        r'^Final Reflection Question',
        r'^Questions to Check Understanding',
        r'^Guided Notes',
        r'^Cornell Notes',
        r'^Essential Question',
        r'^Vocabulary',
    ]

    for line in lines:
        line_stripped = line.strip()
        if not line_stripped:
            continue

        # Extract fill-in-the-blank answers (text between underscores)
        blank_matches = re.findall(r'_{2,}([^_\n]{1,100})_{2,}', line)
        if blank_matches:
            for match in blank_matches:
                answer = match.strip()
                if answer:
                    fill_in_answers.append(answer)
            continue  # Rest of line is template

        # Strip "A: " prefix before checking template patterns so patterns match
        check_text = line_stripped
        if check_text.startswith('A: '):
            check_text = check_text[3:].strip()

        # Check if this line matches a template pattern
        is_template = False
        for pattern in template_patterns:
            if re.match(pattern, check_text, re.IGNORECASE):
                is_template = True
                break

        if is_template:
            continue

        # Skip lines that are just "A: " followed by a question (unanswered)
        if line_stripped.startswith('A: ') and line_stripped.endswith('?'):
            continue

        # Skip Q: labels
        if line_stripped.startswith('Q: '):
            continue

        # Skip vocabulary definitions from template (Term | Definition format)
        if ' | ' in check_text:
            continue

        # Skip quoted text (primary source quotes from the template)
        if check_text.startswith('"') and check_text.endswith('"'):
            continue

        # Skip lines that are template fill-in-blank prompts (not the student's answer)
        # These have the structure "The ___ did something ___" with underscores
        underscore_count = check_text.count('_')
        non_underscore = re.sub(r'_+', '', check_text).strip()
        if underscore_count >= 3 and len(non_underscore) > 20:
            # Line is mostly template text with blank slots — skip it
            continue

        # Keep substantive student-written content (paragraphs, explanations)
        # Must be more than 30 chars to be considered "writing" vs labels
        # Exclude lines that are questions (end with ?) or instruction text
        if len(check_text) > 30 and not check_text.endswith('?'):
            # Skip lines that look like unanswered template text
            instruction_keywords = ['write a few sentences', 'explain in your own words',
                                   'how do you think', 'why do you think', 'what do you think',
                                   'describe how', 'explain how', 'explain why',
                                   'student task:', 'final reflection']
            is_instruction = any(kw in check_text.lower() for kw in instruction_keywords)
            if not is_instruction:
                student_written.append(check_text)

    # Build result
    result_parts = []

    # Only flag fill-in answers if they're suspiciously long (might be copied text)
    # Short answers like "1803", "France", "15 million" are factual, not AI-detectable
    suspicious_fill_ins = [a for a in fill_in_answers if len(a) > 40]
    if suspicious_fill_ins:
        result_parts.append("Suspiciously long fill-in answers:\n" + "\n".join(suspicious_fill_ins))

    # Student written content (main focus for AI detection)
    if student_written:
        result_parts.append("Student written content:\n" + "\n".join(student_written))

    # If no substantial content to analyze, return empty (skip AI detection)
    if not result_parts:
        return ""

    return "\n\n".join(result_parts)


def detect_ai_plagiarism(student_responses: str, grade_level: str = '6', token_tracker: 'TokenTracker' = None) -> dict:
    """
    Dedicated AI/Plagiarism detection using GPT-4o-mini.
    Runs in parallel with grading for speed.

    Returns:
        {
            "ai_detection": {"flag": "none/unlikely/possible/likely", "confidence": 0-100, "reason": "..."},
            "plagiarism_detection": {"flag": "none/possible/likely", "reason": "..."}
        }
    """
    # Skip detection if no substantial content to analyze
    # (e.g., fill-in-the-blank only assignments with short factual answers)
    if not student_responses or len(student_responses.strip()) < 50:
        return {
            "ai_detection": {"flag": "none", "confidence": 0, "reason": "Fill-in-the-blank or short-answer only - exempt from AI detection"},
            "plagiarism_detection": {"flag": "none", "reason": "Fill-in-the-blank or short-answer only - exempt from plagiarism detection"}
        }

    try:
        from openai import OpenAI
    except ImportError:
        return {"ai_detection": {"flag": "none", "confidence": 0, "reason": ""},
                "plagiarism_detection": {"flag": "none", "reason": ""}}

    client = OpenAI(api_key=OPENAI_API_KEY)

    age_range = "11-12" if grade_level == "6" else "12-13" if grade_level == "7" else "13-14"

    detection_prompt = f"""You are an expert at detecting AI-generated content and plagiarism in student work.
Analyze this grade {grade_level} student's responses ({age_range} years old) for signs of AI use or copy-paste.

DETECTION CRITERIA:

1. AI-GENERATED CONTENT - Flag as "likely" if you see:
- Sophisticated vocabulary a {age_range} year old wouldn't use: "ideology", "emphasizing", "implementing", "interconnected", "fostering", "trajectory"
- Perfect grammar and sentence structure throughout
- Academic/formal tone: "This demonstrates...", "It is important to note...", "Furthermore..."
- Phrases like: "transforming a limited mission", "securing vital trade routes", "fundamentally altered"

2. PLAGIARISM/COPY-PASTE - Flag as "likely" if you see:
- Dictionary-perfect definitions (e.g., "exclusive possession or control of the supply of or trade in a commodity or service" for monopoly)
- Wikipedia-style language that doesn't match a child's writing
- Textbook phrases copied verbatim
- Sudden shifts between simple spelling errors and sophisticated paragraphs

3. CONTRAST CHECK (MOST IMPORTANT):
- If student writes simple answers like "it made the US bigger", "idk", misspells words
- BUT also writes "an ideology emphasizing intense loyalty to one's nation"
- That is 100% copy-paste - flag PLAGIARISM as "likely"

REAL {age_range} YEAR OLD WRITES:
- "when one company controls everything" (for monopoly)
- "it helped them trade stuff"
- "so boats could go there"

AI/COPIED TEXT LOOKS LIKE:
- "exclusive possession or control of the supply of or trade in a commodity or service"
- "government-provided financial incentives"
- "implementing three interconnected policies"

STUDENT RESPONSES TO ANALYZE:
{student_responses}

Respond ONLY with this JSON (no other text):
{{
    "ai_detection": {{
        "flag": "<none, unlikely, possible, or likely>",
        "confidence": <0-100>,
        "reason": "<specific phrases that triggered the flag, or empty string if none>"
    }},
    "plagiarism_detection": {{
        "flag": "<none, possible, or likely>",
        "reason": "<specific copied phrases found, or empty string if none>"
    }}
}}"""

    try:
        # Use structured output for guaranteed schema
        try:
            response = client.beta.chat.completions.parse(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": detection_prompt}],
                response_format=DetectionResponse,
                max_tokens=500,
                temperature=0,
                seed=42
            )
            if token_tracker:
                token_tracker.record_openai(response, "gpt-4o-mini")
            parsed = response.choices[0].message.parsed
            if parsed:
                return parsed.model_dump()
        except Exception:
            pass  # Fall through to text fallback

        # Text fallback if structured output fails
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": detection_prompt}],
            max_tokens=500,
            temperature=0,
            seed=42
        )
        if token_tracker:
            token_tracker.record_openai(response, "gpt-4o-mini")
        response_text = response.choices[0].message.content.strip()
        result = _try_parse_json_fallback(response_text)
        if result:
            return result

        return json.loads(response_text)

    except Exception as e:
        print(f"  ⚠️  Detection error: {e}")
        return {
            "ai_detection": {"flag": "none", "confidence": 0, "reason": ""},
            "plagiarism_detection": {"flag": "none", "reason": ""}
        }


def grade_with_ensemble(student_name: str, assignment_data: dict, custom_ai_instructions: str = '',
                        grade_level: str = '6', subject: str = 'Social Studies',
                        ensemble_models: list = None, student_id: str = None,
                        assignment_template: str = None, rubric_prompt: str = None,
                        custom_markers: list = None, exclude_markers: list = None,
                        marker_config: list = None, effort_points: int = 15,
                        extraction_mode: str = 'structured', grading_style: str = 'standard') -> dict:
    """
    Grade assignment using multiple AI models and combine results.

    Args:
        ensemble_models: List of model names to use (e.g., ['gpt-4o-mini', 'claude-haiku', 'gemini-flash'])
        marker_config: List of marker configs with points for section-based grading
        effort_points: Points for effort/engagement category (default 15)

    Returns:
        Combined result with:
        - score: Average of all model scores
        - ensemble_scores: Individual scores from each model
        - ensemble_feedback: Feedback from each model
        - letter_grade: Based on averaged score
    """
    if not ensemble_models or len(ensemble_models) < 2:
        # Fall back to single model grading
        model = ensemble_models[0] if ensemble_models else 'gpt-4o-mini'
        return grade_with_parallel_detection(student_name, assignment_data, custom_ai_instructions,
                                             grade_level, subject, model, student_id,
                                             assignment_template, rubric_prompt, custom_markers, exclude_markers,
                                             marker_config, effort_points, extraction_mode, grading_style)

    print(f"  🎯 Ensemble grading with {len(ensemble_models)} models: {', '.join(ensemble_models)}")

    # Run all models in parallel
    results = {}
    with concurrent.futures.ThreadPoolExecutor(max_workers=len(ensemble_models)) as executor:
        futures = {}
        for model in ensemble_models:
            future = executor.submit(
                grade_assignment, student_name, assignment_data, custom_ai_instructions,
                grade_level, subject, model, student_id, assignment_template, rubric_prompt,
                custom_markers, exclude_markers, marker_config, effort_points, extraction_mode,
                grading_style
            )
            futures[future] = model

        # Collect results
        for future in concurrent.futures.as_completed(futures):
            model = futures[future]
            try:
                result = future.result()
                results[model] = result
                print(f"    ✓ {model}: {result.get('score', 0)} ({result.get('letter_grade', 'N/A')})")
            except Exception as e:
                print(f"    ✗ {model}: Error - {str(e)[:50]}")
                results[model] = {"score": 0, "letter_grade": "ERROR", "feedback": f"Error: {e}"}

    # Calculate ensemble score
    valid_scores = [r.get('score', 0) for r in results.values() if r.get('letter_grade') != 'ERROR']
    if not valid_scores:
        return {"score": 0, "letter_grade": "ERROR", "feedback": "All models failed", "breakdown": {}}

    # Use median for robustness against outliers
    valid_scores.sort()
    if len(valid_scores) % 2 == 0:
        median_score = (valid_scores[len(valid_scores)//2 - 1] + valid_scores[len(valid_scores)//2]) / 2
    else:
        median_score = valid_scores[len(valid_scores)//2]

    avg_score = sum(valid_scores) / len(valid_scores)
    final_score = round(median_score)  # Use median as final

    # Determine letter grade
    if final_score >= 90:
        letter_grade = "A"
    elif final_score >= 80:
        letter_grade = "B"
    elif final_score >= 70:
        letter_grade = "C"
    elif final_score >= 60:
        letter_grade = "D"
    else:
        letter_grade = "F"

    # Pick the best feedback (from the model closest to median score)
    best_model = min(results.keys(), key=lambda m: abs(results[m].get('score', 0) - median_score) if results[m].get('letter_grade') != 'ERROR' else 999)
    best_result = results[best_model]

    # Build ensemble result
    ensemble_result = {
        "score": final_score,
        "letter_grade": letter_grade,
        "feedback": best_result.get("feedback", ""),
        "breakdown": best_result.get("breakdown", {}),
        "ai_detection": best_result.get("ai_detection", {}),
        "plagiarism_detection": best_result.get("plagiarism_detection", {}),
        "student_responses": best_result.get("student_responses", []),
        "unanswered_questions": best_result.get("unanswered_questions", []),
        # Ensemble-specific fields
        "ensemble_grading": True,
        "ensemble_models": ensemble_models,
        "ensemble_scores": {model: results[model].get("score", 0) for model in results},
        "ensemble_grades": {model: results[model].get("letter_grade", "N/A") for model in results},
        "ensemble_avg": round(avg_score, 1),
        "ensemble_median": round(median_score, 1),
        "ensemble_feedback_source": best_model,
    }

    print(f"  📊 Ensemble result: {final_score} ({letter_grade}) - avg: {avg_score:.1f}, median: {median_score:.1f}")

    return ensemble_result


def grade_with_parallel_detection(student_name: str, assignment_data: dict, custom_ai_instructions: str = '',
                                   grade_level: str = '6', subject: str = 'Social Studies',
                                   ai_model: str = 'gpt-4o-mini', student_id: str = None,
                                   assignment_template: str = None, rubric_prompt: str = None,
                                   custom_markers: list = None, exclude_markers: list = None,
                                   marker_config: list = None, effort_points: int = 15,
                                   extraction_mode: str = 'structured', grading_style: str = 'standard',
                                   student_history: str = '') -> dict:
    """
    Grade assignment with parallel AI/plagiarism detection.
    Runs detection (GPT-4o-mini) and grading simultaneously for speed.

    Args:
        rubric_prompt: Custom rubric prompt string from Settings (overrides default)
        marker_config: List of marker configs with points for section-based grading
        effort_points: Points for effort/engagement category (default 15)
    """
    # Extract responses first (needed for both detection and grading context)
    content = assignment_data.get("content", "")
    extracted_text = ""

    # Strip embedded answer key before extraction (handles -- and --- variants)
    if content and "GRAIDER_ANSWER_KEY_START" in content:
        content = content.split("GRAIDER_ANSWER_KEY_START")[0].rstrip().rstrip('-')
        assignment_data = {**assignment_data, "content": content}
        print(f"  🧹 Stripped embedded answer key from document")

    # Priority: Graider structured tables > Graider text fallback > regex extraction
    extraction_result = None
    graider_tables = assignment_data.get("graider_tables")
    if graider_tables:
        print(f"  📊 Parallel detection: Using Graider table extraction ({len(graider_tables)} tables)")
        extraction_result = extract_from_tables(graider_tables, exclude_markers)
    elif assignment_data.get("type") == "text" and content:
        # Try GRAIDER tag plain-text fallback before generic extraction
        if '[GRAIDER:' in content:
            extraction_result = extract_from_graider_text(content, exclude_markers)
        if not extraction_result or not extraction_result.get("extracted_responses"):
            extraction_result = extract_student_responses(content, custom_markers, exclude_markers, assignment_template)

    if extraction_result and extraction_result.get("extracted_responses"):
        # Filter out FITB and vocab items — only send written responses to detection
        # FITB answers (names, dates, facts) and vocab definitions naturally match sources
        skip_types = {'fitb_full', 'vocab_term'}
        written_responses = [
            r for r in extraction_result["extracted_responses"]
            if 'fill_in_blank' not in r.get('type', '') and r.get('type') not in skip_types
        ]
        # If no written responses (pure FITB), extracted_text stays empty → detection skipped
        extracted_text = "\n".join([
            f"Q: {r.get('question', 'Unknown')}\nA: {r.get('answer', '')}"
            for r in written_responses
        ])

    # If no extracted text, can't do parallel detection
    if not extracted_text:
        return grade_assignment(student_name, assignment_data, custom_ai_instructions,
                               grade_level, subject, ai_model, student_id, assignment_template, rubric_prompt,
                               custom_markers, exclude_markers, marker_config, effort_points, extraction_mode,
                               grading_style=grading_style)

    # Multi-pass grading for all providers
    use_multipass = True
    print(f"  🔄 Running parallel detection + multi-pass grading ({ai_model})...")

    # Preprocess text for AI detection (removes template text, focuses on student writing)
    detection_text = preprocess_for_ai_detection(extracted_text)

    # Shared token tracker for both detection and grading
    tracker = TokenTracker()

    # Run detection and grading in parallel using ThreadPoolExecutor
    with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
        # Submit both tasks
        detection_future = executor.submit(detect_ai_plagiarism, detection_text, grade_level, token_tracker=tracker)

        if use_multipass:
            grading_future = executor.submit(grade_multipass, student_name, assignment_data,
                                             custom_ai_instructions, grade_level, subject,
                                             ai_model, student_id, assignment_template, rubric_prompt,
                                             custom_markers, exclude_markers, marker_config, effort_points,
                                             extraction_mode, grading_style, token_tracker=tracker,
                                             student_history=student_history)
        else:
            grading_future = executor.submit(grade_assignment, student_name, assignment_data,
                                             custom_ai_instructions, grade_level, subject,
                                             ai_model, student_id, assignment_template, rubric_prompt,
                                             custom_markers, exclude_markers, marker_config, effort_points,
                                             extraction_mode, grading_style, token_tracker=tracker)

        # Wait for both to complete
        detection_result = detection_future.result()
        grading_result = grading_future.result()

    # Skip detection merging for blank/incomplete submissions — nothing to detect
    if grading_result.get("letter_grade") == "INCOMPLETE" and grading_result.get("score", 0) == 0:
        print(f"  📝 Blank/incomplete submission — skipping AI/plagiarism detection merge")
        grading_result["token_usage"] = tracker.summary()
        return grading_result

    # Merge detection results into grading results
    # Detection agent's flags override grading agent's (more specialized)
    detection_ai = detection_result.get("ai_detection", {})
    detection_plag = detection_result.get("plagiarism_detection", {})

    grading_ai = grading_result.get("ai_detection", {})
    grading_plag = grading_result.get("plagiarism_detection", {})

    # Use the more severe flag from either source
    flag_severity = {"none": 0, "unlikely": 1, "possible": 2, "likely": 3}

    # AI detection - take the higher severity
    if flag_severity.get(detection_ai.get("flag", "none"), 0) >= flag_severity.get(grading_ai.get("flag", "none"), 0):
        grading_result["ai_detection"] = detection_ai
        if detection_ai.get("flag") in ["possible", "likely"]:
            print(f"  🤖 Detection agent flagged AI: {detection_ai.get('flag')} - {detection_ai.get('reason', '')[:100]}")

    # Plagiarism detection - take the higher severity
    if flag_severity.get(detection_plag.get("flag", "none"), 0) >= flag_severity.get(grading_plag.get("flag", "none"), 0):
        grading_result["plagiarism_detection"] = detection_plag
        if detection_plag.get("flag") in ["possible", "likely"]:
            print(f"  📋 Detection agent flagged plagiarism: {detection_plag.get('flag')} - {detection_plag.get('reason', '')[:100]}")

    # Apply score caps based on detection flags
    ai_flag = grading_result.get("ai_detection", {}).get("flag", "none")
    plag_flag = grading_result.get("plagiarism_detection", {}).get("flag", "none")
    original_score = grading_result.get("score", 0)

    # Determine cap based on flags
    cap = 100
    cap_reason = ""

    if ai_flag == "likely" and plag_flag == "likely":
        cap = 40
        cap_reason = "AI + Plagiarism detected"
    elif ai_flag == "likely":
        cap = 50
        cap_reason = "Likely AI-generated"
    elif plag_flag == "likely":
        cap = 50
        cap_reason = "Likely plagiarized"
    elif ai_flag == "possible":
        cap = 65
        cap_reason = "Possible AI use"
    elif plag_flag == "possible":
        cap = 65
        cap_reason = "Possible plagiarism"

    # Apply cap if needed
    if original_score > cap:
        grading_result["score"] = cap
        grading_result["score_capped"] = True
        grading_result["original_score"] = original_score
        grading_result["cap_reason"] = cap_reason
        # Update letter grade
        if cap <= 59:
            grading_result["letter_grade"] = "F"
        elif cap <= 69:
            grading_result["letter_grade"] = "D"
        elif cap <= 79:
            grading_result["letter_grade"] = "C"
        print(f"  ⚠️  Score capped: {original_score} → {cap} ({cap_reason})")

    # Replace feedback with academic integrity message for high AI/plagiarism likelihood
    # But NOT for blank submissions (they get their own feedback)
    ai_confidence = grading_result.get("ai_detection", {}).get("confidence", 0)
    student_responses = grading_result.get("student_responses", [])
    ai_score = grading_result.get("score", 0)
    # Only treat as blank if: no student_responses AND AI gave score 0 AND no JSON recovery.
    # Previously this fired whenever student_responses was empty, which could override
    # legitimate grades if the AI simply didn't populate that field.
    is_blank = (not student_responses and not grading_result.get("json_recovery")
                and (ai_score == 0 or ai_score is None))

    if is_blank:
        # Blank submission — zero score, clear feedback, no AI/plagiarism flags
        grading_result["feedback"] = "You submitted a blank assignment with no responses. Please complete all sections and resubmit."
        grading_result["score"] = 0
        grading_result["letter_grade"] = "INCOMPLETE"
        grading_result["ai_detection"] = {"flag": "none", "confidence": 0, "reason": "Blank submission — no content to evaluate."}
        grading_result["plagiarism_detection"] = {"flag": "none", "reason": "Blank submission — no content to evaluate."}
        grading_result.pop("academic_integrity_flag", None)
        print(f"  📝 Blank submission detected — scored as 0/INCOMPLETE")
    elif ai_confidence >= 50 or plag_flag in ["possible", "likely"]:
        grading_result["original_feedback"] = grading_result.get("feedback", "")
        grading_result["feedback"] = "Please resubmit using your own words. Copying and pasting from Google (plagiarism) or use of AI is considered a violation of academic integrity."
        grading_result["academic_integrity_flag"] = True
        print(f"  🚨 Academic integrity concern - feedback replaced")

    # Update token_usage with final tracker summary (includes both detection + grading)
    grading_result["token_usage"] = tracker.summary()

    return grading_result


# =============================================================================
# MULTI-PASS GRADING PIPELINE
# =============================================================================

class QuestionGrade(BaseModel):
    score: int
    possible: int
    reasoning: str
    is_correct: bool
    quality: str  # "excellent", "good", "adequate", "developing", "insufficient"


class PerQuestionResponse(BaseModel):
    grade: QuestionGrade
    excellent: bool
    improvement_note: str


class FeedbackResponse(BaseModel):
    feedback: str
    excellent_answers: List[str]
    needs_improvement: List[str]
    skills_demonstrated: SkillsDemonstrated


def _parse_expected_answers(custom_instructions: str) -> dict:
    """Parse expected answers from gradingNotes/custom instructions.

    Returns dict mapping question index (int) or question text (str) to expected answer.
    """
    answers = {}
    if not custom_instructions:
        return answers

    # Parse "Q1: answer" or "- Q1: answer" patterns
    for match in re.finditer(r'(?:^|\n)\s*-?\s*Q(\d+)\s*:\s*(.+)', custom_instructions):
        idx = int(match.group(1)) - 1  # 0-indexed
        answers[idx] = match.group(2).strip()

    # Parse "VOCABULARY EXPECTED DEFINITIONS:" section
    in_vocab = False
    for line in custom_instructions.split('\n'):
        line = line.strip()
        if 'EXPECTED' in line.upper() and ('DEFINITION' in line.upper() or 'ANSWER' in line.upper()):
            in_vocab = True
            continue
        if in_vocab and line.startswith('- '):
            parts = line[2:].split(':', 1)
            if len(parts) == 2:
                answers[parts[0].strip()] = parts[1].strip()
        elif in_vocab and not line:
            in_vocab = False

    return answers


def _distribute_points(responses: list, marker_config: list, total_points: int) -> list:
    """Distribute point values and section metadata across extracted responses.

    Uses marker_config if available, otherwise distributes evenly.
    Returns list of dicts with 'points', 'section_name', and 'section_type' per response.
    """
    if not responses:
        return []

    # Build lookup: marker_name_lower -> {points, name, type}
    marker_meta = {}
    if marker_config:
        for m in marker_config:
            if isinstance(m, dict):
                marker_meta[m.get('start', '').lower()] = {
                    'points': m.get('points', 10),
                    'name': m.get('start', 'Section'),
                    'type': m.get('type', 'written')
                }
            elif isinstance(m, str):
                marker_meta[m.lower()] = {
                    'points': 10,
                    'name': m,
                    'type': 'written'
                }

    result = []
    default_pts = total_points // max(len(responses), 1)

    for resp in responses:
        question = resp.get("question", "").lower()
        matched = None
        for marker_key, meta in marker_meta.items():
            if marker_key in question:
                matched = meta
                break

        if matched:
            result.append({
                'points': matched['points'],
                'section_name': matched['name'],
                'section_type': matched['type']
            })
        else:
            result.append({
                'points': default_pts,
                'section_name': '',
                'section_type': 'written'
            })

    return result


MATH_SUBJECTS = {
    'math', 'mathematics', 'algebra', 'pre-algebra', 'geometry',
    'calculus', 'pre-calculus', 'trigonometry', 'statistics',
    'ap calculus', 'ap statistics', 'integrated math',
    'math 6', 'math 7', 'math 8',
}


def _is_math_subject(subject: str) -> bool:
    """Check if the teacher-selected subject is a math subject.

    STRICT: Only returns True when subject was explicitly set to a math
    subject in Settings. Never guesses from question content or keywords.
    """
    return subject.strip().lower() in MATH_SUBJECTS


def grade_per_question(question: str, student_answer: str, expected_answer: str,
                       points: int, grade_level: str, subject: str,
                       teacher_instructions: str, grading_style: str,
                       ai_model: str = 'gpt-4o',
                       ai_provider: str = 'openai',
                       response_type: str = 'marker_response',
                       section_name: str = '', section_type: str = 'written',
                       token_tracker: 'TokenTracker' = None) -> dict:
    """Grade a single question/response pair with full section-aware context.

    Args:
        teacher_instructions: FULL untruncated instructions from app.py. Contains:
            global AI notes, assignment-specific gradingNotes, rubric type overrides,
            accommodation prompts, student history, period differentiation + rubric
            weight adjustments. This is the same string the single-pass gets.
        response_type: One of 'vocab_term', 'numbered_question', 'marker_response', 'fitb_full', 'fill_in_blank'
        section_name: The marker/section this response belongs to (e.g., 'VOCABULARY', 'SUMMARY')
        section_type: The section grading type from marker_config (e.g., 'written', 'fill-in-blank')

    Returns dict with score, reasoning, quality label.
    """
    # Build section-specific grading instructions based on response type
    type_instructions = ""
    if response_type == 'vocab_term':
        type_instructions = """SECTION TYPE: VOCABULARY DEFINITION
- The student was asked to define a vocabulary term
- Grade based on whether the definition captures the correct meaning
- Accept age-appropriate paraphrasing — do NOT require textbook-exact definitions
- A definition that shows understanding of the concept = full credit
- A partially correct definition = partial credit
- A blank or completely wrong definition = 0
- IMPORTANT: Check the TEACHER'S GRADING INSTRUCTIONS below — if the teacher has modified
  how vocabulary should be scored (e.g., requesting leniency, accepting general definitions,
  going easy on vocab, relaxing requirements), follow the teacher's instructions INSTEAD of
  these defaults. The teacher knows their students and their intent overrides default rubric."""
    elif response_type == 'numbered_question':
        type_instructions = """SECTION TYPE: NUMBERED QUESTION
- The student answered a specific numbered question from the assignment
- Grade based on content accuracy and completeness of the answer
- Accept age-appropriate language and reasoning"""
    elif response_type in ('fitb_full', 'fill_in_blank'):
        type_instructions = """SECTION TYPE: FILL-IN-THE-BLANK
- The student filled in blanks in a structured worksheet
- Grade based on factual correctness of each filled-in answer
- Accept synonyms and alternate phrasings that are factually correct
- Spelling errors should NOT be penalized if the meaning is clear"""
    elif section_name and 'summary' in section_name.lower():
        type_instructions = """SECTION TYPE: SUMMARY / WRITTEN RESPONSE
- The student wrote a summary or extended response
- Grade based on: Does it capture the key ideas? Is it in their own words?
- Look for evidence of understanding, not just copying
- Evaluate completeness — did they address the main points?"""
    elif response_type == 'math_equation' or _is_math_subject(subject):
        type_instructions = """SECTION TYPE: MATH EQUATION / CALCULATION
- Accept mathematically equivalent forms (e.g., x(x+2) = x^2+2x)
- Accept equivalent fractions, decimals, and percentages (1/2 = 0.5 = 50%)
- Award partial credit for correct method with arithmetic errors
- If student shows work, evaluate the process even if final answer is wrong
- Do NOT penalize notation differences (2x vs 2*x vs 2·x)
- Evaluate mathematical correctness, not formatting"""
    elif section_type == 'written' and section_name:
        type_instructions = f"""SECTION TYPE: WRITTEN RESPONSE ({section_name})
- Grade based on quality, completeness, and demonstrated understanding
- Look for genuine engagement with the material"""

    excellent_min = int(points * 0.9)
    good_min = int(points * 0.75)
    adequate_min = int(points * 0.6)
    developing_min = int(points * 0.4)

    # Build grading style instructions (matches single-pass behavior)
    if grading_style == 'lenient':
        style_instructions = """GRADING APPROACH: LENIENT
- Prioritize EFFORT over perfection. If a student attempted this, give significant credit.
- Brief answers that show understanding should receive 70-80% of points.
- Do NOT penalize short answers if they demonstrate understanding.
- When in doubt, give the student the benefit of the doubt."""
    elif grading_style == 'strict':
        style_instructions = """GRADING APPROACH: STRICT
- Hold students to high standards for their grade level.
- Brief, underdeveloped answers should be penalized.
- Full credit requires thorough responses demonstrating deep understanding.
- Partial answers receive proportionally reduced credit."""
    else:
        style_instructions = """GRADING APPROACH: STANDARD
- Balance accuracy, completeness, and effort evenly.
- Brief answers receive partial credit proportional to quality.
- Hold to grade-level expectations."""

    section_context = f"\nSECTION: {section_name}" if section_name else ""

    prompt = f"""Grade this single student response.
{section_context}
QUESTION: {question}
STUDENT ANSWER: "{student_answer}"
{f'EXPECTED ANSWER: {expected_answer}' if expected_answer else ''}
POINTS POSSIBLE: {points}

{type_instructions}

{style_instructions}

CONTEXT: Grade {grade_level} {subject} student.

DEFAULT SCORE ANCHORS for {points} points (teacher instructions below may override these):
- Excellent ({excellent_min}-{points}): Correct, complete, shows understanding
- Good ({good_min}-{excellent_min - 1}): Mostly correct, minor gaps
- Adequate ({adequate_min}-{good_min - 1}): Partial understanding shown
- Developing ({developing_min}-{adequate_min - 1}): Minimal understanding
- Insufficient (0-{developing_min - 1}): Incorrect or no meaningful attempt

RULES:
- Accept synonyms and age-appropriate language
- Do NOT penalize spelling if meaning is clear
- Grade the CONTENT, not the writing style
- If blank/empty, score is 0
- If the student answer is template/instruction text (e.g., starts with "Summarize", "Define", "Explain", "Write in complete sentences", "Use evidence from the reading"), this is NOT a student response — it is leftover assignment directions. Score it 0.

---
TEACHER'S GRADING INSTRUCTIONS — these are the HIGHEST PRIORITY and override the score anchors above.
Read these FIRST, then score accordingly:
{teacher_instructions}
---"""

    system_msg = f"You are a grade {grade_level} {subject} teacher grading student work. IMPORTANT: The teacher has provided custom grading instructions in the prompt. You MUST follow them exactly — they override all default scoring rules and anchors. If the teacher says to be lenient, score generously. If the teacher says to accept basic answers, do not penalize simplicity."

    json_schema = '''Respond with ONLY valid JSON in this exact format:
{
    "grade": {
        "score": <integer 0 to ''' + str(points) + '''>,
        "possible": ''' + str(points) + ''',
        "reasoning": "<1-2 sentence explanation>",
        "is_correct": <true or false>,
        "quality": "<excellent|good|adequate|developing|insufficient>"
    },
    "excellent": <true if score >= ''' + str(int(points * 0.9)) + '''>,
    "improvement_note": "<suggestion if not full credit, else empty string>"
}'''

    try:
        if ai_provider == "anthropic":
            import anthropic
            client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            claude_model_map = {
                "claude-haiku": "claude-3-5-haiku-latest",
                "claude-sonnet": "claude-sonnet-4-20250514",
                "claude-opus": "claude-opus-4-20250514",
            }
            actual_model = claude_model_map.get(ai_model, "claude-3-5-haiku-latest")

            response = client.messages.create(
                model=actual_model,
                max_tokens=300,
                system=system_msg + "\n\n" + json_schema,
                messages=[{"role": "user", "content": prompt}]
            )
            if token_tracker:
                token_tracker.record_anthropic(response, actual_model)
            result = _try_parse_json_fallback(response.content[0].text.strip())
            if result and "grade" in result:
                return result

        elif ai_provider == "gemini":
            import google.generativeai as genai
            genai.configure(api_key=GEMINI_API_KEY)
            gemini_model_map = {
                "gemini-flash": "gemini-2.0-flash",
                "gemini-pro": "gemini-2.0-pro-exp",
            }
            actual_model = gemini_model_map.get(ai_model, "gemini-2.0-flash")
            gemini_client = genai.GenerativeModel(actual_model)

            full_prompt = system_msg + "\n\n" + json_schema + "\n\n---\n\n" + prompt
            response = gemini_client.generate_content(full_prompt)
            if token_tracker:
                token_tracker.record_gemini(response, actual_model)
            result = _try_parse_json_fallback(response.text.strip())
            if result and "grade" in result:
                return result

        else:  # OpenAI — use structured output
            from openai import OpenAI
            client = OpenAI(api_key=OPENAI_API_KEY)
            response = client.beta.chat.completions.parse(
                model=ai_model,
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": prompt}
                ],
                response_format=PerQuestionResponse,
                max_tokens=300,
                temperature=0,
                seed=42
            )
            if token_tracker:
                token_tracker.record_openai(response, ai_model)
            parsed = response.choices[0].message.parsed
            if parsed:
                return parsed.model_dump()

    except Exception as e:
        print(f"    ⚠️ Per-question grading error ({ai_provider}): {e}")

    return {
        "grade": {"score": 0, "possible": points,
                  "reasoning": f"Grading error - could not evaluate response ({ai_provider})",
                  "is_correct": False, "quality": "insufficient"},
        "excellent": False,
        "improvement_note": "This response could not be evaluated due to a grading error."
    }


def generate_feedback(question_results: list, total_score: int, total_possible: int,
                      letter_grade: str, grade_level: str, subject: str,
                      teacher_instructions: str = '', ell_language: str = None,
                      ai_model: str = 'gpt-4o-mini',
                      ai_provider: str = 'openai',
                      student_responses: list = None,
                      rubric_breakdown: dict = None,
                      blank_questions: list = None,
                      missing_sections: list = None,
                      token_tracker: 'TokenTracker' = None,
                      student_history: str = '') -> dict:
    """Generate encouraging, improvement-focused teacher feedback from per-question grades.

    Args:
        question_results: Per-question grading results with scores and reasoning.
        student_responses: List of dicts with 'question' and 'answer' keys — the actual
            student text, so the AI can quote specific answers in feedback.
        teacher_instructions: FULL untruncated custom_ai_instructions from app.py.
            Contains student history, accommodations, period differentiation, etc.
        rubric_breakdown: Dict with rubric category scores (content_accuracy, completeness,
            writing_quality, effort_engagement) so feedback can reference rubric performance.
        blank_questions: List of questions the student left blank/unanswered.
        missing_sections: List of required sections entirely missing from submission.
    """

    summary_lines = []
    responses_list = student_responses or []
    for i, qr in enumerate(question_results, 1):
        if not qr:
            continue
        g = qr.get("grade", {})
        # Include the student's actual answer so the AI can quote it
        student_answer = ""
        if i - 1 < len(responses_list):
            resp = responses_list[i - 1]
            student_answer = resp.get("answer", "")
        line = f"Q{i}: {g.get('score', 0)}/{g.get('possible', 10)} ({g.get('quality', 'unknown')})"
        line += f"\n  Reasoning: {g.get('reasoning', '')}"
        if student_answer:
            line += f"\n  Student wrote: \"{student_answer}\""
        summary_lines.append(line)

    # Build grade-scaled feedback instructions
    # ALL grades: encouraging tone, always focused on what the student needs to improve.
    # Higher grades = more time on strengths, but still point to growth areas.
    # Lower grades = more time on improvement guidance, but still acknowledge effort.
    if letter_grade == 'A':
        tone_instructions = """FEEDBACK STRUCTURE FOR AN A (90-100):
Paragraph 1: Celebrate specific strong answers — highlight 2-3 answers that were excellent and explain WHY (e.g., "Your definition of Treaty of Moultrie Creek showed you understood not just what it was, but why it mattered — that's higher-level thinking!").
Paragraph 2: Even A students need growth targets. Identify 1-2 specific areas where they could push deeper or refine their work. Quote their answer and show what a next-level response would look like. Focus on helping them develop more advanced skills (analysis, connections between events, stronger evidence use).
Paragraph 3: If student history is available, connect this to their trajectory and set a challenge for next time.
BALANCE: ~40% strengths, ~60% improvement guidance. An A student should leave knowing exactly what to work on next."""
    elif letter_grade == 'B':
        tone_instructions = """FEEDBACK STRUCTURE FOR A B (80-89):
Paragraph 1: Acknowledge solid work with 1-2 specific strong answers. Quote what they wrote and explain what made it good.
Paragraph 2: Focus here — identify 2-3 specific answers where points were lost. Quote what the student wrote, explain what was missing or incomplete, and tell them exactly what a full-credit answer would have included. Be specific: "You wrote that the Treaty was 'an agreement between the US and Seminoles,' but to earn full credit you needed to mention it was supposed to give them control of millions of acres of Florida land in exchange for allowing roads."
Paragraph 3: Give 1-2 concrete goals for next time. Reference history if available ("You've been hovering around a B — here's what will push you to an A").
BALANCE: ~30% strengths, ~70% improvement guidance."""
    elif letter_grade == 'C':
        tone_instructions = """FEEDBACK STRUCTURE FOR A C (70-79):
Paragraph 1: Briefly acknowledge 1-2 things the student did right. Quote a specific answer that showed some understanding.
Paragraph 2: Main focus — walk through 3-4 specific answers that lost significant points. For each: quote what they wrote, explain what was wrong or missing, and provide the correct information. Example: "For the question about Andrew Jackson's role, you wrote 'he was a president,' but the answer needed to explain that he led troops into Florida during the First Seminole War and later pressured the Seminoles to relocate as president."
Paragraph 3: Give 2-3 specific, actionable steps to improve (re-read specific pages, focus on cause-and-effect, use details from the text). Reference history if available.
BALANCE: ~20% strengths, ~80% improvement guidance. The student needs to understand exactly where they went wrong and what to do differently."""
    elif letter_grade == 'D':
        tone_instructions = """FEEDBACK STRUCTURE FOR A D (60-69):
Paragraph 1: Acknowledge the effort of attempting the assignment, then identify the 2-3 biggest gaps — incomplete answers, missing sections, or incorrect content. Quote specific answers and explain what the correct response should have been.
Paragraph 2: Walk through the most important questions they missed. For each: show what they wrote (or that it was blank), then teach them the answer. This feedback should help them actually learn the material: "The question asked about the causes of the First Seminole War. The key causes were attacks on white settlers, alliances with escaped enslaved people, and southern plantation owners' anger over the Seminole practice of harboring runaways."
Paragraph 3: Give specific recovery steps — re-read certain pages, redo specific questions, come in for help. If history shows a declining trend, address it encouragingly ("I know you can do better — your [previous grade] on [previous assignment] showed you're capable of stronger work").
BALANCE: ~10% strengths, ~90% improvement guidance. Be warm but make sure they leave with the knowledge they were missing."""
    else:  # F or INCOMPLETE
        tone_instructions = """FEEDBACK STRUCTURE FOR AN F (below 60):
Paragraph 1: Identify what went wrong — blank sections, incorrect answers, or missing content. Quote 2-3 specific answers that were wrong or missing and teach the correct answer for each one. The student should learn from reading this feedback.
Paragraph 2: If there is anything the student got partially right, acknowledge it and build on it. Use it as a bridge: "You mentioned 'U.S. efforts to reclaim runaway slaves' as a cause of the First Seminole War — that's a start. To complete the picture, you also needed to include the border conflicts with white settlers and the anger over the Seminole practice of harboring escaped enslaved people."
Paragraph 3: Provide a clear, specific study plan. What pages to re-read, what questions to retry, what concepts to focus on. Make it feel achievable, not overwhelming.
Paragraph 4: If history shows a pattern, address it with care and a path forward ("I've noticed the last few assignments have been tough. Let's figure out what's getting in the way — I'm here to help you turn this around").
BALANCE: ~5% strengths, ~95% improvement guidance. Every sentence should either teach the student something they missed or give them a concrete step to improve. Be encouraging — but the encouragement comes from showing them exactly HOW to do better, not from empty praise."""

    # Build rubric performance summary
    rubric_summary = ""
    if rubric_breakdown:
        rb = rubric_breakdown
        rubric_summary = f"""
RUBRIC BREAKDOWN (address each area in your feedback):
- Content Accuracy: {rb.get('content_accuracy', {}).get('score', 0)}/{rb.get('content_accuracy', {}).get('possible', 40)} — How factually correct were the answers?
- Completeness: {rb.get('completeness', {}).get('score', 0)}/{rb.get('completeness', {}).get('possible', 25)} — Did the student attempt all sections?
- Writing Quality: {rb.get('writing_quality', {}).get('score', 0)}/{rb.get('writing_quality', {}).get('possible', 20)} — Was the writing clear and well-developed?
- Effort & Engagement: {rb.get('effort_engagement', {}).get('score', 0)}/{rb.get('effort_engagement', {}).get('possible', 15)} — Did the student show genuine effort?
"""

    # Build blank/missing sections summary
    blanks_list = blank_questions or []
    missing_list = missing_sections or []
    missing_summary = ""
    if blanks_list or missing_list:
        missing_parts = []
        if missing_list:
            missing_parts.append("MISSING SECTIONS (entire sections not found in submission):")
            for s in missing_list:
                missing_parts.append(f"  - {s}")
        if blanks_list:
            missing_parts.append("BLANK/UNANSWERED QUESTIONS (student left these empty):")
            for q in blanks_list:
                missing_parts.append(f"  - {q}")
        missing_summary = "\n" + "\n".join(missing_parts) + "\n"

    # Build dedicated history section (kept separate from teacher_instructions for prominence)
    history_section = ""
    if student_history:
        history_section = f"\n---\nSTUDENT PERFORMANCE HISTORY (MANDATORY -- you MUST reference this):\n{student_history}\n---\n"

    prompt = f"""Write personalized, encouraging teacher feedback for a grade {grade_level} {subject} student. Focus on what the student needs to know to improve.

SCORE: {total_score}/{total_possible} ({letter_grade})
{rubric_summary}{missing_summary}
PER-QUESTION RESULTS (with the student's actual answers):
{chr(10).join(summary_lines)}

---
TEACHER'S INSTRUCTIONS & GRADING CONTEXT:
{teacher_instructions}
---
{history_section}
NOTE ON TEACHER LENIENCY: If the per-question reasoning above says "Teacher accepts general definitions"
for vocabulary terms, do NOT criticize those vocab answers as "too basic" or "lacking context" in your
feedback. Briefly acknowledge them positively, then move on to the questions and sections that actually
need improvement. The teacher's leniency applies ONLY to vocabulary — you must STILL provide full,
detailed constructive feedback on all other sections (questions, summary, missing work, etc.).

{tone_instructions}

UNIVERSAL RULES:
- Quote or paraphrase the student's SPECIFIC answers — never give generic feedback
- For every wrong answer you mention, explain what the correct answer is or what was missing
- MISSING WORK: If there are blank questions or missing sections listed above, you MUST call them out specifically in your feedback. Name the exact questions or sections that were left blank and explain what the student should have written. Example: "You left the SUMMARY section blank — this section asked you to summarize the key events of the Seminole Wars in 4-5 sentences. To complete it, you'd want to cover the three wars (1817-1858), the role of Andrew Jackson, and the eventual forced relocation to Indian Territory."
- Reference the actual assignment content (topic, questions, vocabulary terms)
- You MUST reference the STUDENT PERFORMANCE HISTORY section above if present. Compare this score to previous scores. Check if the student improved on previously flagged areas. If no history section is provided, skip this
- RUBRIC PERFORMANCE: Address the student's performance on ALL aspects of the rubric — content accuracy, completeness, writing quality, and effort/engagement. For each rubric area, note whether it was a strength or weakness and explain why. Example: "Your content accuracy was strong — most of your answers were factually correct. But your completeness needs work — you left two questions blank, which cost you a full letter grade."
- Do NOT use the student's name — say "you" or "your"
- Sound like a real teacher — use contractions, natural language
- Write feedback in English only
- The feedback must be USEFUL — a parent reading this should understand exactly what their child got right, what they got wrong, and what they need to do to improve
- NEVER write vague transitions like "there are several areas where you can improve" or "however, improvements can be made" without IMMEDIATELY listing the specific areas. If you say improvements are needed, you MUST list each one with the specific question, what the student wrote, and what the correct answer is. Generic "areas to improve" statements with no details are UNACCEPTABLE.
- Do NOT include teacher sign-offs, signatures, or closing lines like "Warm Regards, Mr. Smith" — those are added separately by the system

Also identify:
- Excellent answers: Quote 1-4 specific student answers that earned high marks (fewer for low grades)
- Areas needing improvement: Quote 1-4 specific answers that lost points, with what the correct answer should include
- Strengths: 1-4 specific skills the student demonstrated (tied to rubric areas)
- Developing: 1-3 specific skills the student needs to work on (tied to rubric areas)"""

    system_msg = f"You are an encouraging grade {grade_level} {subject} teacher writing specific, actionable feedback focused on helping the student improve. Always reference actual student answers. Scale the balance of praise vs improvement guidance based on the grade, but every grade level should focus primarily on what the student needs to do to get better."

    json_schema = '''Respond with ONLY valid JSON in this exact format:
{
    "feedback": "<3-4 paragraphs of personalized feedback>",
    "excellent_answers": ["<specific student answers that earned high marks>"],
    "needs_improvement": ["<specific answers that lost points, with corrections>"],
    "skills_demonstrated": {
        "strengths": ["<specific skill demonstrated>"],
        "developing": ["<specific skill to work on>"]
    }
}'''

    try:
        if ai_provider == "anthropic":
            import anthropic
            client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            claude_model_map = {
                "claude-haiku": "claude-3-5-haiku-latest",
                "claude-sonnet": "claude-sonnet-4-20250514",
                "claude-opus": "claude-opus-4-20250514",
            }
            actual_model = claude_model_map.get(ai_model, "claude-3-5-haiku-latest")

            response = client.messages.create(
                model=actual_model,
                max_tokens=3500,
                system=system_msg + "\n\n" + json_schema,
                messages=[{"role": "user", "content": prompt}]
            )
            if token_tracker:
                token_tracker.record_anthropic(response, actual_model)
            result = _try_parse_json_fallback(response.content[0].text.strip())
            if result and "feedback" in result:
                if "skills_demonstrated" not in result or not isinstance(result["skills_demonstrated"], dict):
                    result["skills_demonstrated"] = {"strengths": [], "developing": []}
                if ell_language and result.get("feedback"):
                    translated = _translate_feedback(result["feedback"], ell_language, ai_model, token_tracker=token_tracker)
                    if translated:
                        result["feedback"] = result["feedback"] + "\n\n---\n\n" + translated
                return result

        elif ai_provider == "gemini":
            import google.generativeai as genai
            genai.configure(api_key=GEMINI_API_KEY)
            gemini_model_map = {
                "gemini-flash": "gemini-2.0-flash",
                "gemini-pro": "gemini-2.0-pro-exp",
            }
            actual_model = gemini_model_map.get(ai_model, "gemini-2.0-flash")
            gemini_client = genai.GenerativeModel(actual_model)

            full_prompt = system_msg + "\n\n" + json_schema + "\n\n---\n\n" + prompt
            response = gemini_client.generate_content(full_prompt)
            if token_tracker:
                token_tracker.record_gemini(response, actual_model)
            result = _try_parse_json_fallback(response.text.strip())
            if result and "feedback" in result:
                if "skills_demonstrated" not in result or not isinstance(result["skills_demonstrated"], dict):
                    result["skills_demonstrated"] = {"strengths": [], "developing": []}
                if ell_language and result.get("feedback"):
                    translated = _translate_feedback(result["feedback"], ell_language, ai_model, token_tracker=token_tracker)
                    if translated:
                        result["feedback"] = result["feedback"] + "\n\n---\n\n" + translated
                return result

        else:  # OpenAI — use structured output
            from openai import OpenAI
            client = OpenAI(api_key=OPENAI_API_KEY)
            response = client.beta.chat.completions.parse(
                model=ai_model,
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": prompt}
                ],
                response_format=FeedbackResponse,
                max_tokens=3500,
                temperature=0,
                seed=42
            )
            if token_tracker:
                token_tracker.record_openai(response, ai_model)
            parsed = response.choices[0].message.parsed
            if parsed:
                result = parsed.model_dump()
                if ell_language and result.get("feedback"):
                    translated = _translate_feedback(result["feedback"], ell_language, ai_model, token_tracker=token_tracker)
                    if translated:
                        result["feedback"] = result["feedback"] + "\n\n---\n\n" + translated
                return result

    except Exception as e:
        print(f"  ⚠️ Feedback generation error ({ai_provider}): {e}")

    return {
        "feedback": "Good effort on this assignment. Keep working hard!",
        "excellent_answers": [],
        "needs_improvement": [],
        "skills_demonstrated": {"strengths": [], "developing": []}
    }


def grade_multipass(student_name: str, assignment_data: dict, custom_ai_instructions: str = '',
                    grade_level: str = '6', subject: str = 'Social Studies',
                    ai_model: str = 'gpt-4o-mini', student_id: str = None,
                    assignment_template: str = None, rubric_prompt: str = None,
                    custom_markers: list = None, exclude_markers: list = None,
                    marker_config: list = None, effort_points: int = 15,
                    extraction_mode: str = 'structured', grading_style: str = 'standard',
                    token_tracker: 'TokenTracker' = None,
                    student_history: str = '') -> dict:
    """Multi-pass grading pipeline for consistent, robust scoring.

    Pass 1: Extract responses (reuses existing extraction logic)
    Pass 2: Grade each question individually (parallel, structured output)
    Pass 3: Generate feedback (cheaper model)
    Final: Aggregate scores, apply caps, build result
    """
    # Determine provider from model name
    if ai_model.startswith("claude"):
        provider = "anthropic"
    elif ai_model.startswith("gemini"):
        provider = "gemini"
    else:
        provider = "openai"

    tracker = token_tracker or TokenTracker()
    content = assignment_data.get("content", "")

    # === EXTRACTION ===
    # Priority: Graider structured tables > Graider text fallback > regex extraction
    extraction_result = None

    # Check for Graider table data (structured worksheets)
    graider_tables = assignment_data.get("graider_tables")
    if graider_tables:
        print(f"  📊 Multi-pass: Using Graider table extraction ({len(graider_tables)} tables)")
        extraction_result = extract_from_tables(graider_tables, exclude_markers)
    elif assignment_data.get("type") == "text" and content:
        # Try GRAIDER tag plain-text fallback before generic extraction
        if '[GRAIDER:' in content:
            extraction_result = extract_from_graider_text(content, exclude_markers)
        if not extraction_result or not extraction_result.get("extracted_responses"):
            extraction_result = extract_student_responses(content, custom_markers, exclude_markers, assignment_template)

    if extraction_result:
        answered = extraction_result.get("answered_questions", 0)
        total = extraction_result.get("total_questions", 0)
        print(f"  📋 Multi-pass: Extracted {answered}/{total} responses")

        if answered == 0:
            return {
                "score": 0, "letter_grade": "INCOMPLETE",
                "breakdown": {"content_accuracy": 0, "completeness": 0, "writing_quality": 0, "effort_engagement": 0},
                "feedback": "You submitted a blank assignment with no responses. Please complete all sections and resubmit.",
                "student_responses": [], "unanswered_questions": extraction_result.get("blank_questions", []) + extraction_result.get("missing_sections", []),
                "ai_detection": {"flag": "none", "confidence": 0, "reason": "Blank submission — no content to evaluate."},
                "plagiarism_detection": {"flag": "none", "reason": "Blank submission — no content to evaluate."},
                "skills_demonstrated": {"strengths": [], "developing": []},
                "excellent_answers": [], "needs_improvement": []
            }

        # Force zero if 80%+ of questions are blank — prevents template text inflation
        total_questions = extraction_result.get("total_questions", 0)
        blank_questions_count = len(extraction_result.get("blank_questions", [])) + len(extraction_result.get("missing_sections", []))
        if total_questions > 0 and blank_questions_count / total_questions >= 0.8:
            print(f"  ⚠️  NEARLY BLANK: {blank_questions_count}/{total_questions} questions blank (≥80%)")
            return {
                "score": 0, "letter_grade": "INCOMPLETE",
                "breakdown": {"content_accuracy": 0, "completeness": 0, "writing_quality": 0, "effort_engagement": 0},
                "feedback": f"Your assignment is nearly blank — {blank_questions_count} out of {total_questions} questions have no response. Please complete all sections and resubmit.",
                "student_responses": [], "unanswered_questions": extraction_result.get("blank_questions", []) + extraction_result.get("missing_sections", []),
                "ai_detection": {"flag": "none", "confidence": 0, "reason": "Nearly blank submission."},
                "plagiarism_detection": {"flag": "none", "reason": "Nearly blank submission."},
                "skills_demonstrated": {"strengths": [], "developing": []},
                "excellent_answers": [], "needs_improvement": []
            }

    if not extraction_result or not extraction_result.get("extracted_responses"):
        # Fall back to single-pass for edge cases
        print(f"  ⚠️ Multi-pass: No extracted responses, falling back to single-pass")
        return grade_assignment(student_name, assignment_data, custom_ai_instructions,
                               grade_level, subject, ai_model, student_id, assignment_template,
                               rubric_prompt, custom_markers, exclude_markers, marker_config,
                               effort_points, extraction_mode, grading_style)

    responses = extraction_result["extracted_responses"]

    # Build expected answers lookup from gradingNotes within custom_ai_instructions
    expected_answers = _parse_expected_answers(custom_ai_instructions)

    # NOTE: accommodation context, student history, period differentiation, and rubric
    # type overrides are ALREADY embedded in custom_ai_instructions by app.py (lines 975-1119).
    # We pass the full string untruncated to each per-question call.

    # Append the custom rubric prompt (from Settings) so per-question graders see it.
    # In single-pass, rubric_prompt overrides GRADING_RUBRIC. In multipass, we append it
    # to the teacher instructions so each per-question call gets the rubric categories/weights.
    effective_instructions = custom_ai_instructions
    if rubric_prompt:
        effective_instructions += "\n\n" + rubric_prompt

    # === PASS 2: PER-QUESTION GRADING (parallel) ===
    total_content_points = 100 - effort_points
    question_meta = _distribute_points(responses, marker_config, total_content_points)

    # Use the selected model for per-question grading (no auto-upgrade)
    grading_model = ai_model

    print(f"  🔄 Multi-pass: Grading {len(responses)} questions with {grading_model}...")

    # Submit all questions in parallel, track by index
    question_results = [None] * len(responses)
    with concurrent.futures.ThreadPoolExecutor(max_workers=5) as executor:
        future_to_idx = {}
        for i, resp in enumerate(responses):
            question = resp.get("question", f"Question {i+1}")
            answer = resp.get("answer", "")
            resp_type = resp.get("type", "marker_response")
            meta = question_meta[i] if i < len(question_meta) else {'points': 10, 'section_name': '', 'section_type': 'written'}

            # Match expected answer by multiple strategies:
            # 1. Question number from text (e.g., "1) What was..." → Q1 → index 0)
            # 2. Term/question text match (for vocab: "Seminole Wars" → key match)
            # 3. Section name match
            # 4. Response list index (only works if no vocab terms shift indices)
            expected = ""

            # Strategy 1: Extract question number and match to Q-index
            q_num_match = re.match(r'^(\d+)', question.strip())
            if q_num_match:
                q_idx = int(q_num_match.group(1)) - 1  # "1)" → index 0
                expected = expected_answers.get(q_idx, "") or expected_answers.get(f"Q{q_num_match.group(1)}", "")

            # Strategy 2: Match by term/question text or section name
            if not expected:
                expected = (expected_answers.get(question, "") or
                            expected_answers.get(question.split(':')[0].strip(), "") or
                            expected_answers.get(meta['section_name'], ""))

            # Strategy 3: Fall back to list index
            if not expected:
                expected = expected_answers.get(i, "")

            # SymPy pre-check: if math subject with expected answer, try exact match first
            if _is_math_subject(subject) and expected and answer:
                try:
                    from backend.services.stem_grading import check_math_equivalence
                    equiv = check_math_equivalence(answer, expected)
                    if equiv.get('equivalent'):
                        pts = meta['points']
                        question_results[i] = {
                            "grade": {"score": pts, "possible": pts,
                                      "reasoning": f"Mathematically equivalent ({equiv['method']})",
                                      "is_correct": True, "quality": "excellent"},
                            "excellent": True, "improvement_note": ""
                        }
                        continue  # Skip LLM call — instant correct, zero cost
                except Exception:
                    pass  # SymPy failed — fall through to normal AI grading

            f = executor.submit(
                grade_per_question,
                question=question,
                student_answer=answer,
                expected_answer=expected,
                points=meta['points'],
                grade_level=grade_level,
                subject=subject,
                teacher_instructions=effective_instructions,  # FULL — includes rubric
                grading_style=grading_style,
                ai_model=grading_model,
                ai_provider=provider,
                response_type=resp_type,
                section_name=meta['section_name'],
                section_type=meta['section_type'],
                token_tracker=tracker
            )
            future_to_idx[f] = i

        for future in concurrent.futures.as_completed(future_to_idx):
            idx = future_to_idx[future]
            try:
                question_results[idx] = future.result()
            except Exception as e:
                print(f"    ⚠️ Question {idx+1} grading failed: {e}")
                meta = question_meta[idx] if idx < len(question_meta) else {'points': 10}
                question_results[idx] = {
                    "grade": {"score": int(meta['points'] * 0.7), "possible": meta['points'],
                              "reasoning": "Error during grading", "is_correct": True, "quality": "adequate"},
                    "excellent": False, "improvement_note": ""
                }

    # === TEACHER LENIENCY POST-PROCESSING ===
    # If the teacher requested leniency for specific section types, apply score floors in code.
    # This is more reliable than prompt engineering — the AI scores normally, then we adjust.
    _ei_lower = (effective_instructions or '').lower()
    _has_vocab_leniency = any(phrase in _ei_lower for phrase in [
        'lenient', 'accept general', 'accept basic', 'go easy', 'be generous',
        'not strict', 'relaxed', 'don\'t be harsh', 'accept simple'
    ]) and any(w in _ei_lower for w in ['vocab', 'definition', 'terms'])

    if _has_vocab_leniency:
        adjusted_count = 0
        for i, resp in enumerate(responses):
            if resp.get("type") == "vocab_term" and resp.get("answer", "").strip():
                qr = question_results[i]
                if qr:
                    grade = qr.get("grade", {})
                    pts = grade.get("possible", 9)
                    current_score = grade.get("score", 0)
                    min_score = int(pts * 0.65)  # At least 65% for any non-blank vocab answer
                    if current_score < min_score:
                        grade["score"] = min_score
                        grade["quality"] = "adequate"
                        # REPLACE reasoning entirely — the old reasoning says "too basic"
                        # and the feedback generator echoes it. Clean reasoning = clean feedback.
                        term = resp.get("question", "this term")
                        grade["reasoning"] = f"Student provided a basic definition for {term} that shows general understanding. Teacher accepts general/dictionary-level definitions for vocabulary on this assignment."
                        adjusted_count += 1
        if adjusted_count > 0:
            print(f"  📌 Vocab leniency: adjusted {adjusted_count} vocab scores to minimum 65%")

    # === AGGREGATE SCORES ===
    total_earned = sum(qr.get("grade", {}).get("score", 0) for qr in question_results if qr)
    total_possible = sum(qr.get("grade", {}).get("possible", 10) for qr in question_results if qr)

    blank_count = len(extraction_result.get("blank_questions", [])) + len(extraction_result.get("missing_sections", []))
    if blank_count == 0:
        effort_earned = effort_points
    elif blank_count == 1:
        effort_earned = int(effort_points * 0.7)
    elif blank_count == 2:
        effort_earned = int(effort_points * 0.4)
    else:
        effort_earned = 0  # 3+ blanks = no effort credit

    raw_score = int(round((total_earned / max(total_possible, 1)) * (100 - effort_points) + effort_earned))
    raw_score = max(0, min(100, raw_score))

    # Completeness caps by grading style — each missing section drops max possible grade
    if grading_style == 'strict':
        caps = {0: 100, 1: 85, 2: 75, 3: 65, 4: 55, 5: 45, 6: 35, 7: 25, 8: 15}
    elif grading_style == 'lenient':
        caps = {0: 100, 1: 95, 2: 89, 3: 79, 4: 69, 5: 59, 6: 49, 7: 39, 8: 29}
    else:
        caps = {0: 100, 1: 89, 2: 79, 3: 69, 4: 59, 5: 49, 6: 39, 7: 29, 8: 19}
    if blank_count >= len(caps):
        cap = 0  # More blanks than cap table entries → zero
    else:
        cap = caps.get(blank_count, 0)
    final_score = min(raw_score, cap)
    if blank_count > 0:
        print(f"  📉 Completeness: {blank_count} blank/missing → cap at {cap}")

    if final_score >= 90: letter_grade = "A"
    elif final_score >= 80: letter_grade = "B"
    elif final_score >= 70: letter_grade = "C"
    elif final_score >= 60: letter_grade = "D"
    else: letter_grade = "F"

    per_q_scores = [qr.get("grade", {}).get("score", 0) for qr in question_results if qr]
    print(f"  📊 Per-question: {per_q_scores}")
    print(f"  📊 Raw: {raw_score}, Cap: {cap}, Final: {final_score} ({letter_grade})")

    # === PASS 3: FEEDBACK GENERATION ===
    ell_language = None
    if student_id and student_id != "UNKNOWN":
        ell_file = os.path.expanduser("~/.graider_data/ell_students.json")
        if os.path.exists(ell_file):
            try:
                with open(ell_file, 'r', encoding='utf-8') as f:
                    ell_data = json.load(f)
                ell_entry = ell_data.get(student_id, {})
                lang = ell_entry.get("language")
                if lang and lang != "none":
                    ell_language = lang
            except Exception:
                pass

    # === BUILD BREAKDOWN (before feedback so we can pass rubric scores) ===
    content_pts = int(round((total_earned / max(total_possible, 1)) * 40))
    completeness_pts = max(0, 25 - (blank_count * 6))
    qualities = [qr.get("grade", {}).get("quality", "adequate") for qr in question_results if qr]
    if qualities.count("excellent") + qualities.count("good") > len(qualities) * 0.7:
        writing_pts = 18
    elif qualities.count("developing") + qualities.count("insufficient") > len(qualities) * 0.5:
        writing_pts = 10
    else:
        writing_pts = 15

    rubric_breakdown = {
        "content_accuracy": {"score": content_pts, "possible": 40},
        "completeness": {"score": completeness_pts, "possible": 25},
        "writing_quality": {"score": writing_pts, "possible": 20},
        "effort_engagement": {"score": effort_earned, "possible": effort_points},
    }

    # Collect blank/missing info for feedback
    blank_questions = extraction_result.get("blank_questions", [])
    missing_sections = extraction_result.get("missing_sections", [])

    # Use gpt-4o for feedback — it's 1 call per student and the most important output
    feedback_model = ai_model
    if provider == "openai":
        feedback_model = "gpt-4o"  # Feedback is what teachers/parents read — needs quality
    # Claude/Gemini: use the teacher's selected model

    print(f"  🔄 Multi-pass: Generating feedback ({feedback_model})...")
    feedback_result = generate_feedback(
        question_results=question_results,
        total_score=final_score, total_possible=100,
        letter_grade=letter_grade,
        grade_level=grade_level, subject=subject,
        teacher_instructions=effective_instructions,
        ell_language=ell_language,
        ai_model=feedback_model,
        ai_provider=provider,
        student_responses=responses,
        rubric_breakdown=rubric_breakdown,
        blank_questions=blank_questions,
        missing_sections=missing_sections,
        token_tracker=tracker,
        student_history=student_history
    )

    # === BUILD RESULT ===
    student_response_texts = [resp.get("answer", "")[:500] for resp in responses if resp.get("answer")]

    result = {
        "score": final_score,
        "letter_grade": letter_grade,
        "breakdown": {
            "content_accuracy": min(content_pts, 40),
            "completeness": min(completeness_pts, 25),
            "writing_quality": min(writing_pts, 20),
            "effort_engagement": effort_earned
        },
        "student_responses": student_response_texts,
        "unanswered_questions": extraction_result.get("blank_questions", []) + extraction_result.get("missing_sections", []),
        "excellent_answers": feedback_result.get("excellent_answers", []),
        "needs_improvement": feedback_result.get("needs_improvement", []),
        "skills_demonstrated": feedback_result.get("skills_demonstrated", {"strengths": [], "developing": []}),
        "ai_detection": {"flag": "none", "confidence": 0, "reason": ""},
        "plagiarism_detection": {"flag": "none", "reason": ""},
        "feedback": feedback_result.get("feedback", ""),
        "multipass_grading": True,
        "per_question_scores": [
            {"question": responses[i].get("question", "")[:60],
             "score": qr.get("grade", {}).get("score", 0),
             "possible": qr.get("grade", {}).get("possible", 10),
             "quality": qr.get("grade", {}).get("quality", "")}
            for i, qr in enumerate(question_results) if qr
        ],
        "token_usage": tracker.summary()
    }

    # Add audit trail for AI Reasoning / Raw API Output
    audit_input_parts = []
    for i, resp in enumerate(responses):
        q = resp.get("question", f"Q{i+1}")
        a = resp.get("answer", "")
        audit_input_parts.append(f"[{q}]\n{a}")
    audit_response_parts = []
    for i, qr in enumerate(question_results):
        if qr:
            g = qr.get("grade", {})
            audit_response_parts.append(
                f"Q{i+1}: {g.get('score', 0)}/{g.get('possible', 10)} "
                f"({g.get('quality', 'N/A')}) - {g.get('reasoning', '')}"
            )
    result["_audit"] = {
        "ai_input": "\n\n".join(audit_input_parts),
        "ai_response": "\n".join(audit_response_parts) + "\n\n--- FEEDBACK ---\n" + feedback_result.get("feedback", "")
    }

    # Update writing profile
    if student_id and student_id != "UNKNOWN" and content:
        current_writing_style = analyze_writing_style(content)
        if current_writing_style:
            ai_flag = result.get("ai_detection", {}).get("flag", "none")
            if ai_flag not in ["likely", "possible"]:
                try:
                    update_writing_profile(student_id, current_writing_style, student_name)
                except Exception:
                    pass

    print(f"  ✅ Multi-pass grading complete: {final_score} ({letter_grade})")
    return result


# =============================================================================
# SECTION-BASED RUBRIC BUILDER
# =============================================================================

def build_section_rubric(marker_config: list, effort_points: int = 15) -> str:
    """Build a section-based rubric from marker configuration."""
    if not marker_config:
        return ""  # Use default rubric

    rubric_lines = ["""
You are grading a student assignment with SECTION-BASED POINT VALUES.
Be ENCOURAGING and GENEROUS - these are middle school students.

SECTION POINT VALUES:"""]

    total = 0
    for m in marker_config:
        if isinstance(m, dict):
            name = m.get('start', 'Section')
            points = m.get('points', 10)
            section_type = m.get('type', 'written')
            rubric_lines.append(f"- {name}: {points} points ({section_type})")
            total += points
        elif isinstance(m, str):
            rubric_lines.append(f"- {m}: 10 points")
            total += 10

    rubric_lines.append(f"- Effort & Engagement: {effort_points} points")
    total += effort_points
    rubric_lines.append(f"\nTOTAL: {total} points")

    # Build section names list for JSON example
    section_names = [m.get('start', 'Section') if isinstance(m, dict) else m for m in marker_config]
    section_json_example = ',\n        '.join([f'"{name}": {{"earned": <pts>, "possible": {m.get("points", 10) if isinstance(m, dict) else 10}}}' for name, m in zip(section_names, marker_config)])

    rubric_lines.append(f"""
GRADING RULES:
- Grade each section out of its assigned points
- BLANK SECTION = 0 POINTS for that section (no partial credit)
- For fill-blank sections: each correct answer is worth proportional points
- For written sections: grade on quality, completeness, and effort
- Effort & Engagement is based on overall presentation and engagement
- Accept reasonable synonyms and alternate phrasings
- Minor spelling errors should NOT be penalized if meaning is clear

IMPORTANT: Your JSON output MUST include a "section_scores" field showing points for each section:
"section_scores": {{
        {section_json_example},
        "Effort & Engagement": {{"earned": <pts>, "possible": {effort_points}}}
    }}

The "score" field should equal the sum of all section_scores earned values.
""")

    return "\n".join(rubric_lines)


# =============================================================================
# BILINGUAL FEEDBACK TRANSLATION (two-pass system for ELL students)
# =============================================================================

def _translate_feedback(feedback: str, target_language: str, ai_model: str = 'gpt-4o-mini', token_tracker: 'TokenTracker' = None) -> str:
    """
    Translate grading feedback into the target language using a dedicated API call.
    This is a separate, focused call that produces consistent results because
    translation is the ONLY task — no competing grading instructions.

    Returns the translated text, or empty string on failure.
    """
    if not feedback or not target_language:
        return ""

    prompt = f"""Translate the following teacher feedback into {target_language}.
Keep the same warm, encouraging tone. Do not add or remove content — translate everything faithfully.
Do not include any English text in your response — only the {target_language} translation.

FEEDBACK TO TRANSLATE:
{feedback}"""

    try:
        if ai_model.startswith("claude"):
            import anthropic
            client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
            claude_model_map = {
                "claude-haiku": "claude-3-5-haiku-latest",
                "claude-sonnet": "claude-sonnet-4-20250514",
                "claude-opus": "claude-opus-4-20250514",
            }
            model = claude_model_map.get(ai_model, "claude-3-5-haiku-latest")
            response = client.messages.create(
                model=model,
                max_tokens=2000,
                messages=[{"role": "user", "content": prompt}]
            )
            if token_tracker:
                token_tracker.record_anthropic(response, model)
            return response.content[0].text.strip()

        elif ai_model.startswith("gemini"):
            import google.generativeai as genai
            genai.configure(api_key=GEMINI_API_KEY)
            gemini_model_map = {
                "gemini-flash": "gemini-2.0-flash",
                "gemini-pro": "gemini-2.0-pro-exp",
            }
            model = gemini_model_map.get(ai_model, "gemini-2.0-flash")
            client = genai.GenerativeModel(model)
            response = client.generate_content(prompt)
            if token_tracker:
                token_tracker.record_gemini(response, model)
            return response.text.strip()

        else:
            from openai import OpenAI
            client = OpenAI(api_key=OPENAI_API_KEY)
            response = client.chat.completions.create(
                model=ai_model,
                messages=[{"role": "user", "content": prompt}],
                max_tokens=2000,
                temperature=0.3
            )
            if token_tracker:
                token_tracker.record_openai(response, ai_model)
            return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"  ⚠️  Translation to {target_language} failed: {e}")
        return ""


# =============================================================================
# JSON PARSING HELPERS
# =============================================================================

def _try_parse_json_fallback(text: str) -> dict:
    """Attempt to parse JSON from LLM response text with repair logic.

    Used as fallback for Claude/Gemini responses and OpenAI text fallback.
    Returns parsed dict or None if unrecoverable.
    """
    if not text:
        return None

    # First try direct parse
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    # Strip markdown code blocks
    cleaned = text.strip()
    if cleaned.startswith("```"):
        lines = cleaned.split('\n')
        start = 1 if lines[0].startswith("```") else 0
        end = len(lines)
        for i in range(len(lines) - 1, -1, -1):
            if lines[i].strip() == "```":
                end = i
                break
        cleaned = '\n'.join(lines[start:end]).strip()

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # Repair common LLM JSON issues
    fixed = cleaned

    # Fix malformed "reason" fields
    fixed = re.sub(r'"reason":\s*"[^"]*\\n[^"]*"', '"reason": ""', fixed)
    fixed = re.sub(r'"reason":\s*"[\s\n]*\}', '"reason": ""}', fixed)
    fixed = re.sub(r'"reason":\s*"[\s\n]*,', '"reason": "",', fixed)
    fixed = re.sub(r'"reason":\s*"[^"]*\{[^"]*"', '"reason": ""', fixed)

    # Remove parenthetical comments after closing quotes
    fixed = re.sub(r'"\s*\([^)]+\)', '"', fixed)

    # Add missing commas
    fixed = re.sub(r'"\s*\n(\s*)"', r'",\n\1"', fixed)
    fixed = re.sub(r'(\d)\s*\n(\s*)"', r'\1,\n\2"', fixed)
    fixed = re.sub(r'(true|false|null)\s*\n(\s*)"', r'\1,\n\2"', fixed)
    fixed = re.sub(r'(\]|\})\s*\n(\s*)"', r'\1,\n\2"', fixed)

    # Escape unescaped newlines inside strings
    result_chars = []
    in_string = False
    i = 0
    while i < len(fixed):
        char = fixed[i]
        if char == '\\' and i + 1 < len(fixed):
            result_chars.append(char)
            result_chars.append(fixed[i + 1])
            i += 2
            continue
        if char == '"':
            in_string = not in_string
            result_chars.append(char)
            i += 1
            continue
        if in_string:
            if char == '\n':
                result_chars.append('\\n')
            elif char == '\r':
                pass
            elif char == '\t':
                result_chars.append('\\t')
            else:
                result_chars.append(char)
        else:
            result_chars.append(char)
        i += 1

    fixed = ''.join(result_chars)

    try:
        return json.loads(fixed)
    except json.JSONDecodeError:
        return None


# =============================================================================
# AI GRADING WITH CLAUDE
# =============================================================================

def grade_assignment(student_name: str, assignment_data: dict, custom_ai_instructions: str = '', grade_level: str = '6', subject: str = 'Social Studies', ai_model: str = 'gpt-4o-mini', student_id: str = None, assignment_template: str = None, rubric_prompt: str = None, custom_markers: list = None, exclude_markers: list = None, marker_config: list = None, effort_points: int = 15, extraction_mode: str = 'structured', grading_style: str = 'standard', token_tracker: 'TokenTracker' = None) -> dict:
    """
    Use OpenAI GPT to grade a student assignment.

    FERPA COMPLIANCE: Student name is NOT sent to OpenAI.
    We use "Student" as a placeholder to protect privacy.

    Supports both text and image inputs.

    Parameters:
    - student_name: Name of the student (kept local, not sent to API)
    - assignment_data: dict with "type" ("text" or "image") and "content"
    - custom_ai_instructions: Additional grading instructions from the teacher
    - grade_level: The student's grade level (e.g., '6', '7', '8')
    - subject: The subject being graded (e.g., 'Social Studies', 'English/ELA')
    - ai_model: OpenAI model to use ('gpt-4o' or 'gpt-4o-mini')
    - assignment_template: The original assignment template with all questions (for context)

    Returns dict with:
    - score: numeric grade (0-100)
    - letter_grade: A, B, C, D, or F
    - feedback: detailed feedback for the student
    - breakdown: points for each rubric category
    - authenticity_flag: 'clean', 'review', or 'flagged'
    - authenticity_reason: Explanation for flagged or review status
    """
    # Determine which API to use based on model name
    if ai_model.startswith("claude"):
        provider = "anthropic"
    elif ai_model.startswith("gemini"):
        provider = "gemini"
    else:
        provider = "openai"

    # Initialize the appropriate client
    if provider == "anthropic":
        try:
            import anthropic
        except ImportError:
            print("❌ anthropic not installed. Run: pip install anthropic")
            return {"score": 0, "letter_grade": "ERROR", "breakdown": {}, "feedback": "Anthropic API not available - pip install anthropic"}

        if not ANTHROPIC_API_KEY:
            return {"score": 0, "letter_grade": "ERROR", "breakdown": {}, "feedback": "Anthropic API key not configured"}

        claude_client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        claude_model_map = {
            "claude-haiku": "claude-3-5-haiku-latest",
            "claude-sonnet": "claude-sonnet-4-20250514",
            "claude-opus": "claude-opus-4-20250514",
        }
        actual_model = claude_model_map.get(ai_model, "claude-3-5-haiku-latest")
        print(f"  🤖 Using Claude model: {actual_model}")

    elif provider == "gemini":
        try:
            import google.generativeai as genai
        except ImportError:
            print("❌ google-generativeai not installed. Run: pip install google-generativeai")
            return {"score": 0, "letter_grade": "ERROR", "breakdown": {}, "feedback": "Google AI not available - pip install google-generativeai"}

        if not GEMINI_API_KEY:
            return {"score": 0, "letter_grade": "ERROR", "breakdown": {}, "feedback": "Gemini API key not configured"}

        genai.configure(api_key=GEMINI_API_KEY)
        gemini_model_map = {
            "gemini-flash": "gemini-2.0-flash",
            "gemini-pro": "gemini-2.0-pro-exp",
        }
        actual_model = gemini_model_map.get(ai_model, "gemini-2.0-flash")
        gemini_client = genai.GenerativeModel(actual_model)
        print(f"  🤖 Using Gemini model: {actual_model}")

    else:  # OpenAI
        try:
            from openai import OpenAI
        except ImportError:
            print("❌ openai not installed. Run: pip install openai")
            return {"score": 0, "letter_grade": "ERROR", "breakdown": {}, "feedback": "OpenAI API not available"}

        openai_client = OpenAI(api_key=OPENAI_API_KEY)

    content = assignment_data.get("content", "")

    # Strip embedded answer key from generated worksheets (handles -- and --- variants)
    if content and "GRAIDER_ANSWER_KEY_START" in content:
        content = content.split("GRAIDER_ANSWER_KEY_START")[0].rstrip().rstrip('-')
        assignment_data = {**assignment_data, "content": content}
        print(f"  🧹 Stripped embedded answer key from document")

    # Check for empty/blank student submissions before sending to API
    if assignment_data.get("type") == "text" and content:
        import re

        # Method 1: Check for filled-in blanks (text between underscores like ___answer___)
        filled_blanks = re.findall(r'_{2,}([^_\n]+)_{2,}', content)
        filled_blanks = [b.strip() for b in filled_blanks if b.strip() and len(b.strip()) > 1]

        # Method 2: Check for content after colons that isn't just blanks
        # e.g., "Nationalism: the belief that..." vs "Nationalism: ___"
        after_colons = re.findall(r':\s*([^_\n:]{10,})', content)
        # Filter out template instruction text that isn't student writing
        _instruction_patterns = re.compile(
            r'^(define|summarize|explain|describe|write|use|answer|identify|list|compare|analyze|discuss|'
            r'read|complete|fill|circle|match|select|choose|highlight|underline|review|include)\b',
            re.IGNORECASE
        )
        after_colons = [
            a.strip() for a in after_colons
            if a.strip()
            and not a.strip().startswith('_')
            and not _instruction_patterns.match(a.strip())
            and not a.strip().endswith('?')
            and 'complete sentences' not in a.lower()
            and 'using evidence' not in a.lower()
            and 'in your own words' not in a.lower()
            and 'from the reading' not in a.lower()
            and 'pp ' not in a.strip()[:5]  # page references like "pp 348-349"
        ]

        # Method 3: Look for paragraph-length responses (likely written answers)
        paragraphs = [p.strip() for p in content.split('\n\n') if len(p.strip()) > 50]
        # Filter out paragraphs that are mostly underscores
        real_paragraphs = [p for p in paragraphs if p.count('_') < len(p) * 0.3]

        # Method 4: Count lines that are JUST underscores (blank response lines)
        blank_lines = len(re.findall(r'^[\s_]+$', content, re.MULTILINE))
        total_lines = len([l for l in content.split('\n') if l.strip()])
        blank_ratio = blank_lines / max(total_lines, 1)

        # Method 5: Check for questions followed by no response
        # Look for question patterns and check if there's content after them
        lines = content.split('\n')
        unanswered_questions = []
        question_indices = []

        for i, line in enumerate(lines):
            line_stripped = line.strip()
            # Check if line is a question (ends with ? or starts with number/bullet or is a vocab term with colon)
            is_question = (
                line_stripped.endswith('?') or
                re.match(r'^[•\*\-\u2022\u2023\u25E6\u2043\u2219\s]*\d+[\.\)]\s*\w', line_stripped) or  # "• 1. Question" or "1) Question"
                re.match(r'^[•\*\-\u2022\u2023\u25E6\u2043\u2219\s]*[a-zA-Z][\.\)]\s*\w', line_stripped) or  # "• a. Question" or "a) Question"
                (line_stripped.endswith(':') and len(line_stripped) > 5 and '_' not in line_stripped)  # "Nationalism:"
            )
            if is_question and len(line_stripped) > 5:
                question_indices.append(i)

        # Check content between consecutive questions
        for idx, q_idx in enumerate(question_indices):
            line_stripped = lines[q_idx].strip()
            # Determine where the next question starts (or end of document)
            next_q_idx = question_indices[idx + 1] if idx + 1 < len(question_indices) else len(lines)

            # Get content between this question and the next
            content_between = []
            for j in range(q_idx + 1, min(next_q_idx, q_idx + 6)):  # Check up to 5 lines after
                if j < len(lines):
                    between_line = lines[j].strip()
                    # Skip empty lines and lines that are just underscores
                    if between_line and not re.match(r'^[_\s\-\.]+$', between_line):
                        # Check if this line has actual content (not just template markers)
                        if len(between_line) > 3 and between_line.count('_') < len(between_line) * 0.5:
                            content_between.append(between_line)

            # If no substantive content found after this question, mark as unanswered
            if not content_between:
                unanswered_questions.append(line_stripped[:60] + "..." if len(line_stripped) > 60 else line_stripped)

        # Determine if submission is blank
        has_filled_blanks = len(filled_blanks) >= 2
        has_written_responses = len(after_colons) >= 2 or len(real_paragraphs) >= 1
        mostly_blank_lines = blank_ratio > 0.4
        many_unanswered = len(unanswered_questions) >= 3

        is_blank = (not has_filled_blanks and not has_written_responses and mostly_blank_lines) or \
                   (many_unanswered and not has_filled_blanks and not has_written_responses)

        if is_blank:
            print(f"  ⚠️  BLANK/EMPTY SUBMISSION DETECTED")
            print(f"      Filled blanks: {len(filled_blanks)}, Written responses: {len(after_colons)}")
            print(f"      Blank line ratio: {blank_ratio:.1%}, Unanswered questions: {len(unanswered_questions)}")
            return {
                "score": 0,
                "letter_grade": "INCOMPLETE",
                "breakdown": {
                    "content_accuracy": 0,
                    "completeness": 0,
                    "critical_thinking": 0,
                    "communication": 0
                },
                "feedback": "You submitted a blank assignment with no responses. Please complete all sections and resubmit.",
                "student_responses": [],
                "unanswered_questions": unanswered_questions[:10],
                "ai_detection": {"flag": "none", "confidence": 0, "reason": "Blank submission — no content to evaluate."},
                "plagiarism_detection": {"flag": "none", "reason": "Blank submission — no content to evaluate."},
                "authenticity_flag": "clean",
                "authenticity_reason": "",
                "skills_demonstrated": {}
            }

    # FERPA: Use anonymous placeholder instead of real student name
    anonymous_name = "Student"
    
    # Build custom instructions section if provided
    custom_section = ''
    if custom_ai_instructions:
        custom_section = f"""
---
TEACHER'S GRADING INSTRUCTIONS (FOLLOW THESE CAREFULLY):
{custom_ai_instructions}
---
"""

    # Build student history context for personalized feedback
    history_context = ''
    if student_id and student_id != "UNKNOWN":
        try:
            history_context = build_history_context(student_id)
        except Exception as e:
            print(f"  Note: Could not load student history: {e}")

    # Build accommodation context for IEP/504 students (FERPA compliant)
    # NOTE: Only accommodation TYPE is sent to AI - no student identifying info
    accommodation_context = ''
    if student_id and student_id != "UNKNOWN":
        try:
            accommodation_context = build_accommodation_prompt(student_id)
            if accommodation_context:
                print(f"  Applying accommodations for student")
        except Exception as e:
            print(f"  Note: Could not load accommodations: {e}")

    # Check if this is a fill-in-the-blank assignment
    is_fitb = False
    content_lower = content.lower() if content else ''
    if 'fill-in' in content_lower or 'fill in the blank' in content_lower or 'fillintheblank' in content_lower.replace(' ', '').replace('-', ''):
        is_fitb = True
    # Also check custom AI instructions for FITB rubric type
    if 'FILL-IN-THE-BLANK' in (custom_ai_instructions or '').upper():
        is_fitb = True
    # Also detect FITB by timestamp pattern (e.g., "1. (0:00)" format common in video worksheets)
    if not is_fitb and content:
        import re as _re
        has_timestamps = bool(_re.search(r'\d+\.\s*\(\d+:\d+', content))
        has_filled_underscores = len(_re.findall(r'_{2,}([^_\n]+)_{2,}', content)) >= 2
        if has_timestamps and has_filled_underscores:
            is_fitb = True
            print(f"  📝 FITB detected via timestamps + filled underscores")

    # PRE-EXTRACT student responses to prevent AI hallucination
    extraction_result = None
    extracted_responses_text = ''
    if assignment_data.get("type") == "text" and content:
        if is_fitb:
            # FITB assignment — send full content for grading (works with or without markers)
            print(f"  📝 FITB assignment - sending full content for grading")
            extracted_responses_text = f"""
==================================================
FILL-IN-THE-BLANK SUBMISSION
==================================================

STUDENT'S COMPLETED DOCUMENT:
{content}

==================================================
Grade based on:
1. ACCURACY: Did the student fill in the blanks correctly?
2. COMPLETENESS: Did the student attempt all blanks?
Do NOT penalize for AI/plagiarism - these are factual answers.
==================================================
"""
            extraction_result = {
                "extracted_responses": [{"question": "Fill-in-the-blank", "answer": content, "type": "fitb_full"}],
                "blank_questions": [],
                "total_questions": 1,
                "answered_questions": 1,
                "extraction_summary": "FITB - full content submitted for grading"
            }
        else:
            # Normal extraction for non-FITB assignments
            # Priority: GRAIDER tag plain-text fallback > custom marker extraction
            if '[GRAIDER:' in content:
                extraction_result = extract_from_graider_text(content, exclude_markers)

            if not extraction_result or not extraction_result.get("extracted_responses"):
                # Debug: Log markers being used
                marker_count = len(custom_markers) if custom_markers else 0
                print(f"  🔍 Extraction using {marker_count} markers")
                if custom_markers and marker_count > 0:
                    for i, m in enumerate(custom_markers[:3]):  # Show first 3
                        marker_text = m.get('start', m) if isinstance(m, dict) else m
                        print(f"      Marker {i+1}: {marker_text[:50]}...")

                extraction_result = extract_student_responses(content, custom_markers, exclude_markers, assignment_template)
            if extraction_result:
                extracted_responses_text = format_extracted_for_grading(extraction_result, marker_config, extraction_mode)
                answered = extraction_result.get("answered_questions", 0)
                total = extraction_result.get("total_questions", 0)
                print(f"  📋 Pre-extracted {answered}/{total} responses")

                # DEBUG: Show what was extracted
                for i, resp in enumerate(extraction_result.get("extracted_responses", [])):
                    q_label = resp.get("question", "?")[:60]
                    ans_preview = resp.get("answer", "")[:100].replace('\n', ' ')
                    print(f"      [{i+1}] {q_label}...")
                    print(f"          Answer: {ans_preview}...")

                # If no responses found, return early with 0 score
                if answered == 0:
                    print(f"  ⚠️  NO RESPONSES EXTRACTED - Document is blank or markers don't match")
                    return {
                        "score": 0,
                        "letter_grade": "INCOMPLETE",
                        "breakdown": {
                            "content_accuracy": 0,
                            "completeness": 0,
                            "critical_thinking": 0,
                            "communication": 0
                        },
                        "feedback": "You submitted a blank assignment with no responses. Please complete all sections and resubmit.",
                        "student_responses": [],
                        "unanswered_questions": extraction_result.get("blank_questions", []) + extraction_result.get("missing_sections", []),
                        "ai_detection": {"flag": "none", "confidence": 0, "reason": "Blank submission — no content to evaluate."},
                        "plagiarism_detection": {"flag": "none", "reason": "Blank submission — no content to evaluate."},
                        "authenticity_flag": "clean",
                        "authenticity_reason": "",
                        "skills_demonstrated": {},
                        "extraction_result": extraction_result
                    }

    # Analyze current submission's writing style for AI detection
    writing_style_context = ''
    current_writing_style = None
    style_comparison = None
    if assignment_data.get("type") == "text" and content:
        current_writing_style = analyze_writing_style(content)
        if current_writing_style:
            # Get student's historical writing profile
            historical_profile = get_writing_profile(student_id) if student_id and student_id != "UNKNOWN" else None

            if historical_profile and historical_profile.get("sample_count", 0) >= 2:
                # Compare current vs historical style
                style_comparison = compare_writing_styles(current_writing_style, historical_profile)

                if style_comparison.get("ai_likelihood") in ["likely", "possible"]:
                    print(f"  ⚠️  Writing style deviation detected: {style_comparison.get('deviation')}")
                    writing_style_context = f"""
---
WRITING STYLE ANALYSIS (COMPARE TO STUDENT'S HISTORY):
This student's historical writing profile (based on {historical_profile.get('sample_count', 0)} previous assignments):
- Average complexity score: {historical_profile.get('avg_complexity_score', 'N/A')}/10
- Average sentence length: {historical_profile.get('avg_sentence_length', 'N/A')} words
- Average word length: {historical_profile.get('avg_word_length', 'N/A')} characters
- Typical academic vocabulary: {historical_profile.get('avg_academic_words', 0):.1f} words per submission

Current submission analysis:
- Complexity score: {current_writing_style.get('complexity_score', 'N/A')}/10
- Sentence length: {current_writing_style.get('avg_sentence_length', 'N/A')} words
- Word length: {current_writing_style.get('avg_word_length', 'N/A')} characters
- Academic vocabulary count: {current_writing_style.get('academic_word_count', 0)}

DEVIATION ALERT: {'; '.join(style_comparison.get('deviations', []))}
This suggests possible AI use - be extra vigilant in your authenticity check!
---
"""

    # Map grade level to age range for context
    grade_age_map = {
        'K': '5-6', '1': '6-7', '2': '7-8', '3': '8-9', '4': '9-10', '5': '10-11',
        '6': '11-12', '7': '12-13', '8': '13-14', '9': '14-15', '10': '15-16',
        '11': '16-17', '12': '17-18'
    }
    age_range = grade_age_map.get(str(grade_level), '11-12')

    # Build extracted responses section for the prompt
    extracted_responses_section = ""
    if extracted_responses_text:
        extracted_responses_section = f"""
---
{extracted_responses_text}
---
"""

    # Build assignment template section (provides question context)
    assignment_template_section = ""
    if assignment_template:
        # Truncate if very long to save tokens
        template_text = assignment_template[:8000] if len(assignment_template) > 8000 else assignment_template
        assignment_template_section = f"""
---
ASSIGNMENT TEMPLATE (The questions/prompts the student was asked to answer):
{template_text}
---
"""

    # Use custom rubric if provided, otherwise use default
    # If marker_config is provided, build a section-based rubric
    section_rubric = ""
    if marker_config:
        section_rubric = build_section_rubric(marker_config, effort_points)

    effective_rubric = rubric_prompt if rubric_prompt else GRADING_RUBRIC

    # Build extraction mode-specific instructions
    if extraction_mode == 'ai':
        extraction_instructions = """
CRITICAL - AI EXTRACTION MODE:
The content above contains RAW section content that includes BOTH prompts/questions AND student responses.
YOUR JOB is to identify what is a prompt/question vs what the student actually wrote.

IDENTIFYING STUDENT RESPONSES:
- Questions end with "?" - anything AFTER the "?" on the same line is the student's answer
- Vocabulary format "Term:" - anything after the ":" is the student's definition
- Prompts like "Write your answer:", "Explain...", etc. are instructions, not student content
- Look for the actual student-written text, which is usually less formal and may have spelling errors
- If a section has only a prompt with no student response, mark it as BLANK/UNANSWERED

For example:
- "1. What year was the Louisiana Purchase? 1803" → Student answer is "1803"
- "Antebellum: the period before the war" → Student answer is "the period before the war"
- "Antebellum:" with nothing after → BLANK, student didn't answer
- "Write your answer:" with nothing after → BLANK

Be thorough in separating prompts from responses. Students often write answers on the same line as questions.
"""
    else:
        extraction_instructions = """
CRITICAL - PRE-EXTRACTED RESPONSES:
The student responses have been PRE-EXTRACTED from the document and listed above.
DO NOT invent or hallucinate any responses that are not in the VERIFIED STUDENT RESPONSES section.
ONLY grade the responses that were explicitly extracted and shown to you.
If a question is listed as "UNANSWERED", it means the student left it blank - do not imagine an answer.
"""

    # Build grading style instructions
    if grading_style == 'lenient':
        grading_style_instructions = """
GRADING APPROACH: LENIENT
- Prioritize EFFORT over perfection. If a student attempted a section, they deserve significant credit.
- Brief answers that show understanding should receive 70-80% credit even if not fully developed.
- Do NOT penalize short answers if they demonstrate the student understood the material.
- A student who attempts ALL sections with genuine effort should receive at minimum a B (80+).
- Writing quality expectations should be relaxed - focus on content understanding, not eloquence.
- One-sentence answers to open-ended questions are acceptable if they show comprehension.
- Reserve grades below C only for students who clearly did not try or left sections truly blank/missing.
"""
    elif grading_style == 'strict':
        grading_style_instructions = """
GRADING APPROACH: STRICT
- Hold students to high standards for their grade level.
- Brief, underdeveloped answers to open-ended questions should be penalized.
- Writing quality matters - expect complete sentences and proper grammar.
- Full credit requires thorough, well-developed responses that demonstrate deep understanding.
- Partial answers receive proportionally reduced credit.
"""
    else:
        grading_style_instructions = """
GRADING APPROACH: STANDARD
- Balance accuracy, completeness, writing quality, and effort evenly per the rubric.
- Brief answers should receive partial credit proportional to their quality.
- Students should be encouraged but held to grade-level expectations.
"""

    # Build authenticity/detection section — per-section FITB awareness
    if is_fitb and not custom_markers:
        # Pure FITB (no markers) — skip all detection
        fitb_authenticity_section = f"""AUTHENTICITY CHECKS - FILL-IN-THE-BLANK EXEMPTION:
This is a fill-in-the-blank assignment. Students are expected to write short factual answers (names, dates, places, vocabulary terms).
These answers will naturally match textbook/source material — that is the CORRECT behavior, NOT plagiarism or AI use.
DO NOT flag any answers for AI use or plagiarism. Set ai_detection to "none" and plagiarism_detection to "none".
Focus ONLY on whether the answers are factually correct and complete.

HARD CAPS FOR INCOMPLETE WORK (MANDATORY - each skipped section = one letter grade drop):
Count the number of skipped/unanswered blanks or sections:
- 0 skipped = eligible for A (up to 100)
- 1 skipped = MAXIMUM 89 (B) - NO EXCEPTIONS
- 2 skipped = MAXIMUM 79 (C) - NO EXCEPTIONS
- 3 skipped = MAXIMUM 69 (D) - NO EXCEPTIONS
- 4+ skipped = MAXIMUM 59 (F)
- 5+ skipped = MAXIMUM 49 (F)
- 6+ skipped = MAXIMUM 39 (F)

NEARLY BLANK SUBMISSIONS - Score based on effort shown:
- If student answered only 1-2 blanks out of 5+ = score 10-25 (F)
- If student answered only 3-4 blanks out of 10+ = score 25-40 (F)
- Empty or nearly empty = score in the 0-30 range"""
    else:
        # Full detection — with FITB exemption note for hybrid assignments
        fitb_exemption_note = ""
        if is_fitb:
            fitb_exemption_note = """IMPORTANT - FILL-IN-THE-BLANK EXEMPTION:
This assignment contains fill-in-the-blank sections mixed with written response sections.
Fill-in-the-blank answers (short factual responses like names, dates, places, vocabulary terms) are EXEMPT from AI/plagiarism detection.
These answers are EXPECTED to match textbook/source material — that is correct behavior, NOT cheating.
ONLY apply AI/plagiarism checks to WRITTEN responses (paragraphs, reflections, summaries, explanations).
Do NOT use fill-in-the-blank answers as evidence of AI use or plagiarism.

"""
        fitb_authenticity_section = f"""{fitb_exemption_note}CRITICAL - AUTHENTICITY CHECKS (YOU MUST CHECK THIS CAREFULLY!):

1. AI DETECTION - Compare the student's simple answers to their written paragraphs:
STEP 1: Look at their short answers (fill-in-blanks, one-word responses). Note the vocabulary level.
STEP 2: Look at their paragraph responses. Compare the vocabulary and complexity.
STEP 3: If there's a MISMATCH (simple short answers but sophisticated paragraphs), flag as "likely" AI.

AUTOMATIC "likely" AI FLAGS - if you see ANY of these phrases, it's 100% AI:
- "transformed the nation into a continental power"
- "transforming a limited mission"
- "historic deal that doubled"
- "fueling westward expansion"
- "triggered intense political debates"
- "spurred exploration"
- "fundamentally altered the trajectory"
- "establishing the precedent for"
- "constitutional questions regarding federal authority"
- "resonate through subsequent decades"
- "vital for trade and growth"
- "securing vital trade routes"
- "manifest destiny"
- "territorial expansion"
- "abundant natural resources"
- Any phrase starting with "Transforming...", "Establishing...", "Securing..."
- Any phrase a {age_range} year old would NEVER write

CRITICAL CONTRAST CHECK - THIS IS THE MOST IMPORTANT CHECK:
Look at the student's spelling and grammar in simple answers. If they write:
- Misspellings like "Tomas Jefferson", "the u's", "france" (lowercase)
- Simple phrases like "It doubled in size", "idk"
- Basic vocabulary and short sentences

BUT THEN write sophisticated phrases like:
- "Transforming a limited mission to buy New Orleans into a historic deal"
- Any sentence with words like "vital", "securing", "expanding", "historic deal"

That is 100% AI or copied - flag as "likely" IMMEDIATELY. A student who misspells "Thomas" does NOT write "transforming a limited mission into a historic deal."

Real grade {grade_level} students write: "it made the US bigger", "they needed the river for boats", "so ships could go there"
AI writes: "it transformed the nation into a continental power", "securing vital trade routes"

OBVIOUS COPY-PASTE DEFINITIONS (flag as PLAGIARISM "likely" IMMEDIATELY):
- "exclusive possession or control of the supply of or trade in a commodity or service" = Google definition of monopoly
- "an ideology emphasizing intense loyalty to one's nation" = textbook definition
- "government-provided financial incentives" = too sophisticated
- "implementing three interconnected policies" = not student language
- "authority not explicitly stated in the U.S. Constitution but deemed necessary" = textbook definition of implied powers
- ANY definition that sounds like it was copied from a dictionary or Wikipedia = PLAGIARISM

A real {age_range} year old defines monopoly as: "when one company controls everything" or "only one person sells something"
NOT: "exclusive possession or control of the supply of or trade in a commodity or service"

2. PLAGIARISM DETECTION - Look for:
- SUDDEN SHIFTS in writing quality (simple answers + sophisticated paragraphs = copied/AI)
- Textbook-perfect definitions that don't match the student's other answers
- Phrases that sound memorized or copied verbatim
- Statistics or specific numbers not in the reading (like "828,000 square miles")
- DICTIONARY DEFINITIONS - If a vocabulary definition sounds like it was copied from Google/dictionary (e.g., "exclusive possession or control of the supply of or trade in a commodity or service" for monopoly), flag as PLAGIARISM "likely"
- SOPHISTICATED VOCABULARY MISMATCH - Words like "ideology", "emphasizing", "fostering", "interconnected", "implementing" are NOT how grade {grade_level} students write definitions
- FRAGMENT ANSWERS that are clearly pasted (no complete sentence, just a definition dump)
- If student writes simple answers for some questions but sophisticated definitions for vocabulary = COPY/PASTE from Google

HARD CAPS FOR AI USE / PLAGIARISM (apply FIRST, before other caps):
- AI flag "likely" = MAX score is 50 (F) - this is cheating
- AI flag "possible" = MAX score is 65 (D) - suspicious, needs verification
- Plagiarism flag "likely" = MAX score is 50 (F) - this is cheating
- Plagiarism flag "possible" = MAX score is 65 (D) - suspicious
- If BOTH AI and plagiarism are flagged = MAX score is 40 (F)

In feedback for AI/plagiarism flags:
- Clearly state the work appears to be AI-generated or copied
- Explain that academic integrity is important
- Recommend the student redo the assignment in their own words
- Note this will be reviewed by the teacher

THEN apply HARD CAPS FOR INCOMPLETE WORK (MANDATORY - each skipped section = one letter grade drop):
Count the number of skipped/unanswered written sections (Student Task, Reflection, Explain, etc.):
- 0 skipped = eligible for A (up to 100)
- 1 skipped = MAXIMUM 89 (B) - NO EXCEPTIONS
- 2 skipped = MAXIMUM 79 (C) - NO EXCEPTIONS
- 3 skipped = MAXIMUM 69 (D) - NO EXCEPTIONS
- 4+ skipped = MAXIMUM 59 (F)
- 5+ skipped = MAXIMUM 49 (F)
- 6+ skipped = MAXIMUM 39 (F)

NEARLY BLANK SUBMISSIONS - Score based on effort shown:
- If student answered only 1-2 questions out of 5+ = score 10-25 (F)
- If student answered only 3-4 questions out of 10+ = score 25-40 (F)
- A single poor answer like "idk" or "going again" does NOT earn 50 points
- The score should reflect ACTUAL WORK DONE, not a default minimum
- Empty or nearly empty = score in the 0-30 range

YOU MUST APPLY THESE CAPS. If a student skipped 2 written sections, their score CANNOT be above 79 even if their fill-in-the-blanks were perfect.

The LOWEST cap wins. Example: AI "likely" (cap 50) + 6 sections skipped (cap 39) = final cap is 39."""

    # Load ELL designation for this student (teacher-controlled bilingual feedback)
    ell_language = None
    if student_id and student_id != "UNKNOWN":
        ell_file = os.path.expanduser("~/.graider_data/ell_students.json")
        if os.path.exists(ell_file):
            try:
                with open(ell_file, 'r', encoding='utf-8') as f:
                    ell_data = json.load(f)
                ell_entry = ell_data.get(student_id, {})
                lang = ell_entry.get("language")
                if lang and lang != "none":
                    ell_language = lang
            except Exception:
                pass

    # Always grade in English only — bilingual translation is handled as a separate post-grading step
    ell_instruction = "Write feedback in English only."

    # Repeat teacher instructions at end of prompt — placed here so the AI sees them LAST
    # before generating its response. Earlier placement gets buried under 80+ lines of default rules.
    teacher_override_section = ""
    if custom_ai_instructions and custom_ai_instructions.strip():
        teacher_override_section = f"""

FINAL AUTHORITY — TEACHER'S GRADING INSTRUCTIONS (repeated here because they override ALL defaults above):
{custom_ai_instructions}

^^^ THESE INSTRUCTIONS OVERRIDE the default scoring rules. If the teacher says to be lenient, be lenient.
If the teacher says to accept general definitions, accept them. Do NOT contradict the teacher's instructions
in your scoring or feedback. The teacher knows their students better than any rubric."""

    prompt_text = f"""
{effective_rubric}

{section_rubric}

{ASSIGNMENT_INSTRUCTIONS}
{custom_section}
{grading_style_instructions}
{accommodation_context}
{history_context}
{writing_style_context}
{assignment_template_section}
---

STUDENT CONTEXT:
- Grade Level: {grade_level}
- Subject: {subject}
- Expected Age Range: {age_range} years old

{extracted_responses_section}

{extraction_instructions}

IMPORTANT: Assess ALL sections that appear in the EXTRACTED RESPONSES, UNANSWERED QUESTIONS, and MISSING SECTIONS above.
If a section appears in MISSING SECTIONS, the student entirely omitted a required part of the assignment — penalize accordingly.
If a section appears in UNANSWERED QUESTIONS, the student left a required section blank — penalize accordingly.
Only the extracted/marked sections, unanswered questions, and missing sections count toward the grade.

Your "student_responses" field MUST contain ONLY the raw answer text from each "STUDENT ANSWER:" line above.
Do NOT include question numbers, section names, or labels like "[1] Summary:" - just the student's actual written text.
Example: If the verified response shows 'STUDENT ANSWER: "The treaty was signed in 1803"', your student_responses should contain "The treaty was signed in 1803" - not "Summary: The treaty was signed in 1803".
If no responses were extracted, the student gets a 0.

For MATCHING exercises specifically:
- Look for numbers placed next to vocabulary terms in the extracted responses
- The number indicates which definition the student chose
- Grade whether they matched correctly

GRADING GUIDELINES:
- Assess EVERY answer the student provided.
- For fill-in-the-blank: check if the answer is factually correct or close enough.
- Accept multiple valid answers and synonyms.
- DO NOT penalize spelling mistakes if the meaning is clear.
- Be age-appropriate - these are grade {grade_level} students ({age_range} years old).
- IMPORTANT: If the teacher provided custom grading instructions above, follow them carefully.

CRITICAL - COMPLETENESS REQUIREMENTS:
- Check the EXTRACTED RESPONSES, UNANSWERED QUESTIONS, and MISSING SECTIONS lists above.
- MISSING SECTIONS are sections the teacher REQUIRED but the student ENTIRELY OMITTED from their submission.
  Each missing section must be treated the SAME as a skipped section — it lowers the grade by one full letter.
- UNANSWERED QUESTIONS are sections the student included but left blank — also penalize.
- IMPORTANT: Individual vocabulary terms or bullet points WITHIN a section that HAS a student answer are NOT unanswered questions. Only count items explicitly listed in the UNANSWERED QUESTIONS section above.
- For the sections that WERE extracted, check if the student answered them adequately, especially:
  * "Explain in your own words" sections - these require written responses, not blank
  * "Reflection" or "Final Reflection" questions - these MUST be answered
  * "Student Task" sections - these are major components requiring written responses
  * Any prompt asking students to "Write a few sentences" or "Describe" or "Explain"
  * Summary sections
  * Primary source analysis tasks
- Skipping or omitting written sections shows AVOIDANCE OF EFFORT and must be penalized!
- Count total skipped = UNANSWERED QUESTIONS + MISSING SECTIONS.
{f"""- LENIENT PENALTY SCALE (teacher selected lenient grading):
  * 0 sections skipped/missing = eligible for A (90-100)
  * 1 section skipped/missing = eligible for A/B (85-95) - minor deduction only
  * 2 sections skipped/missing = maximum B (80-89)
  * 3 sections skipped/missing = maximum C (70-79)
  * 4+ sections skipped/missing = maximum D (60-69)
  * ALL or nearly all sections blank = 0 (INCOMPLETE) - student did not attempt the assignment
- Students who attempt most sections with genuine effort should receive at minimum a C (70+).
- An "A" grade (90+) is possible if ALL answered sections show genuine effort and understanding.""" if grading_style == 'lenient' else f"""- STRICT PENALTY SCALE (teacher selected strict grading):
  * 0 sections skipped/missing = eligible for A (90-100)
  * 1 section skipped/missing = maximum B- (80-85) - dropped one letter
  * 2 sections skipped/missing = maximum C (70-75) - dropped two letters
  * 3 sections skipped/missing = maximum D (60-65) - dropped three letters
  * 4+ sections skipped/missing = F (below 60) - shows no effort
  * ALL or nearly all sections blank = 0 (INCOMPLETE) - student did not attempt the assignment
- Students who ONLY do fill-in-the-blanks and skip ALL written responses = maximum D (65)
- An "A" grade (90+) is ONLY possible if ALL sections are completed with thorough, well-developed responses.""" if grading_style == 'strict' else """- STANDARD PENALTY SCALE:
  * 0 sections skipped/missing = eligible for A (90-100)
  * 1 section skipped/missing = maximum B (80-89) - dropped one letter
  * 2 sections skipped/missing = maximum C (70-79) - dropped two letters
  * 3 sections skipped/missing = maximum D (60-69) - dropped three letters
  * 4+ sections skipped/missing = F (below 60) - shows no effort on written work
  * ALL or nearly all sections blank = 0 (INCOMPLETE) - student did not attempt the assignment
- Students who ONLY do fill-in-the-blanks and skip ALL written responses = maximum C (75)
- An "A" grade (90+) is ONLY possible if ALL sections are completed with quality responses."""}
- This applies to ALL assignments - skipping reflections, explanations, or analysis tasks is unacceptable
- In the "unanswered_questions" field, ONLY list items from the UNANSWERED QUESTIONS and MISSING SECTIONS lists above — do NOT invent new unanswered items from individual vocab terms or bullet points within answered sections

{fitb_authenticity_section}
{teacher_override_section}
Provide your response in the following JSON format ONLY (no other text):
{{
    "score": <FIRST calculate raw score, THEN apply the caps above. If 2 sections skipped, max is 79>,
    "letter_grade": "<A, B, C, D, or F - must match the capped score>",
    "breakdown": {{
        "content_accuracy": <points out of 40 - correctness of answers>,
        "completeness": <points out of 25 - ALL sections must be attempted. MISSING SECTIONS and UNANSWERED QUESTIONS both count as skipped. Written responses (reflections, explanations, Student Tasks) count heavily! 0-5 if 2+ major sections skipped/missing, 6-12 if 1 major section skipped/missing, 13-20 if minor gaps only, 21-25 only if ALL parts fully completed>,
        "writing_quality": <points out of 20>,
        "effort_engagement": <points out of 15>
    }},
    "student_responses": ["<EXTRACT ONLY the actual answer text that appears after 'STUDENT ANSWER:' in the verified responses above. Do NOT include the question/section name, number, or label. WRONG: 'Summary: The treaty was...' or '[1] Summary: The treaty...' - RIGHT: 'The treaty was signed in 1803 and...' - just the raw answer text the student wrote>"],
    "unanswered_questions": ["<ONLY list sections/questions from the UNANSWERED QUESTIONS and MISSING SECTIONS lists above. Do NOT list individual vocab terms or bullet points that appear WITHIN a section the student completed — those are part of the student's response, not separate unanswered questions. If a section has a STUDENT ANSWER with content, it is NOT unanswered even if individual terms within it seem brief.>"],
    "excellent_answers": ["<Quote 2-4 specific answers that were particularly strong, accurate, or showed great understanding. Include the exact text the student wrote.>"],
    "needs_improvement": ["<Quote 1-3 specific answers that were incorrect or incomplete, along with what the correct/better answer would be. Format: 'You wrote [X] but [correct info]' or 'For the question about [topic], [guidance]'>"],
    "skills_demonstrated": {{
        "strengths": ["<List 2-4 specific skills the student showed strength in. Go BEYOND the rubric categories - identify skills like: reading comprehension, critical thinking, source analysis, making connections, vocabulary usage, following directions, organization, creativity, historical thinking, cause-and-effect reasoning, comparing/contrasting, using evidence, drawing conclusions, summarizing, note-taking, attention to detail, etc. Only include skills clearly demonstrated in THIS assignment.>"],
        "developing": ["<List 1-2 skills the student is still developing or struggled with. Same skill types as above. Be specific about what skill needs work based on their answers.>"]
    }},
    "ai_detection": {{
        "flag": "<none, unlikely, possible, or likely>",
        "confidence": <number 0-100 representing confidence in the assessment>,
        "reason": "<Brief explanation if not 'none', otherwise empty string>"
    }},
    "plagiarism_detection": {{
        "flag": "<none, possible, or likely>",
        "reason": "<Brief explanation if not 'none', otherwise empty string>"
    }},
    "feedback": "<Write 3-4 paragraphs of thorough, personalized feedback that sounds like a real teacher wrote it - warm, encouraging, and specific. IMPORTANT GUIDELINES: 1) VARY your sentence structure and openings - don't start every sentence the same way. Mix short punchy sentences with longer ones. 2) QUOTE specific answers from the student's work when praising them (e.g., 'I loved how you explained that [quote their answer]' or 'Your answer about [topic] - '[their exact words]' - shows real understanding'). 3) When mentioning areas to improve, be gentle and constructive - reference specific questions they struggled with and give them a hint or the right direction. 4) Sound HUMAN - use contractions (you're, that's, I'm), occasional casual phrases ('Nice!', 'Great thinking here'), and vary your enthusiasm. 5) End with genuine encouragement that connects to something specific they did well. 6) Do NOT use the student's name - say 'you' or 'your'. 7) Avoid repetitive phrases like 'Great job!' at the start of every paragraph - mix it up! 8) IF STUDENT HISTORY IS PROVIDED ABOVE: Reference their progress! Mention streaks, acknowledge CONSISTENT SKILLS (e.g., 'Your reading comprehension continues to be a real strength!'), celebrate IMPROVING SKILLS (e.g., 'I notice your critical thinking is getting sharper - great progress!'), and gently encourage SKILLS TO DEVELOP (e.g., 'Keep working on making connections between ideas'). Connect current work to past achievements when relevant. 9) BILINGUAL FEEDBACK: {ell_instruction}>"
}}
"""

    print(f"  🤖 Grading with AI...")

    try:
        # FERPA COMPLIANCE: Sanitize PII from text content before sending to AI
        if assignment_data["type"] == "text":
            original_content = assignment_data['content']
            anon_id, sanitized_content = sanitize_pii_for_ai(student_name, original_content)

            # Log if any PII was removed (for audit trail)
            if sanitized_content != original_content:
                print(f"  🔒 PII sanitized from submission before AI processing")

            # HARD BLOCK: Only send extracted responses to prevent hallucination
            # If extraction succeeded, use ONLY extracted responses (not raw content)
            if extracted_responses_text:
                # Send only the pre-extracted verified responses
                print(f"  ✅ Using ONLY pre-extracted responses (hallucination prevention)")
                full_prompt = prompt_text + f"\n\nSTUDENT'S VERIFIED RESPONSES (extracted from document):\n{extracted_responses_text}"
            else:
                # Extraction failed or found nothing - REQUIRES MANUAL REVIEW
                print(f"  ⚠️  HARD BLOCK: No responses extracted - flagging for manual review")
                return {
                    "score": 0,
                    "letter_grade": "MANUAL REVIEW",
                    "breakdown": {
                        "content_accuracy": 0,
                        "completeness": 0,
                        "critical_thinking": 0,
                        "communication": 0
                    },
                    "feedback": "⚠️ MANUAL REVIEW REQUIRED: The automated extraction could not find student responses in this document. This could mean:\n\n1. The document is blank or nearly blank\n2. The formatting is unusual\n3. The student wrote in unexpected locations\n\nPlease open the original document and grade manually to prevent AI hallucination.",
                    "student_responses": [],
                    "unanswered_questions": [],
                    "authenticity_flag": "manual_review",
                    "authenticity_reason": "Extraction failed - cannot verify responses",
                    "skills_demonstrated": {},
                    "requires_manual_review": True
                }
            messages = [{"role": "user", "content": full_prompt}]

        elif assignment_data["type"] == "image":
            # Image-based assignment - use vision
            # WARNING: Cannot extract/verify responses from images - higher hallucination risk
            print(f"  ⚠️  IMAGE SUBMISSION: Cannot pre-extract responses - recommend spot-checking")
            messages = [{
                "role": "user",
                "content": [
                    {"type": "text", "text": prompt_text + "\n\nSTUDENT'S WORK (see attached image):\nIMPORTANT: Only grade what you can CLEARLY see in the image. If text is unclear or cut off, mark as incomplete rather than guessing."},
                    {
                        "type": "image_url",
                        "image_url": {
                            "url": f"data:{assignment_data['media_type']};base64,{assignment_data['content']}"
                        }
                    }
                ]
            }]
            # Flag that this is an unverified image submission
            extraction_result = {"type": "image", "verified": False}
        else:
            return {"score": 0, "letter_grade": "ERROR", "breakdown": {}, "feedback": "Unknown content type"}

        # Make API call based on provider
        if provider == "anthropic":
            # Claude API call
            if assignment_data.get("type") == "image":
                claude_content = [
                    {"type": "text", "text": prompt_text + "\n\nSTUDENT'S WORK (see attached image):\nIMPORTANT: Only grade what you can CLEARLY see in the image. If text is unclear or cut off, mark as incomplete rather than guessing."},
                    {
                        "type": "image",
                        "source": {
                            "type": "base64",
                            "media_type": assignment_data['media_type'],
                            "data": assignment_data['content']
                        }
                    }
                ]
            else:
                claude_content = messages[0]["content"] if isinstance(messages[0]["content"], str) else messages[0]["content"][0]["text"]

            response = claude_client.messages.create(
                model=actual_model,
                max_tokens=2000,
                messages=[{"role": "user", "content": claude_content}]
            )
            if token_tracker:
                token_tracker.record_anthropic(response, actual_model)
            response_text = response.content[0].text.strip()

        elif provider == "gemini":
            # Gemini API call with retry for rate limits
            import time
            max_retries = 3
            retry_delay = 2  # seconds

            for attempt in range(max_retries):
                try:
                    if assignment_data.get("type") == "image":
                        import base64
                        image_data = base64.b64decode(assignment_data['content'])
                        image_part = {
                            "mime_type": assignment_data['media_type'],
                            "data": image_data
                        }
                        full_prompt = prompt_text + "\n\nSTUDENT'S WORK (see attached image):\nIMPORTANT: Only grade what you can CLEARLY see in the image. If text is unclear or cut off, mark as incomplete rather than guessing."
                        response = gemini_client.generate_content([full_prompt, image_part])
                    else:
                        text_content = messages[0]["content"] if isinstance(messages[0]["content"], str) else messages[0]["content"][0]["text"]
                        response = gemini_client.generate_content(text_content)
                    if token_tracker:
                        token_tracker.record_gemini(response, actual_model)
                    response_text = response.text.strip()
                    break  # Success, exit retry loop
                except Exception as e:
                    if "429" in str(e) and attempt < max_retries - 1:
                        print(f"  ⏳ Rate limited, retrying in {retry_delay}s... (attempt {attempt + 1}/{max_retries})")
                        time.sleep(retry_delay)
                        retry_delay *= 2  # Exponential backoff
                    else:
                        raise  # Re-raise if not rate limit or out of retries

        else:
            # OpenAI API call with structured output for guaranteed schema
            try:
                response = openai_client.beta.chat.completions.parse(
                    model=ai_model,
                    messages=messages,
                    response_format=GradingResponse,
                    max_tokens=2000,
                    temperature=0,
                    seed=42
                )
                if token_tracker:
                    token_tracker.record_openai(response, ai_model)
                parsed = response.choices[0].message.parsed
                if parsed:
                    result = parsed.model_dump()
                    original_text = json.dumps(result)
                    print(f"  ✅ Structured output parsed successfully")
                    # Skip all JSON cleanup — jump straight to post-processing below
                else:
                    # Model refused or structured parse failed — fall back to text
                    response_text = response.choices[0].message.content or ""
                    print(f"  ⚠️  Structured output empty, falling back to text parse")
                    result = None
            except Exception as structured_err:
                # Structured output not supported for this model — fall back to standard call
                print(f"  ⚠️  Structured output failed ({structured_err}), falling back to standard API")
                response = openai_client.chat.completions.create(
                    model=ai_model,
                    messages=messages,
                    max_tokens=2000,
                    temperature=0,
                    seed=42
                )
                if token_tracker:
                    token_tracker.record_openai(response, ai_model)
                response_text = response.choices[0].message.content.strip()
                result = None

            # If structured output succeeded, skip text parsing
            if result is not None:
                pass  # result already set from parsed.model_dump()
            else:
                # Text fallback: clean up and parse JSON manually
                if response_text.startswith("```"):
                    lines = response_text.split('\n')
                    start = 1 if lines[0].startswith("```") else 0
                    end = len(lines)
                    for i in range(len(lines)-1, -1, -1):
                        if lines[i].strip() == "```":
                            end = i
                            break
                    response_text = '\n'.join(lines[start:end])
                response_text = response_text.strip()
                original_text = response_text

                try:
                    result = json.loads(response_text)
                except json.JSONDecodeError:
                    # Try basic JSON repair for text fallback
                    result = _try_parse_json_fallback(response_text)
                    if result is None:
                        raise

        # For Claude/Gemini providers, parse their text response
        if provider in ("anthropic", "gemini"):
            # Clean up response (remove markdown code blocks if present)
            if response_text.startswith("```"):
                lines = response_text.split('\n')
                start = 1 if lines[0].startswith("```") else 0
                end = len(lines)
                for i in range(len(lines)-1, -1, -1):
                    if lines[i].strip() == "```":
                        end = i
                        break
                response_text = '\n'.join(lines[start:end])
            response_text = response_text.strip()
            original_text = response_text

            result = _try_parse_json_fallback(response_text)
            if result is None:
                raise json.JSONDecodeError("Failed to parse response", response_text, 0)

        # Post-processing: fix double-escaped newlines from some AI providers
        feedback = result.get("feedback", "")
        if "\\n" in feedback:
            feedback = feedback.replace("\\n", "\n")
            result["feedback"] = feedback

        # Strip any accidental bilingual sections from grading response
        if "\n---\n" in feedback:
            feedback = feedback.split("\n---\n")[0].strip()
            result["feedback"] = feedback

        # Two-pass bilingual feedback: translate via separate dedicated API call
        if ell_language and result.get("feedback"):
            print(f"  🌐 Translating feedback to {ell_language}...")
            translated = _translate_feedback(result["feedback"], ell_language, ai_model, token_tracker=token_tracker)
            if translated:
                result["feedback"] = result["feedback"] + "\n\n---\n\n" + translated
                print(f"  ✅ Bilingual feedback added ({ell_language})")
            else:
                print(f"  ⚠️  Translation failed, feedback remains English only")

        # Update student's writing profile (only if not flagged as AI)
        # This builds their baseline for future AI detection
        if student_id and student_id != "UNKNOWN" and current_writing_style:
            ai_flag = result.get("ai_detection", {}).get("flag", "none")
            if ai_flag not in ["likely", "possible"]:
                try:
                    update_writing_profile(student_id, current_writing_style, student_name)
                    print(f"  📊 Updated writing profile for {student_name}")
                except Exception as e:
                    print(f"  Note: Could not update writing profile: {e}")

        # Add style comparison info to result for transparency
        if style_comparison and style_comparison.get("ai_likelihood") in ["likely", "possible"]:
            result["writing_style_deviation"] = style_comparison

        # Add audit trail data
        result["_audit"] = {
            "ai_input": extracted_responses_section,
            "ai_response": original_text
        }

        # Add token usage tracking
        if token_tracker:
            result["token_usage"] = token_tracker.summary()

        return result

    except json.JSONDecodeError as e:
        print(f"  ⚠️  Error parsing AI response: {e}")
        # Try to show response content for debugging
        try:
            raw_preview = response_text[:800] if len(response_text) > 800 else response_text
            print(f"  ⚠️  Raw response preview:\n{raw_preview}")
            # Write full response to temp file for debugging
            import tempfile
            debug_file = tempfile.NamedTemporaryFile(mode='w', suffix='_graider_debug.json', delete=False)
            debug_file.write(response_text)
            debug_file.close()
            print(f"  ⚠️  Full response saved to: {debug_file.name}")
        except:
            print(f"  ⚠️  Could not display response")

        # Try to extract key fields with regex as fallback
        try:
            score_match = re.search(r'"score":\s*(\d+)', response_text)
            grade_match = re.search(r'"letter_grade":\s*"([A-F])"', response_text)
            feedback_match = re.search(r'"feedback":\s*"([^"]{20,500})', response_text)

            if score_match and grade_match:
                print(f"  ✅ Recovered score/grade from malformed JSON")
                return {
                    "score": int(score_match.group(1)),
                    "letter_grade": grade_match.group(1),
                    "breakdown": {"content_accuracy": 0, "completeness": 0, "writing_quality": 0, "effort_engagement": 0},
                    "feedback": feedback_match.group(1) + "..." if feedback_match else "Grading completed but response was malformed.",
                    "student_responses": [],
                    "unanswered_questions": [],
                    "ai_detection": {"flag": "none", "confidence": 0, "reason": ""},
                    "plagiarism_detection": {"flag": "none", "reason": ""},
                    "skills_demonstrated": {"strengths": [], "developing": []},
                    "json_recovery": True
                }
        except:
            pass

        return {
            "score": 0,
            "letter_grade": "ERROR",
            "breakdown": {},
            "feedback": f"Error grading - AI returned invalid JSON. Please review manually."
        }
    except Exception as e:
        print(f"  ⚠️  API error: {e}")
        import traceback
        traceback.print_exc()
        return {
            "score": 0,
            "letter_grade": "ERROR",
            "breakdown": {},
            "feedback": f"API error: {e}"
        }


# =============================================================================
# EMAIL GENERATION
# =============================================================================

def generate_email_content(student_info: dict, grade_result: dict, assignment_name: str) -> tuple:
    """
    Generate email subject and body for a student.
    
    Returns: (subject, body)
    """
    first_name = student_info.get('first_name', 'Student').split()[0]  # Just first name
    
    subject = f"Grade for {assignment_name}: {grade_result['letter_grade']}"
    
    body = f"""Hi {first_name},

Here is your grade and feedback for {assignment_name}:

GRADE: {grade_result['score']}/100 ({grade_result['letter_grade']})

FEEDBACK:
{grade_result['feedback']}

If you have any questions about your grade, please see me during class.

- Mr. Crionas US History
"""
    return subject, body


def save_emails_to_folder(grades: list, output_folder: str, teacher_name: str = '', subject: str = '', school_name: str = ''):
    """
    Save emails as individual text files - ONE EMAIL PER STUDENT
    with feedback for ALL their assignments combined.
    """
    email_folder = Path(output_folder) / "emails"
    email_folder.mkdir(parents=True, exist_ok=True)
    
    # Group grades by student
    students = {}
    for grade in grades:
        student_name = grade.get('student_name', 'Unknown')
        if student_name not in students:
            # Get only first name (no middle initial)
            full_first = grade.get('first_name', student_name.split()[0])
            first_only = full_first.split()[0] if full_first else student_name.split()[0]
            students[student_name] = {
                'email': grade.get('email', ''),
                'first_name': first_only,
                'assignments': []
            }
        students[student_name]['assignments'].append(grade)
    
    email_count = 0
    for student_name, data in students.items():
        if not data['email']:
            continue
        
        # Build combined email
        assignments = data['assignments']
        first_name = data['first_name']
        
        # Email subject line
        if len(assignments) == 1:
            email_subject = f"Grade for {assignments[0].get('assignment', 'Assignment')}: {assignments[0]['letter_grade']}"
        else:
            email_subject = f"Grades for {len(assignments)} Assignments"
        
        # Body
        body = f"Hi {first_name},\n\n"
        
        if len(assignments) == 1:
            a = assignments[0]
            body += f"Here is your grade and feedback for {a.get('assignment', 'your assignment')}:\n\n"
            body += f"GRADE: {a['score']}/100 ({a['letter_grade']})\n\n"
            body += f"FEEDBACK:\n{a.get('feedback', 'No feedback available.')}\n"
        else:
            body += "Here are your grades and feedback:\n\n"
            for a in assignments:
                assignment_name = a.get('assignment', 'Assignment')
                body += f"{'='*50}\n"
                body += f"📝 {assignment_name}\n"
                body += f"{'='*50}\n"
                body += f"GRADE: {a['score']}/100 ({a['letter_grade']})\n\n"
                body += f"FEEDBACK:\n{a.get('feedback', 'No feedback available.')}\n\n"
        
        body += "\nIf you have any questions about your grades, please see me during class.\n\n"
        signature = teacher_name if teacher_name else "Your Teacher"
        if subject:
            signature += f" {subject}"
        body += f"- {signature}\n"
        
        # Save file
        safe_name = re.sub(r'[^\w\s-]', '', student_name).replace(' ', '_')
        filepath = email_folder / f"{safe_name}_email.txt"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"TO: {data['email']}\n")
            f.write(f"SUBJECT: {email_subject}\n")
            f.write(f"{'='*50}\n\n")
            f.write(body)
        
        email_count += 1
    
    print(f"📧 Saved {email_count} email files to: {email_folder}")


def create_outlook_drafts(grades: list):
    """
    Create draft emails in Outlook desktop app (Windows only).
    This lets you review each email before sending.
    """
    try:
        import win32com.client
    except ImportError:
        print("⚠️  pywin32 not installed. Run: pip install pywin32")
        print("   Falling back to saving emails as files.")
        return False
    
    try:
        outlook = win32com.client.Dispatch("Outlook.Application")
        count = 0
        
        for grade in grades:
            if not grade.get('email'):
                continue
                
            subject, body = generate_email_content(grade, grade, ASSIGNMENT_NAME)
            
            mail = outlook.CreateItem(0)  # 0 = mail item
            mail.To = grade['email']
            mail.Subject = subject
            mail.Body = body
            mail.Save()  # Save as draft
            count += 1
        
        print(f"📧 Created {count} draft emails in Outlook")
        return True
        
    except Exception as e:
        print(f"⚠️  Outlook error: {e}")
        return False


# =============================================================================
# CSV EXPORT FOR FOCUS
# =============================================================================

def export_focus_csv(grades: list, output_folder: str, assignment_name: str) -> list:
    """
    Create CSV files formatted for Focus import, SEPARATED BY ASSIGNMENT.
    
    Groups students by the assignment extracted from their filename,
    creates one CSV per assignment type.
    
    Format:
    Student ID,Score,Comment
    1950304,85,"Great job, Jackson! You correctly identified..."
    1956701,92,"Excellent work, Maria!..."
    
    Returns list of created file paths.
    """
    Path(output_folder).mkdir(parents=True, exist_ok=True)
    
    # Group grades by assignment
    assignments = {}
    for grade in grades:
        # Extract assignment name from filename (everything after FirstName_LastName_)
        filename = grade.get('filename', '')
        parts = Path(filename).stem.split('_')
        if len(parts) >= 3:
            # Assignment is everything after first two parts (first_last)
            assignment = '_'.join(parts[2:])
        else:
            assignment = "Unknown_Assignment"
        
        # Clean up assignment name for display (replace underscores with spaces)
        assignment = assignment.strip().replace('_', ' ')
        
        if assignment not in assignments:
            assignments[assignment] = []
        assignments[assignment].append(grade)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    created_files = []
    
    print(f"\n  Creating {len(assignments)} separate Focus CSVs by assignment:")
    
    for assignment, assignment_grades in assignments.items():
        # Clean assignment name for filename
        safe_name = re.sub(r'[^\w\s-]', '', assignment).replace(' ', '_')[:50]
        filepath = Path(output_folder) / f"{safe_name}_{timestamp}.csv"
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            # Column headers: Student ID, Score, Comment
            writer.writerow(['Student ID', 'Score', 'Comment'])
            
            for grade in assignment_grades:
                if grade['student_id'] != "UNKNOWN":
                    # Get student's first name only (no middle initial)
                    full_first = grade.get('first_name', '')
                    first_name = full_first.split()[0] if full_first else ''
                    
                    # Clean up feedback
                    feedback = grade.get('feedback', '')
                    feedback_clean = ' '.join(feedback.split())
                    
                    # Varied natural phrases based on grade
                    score = grade.get('score', 0)
                    
                    if first_name:
                        if score >= 90:
                            openers = [
                                f"Great job, {first_name}!",
                                f"Excellent work, {first_name}!",
                                f"{first_name}, this was fantastic!",
                                f"Well done, {first_name}!",
                                f"{first_name}, you nailed this!"
                            ]
                        elif score >= 80:
                            openers = [
                                f"Nice work, {first_name}!",
                                f"Good job, {first_name}!",
                                f"{first_name}, solid effort here!",
                                f"Well done, {first_name}!",
                                f"{first_name}, this was good work!"
                            ]
                        elif score >= 70:
                            openers = [
                                f"{first_name}, decent effort!",
                                f"Good start, {first_name}!",
                                f"{first_name}, you're on the right track!",
                                f"Keep it up, {first_name}!",
                                f"{first_name}, nice try!"
                            ]
                        else:
                            openers = [
                                f"{first_name}, let's work on this together.",
                                f"{first_name}, I know you can do better!",
                                f"Keep trying, {first_name}!",
                                f"{first_name}, don't give up!",
                                f"{first_name}, let's review this."
                            ]
                        
                        opener = random.choice(openers)
                        comment = f"{opener} {feedback_clean}"
                    else:
                        comment = feedback_clean
                    
                    writer.writerow([
                        grade['student_id'], 
                        grade['score'],
                        comment
                    ])
        
        print(f"    📊 {assignment}: {len(assignment_grades)} students → {filepath.name}")
        created_files.append(str(filepath))
    
    return created_files


def save_to_master_csv(grades: list, output_folder: str):
    """
    Save grades to a master CSV file that tracks ALL grades over time.
    This enables progress tracking across the entire school year.

    DEDUPLICATION: If a student is re-graded on the same assignment,
    the old row is REPLACED (not duplicated). The unique key is
    (Student ID, Assignment).

    Columns:
    - Date, Student ID, Student Name, Period, Assignment, Unit, Quarter
    - Overall Score, Letter Grade
    - Content Accuracy, Completeness, Writing Quality, Effort & Engagement
    """
    master_file = Path(output_folder) / "master_grades.csv"

    HEADER = [
        'Date', 'Student ID', 'Student Name', 'First Name', 'Last Name',
        'Period', 'Assignment', 'Unit', 'Quarter',
        'Overall Score', 'Letter Grade',
        'Content Accuracy', 'Completeness', 'Writing Quality', 'Effort Engagement',
        'Feedback', 'Approved',
        'API Cost', 'Input Tokens', 'Output Tokens', 'API Calls', 'AI Model'
    ]

    # Determine current quarter based on date
    today = datetime.now()
    month = today.month
    if month >= 8 and month <= 10:
        quarter = "Q1"
    elif month >= 11 or month <= 1:
        quarter = "Q2"
    elif month >= 2 and month <= 4:
        quarter = "Q3"
    else:
        quarter = "Q4"

    # Normalize assignment name for dedup matching
    # Handles: trailing "(1)", ".docx", extra whitespace, etc.
    def _normalize_assignment(name):
        n = name.strip()
        n = re.sub(r'\s*\(\d+\)\s*$', '', n)      # Remove trailing (1), (2), etc.
        n = re.sub(r'\.docx?\s*$', '', n, flags=re.IGNORECASE)  # Remove .docx/.doc
        n = re.sub(r'\.pdf\s*$', '', n, flags=re.IGNORECASE)    # Remove .pdf
        n = n.strip().lower()
        return n

    # Build set of (student_id, normalized_assignment) keys being written now
    new_keys = set()
    for grade in grades:
        sid = grade.get('student_id', '')
        assignment = grade.get('assignment', '')
        if sid and sid != "UNKNOWN" and assignment:
            new_keys.add((sid, _normalize_assignment(assignment)))

    # Read existing rows, filtering out any that will be replaced (only if new score >= old)
    existing_rows = []
    kept_old_keys = set()
    if master_file.exists():
        try:
            with open(master_file, 'r', newline='', encoding='utf-8') as f:
                reader = csv.reader(f)
                header = next(reader, None)  # Skip header
                for row in reader:
                    if len(row) >= 7:
                        row_sid = row[1]       # Student ID column
                        row_assign = row[6]    # Assignment column
                        key = (row_sid, _normalize_assignment(row_assign))
                        if key in new_keys:
                            # Compare scores — only replace if new is higher or equal
                            old_score = float(row[9]) if len(row) > 9 and row[9] else 0
                            new_grade = next((g for g in grades if g.get('student_id') == row_sid and _normalize_assignment(g.get('assignment', '')) == key[1]), None)
                            new_score = float(new_grade.get('score', 0) or 0) if new_grade else 0
                            if new_score >= old_score:
                                continue  # Replace — new is higher or equal
                            else:
                                kept_old_keys.add(key)
                                # Keep old row, skip the new grade
                    existing_rows.append(row)
        except Exception as e:
            print(f"  Note: Could not read existing master CSV: {e}")

    # Filter out grades where old score was kept
    if kept_old_keys:
        grades = [g for g in grades if (g.get('student_id', ''), _normalize_assignment(g.get('assignment', ''))) not in kept_old_keys]

    # Write full file: header + existing (deduplicated) + new grades
    with open(master_file, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow(HEADER)

        # Write back existing rows (minus duplicates)
        for row in existing_rows:
            writer.writerow(row)

        # Write new grades
        for grade in grades:
            if grade.get('student_id') == "UNKNOWN":
                continue

            breakdown = grade.get('breakdown', {})

            token_usage = grade.get('token_usage', {})

            writer.writerow([
                today.strftime('%Y-%m-%d'),
                grade.get('student_id', ''),
                grade.get('student_name', ''),
                grade.get('first_name', ''),
                grade.get('last_name', ''),
                grade.get('period', ''),
                grade.get('assignment', ''),
                grade.get('unit', ''),
                grade.get('grading_period', quarter),
                grade.get('score', 0),
                grade.get('letter_grade', ''),
                breakdown.get('content_accuracy', 0),
                breakdown.get('completeness', 0),
                breakdown.get('writing_quality', 0),
                breakdown.get('effort_engagement', 0),
                grade.get('feedback', '').replace('\r', ' ').replace('\n', ' ')[:500],
                grade.get('email_approval', 'pending'),
                token_usage.get('total_cost', ''),
                token_usage.get('total_input_tokens', ''),
                token_usage.get('total_output_tokens', ''),
                token_usage.get('api_calls', ''),
                grade.get('ai_model', '')
            ])
    
    print(f"📊 Updated master grades file: {master_file}")


def export_detailed_report(grades: list, output_folder: str, assignment_name: str) -> str:
    """
    Create detailed CSV with all grading information for your records.
    """
    safe_name = re.sub(r'[^\w\s-]', '', assignment_name).replace(' ', '_')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filepath = Path(output_folder) / f"Detailed_Report_{safe_name}_{timestamp}.csv"
    
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow([
            'Student ID', 'Student Name', 'Email', 'Assignment', 'Score', 'Letter Grade',
            'Content (40)', 'Completeness (25)', 'Writing (20)', 'Effort (15)',
            'Feedback', 'Filename'
        ])
        
        for grade in grades:
            breakdown = grade.get('breakdown', {})
            writer.writerow([
                grade.get('student_id', ''),
                grade.get('student_name', ''),
                grade.get('email', ''),
                grade.get('assignment', ''),
                grade.get('score', ''),
                grade.get('letter_grade', ''),
                breakdown.get('content_accuracy', ''),
                breakdown.get('completeness', ''),
                breakdown.get('writing_quality', ''),
                breakdown.get('effort_engagement', breakdown.get('critical_thinking', '')),
                grade.get('feedback', ''),
                grade.get('filename', '')
            ])
    
    print(f"📋 Detailed report saved: {filepath}")
    return str(filepath)


# =============================================================================
# MAIN GRADING WORKFLOW
# =============================================================================

def run_grading(
    assignment_folder: str = ASSIGNMENT_FOLDER,
    output_folder: str = OUTPUT_FOLDER,
    roster_file: str = ROSTER_FILE,
    assignment_name: str = ASSIGNMENT_NAME,
    create_outlook_emails: bool = False  # Set True if you have Outlook on Windows
):
    """
    Main function - runs the complete grading workflow.
    
    1. Loads student roster
    2. Reads each .docx file from assignment folder
    3. Grades each assignment with AI
    4. Generates emails (saves to files or creates Outlook drafts)
    5. Creates Focus CSV for grade import
    6. Creates detailed report for your records
    """
    print("=" * 60)
    print("📚 ASSIGNMENT GRADER - 6th Grade Social Studies")
    print("=" * 60)
    print(f"📁 Assignment folder: {assignment_folder}")
    print(f"💾 Output folder: {output_folder}")
    print(f"📝 Assignment: {assignment_name}")
    print()
    
    # Load roster
    roster = load_roster(roster_file)
    
    # Get all .docx files
    assignment_path = Path(assignment_folder)
    if not assignment_path.exists():
        print(f"❌ Assignment folder not found: {assignment_folder}")
        return []
    
    # Find all supported files (docx, txt, and images)
    supported_extensions = ['*.docx', '*.txt', '*.jpg', '*.jpeg', '*.png', '*.gif', '*.webp']
    all_files = []
    for ext in supported_extensions:
        all_files.extend(assignment_path.glob(ext))
    
    print(f"📄 Found {len(all_files)} files ({', '.join(supported_extensions)})")
    print()
    
    # Process each file
    all_grades = []
    
    for i, filepath in enumerate(all_files, 1):
        print(f"[{i}/{len(all_files)}] {filepath.name}")
        
        # Parse filename to get student name
        parsed = parse_filename(filepath.name)
        student_name = f"{parsed['first_name']} {parsed['last_name']}"
        lookup_key = parsed['lookup_key']
        
        # Look up student in roster
        if lookup_key in roster:
            student_info = roster[lookup_key].copy()
            print(f"  👤 {student_info['student_name']} (ID: {student_info['student_id']})")
        else:
            # Not found in roster - use parsed name
            student_info = {
                "student_id": "UNKNOWN",
                "student_name": student_name,
                "first_name": parsed['first_name'],
                "last_name": parsed['last_name'],
                "email": ""
            }
            print(f"  👤 {student_name} (⚠️ NOT IN ROSTER)")
        
        # Read file content
        file_data = read_assignment_file(filepath)
        if not file_data:
            print(f"  ❌ Could not read file")
            continue
        
        # Handle based on file type
        markers_found = []
        
        if file_data["type"] == "text":
            # Text-based file - check for markers
            content = file_data["content"]
            
            if len(content.strip()) < 20:
                print(f"  ⚠️  File appears empty ({len(content)} chars)")
                continue
            
            # Extract only the student work portion
            student_work, markers_found = extract_student_work(content)
            
            if markers_found:
                print(f"  📝 Found marker(s): {', '.join(markers_found[:2])}{'...' if len(markers_found) > 2 else ''}")
            else:
                print(f"  ⚠️  NO MARKERS FOUND - Check if student uploaded wrong document!")
                print(f"      → Review manually: {filepath.name}")
            
            if len(student_work.strip()) < 10:
                print(f"  ⚠️  No student work found after marker")
                continue
            
            # Prepare data for grading
            grade_data = {"type": "text", "content": student_work}
        
        elif file_data["type"] == "image":
            # Image file - send entire image to AI for grading
            print(f"  🖼️  Image file - sending to AI for visual grading")
            grade_data = file_data
            markers_found = ["image"]  # Mark as having content
        
        else:
            print(f"  ❌ Unknown file type")
            continue
        
        # Grade with AI
        grade_result = grade_assignment(student_info['student_name'], grade_data)
        
        # Combine all info
        # Extract assignment name from filename
        parts = Path(filepath.name).stem.split('_')
        if len(parts) >= 3:
            assignment_from_file = ' '.join(parts[2:])
        else:
            assignment_from_file = ASSIGNMENT_NAME
        
        grade_record = {
            **student_info,
            **grade_result,
            "filename": filepath.name,
            "assignment": assignment_from_file,
            "has_markers": len(markers_found) > 0
        }
        all_grades.append(grade_record)
        
        print(f"  ✅ Score: {grade_result['score']} ({grade_result['letter_grade']})")
        print()
    
    # Export results
    print("=" * 60)
    print("📊 EXPORTING RESULTS")
    print("=" * 60)
    
    # Focus CSVs (separated by assignment)
    focus_files = export_focus_csv(all_grades, output_folder, assignment_name)
    
    # Detailed report (one file with all grades)
    export_detailed_report(all_grades, output_folder, assignment_name)
    
    # Emails
    if create_outlook_emails:
        if not create_outlook_drafts(all_grades):
            save_emails_to_folder(all_grades, output_folder)
    else:
        save_emails_to_folder(all_grades, output_folder)
    
    # Summary
    print()
    print("=" * 60)
    print("📈 GRADING SUMMARY")
    print("=" * 60)
    
    if all_grades:
        scores = [g['score'] for g in all_grades if g['score'] > 0]
        if scores:
            print(f"Total graded: {len(all_grades)}")
            print(f"Average score: {sum(scores)/len(scores):.1f}")
            print(f"Highest: {max(scores)}")
            print(f"Lowest: {min(scores)}")
            
            # Grade distribution
            grade_dist = {}
            for g in all_grades:
                letter = g['letter_grade']
                grade_dist[letter] = grade_dist.get(letter, 0) + 1
            print(f"Distribution: {dict(sorted(grade_dist.items()))}")
            
            # Per-assignment breakdown
            print(f"\n📚 By Assignment:")
            assignments = {}
            for g in all_grades:
                a = g.get('assignment', 'Unknown')
                if a not in assignments:
                    assignments[a] = []
                assignments[a].append(g['score'])
            
            for assignment, scores_list in sorted(assignments.items()):
                valid_scores = [s for s in scores_list if s > 0]
                if valid_scores:
                    avg = sum(valid_scores) / len(valid_scores)
                    print(f"   • {assignment[:40]}: {len(scores_list)} students, avg {avg:.1f}")
        
        # List students not in roster
        unknown = [g for g in all_grades if g['student_id'] == 'UNKNOWN']
        if unknown:
            print(f"\n⚠️  {len(unknown)} students NOT FOUND in roster:")
            for g in unknown:
                print(f"   - {g['student_name']} ({g['filename']})")
        
        # List documents with no markers (possible wrong uploads)
        no_markers = [g for g in all_grades if not g.get('has_markers', True)]
        if no_markers:
            print(f"\n🚨 {len(no_markers)} documents had NO MARKERS - review for possible wrong uploads:")
            for g in no_markers:
                print(f"   - {g['student_name']}: {g['filename']}")
    
    print()
    print("✅ GRADING COMPLETE!")
    return all_grades


# =============================================================================
# RUN THE SCRIPT
# =============================================================================

if __name__ == "__main__":
    # Update the paths at the top of this file, then run!
    results = run_grading(
        create_outlook_emails=False  # Set True if you have Outlook desktop on Windows
    )
