#!/usr/bin/env python3
"""
Graider - AI-Powered Assignment Grading
=======================================
Just run: python3 grader_app.py
Then open: http://localhost:3000
"""

from flask import Flask, jsonify, request, Response
from flask_cors import CORS
import os
import sys
import threading
from pathlib import Path
from dotenv import load_dotenv

# Load .env from this file's directory, override any existing env vars
_app_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(_app_dir, '.env'), override=True)

app = Flask(__name__)
CORS(app)

# Global state
grading_state = {
    "is_running": False,
    "stop_requested": False,
    "progress": 0,
    "total": 0,
    "current_file": "",
    "log": [],
    "results": [],
    "complete": False,
    "error": None
}

def reset_state():
    global grading_state
    grading_state = {
        "is_running": False,
        "stop_requested": False,
        "progress": 0,
        "total": 0,
        "current_file": "",
        "log": [],
        "results": [],
        "complete": False,
        "error": None
    }

# ══════════════════════════════════════════════════════════════
# HTML TEMPLATE - Beautiful React UI embedded directly
# ══════════════════════════════════════════════════════════════

HTML_TEMPLATE = '''<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Graider - AI-Powered Grading</title>
    <script src="https://cdn.jsdelivr.net/npm/react@18/umd/react.development.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/react-dom@18/umd/react-dom.development.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/@babel/standalone@7/babel.min.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/lucide@0.263.1/dist/umd/lucide.min.js"></script>
    <script src="https://cdn.jsdelivr.net/npm/recharts@2.1.9/umd/Recharts.min.js"></script>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700;800&display=swap" rel="stylesheet">
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body {
            font-family: 'Inter', -apple-system, sans-serif;
            background: linear-gradient(135deg, #1a1a2e 0%, #16213e 50%, #0f3460 100%);
            min-height: 100vh;
            color: #fff;
        }
        @keyframes spin { from { transform: rotate(0deg); } to { transform: rotate(360deg); } }
        @keyframes pulse { 0%, 100% { opacity: 1; } 50% { opacity: 0.5; } }
        @keyframes fadeIn { from { opacity: 0; transform: translateY(10px); } to { opacity: 1; transform: translateY(0); } }
        .fade-in { animation: fadeIn 0.4s ease-out; }
        ::-webkit-scrollbar { width: 8px; }
        ::-webkit-scrollbar-track { background: rgba(255,255,255,0.05); border-radius: 4px; }
        ::-webkit-scrollbar-thumb { background: rgba(255,255,255,0.2); border-radius: 4px; }
    </style>
</head>
<body>
    <div id="root"></div>
    <script type="text/babel">
        const { useState, useEffect, useRef } = React;

        const Icon = ({ name, size = 24 }) => {
            const ref = useRef(null);
            useEffect(() => {
                if (ref.current && lucide[name]) {
                    ref.current.innerHTML = '';
                    const svg = lucide.createElement(lucide[name]);
                    svg.setAttribute('width', size);
                    svg.setAttribute('height', size);
                    ref.current.appendChild(svg);
                }
            }, [name, size]);
            return <span ref={ref} style={{ display: 'inline-flex', alignItems: 'center' }} />;
        };

        const StandardCard = ({ standard, isSelected, onToggle }) => {
            return (
                <div 
                    onClick={onToggle}
                    style={{ 
                        background: isSelected ? 'rgba(99,102,241,0.2)' : 'rgba(255,255,255,0.03)', 
                        border: isSelected ? '1px solid #6366f1' : '1px solid rgba(255,255,255,0.1)',
                        borderRadius: '12px',
                        padding: '15px',
                        cursor: 'pointer',
                        transition: 'all 0.2s',
                        position: 'relative',
                        marginBottom: '10px'
                    }}
                >
                    <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'flex-start', marginBottom: '8px' }}>
                        <span style={{ fontWeight: 700, color: isSelected ? '#a5b4fc' : '#fff', fontSize: '0.9rem' }}>{standard.code}</span>
                        {isSelected && <Icon name="CheckCircle" size={18} style={{ color: '#6366f1' }} />}
                    </div>
                    <p style={{ fontSize: '0.9rem', color: 'rgba(255,255,255,0.8)', lineHeight: '1.5', margin: '0 0 10px 0' }}>{standard.benchmark}</p>
                    <div style={{ display: 'flex', flexWrap: 'wrap', gap: '6px' }}>
                        {standard.topics.map(topic => (
                            <span key={topic} style={{ fontSize: '0.75rem', padding: '3px 8px', borderRadius: '4px', background: 'rgba(255,255,255,0.1)', color: 'rgba(255,255,255,0.6)' }}>
                                {topic}
                            </span>
                        ))}
                    </div>
                </div>
            );
        };

        const App = () => {
            const [config, setConfig] = useState({
                assignments_folder: "/Users/alexc/Library/CloudStorage/OneDrive-VolusiaCountySchools/Assignments",
                output_folder: "/Users/alexc/Downloads/Graider/Results",
                roster_file: "/Users/alexc/Downloads/Graider/all_students_updated.xlsx",
                grading_period: "Q3"
            });
            const [status, setStatus] = useState({
                is_running: false, progress: 0, total: 0, current_file: "",
                log: [], results: [], complete: false, error: null
            });
            const [activeTab, setActiveTab] = useState('grade');
            const [analytics, setAnalytics] = useState(null);
            const [selectedStudent, setSelectedStudent] = useState(null);
            const [analyticsPeriod, setAnalyticsPeriod] = useState('all');
            const [autoGrade, setAutoGrade] = useState(false);
            const [watchStatus, setWatchStatus] = useState({ watching: false, lastCheck: null, newFiles: 0 });
            const logRef = useRef(null);

            // Planner State
            const [plannerConfig, setPlannerConfig] = useState({
                state: 'FL',
                grade: '7',
                subject: 'Civics'
            });
            const [standards, setStandards] = useState([]);
            const [selectedStandards, setSelectedStandards] = useState([]);
            const [lessonPlan, setLessonPlan] = useState(null);
            const [plannerLoading, setPlannerLoading] = useState(false);
            const [unitConfig, setUnitConfig] = useState({
                title: '',
                duration: 1,
                periodLength: 50,
                startDate: '',
                type: 'Lesson Plan',
                format: 'Word'
            });
            
            // Recharts components - safely get from window
            const RC = window.Recharts || {};
            const { LineChart, Line, BarChart, Bar, PieChart, Pie, Cell, XAxis, YAxis, CartesianGrid, Tooltip, Legend, ResponsiveContainer } = RC;
            const chartsLoaded = !!ResponsiveContainer;

            useEffect(() => {
                const interval = setInterval(async () => {
                    try {
                        const res = await fetch('/api/status');
                        const data = await res.json();
                        setStatus(data);
                    } catch (e) {}
                }, 500);
                return () => clearInterval(interval);
            }, []);

            // Load saved settings on startup
            const [settingsLoaded, setSettingsLoaded] = React.useState(false);
            useEffect(() => {
                // Load rubric
                fetch('/api/load-rubric')
                    .then(res => res.json())
                    .then(data => {
                        if (data.rubric) setRubric(data.rubric);
                    })
                    .catch(e => console.error('Failed to load rubric:', e));

                // Load global settings (AI notes)
                fetch('/api/load-global-settings')
                    .then(res => res.json())
                    .then(data => {
                        if (data.settings?.globalAINotes) {
                            setGlobalAINotes(data.settings.globalAINotes);
                        }
                    })
                    .catch(e => console.error('Failed to load settings:', e));

                // Load saved assignments list
                fetch('/api/list-assignments')
                    .then(res => res.json())
                    .then(data => {
                        if (data.assignments) setSavedAssignments(data.assignments);
                    })
                    .catch(e => console.error('Failed to load assignments:', e));

                // Mark settings as loaded after a short delay
                setTimeout(() => setSettingsLoaded(true), 500);
            }, []);

            // Auto-save rubric when it changes (after initial load)
            useEffect(() => {
                if (!settingsLoaded) return;
                const timer = setTimeout(() => {
                    fetch('/api/save-rubric', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify(rubric)
                    }).catch(e => console.error('Failed to save rubric:', e));
                }, 500);
                return () => clearTimeout(timer);
            }, [rubric, settingsLoaded]);

            // Auto-grade watcher
            useEffect(() => {
                if (!autoGrade) return;
                
                const watchInterval = setInterval(async () => {
                    if (status.is_running) return; // Don't check while grading
                    
                    try {
                        const res = await fetch('/api/check-new-files', {
                            method: 'POST',
                            headers: { 'Content-Type': 'application/json' },
                            body: JSON.stringify({ folder: config.assignments_folder, output_folder: config.output_folder })
                        });
                        const data = await res.json();
                        setWatchStatus({ watching: true, lastCheck: new Date().toLocaleTimeString(), newFiles: data.new_files || 0 });
                        
                        if (data.new_files > 0 && !status.is_running) {
                            // Auto-start grading
                            startGrading();
                        }
                    } catch (e) {
                        console.error('Watch error:', e);
                    }
                }, 10000); // Check every 10 seconds
                
                setWatchStatus({ watching: true, lastCheck: 'Starting...', newFiles: 0 });
                
                return () => {
                    clearInterval(watchInterval);
                    setWatchStatus({ watching: false, lastCheck: null, newFiles: 0 });
                };
            }, [autoGrade, config.assignments_folder, config.output_folder]);
            
            // Fetch analytics when tab opens or period changes
            useEffect(() => {
                if (activeTab === 'analytics') {
                    const url = analyticsPeriod === 'all' ? '/api/analytics' : '/api/analytics?period=' + encodeURIComponent(analyticsPeriod);
                    fetch(url)
                        .then(res => res.json())
                        .then(data => setAnalytics(data))
                        .catch(e => console.error(e));
                }
            }, [activeTab, analyticsPeriod]);

            useEffect(() => {
                if (logRef.current) logRef.current.scrollTop = logRef.current.scrollHeight;
            }, [status.log]);

            const startGrading = async () => {
                const res = await fetch('/api/grade', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({
                        ...config,
                        assignmentConfig: selectedGradingConfig,
                        globalAINotes: globalAINotes
                    })
                });
                const data = await res.json();
                if (data.error && !data.error.includes('already in progress')) alert(data.error);
            };

            const loadGradingConfig = async (name) => {
                if (!name) {
                    setSelectedGradingConfig(null);
                    return;
                }

                // Special case: load ALL configs and combine markers + notes
                if (name === '__ALL__') {
                    try {
                        let allMarkers = [];
                        let allNotes = [];
                        for (const assignmentName of savedAssignments) {
                            const res = await fetch('/api/load-assignment?name=' + encodeURIComponent(assignmentName));
                            const data = await res.json();
                            if (data.assignment) {
                                if (data.assignment.customMarkers) {
                                    allMarkers = [...allMarkers, ...data.assignment.customMarkers];
                                }
                                if (data.assignment.gradingNotes) {
                                    allNotes.push('[' + assignmentName + ']: ' + data.assignment.gradingNotes);
                                }
                            }
                        }
                        // Remove duplicate markers
                        const uniqueMarkers = [...new Set(allMarkers)];
                        // Combine global notes + all individual notes
                        const separator = String.fromCharCode(10) + String.fromCharCode(10) + '--- Per-Assignment Notes ---' + String.fromCharCode(10) + String.fromCharCode(10);
                        const combinedNotes = globalAINotes
                            ? globalAINotes + separator + allNotes.join(String.fromCharCode(10) + String.fromCharCode(10))
                            : allNotes.join(String.fromCharCode(10) + String.fromCharCode(10));
                        setSelectedGradingConfig({
                            title: 'All Assignments Combined',
                            customMarkers: uniqueMarkers,
                            gradingNotes: combinedNotes
                        });
                    } catch (e) {
                        console.error('Error loading all configs:', e);
                    }
                    return;
                }

                try {
                    const res = await fetch('/api/load-assignment?name=' + encodeURIComponent(name));
                    const data = await res.json();
                    if (data.assignment) {
                        setSelectedGradingConfig(data.assignment);
                    }
                } catch (e) {
                    console.error('Error loading config:', e);
                }
            };

            const openResults = () => fetch('/api/open-folder', {
                method: 'POST',
                headers: { 'Content-Type': 'application/json' },
                body: JSON.stringify({ folder: config.output_folder })
            });

            const stopGrading = async () => {
                try {
                    const res = await fetch('/api/stop-grading', { method: 'POST' });
                    const data = await res.json();
                    if (data.stopped) {
                        // Turn off auto-grade if it was on
                        setAutoGrade(false);
                    }
                } catch (e) {
                    console.error('Stop error:', e);
                }
            };

            const [emailStatus, setEmailStatus] = React.useState({ sending: false, sent: 0, failed: 0, message: '' });
            const [emailPreview, setEmailPreview] = React.useState({ show: false, emails: [] });
            const [reviewModal, setReviewModal] = React.useState({ show: false, index: -1 });
            const [editedResults, setEditedResults] = React.useState([]);
            const [rubric, setRubric] = React.useState({
                categories: [
                    { name: 'Content Accuracy', points: 40, description: 'Are answers factually correct?' },
                    { name: 'Completeness', points: 25, description: 'Did student attempt all questions?' },
                    { name: 'Writing Quality', points: 20, description: 'Is writing clear and readable?' },
                    { name: 'Effort & Engagement', points: 15, description: 'Did student show genuine effort?' }
                ],
                gradeLevel: '6th',
                subject: 'Social Studies',
                generous: true
            });
            
            // Assignment Builder state
            const [assignment, setAssignment] = React.useState({
                title: '',
                subject: 'Social Studies',
                totalPoints: 100,
                instructions: '',
                questions: [],
                customMarkers: [],
                gradingNotes: ''
            });
            const [savedAssignments, setSavedAssignments] = React.useState([]);
            const [loadedAssignmentName, setLoadedAssignmentName] = React.useState('');
            const [globalAINotes, setGlobalAINotes] = React.useState('');
            const [selectedGradingConfig, setSelectedGradingConfig] = React.useState(null);
            const [importedDoc, setImportedDoc] = React.useState({ text: '', html: '', filename: '', loading: false });
            const [docEditorModal, setDocEditorModal] = React.useState({ show: false, editedText: '', editedHtml: '', viewMode: 'formatted' });
            const [selectedMarkers, setSelectedMarkers] = React.useState([]);
            const fileInputRef = React.useRef(null);
            const docHtmlRef = React.useRef(null);
            
            // Import document handler
            const handleDocImport = async (e) => {
                const file = e.target.files[0];
                if (!file) return;
                
                setImportedDoc({ text: '', html: '', filename: file.name, loading: true });
                
                const formData = new FormData();
                formData.append('file', file);
                
                try {
                    const res = await fetch('/api/parse-document', {
                        method: 'POST',
                        body: formData
                    });
                    const data = await res.json();
                    if (data.error) {
                        alert('Error parsing document: ' + data.error);
                        setImportedDoc({ text: '', html: '', filename: '', loading: false });
                    } else {
                        setImportedDoc({
                            text: data.text || '',
                            html: data.html || '',
                            filename: file.name,
                            loading: false
                        });
                        // Clear loadedAssignmentName since we're importing a new document
                        // This ensures save will warn if the title matches an existing assignment
                        setLoadedAssignmentName('');
                        setDocEditorModal({ show: true, editedText: data.text || '', editedHtml: data.html || '', viewMode: 'formatted' });
                        // Auto-fill title from filename
                        if (!assignment.title) {
                            const title = file.name.replace(/\.(docx|pdf|doc)$/i, '').replace(/_/g, ' ');
                            setAssignment({ ...assignment, title });
                        }
                    }
                } catch (err) {
                    alert('Error: ' + err.message);
                    setImportedDoc({ text: '', html: '', filename: '', loading: false });
                }
            };
            
            // Open editor modal with existing doc
            const openDocEditor = () => {
                if (importedDoc.text || importedDoc.html) {
                    setDocEditorModal({ show: true, editedText: importedDoc.text, editedHtml: importedDoc.html, viewMode: 'formatted' });
                }
            };
            
            // Add selected text as marker
            const addSelectedAsMarker = () => {
                let text = '';
                
                // Try to get selection from iframe first
                try {
                    if (docHtmlRef.current && docHtmlRef.current.contentDocument) {
                        const iframeSelection = docHtmlRef.current.contentDocument.getSelection();
                        if (iframeSelection) {
                            text = iframeSelection.toString().trim();
                        }
                    }
                } catch (e) { console.log('Iframe selection error:', e); }
                
                // Fall back to main window selection
                if (!text) {
                    const selection = window.getSelection();
                    text = selection ? selection.toString().trim() : '';
                }
                
                if (text && text.length > 2 && text.length < 500) {
                    if (!(assignment.customMarkers || []).includes(text)) {
                        setSelectedMarkers([...selectedMarkers, text]);
                        setAssignment({
                            ...assignment,
                            customMarkers: [...(assignment.customMarkers || []), text]
                        });
                    }
                } else if (text.length <= 2) {
                    alert('Please select more text (at least 3 characters)');
                } else if (text.length >= 500) {
                    alert('Selection too long. Please select less text (under 500 characters)');
                }
            };
            
            // Remove marker
            const removeMarker = (marker) => {
                setSelectedMarkers(selectedMarkers.filter(m => m !== marker));
                setAssignment({
                    ...assignment,
                    customMarkers: (assignment.customMarkers || []).filter(m => m !== marker)
                });
            };
            
            // Marker libraries by subject
            const markerLibrary = {
                'Social Studies': ['Explain:', 'Describe the significance of:', 'Compare and contrast:', 'What were the causes of:', 'What were the effects of:', 'Analyze:', 'In your own words:', 'Why do you think:'],
                'English/ELA': ['Write your response:', 'Your thesis statement:', 'Analyze the text:', 'Provide evidence:', 'Explain the theme:', 'Character analysis:', 'Authors purpose:'],
                'Math': ['Show your work:', 'Solve:', 'Calculate:', 'Prove:', 'Find the value of:', 'Graph:', 'Simplify:', 'Word Problem:'],
                'Science': ['Hypothesis:', 'Data/Observations:', 'Conclusion:', 'Procedure:', 'Variables:', 'Analysis:', 'Explain the results:'],
                'History': ['Explain:', 'Describe:', 'What was the impact of:', 'Primary source analysis:', 'Timeline:', 'Cause and effect:', 'Historical significance:'],
                'Foreign Language': ['Translate:', 'Write in target language:', 'Conjugate:', 'Vocabulary:', 'Reading comprehension:', 'Conversation:'],
                'Other': ['Answer:', 'Explain:', 'Describe:', 'Your response:', 'Short answer:', 'Essay:']
            };
            
            const addQuestion = () => {
                setAssignment({
                    ...assignment,
                    questions: [...assignment.questions, {
                        id: Date.now(),
                        type: 'short_answer',
                        prompt: '',
                        points: 10,
                        marker: markerLibrary[assignment.subject]?.[0] || 'Answer:'
                    }]
                });
            };
            
            const updateQuestion = (index, field, value) => {
                const updated = [...assignment.questions];
                updated[index] = { ...updated[index], [field]: value };
                setAssignment({ ...assignment, questions: updated });
            };
            
            const removeQuestion = (index) => {
                setAssignment({
                    ...assignment,
                    questions: assignment.questions.filter((_, i) => i !== index)
                });
            };
            
            const exportAssignment = async (format) => {
                try {
                    const res = await fetch('/api/export-assignment', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({ assignment, format })
                    });
                    const data = await res.json();
                    if (data.error) {
                        alert('Error: ' + data.error);
                    } else {
                        alert('Assignment exported! Check your output folder.');
                    }
                } catch (e) {
                    alert('Error exporting: ' + e.message);
                }
            };
            
            const saveAssignmentConfig = async () => {
                try {
                    let titleToSave = assignment.title;

                    // Check if this title already exists and we're not editing that same assignment
                    // Only check for exact match (case-insensitive)
                    const titleExists = savedAssignments.some(name =>
                        name.toLowerCase() === titleToSave.toLowerCase()
                    );

                    if (titleExists && loadedAssignmentName !== titleToSave) {
                        const choice = confirm(
                            `"${titleToSave}" already exists.\n\nClick OK to OVERWRITE it, or Cancel to enter a new name.`
                        );
                        if (!choice) {
                            const newTitle = prompt('Enter a new assignment name:', titleToSave + ' (new)');
                            if (!newTitle || !newTitle.trim()) {
                                return; // User cancelled
                            }
                            titleToSave = newTitle.trim();
                            setAssignment({ ...assignment, title: titleToSave });
                        }
                    }

                    // Include the imported document content so it can be restored when loading
                    const dataToSave = {
                        ...assignment,
                        title: titleToSave,
                        importedDoc: importedDoc
                    };
                    const res = await fetch('/api/save-assignment-config', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify(dataToSave)
                    });
                    const data = await res.json();
                    if (data.error) {
                        alert('Error: ' + data.error);
                    } else {
                        alert('Assignment saved! Markers will be used when grading.');
                        setLoadedAssignmentName(titleToSave);
                        loadSavedAssignments();
                    }
                } catch (e) {
                    alert('Error saving: ' + e.message);
                }
            };
            
            const loadSavedAssignments = async () => {
                try {
                    const res = await fetch('/api/list-assignments');
                    const data = await res.json();
                    setSavedAssignments(data.assignments || []);
                } catch (e) {
                    console.error('Error loading assignments:', e);
                }
            };
            
            const loadStandards = async () => {
                setPlannerLoading(true);
                try {
                    const res = await fetch('/api/get-standards', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify(plannerConfig)
                    });
                    const data = await res.json();
                    setStandards(data.standards || []);
                } catch (e) {
                    alert('Error loading standards: ' + e.message);
                } finally {
                    setPlannerLoading(false);
                }
            };

            useEffect(() => {
                loadStandards();
            }, [plannerConfig.state, plannerConfig.grade, plannerConfig.subject, activeTab]);

            const toggleStandard = (code) => {
                const newSelected = selectedStandards.includes(code)
                    ? selectedStandards.filter(c => c !== code)
                    : [...selectedStandards, code];
                setSelectedStandards(newSelected);
            };

            const generateLessonPlan = async () => {
                if (selectedStandards.length === 0) {
                    alert('Please select at least one standard.');
                    return;
                }
                if (!unitConfig.title) {
                    alert('Please enter a title.');
                    return;
                }
                
                setPlannerLoading(true);
                try {
                    const res = await fetch('/api/generate-lesson-plan', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({
                            standards: selectedStandards,
                            config: { ...plannerConfig, ...unitConfig }
                        })
                    });
                    const data = await res.json();
                    if (data.error) {
                        alert('Error: ' + data.error);
                    } else {
                        setLessonPlan(data.plan || data);
                    }
                } catch (e) {
                    alert('Error generating plan: ' + e.message);
                } finally {
                    setPlannerLoading(false);
                }
            };

            const exportLessonPlan = async () => {
                if (!lessonPlan) return;
                try {
                    const res = await fetch('/api/export-lesson-plan', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify(lessonPlan)
                    });
                    const data = await res.json();
                    if (data.error) {
                        alert('Error exporting: ' + data.error);
                    } else {
                        alert('Lesson plan exported to: ' + data.path);
                    }
                } catch (e) {
                    alert('Error exporting: ' + e.message);
                }
            };

            const loadAssignment = async (name) => {
                try {
                    const res = await fetch('/api/load-assignment?name=' + encodeURIComponent(name));
                    const data = await res.json();
                    if (data.error) {
                        alert('Error: ' + data.error);
                        return;
                    }
                    if (data.assignment) {
                        // Fully replace assignment with loaded data, ensuring all fields are set
                        setAssignment({
                            title: data.assignment.title || '',
                            subject: data.assignment.subject || 'Social Studies',
                            totalPoints: data.assignment.totalPoints || 100,
                            instructions: data.assignment.instructions || '',
                            questions: data.assignment.questions || [],
                            customMarkers: data.assignment.customMarkers || [],
                            gradingNotes: data.assignment.gradingNotes || ''
                        });
                        setLoadedAssignmentName(name);
                        // Also load the imported document if it was saved
                        if (data.assignment.importedDoc) {
                            setImportedDoc(data.assignment.importedDoc);
                        } else {
                            setImportedDoc({ text: '', html: '', filename: '', loading: false });
                        }
                    }
                } catch (e) {
                    alert('Error loading: ' + e.message);
                }
            };
            
            // Load saved assignments on mount
            React.useEffect(() => {
                loadSavedAssignments();
            }, []);
            
            // Sync editedResults with status.results
            React.useEffect(() => {
                if (status.results.length > 0 && editedResults.length !== status.results.length) {
                    setEditedResults(status.results.map(r => ({ ...r, edited: false })));
                }
            }, [status.results]);
            
            const openReview = (index) => {
                setReviewModal({ show: true, index });
            };
            
            const updateGrade = (index, field, value) => {
                const updated = [...editedResults];
                updated[index] = { ...updated[index], [field]: value, edited: true };
                // Recalculate letter grade if score changed
                if (field === 'score') {
                    const score = parseInt(value) || 0;
                    updated[index].letter_grade = score >= 90 ? 'A' : score >= 80 ? 'B' : score >= 70 ? 'C' : score >= 60 ? 'D' : 'F';
                }
                setEditedResults(updated);
            };
            
            const saveRubric = async () => {
                try {
                    await fetch('/api/save-rubric', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify(rubric)
                    });
                    alert('Rubric saved!');
                } catch (e) {
                    alert('Error saving rubric: ' + e.message);
                }
            };
            
            const addCategory = () => {
                setRubric({
                    ...rubric,
                    categories: [...rubric.categories, { name: 'New Category', points: 10, description: '' }]
                });
            };
            
            const removeCategory = (index) => {
                setRubric({
                    ...rubric,
                    categories: rubric.categories.filter((_, i) => i !== index)
                });
            };
            
            const updateCategory = (index, field, value) => {
                const cats = [...rubric.categories];
                cats[index] = { ...cats[index], [field]: field === 'points' ? parseInt(value) || 0 : value };
                setRubric({ ...rubric, categories: cats });
            };
            
            const previewEmails = () => {
                const results = editedResults.length > 0 ? editedResults : status.results;
                if (results.length === 0) {
                    alert('No results to email. Run grading first.');
                    return;
                }
                
                // Group by student email
                const students = {};
                results.forEach(r => {
                    const email = r.email || '';
                    if (email && email.includes('@') && r.student_id !== 'UNKNOWN') {
                        if (!students[email]) {
                            students[email] = { name: r.student_name, email: email, grades: [] };
                        }
                        students[email].grades.push(r);
                    }
                });
                
                const NL = String.fromCharCode(10);
                const emailList = Object.values(students).map(s => {
                    const firstName = s.name.split(' ')[0];
                    const subject = s.grades.length === 1 
                        ? 'Grade for ' + s.grades[0].assignment + ': ' + s.grades[0].letter_grade
                        : 'Grades for ' + s.grades.length + ' Assignments';
                    
                    let body = 'Hi ' + firstName + ',' + NL + NL;
                    if (s.grades.length === 1) {
                        const g = s.grades[0];
                        body += 'Here is your grade and feedback for ' + g.assignment + ':' + NL + NL;
                        body += 'GRADE: ' + g.score + '/100 (' + g.letter_grade + ')' + NL + NL;
                        body += 'FEEDBACK:' + NL + g.feedback;
                    } else {
                        body += 'Here are your grades and feedback:' + NL + NL;
                        s.grades.forEach(g => {
                            body += '━━━━━━━━━━━━━━━━━━━━━━━━━━' + NL;
                            body += g.assignment + NL + 'GRADE: ' + g.score + '/100 (' + g.letter_grade + ')' + NL + NL;
                            body += 'FEEDBACK:' + NL + g.feedback + NL + NL;
                        });
                    }
                    
                    return { to: s.email, name: s.name, subject: subject, body: body, assignments: s.grades.length };
                });
                
                setEmailPreview({ show: true, emails: emailList });
            };
            
            const sendEmails = async () => {
                setEmailPreview({ ...emailPreview, show: false });
                
                const results = editedResults.length > 0 ? editedResults : status.results;
                if (results.length === 0) {
                    alert('No results to email. Run grading first.');
                    return;
                }
                
                setEmailStatus({ sending: true, sent: 0, failed: 0, message: 'Sending emails...' });
                
                try {
                    const res = await fetch('/api/send-emails', {
                        method: 'POST',
                        headers: { 'Content-Type': 'application/json' },
                        body: JSON.stringify({ results })
                    });
                    const data = await res.json();
                    setEmailStatus({ 
                        sending: false, 
                        sent: data.sent || 0, 
                        failed: data.failed || 0,
                        message: data.error ? `Error: ${data.error}` : `✅ Sent ${data.sent} emails${data.failed > 0 ? `, ${data.failed} failed` : ''}`
                    });
                } catch (e) {
                    setEmailStatus({ sending: false, sent: 0, failed: 0, message: `Error: ${e.message}` });
                }
            };

            const pct = status.total > 0 ? (status.progress / status.total) * 100 : 0;

            return (
                <div style={{ minHeight: '100vh', padding: '20px' }}>
                    {/* Email Preview Modal */}
                    {emailPreview.show && (
                        <div style={{ position: 'fixed', top: 0, left: 0, right: 0, bottom: 0, background: 'rgba(0,0,0,0.8)', zIndex: 1000, display: 'flex', alignItems: 'center', justifyContent: 'center', padding: '20px' }}>
                            <div style={{ background: '#1a1a2e', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', width: '100%', maxWidth: '800px', maxHeight: '90vh', overflow: 'hidden', display: 'flex', flexDirection: 'column' }}>
                                <div style={{ padding: '20px 25px', borderBottom: '1px solid rgba(255,255,255,0.1)', display: 'flex', justifyContent: 'space-between', alignItems: 'center' }}>
                                    <h2 style={{ fontSize: '1.3rem', fontWeight: 700, display: 'flex', alignItems: 'center', gap: '10px' }}><Icon name="Mail" size={24} />Email Preview ({emailPreview.emails.length} students)</h2>
                                    <button onClick={() => setEmailPreview({ show: false, emails: [] })} style={{ background: 'none', border: 'none', color: 'rgba(255,255,255,0.6)', cursor: 'pointer', padding: '5px' }}><Icon name="X" size={24} /></button>
                                </div>
                                <div style={{ flex: 1, overflowY: 'auto', padding: '20px' }}>
                                    {emailPreview.emails.map((email, i) => (
                                        <div key={i} style={{ background: 'rgba(255,255,255,0.03)', borderRadius: '12px', border: '1px solid rgba(255,255,255,0.1)', marginBottom: '15px', overflow: 'hidden' }}>
                                            <div style={{ padding: '15px 20px', borderBottom: '1px solid rgba(255,255,255,0.05)', display: 'flex', justifyContent: 'space-between', alignItems: 'center' }}>
                                                <div>
                                                    <div style={{ fontWeight: 600, marginBottom: '4px' }}>{email.name}</div>
                                                    <div style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.5)' }}>{email.to}</div>
                                                </div>
                                                <span style={{ background: 'rgba(99,102,241,0.2)', color: '#a5b4fc', padding: '4px 12px', borderRadius: '20px', fontSize: '0.8rem' }}>{email.assignments} assignment{email.assignments > 1 ? 's' : ''}</span>
                                            </div>
                                            <div style={{ padding: '15px 20px' }}>
                                                <div style={{ fontSize: '0.9rem', color: '#a5b4fc', marginBottom: '10px' }}><strong>Subject:</strong> {email.subject}</div>
                                                <div style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.7)', whiteSpace: 'pre-wrap', maxHeight: '150px', overflowY: 'auto', background: 'rgba(0,0,0,0.2)', padding: '12px', borderRadius: '8px', fontFamily: 'Monaco, monospace', lineHeight: '1.5' }}>{email.body}</div>
                                            </div>
                                        </div>
                                    ))}
                                </div>
                                <div style={{ padding: '20px 25px', borderTop: '1px solid rgba(255,255,255,0.1)', display: 'flex', gap: '15px', justifyContent: 'flex-end' }}>
                                    <button onClick={() => setEmailPreview({ show: false, emails: [] })} style={{ padding: '12px 24px', borderRadius: '10px', border: '1px solid rgba(255,255,255,0.2)', background: 'transparent', color: '#fff', cursor: 'pointer', fontWeight: 600 }}>Cancel</button>
                                    <button onClick={sendEmails} style={{ padding: '12px 24px', borderRadius: '10px', border: 'none', background: 'linear-gradient(135deg, #10b981, #059669)', color: '#fff', cursor: 'pointer', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}><Icon name="Send" size={18} />Send All Emails</button>
                                </div>
                            </div>
                        </div>
                    )}
                    
                    {/* Document Editor Modal - Full Screen */}
                    {docEditorModal.show && (
                        <div style={{ position: 'fixed', top: 0, left: 0, right: 0, bottom: 0, background: 'rgba(0,0,0,0.95)', zIndex: 1000, display: 'flex', flexDirection: 'column' }}>
                            {/* Header */}
                            <div style={{ padding: '15px 25px', borderBottom: '1px solid rgba(255,255,255,0.1)', display: 'flex', justifyContent: 'space-between', alignItems: 'center', background: '#1a1a2e' }}>
                                <div style={{ display: 'flex', alignItems: 'center', gap: '15px' }}>
                                    <h2 style={{ fontSize: '1.2rem', fontWeight: 700, display: 'flex', alignItems: 'center', gap: '10px', margin: 0 }}>
                                        <Icon name="FileEdit" size={22} />Document Editor
                                    </h2>
                                    <span style={{ fontSize: '0.9rem', color: 'rgba(255,255,255,0.5)' }}>{importedDoc.filename}</span>
                                    <span style={{ fontSize: '0.8rem', padding: '4px 10px', background: 'rgba(16,185,129,0.2)', color: '#10b981', borderRadius: '4px' }}>
                                        Editable
                                    </span>
                                </div>
                                <div style={{ display: 'flex', alignItems: 'center', gap: '15px' }}>
                                    <span style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.5)' }}>
                                        <Icon name="Target" size={14} style={{ marginRight: '6px' }} />
                                        {(assignment.customMarkers || []).length} sections marked
                                    </span>
                                    <button onClick={() => {
                                        // Save edited HTML content from iframe
                                        let htmlContent = docEditorModal.editedHtml || importedDoc.html;
                                        let textContent = docEditorModal.editedText || importedDoc.text;
                                        try {
                                            if (docHtmlRef.current && docHtmlRef.current.contentDocument) {
                                                htmlContent = docHtmlRef.current.contentDocument.body.innerHTML;
                                                textContent = docHtmlRef.current.contentDocument.body.innerText;
                                            }
                                        } catch (e) { console.log('Could not get iframe content:', e); }
                                        setImportedDoc({ ...importedDoc, text: textContent, html: htmlContent });
                                        setDocEditorModal({ show: false, editedText: '', editedHtml: '', viewMode: 'formatted' });
                                    }} style={{ padding: '10px 25px', borderRadius: '8px', border: 'none', background: 'linear-gradient(135deg, #10b981, #059669)', color: '#fff', cursor: 'pointer', fontWeight: 600 }}>
                                        Save & Close
                                    </button>
                                    <button onClick={() => setDocEditorModal({ show: false, editedText: '', editedHtml: '', viewMode: 'formatted' })} style={{ background: 'none', border: 'none', color: 'rgba(255,255,255,0.6)', cursor: 'pointer', padding: '5px' }}>
                                        <Icon name="X" size={24} />
                                    </button>
                                </div>
                            </div>
                            
                            {/* Main Content */}
                            <div style={{ flex: 1, display: 'flex', overflow: 'hidden' }}>
                                {/* Document Area */}
                                <div style={{ flex: 1, display: 'flex', flexDirection: 'column', borderRight: '1px solid rgba(255,255,255,0.1)' }}>
                                    {/* Toolbar */}
                                    <div style={{ padding: '12px 20px', background: 'rgba(251,191,36,0.1)', borderBottom: '1px solid rgba(251,191,36,0.2)', display: 'flex', alignItems: 'center', gap: '15px', flexWrap: 'wrap' }}>
                                        <button onClick={addSelectedAsMarker} style={{ padding: '10px 20px', borderRadius: '8px', border: 'none', background: 'linear-gradient(135deg, #f59e0b, #d97706)', color: '#fff', cursor: 'pointer', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}>
                                            <Icon name="Highlighter" size={16} />Mark Selected Text
                                        </button>
                                        
                                        <span style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.5)' }}>
                                            <Icon name="Edit3" size={14} style={{ marginRight: '4px' }} />
                                            Click anywhere to edit • Select text to mark as gradeable
                                        </span>
                                    </div>
                                    
                                    {/* Document Content - EDITABLE via iframe */}
                                    <iframe
                                        ref={docHtmlRef}
                                        style={{ flex: 1, border: 'none', background: '#fff' }}
                                        srcDoc={'<!DOCTYPE html><html><head><style>body{font-family:Georgia,serif;font-size:16px;line-height:1.7;color:#1a1a2e;padding:40px 50px;margin:0;min-height:100vh;box-sizing:border-box;}table{border-collapse:collapse;width:100%;margin:15px 0;}td,th{border:1px solid #ccc;padding:8px 12px;}th{background:#f5f5f5;}p{margin:10px 0;}h1,h2,h3{margin:20px 0 10px 0;}img{max-width:100%;}</style></head><body contenteditable="true">' + (docEditorModal.editedHtml || importedDoc.html || '<p>No document loaded</p>') + '</body></html>'}
                                    />
                                </div>
                                
                                {/* Right Panel - Marked Sections + Grading Notes */}
                                <div style={{ width: '380px', display: 'flex', flexDirection: 'column', background: '#1a1a2e' }}>
                                    {/* Global AI Notes - Applies to ALL assignments */}
                                    <div style={{ borderBottom: '1px solid rgba(255,255,255,0.1)' }}>
                                        <div style={{ padding: '12px 20px', borderBottom: '1px solid rgba(255,255,255,0.05)', background: 'rgba(16,185,129,0.1)' }}>
                                            <h3 style={{ fontSize: '0.95rem', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px', margin: 0 }}>
                                                <Icon name="Globe" size={16} style={{ color: '#10b981' }} />
                                                Global AI Notes
                                            </h3>
                                            <p style={{ fontSize: '0.75rem', color: 'rgba(255,255,255,0.4)', margin: '4px 0 0 0' }}>
                                                Applies to ALL assignments (persists across restarts)
                                            </p>
                                        </div>
                                        <div style={{ padding: '12px' }}>
                                            <textarea
                                                value={globalAINotes}
                                                onChange={e => setGlobalAINotes(e.target.value)}
                                                onBlur={() => {
                                                    fetch('/api/save-global-settings', {
                                                        method: 'POST',
                                                        headers: { 'Content-Type': 'application/json' },
                                                        body: JSON.stringify({ globalAINotes })
                                                    });
                                                }}
                                                placeholder="Global instructions for AI grading:&#10;• Always accept student's interpretation&#10;• Be lenient on spelling/grammar&#10;• Focus on content, not formatting"
                                                style={{
                                                    width: '100%',
                                                    minHeight: '80px',
                                                    padding: '10px',
                                                    borderRadius: '8px',
                                                    border: '1px solid rgba(16,185,129,0.3)',
                                                    background: 'rgba(0,0,0,0.3)',
                                                    color: '#fff',
                                                    fontSize: '0.85rem',
                                                    lineHeight: '1.4',
                                                    resize: 'vertical',
                                                    outline: 'none',
                                                    boxSizing: 'border-box'
                                                }}
                                            />
                                        </div>
                                    </div>

                                    {/* Assignment-Specific Grading Notes */}
                                    <div style={{ borderBottom: '1px solid rgba(255,255,255,0.1)' }}>
                                        <div style={{ padding: '12px 20px', borderBottom: '1px solid rgba(255,255,255,0.05)', background: 'rgba(99,102,241,0.1)' }}>
                                            <h3 style={{ fontSize: '0.95rem', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px', margin: 0 }}>
                                                <Icon name="MessageSquare" size={16} style={{ color: '#a5b4fc' }} />
                                                Assignment-Specific Notes
                                            </h3>
                                            <p style={{ fontSize: '0.75rem', color: 'rgba(255,255,255,0.4)', margin: '4px 0 0 0' }}>
                                                Saved with this assignment configuration
                                            </p>
                                        </div>
                                        <div style={{ padding: '12px' }}>
                                            <textarea
                                                value={assignment.gradingNotes || ''}
                                                onChange={e => setAssignment({ ...assignment, gradingNotes: e.target.value })}
                                                placeholder="Examples:&#10;• Full credit if they mention 3+ causes&#10;• Deduct points for missing dates&#10;• Look for keywords: treaty, territory"
                                                style={{
                                                    width: '100%',
                                                    minHeight: '100px',
                                                    padding: '10px',
                                                    borderRadius: '8px',
                                                    border: '1px solid rgba(99,102,241,0.3)',
                                                    background: 'rgba(0,0,0,0.3)',
                                                    color: '#fff',
                                                    fontSize: '0.85rem',
                                                    lineHeight: '1.4',
                                                    resize: 'vertical',
                                                    outline: 'none',
                                                    boxSizing: 'border-box'
                                                }}
                                            />
                                        </div>
                                    </div>
                                    
                                    {/* Marked Sections */}
                                    <div style={{ padding: '15px 20px', borderBottom: '1px solid rgba(255,255,255,0.1)' }}>
                                        <h3 style={{ fontSize: '1rem', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px', margin: 0 }}>
                                            <Icon name="Target" size={18} style={{ color: '#fbbf24' }} />
                                            Marked Sections ({(assignment.customMarkers || []).length})
                                        </h3>
                                    </div>
                                    <div style={{ flex: 1, overflow: 'auto', padding: '15px' }}>
                                        {(assignment.customMarkers || []).length === 0 ? (
                                            <div style={{ textAlign: 'center', padding: '30px 15px', color: 'rgba(255,255,255,0.3)' }}>
                                                <Icon name="MousePointer" size={32} />
                                                <p style={{ marginTop: '10px', fontSize: '0.85rem' }}>Select text in the document and click "Mark Selected Text"</p>
                                            </div>
                                        ) : (
                                            <div style={{ display: 'flex', flexDirection: 'column', gap: '10px' }}>
                                                {assignment.customMarkers.map((marker, i) => (
                                                    <div key={i} style={{ padding: '12px 15px', background: 'rgba(251,191,36,0.15)', borderRadius: '10px', border: '1px solid rgba(251,191,36,0.3)' }}>
                                                        <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'flex-start', marginBottom: '8px' }}>
                                                            <span style={{ fontSize: '0.75rem', color: '#fbbf24', fontWeight: 600 }}>Section {i + 1}</span>
                                                            <button onClick={() => removeMarker(marker)} style={{ background: 'none', border: 'none', color: 'rgba(255,255,255,0.4)', cursor: 'pointer', padding: '0' }}>
                                                                <Icon name="Trash2" size={14} />
                                                            </button>
                                                        </div>
                                                        <p style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.8)', margin: 0, lineHeight: '1.5' }}>
                                                            {marker.length > 120 ? marker.substring(0, 120) + '...' : marker}
                                                        </p>
                                                    </div>
                                                ))}
                                            </div>
                                        )}
                                    </div>
                                </div>
                            </div>
                        </div>
                    )}
                    
                    {/* Review/Edit Grade Modal */}
                    {reviewModal.show && editedResults[reviewModal.index] && (
                        <div style={{ position: 'fixed', top: 0, left: 0, right: 0, bottom: 0, background: 'rgba(0,0,0,0.9)', zIndex: 1000, display: 'flex', alignItems: 'center', justifyContent: 'center', padding: '20px' }}>
                            <div style={{ background: '#1a1a2e', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', width: '100%', maxWidth: '1400px', maxHeight: '95vh', overflow: 'hidden', display: 'flex', flexDirection: 'column' }}>
                                <div style={{ padding: '15px 25px', borderBottom: '1px solid rgba(255,255,255,0.1)', display: 'flex', justifyContent: 'space-between', alignItems: 'center' }}>
                                    <h2 style={{ fontSize: '1.2rem', fontWeight: 700, display: 'flex', alignItems: 'center', gap: '10px' }}><Icon name="FileEdit" size={22} />Review & Edit: {editedResults[reviewModal.index].student_name} - {editedResults[reviewModal.index].assignment || editedResults[reviewModal.index].filename}</h2>
                                    <button onClick={() => setReviewModal({ show: false, index: -1 })} style={{ background: 'none', border: 'none', color: 'rgba(255,255,255,0.6)', cursor: 'pointer', padding: '5px' }}><Icon name="X" size={24} /></button>
                                </div>
                                <div style={{ flex: 1, overflowY: 'auto', display: 'grid', gridTemplateColumns: '1fr 1fr', gap: '0' }}>
                                    {(() => {
                                        const r = editedResults[reviewModal.index];
                                        return (
                                            <>
                                                {/* LEFT: Student Work */}
                                                <div style={{ borderRight: '1px solid rgba(255,255,255,0.1)', display: 'flex', flexDirection: 'column' }}>
                                                    <div style={{ padding: '15px 20px', borderBottom: '1px solid rgba(255,255,255,0.1)', background: 'rgba(99,102,241,0.1)' }}>
                                                        <h3 style={{ fontSize: '1rem', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}><Icon name="FileText" size={18} />Student Submission</h3>
                                                    </div>
                                                    <div style={{ flex: 1, padding: '20px', overflowY: 'auto' }}>
                                                        <pre style={{ whiteSpace: 'pre-wrap', wordWrap: 'break-word', fontFamily: 'inherit', fontSize: '0.9rem', lineHeight: '1.7', color: 'rgba(255,255,255,0.85)', margin: 0 }}>{r.student_content || 'No content available. This may be an image file.'}</pre>
                                                    </div>
                                                </div>
                                                
                                                {/* RIGHT: AI Assessment */}
                                                <div style={{ display: 'flex', flexDirection: 'column' }}>
                                                    <div style={{ padding: '15px 20px', borderBottom: '1px solid rgba(255,255,255,0.1)', background: 'rgba(16,185,129,0.1)' }}>
                                                        <h3 style={{ fontSize: '1rem', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}><Icon name="Bot" size={18} />AI Assessment</h3>
                                                    </div>
                                                    <div style={{ flex: 1, padding: '20px', overflowY: 'auto' }}>
                                                        {/* Score and Grade */}
                                                        <div style={{ display: 'grid', gridTemplateColumns: '1fr 1fr', gap: '15px', marginBottom: '20px' }}>
                                                            <div>
                                                                <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Score (0-100)</label>
                                                                <input type="number" min="0" max="100" value={r.score} onChange={e => updateGrade(reviewModal.index, 'score', e.target.value)}
                                                                    style={{ width: '100%', padding: '12px 15px', borderRadius: '10px', border: '2px solid rgba(99,102,241,0.3)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '1.3rem', fontWeight: 700, textAlign: 'center' }} />
                                                            </div>
                                                            <div>
                                                                <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Letter Grade</label>
                                                                <div style={{ padding: '12px 15px', borderRadius: '10px', fontSize: '1.3rem', fontWeight: 700, textAlign: 'center',
                                                                    background: r.letter_grade === 'A' ? 'rgba(74,222,128,0.2)' : r.letter_grade === 'B' ? 'rgba(96,165,250,0.2)' : r.letter_grade === 'C' ? 'rgba(251,191,36,0.2)' : 'rgba(248,113,113,0.2)',
                                                                    color: r.letter_grade === 'A' ? '#4ade80' : r.letter_grade === 'B' ? '#60a5fa' : r.letter_grade === 'C' ? '#fbbf24' : '#f87171' }}>{r.letter_grade}</div>
                                                            </div>
                                                        </div>
                                                        
                                                        {/* Breakdown if available */}
                                                        {r.breakdown && Object.keys(r.breakdown).length > 0 && (
                                                            <div style={{ marginBottom: '20px', padding: '15px', background: 'rgba(255,255,255,0.03)', borderRadius: '10px', border: '1px solid rgba(255,255,255,0.1)' }}>
                                                                <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '10px' }}>Score Breakdown</label>
                                                                <div style={{ display: 'grid', gridTemplateColumns: 'repeat(2, 1fr)', gap: '8px' }}>
                                                                    {Object.entries(r.breakdown).map(([key, value]) => (
                                                                        <div key={key} style={{ display: 'flex', justifyContent: 'space-between', padding: '6px 10px', background: 'rgba(0,0,0,0.2)', borderRadius: '6px', fontSize: '0.85rem' }}>
                                                                            <span style={{ color: 'rgba(255,255,255,0.6)' }}>{key}</span>
                                                                            <span style={{ fontWeight: 600 }}>{value}</span>
                                                                        </div>
                                                                    ))}
                                                                </div>
                                                            </div>
                                                        )}
                                                        
                                                        {/* Feedback */}
                                                        <div>
                                                            <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>AI Feedback (editable)</label>
                                                            <textarea value={r.feedback} onChange={e => updateGrade(reviewModal.index, 'feedback', e.target.value)}
                                                                style={{ width: '100%', minHeight: '250px', padding: '15px', borderRadius: '10px', border: '2px solid rgba(99,102,241,0.3)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem', lineHeight: '1.6', resize: 'vertical', fontFamily: 'inherit' }} />
                                                        </div>
                                                        
                                                        {r.edited && (
                                                            <div style={{ marginTop: '15px', padding: '10px 15px', background: 'rgba(251,191,36,0.1)', borderRadius: '8px', border: '1px solid rgba(251,191,36,0.3)', color: '#fbbf24', fontSize: '0.85rem', display: 'flex', alignItems: 'center', gap: '8px' }}>
                                                                <Icon name="AlertCircle" size={16} /> Modified - changes will be reflected in emails
                                                            </div>
                                                        )}
                                                    </div>
                                                </div>
                                            </>
                                        );
                                    })()}
                                </div>
                                <div style={{ padding: '15px 25px', borderTop: '1px solid rgba(255,255,255,0.1)', display: 'flex', gap: '15px', justifyContent: 'space-between', background: 'rgba(0,0,0,0.2)' }}>
                                    <div style={{ display: 'flex', gap: '10px', alignItems: 'center' }}>
                                        <button onClick={() => reviewModal.index > 0 && setReviewModal({ ...reviewModal, index: reviewModal.index - 1 })} disabled={reviewModal.index === 0}
                                            style={{ padding: '10px 16px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'transparent', color: reviewModal.index === 0 ? 'rgba(255,255,255,0.3)' : '#fff', cursor: reviewModal.index === 0 ? 'not-allowed' : 'pointer' }}><Icon name="ChevronLeft" size={18} /></button>
                                        <button onClick={() => reviewModal.index < editedResults.length - 1 && setReviewModal({ ...reviewModal, index: reviewModal.index + 1 })} disabled={reviewModal.index >= editedResults.length - 1}
                                            style={{ padding: '10px 16px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'transparent', color: reviewModal.index >= editedResults.length - 1 ? 'rgba(255,255,255,0.3)' : '#fff', cursor: reviewModal.index >= editedResults.length - 1 ? 'not-allowed' : 'pointer' }}><Icon name="ChevronRight" size={18} /></button>
                                        <span style={{ padding: '10px 16px', color: 'rgba(255,255,255,0.5)', fontSize: '0.9rem' }}>{reviewModal.index + 1} of {editedResults.length}</span>
                                    </div>
                                    <button onClick={() => setReviewModal({ show: false, index: -1 })} style={{ padding: '12px 24px', borderRadius: '10px', border: 'none', background: 'linear-gradient(135deg, #6366f1, #8b5cf6)', color: '#fff', cursor: 'pointer', fontWeight: 600 }}>Done Reviewing</button>
                                </div>
                            </div>
                        </div>
                    )}
                    
                    <header style={{ textAlign: 'center', marginBottom: '30px', paddingTop: '20px' }}>
                        <div style={{ display: 'inline-flex', alignItems: 'center', gap: '15px', marginBottom: '10px' }}>
                            <div style={{ background: 'linear-gradient(135deg, #6366f1, #8b5cf6)', borderRadius: '16px', padding: '12px', boxShadow: '0 10px 40px rgba(99,102,241,0.3)' }}>
                                <Icon name="GraduationCap" size={36} />
                            </div>
                            <h1 style={{ fontSize: '2.5rem', fontWeight: 800, background: 'linear-gradient(135deg, #fff, #a5b4fc)', WebkitBackgroundClip: 'text', WebkitTextFillColor: 'transparent' }}>
                                Graider
                            </h1>
                        </div>
                        <p style={{ color: 'rgba(255,255,255,0.6)', fontSize: '1.1rem' }}>Your AI-Powered Teaching Assistant</p>
                    </header>

                    <div style={{ maxWidth: '900px', margin: '0 auto' }}>
                        <div style={{ display: 'flex', gap: '10px', marginBottom: '20px' }}>
                            {[{id:'grade',icon:'Home',label:'Home'},{id:'results',icon:'BarChart3',label:'Results'},{id:'builder',icon:'FileEdit',label:'Builder'},{id:'planner',icon:'BookOpen',label:'Planner'},{id:'analytics',icon:'TrendingUp',label:'Analytics'},{id:'settings',icon:'Settings',label:'Settings'}].map(tab => (
                                <button key={tab.id} onClick={() => setActiveTab(tab.id)}
                                    style={{ display: 'flex', alignItems: 'center', gap: '8px', padding: '12px 24px', borderRadius: '12px', border: 'none',
                                        background: activeTab === tab.id ? 'linear-gradient(135deg, #6366f1, #8b5cf6)' : 'rgba(255,255,255,0.05)',
                                        color: activeTab === tab.id ? '#fff' : 'rgba(255,255,255,0.6)', fontSize: '0.95rem', fontWeight: 600, cursor: 'pointer',
                                        boxShadow: activeTab === tab.id ? '0 10px 30px rgba(99,102,241,0.3)' : 'none' }}>
                                    <Icon name={tab.icon} size={18} />{tab.label}
                                </button>
                            ))}
                        </div>

                        {activeTab === 'grade' && (
                            <div className="fade-in">
                                <div style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '30px', marginBottom: '20px' }}>
                                    {status.is_running && (
                                        <div style={{ marginBottom: '25px' }}>
                                            <div style={{ display: 'flex', justifyContent: 'space-between', marginBottom: '10px' }}>
                                                <span style={{ color: 'rgba(255,255,255,0.7)' }}>Grading: {status.current_file}</span>
                                                <span style={{ color: '#6366f1', fontWeight: 600 }}>{status.progress} / {status.total}</span>
                                            </div>
                                            <div style={{ height: '8px', background: 'rgba(255,255,255,0.1)', borderRadius: '4px', overflow: 'hidden' }}>
                                                <div style={{ height: '100%', width: pct+'%', background: 'linear-gradient(90deg, #6366f1, #8b5cf6)', borderRadius: '4px', transition: 'width 0.3s' }} />
                                            </div>
                                        </div>
                                    )}
                                    {/* Assignment Config Selector */}
                                    {savedAssignments.length > 0 && (
                                        <div style={{ marginBottom: '20px', padding: '15px', background: 'rgba(251,191,36,0.1)', borderRadius: '12px', border: '1px solid rgba(251,191,36,0.3)' }}>
                                            <div style={{ display: 'flex', alignItems: 'center', gap: '15px', flexWrap: 'wrap' }}>
                                                <div style={{ display: 'flex', alignItems: 'center', gap: '8px' }}>
                                                    <Icon name="FileCheck" size={18} style={{ color: '#fbbf24' }} />
                                                    <span style={{ fontWeight: 600, color: '#fbbf24' }}>Assignment Config:</span>
                                                </div>
                                                <select
                                                    onChange={e => loadGradingConfig(e.target.value)}
                                                    style={{ padding: '8px 15px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem', minWidth: '200px' }}
                                                >
                                                    <option value="">Use default markers</option>
                                                    <option value="__ALL__">Use ALL markers (all assignments)</option>
                                                    {savedAssignments.map(a => <option key={a} value={a}>{a}</option>)}
                                                </select>
                                                {selectedGradingConfig && (
                                                    <span style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.6)' }}>
                                                        {(selectedGradingConfig.customMarkers || []).length} markers, {selectedGradingConfig.gradingNotes ? 'has notes' : 'no notes'}
                                                    </span>
                                                )}
                                            </div>
                                        </div>
                                    )}

                                    <div style={{ display: 'flex', gap: '15px', justifyContent: 'center', flexWrap: 'wrap' }}>
                                        {!status.is_running ? (
                                            <button onClick={startGrading} disabled={autoGrade}
                                                style={{ display: 'flex', alignItems: 'center', gap: '10px', padding: '16px 40px', borderRadius: '14px', border: 'none',
                                                    background: autoGrade ? 'rgba(255,255,255,0.1)' : 'linear-gradient(135deg, #10b981, #059669)',
                                                    color: '#fff', fontSize: '1.1rem', fontWeight: 700, cursor: autoGrade ? 'not-allowed' : 'pointer',
                                                    boxShadow: autoGrade ? 'none' : '0 10px 40px rgba(16,185,129,0.3)' }}>
                                                <Icon name="Rocket" size={22} />Start Grading
                                            </button>
                                        ) : (
                                            <>
                                                <button disabled
                                                    style={{ display: 'flex', alignItems: 'center', gap: '10px', padding: '16px 40px', borderRadius: '14px', border: 'none',
                                                        background: 'rgba(255,255,255,0.1)', color: '#fff', fontSize: '1.1rem', fontWeight: 700, cursor: 'not-allowed' }}>
                                                    <span style={{ animation: 'spin 1s linear infinite', display: 'inline-flex' }}><Icon name="Loader2" size={22} /></span>Grading...
                                                </button>
                                                <button onClick={stopGrading}
                                                    style={{ display: 'flex', alignItems: 'center', gap: '10px', padding: '16px 30px', borderRadius: '14px', border: 'none',
                                                        background: 'linear-gradient(135deg, #ef4444, #dc2626)', color: '#fff', fontSize: '1rem', fontWeight: 700, cursor: 'pointer',
                                                        boxShadow: '0 10px 40px rgba(239,68,68,0.3)' }}>
                                                    <Icon name="Square" size={18} />Stop
                                                </button>
                                            </>
                                        )}
                                        <button onClick={openResults}
                                            style={{ display: 'flex', alignItems: 'center', gap: '10px', padding: '16px 30px', borderRadius: '14px',
                                                border: '2px solid rgba(99,102,241,0.5)', background: 'transparent', color: '#a5b4fc', fontSize: '1rem', fontWeight: 600, cursor: 'pointer' }}>
                                            <Icon name="FolderOpen" size={20} />Open Results
                                        </button>
                                    </div>
                                    
                                    {/* Auto-Grade Toggle */}
                                    <div style={{ marginTop: '20px', padding: '20px', background: autoGrade ? 'rgba(16,185,129,0.15)' : 'rgba(255,255,255,0.03)', borderRadius: '14px', border: autoGrade ? '2px solid rgba(16,185,129,0.4)' : '1px solid rgba(255,255,255,0.1)', transition: 'all 0.3s' }}>
                                        <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center' }}>
                                            <div style={{ display: 'flex', alignItems: 'center', gap: '12px' }}>
                                                <div style={{ width: '44px', height: '24px', borderRadius: '12px', background: autoGrade ? '#10b981' : 'rgba(255,255,255,0.2)', cursor: 'pointer', position: 'relative', transition: 'background 0.3s' }} onClick={() => setAutoGrade(!autoGrade)}>
                                                    <div style={{ width: '20px', height: '20px', borderRadius: '50%', background: '#fff', position: 'absolute', top: '2px', left: autoGrade ? '22px' : '2px', transition: 'left 0.3s', boxShadow: '0 2px 4px rgba(0,0,0,0.2)' }} />
                                                </div>
                                                <div>
                                                    <div style={{ fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}>
                                                        <Icon name={autoGrade ? 'Zap' : 'ZapOff'} size={18} />
                                                        Auto-Grade Mode
                                                    </div>
                                                    <div style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.5)', marginTop: '2px' }}>
                                                        {autoGrade ? 'Watching for new files... grades automatically' : 'Click to enable automatic grading'}
                                                    </div>
                                                </div>
                                            </div>
                                            {autoGrade && watchStatus.watching && (
                                                <div style={{ textAlign: 'right', fontSize: '0.85rem' }}>
                                                    <div style={{ color: '#4ade80', display: 'flex', alignItems: 'center', gap: '6px', justifyContent: 'flex-end' }}>
                                                        <span style={{ width: '8px', height: '8px', borderRadius: '50%', background: '#4ade80', animation: 'pulse 2s infinite' }} />
                                                        Watching
                                                    </div>
                                                    <div style={{ color: 'rgba(255,255,255,0.4)', marginTop: '2px' }}>Last check: {watchStatus.lastCheck}</div>
                                                </div>
                                            )}
                                        </div>
                                    </div>
                                </div>

                                <div style={{ background: '#0d1117', borderRadius: '16px', border: '1px solid rgba(255,255,255,0.1)', overflow: 'hidden' }}>
                                    <div style={{ padding: '12px 20px', borderBottom: '1px solid rgba(255,255,255,0.1)', display: 'flex', alignItems: 'center', gap: '10px' }}>
                                        <Icon name="Terminal" size={16} /><span style={{ fontSize: '0.9rem', fontWeight: 600 }}>Output Log</span>
                                        <div style={{ marginLeft: 'auto', display: 'flex', gap: '6px' }}>
                                            <div style={{ width: 12, height: 12, borderRadius: '50%', background: '#ff5f57' }} />
                                            <div style={{ width: 12, height: 12, borderRadius: '50%', background: '#febc2e' }} />
                                            <div style={{ width: 12, height: 12, borderRadius: '50%', background: '#28c840' }} />
                                        </div>
                                    </div>
                                    <div ref={logRef} style={{ height: '300px', overflowY: 'auto', padding: '20px', fontFamily: 'Monaco, monospace', fontSize: '13px', lineHeight: 1.6 }}>
                                        {status.log.length === 0 ? (
                                            <div style={{ color: 'rgba(255,255,255,0.3)' }}>Ready to grade. Click "Start Grading" to begin...</div>
                                        ) : status.log.map((line, i) => (
                                            <div key={i} style={{ color: line.includes('❌') ? '#f87171' : line.includes('✅') ? '#4ade80' : line.includes('📝') ? '#60a5fa' : line.includes('🤖') ? '#c084fc' : 'rgba(255,255,255,0.8)' }}>{line}</div>
                                        ))}
                                    </div>
                                </div>
                            </div>
                        )}

                        {activeTab === 'results' && (
                            <div className="fade-in" style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '30px' }}>
                                <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center', marginBottom: '20px' }}>
                                    <h2 style={{ fontSize: '1.3rem', fontWeight: 700, display: 'flex', alignItems: 'center', gap: '10px' }}><Icon name="Trophy" size={24} />Grading Results</h2>
                                    {status.results.length > 0 && (
                                        <button onClick={previewEmails} disabled={emailStatus.sending}
                                            style={{ display: 'flex', alignItems: 'center', gap: '8px', padding: '12px 24px', borderRadius: '12px', border: 'none',
                                                background: emailStatus.sending ? 'rgba(255,255,255,0.1)' : 'linear-gradient(135deg, #6366f1, #8b5cf6)',
                                                color: '#fff', fontSize: '0.95rem', fontWeight: 600, cursor: emailStatus.sending ? 'not-allowed' : 'pointer',
                                                boxShadow: '0 4px 15px rgba(99,102,241,0.3)' }}>
                                            {emailStatus.sending ? (<><span style={{ animation: 'spin 1s linear infinite', display: 'inline-flex' }}><Icon name="Loader2" size={18} /></span>Sending...</>)
                                                : (<><Icon name="Mail" size={18} />Email Students</>)}
                                        </button>
                                    )}
                                </div>
                                {emailStatus.message && (
                                    <div style={{ marginBottom: '20px', padding: '12px 16px', borderRadius: '10px', 
                                        background: emailStatus.message.includes('Error') ? 'rgba(248,113,113,0.1)' : 'rgba(74,222,128,0.1)',
                                        border: `1px solid ${emailStatus.message.includes('Error') ? 'rgba(248,113,113,0.3)' : 'rgba(74,222,128,0.3)'}`,
                                        color: emailStatus.message.includes('Error') ? '#f87171' : '#4ade80' }}>
                                        {emailStatus.message}
                                    </div>
                                )}
                                {status.results.length === 0 ? (
                                    <div style={{ textAlign: 'center', padding: '60px 20px', color: 'rgba(255,255,255,0.4)' }}>
                                        <Icon name="FileQuestion" size={48} /><p style={{ marginTop: '15px' }}>No results yet. Start grading to see results here.</p>
                                    </div>
                                ) : (
                                    <table style={{ width: '100%', borderCollapse: 'collapse' }}>
                                        <thead><tr style={{ borderBottom: '2px solid rgba(255,255,255,0.1)' }}>
                                            <th style={{ padding: '12px', textAlign: 'left', color: 'rgba(255,255,255,0.6)' }}>Student</th>
                                            <th style={{ padding: '12px', textAlign: 'left', color: 'rgba(255,255,255,0.6)' }}>Assignment</th>
                                            <th style={{ padding: '12px', textAlign: 'center', color: 'rgba(255,255,255,0.6)' }}>Score</th>
                                            <th style={{ padding: '12px', textAlign: 'center', color: 'rgba(255,255,255,0.6)' }}>Grade</th>
                                            <th style={{ padding: '12px', textAlign: 'center', color: 'rgba(255,255,255,0.6)' }}>Action</th>
                                        </tr></thead>
                                        <tbody>
                                            {(editedResults.length > 0 ? editedResults : status.results).map((r, i) => (
                                                <tr key={i} style={{ borderBottom: '1px solid rgba(255,255,255,0.05)', cursor: 'pointer', transition: 'background 0.2s' }} 
                                                    onMouseOver={e => e.currentTarget.style.background = 'rgba(99,102,241,0.1)'} 
                                                    onMouseOut={e => e.currentTarget.style.background = 'transparent'}>
                                                    <td style={{ padding: '14px 12px' }} onClick={() => openReview(i)}>
                                                        <div style={{ fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}>
                                                            {r.student_name}
                                                            {r.edited && <span style={{ background: 'rgba(251,191,36,0.2)', color: '#fbbf24', padding: '2px 6px', borderRadius: '4px', fontSize: '0.7rem' }}>edited</span>}
                                                        </div>
                                                        <div style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.4)' }}>ID: {r.student_id}</div>
                                                    </td>
                                                    <td style={{ padding: '14px 12px', color: 'rgba(255,255,255,0.7)', fontSize: '0.9rem' }} onClick={() => openReview(i)}>{r.assignment || r.filename}</td>
                                                    <td style={{ padding: '14px 12px', textAlign: 'center' }} onClick={() => openReview(i)}>
                                                        <span style={{ background: r.score >= 90 ? 'rgba(74,222,128,0.2)' : r.score >= 80 ? 'rgba(96,165,250,0.2)' : r.score >= 70 ? 'rgba(251,191,36,0.2)' : 'rgba(248,113,113,0.2)',
                                                            color: r.score >= 90 ? '#4ade80' : r.score >= 80 ? '#60a5fa' : r.score >= 70 ? '#fbbf24' : '#f87171', padding: '6px 14px', borderRadius: '20px', fontWeight: 700 }}>{r.score}</span>
                                                    </td>
                                                    <td style={{ padding: '14px 12px', textAlign: 'center' }} onClick={() => openReview(i)}>
                                                        <span style={{ fontSize: '1.2rem', fontWeight: 800, color: r.letter_grade === 'A' ? '#4ade80' : r.letter_grade === 'B' ? '#60a5fa' : r.letter_grade === 'C' ? '#fbbf24' : '#f87171' }}>{r.letter_grade}</span>
                                                    </td>
                                                    <td style={{ padding: '14px 12px', textAlign: 'center' }}>
                                                        <button onClick={() => openReview(i)} style={{ padding: '6px 12px', borderRadius: '6px', border: 'none', background: 'rgba(99,102,241,0.2)', color: '#a5b4fc', cursor: 'pointer', fontSize: '0.85rem', display: 'inline-flex', alignItems: 'center', gap: '5px' }}>
                                                            <Icon name="Edit" size={14} /> Review
                                                        </button>
                                                    </td>
                                                </tr>
                                            ))}
                                        </tbody>
                                    </table>
                                )}
                            </div>
                        )}

                        {activeTab === 'settings' && (
                            <div className="fade-in" style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '30px' }}>
                                <h2 style={{ fontSize: '1.3rem', fontWeight: 700, marginBottom: '25px', display: 'flex', alignItems: 'center', gap: '10px' }}><Icon name="Settings" size={24} />Configuration</h2>
                                {[{key:'assignments_folder',label:'Assignments Folder',icon:'FolderInput',type:'folder'},{key:'output_folder',label:'Output Folder',icon:'FolderOutput',type:'folder'},{key:'roster_file',label:'Roster File (Excel or CSV)',icon:'FileSpreadsheet',type:'file'}].map(f => (
                                    <div key={f.key} style={{ marginBottom: '20px' }}>
                                        <label style={{ display: 'flex', alignItems: 'center', gap: '8px', marginBottom: '8px', color: 'rgba(255,255,255,0.7)', fontWeight: 500 }}><Icon name={f.icon} size={18} />{f.label}</label>
                                        <div style={{ display: 'flex', gap: '10px' }}>
                                            <input type="text" value={config[f.key]} onChange={e => setConfig({ ...config, [f.key]: e.target.value })}
                                                style={{ flex: 1, padding: '14px 18px', borderRadius: '12px', border: '2px solid rgba(255,255,255,0.1)', background: 'rgba(0,0,0,0.2)', color: '#fff', fontSize: '0.95rem', outline: 'none' }} />
                                            <button onClick={async () => {
                                                try {
                                                    const res = await fetch('/api/browse?type=' + f.type);
                                                    const data = await res.json();
                                                    if (data.path) setConfig({ ...config, [f.key]: data.path });
                                                } catch (e) { console.error(e); }
                                            }} style={{ padding: '14px 20px', borderRadius: '12px', border: 'none', background: 'linear-gradient(135deg, #6366f1, #8b5cf6)', color: '#fff', cursor: 'pointer', fontWeight: 600, whiteSpace: 'nowrap' }}>
                                                Browse
                                            </button>
                                        </div>
                                    </div>
                                ))}
                                <div style={{ marginBottom: '20px' }}>
                                    <label style={{ display: 'flex', alignItems: 'center', gap: '8px', marginBottom: '8px', color: 'rgba(255,255,255,0.7)', fontWeight: 500 }}><Icon name="Calendar" size={18} />Grading Period</label>
                                    <select value={config.grading_period} onChange={e => setConfig({ ...config, grading_period: e.target.value })} style={{ padding: '14px 18px', borderRadius: '12px', border: '2px solid rgba(255,255,255,0.1)', background: 'rgba(0,0,0,0.2)', color: '#fff', fontSize: '0.95rem', minWidth: '200px' }}>
                                        <option value="Q1">Q1 - Quarter 1</option>
                                        <option value="Q2">Q2 - Quarter 2</option>
                                        <option value="Q3">Q3 - Quarter 3</option>
                                        <option value="Q4">Q4 - Quarter 4</option>
                                        <option value="S1">S1 - Semester 1</option>
                                        <option value="S2">S2 - Semester 2</option>
                                    </select>
                                </div>

                                <div style={{ marginTop: '30px', padding: '20px', background: 'rgba(99,102,241,0.1)', borderRadius: '12px', border: '1px solid rgba(99,102,241,0.2)' }}>
                                    <h3 style={{ fontSize: '0.95rem', fontWeight: 600, marginBottom: '12px', display: 'flex', alignItems: 'center', gap: '8px' }}><Icon name="FileCheck" size={18} />Supported File Types</h3>
                                    <div style={{ display: 'flex', flexWrap: 'wrap', gap: '8px' }}>
                                        {['.docx','.txt','.jpg','.jpeg','.png'].map(ext => (
                                            <span key={ext} style={{ padding: '6px 12px', background: 'rgba(255,255,255,0.1)', borderRadius: '6px', fontSize: '0.85rem', fontFamily: 'Monaco, monospace' }}>{ext}</span>
                                        ))}
                                    </div>
                                </div>

                                {/* Rubric Editor */}
                                <div style={{ marginTop: '30px', padding: '25px', background: 'rgba(16,185,129,0.1)', borderRadius: '12px', border: '1px solid rgba(16,185,129,0.2)' }}>
                                    <h3 style={{ fontSize: '1.1rem', fontWeight: 600, marginBottom: '20px', display: 'flex', alignItems: 'center', gap: '10px' }}><Icon name="ClipboardList" size={20} />Grading Rubric</h3>
                                    
                                    <div style={{ display: 'grid', gridTemplateColumns: '1fr 1fr 1fr', gap: '15px', marginBottom: '20px' }}>
                                        <div>
                                            <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Grade Level</label>
                                            <select value={rubric.gradeLevel} onChange={e => setRubric({ ...rubric, gradeLevel: e.target.value })}
                                                style={{ width: '100%', padding: '10px 12px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem' }}>
                                                {['4th', '5th', '6th', '7th', '8th', '9th', '10th', '11th', '12th'].map(g => <option key={g} value={g}>{g} Grade</option>)}
                                            </select>
                                        </div>
                                        <div>
                                            <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Subject</label>
                                            <select value={rubric.subject} onChange={e => setRubric({ ...rubric, subject: e.target.value })}
                                                style={{ width: '100%', padding: '10px 12px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem' }}>
                                                {['Social Studies', 'English/ELA', 'Math', 'Science', 'History', 'Foreign Language', 'Other'].map(s => <option key={s} value={s}>{s}</option>)}
                                            </select>
                                        </div>
                                        <div>
                                            <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Grading Style</label>
                                            <label style={{ display: 'flex', alignItems: 'center', gap: '10px', padding: '10px 12px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', cursor: 'pointer' }}>
                                                <input type="checkbox" checked={rubric.generous} onChange={e => setRubric({ ...rubric, generous: e.target.checked })} />
                                                <span style={{ fontSize: '0.9rem' }}>Generous grading</span>
                                            </label>
                                        </div>
                                    </div>
                                    
                                    <div style={{ marginBottom: '15px' }}>
                                        <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center', marginBottom: '10px' }}>
                                            <label style={{ fontSize: '0.9rem', fontWeight: 600 }}>Categories (Total: {rubric.categories.reduce((sum, c) => sum + c.points, 0)} points)</label>
                                            <button onClick={addCategory} style={{ padding: '6px 12px', borderRadius: '6px', border: 'none', background: 'rgba(16,185,129,0.3)', color: '#4ade80', cursor: 'pointer', fontSize: '0.85rem', display: 'flex', alignItems: 'center', gap: '5px' }}>
                                                <Icon name="Plus" size={14} /> Add
                                            </button>
                                        </div>
                                        
                                        {rubric.categories.map((cat, i) => (
                                            <div key={i} style={{ display: 'grid', gridTemplateColumns: '2fr 80px 3fr 40px', gap: '10px', marginBottom: '10px', alignItems: 'center' }}>
                                                <input type="text" value={cat.name} onChange={e => updateCategory(i, 'name', e.target.value)} placeholder="Category name"
                                                    style={{ padding: '10px 12px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem' }} />
                                                <input type="number" value={cat.points} onChange={e => updateCategory(i, 'points', e.target.value)} min="0" max="100"
                                                    style={{ padding: '10px 12px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem', textAlign: 'center' }} />
                                                <input type="text" value={cat.description} onChange={e => updateCategory(i, 'description', e.target.value)} placeholder="Description"
                                                    style={{ padding: '10px 12px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.85rem' }} />
                                                <button onClick={() => removeCategory(i)} style={{ padding: '8px', borderRadius: '6px', border: 'none', background: 'rgba(248,113,113,0.2)', color: '#f87171', cursor: 'pointer' }}>
                                                    <Icon name="Trash2" size={16} />
                                                </button>
                                            </div>
                                        ))}
                                    </div>
                                    
                                    <button onClick={saveRubric} style={{ padding: '12px 24px', borderRadius: '10px', border: 'none', background: 'linear-gradient(135deg, #10b981, #059669)', color: '#fff', cursor: 'pointer', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}>
                                        <Icon name="Save" size={18} /> Save Rubric
                                    </button>
                                </div>
                            </div>
                        )}
                        
                        {activeTab === 'builder' && (
                            <div className="fade-in">
                                {/* Saved Assignments Section */}
                                <div style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '25px', marginBottom: '20px' }}>
                                    <button
                                        onClick={() => {
                                            setAssignment({ title: '', subject: 'Social Studies', totalPoints: 100, instructions: '', questions: [], customMarkers: [], gradingNotes: '' });
                                            setImportedDoc({ text: '', html: '', filename: '', loading: false });
                                            setLoadedAssignmentName('');
                                        }}
                                        style={{ padding: '8px 16px', borderRadius: '8px', border: 'none', background: 'linear-gradient(135deg, #6366f1, #8b5cf6)', color: '#fff', cursor: 'pointer', fontWeight: 600, fontSize: '0.85rem', marginBottom: '15px' }}
                                    >+ New Assignment</button>
                                    <h3 style={{ fontSize: '1.1rem', fontWeight: 700, display: 'flex', alignItems: 'center', gap: '10px', margin: '0 0 15px 0' }}>
                                        <Icon name="FolderOpen" size={20} style={{ color: '#10b981' }} />
                                        Saved Assignment Configs ({savedAssignments.length})
                                    </h3>

                                    {savedAssignments.length === 0 ? (
                                        <div style={{ padding: '30px', textAlign: 'center', color: 'rgba(255,255,255,0.4)' }}>
                                            <Icon name="FileQuestion" size={40} style={{ marginBottom: '10px', opacity: 0.5 }} />
                                            <p style={{ margin: 0 }}>No saved assignments yet. Create one below!</p>
                                        </div>
                                    ) : (
                                        <div style={{ display: 'grid', gridTemplateColumns: 'repeat(auto-fill, minmax(280px, 1fr))', gap: '12px' }}>
                                            {savedAssignments.map(name => (
                                                <div
                                                    key={name}
                                                    style={{
                                                        padding: '15px',
                                                        background: loadedAssignmentName === name ? 'rgba(99,102,241,0.2)' : 'rgba(0,0,0,0.2)',
                                                        borderRadius: '12px',
                                                        border: loadedAssignmentName === name ? '2px solid rgba(99,102,241,0.5)' : '1px solid rgba(255,255,255,0.1)',
                                                        cursor: 'pointer',
                                                        transition: 'all 0.2s'
                                                    }}
                                                    onClick={() => loadAssignment(name)}
                                                >
                                                    <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'flex-start' }}>
                                                        <div style={{ flex: 1 }}>
                                                            <div style={{ fontWeight: 600, marginBottom: '5px', display: 'flex', alignItems: 'center', gap: '8px' }}>
                                                                <Icon name="FileText" size={16} style={{ color: '#a5b4fc' }} />
                                                                {name}
                                                            </div>
                                                            <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)' }}>
                                                                Click to load and edit
                                                            </div>
                                                        </div>
                                                        <button
                                                            onClick={(e) => {
                                                                e.stopPropagation();
                                                                if (confirm(`Delete "${name}"?`)) {
                                                                    fetch('/api/delete-assignment?name=' + encodeURIComponent(name), { method: 'DELETE' })
                                                                        .then(() => {
                                                                            setSavedAssignments(savedAssignments.filter(a => a !== name));
                                                                            if (loadedAssignmentName === name) {
                                                                                setAssignment({ title: '', subject: 'Social Studies', totalPoints: 100, instructions: '', questions: [], customMarkers: [], gradingNotes: '' });
                                                                                setLoadedAssignmentName('');
                                                                            }
                                                                        });
                                                                }
                                                            }}
                                                            style={{ padding: '4px', background: 'none', border: 'none', color: 'rgba(255,255,255,0.3)', cursor: 'pointer' }}
                                                        >
                                                            <Icon name="Trash2" size={14} />
                                                        </button>
                                                    </div>
                                                </div>
                                            ))}
                                        </div>
                                    )}
                                </div>

                                {/* Assignment Editor */}
                                <div style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '30px' }}>
                                <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center', marginBottom: '25px' }}>
                                    <h2 style={{ fontSize: '1.3rem', fontWeight: 700, display: 'flex', alignItems: 'center', gap: '10px' }}>
                                        <Icon name="FileEdit" size={24} />
                                        {assignment.title ? `Editing: ${assignment.title}` : 'New Assignment'}
                                    </h2>
                                    {assignment.title && (
                                        <span style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.5)' }}>
                                            {(assignment.customMarkers || []).length} markers • {assignment.gradingNotes ? 'Has notes' : 'No notes'}
                                        </span>
                                    )}
                                </div>
                                
                                {/* Assignment Details */}
                                <div style={{ display: 'grid', gridTemplateColumns: '2fr 1fr 1fr', gap: '15px', marginBottom: '25px' }}>
                                    <div>
                                        <label style={{ display: 'block', fontSize: '0.85rem', color: 'rgba(255,255,255,0.6)', marginBottom: '6px' }}>Assignment Title</label>
                                        <input type="text" value={assignment.title} onChange={e => setAssignment({ ...assignment, title: e.target.value })} placeholder="e.g., Louisiana Purchase Quiz"
                                            style={{ width: '100%', padding: '12px 15px', borderRadius: '10px', border: '2px solid rgba(99,102,241,0.3)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '1rem' }} />
                                    </div>
                                    <div>
                                        <label style={{ display: 'block', fontSize: '0.85rem', color: 'rgba(255,255,255,0.6)', marginBottom: '6px' }}>Subject</label>
                                        <select value={assignment.subject} onChange={e => setAssignment({ ...assignment, subject: e.target.value })}
                                            style={{ width: '100%', padding: '12px 15px', borderRadius: '10px', border: '2px solid rgba(99,102,241,0.3)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '1rem' }}>
                                            {Object.keys(markerLibrary).map(s => <option key={s} value={s}>{s}</option>)}
                                        </select>
                                    </div>
                                    <div>
                                        <label style={{ display: 'block', fontSize: '0.85rem', color: 'rgba(255,255,255,0.6)', marginBottom: '6px' }}>Total Points</label>
                                        <input type="number" value={assignment.totalPoints} onChange={e => setAssignment({ ...assignment, totalPoints: parseInt(e.target.value) || 100 })}
                                            style={{ width: '100%', padding: '12px 15px', borderRadius: '10px', border: '2px solid rgba(99,102,241,0.3)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '1rem', textAlign: 'center' }} />
                                    </div>
                                </div>

                                {/* Import Document Section */}
                                <div style={{ marginBottom: '25px', padding: '20px', background: 'rgba(251,191,36,0.1)', borderRadius: '12px', border: '1px solid rgba(251,191,36,0.3)' }}>
                                    <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center' }}>
                                        <div>
                                            <h3 style={{ fontSize: '1rem', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px', marginBottom: '5px' }}>
                                                <Icon name="FileUp" size={20} />Import Document & Mark Sections
                                            </h3>
                                            <p style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.5)', margin: 0 }}>
                                                {importedDoc.text ? (
                                                    <><strong style={{ color: '#fbbf24' }}>{importedDoc.filename}</strong> loaded • {(assignment.customMarkers || []).length} sections marked</>
                                                ) : 'Import a Word or PDF to highlight gradeable sections'}
                                            </p>
                                        </div>
                                        <div style={{ display: 'flex', gap: '10px' }}>
                                            <input type="file" ref={fileInputRef} onChange={handleDocImport} accept=".docx,.pdf,.doc,.txt" style={{ display: 'none' }} />
                                            {importedDoc.text && (
                                                <button onClick={openDocEditor} style={{ padding: '10px 20px', borderRadius: '8px', border: '2px solid rgba(251,191,36,0.5)', background: 'transparent', color: '#fbbf24', cursor: 'pointer', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}>
                                                    <Icon name="Edit" size={16} />Edit & Mark
                                                </button>
                                            )}
                                            <button onClick={() => fileInputRef.current?.click()} style={{ padding: '10px 20px', borderRadius: '8px', border: 'none', background: 'linear-gradient(135deg, #f59e0b, #d97706)', color: '#fff', cursor: 'pointer', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}>
                                                <Icon name="Upload" size={16} />{importedDoc.loading ? 'Loading...' : importedDoc.text ? 'Import New' : 'Import Word/PDF'}
                                            </button>
                                            <button
                                                onClick={() => {
                                                    const nl = String.fromCharCode(10);
                                                    const blankHtml = '<h1>' + (assignment.title || 'New Assignment') + '</h1><p>Name: ___________________ Date: ___________</p><h2>Directions:</h2><p>Enter your directions here...</p><h2>Question 1:</h2><p></p><h2>Question 2:</h2><p></p>';
                                                    const blankText = 'New Assignment' + nl + 'Name: ___________________ Date: ___________' + nl + 'Directions:' + nl + 'Enter your directions here...' + nl + 'Question 1:' + nl + nl + 'Question 2:' + nl;
                                                    setImportedDoc({ text: blankText, html: blankHtml, filename: 'New Document', loading: false });
                                                    setLoadedAssignmentName('');
                                                    setDocEditorModal({ show: true, editedText: '', editedHtml: blankHtml, viewMode: 'formatted' });
                                                }}
                                                style={{ padding: '10px 20px', borderRadius: '8px', border: '2px solid rgba(99,102,241,0.5)', background: 'transparent', color: '#a5b4fc', cursor: 'pointer', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}
                                            >
                                                <Icon name="FilePlus" size={16} />Create Blank
                                            </button>
                                        </div>
                                    </div>
                                    
                                    {/* Manual Marker Input */}
                                    <div style={{ marginTop: '15px', display: 'flex', gap: '10px', alignItems: 'center' }}>
                                        <input
                                            type="text"
                                            id="manualMarkerInput"
                                            placeholder="Type a marker phrase and press Add..."
                                            style={{ flex: 1, padding: '10px 15px', borderRadius: '8px', border: '1px solid rgba(251,191,36,0.3)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem' }}
                                            onKeyDown={(e) => {
                                                if (e.key === 'Enter' && e.target.value.trim()) {
                                                    const newMarker = e.target.value.trim();
                                                    if (!(assignment.customMarkers || []).includes(newMarker)) {
                                                        setAssignment({ ...assignment, customMarkers: [...(assignment.customMarkers || []), newMarker] });
                                                    }
                                                    e.target.value = '';
                                                }
                                            }}
                                        />
                                        <button
                                            onClick={() => {
                                                const input = document.getElementById('manualMarkerInput');
                                                if (input && input.value.trim()) {
                                                    const newMarker = input.value.trim();
                                                    if (!(assignment.customMarkers || []).includes(newMarker)) {
                                                        setAssignment({ ...assignment, customMarkers: [...(assignment.customMarkers || []), newMarker] });
                                                    }
                                                    input.value = '';
                                                }
                                            }}
                                            style={{ padding: '10px 16px', borderRadius: '8px', border: 'none', background: 'rgba(251,191,36,0.3)', color: '#fbbf24', cursor: 'pointer', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '6px' }}
                                        >
                                            <Icon name="Plus" size={16} />Add
                                        </button>
                                    </div>

                                    {/* Custom Markers Summary */}
                                    {(assignment.customMarkers || []).length > 0 && (
                                        <div style={{ marginTop: '15px', display: 'flex', flexWrap: 'wrap', gap: '8px' }}>
                                            {assignment.customMarkers.map((marker, i) => (
                                                <div key={i} style={{ display: 'flex', alignItems: 'center', gap: '6px', padding: '6px 12px', background: 'rgba(251,191,36,0.2)', borderRadius: '6px', border: '1px solid rgba(251,191,36,0.3)' }}>
                                                    <Icon name="Target" size={12} style={{ color: '#fbbf24' }} />
                                                    <span style={{ fontSize: '0.8rem', color: '#fff', maxWidth: '200px', overflow: 'hidden', textOverflow: 'ellipsis', whiteSpace: 'nowrap' }}>{marker}</span>
                                                    <button onClick={() => removeMarker(marker)} style={{ background: 'none', border: 'none', color: 'rgba(255,255,255,0.4)', cursor: 'pointer', padding: '0', marginLeft: '4px' }}>
                                                        <Icon name="X" size={12} />
                                                    </button>
                                                </div>
                                            ))}
                                        </div>
                                    )}
                                </div>
                                
                                {/* Marker Library */}
                                <div style={{ marginBottom: '25px', padding: '15px 20px', background: 'rgba(99,102,241,0.1)', borderRadius: '12px', border: '1px solid rgba(99,102,241,0.2)' }}>
                                    <label style={{ display: 'block', fontSize: '0.9rem', fontWeight: 600, marginBottom: '10px' }}><Icon name="Bookmark" size={16} style={{ marginRight: '8px' }} />Suggested Markers for {assignment.subject}</label>
                                    <div style={{ display: 'flex', flexWrap: 'wrap', gap: '8px' }}>
                                        {(markerLibrary[assignment.subject] || []).map((marker, i) => (
                                            <span key={i} style={{ padding: '6px 12px', background: 'rgba(255,255,255,0.1)', borderRadius: '6px', fontSize: '0.85rem', cursor: 'pointer' }}
                                                onClick={() => navigator.clipboard.writeText(marker)} title="Click to copy">{marker}</span>
                                        ))}
                                    </div>
                                </div>
                                
                                {/* Questions */}
                                <div style={{ marginBottom: '20px' }}>
                                    <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center', marginBottom: '15px' }}>
                                        <h3 style={{ fontSize: '1rem', fontWeight: 600 }}>Questions ({assignment.questions.length}) - {assignment.questions.reduce((sum, q) => sum + (q.points || 0), 0)} pts</h3>
                                        <button onClick={addQuestion} style={{ padding: '8px 16px', borderRadius: '8px', border: 'none', background: 'linear-gradient(135deg, #6366f1, #8b5cf6)', color: '#fff', cursor: 'pointer', fontSize: '0.9rem', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '6px' }}>
                                            <Icon name="Plus" size={16} /> Add Question
                                        </button>
                                    </div>
                                    
                                    {assignment.questions.length === 0 ? (
                                        <div style={{ textAlign: 'center', padding: '40px', background: 'rgba(0,0,0,0.2)', borderRadius: '12px', color: 'rgba(255,255,255,0.4)' }}>
                                            <Icon name="FileQuestion" size={40} />
                                            <p style={{ marginTop: '10px' }}>No questions yet. Click "Add Question" to start building.</p>
                                        </div>
                                    ) : (
                                        <div style={{ display: 'flex', flexDirection: 'column', gap: '15px' }}>
                                            {assignment.questions.map((q, i) => (
                                                <div key={q.id} style={{ background: 'rgba(255,255,255,0.03)', borderRadius: '12px', border: '1px solid rgba(255,255,255,0.1)', padding: '20px' }}>
                                                    <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center', marginBottom: '15px' }}>
                                                        <span style={{ fontSize: '0.9rem', fontWeight: 600, color: '#a5b4fc' }}>Question {i + 1}</span>
                                                        <button onClick={() => removeQuestion(i)} style={{ padding: '6px 10px', borderRadius: '6px', border: 'none', background: 'rgba(248,113,113,0.2)', color: '#f87171', cursor: 'pointer' }}>
                                                            <Icon name="Trash2" size={14} />
                                                        </button>
                                                    </div>
                                                    <div style={{ display: 'grid', gridTemplateColumns: '1fr 150px 100px', gap: '12px', marginBottom: '12px' }}>
                                                        <div>
                                                            <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '4px' }}>Marker (how AI identifies this section)</label>
                                                            <select value={q.marker} onChange={e => updateQuestion(i, 'marker', e.target.value)}
                                                                style={{ width: '100%', padding: '10px 12px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem' }}>
                                                                {(markerLibrary[assignment.subject] || markerLibrary['Other']).map(m => <option key={m} value={m}>{m}</option>)}
                                                                <option value="custom">Custom marker...</option>
                                                            </select>
                                                        </div>
                                                        <div>
                                                            <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '4px' }}>Type</label>
                                                            <select value={q.type} onChange={e => updateQuestion(i, 'type', e.target.value)}
                                                                style={{ width: '100%', padding: '10px 12px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem' }}>
                                                                <option value="short_answer">Short Answer</option>
                                                                <option value="essay">Essay</option>
                                                                <option value="fill_blank">Fill in Blank</option>
                                                                <option value="multiple_choice">Multiple Choice</option>
                                                            </select>
                                                        </div>
                                                        <div>
                                                            <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '4px' }}>Points</label>
                                                            <input type="number" value={q.points} onChange={e => updateQuestion(i, 'points', parseInt(e.target.value) || 0)} min="0"
                                                                style={{ width: '100%', padding: '10px 12px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem', textAlign: 'center' }} />
                                                        </div>
                                                    </div>
                                                    <div>
                                                        <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '4px' }}>Question/Prompt</label>
                                                        <textarea value={q.prompt} onChange={e => updateQuestion(i, 'prompt', e.target.value)} placeholder="Enter the question or prompt..."
                                                            style={{ width: '100%', minHeight: '60px', padding: '10px 12px', borderRadius: '8px', border: '1px solid rgba(255,255,255,0.2)', background: 'rgba(0,0,0,0.3)', color: '#fff', fontSize: '0.9rem', resize: 'vertical' }} />
                                                    </div>
                                                </div>
                                            ))}
                                        </div>
                                    )}
                                </div>
                                
                                {/* Export Buttons */}
                                <div style={{ display: 'flex', gap: '15px', flexWrap: 'wrap' }}>
                                    <button onClick={saveAssignmentConfig} disabled={!assignment.title} style={{ padding: '12px 24px', borderRadius: '10px', border: 'none', background: !assignment.title ? 'rgba(255,255,255,0.1)' : 'linear-gradient(135deg, #10b981, #059669)', color: '#fff', cursor: !assignment.title ? 'not-allowed' : 'pointer', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px' }}>
                                        <Icon name="Save" size={18} /> Save for Grading
                                    </button>
                                    <button onClick={() => exportAssignment('docx')} disabled={!assignment.title || assignment.questions.length === 0} style={{ padding: '12px 24px', borderRadius: '10px', border: '2px solid rgba(99,102,241,0.5)', background: 'transparent', color: '#a5b4fc', cursor: (!assignment.title || assignment.questions.length === 0) ? 'not-allowed' : 'pointer', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px', opacity: (!assignment.title || assignment.questions.length === 0) ? 0.5 : 1 }}>
                                        <Icon name="FileText" size={18} /> Export Word Doc
                                    </button>
                                    <button onClick={() => exportAssignment('pdf')} disabled={!assignment.title || assignment.questions.length === 0} style={{ padding: '12px 24px', borderRadius: '10px', border: '2px solid rgba(248,113,113,0.5)', background: 'transparent', color: '#f87171', cursor: (!assignment.title || assignment.questions.length === 0) ? 'not-allowed' : 'pointer', fontWeight: 600, display: 'flex', alignItems: 'center', gap: '8px', opacity: (!assignment.title || assignment.questions.length === 0) ? 0.5 : 1 }}>
                                        <Icon name="FileType" size={18} /> Export PDF
                                    </button>
                                </div>
                                
                                {/* Preview */}
                                {assignment.questions.length > 0 && (
                                    <div style={{ marginTop: '30px', padding: '25px', background: 'rgba(255,255,255,0.02)', borderRadius: '12px', border: '1px solid rgba(255,255,255,0.1)' }}>
                                        <h3 style={{ fontSize: '1rem', fontWeight: 600, marginBottom: '15px', display: 'flex', alignItems: 'center', gap: '8px' }}><Icon name="Eye" size={18} />Preview</h3>
                                        <div style={{ background: '#fff', color: '#000', padding: '30px', borderRadius: '8px', fontFamily: 'Georgia, serif' }}>
                                            <h1 style={{ fontSize: '1.5rem', marginBottom: '5px', textAlign: 'center' }}>{assignment.title || 'Untitled Assignment'}</h1>
                                            <p style={{ textAlign: 'center', color: '#666', marginBottom: '20px' }}>Name: _________________ Date: _________</p>
                                            {assignment.instructions && <p style={{ marginBottom: '20px', fontStyle: 'italic' }}>{assignment.instructions}</p>}
                                            {assignment.questions.map((q, i) => (
                                                <div key={q.id} style={{ marginBottom: '25px' }}>
                                                    <p style={{ fontWeight: 'bold', marginBottom: '8px' }}>{q.marker} ({q.points} pts)</p>
                                                    <p style={{ marginBottom: '10px' }}>{q.prompt || '[Question prompt]'}</p>
                                                    <div style={{ borderBottom: '1px solid #ccc', height: q.type === 'essay' ? '100px' : '40px' }}></div>
                                                </div>
                                            ))}
                                        </div>
                                    </div>
                                )}
                            </div>
                            </div>
                        )}

                        {activeTab === 'analytics' && (
                            <div className="fade-in">
                                {!analytics || analytics.error ? (
                                    <div style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '60px', textAlign: 'center' }}>
                                        <Icon name="BarChart3" size={64} />
                                        <h2 style={{ marginTop: '20px', fontSize: '1.5rem' }}>No Data Yet</h2>
                                        <p style={{ color: 'rgba(255,255,255,0.5)', marginTop: '10px' }}>Grade some assignments to see analytics here.</p>
                                    </div>
                                ) : (
                                    <>
                                        {/* Period Filter */}
                                        <div style={{ display: 'flex', alignItems: 'center', justifyContent: 'space-between', marginBottom: '20px' }}>
                                            <h2 style={{ fontSize: '1.3rem', fontWeight: 700, display: 'flex', alignItems: 'center', gap: '10px' }}>
                                                <Icon name="BarChart3" size={24} /> Class Analytics
                                            </h2>
                                            <div style={{ display: 'flex', alignItems: 'center', gap: '10px' }}>
                                                <label style={{ fontSize: '0.9rem', color: 'rgba(255,255,255,0.6)' }}>Filter by Period:</label>
                                                <select
                                                    value={analyticsPeriod}
                                                    onChange={e => setAnalyticsPeriod(e.target.value)}
                                                    style={{ padding: '8px 16px', borderRadius: '10px', background: 'rgba(0,0,0,0.3)', border: '1px solid rgba(255,255,255,0.2)', color: '#fff', fontSize: '0.9rem', cursor: 'pointer' }}
                                                >
                                                    <option value="all">All Periods</option>
                                                    {(analytics.available_periods || []).map(p => (
                                                        <option key={p} value={p}>{p}</option>
                                                    ))}
                                                </select>
                                            </div>
                                        </div>

                                        {/* Class Stats Cards */}
                                        <div style={{ display: 'grid', gridTemplateColumns: 'repeat(4, 1fr)', gap: '15px', marginBottom: '20px' }}>
                                            {[
                                                { label: 'Total Graded', value: analytics.class_stats?.total_assignments || 0, icon: 'FileCheck', color: '#6366f1' },
                                                { label: 'Students', value: analytics.class_stats?.total_students || 0, icon: 'Users', color: '#8b5cf6' },
                                                { label: 'Class Average', value: (analytics.class_stats?.class_average || 0) + '%', icon: 'TrendingUp', color: '#10b981' },
                                                { label: 'Highest Score', value: (analytics.class_stats?.highest || 0) + '%', icon: 'Trophy', color: '#f59e0b' }
                                            ].map((stat, i) => (
                                                <div key={i} style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '16px', border: '1px solid rgba(255,255,255,0.1)', padding: '20px' }}>
                                                    <div style={{ display: 'flex', alignItems: 'center', gap: '10px', marginBottom: '10px' }}>
                                                        <div style={{ background: stat.color + '20', padding: '8px', borderRadius: '10px' }}><Icon name={stat.icon} size={20} /></div>
                                                        <span style={{ color: 'rgba(255,255,255,0.6)', fontSize: '0.9rem' }}>{stat.label}</span>
                                                    </div>
                                                    <div style={{ fontSize: '2rem', fontWeight: 800, color: stat.color }}>{stat.value}</div>
                                                </div>
                                            ))}
                                        </div>
                                        
                                        {/* Grade Distribution Pie + Assignment Averages Bar */}
                                        {chartsLoaded ? (
                                        <div style={{ display: 'grid', gridTemplateColumns: '1fr 2fr', gap: '20px', marginBottom: '20px' }}>
                                            {/* Grade Distribution */}
                                            <div style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '25px' }}>
                                                <h3 style={{ fontSize: '1.1rem', fontWeight: 700, marginBottom: '20px', display: 'flex', alignItems: 'center', gap: '10px' }}><Icon name="PieChart" size={20} />Grade Distribution</h3>
                                                <ResponsiveContainer width="100%" height={200}>
                                                    <PieChart>
                                                        <Pie data={[
                                                            { name: 'A', value: analytics.class_stats?.grade_distribution?.A || 0 },
                                                            { name: 'B', value: analytics.class_stats?.grade_distribution?.B || 0 },
                                                            { name: 'C', value: analytics.class_stats?.grade_distribution?.C || 0 },
                                                            { name: 'D', value: analytics.class_stats?.grade_distribution?.D || 0 },
                                                            { name: 'F', value: analytics.class_stats?.grade_distribution?.F || 0 }
                                                        ].filter(d => d.value > 0)} cx="50%" cy="50%" outerRadius={70} dataKey="value" label={({name, value}) => name + ': ' + value}>
                                                            {['#4ade80','#60a5fa','#fbbf24','#f97316','#ef4444'].map((c, i) => <Cell key={i} fill={c} />)}
                                                        </Pie>
                                                        <Tooltip />
                                                    </PieChart>
                                                </ResponsiveContainer>
                                            </div>
                                            
                                            {/* Assignment Averages */}
                                            <div style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '25px' }}>
                                                <h3 style={{ fontSize: '1.1rem', fontWeight: 700, marginBottom: '20px', display: 'flex', alignItems: 'center', gap: '10px' }}><Icon name="BarChart3" size={20} />Assignment Averages</h3>
                                                <ResponsiveContainer width="100%" height={200}>
                                                    <BarChart data={analytics.assignment_stats || []}>
                                                        <CartesianGrid strokeDasharray="3 3" stroke="rgba(255,255,255,0.1)" />
                                                        <XAxis dataKey="name" tick={{ fill: 'rgba(255,255,255,0.6)', fontSize: 11 }} />
                                                        <YAxis domain={[0, 100]} tick={{ fill: 'rgba(255,255,255,0.6)' }} />
                                                        <Tooltip contentStyle={{ background: '#1a1a2e', border: '1px solid rgba(255,255,255,0.1)', borderRadius: '8px' }} />
                                                        <Bar dataKey="average" fill="#6366f1" radius={[4, 4, 0, 0]} />
                                                    </BarChart>
                                                </ResponsiveContainer>
                                            </div>
                                        </div>
                                        ) : (
                                        <div style={{ background: 'rgba(255,255,255,0.03)', borderRadius: '20px', padding: '30px', marginBottom: '20px', textAlign: 'center' }}>
                                            <p style={{ color: 'rgba(255,255,255,0.5)' }}>Charts loading... If this persists, refresh the page.</p>
                                        </div>
                                        )}
                                        
                                        {/* Student Progress Line Chart */}
                                        <div id="student-progress-section" style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: selectedStudent ? '2px solid #6366f1' : '1px solid rgba(255,255,255,0.1)', padding: '25px', marginBottom: '20px', transition: 'border 0.3s' }}>
                                            <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center', marginBottom: '15px' }}>
                                                <h3 style={{ fontSize: '1.1rem', fontWeight: 700, display: 'flex', alignItems: 'center', gap: '10px' }}>
                                                    <Icon name="TrendingUp" size={20} />
                                                    {selectedStudent ? selectedStudent + "'s Progress" : 'Student Progress Over Time'}
                                                </h3>
                                                {selectedStudent && (
                                                    <button onClick={() => setSelectedStudent(null)} style={{ padding: '6px 12px', borderRadius: '8px', border: 'none', background: 'rgba(255,255,255,0.1)', color: '#fff', fontSize: '0.85rem', cursor: 'pointer', display: 'flex', alignItems: 'center', gap: '6px' }}>
                                                        <Icon name="X" size={14} /> Clear Selection
                                                    </button>
                                                )}
                                            </div>
                                            {selectedStudent && (() => {
                                                const studentData = (analytics.student_progress || []).find(s => s.name === selectedStudent);
                                                if (!studentData) return null;
                                                const grades = studentData.grades || [];
                                                const highest = grades.length > 0 ? Math.max(...grades.map(g => g.score)) : 0;
                                                const lowest = grades.length > 0 ? Math.min(...grades.map(g => g.score)) : 0;
                                                return (
                                                    <div style={{ display: 'grid', gridTemplateColumns: 'repeat(4, 1fr)', gap: '15px', marginBottom: '20px' }}>
                                                        <div style={{ background: 'rgba(99,102,241,0.1)', borderRadius: '12px', padding: '15px', textAlign: 'center' }}>
                                                            <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '5px' }}>Average</div>
                                                            <div style={{ fontSize: '1.5rem', fontWeight: 700, color: '#6366f1' }}>{studentData.average}%</div>
                                                        </div>
                                                        <div style={{ background: 'rgba(74,222,128,0.1)', borderRadius: '12px', padding: '15px', textAlign: 'center' }}>
                                                            <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '5px' }}>Highest</div>
                                                            <div style={{ fontSize: '1.5rem', fontWeight: 700, color: '#4ade80' }}>{highest}%</div>
                                                        </div>
                                                        <div style={{ background: 'rgba(248,113,113,0.1)', borderRadius: '12px', padding: '15px', textAlign: 'center' }}>
                                                            <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '5px' }}>Lowest</div>
                                                            <div style={{ fontSize: '1.5rem', fontWeight: 700, color: '#f87171' }}>{lowest}%</div>
                                                        </div>
                                                        <div style={{ background: 'rgba(251,191,36,0.1)', borderRadius: '12px', padding: '15px', textAlign: 'center' }}>
                                                            <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '5px' }}>Assignments</div>
                                                            <div style={{ fontSize: '1.5rem', fontWeight: 700, color: '#fbbf24' }}>{grades.length}</div>
                                                        </div>
                                                    </div>
                                                );
                                            })()}
                                            {!selectedStudent && (
                                            <div style={{ display: 'flex', flexWrap: 'wrap', gap: '8px', marginBottom: '15px' }}>
                                                <span style={{ padding: '6px 12px', fontSize: '0.85rem', color: 'rgba(255,255,255,0.5)' }}>Click a student name below to view details</span>
                                            </div>
                                            )}
                                            {chartsLoaded && (
                                            <ResponsiveContainer width="100%" height={250}>
                                                <LineChart data={(() => {
                                                    const filtered = selectedStudent
                                                        ? (analytics.student_progress || []).filter(s => s.name === selectedStudent)
                                                        : (analytics.student_progress || []);
                                                    const allGrades = filtered.flatMap(s => s.grades.map(g => ({...g, student: s.name.split(' ')[0]})));
                                                    return allGrades.sort((a,b) => a.date.localeCompare(b.date));
                                                })()}>
                                                    <CartesianGrid strokeDasharray="3 3" stroke="rgba(255,255,255,0.1)" />
                                                    <XAxis dataKey="assignment" tick={{ fill: 'rgba(255,255,255,0.6)', fontSize: 10 }} angle={-20} textAnchor="end" height={60} />
                                                    <YAxis domain={[0, 100]} tick={{ fill: 'rgba(255,255,255,0.6)' }} />
                                                    <Tooltip contentStyle={{ background: '#1a1a2e', border: '1px solid rgba(255,255,255,0.1)', borderRadius: '8px' }} />
                                                    <Line type="monotone" dataKey="score" stroke="#6366f1" strokeWidth={3} dot={{ fill: '#6366f1', r: 5 }} />
                                                </LineChart>
                                            </ResponsiveContainer>
                                            )}
                                            {selectedStudent && (() => {
                                                const studentData = (analytics.student_progress || []).find(s => s.name === selectedStudent);
                                                if (!studentData || !studentData.grades || studentData.grades.length === 0) return null;
                                                return (
                                                    <div style={{ marginTop: '20px' }}>
                                                        <h4 style={{ fontSize: '0.95rem', fontWeight: 600, marginBottom: '10px', color: 'rgba(255,255,255,0.7)' }}>Assignment History</h4>
                                                        <div style={{ display: 'grid', gridTemplateColumns: 'repeat(auto-fill, minmax(200px, 1fr))', gap: '10px' }}>
                                                            {studentData.grades.map((g, i) => (
                                                                <div key={i} style={{ background: 'rgba(0,0,0,0.2)', borderRadius: '10px', padding: '12px', display: 'flex', justifyContent: 'space-between', alignItems: 'center' }}>
                                                                    <div>
                                                                        <div style={{ fontWeight: 600, fontSize: '0.9rem' }}>{g.assignment}</div>
                                                                        <div style={{ fontSize: '0.75rem', color: 'rgba(255,255,255,0.4)' }}>{g.date}</div>
                                                                    </div>
                                                                    <span style={{ padding: '4px 10px', borderRadius: '8px', fontWeight: 700, fontSize: '0.9rem',
                                                                        background: g.score >= 90 ? 'rgba(74,222,128,0.2)' : g.score >= 80 ? 'rgba(96,165,250,0.2)' : g.score >= 70 ? 'rgba(251,191,36,0.2)' : 'rgba(248,113,113,0.2)',
                                                                        color: g.score >= 90 ? '#4ade80' : g.score >= 80 ? '#60a5fa' : g.score >= 70 ? '#fbbf24' : '#f87171'
                                                                    }}>{g.score}%</span>
                                                                </div>
                                                            ))}
                                                        </div>
                                                    </div>
                                                );
                                            })()}
                                        </div>
                                        
                                        {/* Two columns: Attention Needed + Top Performers */}
                                        <div style={{ display: 'grid', gridTemplateColumns: '1fr 1fr', gap: '20px', marginBottom: '20px' }}>
                                            {/* Needs Attention */}
                                            <div style={{ background: 'rgba(239,68,68,0.1)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(239,68,68,0.3)', padding: '25px' }}>
                                                <h3 style={{ fontSize: '1.1rem', fontWeight: 700, marginBottom: '15px', display: 'flex', alignItems: 'center', gap: '10px', color: '#f87171' }}><Icon name="AlertTriangle" size={20} />Needs Attention</h3>
                                                {(analytics.attention_needed || []).length === 0 ? (
                                                    <p style={{ color: 'rgba(255,255,255,0.5)' }}>All students are doing well! 🎉</p>
                                                ) : (
                                                    <div style={{ display: 'flex', flexDirection: 'column', gap: '10px' }}>
                                                        {(analytics.attention_needed || []).slice(0, 5).map((s, i) => (
                                                            <div key={i} style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center', padding: '10px 15px', background: 'rgba(0,0,0,0.2)', borderRadius: '10px', cursor: 'pointer', transition: 'background 0.2s' }} onClick={() => setSelectedStudent(s.name)} onMouseOver={e => e.currentTarget.style.background = 'rgba(0,0,0,0.4)'} onMouseOut={e => e.currentTarget.style.background = 'rgba(0,0,0,0.2)'}>
                                                                <span style={{ textDecoration: 'underline', textDecorationStyle: 'dotted' }}>{s.name}</span>
                                                                <div style={{ display: 'flex', alignItems: 'center', gap: '10px' }}>
                                                                    <span style={{ color: '#f87171', fontWeight: 700 }}>{s.average}%</span>
                                                                    <span style={{ fontSize: '0.8rem', padding: '2px 8px', borderRadius: '4px', background: s.trend === 'declining' ? 'rgba(239,68,68,0.3)' : 'rgba(251,191,36,0.3)', color: s.trend === 'declining' ? '#f87171' : '#fbbf24' }}>{s.trend}</span>
                                                                </div>
                                                            </div>
                                                        ))}
                                                    </div>
                                                )}
                                            </div>
                                            
                                            {/* Top Performers */}
                                            <div style={{ background: 'rgba(74,222,128,0.1)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(74,222,128,0.3)', padding: '25px' }}>
                                                <h3 style={{ fontSize: '1.1rem', fontWeight: 700, marginBottom: '15px', display: 'flex', alignItems: 'center', gap: '10px', color: '#4ade80' }}><Icon name="Award" size={20} />Top Performers</h3>
                                                <div style={{ display: 'flex', flexDirection: 'column', gap: '10px' }}>
                                                    {(analytics.top_performers || []).map((s, i) => (
                                                        <div key={i} style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center', padding: '10px 15px', background: 'rgba(0,0,0,0.2)', borderRadius: '10px', cursor: 'pointer', transition: 'background 0.2s' }} onClick={() => setSelectedStudent(s.name)} onMouseOver={e => e.currentTarget.style.background = 'rgba(0,0,0,0.4)'} onMouseOut={e => e.currentTarget.style.background = 'rgba(0,0,0,0.2)'}>
                                                            <div style={{ display: 'flex', alignItems: 'center', gap: '10px' }}>
                                                                <span style={{ width: '24px', height: '24px', borderRadius: '50%', background: i === 0 ? '#fbbf24' : i === 1 ? '#94a3b8' : i === 2 ? '#cd7f32' : 'rgba(255,255,255,0.1)', display: 'flex', alignItems: 'center', justifyContent: 'center', fontSize: '0.75rem', fontWeight: 700 }}>{i + 1}</span>
                                                                <span style={{ textDecoration: 'underline', textDecorationStyle: 'dotted' }}>{s.name}</span>
                                                            </div>
                                                            <span style={{ color: '#4ade80', fontWeight: 700 }}>{s.average}%</span>
                                                        </div>
                                                    ))}
                                                </div>
                                            </div>
                                        </div>
                                        
                                        {/* All Students Table */}
                                        <div style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '25px' }}>
                                            <h3 style={{ fontSize: '1.1rem', fontWeight: 700, marginBottom: '15px', display: 'flex', alignItems: 'center', gap: '10px' }}><Icon name="Users" size={20} />All Students Overview</h3>
                                            <div style={{ overflowX: 'auto' }}>
                                                <table style={{ width: '100%', borderCollapse: 'collapse' }}>
                                                    <thead>
                                                        <tr style={{ borderBottom: '2px solid rgba(255,255,255,0.1)' }}>
                                                            <th style={{ padding: '12px', textAlign: 'left', color: 'rgba(255,255,255,0.6)' }}>Student</th>
                                                            <th style={{ padding: '12px', textAlign: 'center', color: 'rgba(255,255,255,0.6)' }}>Assignments</th>
                                                            <th style={{ padding: '12px', textAlign: 'center', color: 'rgba(255,255,255,0.6)' }}>Average</th>
                                                            <th style={{ padding: '12px', textAlign: 'center', color: 'rgba(255,255,255,0.6)' }}>Trend</th>
                                                        </tr>
                                                    </thead>
                                                    <tbody>
                                                        {(analytics.student_progress || []).map((s, i) => (
                                                            <tr key={i} style={{ borderBottom: '1px solid rgba(255,255,255,0.05)', cursor: 'pointer', background: selectedStudent === s.name ? 'rgba(99,102,241,0.2)' : 'transparent', transition: 'background 0.2s' }} onClick={() => setSelectedStudent(s.name)} onMouseOver={e => { if (selectedStudent !== s.name) e.currentTarget.style.background = 'rgba(255,255,255,0.05)'; }} onMouseOut={e => { if (selectedStudent !== s.name) e.currentTarget.style.background = 'transparent'; }}>
                                                                <td style={{ padding: '12px', fontWeight: 600, textDecoration: 'underline', textDecorationStyle: 'dotted' }}>{s.name}</td>
                                                                <td style={{ padding: '12px', textAlign: 'center' }}>{s.grades.length}</td>
                                                                <td style={{ padding: '12px', textAlign: 'center' }}>
                                                                    <span style={{ padding: '4px 12px', borderRadius: '20px', fontWeight: 700,
                                                                        background: s.average >= 90 ? 'rgba(74,222,128,0.2)' : s.average >= 80 ? 'rgba(96,165,250,0.2)' : s.average >= 70 ? 'rgba(251,191,36,0.2)' : 'rgba(248,113,113,0.2)',
                                                                        color: s.average >= 90 ? '#4ade80' : s.average >= 80 ? '#60a5fa' : s.average >= 70 ? '#fbbf24' : '#f87171'
                                                                    }}>{s.average}%</span>
                                                                </td>
                                                                <td style={{ padding: '12px', textAlign: 'center' }}>
                                                                    <span style={{ display: 'inline-flex', alignItems: 'center', gap: '4px', color: s.trend === 'improving' ? '#4ade80' : s.trend === 'declining' ? '#f87171' : '#94a3b8' }}>
                                                                        <Icon name={s.trend === 'improving' ? 'TrendingUp' : s.trend === 'declining' ? 'TrendingDown' : 'Minus'} size={16} />
                                                                        {s.trend}
                                                                    </span>
                                                                </td>
                                                            </tr>
                                                        ))}
                                                    </tbody>
                                                </table>
                                            </div>
                                        </div>
                                    </>
                                )}
                            </div>
                        )}
                        
                        {activeTab === 'planner' && (
                            <div className="fade-in">
                                <div style={{ display: 'grid', gridTemplateColumns: '300px 1fr', gap: '25px' }}>
                                    {/* Sidebar: Configuration */}
                                    <div style={{ display: 'flex', flexDirection: 'column', gap: '20px' }}>
                                        {/* Basic Settings */}
                                        <div style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '20px' }}>
                                            <h3 style={{ fontSize: '1.1rem', fontWeight: 700, marginBottom: '15px', display: 'flex', alignItems: 'center', gap: '10px' }}>
                                                <Icon name="Settings2" size={20} /> Configuration
                                            </h3>
                                            
                                            <div style={{ display: 'flex', flexDirection: 'column', gap: '15px' }}>
                                                <div>
                                                    <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>State</label>
                                                    <select 
                                                        value={plannerConfig.state}
                                                        onChange={e => setPlannerConfig({...plannerConfig, state: e.target.value})}
                                                        style={{ width: '100%', padding: '10px 12px', borderRadius: '10px', background: 'rgba(0,0,0,0.3)', border: '1px solid rgba(255,255,255,0.1)', color: '#fff', fontSize: '0.9rem' }}
                                                    >
                                                        <option value="FL">Florida</option>
                                                        <option value="TX">Texas</option>
                                                        <option value="NY">New York</option>
                                                        <option value="CA">California</option>
                                                    </select>
                                                </div>
                                                
                                                <div>
                                                    <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Grade Level</label>
                                                    <select 
                                                        value={plannerConfig.grade}
                                                        onChange={e => setPlannerConfig({...plannerConfig, grade: e.target.value})}
                                                        style={{ width: '100%', padding: '10px 12px', borderRadius: '10px', background: 'rgba(0,0,0,0.3)', border: '1px solid rgba(255,255,255,0.1)', color: '#fff', fontSize: '0.9rem' }}
                                                    >
                                                        <option value="K">Kindergarten</option>
                                                        <option value="1">1st Grade</option>
                                                        <option value="2">2nd Grade</option>
                                                        <option value="3">3rd Grade</option>
                                                        <option value="4">4th Grade</option>
                                                        <option value="5">5th Grade</option>
                                                        <option value="6">6th Grade</option>
                                                        <option value="7">7th Grade</option>
                                                        <option value="8">8th Grade</option>
                                                        <option value="9">9th Grade</option>
                                                        <option value="10">10th Grade</option>
                                                        <option value="11">11th Grade</option>
                                                        <option value="12">12th Grade</option>
                                                    </select>
                                                </div>
                                                
                                                <div>
                                                    <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Subject</label>
                                                    <select 
                                                        value={plannerConfig.subject}
                                                        onChange={e => setPlannerConfig({...plannerConfig, subject: e.target.value})}
                                                        style={{ width: '100%', padding: '10px 12px', borderRadius: '10px', background: 'rgba(0,0,0,0.3)', border: '1px solid rgba(255,255,255,0.1)', color: '#fff', fontSize: '0.9rem' }}
                                                    >
                                                        <option value="Civics">Civics</option>
                                                        <option value="History">History</option>
                                                        <option value="Geography">Geography</option>
                                                        <option value="Economics">Economics</option>
                                                        <option value="Math">Math</option>
                                                        <option value="Science">Science</option>
                                                        <option value="ELA">English / ELA</option>
                                                    </select>
                                                </div>
                                            </div>
                                        </div>
                                        




                                        {/* Unit Details */}
                                        <div style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '20px' }}>
                                            <h3 style={{ fontSize: '1.1rem', fontWeight: 700, marginBottom: '15px', display: 'flex', alignItems: 'center', gap: '10px' }}>
                                                <Icon name="FileText" size={20} /> Details
                                            </h3>
                                            
                                            <div style={{ display: 'flex', flexDirection: 'column', gap: '15px' }}>
                                                <div>
                                                    <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Content Type</label>
                                                    <select 
                                                        value={unitConfig.type}
                                                        onChange={e => setUnitConfig({...unitConfig, type: e.target.value})}
                                                        style={{ width: '100%', padding: '10px 12px', borderRadius: '10px', background: 'rgba(0,0,0,0.3)', border: '1px solid rgba(255,255,255,0.1)', color: '#fff', fontSize: '0.9rem' }}
                                                    >
                                                        <option value="Unit Plan">Unit Plan</option>
                                                        <option value="Lesson Plan">Lesson Plan</option>
                                                        <option value="Assignment">Assignment</option>
                                                        <option value="Project">Project</option>
                                                    </select>
                                                </div>

                                                <div>
                                                    <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Title</label>
                                                    <input 
                                                        type="text"
                                                        value={unitConfig.title}
                                                        onChange={e => setUnitConfig({...unitConfig, title: e.target.value})}
                                                        placeholder="e.g., Foundations of Government"
                                                        style={{ width: '100%', padding: '10px 12px', borderRadius: '10px', background: 'rgba(0,0,0,0.3)', border: '1px solid rgba(255,255,255,0.1)', color: '#fff', fontSize: '0.9rem' }}
                                                    />
                                                </div>
                                                
                                                <div style={{ display: 'grid', gridTemplateColumns: '1fr 1fr', gap: '12px' }}>
                                                    <div>
                                                        <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Duration (Days)</label>
                                                        <input
                                                            type="number"
                                                            value={unitConfig.duration}
                                                            onChange={e => setUnitConfig({...unitConfig, duration: parseInt(e.target.value) || 1})}
                                                            min="1" max="20"
                                                            style={{ width: '100%', padding: '10px 12px', borderRadius: '10px', background: 'rgba(0,0,0,0.3)', border: '1px solid rgba(255,255,255,0.1)', color: '#fff', fontSize: '0.9rem' }}
                                                        />
                                                    </div>
                                                    <div>
                                                        <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Period Length (min)</label>
                                                        <input
                                                            type="number"
                                                            value={unitConfig.periodLength}
                                                            onChange={e => setUnitConfig({...unitConfig, periodLength: parseInt(e.target.value) || 50})}
                                                            min="20" max="120"
                                                            style={{ width: '100%', padding: '10px 12px', borderRadius: '10px', background: 'rgba(0,0,0,0.3)', border: '1px solid rgba(255,255,255,0.1)', color: '#fff', fontSize: '0.9rem' }}
                                                        />
                                                    </div>
                                                </div>

                                                <div>
                                                    <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Output Format</label>
                                                    <select 
                                                        value={unitConfig.format}
                                                        onChange={e => setUnitConfig({...unitConfig, format: e.target.value})}
                                                        style={{ width: '100%', padding: '10px 12px', borderRadius: '10px', background: 'rgba(0,0,0,0.3)', border: '1px solid rgba(255,255,255,0.1)', color: '#fff', fontSize: '0.9rem' }}
                                                    >
                                                        <option value="Word">Word Document (.docx)</option>
                                                        <option value="Text">Plain Text (.txt)</option>
                                                    </select>
                                                </div>

                                                <div>
                                                    <label style={{ display: 'block', fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '6px' }}>Additional Requirements / Directions</label>
                                                    <textarea 
                                                        value={unitConfig.requirements || ''}
                                                        onChange={e => setUnitConfig({...unitConfig, requirements: e.target.value})}
                                                        placeholder="e.g. Focus on primary sources, make it a group project, include a debate..."
                                                        style={{ width: '100%', minHeight: '80px', padding: '10px 12px', borderRadius: '10px', background: 'rgba(0,0,0,0.3)', border: '1px solid rgba(255,255,255,0.1)', color: '#fff', fontSize: '0.9rem', resize: 'vertical' }}
                                                    />
                                                </div>
                                                
                                                <button 
                                                    onClick={generateLessonPlan}
                                                    disabled={plannerLoading || selectedStandards.length === 0}
                                                    style={{ 
                                                        width: '100%', padding: '12px', borderRadius: '10px', border: 'none', 
                                                        background: plannerLoading || selectedStandards.length === 0 ? 'rgba(255,255,255,0.1)' : 'linear-gradient(135deg, #10b981, #059669)',
                                                        color: '#fff', fontSize: '0.95rem', fontWeight: 600, cursor: plannerLoading || selectedStandards.length === 0 ? 'not-allowed' : 'pointer',
                                                        display: 'flex', alignItems: 'center', justifyContent: 'center', gap: '8px', marginTop: '10px'
                                                    }}
                                                >
                                                    {plannerLoading ? <Icon name="Loader2" size={18} style={{ animation: 'spin 1s linear infinite' }} /> : <Icon name="Sparkles" size={18} />}
                                                    {plannerLoading ? 'Generating...' : 'Generate Plan'}
                                                </button>
                                            </div>
                                        </div>
                                    </div>

                                    {/* Main Content: Standards & Preview */}
                                    <div style={{ display: 'flex', flexDirection: 'column', gap: '20px' }}>
                                        
                                        {lessonPlan ? (
                                            <div style={{ background: 'rgba(255,255,255,0.03)', backdropFilter: 'blur(20px)', borderRadius: '20px', border: '1px solid rgba(255,255,255,0.1)', padding: '30px', maxHeight: '80vh', overflowY: 'auto' }}>
                                                {/* Header */}
                                                <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'flex-start', marginBottom: '25px', borderBottom: '1px solid rgba(255,255,255,0.1)', paddingBottom: '20px' }}>
                                                    <div style={{ flex: 1 }}>
                                                        <div style={{ display: 'flex', alignItems: 'center', gap: '10px', marginBottom: '10px' }}>
                                                            <h2 style={{ fontSize: '1.8rem', fontWeight: 700, margin: 0 }}>{lessonPlan.title}</h2>
                                                            <span style={{
                                                                background: (lessonPlan.overview || '').includes('MOCK MODE') ? 'rgba(251,191,36,0.2)' : 'rgba(74,222,128,0.2)',
                                                                color: (lessonPlan.overview || '').includes('MOCK MODE') ? '#fbbf24' : '#4ade80',
                                                                fontSize: '0.7rem', fontWeight: 700, padding: '4px 8px', borderRadius: '4px',
                                                                border: (lessonPlan.overview || '').includes('MOCK MODE') ? '1px solid rgba(251,191,36,0.3)' : '1px solid rgba(74,222,128,0.3)'
                                                            }}>
                                                                {(lessonPlan.overview || '').includes('MOCK MODE') ? 'MOCK MODE' : 'AI GENERATED'}
                                                            </span>
                                                        </div>
                                                        <p style={{ color: 'rgba(255,255,255,0.6)', lineHeight: '1.6', marginBottom: '15px' }}>{lessonPlan.overview}</p>

                                                        {/* Essential Questions */}
                                                        {lessonPlan.essential_questions && lessonPlan.essential_questions.length > 0 && (
                                                            <div style={{ background: 'rgba(99,102,241,0.1)', border: '1px solid rgba(99,102,241,0.3)', borderRadius: '10px', padding: '15px', marginTop: '15px' }}>
                                                                <h4 style={{ fontSize: '0.85rem', color: '#a5b4fc', marginBottom: '10px', fontWeight: 600 }}>Essential Questions</h4>
                                                                <ul style={{ margin: 0, paddingLeft: '20px' }}>
                                                                    {lessonPlan.essential_questions.map((q, i) => (
                                                                        <li key={i} style={{ fontSize: '0.9rem', color: 'rgba(255,255,255,0.8)', marginBottom: '5px' }}>{q}</li>
                                                                    ))}
                                                                </ul>
                                                            </div>
                                                        )}
                                                    </div>
                                                    <div style={{ display: 'flex', gap: '10px', marginLeft: '20px' }}>
                                                        <button onClick={exportLessonPlan} style={{ padding: '10px 20px', borderRadius: '10px', border: '1px solid rgba(255,255,255,0.2)', background: 'transparent', color: '#fff', cursor: 'pointer', display: 'flex', alignItems: 'center', gap: '8px', fontWeight: 500 }}>
                                                            <Icon name="Download" size={16} /> Export
                                                        </button>
                                                        <button onClick={() => setLessonPlan(null)} style={{ padding: '10px 20px', borderRadius: '10px', border: 'none', background: 'rgba(255,255,255,0.1)', color: '#fff', cursor: 'pointer', fontWeight: 500 }}>
                                                            Close
                                                        </button>
                                                    </div>
                                                </div>

                                                {/* Days */}
                                                <div style={{ display: 'flex', flexDirection: 'column', gap: '30px' }}>
                                                    {(lessonPlan.days || []).map((day, i) => (
                                                        <div key={i} style={{ background: 'rgba(0,0,0,0.3)', borderRadius: '16px', padding: '25px', border: '1px solid rgba(255,255,255,0.05)' }}>
                                                            {/* Day Header */}
                                                            <div style={{ display: 'flex', alignItems: 'flex-start', gap: '15px', marginBottom: '20px', paddingBottom: '15px', borderBottom: '1px solid rgba(255,255,255,0.1)' }}>
                                                                <div style={{ width: '50px', height: '50px', borderRadius: '12px', background: 'linear-gradient(135deg, #6366f1, #8b5cf6)', display: 'flex', alignItems: 'center', justifyContent: 'center', fontWeight: 700, fontSize: '1.2rem', flexShrink: 0 }}>
                                                                    {day.day}
                                                                </div>
                                                                <div style={{ flex: 1 }}>
                                                                    <h3 style={{ fontSize: '1.3rem', fontWeight: 600, margin: '0 0 8px 0' }}>{day.topic}</h3>
                                                                    <div style={{ display: 'flex', alignItems: 'center', gap: '8px', marginBottom: '8px' }}>
                                                                        <Icon name="Target" size={14} style={{ color: '#10b981' }} />
                                                                        <span style={{ fontSize: '0.9rem', color: '#10b981', fontWeight: 500 }}>Learning Objective:</span>
                                                                    </div>
                                                                    <p style={{ fontSize: '0.9rem', color: 'rgba(255,255,255,0.8)', margin: 0, lineHeight: '1.5' }}>{day.objective}</p>

                                                                    {/* Standards Addressed */}
                                                                    {day.standards_addressed && day.standards_addressed.length > 0 && (
                                                                        <div style={{ display: 'flex', flexWrap: 'wrap', gap: '6px', marginTop: '10px' }}>
                                                                            {day.standards_addressed.map((std, idx) => (
                                                                                <span key={idx} style={{ fontSize: '0.75rem', padding: '3px 8px', borderRadius: '4px', background: 'rgba(99,102,241,0.2)', color: '#a5b4fc', border: '1px solid rgba(99,102,241,0.3)' }}>{std}</span>
                                                                            ))}
                                                                        </div>
                                                                    )}
                                                                </div>
                                                            </div>

                                                            {/* Vocabulary */}
                                                            {day.vocabulary && day.vocabulary.length > 0 && (
                                                                <div style={{ marginBottom: '20px', background: 'rgba(255,255,255,0.03)', padding: '15px', borderRadius: '10px' }}>
                                                                    <h4 style={{ fontSize: '0.9rem', color: '#fbbf24', marginBottom: '12px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="BookOpen" size={14} /> Vocabulary</h4>
                                                                    <div style={{ display: 'grid', gridTemplateColumns: 'repeat(auto-fill, minmax(250px, 1fr))', gap: '10px' }}>
                                                                        {day.vocabulary.map((v, idx) => (
                                                                            <div key={idx} style={{ background: 'rgba(0,0,0,0.2)', padding: '10px', borderRadius: '8px' }}>
                                                                                <span style={{ fontWeight: 600, color: '#fbbf24' }}>{typeof v === 'object' ? v.term : v}</span>
                                                                                {typeof v === 'object' && v.definition && (
                                                                                    <span style={{ color: 'rgba(255,255,255,0.6)', marginLeft: '8px' }}>- {v.definition}</span>
                                                                                )}
                                                                            </div>
                                                                        ))}
                                                                    </div>
                                                                </div>
                                                            )}

                                                            {/* Timing Schedule */}
                                                            {day.timing && day.timing.length > 0 && (
                                                                <div style={{ marginBottom: '20px', background: 'rgba(255,255,255,0.03)', padding: '15px', borderRadius: '10px' }}>
                                                                    <h4 style={{ fontSize: '0.9rem', color: '#60a5fa', marginBottom: '12px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="Clock" size={14} /> Lesson Timing</h4>
                                                                    <div style={{ display: 'flex', flexDirection: 'column', gap: '8px' }}>
                                                                        {day.timing.map((t, idx) => (
                                                                            <div key={idx} style={{ display: 'flex', alignItems: 'flex-start', gap: '12px', padding: '10px', background: 'rgba(0,0,0,0.2)', borderRadius: '8px' }}>
                                                                                <div style={{ minWidth: '70px', fontFamily: 'monospace', color: '#60a5fa', fontWeight: 600 }}>{t.minutes || t.duration}</div>
                                                                                <div style={{ flex: 1 }}>
                                                                                    <div style={{ fontWeight: 600, marginBottom: '4px' }}>{t.activity}</div>
                                                                                    <div style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.6)' }}>{t.description}</div>
                                                                                </div>
                                                                            </div>
                                                                        ))}
                                                                    </div>
                                                                </div>
                                                            )}

                                                            {/* Bell Ringer Detail */}
                                                            {day.bell_ringer && typeof day.bell_ringer === 'object' && (
                                                                <div style={{ marginBottom: '20px', background: 'rgba(165,180,252,0.1)', padding: '15px', borderRadius: '10px', border: '1px solid rgba(165,180,252,0.2)' }}>
                                                                    <h4 style={{ fontSize: '0.9rem', color: '#a5b4fc', marginBottom: '12px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="Zap" size={14} /> Bell Ringer</h4>
                                                                    <p style={{ fontWeight: 500, marginBottom: '10px', fontSize: '0.95rem' }}>{day.bell_ringer.prompt}</p>
                                                                    {day.bell_ringer.expected_responses && (
                                                                        <div style={{ marginTop: '10px' }}>
                                                                            <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '5px' }}>Expected Responses:</div>
                                                                            <ul style={{ margin: 0, paddingLeft: '20px', color: 'rgba(255,255,255,0.7)', fontSize: '0.85rem' }}>
                                                                                {day.bell_ringer.expected_responses.map((r, idx) => <li key={idx}>{r}</li>)}
                                                                            </ul>
                                                                        </div>
                                                                    )}
                                                                </div>
                                                            )}

                                                            {/* Simple Bell Ringer */}
                                                            {day.bell_ringer && typeof day.bell_ringer === 'string' && (
                                                                <div style={{ marginBottom: '20px', background: 'rgba(165,180,252,0.1)', padding: '15px', borderRadius: '10px', border: '1px solid rgba(165,180,252,0.2)' }}>
                                                                    <h4 style={{ fontSize: '0.9rem', color: '#a5b4fc', marginBottom: '10px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="Zap" size={14} /> Bell Ringer</h4>
                                                                    <p style={{ margin: 0, fontSize: '0.9rem' }}>{day.bell_ringer}</p>
                                                                </div>
                                                            )}

                                                            {/* Direct Instruction */}
                                                            {day.direct_instruction && (
                                                                <div style={{ marginBottom: '20px', background: 'rgba(251,191,36,0.1)', padding: '15px', borderRadius: '10px', border: '1px solid rgba(251,191,36,0.2)' }}>
                                                                    <h4 style={{ fontSize: '0.9rem', color: '#fbbf24', marginBottom: '12px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="Presentation" size={14} /> Direct Instruction</h4>
                                                                    {day.direct_instruction.key_points && (
                                                                        <div style={{ marginBottom: '12px' }}>
                                                                            <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '8px' }}>Key Points:</div>
                                                                            <ul style={{ margin: 0, paddingLeft: '20px' }}>
                                                                                {day.direct_instruction.key_points.map((p, idx) => (
                                                                                    <li key={idx} style={{ marginBottom: '6px', fontSize: '0.9rem', color: 'rgba(255,255,255,0.8)' }}>{p}</li>
                                                                                ))}
                                                                            </ul>
                                                                        </div>
                                                                    )}
                                                                    {day.direct_instruction.check_for_understanding && (
                                                                        <div>
                                                                            <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '8px' }}>Check for Understanding:</div>
                                                                            <ul style={{ margin: 0, paddingLeft: '20px' }}>
                                                                                {day.direct_instruction.check_for_understanding.map((q, idx) => (
                                                                                    <li key={idx} style={{ marginBottom: '4px', fontSize: '0.85rem', color: 'rgba(255,255,255,0.7)', fontStyle: 'italic' }}>"{q}"</li>
                                                                                ))}
                                                                            </ul>
                                                                        </div>
                                                                    )}
                                                                </div>
                                                            )}

                                                            {/* Activity Detail */}
                                                            {day.activity && typeof day.activity === 'object' && (
                                                                <div style={{ marginBottom: '20px', background: 'rgba(74,222,128,0.1)', padding: '15px', borderRadius: '10px', border: '1px solid rgba(74,222,128,0.2)' }}>
                                                                    <h4 style={{ fontSize: '0.9rem', color: '#4ade80', marginBottom: '12px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="Activity" size={14} /> {day.activity.name || 'Main Activity'}</h4>
                                                                    <p style={{ marginBottom: '12px', fontSize: '0.9rem' }}>{day.activity.description}</p>
                                                                    {day.activity.grouping && (
                                                                        <div style={{ display: 'inline-block', padding: '4px 10px', background: 'rgba(74,222,128,0.2)', borderRadius: '6px', fontSize: '0.8rem', marginBottom: '12px' }}>
                                                                            Grouping: {day.activity.grouping}
                                                                        </div>
                                                                    )}
                                                                    {day.activity.student_tasks && (
                                                                        <div style={{ marginTop: '10px' }}>
                                                                            <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '8px' }}>Student Tasks:</div>
                                                                            <ol style={{ margin: 0, paddingLeft: '20px' }}>
                                                                                {day.activity.student_tasks.map((t, idx) => (
                                                                                    <li key={idx} style={{ marginBottom: '4px', fontSize: '0.85rem', color: 'rgba(255,255,255,0.8)' }}>{t}</li>
                                                                                ))}
                                                                            </ol>
                                                                        </div>
                                                                    )}
                                                                    {day.activity.differentiation && (
                                                                        <div style={{ marginTop: '12px', display: 'grid', gridTemplateColumns: '1fr 1fr', gap: '10px' }}>
                                                                            {day.activity.differentiation.struggling && (
                                                                                <div style={{ background: 'rgba(0,0,0,0.2)', padding: '10px', borderRadius: '8px' }}>
                                                                                    <div style={{ fontSize: '0.75rem', color: '#f87171', marginBottom: '4px' }}>Support for Struggling:</div>
                                                                                    <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.7)' }}>{day.activity.differentiation.struggling}</div>
                                                                                </div>
                                                                            )}
                                                                            {day.activity.differentiation.advanced && (
                                                                                <div style={{ background: 'rgba(0,0,0,0.2)', padding: '10px', borderRadius: '8px' }}>
                                                                                    <div style={{ fontSize: '0.75rem', color: '#a78bfa', marginBottom: '4px' }}>Extension for Advanced:</div>
                                                                                    <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.7)' }}>{day.activity.differentiation.advanced}</div>
                                                                                </div>
                                                                            )}
                                                                        </div>
                                                                    )}
                                                                </div>
                                                            )}

                                                            {/* Simple Activity */}
                                                            {day.activity && typeof day.activity === 'string' && (
                                                                <div style={{ marginBottom: '20px', background: 'rgba(74,222,128,0.1)', padding: '15px', borderRadius: '10px', border: '1px solid rgba(74,222,128,0.2)' }}>
                                                                    <h4 style={{ fontSize: '0.9rem', color: '#4ade80', marginBottom: '10px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="Activity" size={14} /> Main Activity</h4>
                                                                    <p style={{ margin: 0, fontSize: '0.9rem' }}>{day.activity}</p>
                                                                </div>
                                                            )}

                                                            {/* Assessment */}
                                                            {day.assessment && (
                                                                <div style={{ marginBottom: '20px', background: 'rgba(248,113,113,0.1)', padding: '15px', borderRadius: '10px', border: '1px solid rgba(248,113,113,0.2)' }}>
                                                                    <h4 style={{ fontSize: '0.9rem', color: '#f87171', marginBottom: '10px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="CheckCircle" size={14} /> Assessment</h4>
                                                                    {typeof day.assessment === 'object' ? (
                                                                        <>
                                                                            {day.assessment.type && <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '8px' }}>Type: {day.assessment.type}</div>}
                                                                            <p style={{ margin: '0 0 10px 0', fontSize: '0.9rem' }}>{day.assessment.description}</p>
                                                                            {day.assessment.exit_ticket && (
                                                                                <div style={{ background: 'rgba(0,0,0,0.2)', padding: '10px', borderRadius: '8px', marginTop: '10px' }}>
                                                                                    <div style={{ fontSize: '0.75rem', color: 'rgba(255,255,255,0.5)', marginBottom: '4px' }}>Exit Ticket:</div>
                                                                                    <div style={{ fontSize: '0.85rem', fontStyle: 'italic' }}>"{day.assessment.exit_ticket}"</div>
                                                                                </div>
                                                                            )}
                                                                        </>
                                                                    ) : (
                                                                        <p style={{ margin: 0, fontSize: '0.9rem' }}>{day.assessment}</p>
                                                                    )}
                                                                </div>
                                                            )}

                                                            {/* Bottom Row: Materials, Homework, Teacher Notes */}
                                                            <div style={{ display: 'grid', gridTemplateColumns: 'repeat(auto-fit, minmax(200px, 1fr))', gap: '15px' }}>
                                                                {day.materials && day.materials.length > 0 && (
                                                                    <div style={{ background: 'rgba(255,255,255,0.03)', padding: '12px', borderRadius: '10px' }}>
                                                                        <h4 style={{ fontSize: '0.85rem', color: '#94a3b8', marginBottom: '8px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="Box" size={14} /> Materials</h4>
                                                                        <div style={{ display: 'flex', flexWrap: 'wrap', gap: '6px' }}>
                                                                            {day.materials.map((m, idx) => (
                                                                                <span key={idx} style={{ fontSize: '0.8rem', padding: '3px 8px', borderRadius: '4px', background: 'rgba(255,255,255,0.1)' }}>{m}</span>
                                                                            ))}
                                                                        </div>
                                                                    </div>
                                                                )}
                                                                {day.homework && (
                                                                    <div style={{ background: 'rgba(255,255,255,0.03)', padding: '12px', borderRadius: '10px' }}>
                                                                        <h4 style={{ fontSize: '0.85rem', color: '#94a3b8', marginBottom: '8px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="Home" size={14} /> Homework</h4>
                                                                        <p style={{ margin: 0, fontSize: '0.85rem', color: 'rgba(255,255,255,0.7)' }}>{day.homework}</p>
                                                                    </div>
                                                                )}
                                                                {day.teacher_notes && (
                                                                    <div style={{ background: 'rgba(255,255,255,0.03)', padding: '12px', borderRadius: '10px' }}>
                                                                        <h4 style={{ fontSize: '0.85rem', color: '#94a3b8', marginBottom: '8px', display: 'flex', alignItems: 'center', gap: '6px' }}><Icon name="StickyNote" size={14} /> Teacher Notes</h4>
                                                                        <p style={{ margin: 0, fontSize: '0.85rem', color: 'rgba(255,255,255,0.7)' }}>{day.teacher_notes}</p>
                                                                    </div>
                                                                )}
                                                            </div>
                                                        </div>
                                                    ))}
                                                </div>

                                                {/* Unit Assessment */}
                                                {lessonPlan.unit_assessment && (
                                                    <div style={{ marginTop: '30px', background: 'rgba(99,102,241,0.1)', border: '1px solid rgba(99,102,241,0.3)', borderRadius: '16px', padding: '20px' }}>
                                                        <h3 style={{ fontSize: '1.1rem', color: '#a5b4fc', marginBottom: '15px', display: 'flex', alignItems: 'center', gap: '8px' }}><Icon name="Award" size={18} /> Summative Assessment</h3>
                                                        {typeof lessonPlan.unit_assessment === 'object' ? (
                                                            <>
                                                                {lessonPlan.unit_assessment.type && <div style={{ fontSize: '0.85rem', color: 'rgba(255,255,255,0.5)', marginBottom: '10px' }}>Type: {lessonPlan.unit_assessment.type}</div>}
                                                                <p style={{ marginBottom: '15px', fontSize: '0.95rem' }}>{lessonPlan.unit_assessment.description}</p>
                                                                {lessonPlan.unit_assessment.components && (
                                                                    <div style={{ marginTop: '10px' }}>
                                                                        <div style={{ fontSize: '0.8rem', color: 'rgba(255,255,255,0.5)', marginBottom: '8px' }}>Components:</div>
                                                                        <ul style={{ margin: 0, paddingLeft: '20px' }}>
                                                                            {lessonPlan.unit_assessment.components.map((c, idx) => (
                                                                                <li key={idx} style={{ marginBottom: '4px', fontSize: '0.85rem' }}>{c}</li>
                                                                            ))}
                                                                        </ul>
                                                                    </div>
                                                                )}
                                                            </>
                                                        ) : (
                                                            <p style={{ margin: 0, fontSize: '0.95rem' }}>{lessonPlan.unit_assessment}</p>
                                                        )}
                                                    </div>
                                                )}

                                                {/* Resources */}
                                                {lessonPlan.resources && lessonPlan.resources.length > 0 && (
                                                    <div style={{ marginTop: '20px', padding: '15px', background: 'rgba(255,255,255,0.03)', borderRadius: '12px' }}>
                                                        <h4 style={{ fontSize: '0.9rem', color: '#94a3b8', marginBottom: '10px' }}>Resources</h4>
                                                        <div style={{ display: 'flex', flexWrap: 'wrap', gap: '8px' }}>
                                                            {lessonPlan.resources.map((r, idx) => (
                                                                <span key={idx} style={{ fontSize: '0.8rem', padding: '4px 10px', borderRadius: '6px', background: 'rgba(255,255,255,0.1)' }}>{r}</span>
                                                            ))}
                                                        </div>
                                                    </div>
                                                )}
                                            </div>
                                        ) : (
                                        /* Standards List */
                                        <div>
                                            <div style={{ display: 'flex', justifyContent: 'space-between', alignItems: 'center', marginBottom: '15px' }}>
                                                <h3 style={{ fontSize: '1.1rem', fontWeight: 700, display: 'flex', alignItems: 'center', gap: '10px' }}>
                                                    <Icon name="Library" size={20} /> Select Standards ({selectedStandards.length})
                                                </h3>
                                                <div style={{ fontSize: '0.9rem', color: 'rgba(255,255,255,0.5)' }}>
                                                    {standards.length} standards available
                                                </div>
                                            </div>
                                            
                                            <div style={{ maxHeight: '400px', overflowY: 'auto', paddingRight: '5px' }}>
                                                {plannerLoading ? (
                                                    <div style={{ textAlign: 'center', padding: '40px', color: 'rgba(255,255,255,0.5)' }}>
                                                        <Icon name="Loader2" size={30} style={{ animation: 'spin 1s linear infinite' }} />
                                                        <p style={{ marginTop: '10px' }}>Loading standards...</p>
                                                    </div>
                                                ) : standards.length > 0 ? (
                                                    standards.map(std => (
                                                        <StandardCard
                                                            key={std.code}
                                                            standard={std}
                                                            isSelected={selectedStandards.includes(std.code)}
                                                            onToggle={() => toggleStandard(std.code)}
                                                        />
                                                    ))
                                                ) : (
                                                    <div style={{ textAlign: 'center', padding: '40px', background: 'rgba(255,255,255,0.03)', borderRadius: '12px' }}>
                                                        <p style={{ color: 'rgba(255,255,255,0.5)' }}>No standards found for this configuration.</p>
                                                    </div>
                                                )}
                                            </div>
                                        </div>
                                        )}
                                    </div>
                                </div>
                            </div>
                        )}


                        <footer style={{ textAlign: 'center', marginTop: '30px', padding: '20px', color: 'rgba(255,255,255,0.3)', fontSize: '0.85rem' }}>
                            Powered by OpenAI GPT-4o • Built with ❤️ for education
                        </footer>
                    </div>
                </div>
            );
        };

        ReactDOM.createRoot(document.getElementById('root')).render(<App />);
    </script>
</body>
</html>'''

# ══════════════════════════════════════════════════════════════
# ROUTES
# ══════════════════════════════════════════════════════════════

@app.route('/')
def index():
    return Response(HTML_TEMPLATE, mimetype='text/html')

@app.route('/api/status')
def get_status():
    return jsonify(grading_state)

@app.route('/api/browse')
def browse_for_path():
    """Open a file/folder picker dialog."""
    import subprocess
    
    browse_type = request.args.get('type', 'folder')
    
    try:
        if browse_type == 'folder':
            # Use osascript to open folder picker on Mac
            script = '''
            tell application "System Events"
                activate
                set folderPath to POSIX path of (choose folder with prompt "Select Folder")
                return folderPath
            end tell
            '''
        else:
            # File picker
            script = '''
            tell application "System Events"
                activate
                set filePath to POSIX path of (choose file with prompt "Select File")
                return filePath
            end tell
            '''
        
        result = subprocess.run(['osascript', '-e', script], capture_output=True, text=True, timeout=60)
        
        if result.returncode == 0 and result.stdout.strip():
            path = result.stdout.strip()
            return jsonify({"path": path})
        else:
            return jsonify({"path": None, "error": "Cancelled or no selection"})
    
    except subprocess.TimeoutExpired:
        return jsonify({"path": None, "error": "Timeout"})
    except Exception as e:
        return jsonify({"path": None, "error": str(e)})

@app.route('/api/parse-document', methods=['POST'])
def parse_document():
    """Parse an uploaded Word/PDF document and convert to HTML for full formatting."""
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"})
    
    file = request.files['file']
    filename = file.filename.lower()
    file_data = file.read()
    
    try:
        if filename.endswith('.docx'):
            # Try mammoth first for best HTML conversion
            try:
                import mammoth
                import io
                
                result = mammoth.convert_to_html(io.BytesIO(file_data))
                html = result.value
                
                # Add some basic styling
                styled_html = f'''
                <style>
                    body {{ font-family: Georgia, serif; line-height: 1.6; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
                    td, th {{ border: 1px solid #ccc; padding: 8px 12px; text-align: left; }}
                    th {{ background: #f5f5f5; font-weight: bold; }}
                    p {{ margin: 10px 0; }}
                    h1, h2, h3 {{ margin: 20px 0 10px 0; }}
                    ul, ol {{ margin: 10px 0; padding-left: 25px; }}
                </style>
                {html}
                '''
                
                # Also extract plain text for marking
                from docx import Document
                doc = Document(io.BytesIO(file_data))
                plain_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        plain_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            plain_text.append(row_text)
                
                return jsonify({
                    "html": styled_html,
                    "text": '\n'.join(plain_text),
                    "filename": file.filename,
                    "type": "html"
                })
                
            except ImportError:
                # Fallback to python-docx only
                from docx import Document
                import io
                
                doc = Document(io.BytesIO(file_data))
                
                html_parts = ['<div style="font-family: Georgia, serif; line-height: 1.6;">']
                
                from docx.oxml.table import CT_Tbl
                from docx.oxml.text.paragraph import CT_P
                from docx.table import Table
                from docx.text.paragraph import Paragraph
                
                for child in doc.element.body.iterchildren():
                    if isinstance(child, CT_P):
                        para = Paragraph(child, doc)
                        if para.text.strip():
                            style = para.style.name if para.style else ''
                            if 'Heading 1' in style:
                                html_parts.append(f'<h1>{para.text}</h1>')
                            elif 'Heading 2' in style:
                                html_parts.append(f'<h2>{para.text}</h2>')
                            elif 'Heading' in style:
                                html_parts.append(f'<h3>{para.text}</h3>')
                            else:
                                html_parts.append(f'<p>{para.text}</p>')
                    elif isinstance(child, CT_Tbl):
                        table = Table(child, doc)
                        html_parts.append('<table style="border-collapse: collapse; width: 100%; margin: 15px 0;">')
                        for row_idx, row in enumerate(table.rows):
                            html_parts.append('<tr>')
                            for cell in row.cells:
                                tag = 'th' if row_idx == 0 else 'td'
                                style = 'border: 1px solid #ccc; padding: 8px 12px; background: #f5f5f5;' if row_idx == 0 else 'border: 1px solid #ccc; padding: 8px 12px;'
                                html_parts.append(f'<{tag} style="{style}">{cell.text}</{tag}>')
                            html_parts.append('</tr>')
                        html_parts.append('</table>')
                
                html_parts.append('</div>')
                
                # Plain text version
                plain_text = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        plain_text.append(para.text)
                for table in doc.tables:
                    for row in table.rows:
                        row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            plain_text.append(row_text)
                
                return jsonify({
                    "html": ''.join(html_parts),
                    "text": '\n'.join(plain_text),
                    "filename": file.filename,
                    "type": "html"
                })
        
        elif filename.endswith('.pdf'):
            try:
                import fitz  # PyMuPDF
                import base64
                import io
                
                doc = fitz.open(stream=file_data, filetype="pdf")
                
                # Convert each page to an image for perfect rendering
                images_html = []
                plain_text = []
                
                for page_num, page in enumerate(doc):
                    # Render page to image at 150 DPI for good quality
                    mat = fitz.Matrix(1.5, 1.5)  # 150 DPI
                    pix = page.get_pixmap(matrix=mat)
                    img_data = pix.tobytes("png")
                    img_base64 = base64.b64encode(img_data).decode('utf-8')
                    
                    images_html.append(f'''
                        <div style="margin-bottom: 20px; box-shadow: 0 2px 10px rgba(0,0,0,0.1);">
                            <div style="background: #6366f1; color: white; padding: 5px 15px; font-size: 12px;">Page {page_num + 1}</div>
                            <img src="data:image/png;base64,{img_base64}" style="width: 100%; display: block;" />
                        </div>
                    ''')
                    
                    # Also get text for marking
                    plain_text.append(page.get_text())
                
                doc.close()
                
                full_html = f'''
                <div style="background: #e5e5e5; padding: 20px;">
                    {''.join(images_html)}
                </div>
                '''
                
                return jsonify({
                    "html": full_html,
                    "text": '\n\n'.join(plain_text),
                    "filename": file.filename,
                    "type": "html"
                })
                
            except ImportError:
                return jsonify({"error": "PDF support requires PyMuPDF. Run: pip3 install pymupdf"})
        
        elif filename.endswith('.txt'):
            text = file_data.decode('utf-8', errors='ignore')
            html = f'<pre style="font-family: Monaco, monospace; white-space: pre-wrap; line-height: 1.6;">{text}</pre>'
            return jsonify({
                "html": html,
                "text": text,
                "filename": file.filename,
                "type": "html"
            })
        
        else:
            return jsonify({"error": "Unsupported file type. Use .docx, .pdf, or .txt"})
    
    except Exception as e:
        import traceback
        return jsonify({"error": f"{str(e)}\n{traceback.format_exc()}"})

@app.route('/api/check-new-files', methods=['POST'])
def check_new_files():
    """Check for new files that haven't been graded yet."""
    import csv
    
    data = request.json
    assignments_folder = data.get('folder', '/Users/alexc/Library/CloudStorage/OneDrive-VolusiaCountySchools/Assignments')
    output_folder = data.get('output_folder', '/Users/alexc/Downloads/Graider/Results')
    
    if not os.path.exists(assignments_folder):
        return jsonify({"error": "Folder not found", "new_files": 0})
    
    # Load already graded files from master CSV
    already_graded = set()
    master_file = os.path.join(output_folder, "master_grades.csv")
    if os.path.exists(master_file):
        try:
            with open(master_file, 'r', encoding='utf-8') as f:
                reader = csv.DictReader(f)
                for row in reader:
                    filename = row.get('Filename', '')
                    if filename:
                        already_graded.add(filename)
        except:
            pass
    
    # Count new files
    all_files = []
    for ext in ['*.docx', '*.txt', '*.jpg', '*.jpeg', '*.png']:
        all_files.extend(Path(assignments_folder).glob(ext))
    
    new_files = [f for f in all_files if f.name not in already_graded]
    
    return jsonify({
        "total_files": len(all_files),
        "already_graded": len(already_graded),
        "new_files": len(new_files),
        "new_file_names": [f.name for f in new_files[:5]]  # Return first 5 names
    })

@app.route('/api/stop-grading', methods=['POST'])
def stop_grading():
    """Stop grading and save progress."""
    global grading_state
    
    if grading_state["is_running"]:
        grading_state["stop_requested"] = True
        grading_state["log"].append("")
        grading_state["log"].append("⏹️  Stop requested... saving progress...")
        return jsonify({"stopped": True, "message": "Stop requested, saving progress..."})
    
    return jsonify({"stopped": False, "message": "Grading not running"})

@app.route('/api/grade', methods=['POST'])
def start_grading():
    global grading_state

    if grading_state["is_running"]:
        return jsonify({"error": "Grading already in progress"}), 400

    data = request.json
    assignments_folder = data.get('assignments_folder', '/Users/alexc/Library/CloudStorage/OneDrive-VolusiaCountySchools/Assignments')
    output_folder = data.get('output_folder', '/Users/alexc/Downloads/Graider/Results')
    roster_file = data.get('roster_file', '/Users/alexc/Downloads/Graider/all_students_updated.xlsx')
    grading_period = data.get('grading_period', 'Q3')

    # Get custom assignment config and global AI notes
    assignment_config = data.get('assignmentConfig')  # Contains customMarkers, gradingNotes
    global_ai_notes = data.get('globalAINotes', '')

    if not os.path.exists(assignments_folder):
        return jsonify({"error": f"Assignments folder not found: {assignments_folder}"}), 400
    if not os.path.exists(roster_file):
        return jsonify({"error": f"Roster file not found: {roster_file}"}), 400

    reset_state()
    grading_state["is_running"] = True

    thread = threading.Thread(target=run_grading_thread, args=(assignments_folder, output_folder, roster_file, assignment_config, global_ai_notes, grading_period))
    thread.start()

    return jsonify({"status": "started"})

def run_grading_thread(assignments_folder, output_folder, roster_file, assignment_config=None, global_ai_notes='', grading_period='Q3'):
    global grading_state

    # Load ALL saved assignment configs for auto-matching
    all_configs = {}
    assignments_dir = os.path.expanduser("~/.graider_assignments")
    if os.path.exists(assignments_dir):
        for f in os.listdir(assignments_dir):
            if f.endswith('.json'):
                config_name = f.replace('.json', '')
                try:
                    with open(os.path.join(assignments_dir, f), 'r') as cf:
                        all_configs[config_name.lower()] = json.load(cf)
                except:
                    pass

    # Function to find matching config for a filename
    def find_matching_config(filename):
        import re
        filename_lower = filename.lower()
        # Remove student name prefix (everything before " - " or "_")
        if ' - ' in filename_lower:
            assignment_part = filename_lower.split(' - ', 1)[1]
        elif '_' in filename_lower:
            parts = filename_lower.split('_')
            assignment_part = '_'.join(parts[2:]) if len(parts) > 2 else filename_lower
        else:
            assignment_part = filename_lower

        # Remove extension
        assignment_part = os.path.splitext(assignment_part)[0]

        # Normalize for comparison: replace & with "and", remove special chars
        def normalize(s):
            s = s.replace('&', 'and').replace('–', '-').replace('—', '-')
            s = re.sub(r'[^\w\s]', ' ', s)  # Remove special chars except spaces
            s = re.sub(r'\s+', ' ', s).strip()  # Normalize whitespace
            return s

        assignment_normalized = normalize(assignment_part)
        assignment_words = set(assignment_normalized.split())

        # Try to match against saved config names
        best_match = None
        best_score = 0
        for config_name, config_data in all_configs.items():
            config_normalized = normalize(config_name)
            config_words = set(config_normalized.split())

            # Check if config name appears in assignment part or vice versa
            if config_normalized in assignment_normalized or assignment_normalized in config_normalized:
                score = len(config_normalized)
                if score > best_score:
                    best_score = score
                    best_match = config_data

            # Word overlap scoring - count matching words
            common_words = assignment_words & config_words
            # Filter out common words
            common_words -= {'the', 'a', 'an', 'and', 'or', 'of', 'in', 'to', 'for', 'on', 'pp', 'notes', 'cornell'}
            if len(common_words) >= 2:  # At least 2 meaningful words match
                word_score = len(common_words) * 10  # Weight word matches heavily
                if word_score > best_score:
                    best_score = word_score
                    best_match = config_data

            # Also check the title field inside the config
            config_title = config_data.get('title', '').lower()
            if config_title:
                title_normalized = normalize(config_title)
                title_words = set(title_normalized.split())
                if title_normalized in assignment_normalized or assignment_normalized in title_normalized:
                    score = len(title_normalized)
                    if score > best_score:
                        best_score = score
                        best_match = config_data
                # Word overlap with title
                common_title_words = assignment_words & title_words
                common_title_words -= {'the', 'a', 'an', 'and', 'or', 'of', 'in', 'to', 'for', 'on', 'pp', 'notes', 'cornell'}
                if len(common_title_words) >= 2:
                    word_score = len(common_title_words) * 10
                    if word_score > best_score:
                        best_score = word_score
                        best_match = config_data

            # Check aliases
            aliases = config_data.get('aliases', [])
            for alias in aliases:
                alias_normalized = normalize(alias.lower())
                if alias_normalized in assignment_normalized or assignment_normalized in alias_normalized:
                    score = len(alias_normalized) + 5  # Bonus for alias match
                    if score > best_score:
                        best_score = score
                        best_match = config_data

        return best_match

    # Extract custom markers and notes from selected config (fallback)
    fallback_markers = []
    fallback_notes = ''
    if assignment_config:
        fallback_markers = assignment_config.get('customMarkers', [])
        fallback_notes = assignment_config.get('gradingNotes', '')

    try:
        from assignment_grader import (
            load_roster, parse_filename, read_assignment_file,
            extract_student_work, grade_assignment, export_focus_csv,
            export_detailed_report, save_emails_to_folder, save_to_master_csv,
            ASSIGNMENT_NAME, STUDENT_WORK_MARKERS
        )
        import csv

        if all_configs:
            grading_state["log"].append(f"📚 Loaded {len(all_configs)} assignment configs for auto-matching")

        if global_ai_notes:
            grading_state["log"].append(f"📝 Global AI notes loaded")

        os.makedirs(output_folder, exist_ok=True)
        
        # Load already graded files from master CSV
        already_graded = set()
        master_file = os.path.join(output_folder, "master_grades.csv")
        if os.path.exists(master_file):
            try:
                with open(master_file, 'r', encoding='utf-8') as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        filename = row.get('Filename', '')
                        if filename:
                            already_graded.add(filename)
                grading_state["log"].append(f"📚 Found {len(already_graded)} previously graded files")
            except:
                pass
        
        grading_state["log"].append("📋 Loading student roster...")
        roster = load_roster(roster_file)
        grading_state["log"].append(f"✅ Loaded {len(roster)//2} students")
        
        assignment_path = Path(assignments_folder)
        all_files = []
        for ext in ['*.docx', '*.txt', '*.jpg', '*.jpeg', '*.png']:
            all_files.extend(assignment_path.glob(ext))
        
        # Filter out already graded files
        new_files = [f for f in all_files if f.name not in already_graded]
        skipped = len(all_files) - len(new_files)
        
        if skipped > 0:
            grading_state["log"].append(f"⏭️  Skipping {skipped} already-graded files")
        
        grading_state["total"] = len(new_files)
        grading_state["log"].append(f"📄 Found {len(new_files)} NEW files to grade")
        
        if len(new_files) == 0:
            grading_state["log"].append("")
            grading_state["log"].append("✨ All files have already been graded!")
            grading_state["complete"] = True
            grading_state["is_running"] = False
            return
        
        all_grades = []
        
        for i, filepath in enumerate(new_files, 1):
            # Check if stop was requested
            if grading_state.get("stop_requested", False):
                grading_state["log"].append("")
                grading_state["log"].append(f"⏹️  Stopped at {i-1}/{len(new_files)} files")
                grading_state["log"].append(f"✅ Progress saved! {len(all_grades)} grades completed.")
                break
            
            grading_state["progress"] = i
            grading_state["current_file"] = filepath.name
            
            parsed = parse_filename(filepath.name)
            student_name = f"{parsed['first_name']} {parsed['last_name']}"
            lookup_key = parsed['lookup_key']
            
            if lookup_key in roster:
                student_info = roster[lookup_key].copy()
            else:
                student_info = {"student_id": "UNKNOWN", "student_name": student_name, "first_name": parsed['first_name'], "last_name": parsed['last_name'], "email": ""}
            
            grading_state["log"].append(f"[{i}/{len(new_files)}] {student_info['student_name']}")

            # Try to auto-match assignment config based on filename
            matched_config = find_matching_config(filepath.name)
            if matched_config:
                file_markers = matched_config.get('customMarkers', [])
                file_notes = matched_config.get('gradingNotes', '')
                matched_title = matched_config.get('title', 'Unknown')
                grading_state["log"].append(f"  📎 Matched config: {matched_title}")
            else:
                file_markers = fallback_markers
                file_notes = fallback_notes

            # Build combined AI notes for this file
            file_ai_notes = ''
            if global_ai_notes:
                file_ai_notes += f"GLOBAL GRADING INSTRUCTIONS:\n{global_ai_notes}\n\n"
            if file_notes:
                file_ai_notes += f"ASSIGNMENT-SPECIFIC INSTRUCTIONS:\n{file_notes}"

            # Add file-specific markers to the markers list temporarily
            original_markers = STUDENT_WORK_MARKERS.copy()
            for marker in file_markers:
                if marker not in STUDENT_WORK_MARKERS:
                    STUDENT_WORK_MARKERS.append(marker)

            file_data = read_assignment_file(filepath)
            if not file_data:
                grading_state["log"].append(f"  ❌ Could not read file")
                STUDENT_WORK_MARKERS.clear()
                STUDENT_WORK_MARKERS.extend(original_markers)
                continue

            markers_found = []
            if file_data["type"] == "text":
                student_work, markers_found = extract_student_work(file_data["content"])
                if markers_found:
                    grading_state["log"].append(f"  📝 Markers: {', '.join(markers_found[:2])}")
                grade_data = {"type": "text", "content": student_work}
            else:
                grading_state["log"].append(f"  🖼️ Image file")
                grade_data = file_data

            grading_state["log"].append(f"  🤖 Grading...")
            grade_result = grade_assignment(student_info['student_name'], grade_data, file_ai_notes)

            # Restore original markers
            STUDENT_WORK_MARKERS.clear()
            STUDENT_WORK_MARKERS.extend(original_markers)
            grading_state["log"].append(f"  ✅ Score: {grade_result['score']} ({grade_result['letter_grade']})")
            
            # Extract assignment name from filename
            parts = Path(filepath.name).stem.split('_')
            if len(parts) >= 3:
                assignment_from_file = ' '.join(parts[2:])
            else:
                assignment_from_file = ASSIGNMENT_NAME
            
            # Get student content for review
            if file_data["type"] == "text":
                student_content = student_work if student_work else file_data.get("content", "")
            else:
                student_content = "[Image file - view in original document]"
            
            # Build full grade record for export
            grade_record = {
                **student_info,
                **grade_result,
                "filename": filepath.name,
                "assignment": assignment_from_file,
                "grading_period": grading_period,
                "has_markers": len(markers_found) > 0
            }
            all_grades.append(grade_record)
            
            grading_state["results"].append({
                "student_name": student_info['student_name'],
                "student_id": student_info['student_id'],
                "email": student_info.get('email', ''),
                "filename": filepath.name,
                "filepath": str(filepath),
                "assignment": assignment_from_file,
                "score": grade_result['score'],
                "letter_grade": grade_result['letter_grade'],
                "feedback": grade_result.get('feedback', ''),
                "student_content": student_content[:5000],  # Limit size for UI
                "breakdown": grade_result.get('breakdown', {})
            })
        
        # Export CSVs and emails
        if len(all_grades) > 0:
            grading_state["log"].append("")
            grading_state["log"].append("📊 Exporting results...")
            
            # Focus CSVs (by assignment)
            export_focus_csv(all_grades, output_folder, ASSIGNMENT_NAME)
            grading_state["log"].append("  ✅ Focus CSVs created")
            
            # Detailed report
            export_detailed_report(all_grades, output_folder, ASSIGNMENT_NAME)
            grading_state["log"].append("  ✅ Detailed report created")
            
            # Email files
            save_emails_to_folder(all_grades, output_folder)
            grading_state["log"].append("  ✅ Email files created")
            
            # Master tracking CSV
            save_to_master_csv(all_grades, output_folder)
            grading_state["log"].append("  ✅ Master grades updated")
        
        grading_state["log"].append("")
        grading_state["log"].append("═" * 50)
        
        if grading_state.get("stop_requested", False):
            grading_state["log"].append(f"⏹️  GRADING STOPPED - {len(all_grades)} files saved")
            grading_state["log"].append("💡 Restart to continue with remaining files")
        else:
            grading_state["log"].append("✅ GRADING COMPLETE!")
        
        grading_state["log"].append(f"📁 Results saved to: {output_folder}")
        grading_state["complete"] = True
        
    except Exception as e:
        grading_state["error"] = str(e)
        grading_state["log"].append(f"❌ Error: {str(e)}")
    finally:
        grading_state["is_running"] = False
        grading_state["stop_requested"] = False

@app.route('/api/open-folder', methods=['POST'])
def open_folder():
    data = request.json
    folder = data.get('folder', '/Users/alexc/Downloads/Graider/Results')
    if os.path.exists(folder):
        os.system(f'open "{folder}"')
        return jsonify({"status": "opened"})
    os.makedirs(folder, exist_ok=True)
    os.system(f'open "{folder}"')
    return jsonify({"status": "created and opened"})

@app.route('/api/send-emails', methods=['POST'])
def send_emails():
    """Send grade emails to students via Resend."""
    try:
        import resend
        from collections import defaultdict

        # Get Resend API key from environment
        resend_api_key = os.getenv('RESEND_API_KEY')
        if not resend_api_key:
            return jsonify({"error": "RESEND_API_KEY not found in .env file"})

        resend.api_key = resend_api_key
        from_email = os.getenv('RESEND_FROM_EMAIL', 'Graider <noreply@graider.live>')

        data = request.json
        results = data.get('results', [])
        teacher_name = data.get('teacher_name', 'Your Teacher')
        teacher_email = data.get('teacher_email', '')  # For reply-to
        email_signature = data.get('email_signature', '')

        if not results:
            return jsonify({"error": "No results to email"})

        # Group by student email for combined emails
        students = defaultdict(list)

        for r in results:
            email = r.get('student_email') or r.get('email', '')
            if email and '@' in email and r.get('student_id') != 'UNKNOWN':
                students[email].append(r)

        sent = 0
        failed = 0

        for email, grades in students.items():
            first_name = grades[0].get('student_name', 'Student').split()[0]

            # Build subject
            if len(grades) == 1:
                assignment = grades[0].get('assignment', 'Assignment')
                subject = f"Grade for {assignment}: {grades[0].get('letter_grade', '')}"
            else:
                subject = f"Grades for {len(grades)} Assignments"

            # Build body
            body = f"Hi {first_name},\n\n"

            if len(grades) == 1:
                g = grades[0]
                # Use custom email body if provided
                if g.get('custom_email_body'):
                    body = g.get('custom_email_body')
                else:
                    body += f"Here is your grade and feedback for {g.get('assignment', 'your assignment')}:\n\n"
                    body += f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                    body += f"GRADE: {g.get('score', 0)}/100 ({g.get('letter_grade', '')})\n"
                    body += f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n\n"
                    body += f"FEEDBACK:\n{g.get('feedback', 'No feedback available.')}\n"
            else:
                body += "Here are your grades and feedback:\n\n"
                for g in grades:
                    body += f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                    body += f"📝 {g.get('assignment', 'Assignment')}\n"
                    body += f"GRADE: {g.get('score', 0)}/100 ({g.get('letter_grade', '')})\n\n"
                    body += f"FEEDBACK:\n{g.get('feedback', 'No feedback available.')}\n\n"

            body += f"\n━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            body += f"\nIf you have any questions, please see me during class.\n\n{teacher_name}"

            if email_signature:
                body += f"\n\n{email_signature}"

            # Save email to file for logging
            safe_name = grades[0].get('student_name', 'Unknown').replace(' ', '_').replace('/', '_')
            email_log_dir = os.path.join(app_dir, 'Results', 'emails')
            os.makedirs(email_log_dir, exist_ok=True)
            email_log_path = os.path.join(email_log_dir, f"{safe_name}_email.txt")
            with open(email_log_path, 'w') as f:
                f.write(f"TO: {email}\n")
                f.write(f"SUBJECT: {subject}\n")
                f.write("=" * 50 + "\n\n")
                f.write(body)

            # Send via Resend
            try:
                params = {
                    "from": from_email,
                    "to": [email],
                    "subject": subject,
                    "text": body,
                }
                if teacher_email:
                    params["reply_to"] = teacher_email

                response = resend.Emails.send(params)

                if response and response.get('id'):
                    print(f"  ✅ Sent to {first_name} ({email}) - ID: {response.get('id')}")
                    sent += 1
                else:
                    print(f"  ❌ Failed to send to {email}: No response ID")
                    failed += 1
            except Exception as e:
                print(f"  ❌ Failed to send to {email}: {e}")
                failed += 1

        return jsonify({"sent": sent, "failed": failed, "total": len(students)})

    except ImportError:
        return jsonify({"error": "resend package not installed. Run: pip install resend"})
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route('/api/save-rubric', methods=['POST'])
def save_rubric():
    """Save rubric configuration to JSON file."""
    import json

    data = request.json
    rubric_path = os.path.expanduser("~/.graider_rubric.json")

    try:
        with open(rubric_path, 'w') as f:
            json.dump(data, f, indent=2)
        return jsonify({"status": "saved"})
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route('/api/load-rubric')
def load_rubric():
    """Load rubric configuration from JSON file."""
    import json

    rubric_path = os.path.expanduser("~/.graider_rubric.json")

    if not os.path.exists(rubric_path):
        return jsonify({"rubric": None})

    try:
        with open(rubric_path, 'r') as f:
            data = json.load(f)
        return jsonify({"rubric": data})
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route('/api/save-global-settings', methods=['POST'])
def save_global_settings():
    """Save global AI notes and settings."""
    import json

    data = request.json
    settings_path = os.path.expanduser("~/.graider_settings.json")

    try:
        with open(settings_path, 'w') as f:
            json.dump(data, f, indent=2)
        return jsonify({"status": "saved"})
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route('/api/load-global-settings')
def load_global_settings():
    """Load global AI notes and settings."""
    import json

    settings_path = os.path.expanduser("~/.graider_settings.json")

    if not os.path.exists(settings_path):
        return jsonify({"settings": None})

    try:
        with open(settings_path, 'r') as f:
            data = json.load(f)
        return jsonify({"settings": data})
    except Exception as e:
        return jsonify({"error": str(e)})

# =============================================================================
# ASSIGNMENT BUILDER API
# =============================================================================

@app.route('/api/save-assignment-config', methods=['POST'])
def save_assignment_config():
    """Save assignment configuration for grading."""
    import json
    
    data = request.json
    assignments_dir = os.path.expanduser("~/.graider_assignments")
    os.makedirs(assignments_dir, exist_ok=True)
    
    # Clean title for filename
    title = data.get('title', 'Untitled')
    safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()
    filepath = os.path.join(assignments_dir, f"{safe_title}.json")
    
    try:
        with open(filepath, 'w') as f:
            json.dump(data, f, indent=2)
        return jsonify({"status": "saved", "path": filepath})
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route('/api/list-assignments')
def list_assignments():
    """List saved assignment configurations."""
    assignments_dir = os.path.expanduser("~/.graider_assignments")
    
    if not os.path.exists(assignments_dir):
        return jsonify({"assignments": []})
    
    assignments = []
    for f in os.listdir(assignments_dir):
        if f.endswith('.json'):
            assignments.append(f.replace('.json', ''))
    
    return jsonify({"assignments": sorted(assignments)})

@app.route('/api/load-assignment')
def load_assignment():
    """Load a saved assignment configuration."""
    import json

    name = request.args.get('name', '')
    assignments_dir = os.path.expanduser("~/.graider_assignments")
    filepath = os.path.join(assignments_dir, f"{name}.json")

    if not os.path.exists(filepath):
        return jsonify({"error": "Assignment not found"})

    try:
        with open(filepath, 'r') as f:
            data = json.load(f)
        return jsonify({"assignment": data})
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route('/api/delete-assignment', methods=['DELETE'])
def delete_assignment():
    """Delete a saved assignment configuration."""
    name = request.args.get('name', '')
    assignments_dir = os.path.expanduser("~/.graider_assignments")
    filepath = os.path.join(assignments_dir, f"{name}.json")

    if not os.path.exists(filepath):
        return jsonify({"error": "Assignment not found"})

    try:
        os.remove(filepath)
        return jsonify({"status": "deleted"})
    except Exception as e:
        return jsonify({"error": str(e)})

@app.route('/api/export-assignment', methods=['POST'])
def export_assignment():
    """Export assignment to Word or PDF format."""
    data = request.json
    assignment = data.get('assignment', {})
    format_type = data.get('format', 'docx')
    
    title = assignment.get('title', 'Untitled Assignment')
    instructions = assignment.get('instructions', '')
    questions = assignment.get('questions', [])
    
    output_folder = "/Users/alexc/Downloads/Graider/Assignments"
    os.makedirs(output_folder, exist_ok=True)
    
    safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()
    
    if format_type == 'docx':
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Name/Date line
            name_para = doc.add_paragraph("Name: _________________________ Date: _____________")
            name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # Instructions
            if instructions:
                inst_para = doc.add_paragraph(instructions)
                inst_para.italic = True
                doc.add_paragraph()
            
            # Questions
            for i, q in enumerate(questions, 1):
                marker = q.get('marker', 'Answer:')
                prompt = q.get('prompt', '')
                points = q.get('points', 10)
                q_type = q.get('type', 'short_answer')
                
                # Question header with marker
                q_para = doc.add_paragraph()
                run = q_para.add_run(f"{marker} ")
                run.bold = True
                q_para.add_run(f"({points} pts)")
                
                # Question prompt
                if prompt:
                    doc.add_paragraph(prompt)
                
                # Answer space
                lines = 3 if q_type == 'short_answer' else 6 if q_type == 'essay' else 2
                for _ in range(lines):
                    doc.add_paragraph("_" * 70)
                
                doc.add_paragraph()
            
            filepath = os.path.join(output_folder, f"{safe_title}.docx")
            doc.save(filepath)
            
            # Open the folder
            os.system(f'open "{output_folder}"')
            
            return jsonify({"status": "exported", "path": filepath})
            
        except ImportError:
            return jsonify({"error": "python-docx not installed. Run: pip3 install python-docx"})
        except Exception as e:
            return jsonify({"error": str(e)})
    
    elif format_type == 'pdf':
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            from reportlab.lib.units import inch
            
            filepath = os.path.join(output_folder, f"{safe_title}.pdf")
            c = canvas.Canvas(filepath, pagesize=letter)
            width, height = letter
            
            y = height - inch
            
            # Title
            c.setFont("Helvetica-Bold", 18)
            c.drawCentredString(width/2, y, title)
            y -= 30
            
            # Name/Date
            c.setFont("Helvetica", 12)
            c.drawCentredString(width/2, y, "Name: _________________________ Date: _____________")
            y -= 40
            
            # Instructions
            if instructions:
                c.setFont("Helvetica-Oblique", 11)
                c.drawString(inch, y, instructions[:80])
                y -= 30
            
            # Questions
            c.setFont("Helvetica", 11)
            for i, q in enumerate(questions, 1):
                if y < 2*inch:
                    c.showPage()
                    y = height - inch
                
                marker = q.get('marker', 'Answer:')
                prompt = q.get('prompt', '')
                points = q.get('points', 10)
                q_type = q.get('type', 'short_answer')
                
                c.setFont("Helvetica-Bold", 11)
                c.drawString(inch, y, f"{marker} ({points} pts)")
                y -= 20
                
                if prompt:
                    c.setFont("Helvetica", 11)
                    # Simple word wrap
                    words = prompt.split()
                    line = ""
                    for word in words:
                        if len(line + word) < 80:
                            line += word + " "
                        else:
                            c.drawString(inch, y, line)
                            y -= 15
                            line = word + " "
                    if line:
                        c.drawString(inch, y, line)
                        y -= 15
                
                # Answer lines
                lines = 3 if q_type == 'short_answer' else 6 if q_type == 'essay' else 2
                for _ in range(lines):
                    y -= 20
                    c.line(inch, y, width - inch, y)
                
                y -= 30
            
            c.save()
            os.system(f'open "{output_folder}"')
            
            return jsonify({"status": "exported", "path": filepath})
            
        except ImportError:
            return jsonify({"error": "reportlab not installed. Run: pip3 install reportlab"})
        except Exception as e:
            return jsonify({"error": str(e)})
    
    return jsonify({"error": "Unknown format"})

@app.route('/api/analytics')
def get_analytics():
    """Load master CSV and return analytics data for charts."""
    import csv
    from collections import defaultdict

    # Get period filter from query params
    period_filter = request.args.get('period', 'all')

    master_file = "/Users/alexc/Downloads/Graider/Results/master_grades.csv"

    if not os.path.exists(master_file):
        return jsonify({"error": "No data yet", "students": [], "assignments": [], "trends": []})

    students = defaultdict(list)
    assignments = defaultdict(list)
    categories = defaultdict(lambda: {"content": [], "completeness": [], "writing": [], "effort": []})
    all_grades = []
    available_periods = set()

    try:
        with open(master_file, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                # Track all available periods
                row_quarter = row.get("Quarter", "")
                if row_quarter:
                    available_periods.add(row_quarter)

                # Filter by period if specified
                if period_filter != 'all' and row_quarter != period_filter:
                    continue

                grade_data = {
                    "date": row.get("Date", ""),
                    "student_id": row.get("Student ID", ""),
                    "student_name": row.get("Student Name", ""),
                    "first_name": row.get("First Name", ""),
                    "assignment": row.get("Assignment", ""),
                    "quarter": row_quarter,
                    "score": int(float(row.get("Overall Score", 0) or 0)),
                    "letter_grade": row.get("Letter Grade", ""),
                    "content": int(float(row.get("Content Accuracy", 0) or 0)),
                    "completeness": int(float(row.get("Completeness", 0) or 0)),
                    "writing": int(float(row.get("Writing Quality", 0) or 0)),
                    "effort": int(float(row.get("Effort Engagement", 0) or 0)),
                }
                all_grades.append(grade_data)
                
                # Group by student
                students[grade_data["student_name"]].append(grade_data)
                
                # Group by assignment
                assignments[grade_data["assignment"]].append(grade_data)
                
                # Track category scores
                categories[grade_data["student_name"]]["content"].append(grade_data["content"])
                categories[grade_data["student_name"]]["completeness"].append(grade_data["completeness"])
                categories[grade_data["student_name"]]["writing"].append(grade_data["writing"])
                categories[grade_data["student_name"]]["effort"].append(grade_data["effort"])
    except Exception as e:
        return jsonify({"error": str(e)})
    
    # Calculate student progress (for line charts)
    student_progress = []
    for name, grades in students.items():
        sorted_grades = sorted(grades, key=lambda x: x["date"])
        student_progress.append({
            "name": name,
            "grades": [{"date": g["date"], "assignment": g["assignment"], "score": g["score"]} for g in sorted_grades],
            "average": round(sum(g["score"] for g in grades) / len(grades), 1) if grades else 0,
            "trend": "improving" if len(sorted_grades) >= 2 and sorted_grades[-1]["score"] > sorted_grades[0]["score"] else 
                    "declining" if len(sorted_grades) >= 2 and sorted_grades[-1]["score"] < sorted_grades[0]["score"] else "stable"
        })
    
    # Calculate assignment averages (for bar charts)
    assignment_stats = []
    for name, grades in assignments.items():
        assignment_stats.append({
            "name": name[:30],
            "average": round(sum(g["score"] for g in grades) / len(grades), 1) if grades else 0,
            "count": len(grades),
            "highest": max(g["score"] for g in grades) if grades else 0,
            "lowest": min(g["score"] for g in grades) if grades else 0
        })
    
    # Calculate category averages per student (for radar charts)
    category_stats = []
    for name, cats in categories.items():
        category_stats.append({
            "name": name,
            "content": round(sum(cats["content"]) / len(cats["content"]) * 2.5, 1) if cats["content"] else 0,  # Scale to 100
            "completeness": round(sum(cats["completeness"]) / len(cats["completeness"]) * 4, 1) if cats["completeness"] else 0,
            "writing": round(sum(cats["writing"]) / len(cats["writing"]) * 5, 1) if cats["writing"] else 0,
            "effort": round(sum(cats["effort"]) / len(cats["effort"]) * 6.67, 1) if cats["effort"] else 0
        })
    
    # Class-wide stats
    all_scores = [g["score"] for g in all_grades]
    class_stats = {
        "total_assignments": len(all_grades),
        "total_students": len(students),
        "class_average": round(sum(all_scores) / len(all_scores), 1) if all_scores else 0,
        "highest": max(all_scores) if all_scores else 0,
        "lowest": min(all_scores) if all_scores else 0,
        "grade_distribution": {
            "A": len([s for s in all_scores if s >= 90]),
            "B": len([s for s in all_scores if 80 <= s < 90]),
            "C": len([s for s in all_scores if 70 <= s < 80]),
            "D": len([s for s in all_scores if 60 <= s < 70]),
            "F": len([s for s in all_scores if s < 60])
        }
    }
    
    # Students needing attention (below 70 average or declining)
    attention_needed = [s for s in student_progress if s["average"] < 70 or s["trend"] == "declining"]
    
    # Top performers
    top_performers = sorted(student_progress, key=lambda x: x["average"], reverse=True)[:5]
    
    return jsonify({
        "class_stats": class_stats,
        "student_progress": sorted(student_progress, key=lambda x: x["name"]),
        "assignment_stats": assignment_stats,
        "category_stats": category_stats,
        "attention_needed": attention_needed,
        "top_performers": top_performers,
        "all_grades": all_grades,
        "available_periods": sorted(list(available_periods))
    })

@app.route('/api/get-standards', methods=['POST'])
def get_standards():
    """Get standards for a specific state, grade, and subject."""
    data = request.json
    state = data.get('state', 'FL')
    grade = data.get('grade', '7')
    subject = data.get('subject', 'Civics')
    
    # In a real app, this would query a database
    # For now, we'll return hardcoded Florida Civics standards
    
    standards = []
    
    if state == 'FL' and subject == 'History':
        standards = [
            # 1. Research & Inquiry
            {
                "code": "SS.8.A.1.1",
                "benchmark": "Provide supporting details for an answer from text, interview for oral history, check validity of information from research/text, and identify strong vs. weak arguments.",
                "topics": ["Research Skills", "Primary Sources", "Analysis"]
            },
            {
                "code": "SS.8.A.1.3",
                "benchmark": "Analyze timelines, charts, graphs, and photographs to determine cause and effect.",
                "topics": ["Data Analysis", "Timelines", "Cause and Effect"]
            },
            
            # 2. Colonial Era
            {
                "code": "SS.8.A.2.1",
                "benchmark": "Compare the relationships among the British, French, Spanish, and Dutch in their struggle for colonization of North America.",
                "topics": ["Colonization", "European Powers", "North America"]
            },
            {
                "code": "SS.8.A.2.2",
                "benchmark": "Compare the characteristics of the New England, Middle, and Southern colonies.",
                "topics": ["Colonial Regions", "Economy", "Geography"]
            },
            {
                "code": "SS.8.A.2.4",
                "benchmark": "Identify the impact of key colonial figures on the economic, political, and social development of the colonies.",
                "topics": ["Colonial Figures", "Leadership"]
            },
            
            # 3. Revolution & Constitution
            {
                "code": "SS.8.A.3.1",
                "benchmark": "Explain the consequences of the French and Indian War in British policies for the American colonies from 1763 - 1774.",
                "topics": ["French and Indian War", "British Policy", "Taxation"]
            },
            {
                "code": "SS.8.A.3.2",
                "benchmark": "Explain American colonial reaction to British policy from 1763 - 1774.",
                "topics": ["Protest", "Revolutionary Spirit", "Sons of Liberty"]
            },
            {
                "code": "SS.8.A.3.6",
                "benchmark": "Examine the causes, course, and consequences of the American Revolution.",
                "topics": ["American Revolution", "War of Independence"]
            },
            {
                "code": "SS.8.A.3.10",
                "benchmark": "Examine the course and consequences of the Constitutional Convention (New Jersey Plan, Virginia Plan, Great Compromise, 3/5 Compromise).",
                "topics": ["Constitution", "Compromises", "Founding Fathers"]
            },
            
            # 4. Expansion & Reform
            {
                "code": "SS.8.A.4.1",
                "benchmark": "Examine the causes, course, and consequences of United States westward expansion and its growing diplomatic assertiveness.",
                "topics": ["Westward Expansion", "Manifest Destiny", "Louisiana Purchase"]
            },
            {
                "code": "SS.8.A.4.2",
                "benchmark": "Describe the debate surrounding the spread of slavery into western territories and Florida.",
                "topics": ["Slavery", "Sectionalism", "Expansion"]
            },
            {
                "code": "SS.8.A.4.6",
                "benchmark": "Identify significant individuals and groups who helped to define American culture and beliefs.",
                "topics": ["Reform Movements", "Culture", "Social Change"]
            },
            
            # 5. Civil War & Reconstruction
            {
                "code": "SS.8.A.5.1",
                "benchmark": "Explain the causes, course, and consequence of the Civil War (sectionalism, slavery, states' rights, family/gender roles).",
                "topics": ["Civil War", "Causes of War", "Sectional Crisis"]
            },
            {
                "code": "SS.8.A.5.3",
                "benchmark": "Explain major domestic and international economic, military, political, and socio-cultural events of Abraham Lincoln's presidency.",
                "topics": ["Lincoln", "Civil War Leadership", "Emancipation"]
            },
            {
                "code": "SS.8.A.5.6",
                "benchmark": "Explain the causes, course, and consequences of the 13th, 14th, and 15th amendments and the rise of Jim Crow laws.",
                "topics": ["Reconstruction", "Amendments", "Civil Rights", "Jim Crow"]
            }
        ]
    
    elif state == 'FL' and subject == 'Civics':
        standards = [
            {
                "code": "SS.7.C.1.1",
                "benchmark": "Recognize how Enlightenment ideas including Montesquieu's view of separation of power and John Locke's theories related to natural law and how John Locke's social contract influenced the Founding Fathers.",
                "topics": ["Enlightenment", "Natural Law", "Social Contract", "Separation of Powers"]
            },
            {
                "code": "SS.7.C.1.2",
                "benchmark": "Trace the impact that the Magna Carta, English Bill of Rights, Mayflower Compact, and Thomas Paine's Common Sense had on colonists' views of government.",
                "topics": ["Founding Documents", "Colonial Government", "Limited Government"]
            },
            {
                "code": "SS.7.C.1.3",
                "benchmark": "Describe how English policies and responses to colonial concerns led to the writing of the Declaration of Independence.",
                "topics": ["Colonial Grievances", "Declaration of Independence", "Revolution"]
            },
            {
                "code": "SS.7.C.1.4",
                "benchmark": "Analyze the ideas (natural rights, role of the government) and complaints set forth in the Declaration of Independence.",
                "topics": ["Natural Rights", "Declaration of Independence", "Grievances"]
            },
            {
                "code": "SS.7.C.1.5",
                "benchmark": "Identify how the weaknesses of the Articles of Confederation led to the writing of the Constitution.",
                "topics": ["Articles of Confederation", "Constitution", "Federalism"]
            },
            {
                "code": "SS.7.C.1.6",
                "benchmark": "Interpret the intentions of the Preamble of the Constitution.",
                "topics": ["Preamble", "Constitution", "Goals of Government"]
            },
            {
                "code": "SS.7.C.1.7",
                "benchmark": "Describe how the Constitution limits the powers of government through separation of powers and checks and balances.",
                "topics": ["Separation of Powers", "Checks and Balances", "Limited Government"]
            },
            {
                "code": "SS.7.C.1.8",
                "benchmark": "Explain the viewpoints of the Federalists and the Anti-Federalists regarding the ratification of the Constitution and inclusion of a bill of rights.",
                "topics": ["Federalists", "Anti-Federalists", "Ratification", "Bill of Rights"]
            },
            {
                "code": "SS.7.C.1.9",
                "benchmark": "Define the rule of law and recognize its influence on the development of the American legal, political, and governmental systems.",
                "topics": ["Rule of Law", "Fairness", "Accountability"]
            },
            {
                "code": "SS.7.C.2.1",
                "benchmark": "Define the term 'citizen,' and identify legal means of becoming a United States citizen.",
                "topics": ["Citizenship", "Naturalization", "Rights and Responsibilities"]
            },
            {
                "code": "SS.7.C.2.2",
                "benchmark": "Evaluate the obligations citizens have to obey laws, pay taxes, defend the nation, and serve on juries.",
                "topics": ["Obligations", "Citizenship", "Civic Duty"]
            },
            {
                "code": "SS.7.C.2.4",
                "benchmark": "Evaluate rights contained in the Bill of Rights and other amendments to the Constitution.",
                "topics": ["Bill of Rights", "Amendments", "Civil Liberties"]
            },
            {
                "code": "SS.7.C.2.8",
                "benchmark": "Identify America's current political parties, and illustrate their ideas about government.",
                "topics": ["Political Parties", "Ideologies", "Two-Party System"]
            },
            {
                "code": "SS.7.C.3.1",
                "benchmark": "Compare different forms of government (direct democracy, representative democracy, socialism, communism, monarchy, oligarchy, autocracy).",
                "topics": ["Forms of Government", "Democracy", "Authoritarianism"]
            },
            {
                "code": "SS.7.C.3.3",
                "benchmark": "Illustrate the structure and function (three branches of government established in Articles I, II, and III with corresponding powers) of government in the United States as established in the Constitution.",
                "topics": ["Three Branches", "Structure of Government", "Articles I-III"]
            },
            {
                "code": "SS.7.C.3.12",
                "benchmark": "Analyze the significance and outcomes of landmark Supreme Court cases including, but not limited to, Marbury v. Madison, Plessy v. Ferguson, Brown v. Board of Education, Gideon v. Wainwright, Miranda v. Arizona, in re Gault, Tinker v. Des Moines, Hazelwood v. Kuhlmeier, United States v. Nixon, and Bush v. Gore.",
                "topics": ["Supreme Court", "Landmark Cases", "Judicial Review"]
            }
        ]
    
    return jsonify({"standards": standards})

@app.route('/api/generate-lesson-plan', methods=['POST'])
def generate_lesson_plan():
    """Generate a lesson plan using AI."""
    data = request.json
    selected_standards = data.get('standards', [])
    config = data.get('config', {})
    
    if not selected_standards:
        return jsonify({"error": "No standards selected"})
        
    try:
        try:
            from openai import OpenAI
            from dotenv import load_dotenv
            import os
            import json
            
            # Load .env from the app directory, override any existing env vars
            app_dir = os.path.dirname(os.path.abspath(__file__))
            load_dotenv(os.path.join(app_dir, '.env'), override=True)
            api_key = os.getenv("OPENAI_API_KEY")
            
            # Use mock if key is missing or explicitly placeholder
            if not api_key or api_key.strip() == "" or "your-key-here" in api_key:
                raise Exception("Missing or placeholder API Key")
                
            client = OpenAI(api_key=api_key)

            period_length = config.get('periodLength', 50)
            content_type = config.get('type', 'Lesson Plan')

            prompt = f"""
You are an expert curriculum developer creating a COMPLETE, READY-TO-USE {content_type} for a {config.get('grade', '7')}th grade {config.get('subject', 'Civics')} class.

Title: "{config.get('title', 'Untitled')}"
Duration: {config.get('duration', 1)} day(s)
Class Period Length: {period_length} minutes

Standards to Cover:
{', '.join(selected_standards)}

Additional Requirements:
{config.get('requirements', 'None specified')}

Create a COMPREHENSIVE, DETAILED plan that a teacher can use immediately without any additional preparation. Include specific content, exact questions to ask, sample answers, and minute-by-minute timing.

Return JSON with this structure:
{{
    "title": "Full descriptive title",
    "overview": "2-3 sentence summary of what students will learn and accomplish",
    "essential_questions": ["Driving question 1", "Driving question 2"],
    "days": [
        {{
            "day": 1,
            "topic": "Specific topic for this day",
            "objective": "Students will be able to... (measurable, specific)",
            "standards_addressed": ["Which standards this day covers"],
            "vocabulary": [
                {{"term": "word", "definition": "student-friendly definition"}}
            ],
            "timing": [
                {{"minutes": "0-5", "duration": "5 min", "activity": "Bell Ringer", "description": "Detailed description of what students do"}},
                {{"minutes": "5-15", "duration": "10 min", "activity": "Direct Instruction", "description": "Exactly what teacher presents"}},
                {{"minutes": "15-35", "duration": "20 min", "activity": "Guided Practice", "description": "Step-by-step activity"}},
                {{"minutes": "35-45", "duration": "10 min", "activity": "Independent Practice", "description": "What students do independently"}},
                {{"minutes": "45-{period_length}", "duration": "5 min", "activity": "Closure", "description": "Exit ticket or wrap-up"}}
            ],
            "bell_ringer": {{
                "prompt": "The exact question or task for students",
                "expected_responses": ["Possible student answers"],
                "discussion_points": ["Follow-up questions to ask"]
            }},
            "direct_instruction": {{
                "key_points": ["Main concept 1 with explanation", "Main concept 2 with explanation"],
                "examples": ["Specific example to share with students"],
                "check_for_understanding": ["Questions to ask during instruction"]
            }},
            "activity": {{
                "name": "Activity name",
                "description": "Detailed step-by-step instructions",
                "grouping": "Individual/Pairs/Small Groups/Whole Class",
                "student_tasks": ["Step 1", "Step 2", "Step 3"],
                "teacher_role": "What teacher does during activity",
                "differentiation": {{
                    "struggling": "Support for struggling students",
                    "advanced": "Extension for advanced students"
                }}
            }},
            "assessment": {{
                "type": "Formative/Summative",
                "description": "How learning is assessed",
                "criteria": ["What demonstrates mastery"],
                "exit_ticket": "Specific exit ticket question if applicable"
            }},
            "materials": ["Specific item 1", "Specific item 2"],
            "homework": "Assignment if any, or null",
            "teacher_notes": "Tips, common misconceptions, or things to watch for"
        }}
    ],
    "unit_assessment": {{
        "type": "Test/Project/Presentation/Essay",
        "description": "Detailed description of summative assessment",
        "components": ["What the assessment includes"],
        "rubric_criteria": ["Grading criteria"]
    }},
    "resources": ["Textbook pages", "Websites", "Handouts needed"]
}}

Make the content SPECIFIC and DETAILED - include actual facts, real examples, specific questions with sample answers. This should be a complete lesson a teacher can walk into class and teach."""
            
            completion = client.chat.completions.create(
                model="gpt-4o",
                messages=[
                    {"role": "system", "content": "You are an expert curriculum developer. Return valid JSON only."},
                    {"role": "user", "content": prompt}
                ],
                response_format={"type": "json_object"}
            )
            
            content = completion.choices[0].message.content
            plan = json.loads(content)
            return jsonify({"plan": plan, "method": "AI"})
            
        except Exception as e:
            error_msg = str(e)
            print(f"OpenAI API Error: {error_msg}. Falling back to Mock Mode.")
            
            # Fallback Mock Plan based on Type
            content_type = config.get('type', 'Unit Plan')
            
            # Use specific mock content if requirements are present
            req_note = f" (Note: Custom requirements '{config.get('requirements')}' were ignored in Mock Mode)" if config.get('requirements') else ""
            
            mock_plan = {
                "title": f"{config.get('title', 'Unit Plan')} ({content_type} - Mock)",
                "overview": f"⚠️ GENERATED IN MOCK MODE. Real AI generation failed.\n\nError Details: {error_msg}\n\nThis is a simulated {content_type} for {config.get('subject')} covering standards: {', '.join(selected_standards)}. Check your API key and connection.{req_note}",
                "days": [],
                "unit_assessment": "Mock Assessment Idea"
            }
            
            # Generate days based on content type
            if content_type == 'Assignment':
                 mock_plan['days'] = [{
                    "day": 1,
                    "topic": "Assignment: Core Concepts",
                    "objective": "Students will demonstrate understanding through this assignment.",
                    "vocabulary": ["Key Term 1", "Key Term 2"],
                    "bell_ringer": "Review instructions.",
                    "activity": "Complete the worksheet/project.",
                    "assessment": "Graded submission.",
                    "materials": ["Worksheet", "Resources"]
                 }]
            else:
                mock_plan['days'] = [
                    {
                        "day": i + 1,
                        "topic": f"Mock Topic {i + 1}: Foundations",
                        "objective": "Students will understand key concepts relating to the selected standards.",
                        "vocabulary": ["Democracy", "Liberty", "Constitution", "Rights", "Government"],
                        "bell_ringer": "Students will answer a prompt on the board.",
                        "activity": "Group activity: Analyze primary source documents.",
                        "assessment": "Exit Ticket.",
                        "materials": ["Textbook", "Worksheet", "Pencils"]
                    } for i in range(int(config.get('duration', 5)))
                ]

            return jsonify({"plan": mock_plan, "method": "Mock", "error": error_msg})
            
    except Exception as e:
        print(f"Error generating plan: {e}")
        return jsonify({"error": str(e)})


@app.route('/api/export-lesson-plan', methods=['POST'])
def export_lesson_plan():
    """Export the lesson plan to a Word document."""
    data = request.json
    plan = data.get('plan', data)  # Handle both {plan: ...} and direct plan object

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import subprocess

        doc = Document()

        # Helper to format vocabulary
        def format_vocab(vocab_list):
            if not vocab_list:
                return ""
            items = []
            for v in vocab_list:
                if isinstance(v, dict):
                    term = v.get('term', '')
                    defn = v.get('definition', '')
                    items.append(f"{term}: {defn}" if defn else term)
                else:
                    items.append(str(v))
            return '\n'.join(items)

        # Helper to format bell ringer
        def format_bell_ringer(br):
            if not br:
                return ""
            if isinstance(br, str):
                return br
            prompt = br.get('prompt', '')
            responses = br.get('expected_responses', [])
            result = prompt
            if responses:
                result += "\n\nExpected Responses:\n" + '\n'.join(f"- {r}" for r in responses)
            return result

        # Helper to format activity
        def format_activity(act):
            if not act:
                return ""
            if isinstance(act, str):
                return act
            parts = []
            if act.get('name'):
                parts.append(f"Activity: {act['name']}")
            if act.get('description'):
                parts.append(act['description'])
            if act.get('grouping'):
                parts.append(f"Grouping: {act['grouping']}")
            if act.get('student_tasks'):
                parts.append("\nStudent Tasks:")
                for i, t in enumerate(act['student_tasks'], 1):
                    parts.append(f"  {i}. {t}")
            if act.get('differentiation'):
                diff = act['differentiation']
                if diff.get('struggling'):
                    parts.append(f"\nSupport for Struggling: {diff['struggling']}")
                if diff.get('advanced'):
                    parts.append(f"Extension for Advanced: {diff['advanced']}")
            return '\n'.join(parts)

        # Helper to format assessment
        def format_assessment(asmt):
            if not asmt:
                return ""
            if isinstance(asmt, str):
                return asmt
            parts = []
            if asmt.get('type'):
                parts.append(f"Type: {asmt['type']}")
            if asmt.get('description'):
                parts.append(asmt['description'])
            if asmt.get('exit_ticket'):
                parts.append(f"\nExit Ticket: \"{asmt['exit_ticket']}\"")
            return '\n'.join(parts)

        # Title
        title = doc.add_heading(plan.get('title', 'Lesson Plan'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Overview
        if plan.get('overview'):
            doc.add_heading('Overview', level=1)
            doc.add_paragraph(plan['overview'])

        # Essential Questions
        if plan.get('essential_questions'):
            doc.add_heading('Essential Questions', level=1)
            for q in plan['essential_questions']:
                doc.add_paragraph(f"• {q}")

        # Daily Plans
        if plan.get('days'):
            doc.add_heading('Daily Lesson Plans', level=1)

            for day in plan['days']:
                doc.add_heading(f"Day {day.get('day')}: {day.get('topic')}", level=2)

                # Objective
                if day.get('objective'):
                    p = doc.add_paragraph()
                    p.add_run('Learning Objective: ').bold = True
                    p.add_run(day['objective'])

                # Standards
                if day.get('standards_addressed'):
                    p = doc.add_paragraph()
                    p.add_run('Standards: ').bold = True
                    p.add_run(', '.join(day['standards_addressed']))

                # Timing
                if day.get('timing'):
                    doc.add_heading('Lesson Timing', level=3)
                    for t in day['timing']:
                        time_str = t.get('minutes') or t.get('duration', '')
                        doc.add_paragraph(f"{time_str} - {t.get('activity', '')}: {t.get('description', '')}")

                # Vocabulary
                vocab_text = format_vocab(day.get('vocabulary'))
                if vocab_text:
                    doc.add_heading('Vocabulary', level=3)
                    doc.add_paragraph(vocab_text)

                # Bell Ringer
                br_text = format_bell_ringer(day.get('bell_ringer'))
                if br_text:
                    doc.add_heading('Bell Ringer', level=3)
                    doc.add_paragraph(br_text)

                # Direct Instruction
                if day.get('direct_instruction'):
                    di = day['direct_instruction']
                    doc.add_heading('Direct Instruction', level=3)
                    if di.get('key_points'):
                        doc.add_paragraph('Key Points:')
                        for kp in di['key_points']:
                            doc.add_paragraph(f"• {kp}")
                    if di.get('check_for_understanding'):
                        doc.add_paragraph('\nCheck for Understanding:')
                        for q in di['check_for_understanding']:
                            doc.add_paragraph(f"• \"{q}\"")

                # Activity
                act_text = format_activity(day.get('activity'))
                if act_text:
                    doc.add_heading('Main Activity', level=3)
                    doc.add_paragraph(act_text)

                # Assessment
                asmt_text = format_assessment(day.get('assessment'))
                if asmt_text:
                    doc.add_heading('Assessment', level=3)
                    doc.add_paragraph(asmt_text)

                # Materials
                if day.get('materials'):
                    doc.add_heading('Materials', level=3)
                    doc.add_paragraph(', '.join(day['materials']))

                # Homework
                if day.get('homework'):
                    doc.add_heading('Homework', level=3)
                    doc.add_paragraph(day['homework'])

                # Teacher Notes
                if day.get('teacher_notes'):
                    doc.add_heading('Teacher Notes', level=3)
                    doc.add_paragraph(day['teacher_notes'])

                doc.add_paragraph()  # Spacer

        # Unit Assessment
        if plan.get('unit_assessment'):
            doc.add_heading('Summative Assessment', level=1)
            ua = plan['unit_assessment']
            if isinstance(ua, dict):
                if ua.get('type'):
                    doc.add_paragraph(f"Type: {ua['type']}")
                if ua.get('description'):
                    doc.add_paragraph(ua['description'])
                if ua.get('components'):
                    doc.add_paragraph('\nComponents:')
                    for c in ua['components']:
                        doc.add_paragraph(f"• {c}")
            else:
                doc.add_paragraph(str(ua))

        # Resources
        if plan.get('resources'):
            doc.add_heading('Resources', level=1)
            for r in plan['resources']:
                doc.add_paragraph(f"• {r}")

        # Save file
        filename = f"Lesson_Plan_{int(time.time())}.docx"
        # Try to get output folder, fall back to Downloads
        output_folder = grading_state.get("config", {}).get("assignments_folder", os.path.expanduser("~/Downloads"))
        filepath = os.path.join(output_folder, filename)
        doc.save(filepath)

        # Open the file
        subprocess.run(['open', filepath])

        return jsonify({"status": "success", "path": filepath})

    except Exception as e:
        import traceback
        traceback.print_exc()
        print(f"Error exporting plan: {e}")
        return jsonify({"error": str(e)})

# ══════════════════════════════════════════════════════════════
# MAIN
# ══════════════════════════════════════════════════════════════

if __name__ == '__main__':
    import webbrowser
    import threading
    
    def open_browser():
        """Open browser after short delay to let server start."""
        import time
        time.sleep(1.5)
        webbrowser.open('http://localhost:3000')
    
    print()
    print("╔══════════════════════════════════════════════════╗")
    print("║  🎓 Graider - AI-Powered Assignment Grading      ║")
    print("╠══════════════════════════════════════════════════╣")
    print("║                                                  ║")
    print("║  Open in browser: http://localhost:3000          ║")
    print("║                                                  ║")
    print("║  Press Ctrl+C to stop                            ║")
    print("╚══════════════════════════════════════════════════╝")
    print()
    
    # Auto-open browser
    threading.Thread(target=open_browser, daemon=True).start()
    
    app.run(host='0.0.0.0', port=3000, debug=False)
