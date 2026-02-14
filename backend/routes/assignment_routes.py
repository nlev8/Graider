"""
Assignment Builder API routes for Graider.
Handles saving, loading, listing, deleting, and exporting assignments.
"""
import os
import json
from flask import Blueprint, request, jsonify, send_from_directory

assignment_bp = Blueprint('assignment', __name__)

ASSIGNMENTS_DIR = os.path.expanduser("~/.graider_assignments")


@assignment_bp.route('/api/save-assignment-config', methods=['POST'])
def save_assignment_config():
    """Save assignment configuration for grading.

    Uses merge-save: loads existing config first, then merges incoming data
    on top. This prevents partial saves (e.g., from the grading tab's auto-save)
    from wiping fields like excludeMarkers, aliases, rubricType, etc.
    """
    data = request.json
    os.makedirs(ASSIGNMENTS_DIR, exist_ok=True)

    # Clean title for filename
    title = data.get('title', 'Untitled')
    safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()
    filepath = os.path.join(ASSIGNMENTS_DIR, f"{safe_title}.json")

    try:
        # Load existing config to preserve fields not in the incoming data
        existing = {}
        if os.path.exists(filepath):
            try:
                with open(filepath, 'r') as f:
                    existing = json.load(f)
            except (json.JSONDecodeError, Exception):
                existing = {}

        # Merge: existing fields are preserved, incoming data overwrites
        merged = {**existing, **data}

        with open(filepath, 'w') as f:
            json.dump(merged, f, indent=2)
        return jsonify({"status": "saved", "path": filepath})
    except Exception as e:
        return jsonify({"error": str(e)})


@assignment_bp.route('/api/list-assignments')
def list_assignments():
    """List saved assignment configurations with aliases."""
    if not os.path.exists(ASSIGNMENTS_DIR):
        return jsonify({"assignments": [], "assignmentData": {}})

    assignments = []
    assignment_data = {}  # Map of name -> {aliases: [...]}

    for f in os.listdir(ASSIGNMENTS_DIR):
        if f.endswith('.json'):
            name = f.replace('.json', '')
            assignments.append(name)
            # Load assignment to get aliases and completion status
            try:
                with open(os.path.join(ASSIGNMENTS_DIR, f), 'r') as af:
                    data = json.load(af)
                    imported_doc = data.get("importedDoc") or {}
                    assignment_data[name] = {
                        "aliases": data.get("aliases", []),
                        "title": data.get("title", name),
                        "completionOnly": data.get("completionOnly", False),
                        "rubricType": data.get("rubricType") or "standard",
                        "countsTowardsGrade": data.get("countsTowardsGrade", True),  # Default to True
                        "importedFilename": imported_doc.get("filename", ""),  # Original filename for matching
                    }
            except:
                assignment_data[name] = {"aliases": [], "title": name, "completionOnly": False, "rubricType": "standard", "countsTowardsGrade": True, "importedFilename": ""}

    return jsonify({"assignments": sorted(assignments), "assignmentData": assignment_data})


@assignment_bp.route('/api/load-assignment')
def load_assignment():
    """Load a saved assignment configuration."""
    name = request.args.get('name', '')
    filepath = os.path.join(ASSIGNMENTS_DIR, f"{name}.json")

    if not os.path.exists(filepath):
        return jsonify({"error": "Assignment not found"})

    try:
        with open(filepath, 'r') as f:
            data = json.load(f)
        return jsonify({"assignment": data})
    except Exception as e:
        return jsonify({"error": str(e)})


@assignment_bp.route('/api/delete-assignment', methods=['DELETE'])
def delete_assignment():
    """Delete a saved assignment configuration."""
    name = request.args.get('name', '')
    filepath = os.path.join(ASSIGNMENTS_DIR, f"{name}.json")

    if not os.path.exists(filepath):
        return jsonify({"error": "Assignment not found"})

    try:
        os.remove(filepath)
        return jsonify({"status": "deleted"})
    except Exception as e:
        return jsonify({"error": str(e)})


@assignment_bp.route('/api/export-assignment', methods=['POST'])
def export_assignment():
    """Export assignment to Word or PDF format."""
    data = request.json
    assignment = data.get('assignment', {})
    format_type = data.get('format', 'docx')

    title = assignment.get('title', 'Untitled Assignment')
    instructions = assignment.get('instructions', '')
    questions = assignment.get('questions', [])
    table_structured = assignment.get('tableStructured', False)

    output_folder = os.path.expanduser("~/Downloads/Graider/Assignments")
    os.makedirs(output_folder, exist_ok=True)

    safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()

    if format_type == 'docx':
        return _export_docx(title, instructions, questions, output_folder, safe_title,
                            table_structured=table_structured)
    elif format_type == 'pdf':
        return _export_pdf(title, instructions, questions, output_folder, safe_title)

    return jsonify({"error": "Unknown format"})


def _export_docx(title, instructions, questions, output_folder, safe_title,
                  table_structured=False):
    """Export assignment to Word format.

    If table_structured is True, uses Graider structured tables for questions.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title_para = doc.add_heading(title, 0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Name/Date line
        name_para = doc.add_paragraph("Name: _________________________ Date: _____________")
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Instructions
        if instructions:
            inst_para = doc.add_paragraph(instructions)
            inst_para.italic = True
            doc.add_paragraph()

        if table_structured:
            # Use Graider structured tables
            from backend.services.document_generator import DEFAULT_STYLE
            from backend.services.worksheet_generator import _add_graider_table, _add_graider_marker
            style = dict(DEFAULT_STYLE)

            for i, q in enumerate(questions, 1):
                marker = q.get('marker', 'Answer:')
                prompt = q.get('prompt', '')
                points = q.get('points', 10)
                q_type = q.get('type', 'short_answer')

                header = str(i) + ") " + (prompt or marker)
                height = 2160 if q_type == 'short_answer' else 4320 if q_type == 'essay' else 1440
                _add_graider_table(
                    doc, header, "GRAIDER:QUESTION:" + str(i),
                    points, style, height
                )

            _add_graider_marker(doc)
        else:
            # Original paragraph-based export
            for i, q in enumerate(questions, 1):
                marker = q.get('marker', 'Answer:')
                prompt = q.get('prompt', '')
                points = q.get('points', 10)
                q_type = q.get('type', 'short_answer')

                # Question header with marker
                q_para = doc.add_paragraph()
                run = q_para.add_run(f"{marker} ")
                run.bold = True
                q_para.add_run(f"({points} pts)")

                # Question prompt
                if prompt:
                    doc.add_paragraph(prompt)

                # Answer space
                lines = 3 if q_type == 'short_answer' else 6 if q_type == 'essay' else 2
                for _ in range(lines):
                    doc.add_paragraph("_" * 70)

                doc.add_paragraph()

        filepath = os.path.join(output_folder, f"{safe_title}.docx")
        doc.save(filepath)

        # Open the folder
        os.system(f'open "{output_folder}"')

        return jsonify({"status": "exported", "path": filepath})

    except ImportError:
        return jsonify({"error": "python-docx not installed. Run: pip3 install python-docx"})
    except Exception as e:
        return jsonify({"error": str(e)})


def _export_pdf(title, instructions, questions, output_folder, safe_title):
    """Export assignment to PDF format."""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
        from reportlab.lib.units import inch

        filepath = os.path.join(output_folder, f"{safe_title}.pdf")
        c = canvas.Canvas(filepath, pagesize=letter)
        width, height = letter

        y = height - inch

        # Title
        c.setFont("Helvetica-Bold", 18)
        c.drawCentredString(width / 2, y, title)
        y -= 30

        # Name/Date
        c.setFont("Helvetica", 12)
        c.drawCentredString(width / 2, y, "Name: _________________________ Date: _____________")
        y -= 40

        # Instructions
        if instructions:
            c.setFont("Helvetica-Oblique", 11)
            c.drawString(inch, y, instructions[:80])
            y -= 30

        # Questions
        c.setFont("Helvetica", 11)
        for i, q in enumerate(questions, 1):
            if y < 2 * inch:
                c.showPage()
                y = height - inch

            marker = q.get('marker', 'Answer:')
            prompt = q.get('prompt', '')
            points = q.get('points', 10)
            q_type = q.get('type', 'short_answer')

            c.setFont("Helvetica-Bold", 11)
            c.drawString(inch, y, f"{marker} ({points} pts)")
            y -= 20

            if prompt:
                c.setFont("Helvetica", 11)
                # Simple word wrap
                words = prompt.split()
                line = ""
                for word in words:
                    if len(line + word) < 80:
                        line += word + " "
                    else:
                        c.drawString(inch, y, line)
                        y -= 15
                        line = word + " "
                if line:
                    c.drawString(inch, y, line)
                    y -= 15

            # Answer lines
            lines = 3 if q_type == 'short_answer' else 6 if q_type == 'essay' else 2
            for _ in range(lines):
                y -= 20
                c.line(inch, y, width - inch, y)

            y -= 30

        c.save()
        os.system(f'open "{output_folder}"')

        return jsonify({"status": "exported", "path": filepath})

    except ImportError:
        return jsonify({"error": "reportlab not installed. Run: pip3 install reportlab"})
    except Exception as e:
        return jsonify({"error": str(e)})


@assignment_bp.route('/api/download-document/<filename>')
def download_document(filename):
    """Serve a generated document for download."""
    docs_dir = os.path.expanduser("~/Downloads/Graider/Documents")
    if not os.path.exists(os.path.join(docs_dir, filename)):
        return jsonify({"error": "Document not found"}), 404
    return send_from_directory(docs_dir, filename, as_attachment=True)


@assignment_bp.route('/api/download-worksheet/<filename>')
def download_worksheet(filename):
    """Serve a generated worksheet for download."""
    worksheets_dir = os.path.expanduser("~/Downloads/Graider/Worksheets")
    if not os.path.exists(os.path.join(worksheets_dir, filename)):
        return jsonify({"error": "Worksheet not found"}), 404
    return send_from_directory(worksheets_dir, filename, as_attachment=True)


@assignment_bp.route('/api/download-export/<filename>')
def download_export(filename):
    """Serve an exported CSV file for download."""
    exports_dir = os.path.expanduser("~/.graider_exports/focus")
    if not os.path.exists(os.path.join(exports_dir, filename)):
        return jsonify({"error": "Export not found"}), 404
    return send_from_directory(exports_dir, filename, as_attachment=True)
