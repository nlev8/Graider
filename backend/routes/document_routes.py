"""
Document handling API routes for Graider.
Handles document parsing and export.
"""
import logging
import os
import base64
from pathlib import Path
from flask import Blueprint, request, jsonify
from backend.utils.auth_decorators import require_teacher
from backend.utils.errors import handle_route_errors
import sentry_sdk

_logger = logging.getLogger(__name__)

document_bp = Blueprint('document', __name__)


@document_bp.route('/api/parse-document', methods=['POST'])
@require_teacher
@handle_route_errors
def parse_document():
    """Parse an uploaded Word/PDF document and convert to HTML."""
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    filename = file.filename.lower()

    ALLOWED_EXTENSIONS = {'.docx', '.pdf', '.txt', '.doc', '.rtf'}
    ext = os.path.splitext(filename)[1].lower()
    if not ext and filename.startswith('.'):
        ext = filename  # Handle bare ".txt" filenames
    if ext not in ALLOWED_EXTENSIONS:
        return jsonify({"error": f"Unsupported file type: {ext}"}), 400

    file_data = file.read()

    if filename.endswith('.docx'):
        return _parse_docx(file_data, file.filename)
    elif filename.endswith('.pdf'):
        return _parse_pdf(file_data, file.filename)
    elif filename.endswith('.txt'):
        return _parse_txt(file_data, file.filename)
    else:
        return jsonify({"error": "Unsupported file type. Use .docx, .pdf, or .txt"}), 400


def _parse_docx(file_data, filename):
    """Parse Word document to HTML."""
    import io

    # Try mammoth first for best HTML conversion
    try:
        import mammoth
        from docx import Document

        result = mammoth.convert_to_html(io.BytesIO(file_data))
        html = result.value

        styled_html = f'''
        <style>
            body {{ font-family: Georgia, serif; line-height: 1.6; }}
            table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
            td, th {{ border: 1px solid #ccc; padding: 8px 12px; text-align: left; }}
            th {{ background: #f5f5f5; font-weight: bold; }}
            p {{ margin: 10px 0; }}
            h1, h2, h3 {{ margin: 20px 0 10px 0; }}
            ul, ol {{ margin: 10px 0; padding-left: 25px; }}
        </style>
        {html}
        '''

        # Extract plain text for marking
        doc = Document(io.BytesIO(file_data))
        plain_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                plain_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    plain_text.append(row_text)

        # Try to extract full document title from metadata or first heading
        doc_title = None
        try:
            # Check document properties for title
            if doc.core_properties.title:
                doc_title = doc.core_properties.title
        except Exception as e:
            sentry_sdk.capture_exception(e)

        # If no metadata title, use first heading or first paragraph
        if not doc_title and plain_text:
            # First non-empty line is likely the title
            doc_title = plain_text[0].strip()

        return jsonify({
            "html": styled_html,
            "text": '\n'.join(plain_text),
            "filename": filename,
            "doc_title": doc_title,  # Full title from document content
            "type": "html"
        })

    except ImportError:
        # Fallback to python-docx only
        from docx import Document
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        doc = Document(io.BytesIO(file_data))
        html_parts = ['<div style="font-family: Georgia, serif; line-height: 1.6;">']

        for child in doc.element.body.iterchildren():
            if isinstance(child, CT_P):
                para = Paragraph(child, doc)
                if para.text.strip():
                    style = para.style.name if para.style else ''
                    if 'Heading 1' in style:
                        html_parts.append(f'<h1>{para.text}</h1>')
                    elif 'Heading 2' in style:
                        html_parts.append(f'<h2>{para.text}</h2>')
                    elif 'Heading' in style:
                        html_parts.append(f'<h3>{para.text}</h3>')
                    else:
                        html_parts.append(f'<p>{para.text}</p>')
            elif isinstance(child, CT_Tbl):
                table = Table(child, doc)
                html_parts.append('<table style="border-collapse: collapse; width: 100%; margin: 15px 0;">')
                for row_idx, row in enumerate(table.rows):
                    html_parts.append('<tr>')
                    for cell in row.cells:
                        tag = 'th' if row_idx == 0 else 'td'
                        style = 'border: 1px solid #ccc; padding: 8px 12px; background: #f5f5f5;' if row_idx == 0 else 'border: 1px solid #ccc; padding: 8px 12px;'
                        html_parts.append(f'<{tag} style="{style}">{cell.text}</{tag}>')
                    html_parts.append('</tr>')
                html_parts.append('</table>')

        html_parts.append('</div>')

        plain_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                plain_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    plain_text.append(row_text)

        return jsonify({
            "html": ''.join(html_parts),
            "text": '\n'.join(plain_text),
            "filename": filename,
            "type": "html"
        })


def _parse_pdf(file_data, filename):
    """Parse PDF document to HTML with images."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_data, filetype="pdf")
        images_html = []
        plain_text = []

        for page_num, page in enumerate(doc):
            mat = fitz.Matrix(1.5, 1.5)  # 150 DPI
            pix = page.get_pixmap(matrix=mat)
            img_data = pix.tobytes("png")
            img_base64 = base64.b64encode(img_data).decode('utf-8')

            images_html.append(f'''
                <div style="margin-bottom: 20px; box-shadow: 0 2px 10px rgba(0,0,0,0.1);">
                    <div style="background: #6366f1; color: white; padding: 5px 15px; font-size: 12px;">Page {page_num + 1}</div>
                    <img src="data:image/png;base64,{img_base64}" style="width: 100%; display: block;" />
                </div>
            ''')
            plain_text.append(page.get_text())

        doc.close()

        full_html = f'''
        <div style="background: #e5e5e5; padding: 20px;">
            {''.join(images_html)}
        </div>
        '''

        return jsonify({
            "html": full_html,
            "text": '\n\n'.join(plain_text),
            "filename": filename,
            "type": "html"
        })

    except ImportError:
        return jsonify({"error": "PDF support requires PyMuPDF. Run: pip3 install pymupdf"})


def _parse_txt(file_data, filename):
    """Parse text file."""
    text = file_data.decode('utf-8', errors='ignore')
    html = f'<pre style="font-family: Monaco, monospace; white-space: pre-wrap; line-height: 1.6;">{text}</pre>'
    return jsonify({
        "html": html,
        "text": text,
        "filename": filename,
        "type": "html"
    })
