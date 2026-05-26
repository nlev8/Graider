"""
Lesson Planner API routes for Graider.
Handles standards retrieval and lesson plan generation/export.
"""
import os
import sys
import json
import time
import math
import re
import logging
import subprocess
from flask import Blueprint, request, jsonify, g, send_file
from backend.services.openai_context import build_openai_context
from werkzeug.utils import secure_filename
from pathlib import Path
from backend.extensions import limiter
from backend.utils.auth_decorators import require_teacher
from backend.utils.errors import handle_route_errors
from backend.retry import with_retry
from backend.paths import graider_export_dir

ALLOWED_DOC_EXTENSIONS = {'.docx', '.pdf', '.txt', '.doc', '.rtf', '.png', '.jpg', '.jpeg'}

# Import MODEL_PRICING for token cost tracking
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))
from assignment_grader import MODEL_PRICING

# Import storage abstraction for saving grading configs to Supabase
try:
    from backend.storage import save as storage_save
except ImportError:
    try:
        from storage import save as storage_save
    except ImportError:
        storage_save = None

# ── Phase 3b1 PR5: shim removed. Post-processing pipeline lives in ──────────
# backend/services/assignment_post_processing.py. Import only the names this
# module still calls directly (names only reached INSIDE the moved orchestrator
# do not need to be re-imported here).
from backend.services.assignment_post_processing import (
    PLANNER_COSTS_FILE,
    _build_question_count_instruction,
    _build_section_categories_prompt,
    _build_subject_boundary_prompt,
    _classify_question_type,
    _extract_usage,
    _hydrate_question,
    _merge_usage,
    _post_process_assignment,
    _record_planner_cost,
    _split_markdown_table,
    _validate_question,
)

# ── Tier 2 PR1: standards loading + matching extracted to ───────────────────
# backend/services/planner_standards.py (pure logic, no Flask). Re-exported
# here so existing `from backend.routes.planner_routes import load_standards`
# (and friends) callers keep resolving. _get_openai_context stays in this
# module because it reads Flask `g`.
from backend.services.planner_standards import (  # noqa: F401
    DATA_DIR,
    DOCUMENTS_DIR,
    _extract_grade_from_code,
    _get_standards_map,
    _grade_matches,
    _load_standards_file,
    load_standards,
    load_support_documents_for_planning,
)

# ── Tier 2 PR2: document / visual / platform-export rendering extracted to ──
# backend/services/planner_export.py (pure logic, no Flask). Re-exported here
# so existing `from backend.routes.planner_routes import _create_visual_for_question`
# (and friends) callers — plus test monkeypatch targets like
# `backend.routes.planner_routes._get_export_dir` — keep resolving unchanged.
# _save_grading_config_for_export stays in this module: it has an inner
# `from flask import g` (best-effort Supabase save) and could not move
# verbatim without violating the no-Flask-import boundary of the service.
from backend.services.planner_export import (  # noqa: F401
    _create_visual_for_question,
    _export_assignment_docx_graider,
    _get_export_dir,
    _question_to_visual_dict,
    generate_qti_xml,
    parse_template_structure,
)
from backend.services.planner_export import build_platform_export
from backend.services.planner_study_aids import (
    generate_study_guide_content,
    generate_flashcards_content,
)
from backend.services.planner_study_aids import generate_slides_payload
from backend.services.planner_content_tools import adjust_reading_level_content
from backend.services.planner_standards import rewrite_for_alignment_content
from backend.services.planner_standards import align_document_to_standards_content
from backend.services.planner_standards import IMAGE_EXTENSIONS, TextExtractionError, extract_text_from_upload
from backend.services.planner_generation import brainstorm_lesson_ideas_content
from backend.services.planner_generation import generate_lesson_plan_content
from backend.services.planner_generation import generate_assessment_content
from backend.services.planner_generation import generate_assignment_from_lesson_content
from backend.services.planner_generation import generate_replacement_questions
from backend.services.planner_assessments import grade_assessment_answers_logic

# ── Tier 2 PR3: prompt construction extracted to ───────────────────────────
# backend/services/planner_prompts.py (pure string building, no Flask).
# Re-exported here so existing `from backend.routes.planner_routes import
# _build_assignment_prompt` (and friends) callers keep resolving unchanged.
from backend.services.planner_prompts import (  # noqa: F401
    _build_assignment_prompt,
    _build_period_differentiation_block,
)

GEOMETRY_SUBTYPES = {
    'geometry', 'triangle', 'rectangle', 'circle', 'trapezoid',
    'parallelogram', 'rectangular_prism', 'cylinder', 'regular_polygon',
    'pythagorean', 'angles', 'similarity', 'trig'
}


# ── Phase 3c orphan patterns (defined but not currently referenced) ─────────
# Patterns for common theorem setups to detect over-determined or inconsistent
# values. Retained at module level to preserve prior compilation semantics.
import re as _re

_TANGENT_SECANT_RE = _re.compile(
    r'tangent.*(?:external|outside).*?(\d+(?:\.\d+)?).*?(?:whole|secant).*?(\d+(?:\.\d+)?)',
    _re.IGNORECASE | _re.DOTALL,
)
_CHORD_CHORD_RE = _re.compile(
    r'chord.*?(\d+(?:\.\d+)?)\s*[×*·]\s*(\d+(?:\.\d+)?)\s*=\s*(\d+(?:\.\d+)?)\s*[×*·]\s*(\d+(?:\.\d+)?)',
    _re.IGNORECASE,
)
_PYTHAGOREAN_RE = _re.compile(
    r'(?:right\s+triangle|hypotenuse|legs?).*?(\d+(?:\.\d+)?).*?(\d+(?:\.\d+)?).*?(\d+(?:\.\d+)?)',
    _re.IGNORECASE | _re.DOTALL,
)


def _get_openai_context():
    """Extract user_id for the post-processing pipeline.

    Returns (user_id, None) — the client slot is kept for call-site
    compatibility but is no longer used now that _auto_fix_flagged_questions
    builds its own LLM adapter internally.
    """
    try:
        user_id = getattr(g, 'user_id', 'local-dev')
        return build_openai_context(user_id)
    except Exception as e:
        _logger.warning("OpenAI context unavailable (non-fatal): %s", e)
        return None, None


planner_bp = Blueprint('planner', __name__)
_logger = logging.getLogger(__name__)

# Standards loading + matching (DATA_DIR, DOCUMENTS_DIR, load_standards,
# _get_standards_map, _load_standards_file, _extract_grade_from_code,
# _grade_matches, load_support_documents_for_planning) extracted to
# backend/services/planner_standards.py — re-imported above (Tier 2 PR1).


@planner_bp.route('/api/get-standards', methods=['POST'])
@require_teacher
@handle_route_errors
def get_standards():
    """Get standards for a specific state, grade, and subject."""
    data = request.json
    state = data.get('state', 'FL')
    grade = data.get('grade', '7')
    subject = data.get('subject', 'Civics')

    # Try to load from JSON files first, filtered by grade
    result = load_standards(state, subject, grade)
    standards = result['standards']

    if standards:
        return jsonify({
            "standards": standards,
            "grade": grade,
            "subject": subject,
            "fallback_used": result.get("fallback_used", False),
            "fallback_framework": result.get("fallback_framework"),
            "no_framework": result.get("no_framework", False),
            "state_note": result.get("state_note"),
        })

    # Fallback to empty if no data file exists
    return jsonify({
        "standards": [],
        "grade": grade,
        "subject": subject,
        "fallback_used": result.get("fallback_used", False),
        "fallback_framework": result.get("fallback_framework"),
        "no_framework": result.get("no_framework", False),
        "state_note": result.get("state_note"),
    })


@planner_bp.route('/api/available-states', methods=['GET'])
def get_available_states():
    """Return list of all supported states with names. No auth required."""
    smap = _get_standards_map()
    states = []
    for code, info in sorted(smap.get('states', {}).items(), key=lambda x: x[1].get('name', '')):
        states.append({'code': code, 'name': info.get('name', code)})
    return jsonify({'states': states})


@planner_bp.route('/api/align-document-to-standards', methods=['POST'])
@require_teacher
@handle_route_errors
def align_document_to_standards():
    """Analyze a document and identify which standards it aligns to."""
    data = request.json
    doc_text = data.get('documentText', '')
    state = data.get('state', 'FL')
    grade = data.get('grade', '7')
    subject = data.get('subject', '')

    if not doc_text or not doc_text.strip():
        return jsonify({"error": "No document text provided"})
    if not subject:
        return jsonify({"error": "Subject is required. Set it in Settings."})

    result = load_standards(state, subject, grade)
    standards = result['standards']
    if not standards:
        return jsonify({"error": f"No standards found for {state} {subject} grade {grade}. Check that a standards file exists in backend/data/."})

    # Build condensed standards reference for AI prompt (limit token usage)
    standards_ref = []
    for s in standards:
        standards_ref.append({
            "code": s.get("code", ""),
            "benchmark": s.get("benchmark", "")[:200],
            "topics": s.get("topics", []),
            "vocabulary": s.get("vocabulary", []),
            "dok": s.get("dok", ""),
        })

    try:
        from backend.api_keys import get_api_key as _gak
        teacher_id = getattr(g, 'user_id', 'local-dev')
        api_key = _gak('openai', teacher_id)

        if not api_key or api_key.strip() == "" or "your-key-here" in api_key:
            return jsonify({"error": "Missing or placeholder OpenAI API Key"})

        return jsonify(align_document_to_standards_content(
            doc_text=doc_text, standards_ref=standards_ref, api_key=api_key,
        ))

    except Exception:
        _logger.exception("Standards alignment failed")
        return jsonify({"error": "An internal error occurred"}), 500


@planner_bp.route('/api/rewrite-for-alignment', methods=['POST'])
@require_teacher
@handle_route_errors
def rewrite_for_alignment():
    """Rewrite specific questions to better align with selected standards."""
    data = request.json
    questions = data.get('questions', [])
    doc_text = data.get('documentText', '')
    grade = data.get('grade', '7')
    subject = data.get('subject', '')
    state = data.get('state', 'FL')

    if not questions:
        return jsonify({"error": "No questions provided for rewriting"})

    result = load_standards(state, subject, grade)
    standards = result['standards']
    standards_by_code = {s.get("code"): s for s in standards} if standards else {}

    # Enrich questions with full standard details
    enriched_questions = []
    for q in questions:
        std_code = q.get('target_standard', '')
        std_detail = standards_by_code.get(std_code, {})
        enriched_questions.append({
            "original_text": q.get('original_text', ''),
            "target_standard_code": std_code,
            "target_benchmark": std_detail.get('benchmark', ''),
            "target_topics": std_detail.get('topics', []),
            "target_vocabulary": std_detail.get('vocabulary', []),
            "essential_questions": std_detail.get('essential_questions', []),
            "rewrite_goal": q.get('rewrite_goal', ''),
        })

    try:
        from backend.api_keys import get_api_key as _gak
        teacher_id = getattr(g, 'user_id', 'local-dev')
        api_key = _gak('openai', teacher_id)

        if not api_key or api_key.strip() == "" or "your-key-here" in api_key:
            return jsonify({"error": "Missing or placeholder OpenAI API Key"})

        return jsonify(rewrite_for_alignment_content(
            enriched_questions=enriched_questions, doc_text=doc_text,
            grade=grade, subject=subject, api_key=api_key,
        ))

    except Exception:
        _logger.exception("Rewrite for alignment failed")
        return jsonify({"error": "An internal error occurred"}), 500


@planner_bp.route('/api/get-lesson-templates', methods=['POST'])
@require_teacher
@handle_route_errors
def get_lesson_templates():
    """Get subject-specific lesson activity templates."""
    data = request.json
    subject = data.get('subject', '').lower().replace(' ', '_').replace('/', '-')

    templates_file = DATA_DIR / 'lesson_templates.json'
    if not templates_file.exists():
        return jsonify({"templates": None, "error": "Templates file not found"})

    try:
        with open(templates_file, 'r') as f:
            all_templates = json.load(f)

        # Try exact match first
        if subject in all_templates:
            return jsonify({"templates": all_templates[subject], "subject": subject})

        # Try partial match (e.g., 'us_history' -> 'social_studies')
        subject_mapping = {
            'us_history': 'social_studies',
            'world_history': 'social_studies',
            'civics': 'social_studies',
            'english-ela': 'social_studies',  # Use social_studies templates as fallback
        }

        mapped_subject = subject_mapping.get(subject)
        if mapped_subject and mapped_subject in all_templates:
            return jsonify({"templates": all_templates[mapped_subject], "subject": mapped_subject})

        # Return all available subjects
        return jsonify({
            "templates": None,
            "available_subjects": list(all_templates.keys()),
            "requested": subject
        })

    except Exception as e:
        _logger.exception("Failed to get lesson templates")
        return jsonify({"error": "An internal error occurred"}), 500


@planner_bp.route('/api/brainstorm-lesson-ideas', methods=['POST'])
@limiter.limit("10 per minute")
@require_teacher
@handle_route_errors
def brainstorm_lesson_ideas():
    """Generate multiple lesson plan ideas/concepts for selected standards."""
    data = request.json
    selected_standards = data.get('standards', [])
    config = data.get('config', {})

    if not selected_standards and not data.get('referenceDocs'):
        return jsonify({"error": "Please select standards or upload reference documents"})

    _subject = config.get('subject', '').strip()
    _grade = config.get('grade', '').strip()
    if not _subject or not _grade:
        from flask import current_app
        current_app.logger.warning(
            f"Brainstorm requested without subject/grade: subject={_subject!r}, grade={_grade!r}")

    try:
        from backend.api_keys import get_api_key
        api_key = get_api_key('openai', getattr(g, 'user_id', 'local-dev'))
        return jsonify(brainstorm_lesson_ideas_content(
            selected_standards=selected_standards, config=config, api_key=api_key))
    except Exception as e:
        error_msg = str(e)
        _logger.exception("Brainstorm Error")
        # Fallback mock ideas
        mock_ideas = {
            "ideas": [
                {"id": 1, "title": "Interactive Discussion", "approach": "Discussion", "brief": "Engage students in guided discussion.", "hook": "Opening question", "key_activity": "Socratic seminar", "assessment_type": "Participation rubric"},
                {"id": 2, "title": "Hands-On Activity", "approach": "Activity-Based", "brief": "Students learn through doing.", "hook": "Mystery item reveal", "key_activity": "Station rotations", "assessment_type": "Exit ticket"},
                {"id": 3, "title": "Research Project", "approach": "Research", "brief": "Students investigate topics independently.", "hook": "Essential question", "key_activity": "Guided research", "assessment_type": "Presentation"},
            ]
        }
        return jsonify({**mock_ideas, "error": error_msg, "method": "Mock"})


@planner_bp.route('/api/generate-lesson-plan', methods=['POST'])
@limiter.limit("10 per minute")
@require_teacher
@handle_route_errors
def generate_lesson_plan():
    """Generate a lesson plan using AI."""
    data = request.json
    selected_standards = data.get('standards', [])
    config = data.get('config', {})
    selected_idea = data.get('selectedIdea')  # Optional: from brainstorming
    generate_variations = data.get('generateVariations', False)  # Generate multiple variations
    reference_docs = data.get('referenceDocs', [])  # Uploaded reference documents

    if not selected_standards and not data.get('referenceDocs'):
        return jsonify({"error": "Please select standards or upload reference documents"})

    _subject = config.get('subject', '').strip()
    _grade = config.get('grade', '').strip()
    if not _subject or not _grade:
        from flask import current_app
        current_app.logger.warning(
            f"Lesson plan requested without subject/grade: subject={_subject!r}, grade={_grade!r}")

    try:
        from backend.api_keys import get_api_key as _gak
        api_key = _gak('openai', getattr(g, 'user_id', 'local-dev'))
        return jsonify(generate_lesson_plan_content(
            selected_standards=selected_standards, config=config, selected_idea=selected_idea,
            generate_variations=generate_variations, reference_docs=reference_docs,
            api_key=api_key, openai_context=_get_openai_context()))
    except Exception as e:
        error_msg = str(e)
        # Keep as warning (recoverable — falls back to mock) but add exc_info for traceback.
        _logger.warning("OpenAI API Error — falling back to Mock Mode", exc_info=True)

        # Fallback Mock Plan
        content_type = config.get('type', 'Unit Plan')

        mock_plan = {
            "title": f"{config.get('title', 'Unit Plan')} ({content_type} - Mock)",
            "overview": f"GENERATED IN MOCK MODE. Error: {error_msg}",
            "days": [],
            "unit_assessment": "Mock Assessment"
        }

        if content_type == 'Assignment':
            mock_plan['days'] = [{
                "day": 1,
                "topic": "Assignment: Core Concepts",
                "objective": "Students will demonstrate understanding.",
                "vocabulary": ["Key Term 1", "Key Term 2"],
                "bell_ringer": "Review instructions.",
                "activity": "Complete the assignment.",
                "assessment": "Graded submission.",
                "materials": ["Worksheet", "Resources"]
            }]
        else:
            mock_plan['days'] = [
                {
                    "day": i + 1,
                    "topic": f"Mock Topic {i + 1}",
                    "objective": "Students will understand key concepts.",
                    "vocabulary": ["Term 1", "Term 2"],
                    "bell_ringer": "Prompt on board.",
                    "activity": "Group activity.",
                    "assessment": "Exit Ticket.",
                    "materials": ["Textbook", "Worksheet"]
                } for i in range(int(config.get('duration', 5)))
            ]

        return jsonify({"plan": mock_plan, "method": "Mock", "error": error_msg})


@planner_bp.route('/api/generate-assignment-from-lesson', methods=['POST'])
@limiter.limit("10 per minute")
@require_teacher
@handle_route_errors
def generate_assignment_from_lesson():
    """Generate an assignment based on an existing lesson plan."""
    data = request.json
    lesson_plan = data.get('lessonPlan', {})
    config = data.get('config', {})
    assignment_type = data.get('assignmentType', 'assignment')

    # Validate assignment_type
    if assignment_type not in ('assignment', 'project', 'essay'):
        assignment_type = 'assignment'

    _subject = config.get('subject', '').strip()
    _grade = config.get('grade', '').strip()
    if not _subject or not _grade:
        from flask import current_app
        current_app.logger.warning(
            f"Assignment-from-lesson requested without subject/grade: subject={_subject!r}, grade={_grade!r}")

    config_standards = config.get('standards', [])
    reference_docs = config.get('referenceDocs', [])
    global_ai_notes = config.get('globalAINotes', '')
    content_only = config.get('contentOnly', False)

    if not lesson_plan:
        return jsonify({"error": "No lesson plan provided"})

    try:
        from backend.api_keys import get_api_key as _gak
        api_key = _gak('openai', getattr(g, 'user_id', 'local-dev'))
        return jsonify(generate_assignment_from_lesson_content(
            lesson_plan=lesson_plan, config=config, assignment_type=assignment_type,
            content_only=content_only, config_standards=config_standards,
            reference_docs=reference_docs, api_key=api_key, openai_context=_get_openai_context()))
    except Exception as e:
        error_msg = str(e)
        _logger.exception("Assignment Generation Error")

        # Detect network-blocked AI provider
        error_lower = error_msg.lower()
        if any(kw in error_lower for kw in ['connection', 'timeout', 'unreachable', 'refused', 'apiconnectionerror', 'connecttimeout', 'name resolution']):
            return jsonify({
                "error": "Unable to connect to the AI provider. This may be due to network restrictions on your school's network. "
                         "Try again from a different network or contact your IT department to allow access to OpenAI/Anthropic services.",
                "network_error": True,
            }), 503

        # Fallback mock assignment
        mock_assignment = {
            "title": f"{assignment_type.title()} - {lesson_plan.get('title', 'Lesson')}",
            "type": assignment_type,
            "instructions": "Complete all sections. Show your work.",
            "time_estimate": "30-45 minutes",
            "total_points": 100,
            "sections": [
                {
                    "name": "Part A: Key Concepts",
                    "type": "short_answer",
                    "points": 50,
                    "questions": [
                        {"number": 1, "question": "Explain the main concept from the lesson.", "points": 25},
                        {"number": 2, "question": "Give an example that demonstrates your understanding.", "points": 25}
                    ]
                },
                {
                    "name": "Part B: Vocabulary",
                    "type": "matching",
                    "points": 50,
                    "questions": [
                        {"number": 1, "question": "Match terms to definitions", "points": 50}
                    ]
                }
            ],
            "error": error_msg,
            "method": "Mock"
        }
        return jsonify({"assignment": mock_assignment, "method": "Mock", "error": error_msg})


@planner_bp.route('/api/export-lesson-plan', methods=['POST'])
@require_teacher
@handle_route_errors
def export_lesson_plan():
    """Export the lesson plan to a Word document."""
    data = request.json
    plan = data.get('plan', data)

    if not isinstance(plan, dict) or not any(
        plan.get(k) for k in (
            'overview', 'essential_questions', 'days',
            'unit_assessment', 'resources', 'sections',
        )
    ):
        return jsonify({"error": "Nothing to export: the lesson plan has no content."}), 400

    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Helper functions
        def format_vocab(vocab_list):
            if not vocab_list:
                return ""
            items = []
            for v in vocab_list:
                if isinstance(v, dict):
                    term = v.get('term', '')
                    defn = v.get('definition', '')
                    items.append(f"{term}: {defn}" if defn else term)
                else:
                    items.append(str(v))
            return '\n'.join(items)

        def format_bell_ringer(br):
            if not br:
                return ""
            if isinstance(br, str):
                return br
            prompt = br.get('prompt', '')
            responses = br.get('expected_responses', [])
            result = prompt
            if responses:
                result += "\n\nExpected Responses:\n" + '\n'.join(f"- {r}" for r in responses)
            return result

        def format_activity(act):
            if not act:
                return ""
            if isinstance(act, str):
                return act
            parts = []
            if act.get('name'):
                parts.append(f"Activity: {act['name']}")
            if act.get('description'):
                parts.append(act['description'])
            if act.get('grouping'):
                parts.append(f"Grouping: {act['grouping']}")
            if act.get('student_tasks'):
                parts.append("\nStudent Tasks:")
                for i, t in enumerate(act['student_tasks'], 1):
                    parts.append(f"  {i}. {t}")
            if act.get('differentiation'):
                diff = act['differentiation']
                if diff.get('struggling'):
                    parts.append(f"\nSupport for Struggling: {diff['struggling']}")
                if diff.get('advanced'):
                    parts.append(f"Extension for Advanced: {diff['advanced']}")
            return '\n'.join(parts)

        def format_assessment(asmt):
            if not asmt:
                return ""
            if isinstance(asmt, str):
                return asmt
            parts = []
            if asmt.get('type'):
                parts.append(f"Type: {asmt['type']}")
            if asmt.get('description'):
                parts.append(asmt['description'])
            if asmt.get('exit_ticket'):
                parts.append(f"\nExit Ticket: \"{asmt['exit_ticket']}\"")
            return '\n'.join(parts)

        # Title
        title = doc.add_heading(plan.get('title', 'Lesson Plan'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Overview
        if plan.get('overview'):
            doc.add_heading('Overview', level=1)
            doc.add_paragraph(plan['overview'])

        # Essential Questions
        if plan.get('essential_questions'):
            doc.add_heading('Essential Questions', level=1)
            for q in plan['essential_questions']:
                doc.add_paragraph(f"* {q}")

        # Daily Plans
        if plan.get('days'):
            doc.add_heading('Daily Lesson Plans', level=1)

            for day in plan['days']:
                doc.add_heading(f"Day {day.get('day')}: {day.get('topic')}", level=2)

                if day.get('objective'):
                    p = doc.add_paragraph()
                    p.add_run('Learning Objective: ').bold = True
                    p.add_run(day['objective'])

                if day.get('standards_addressed'):
                    p = doc.add_paragraph()
                    p.add_run('Standards: ').bold = True
                    p.add_run(', '.join(day['standards_addressed']))

                if day.get('timing'):
                    doc.add_heading('Lesson Timing', level=3)
                    for t in day['timing']:
                        time_str = t.get('minutes') or t.get('duration', '')
                        doc.add_paragraph(f"{time_str} - {t.get('activity', '')}: {t.get('description', '')}")

                vocab_text = format_vocab(day.get('vocabulary'))
                if vocab_text:
                    doc.add_heading('Vocabulary', level=3)
                    doc.add_paragraph(vocab_text)

                br_text = format_bell_ringer(day.get('bell_ringer'))
                if br_text:
                    doc.add_heading('Bell Ringer', level=3)
                    doc.add_paragraph(br_text)

                if day.get('direct_instruction'):
                    di = day['direct_instruction']
                    doc.add_heading('Direct Instruction', level=3)
                    if di.get('key_points'):
                        doc.add_paragraph('Key Points:')
                        for kp in di['key_points']:
                            doc.add_paragraph(f"* {kp}")
                    if di.get('check_for_understanding'):
                        doc.add_paragraph('\nCheck for Understanding:')
                        for q in di['check_for_understanding']:
                            doc.add_paragraph(f"* \"{q}\"")

                act_text = format_activity(day.get('activity'))
                if act_text:
                    doc.add_heading('Main Activity', level=3)
                    doc.add_paragraph(act_text)

                asmt_text = format_assessment(day.get('assessment'))
                if asmt_text:
                    doc.add_heading('Assessment', level=3)
                    doc.add_paragraph(asmt_text)

                if day.get('materials'):
                    doc.add_heading('Materials', level=3)
                    doc.add_paragraph(', '.join(day['materials']))

                if day.get('homework'):
                    doc.add_heading('Homework', level=3)
                    doc.add_paragraph(day['homework'])

                if day.get('teacher_notes'):
                    doc.add_heading('Teacher Notes', level=3)
                    doc.add_paragraph(day['teacher_notes'])

                doc.add_paragraph()

        # Unit Assessment
        if plan.get('unit_assessment'):
            doc.add_heading('Summative Assessment', level=1)
            ua = plan['unit_assessment']
            if isinstance(ua, dict):
                if ua.get('type'):
                    doc.add_paragraph(f"Type: {ua['type']}")
                if ua.get('description'):
                    doc.add_paragraph(ua['description'])
                if ua.get('components'):
                    doc.add_paragraph('\nComponents:')
                    for c in ua['components']:
                        doc.add_paragraph(f"* {c}")
            else:
                doc.add_paragraph(str(ua))

        # Resources
        if plan.get('resources'):
            doc.add_heading('Resources', level=1)
            for r in plan['resources']:
                doc.add_paragraph(f"* {r}")

        # Save file
        filename = f"Lesson_Plan_{int(time.time())}.docx"
        output_folder = graider_export_dir()
        os.makedirs(output_folder, exist_ok=True)
        filepath = os.path.join(output_folder, filename)
        doc.save(filepath)

        # Open the file
        if sys.platform == 'darwin':
            subprocess.run(['open', filepath], check=False)

        return jsonify({"status": "success", "path": filepath})

    except Exception as e:
        _logger.exception("Request failed: %s", request.path)
        return jsonify({"error": "An internal error occurred"}), 500


def _save_grading_config_for_export(assignment):
    """Auto-save a grading config so the Graider extraction pipeline can find expected answers.

    Builds gradingNotes from the answer key and saves to ~/.graider_assignments/.
    """
    try:
        title = assignment.get('title', 'Untitled')
        safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()
        if not safe_title:
            return

        config_dir = os.path.expanduser("~/.graider_assignments")
        os.makedirs(config_dir, exist_ok=True)

        # Build grading notes from answers
        grading_lines = []
        sections = assignment.get('sections', [])
        answer_key = assignment.get('answer_key', {})

        # Try answer_key dict first (assessment format)
        if answer_key:
            for q_num, answer_data in sorted(answer_key.items(), key=lambda x: int(x[0]) if str(x[0]).isdigit() else 0):
                if isinstance(answer_data, dict):
                    ans = answer_data.get('answer', '')
                else:
                    ans = str(answer_data)
                grading_lines.append(f"Q{q_num}: {ans}")
        else:
            # Fall back to extracting from sections (assignment format)
            q_idx = 1
            for section in sections:
                for q in section.get('questions', []):
                    answer = q.get('answer', '')
                    if answer:
                        grading_lines.append(f"Q{q.get('number', q_idx)}: {answer}")
                    q_idx += 1

        # Build plain-text and HTML representations for grading setup display
        doc_lines = [title, '', 'Name: ' + '_' * 50, 'Date: ' + '_' * 50,
                     'Period: ' + '_' * 50, '']
        html_parts = [
            '<style>',
            'body { font-family: Georgia, serif; line-height: 1.6; }',
            'table { border-collapse: collapse; width: 100%; margin: 15px 0; }',
            'td, th { border: 1px solid #ccc; padding: 8px 12px; text-align: left; }',
            'th { background: #f5f5f5; font-weight: bold; }',
            'p { margin: 10px 0; }',
            'h1, h2, h3 { margin: 20px 0 10px 0; }',
            '</style>',
            '<h1><strong>' + title + '</strong></h1>',
            '<p><strong>Name: </strong>' + '_' * 50 + '</p>',
            '<p><strong>Date: </strong>' + '_' * 50 + '</p>',
            '<p><strong>Period: </strong>' + '_' * 50 + '</p>',
        ]
        instructions = assignment.get('instructions', '')
        if instructions:
            doc_lines.append(instructions)
            doc_lines.append('')
            html_parts.append('<p><em>' + instructions + '</em></p>')
        q_num = 1
        for section in sections:
            sec_title = section.get('title', '')
            if sec_title:
                doc_lines.append(sec_title.upper())
                html_parts.append('<h2><strong>' + sec_title.upper() + '</strong></h2>')
            for q in section.get('questions', []):
                qtype = q.get('question_type', '')
                if qtype == 'vocab_term':
                    term = q.get('term', q.get('question', ''))
                    doc_lines.append(term + ': ' + '_' * 60)
                    html_parts.append('<table><tr><td><p><strong>' + term + ':</strong> ' + '_' * 60 + '</p></td></tr></table>')
                elif qtype in ('fill_in_blank', 'fitb'):
                    pts = str(q.get('points', 5))
                    qtxt = q.get('question', '')
                    doc_lines.append(str(q_num) + ') ' + qtxt + '  (' + pts + ' pts)')
                    doc_lines.append('Answer: ' + '_' * 55)
                    html_parts.append(
                        '<table><tr><td><p>[GRAIDER:QUESTION:' + str(q_num) + ']<strong>  '
                        + str(q_num) + ') ' + qtxt + '</strong>  (' + pts + ' pts)</p></td></tr>'
                        + '<tr><td><p><strong>Answer:</strong></p><p><em>Type your answer here...</em></p></td></tr></table>'
                    )
                else:
                    pts = str(q.get('points', 10))
                    qtxt = q.get('question', '')
                    doc_lines.append(str(q_num) + ') ' + qtxt + '  (' + pts + ' pts)')
                    doc_lines.append('Response: ' + '_' * 55)
                    doc_lines.append('_' * 65)
                    html_parts.append(
                        '<table><tr><td><p>[GRAIDER:QUESTION:' + str(q_num) + ']<strong>  '
                        + str(q_num) + ') ' + qtxt + '</strong>  (' + pts + ' pts)</p></td></tr>'
                        + '<tr><td><p><strong>Your Answer:</strong></p><p><em>Type your answer here...</em></p></td></tr></table>'
                    )
                doc_lines.append('')
                q_num += 1

        doc_text = '\n'.join(doc_lines)
        doc_html = '\n'.join(html_parts)

        config = {
            "title": title,
            "gradingNotes": "\n".join(grading_lines),
            "tableStructured": True,
            "tableVersion": "v1",
            "totalPoints": assignment.get('total_points', 100),
            "subject": assignment.get('subject', ''),
            "grade": assignment.get('grade', ''),
            "importedDoc": {
                "text": doc_text,
                "html": doc_html,
                "filename": safe_title + "_Student.docx",
                "loading": False
            },
        }

        config_path = os.path.join(config_dir, f"{safe_title}.json")
        with open(config_path, 'w') as f:
            json.dump(config, f, indent=2)

        # Also save to Supabase storage if available
        if storage_save:
            try:
                from flask import g
                teacher_id = getattr(g, 'user_id', 'local-dev')
                storage_save(f'assignment:{safe_title}', config, teacher_id)
            except Exception:
                pass  # Local save succeeded, Supabase is best-effort

        _logger.info("Saved grading config: %s", config_path)
    except Exception as e:
        _logger.warning("Could not save grading config: %s", e)


@planner_bp.route('/api/export-generated-assignment', methods=['POST'])
@require_teacher
@handle_route_errors
def export_generated_assignment():
    """Export a generated assignment to PDF or DOCX (with Graider tables) format."""
    data = request.json
    assignment = data.get('assignment', {})

    if not isinstance(assignment, dict) or not assignment.get('sections'):
        return jsonify({"error": "Nothing to export: the assignment has no questions."}), 400

    format_type = data.get('format', 'pdf')  # Default to PDF now
    include_answers = data.get('include_answers', False)

    title = assignment.get('title', 'Assignment')
    instructions = assignment.get('instructions', '')
    sections = assignment.get('sections', [])
    total_points = assignment.get('total_points', 100)
    time_estimate = assignment.get('time_estimate', '')
    teacher_name = data.get('teacher_name', '')
    subject_name = data.get('subject', '')

    # Inject teacher/subject into assignment for export functions
    if teacher_name:
        assignment['teacher_name'] = teacher_name
    if subject_name:
        assignment['subject'] = subject_name

    # DOCX with Graider tables for student worksheets
    if format_type == 'docx' and not include_answers:
        try:
            safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()
            output_folder = graider_export_dir("Assignments")
            os.makedirs(output_folder, exist_ok=True)
            filepath = _export_assignment_docx_graider(assignment, output_folder, safe_title)
            _save_grading_config_for_export(assignment)
            if sys.platform == 'darwin':
                subprocess.run(['open', filepath], check=False)
            return jsonify({"status": "success", "path": filepath})
        except Exception as e:
            _logger.exception("Request failed: %s", request.path)
            return jsonify({"error": "An internal error occurred"}), 500

    # PDF path (answer keys and fallback)
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.units import inch
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.colors import black, gray, lightgrey, red, green
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Image,
            Table, TableStyle, PageBreak, KeepTogether, Flowable
        )
        from reportlab.lib.colors import white as rl_white
        import io

        # Bubble circle for MC/TF answer sheets
        class _BubbleCircle(Flowable):
            """Draws a circle bubble — empty (outline) or filled (solid black)."""
            def __init__(self, filled=False, size=9):
                Flowable.__init__(self)
                self.filled = filled
                self.size = size
                self.width = size + 4
                self.height = size + 4
            def draw(self):
                r = self.size / 2
                cx = r + 2
                cy = r + 2
                from reportlab.lib.colors import Color
                self.canv.setStrokeColor(Color(0.3, 0.3, 0.3))
                self.canv.setLineWidth(1.2)
                if self.filled:
                    self.canv.setFillColor(black)
                    self.canv.circle(cx, cy, r, fill=1)
                else:
                    self.canv.setFillColor(rl_white)
                    self.canv.circle(cx, cy, r, fill=1, stroke=1)

        # Set up styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle', parent=styles['Heading1'],
            alignment=TA_CENTER, fontSize=18, spaceAfter=6
        )
        heading_style = ParagraphStyle(
            'CustomHeading', parent=styles['Heading2'],
            fontSize=14, spaceAfter=6, spaceBefore=12
        )
        normal_style = styles['Normal']
        bold_style = ParagraphStyle(
            'Bold', parent=styles['Normal'],
            fontName='Helvetica-Bold'
        )
        center_style = ParagraphStyle(
            'Center', parent=styles['Normal'],
            alignment=TA_CENTER
        )
        answer_style = ParagraphStyle(
            'Answer', parent=styles['Normal'],
            fontName='Helvetica-Bold', textColor=green
        )

        # Helper: convert Unicode subscript/superscript to ReportLab XML tags
        _sub_map = str.maketrans('₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎', '0123456789+-=()')
        _sup_map = str.maketrans('⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾', '0123456789+-=()')
        def _fix_sub_sup(text):
            """Replace Unicode sub/superscript chars with ReportLab <sub>/<sup> tags."""
            if not text:
                return text
            import re as _re
            # Subscripts: ₀-₉ and related
            def _replace_sub(m):
                return '<sub>' + m.group(0).translate(_sub_map) + '</sub>'
            def _replace_sup(m):
                return '<sup>' + m.group(0).translate(_sup_map) + '</sup>'
            text = _re.sub('[₀₁₂₃₄₅₆₇₈₉₊₋₌₍₎]+', _replace_sub, text)
            text = _re.sub('[⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾]+', _replace_sup, text)
            return text

        # Build the PDF
        safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()
        suffix = "_ANSWER_KEY" if include_answers else "_Student"
        filename = f"{safe_title}{suffix}.pdf"
        output_folder = graider_export_dir("Assignments")
        os.makedirs(output_folder, exist_ok=True)
        filepath = os.path.join(output_folder, filename)

        doc = SimpleDocTemplate(
            filepath, pagesize=letter,
            topMargin=0.5*inch, bottomMargin=0.5*inch,
            leftMargin=0.75*inch, rightMargin=0.75*inch
        )

        story = []

        # Teacher name / subject header
        _teacher = assignment.get('teacher_name', '')
        _subject = assignment.get('subject', '')
        if _teacher or _subject:
            header_parts = []
            if _teacher:
                header_parts.append(_teacher)
            if _subject:
                header_parts.append(_subject)
            header_text = "  |  ".join(header_parts)
            header_style = ParagraphStyle(
                'TeacherHeader', parent=styles['Normal'],
                alignment=TA_CENTER, fontSize=11, textColor=gray
            )
            story.append(Paragraph(header_text, header_style))
            story.append(Spacer(1, 0.05*inch))

        # Title
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 0.1*inch))

        # Name/Date/Period or Answer Key header
        if include_answers:
            story.append(Paragraph(
                "<b>ANSWER KEY - FOR TEACHER USE ONLY</b>",
                center_style
            ))
        else:
            story.append(Paragraph(
                "Name: _______________________  Date: _______________  Period: _____",
                center_style
            ))

        # Total points only (no time limit on assignments)
        if total_points:
            story.append(Spacer(1, 0.1*inch))
            story.append(Paragraph(f"Total Points: {total_points}", center_style))

        story.append(Spacer(1, 0.15*inch))

        # Instructions
        if instructions:
            story.append(Paragraph(f"<b>Instructions:</b> {instructions}", normal_style))
            story.append(Spacer(1, 0.15*inch))

        question_num = 1

        # Process sections
        for section in sections:
            section_name = section.get('name', 'Section')
            section_points = section.get('points', 0)
            section_type = section.get('type', 'short_answer')
            questions = section.get('questions', [])

            # Section header
            pts_text = f" ({section_points} points)" if section_points else ""
            story.append(Paragraph(f"<b>{section_name}</b>{pts_text}", heading_style))
            story.append(Spacer(1, 0.1*inch))  # Space between section header and questions

            for q in questions:
                q_number = q.get('number', question_num)
                q_text = _fix_sub_sup(q.get('question', ''))
                q_points = q.get('points', 0)
                q_options = [_fix_sub_sup(o) for o in q.get('options', [])]
                q_answer = q.get('answer', '')
                q_type = q.get('question_type', section_type)
                q_visual = q.get('visual_type', None)  # number_line, coordinate_plane, etc.

                # Inject True/False options if missing (safety net for older
                # assignments generated before the hydrator inject landed; new
                # assignments get options populated in
                # assignment_post_processing._hydrate_question).
                if not q_options and q_type in ('true_false', 'tf'):
                    q_options = ['True', 'False']

                # Question text — detect and render inline markdown tables
                pts_text = f" ({q_points} pts)" if q_points else ""
                table_parts = _split_markdown_table(q_text)
                if table_parts:
                    # Text before table
                    before_text = table_parts['before'].strip()
                    combined_before = f"<b>Question {q_number}:</b> {before_text}{pts_text}" if before_text else f"<b>Question {q_number}:</b>{pts_text}"
                    story.append(Paragraph(combined_before, normal_style))
                    story.append(Spacer(1, 0.05*inch))
                    # Render the table
                    md_table = table_parts['table']
                    t_data = [md_table['headers']] + md_table['rows']
                    col_count = len(md_table['headers'])
                    col_w = min(1.2*inch, (6.5*inch) / max(col_count, 1))
                    from reportlab.lib import colors as rl_colors
                    tbl = Table(t_data, colWidths=[col_w]*col_count)
                    tbl.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), rl_colors.Color(0.9, 0.9, 0.95)),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('GRID', (0, 0), (-1, -1), 0.5, rl_colors.Color(0.6, 0.6, 0.6)),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('TOPPADDING', (0, 0), (-1, -1), 4),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 4),
                    ]))
                    story.append(tbl)
                    # Text after table
                    if table_parts.get('after', '').strip():
                        story.append(Paragraph(table_parts['after'].strip(), normal_style))
                    story.append(Spacer(1, 0.05*inch))
                else:
                    story.append(Paragraph(
                        f"<b>Question {q_number}:</b> {q_text}{pts_text}",
                        normal_style
                    ))
                    story.append(Spacer(1, 0.05*inch))

                # Multiple choice options with bubble circles
                if q_options and q_type != 'matching':
                    is_tf = q_type in ('true_false', 'tf')
                    # Determine correct answer index for answer key
                    correct_idx = None
                    if include_answers and q_answer:
                        if is_tf:
                            for oi, opt in enumerate(q_options):
                                if opt.lower().strip() == str(q_answer).lower().strip():
                                    correct_idx = oi
                                    break
                        else:
                            from backend.services.worksheet_generator import _normalize_correct_answer_to_letter
                            letter = _normalize_correct_answer_to_letter(q_answer, q_options)
                            if letter:
                                correct_idx = ord(letter) - ord('A')

                    for oi, opt in enumerate(q_options):
                        is_filled = (include_answers and correct_idx is not None and oi == correct_idx)
                        bubble = _BubbleCircle(filled=is_filled, size=9)
                        opt_para = Paragraph(opt, normal_style)
                        row_table = Table(
                            [[bubble, opt_para]],
                            colWidths=[0.5*inch, 5.5*inch],
                        )
                        row_table.setStyle(TableStyle([
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('ALIGN', (0, 0), (0, 0), 'CENTER'),
                            ('LEFTPADDING', (0, 0), (0, 0), 20),
                            ('LEFTPADDING', (1, 0), (1, 0), 0),
                            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
                            ('TOPPADDING', (0, 0), (-1, -1), 2),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
                        ]))
                        story.append(row_table)

                # Matching question: render terms and definitions columns
                q_terms = q.get('terms', [])
                q_definitions = q.get('definitions', [])
                if (q_type == 'matching' or (q_terms and q_definitions)):
                    q_terms = [_fix_sub_sup(str(t)) for t in q_terms]
                    q_definitions = [_fix_sub_sup(str(d)) for d in q_definitions]
                    from reportlab.lib import colors as rl_colors
                    import random as _random
                    # Shuffle definitions for student version
                    shuffled_defs = list(q_definitions)
                    if not include_answers:
                        _random.seed(q_number)  # Deterministic shuffle per question
                        _random.shuffle(shuffled_defs)
                    # Build two-column table: numbered terms | lettered definitions
                    # Use Paragraph objects for definitions so they word-wrap
                    def_style = ParagraphStyle('MatchDef', parent=normal_style, fontSize=10)
                    term_style = ParagraphStyle('MatchTerm', parent=normal_style, fontSize=10)
                    max_rows = max(len(q_terms), len(shuffled_defs))
                    match_data = [['', 'Terms', '', 'Definitions']]
                    for ri in range(max_rows):
                        term_num = str(ri + 1) + '.' if ri < len(q_terms) else ''
                        term_text = Paragraph(q_terms[ri], term_style) if ri < len(q_terms) else ''
                        def_letter = chr(65 + ri) + '.' if ri < len(shuffled_defs) else ''
                        def_text = Paragraph(shuffled_defs[ri], def_style) if ri < len(shuffled_defs) else ''
                        match_data.append([term_num, term_text, def_letter, def_text])
                    col_widths = [0.3*inch, 2.2*inch, 0.3*inch, 3.5*inch]
                    match_tbl = Table(match_data, colWidths=col_widths)
                    match_tbl.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), rl_colors.Color(0.9, 0.9, 0.95)),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('ALIGN', (0, 0), (0, -1), 'RIGHT'),
                        ('ALIGN', (2, 0), (2, -1), 'RIGHT'),
                        ('GRID', (0, 0), (-1, -1), 0.5, rl_colors.Color(0.6, 0.6, 0.6)),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                        ('TOPPADDING', (0, 0), (-1, -1), 6),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                    ]))
                    story.append(Spacer(1, 0.05*inch))
                    story.append(match_tbl)
                    story.append(Spacer(1, 0.05*inch))

                # Add visual elements based on question type
                if q_visual or q_type in ['number_line', 'coordinate_plane', 'graph',
                                          'geometry', 'triangle', 'rectangle', 'regular_polygon',
                                          'circle', 'trapezoid', 'parallelogram',
                                          'rectangular_prism', 'cylinder',
                                          'pythagorean', 'trig', 'angles', 'similarity',
                                          'box_plot', 'bar_chart', 'function_graph',
                                          'dot_plot', 'stem_and_leaf', 'unit_circle',
                                          'transformations', 'fraction_model',
                                          'probability_tree', 'tape_diagram',
                                          'venn_diagram', 'protractor', 'angle_protractor',
                                          'histogram', 'pie_chart']:
                    visual_image = _create_visual_for_question(q, include_answers)
                    if visual_image:
                        story.append(Spacer(1, 0.1*inch))
                        story.append(visual_image)
                        story.append(Spacer(1, 0.1*inch))

                # Answer section
                if include_answers:
                    # MC/TF: filled bubble already shows the answer — skip text label
                    if q_options and q_type != 'matching':
                        pass  # Bubble is already filled above
                    elif q_type == 'matching' or (q_terms and q_definitions):
                        if isinstance(q_answer, dict):
                            ans_parts = [f"{k} → {v}" for k, v in q_answer.items()]
                            ans_text = "ANSWERS: " + " | ".join(ans_parts)
                        else:
                            ans_text = f"ANSWER: {q_answer}"
                        story.append(Paragraph(f"<b>{_fix_sub_sup(str(ans_text))}</b>", answer_style))
                    elif q_type == 'coordinates' and isinstance(q_answer, dict):
                        ans_text = f"ANSWER: Lat: {q_answer.get('lat', 0)}, Lng: {q_answer.get('lng', 0)}"
                        story.append(Paragraph(f"<b>{_fix_sub_sup(str(ans_text))}</b>", answer_style))
                    else:
                        ans_text = f"ANSWER: {q_answer}"
                        story.append(Paragraph(f"<b>{_fix_sub_sup(str(ans_text))}</b>", answer_style))

                    if q_type == 'math_equation':
                        story.append(Paragraph("<i>(Equivalent forms accepted)</i>", normal_style))
                    elif q_type == 'coordinates':
                        tolerance_km = q.get('tolerance_km', 50)
                        story.append(Paragraph(f"<i>(Acceptable within {tolerance_km} km)</i>", normal_style))
                else:
                    # Answer space for students
                    if q_type == 'matching' or (q_terms and q_definitions):
                        pass  # Match table already has answer blanks
                    elif q_type == 'math_equation':
                        story.append(Paragraph("Show your work:", normal_style))
                        for _ in range(3):
                            story.append(Paragraph("_" * 85, normal_style))
                        story.append(Paragraph("<b>Final Answer:</b> " + "_" * 50, normal_style))
                    elif q_type == 'coordinates':
                        story.append(Paragraph(
                            "<b>Latitude:</b> _______________°  <b>Longitude:</b> _______________°",
                            normal_style
                        ))
                    elif q_type == 'data_table':
                        # Create empty table — handle both normalized and raw AI field names
                        headers = q.get('headers', q.get('column_headers', ['Column 1', 'Column 2', 'Column 3']))
                        row_labels = q.get('row_labels', [])
                        expected = q.get('expected_data', [])
                        num_rows = q.get('num_rows', len(expected) if expected else 5)
                        if row_labels:
                            table_data = [[''] + headers]
                            for ri in range(num_rows):
                                label = row_labels[ri] if ri < len(row_labels) else ''
                                table_data.append([label] + [''] * len(headers))
                        else:
                            table_data = [headers] + [[''] * len(headers) for _ in range(num_rows)]
                        col_count = len(table_data[0])
                        t = Table(table_data, colWidths=[1.5*inch] * col_count)
                        t.setStyle(TableStyle([
                            ('GRID', (0, 0), (-1, -1), 1, black),
                            ('BACKGROUND', (0, 0), (-1, 0), lightgrey),
                            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                            ('FONTSIZE', (0, 0), (-1, -1), 10),
                            ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                            ('TOPPADDING', (0, 0), (-1, -1), 12),
                        ]))
                        story.append(t)
                    elif section_type in ['essay', 'extended_response']:
                        for _ in range(8):
                            story.append(Paragraph("_" * 85, normal_style))
                    elif section_type == 'short_answer':
                        for _ in range(3):
                            story.append(Paragraph("_" * 85, normal_style))
                    elif section_type in ['multiple_choice', 'true_false']:
                        pass  # Bubbles are the answer — no separate answer line needed
                    else:
                        for _ in range(2):
                            story.append(Paragraph("_" * 85, normal_style))

                story.append(Spacer(1, 0.15*inch))
                question_num += 1

        # Rubric for teacher version
        if include_answers and assignment.get('rubric', {}).get('criteria'):
            story.append(PageBreak())
            story.append(Paragraph("<b>Grading Rubric</b>", heading_style))
            for criterion in assignment['rubric']['criteria']:
                story.append(Paragraph(
                    f"<b>{criterion.get('name', 'Criterion')}:</b> "
                    f"{criterion.get('points', 0)} points - {criterion.get('description', '')}",
                    normal_style
                ))
                story.append(Spacer(1, 0.05*inch))

        # Build PDF
        doc.build(story)

        # Open the file
        if sys.platform == 'darwin':
            subprocess.run(['open', filepath], check=False)

        return jsonify({"status": "success", "path": filepath})

    except Exception as e:
        _logger.exception("Request failed: %s", request.path)
        return jsonify({"error": "An internal error occurred"}), 500


# =============================================================================
# ASSESSMENT GENERATION
# =============================================================================

@planner_bp.route('/api/generate-assessment', methods=['POST'])
@limiter.limit("10 per minute")
@require_teacher
@handle_route_errors
def generate_assessment():
    """
    Generate a standards-aligned assessment with DOK level distribution.

    Request body:
    {
        "standards": [{"code": "SS.8.A.1.1", "benchmark": "...", "dok": 2, ...}],
        "config": {
            "grade": "8",
            "subject": "US History",
            "teacher_name": "Mr. Smith"
        },
        "assessmentConfig": {
            "type": "quiz",  // quiz, test, benchmark, formative
            "title": "Chapter 5 Assessment",
            "totalQuestions": 15,
            "questionTypes": {
                "multiple_choice": 10,
                "short_answer": 3,
                "extended_response": 2
            },
            "dokDistribution": {
                "1": 3,   // 3 DOK 1 questions
                "2": 6,   // 6 DOK 2 questions
                "3": 4,   // 4 DOK 3 questions
                "4": 2    // 2 DOK 4 questions
            },
            "includeAnswerKey": true,
            "includeStandardsReference": true
        }
    }
    """
    data = request.json
    standards = data.get('standards', [])
    config = data.get('config', {})
    assessment_config = data.get('assessmentConfig', {})
    content_only = config.get('contentOnly', False)

    # Normalize standards: accept both dicts and plain code strings
    normalized = []
    for s in standards:
        if isinstance(s, str):
            normalized.append({'code': s})
        elif isinstance(s, dict):
            normalized.append(s)
    standards = normalized

    if not standards:
        return jsonify({"error": "No standards provided"})

    _subject = config.get('subject', '').strip()
    _grade = config.get('grade', '').strip()
    if not _subject or not _grade:
        from flask import current_app
        current_app.logger.warning(
            f"Assessment requested without subject/grade: subject={_subject!r}, grade={_grade!r}")

    try:
        from backend.api_keys import get_api_key as _gak
        api_key = _gak('openai', getattr(g, 'user_id', 'local-dev'))
        return jsonify(generate_assessment_content(
            standards=standards, config=config, assessment_config=assessment_config,
            content_only=content_only, content_sources=data.get('contentSources', []),
            api_key=api_key, openai_context=_get_openai_context()))
    except Exception as e:
        error_msg = str(e)
        _logger.exception("Assessment Generation Error")
        return jsonify({"error": f"Failed to generate assessment: {error_msg}"}), 500


@planner_bp.route('/api/export-assessment', methods=['POST'])
@require_teacher
@handle_route_errors
def export_assessment():
    """Export assessment to Word document with Graider table extraction tags."""
    data = request.json
    assessment = data.get('assessment', {})
    include_answer_key = data.get('includeAnswerKey', False)

    if not assessment:
        return jsonify({"error": "No assessment data provided"})

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from backend.services.worksheet_generator import _add_graider_table, _add_graider_marker
        import tempfile
        import base64

        doc = Document()

        # Graider table style dict (blue header, white text)
        graider_style = {
            "table_header_bg": "#4472C4",
            "table_header_text_color": "#FFFFFF",
        }

        # Title
        title = doc.add_heading(assessment.get('title', 'Assessment'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header info
        header_info = doc.add_paragraph()
        header_info.add_run(f"Subject: {assessment.get('subject', '')}").bold = True
        header_info.add_run(f"    Grade: {assessment.get('grade', '')}")
        header_info.add_run(f"    Time: {assessment.get('time_estimate', '')}")
        header_info.add_run(f"    Total Points: {assessment.get('total_points', '')}")

        # Student info line
        doc.add_paragraph("Name: _________________________    Date: _____________    Period: _____")

        # Instructions
        if assessment.get('instructions'):
            inst = doc.add_paragraph()
            inst.add_run("Instructions: ").bold = True
            inst.add_run(assessment.get('instructions'))

        doc.add_paragraph()  # Space

        # Sections
        for section in assessment.get('sections', []):
            # Section header
            sec_head = doc.add_heading(section.get('name', 'Section'), level=1)

            if section.get('instructions'):
                sec_inst = doc.add_paragraph()
                sec_inst.add_run(section.get('instructions')).italic = True

            # Questions
            for q in section.get('questions', []):
                q_num = q.get('number', '')
                q_text = q.get('question', '')
                q_points = q.get('points', 1)
                q_type = q.get('type', section.get('type', ''))

                if q.get('options'):
                    # MC/TF: render question + options as paragraphs, then small Graider table
                    q_para = doc.add_paragraph()
                    q_para.add_run(f"{q_num}. ").bold = True
                    q_para.add_run(f"{q_text} ")
                    q_para.add_run(f"({q_points} pt{'s' if q_points > 1 else ''})").italic = True

                    for opt in q.get('options', []):
                        opt_para = doc.add_paragraph(f"    {opt}")
                        opt_para.paragraph_format.space_before = Pt(2)
                        opt_para.paragraph_format.space_after = Pt(2)

                    _add_graider_table(doc, f"Answer for Question {q_num}",
                                       f"GRAIDER:QUESTION:{q_num}", q_points,
                                       graider_style, 720)  # 0.5 inch

                elif q.get('terms') and q.get('definitions'):
                    # Matching: two-column table with draw-a-line instruction
                    q_para = doc.add_paragraph()
                    q_para.add_run(f"{q_num}. ").bold = True
                    q_para.add_run(f"{q_text} ")
                    q_para.add_run(f"({q_points} pt{'s' if q_points > 1 else ''})").italic = True

                    inst_para = doc.add_paragraph()
                    inst_para.add_run("Directions: ").bold = True
                    inst_para.add_run("Draw a line from each term to its matching definition.")
                    inst_para.italic = True

                    terms = q.get('terms', [])
                    definitions = q.get('definitions', [])
                    max_len = max(len(terms), len(definitions))

                    tbl = doc.add_table(rows=max_len + 1, cols=3)
                    tbl.style = 'Table Grid'

                    # Header row
                    tbl.rows[0].cells[0].text = "Term"
                    tbl.rows[0].cells[1].text = ""
                    tbl.rows[0].cells[2].text = "Definition"
                    for cell in tbl.rows[0].cells:
                        for paragraph in cell.paragraphs:
                            for run_obj in paragraph.runs:
                                run_obj.bold = True

                    # Set narrow middle column for drawing lines
                    for row in tbl.rows:
                        row.cells[1].width = Inches(0.5)

                    for i in range(max_len):
                        if i < len(terms):
                            tbl.rows[i + 1].cells[0].text = f"{i + 1}. {terms[i]}"
                        if i < len(definitions):
                            letter_char = chr(65 + i)
                            tbl.rows[i + 1].cells[2].text = f"{letter_char}. {definitions[i]}"

                    doc.add_paragraph()  # Space after table

                elif q_type == 'extended_response':
                    # Extended response: Graider table with 3" height
                    _add_graider_table(doc, f"{q_num}. {q_text} ({q_points} pts)",
                                       f"GRAIDER:QUESTION:{q_num}", q_points,
                                       graider_style, 4320)  # 3 inches

                else:
                    # Short answer / default: Graider table with 1.5" height
                    _add_graider_table(doc, f"{q_num}. {q_text} ({q_points} pts)",
                                       f"GRAIDER:QUESTION:{q_num}", q_points,
                                       graider_style, 2160)  # 1.5 inches

        # Add Graider marker before answer key
        _add_graider_marker(doc)

        # Answer Key (separate page) — no Graider tables
        if include_answer_key:
            doc.add_page_break()
            doc.add_heading("Answer Key", 0)

            answer_key = assessment.get('answer_key', {})
            for q_num, answer_data in sorted(answer_key.items(), key=lambda x: int(x[0]) if x[0].isdigit() else 0):
                ans_para = doc.add_paragraph()
                ans_para.add_run(f"{q_num}. ").bold = True

                if isinstance(answer_data, dict):
                    ans_para.add_run(str(answer_data.get('answer', '')))
                    if answer_data.get('explanation'):
                        ans_para.add_run(f" - {answer_data.get('explanation')}")
                else:
                    ans_para.add_run(str(answer_data))

        # Auto-save grading config for structured extraction
        _save_grading_config_for_export(assessment)

        # Save to temp file and encode
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            doc.save(tmp.name)
            tmp.seek(0)
            with open(tmp.name, 'rb') as f:
                doc_bytes = f.read()
            os.unlink(tmp.name)

        doc_base64 = base64.b64encode(doc_bytes).decode('utf-8')

        # Generate filename
        safe_title = ''.join(c if c.isalnum() or c in ' -_' else '' for c in assessment.get('title', 'Assessment'))
        filename = f"{safe_title.replace(' ', '_')}.docx"

        return jsonify({
            "document": doc_base64,
            "filename": filename,
            "format": "docx"
        })

    except Exception as e:
        _logger.exception("Request failed: %s", request.path)
        return jsonify({"error": "An internal error occurred"}), 500


# =============================================================================
# ASSESSMENT PLATFORM TEMPLATES
# =============================================================================

TEMPLATES_DIR = os.path.expanduser("~/.graider_data/assessment_templates")


@planner_bp.route('/api/upload-assessment-template', methods=['POST'])
@require_teacher
@handle_route_errors
def upload_assessment_template():
    """Upload a sample template from an assessment platform (e.g., Wayground, Canvas)."""
    import uuid

    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400

    file = request.files['file']
    platform = request.form.get('platform', 'custom')
    name = request.form.get('name', 'Untitled Template')

    if file.filename == '':
        return jsonify({"error": "No file selected"}), 400

    ext = os.path.splitext(file.filename)[1].lower() if file.filename else ''
    if ext not in ALLOWED_DOC_EXTENSIONS:
        return jsonify({"error": f"Unsupported file type: {ext}"}), 400

    # Create templates directory if it doesn't exist
    os.makedirs(TEMPLATES_DIR, exist_ok=True)

    # Generate unique ID
    template_id = str(uuid.uuid4())[:8]

    # Save the file
    filename = f"{template_id}_{platform}{ext}"
    filepath = os.path.join(TEMPLATES_DIR, secure_filename(filename))
    file.save(filepath)

    # Parse the template to understand its structure
    template_structure = parse_template_structure(filepath, ext)

    # Save metadata
    metadata = {
        "id": template_id,
        "name": name,
        "platform": platform,
        "filename": filename,
        "filepath": filepath,
        "original_filename": file.filename,
        "extension": ext,
        "structure": template_structure,
        "created_at": time.strftime('%Y-%m-%d %H:%M:%S')
    }

    metadata_path = os.path.join(TEMPLATES_DIR, f"{template_id}.meta.json")
    with open(metadata_path, 'w') as f:
        json.dump(metadata, f, indent=2)

    return jsonify({
        "success": True,
        "template": metadata,
        "message": f"Template '{name}' uploaded successfully"
    })


@planner_bp.route('/api/assessment-templates', methods=['GET'])
@require_teacher
@handle_route_errors
def get_assessment_templates():
    """Get all uploaded assessment templates."""
    templates = []

    if not os.path.exists(TEMPLATES_DIR):
        return jsonify({"templates": []})

    for f in os.listdir(TEMPLATES_DIR):
        if f.endswith('.meta.json'):
            try:
                with open(os.path.join(TEMPLATES_DIR, f), 'r') as mf:
                    metadata = json.load(mf)
                    templates.append(metadata)
            except Exception:
                pass

    # Sort by creation date (newest first)
    templates.sort(key=lambda x: x.get('created_at', ''), reverse=True)

    return jsonify({"templates": templates})


@planner_bp.route('/api/assessment-template/<template_id>', methods=['DELETE'])
@require_teacher
@handle_route_errors
def delete_assessment_template(template_id):
    """Delete an assessment template."""
    if not os.path.exists(TEMPLATES_DIR):
        return jsonify({"error": "Template not found"}), 404

    # Find and delete metadata and file
    meta_path = os.path.join(TEMPLATES_DIR, f"{template_id}.meta.json")

    if not os.path.exists(meta_path):
        return jsonify({"error": "Template not found"}), 404

    try:
        with open(meta_path, 'r') as f:
            metadata = json.load(f)

        # Delete the template file
        if metadata.get('filepath') and os.path.exists(metadata['filepath']):
            os.remove(metadata['filepath'])

        # Delete metadata
        os.remove(meta_path)

        return jsonify({"success": True, "message": "Template deleted"})

    except Exception as e:
        _logger.exception("Failed to delete template")
        return jsonify({"error": "An internal error occurred"}), 500


@planner_bp.route('/api/export-assessment-platform', methods=['POST'])
@require_teacher
@handle_route_errors
def export_assessment_for_platform():
    """Export assessment in a specific platform's format."""
    data = request.json
    assessment = data.get('assessment', {})
    platform = data.get('platform', 'csv')
    template_id = data.get('templateId')

    if not assessment:
        return jsonify({"error": "No assessment data provided"}), 400

    try:
        # Get template structure if provided
        template_structure = None
        if template_id:
            meta_path = os.path.join(TEMPLATES_DIR, f"{template_id}.meta.json")
            if os.path.exists(meta_path):
                with open(meta_path, 'r') as f:
                    template_meta = json.load(f)
                    template_structure = template_meta.get('structure', {})

        result = build_platform_export(assessment, platform, template_structure)
        if result is None:
            return jsonify({"error": f"Unknown platform: {platform}"}), 400
        return jsonify(result)

    except Exception as e:
        _logger.exception("Platform export error")
        return jsonify({"error": "An internal error occurred"}), 500


@planner_bp.route('/api/grade-assessment-answers', methods=['POST'])
@require_teacher
@handle_route_errors
def grade_assessment_answers():
    """
    Grade student answers against the assessment using AI for open-ended questions.
    Returns detailed feedback for each question.
    """
    try:
        data = request.json
        assessment = data.get('assessment', {})
        answers = data.get('answers', {})

        if not assessment or not answers:
            return jsonify({"error": "Missing assessment or answers"}), 400

        return jsonify(grade_assessment_answers_logic(assessment, answers))

    except Exception as e:
        _logger.exception("Request failed: %s", request.path)
        return jsonify({"error": "An internal error occurred"}), 500


@planner_bp.route('/api/regenerate-questions', methods=['POST'])
@require_teacher
@handle_route_errors
def regenerate_questions():
    """Regenerate specific questions in an assessment/assignment using AI.

    Expects:
      questions_to_replace: [{section_index, question_index, question_type, points, dok, standard}, ...]
      existing_questions: [str, ...] — question texts to avoid duplicating
      config: {grade, subject, globalAINotes}
    Returns:
      replacements: [{section_index, question_index, question: {...}}, ...]
      usage: cost/token info
    """
    data = request.json
    questions_to_replace = data.get('questions_to_replace', [])
    existing_questions = data.get('existing_questions', [])
    config = data.get('config', {})

    if not questions_to_replace:
        return jsonify({"error": "No questions specified for regeneration"}), 400

    _subject = config.get('subject', '').strip()
    _grade = config.get('grade', config.get('grade_level', '')).strip()
    if not _subject or not _grade:
        from flask import current_app
        current_app.logger.warning(
            f"Regenerate requested without subject/grade: subject={_subject!r}, grade={_grade!r}")

    try:
        from backend.api_keys import get_api_key as _gak
        api_key = _gak('openai', getattr(g, 'user_id', 'local-dev'))
        return jsonify(generate_replacement_questions(
            questions_to_replace=questions_to_replace, existing_questions=existing_questions,
            config=config, api_key=api_key))
    except Exception as e:
        _logger.exception("Request failed: %s", request.path)
        return jsonify({"error": "An internal error occurred"}), 500


@planner_bp.route('/api/planner/costs', methods=['GET'])
@require_teacher
@handle_route_errors
def get_planner_costs():
    """Return planner API cost summary."""
    try:
        with open(PLANNER_COSTS_FILE, 'r') as f:
            return jsonify(json.load(f))
    except (FileNotFoundError, json.JSONDecodeError):
        return jsonify({"total": {"input_tokens": 0, "output_tokens": 0, "total_cost": 0, "api_calls": 0}, "daily": {}})


@planner_bp.route('/api/adjust-reading-level', methods=['POST'])
@require_teacher
@handle_route_errors
def adjust_reading_level():
    """Rewrite text at a target reading level while preserving key terms."""
    from backend.api_keys import get_api_key as _gak
    teacher_id = getattr(g, 'user_id', 'local-dev')
    api_key = _gak('openai', teacher_id)

    if not api_key or api_key.strip() == "" or "your-key-here" in api_key:
        return jsonify({"error": "Missing or placeholder OpenAI API Key"})

    data = request.json or {}
    text = data.get('text', '').strip()
    target_level = data.get('target_level', '6')
    subject = data.get('subject', '')
    preserve_terms = data.get('preserve_terms', [])

    if not text:
        return jsonify({"error": "No text provided"}), 400

    try:
        return jsonify(adjust_reading_level_content(
            text=text, target_level=target_level, subject=subject,
            preserve_terms=preserve_terms, api_key=api_key,
        ))
    except Exception as e:
        _logger.exception("Request failed: %s", request.path)
        return jsonify({"error": "An internal error occurred"}), 500


@planner_bp.route('/api/extract-text', methods=['POST'])
@require_teacher
@handle_route_errors
def extract_text_from_file():
    """Extract plain text from uploaded documents (docx, pdf, txt) or images (png, jpg, etc.)."""
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    filename = file.filename.lower() if file.filename else ''
    ext = os.path.splitext(filename)[1].lower() if filename else ''
    if ext not in ALLOWED_DOC_EXTENSIONS:
        return jsonify({"error": f"Unsupported file type: {ext}"}), 400

    file_data = file.read()

    # Images need an OpenAI key (vision). Resolve it route-side (Flask g) and own
    # the missing-key 400 here; non-image types never touch the key, preserving
    # the original no-extra-key-lookup behavior.
    api_key = None
    if filename.endswith(IMAGE_EXTENSIONS):
        from backend.api_keys import get_api_key as _gak
        teacher_id = getattr(g, 'user_id', 'local-dev')
        api_key = _gak('openai', teacher_id)
        if not api_key:
            return jsonify({"error": "OpenAI API key required for image text extraction"}), 400

    try:
        text = extract_text_from_upload(file_data=file_data, filename=filename, api_key=api_key)
        return jsonify({"text": text})
    except TextExtractionError as te:
        return jsonify({"error": str(te)}), 400
    except Exception:
        _logger.exception("Request failed: %s", request.path)
        return jsonify({"error": "An internal error occurred"}), 500


# ══════════════════════════════════════════════════════════════
# STUDY GUIDE GENERATION
# ══════════════════════════════════════════════════════════════

@planner_bp.route('/api/generate-study-guide', methods=['POST'])
@require_teacher
@handle_route_errors
def generate_study_guide():
    """Generate a structured study guide from content using Gemini Flash."""
    data = request.get_json(silent=True) or {}

    content = data.get('content', '').strip()
    title = data.get('title', 'Study Guide')
    subject = data.get('subject', '')
    grade = data.get('grade', '')
    instructions = data.get('instructions', '')
    global_ai_notes = data.get('globalAINotes', '')
    lesson_plan = data.get('lessonPlan')

    if not content and not lesson_plan:
        return jsonify({"error": "Provide content or a lesson plan to generate a study guide."}), 400

    user_id = getattr(g, 'user_id', 'local-dev')

    try:
        study_guide = generate_study_guide_content(
            content=content, subject=subject, grade=grade,
            instructions=instructions, global_ai_notes=global_ai_notes,
            lesson_plan=lesson_plan, user_id=user_id,
        )
        return jsonify({
            "study_guide": study_guide,
            "title": study_guide.get("title", title),
        })

    except json.JSONDecodeError as e:
        _logger.error("Study guide JSON parse failed: %s", e)
        return jsonify({"error": "Failed to parse study guide. Please try again."}), 500
    except Exception as e:
        _logger.exception("Study guide generation failed")
        return jsonify({"error": f"Generation failed: {str(e)[:200]}"}), 500


# ══════════════════════════════════════════════════════════════
# STUDY GUIDE EXPORT
# ══════════════════════════════════════════════════════════════

# Wave 6 Slice 1: study-aid render helpers extracted to
# backend/services/planner_export.py. Re-imported here so the export route
# bodies and tests that patch _get_export_dir on this module keep working.
from backend.services.planner_export import (  # noqa: F401  (re-export shim)
    _export_study_guide_docx,
    _export_study_guide_pdf,
    _export_flashcards_pdf,
    _export_flashcards_docx,
)


@planner_bp.route('/api/export-study-guide', methods=['POST'])
@require_teacher
@handle_route_errors
def export_study_guide():
    """Export a study guide to DOCX or PDF."""
    data = request.get_json(silent=True) or {}
    study_guide = data.get('study_guide')
    fmt = data.get('format', 'docx').lower()

    if not study_guide:
        return jsonify({"error": "No study guide data provided."}), 400

    title = study_guide.get("title", "Study Guide")
    safe_title = "".join(c for c in title if c.isalnum() or c in " -_").strip()[:80]
    export_dir = _get_export_dir()

    try:
        if fmt == 'pdf':
            filepath = os.path.join(export_dir, f"{safe_title}.pdf")
            _export_study_guide_pdf(study_guide, filepath)
            mimetype = 'application/pdf'
        else:
            filepath = os.path.join(export_dir, f"{safe_title}.docx")
            _export_study_guide_docx(study_guide, filepath)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'

        return send_file(filepath, mimetype=mimetype, as_attachment=True,
                         download_name=os.path.basename(filepath))

    except Exception as e:
        _logger.exception("Study guide export failed")
        return jsonify({"error": f"Export failed: {str(e)[:200]}"}), 500


# ══════════════════════════════════════════════════════════════
# FLASHCARD GENERATION
# ══════════════════════════════════════════════════════════════

@planner_bp.route('/api/generate-flashcards', methods=['POST'])
@require_teacher
@handle_route_errors
def generate_flashcards():
    """Generate flashcards from content using Gemini Flash."""
    data = request.get_json(silent=True) or {}

    content = data.get('content', '').strip()
    title = data.get('title', 'Flashcards')
    subject = data.get('subject', '')
    grade = data.get('grade', '')
    instructions = data.get('instructions', '')
    global_ai_notes = data.get('globalAINotes', '')
    lesson_plan = data.get('lessonPlan')
    card_count = data.get('cardCount', 15)

    if not content and not lesson_plan:
        return jsonify({"error": "Provide content or a lesson plan to generate flashcards."}), 400

    user_id = getattr(g, 'user_id', 'local-dev')

    try:
        flashcards = generate_flashcards_content(
            content=content, subject=subject, grade=grade,
            instructions=instructions, global_ai_notes=global_ai_notes,
            lesson_plan=lesson_plan, card_count=card_count, user_id=user_id,
        )
        return jsonify({
            "flashcards": flashcards,
            "title": flashcards.get("title", title),
        })

    except json.JSONDecodeError as e:
        _logger.error("Flashcard JSON parse failed: %s", e)
        return jsonify({"error": "Failed to parse flashcards. Please try again."}), 500
    except Exception as e:
        _logger.exception("Flashcard generation failed")
        return jsonify({"error": "Generation failed: " + str(e)[:200]}), 500


@planner_bp.route('/api/export-flashcards', methods=['POST'])
@require_teacher
@handle_route_errors
def export_flashcards():
    """Export flashcards to PDF or DOCX."""
    data = request.get_json(silent=True) or {}
    flashcards = data.get('flashcards')
    fmt = data.get('format', 'pdf').lower()

    if not flashcards:
        return jsonify({"error": "No flashcard data provided."}), 400

    title = flashcards.get("title", "Flashcards")
    safe_title = "".join(c for c in title if c.isalnum() or c in " -_").strip()[:80]
    export_dir = _get_export_dir()

    try:
        if fmt == 'docx':
            filepath = os.path.join(export_dir, safe_title + ".docx")
            _export_flashcards_docx(flashcards, filepath)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            filepath = os.path.join(export_dir, safe_title + ".pdf")
            _export_flashcards_pdf(flashcards, filepath)
            mimetype = 'application/pdf'

        return send_file(filepath, mimetype=mimetype, as_attachment=True,
                         download_name=os.path.basename(filepath))

    except Exception as e:
        _logger.exception("Flashcard export failed")
        return jsonify({"error": "Export failed: " + str(e)[:200]}), 500


# ══════════════════════════════════════════════════════════════
# SLIDE DECK GENERATION
# ══════════════════════════════════════════════════════════════

@planner_bp.route('/api/generate-slides', methods=['POST'])
@require_teacher
@handle_route_errors
def generate_slides():
    """Generate a slide deck from content using Gemini Flash."""
    data = request.get_json(silent=True) or {}

    content = data.get('content', '').strip()
    title = data.get('title', 'Slide Deck')
    subject = data.get('subject', '')
    grade = data.get('grade', '')
    instructions = data.get('instructions', '')
    global_ai_notes = data.get('globalAINotes', '')
    lesson_plan = data.get('lessonPlan')
    slide_count = min(data.get('slideCount', 10), 20)
    max_images = min(data.get('maxImages', 5), 10)
    generate_images = data.get('generateImages', True)
    deck_format = data.get('deckFormat', 'detailed')

    if not content and not lesson_plan:
        return jsonify({"error": "Provide content or a lesson plan to generate slides."}), 400

    user_id = getattr(g, 'user_id', 'local-dev')

    try:
        return jsonify(generate_slides_payload(
            content=content, title=title, subject=subject, grade=grade,
            instructions=instructions, global_ai_notes=global_ai_notes,
            lesson_plan=lesson_plan, slide_count=slide_count, max_images=max_images,
            generate_images=generate_images, deck_format=deck_format, user_id=user_id,
        ))

    except json.JSONDecodeError as e:
        _logger.error("Slide content JSON parse failed: %s", e)
        return jsonify({"error": "Failed to parse slide content. Please try again."}), 500
    except Exception as e:
        _logger.exception("Slide generation failed")
        return jsonify({"error": "Generation failed: " + str(e)[:200]}), 500


@planner_bp.route('/api/export-slides', methods=['POST'])
@require_teacher
@handle_route_errors
def export_slides():
    """Export generated slides as PowerPoint (.pptx)."""
    data = request.get_json(silent=True) or {}
    slide_data = data.get('slides')

    if not slide_data or not slide_data.get('slides'):
        return jsonify({"error": "No slide data provided."}), 400

    title = slide_data.get("title", "Slide Deck")
    safe_title = "".join(c for c in title if c.isalnum() or c in " -_").strip()[:80]
    export_dir = _get_export_dir()

    try:
        import base64
        from backend.services.slide_generator import assemble_pptx

        # Decode image data from base64
        images = {}
        for k, v in slide_data.get("_image_data", {}).items():
            try:
                images[int(k)] = base64.b64decode(v)
            except Exception:
                pass

        filepath = os.path.join(export_dir, safe_title + ".pptx")
        assemble_pptx(
            slide_data["slides"], slide_data.get("theme", {}),
            title, images, filepath
        )

        return send_file(
            filepath,
            mimetype='application/vnd.openxmlformats-officedocument.presentationml.presentation',
            as_attachment=True,
            download_name=safe_title + ".pptx",
        )

    except Exception as e:
        _logger.exception("Slide export failed")
        return jsonify({"error": "Export failed: " + str(e)[:200]}), 500
