"""
Assignment Grader for 6th Grade Social Studies
===============================================
Tailored for Southwestern Middle School

FILE NAMING EXPECTED: FirstName_LastName_AssignmentName.docx
ROSTER FORMAT: "LastName; FirstName" with Student ID and Email columns

FERPA COMPLIANCE:
- Student names are NOT sent to OpenAI's API
- Only assignment content is analyzed for grading
- All student identification stays local on your computer
- OpenAI API data is not used to train models (per their policy)
- Consult your district's policies for AI tool usage

SETUP:
1. pip install openai python-docx openpyxl python-dotenv
2. Create .env file with: OPENAI_API_KEY=your-key-here
3. Update the folder paths below
4. Update the ASSIGNMENT_INSTRUCTIONS for each assignment
"""

import os
import csv
import json
import re
import random
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv

# Load environment variables from .env file (override system env vars)
import os
app_dir = os.path.dirname(os.path.abspath(__file__))
load_dotenv(os.path.join(app_dir, '.env'), override=True)

# =============================================================================
# CONFIGURATION - UPDATE THESE FOR EACH GRADING SESSION
# =============================================================================

# Folder containing student assignment files (.docx)
ASSIGNMENT_FOLDER = "/Users/alexc/Downloads/Assignments"

# Output folder for CSV and email files
OUTPUT_FOLDER = "/Users/alexc/Downloads/Assignment Grader/Results"

# Path to your student roster Excel file
ROSTER_FILE = "/Users/alexc/Downloads/Assignment Grader/all_students_updated.xlsx"

# API keys (set in .env file)
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY", "")
ANTHROPIC_API_KEY = os.getenv("ANTHROPIC_API_KEY", "")
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")

# Assignment name (used in output files and emails)
ASSIGNMENT_NAME = "Cornell Notes - Political Parties"  # UPDATE FOR EACH ASSIGNMENT

# Marker phrase(s) that indicate where student work begins
# Only content within the section (until next header) will be graded
STUDENT_WORK_MARKERS = [
    # Direct task indicators
    "student task",
    "your turn",
    "now you try",
    "student responses below",
    
    # Question indicators
    "answer the question",
    "questions to check understanding",
    "questions / summary",
    
    # Activity types
    "answer the",
    "match",
    "fill-in-the-blank",
    "fill in the blank",
    "write",
    "explain",
    "complete",
    "describe",
    "summarize",
    
    # Analysis & thinking verbs
    "analyze",
    "compare",
    "contrast",
    "evaluate",
    "identify",
    "list",
    "define",
    
    # Application verbs
    "apply",
    "demonstrate",
    "illustrate",
    "predict",
    "solve",
    
    # Creation verbs
    "create",
    "design",
    "develop",
    "construct",
    
    # Reflection indicators
    "final reflection",
    "final reflection question",
    "reflect",
    "think about",
    
    # Learning check indicators
    "let's see what you've learned",
    "lets see what you've learned",
    "let's see what you've learned",  # curly apostrophe version
    "check your understanding",
    "show what you know",
    "practice",
]

# Section headers that indicate a NEW section (stops extraction from previous marker)
SECTION_HEADERS = [
    "vocabulary",
    "vocabulary mini-lesson",
    "key vocabulary",
    "notes",
    "guided notes",
    "reading",
    "directions",
    "instructions",
    "primary source",
    "background",
    "overview",
    "introduction",
]


# =============================================================================
# GRADING RUBRIC - Customize point values and criteria as needed
# =============================================================================

GRADING_RUBRIC = """
You are grading 6th grade Social Studies assignments. Be ENCOURAGING and GENEROUS.
These are 11-12 year old students - grade with appropriate expectations for their age.

IMPORTANT GRADING GUIDELINES:
- For FILL-IN-THE-BLANK exercises: Accept any answer that is factually correct or very close. Spelling mistakes should NOT reduce the grade if the intent is clear.
- Accept reasonable synonyms, alternate phrasings, and partial answers that demonstrate understanding.
- If a student gets the main idea right but uses slightly different wording, give FULL CREDIT.
- Minor spelling errors (like "piler" instead of "pillar") should NOT be penalized.
- Be GENEROUS - when in doubt, give the student the benefit of the doubt.

GRADING SCALE (out of 100 points):

1. CONTENT ACCURACY (40 points)
   - Are the answers factually correct or demonstrate understanding?
   - For fill-in-the-blank: Is the completed statement essentially true?
   - 40 pts: Most answers correct (90%+)
   - 35 pts: Good understanding (80-89% correct)
   - 30 pts: Solid effort (70-79% correct)
   - 25 pts: Some understanding (60-69% correct)
   - 20 pts: Partial understanding (50-59% correct)
   - Below 20: Less than half correct

2. COMPLETENESS (25 points)
   - Did the student attempt ALL questions/blanks?
   - 25 pts: All questions attempted
   - 20 pts: Nearly all attempted (90%+)
   - 15 pts: Most attempted (75%+)
   - 10 pts: About half attempted
   - 5 pts: Less than half attempted

3. WRITING QUALITY (20 points)
   - Is the writing legible and understandable?
   - For fill-in-blank, this is less important - be generous
   - 20 pts: Clear and readable
   - 15 pts: Minor issues but understandable
   - 10 pts: Some difficulty but can figure out meaning
   - 5 pts: Hard to understand

4. EFFORT & ENGAGEMENT (15 points)
   - Did the student put in genuine effort?
   - Are answers thoughtful (not random guesses)?
   - 15 pts: Clear effort shown
   - 10 pts: Good effort
   - 5 pts: Minimal effort
   - 0 pts: No real effort

GRADE RANGES:
- A: 90-100 (Great job!)
- B: 80-89 (Good work!)
- C: 70-79 (Solid effort)
- D: 60-69 (Needs improvement)
- F: Below 60 (Significant concerns)

REMEMBER: These are 6th graders. Be kind, encouraging, and generous with grading.
A student who attempts all questions and gets most right should get an A or B.
"""

# UPDATE THIS FOR EACH ASSIGNMENT
ASSIGNMENT_INSTRUCTIONS = """
CURRENT ASSIGNMENT: Cornell Notes - Political Parties

This is a Cornell Notes assignment about Political Parties in early American history.

Requirements:
1. Main notes section should contain key information about political parties
2. Questions/cue column should have relevant questions
3. Summary at bottom should synthesize the main ideas
4. Content should demonstrate understanding of:
   - Federalists vs Democratic-Republicans
   - Key figures (Hamilton, Jefferson, etc.)
   - Main beliefs and positions of each party
"""


# =============================================================================
# ROSTER LOADING - Reads your Excel student list
# =============================================================================

def load_roster(roster_path: str) -> dict:
    """
    Load student roster from Excel or CSV file.

    Excel format: Student, Student ID, Local ID, Email, Grade, Team
    CSV format: FirstName, LastName, StudentID, Email, Period

    Returns dict mapping "firstname lastname" (lowercase) -> student info
    """
    roster = {}

    if not Path(roster_path).exists():
        print(f"⚠️  Roster file not found: {roster_path}")
        return {}

    roster_path_lower = roster_path.lower()

    # Handle CSV files
    if roster_path_lower.endswith('.csv'):
        import csv
        with open(roster_path, 'r', encoding='utf-8') as f:
            reader = csv.DictReader(f)
            for row in reader:
                # Try various column name formats
                first_name = row.get('FirstName') or row.get('First Name') or row.get('first_name') or row.get('First') or ''
                last_name = row.get('LastName') or row.get('Last Name') or row.get('last_name') or row.get('Last') or ''
                student_id = row.get('StudentID') or row.get('Student ID') or row.get('student_id') or row.get('ID') or ''
                email = row.get('Email') or row.get('email') or ''
                period = row.get('Period') or row.get('period') or row.get('Class') or ''

                # If Student column exists with "Last; First" format
                if not first_name and not last_name:
                    student_col = row.get('Student') or row.get('Name') or ''
                    if ';' in student_col:
                        parts = student_col.split(';')
                        last_name = parts[0].strip()
                        first_name = parts[1].strip() if len(parts) > 1 else ''
                    elif ',' in student_col:
                        parts = student_col.split(',')
                        last_name = parts[0].strip()
                        first_name = parts[1].strip() if len(parts) > 1 else ''

                if not first_name and not last_name:
                    continue

                first_name_simple = first_name.split()[0] if first_name else ''
                lookup_key = f"{first_name_simple} {last_name}".lower().strip()

                roster[lookup_key] = {
                    "student_id": str(student_id),
                    "student_name": f"{first_name} {last_name}".strip(),
                    "first_name": first_name,
                    "last_name": last_name,
                    "email": email or "",
                    "period": str(period) if period else ""
                }

                reverse_key = f"{last_name} {first_name_simple}".lower().strip()
                roster[reverse_key] = roster[lookup_key]

        print(f"📋 Loaded {len(roster)//2} students from CSV roster")
        return roster

    # Handle Excel files
    try:
        import openpyxl
    except ImportError:
        print("❌ openpyxl not installed. Run: pip install openpyxl")
        return {}

    wb = openpyxl.load_workbook(roster_path)
    sheet = wb.active

    # Get header row to find column indices
    headers = [str(cell.value).lower() if cell.value else '' for cell in sheet[1]]

    # Try to find period column
    period_col = None
    for i, h in enumerate(headers):
        if 'period' in h or 'class' in h or 'team' in h:
            period_col = i
            break

    # Skip header row
    for row in sheet.iter_rows(min_row=2, values_only=True):
        if not row[0]:
            continue

        student_name = row[0]
        student_id = row[1]
        email = row[3] if len(row) > 3 else ''
        period = row[period_col] if period_col is not None and len(row) > period_col else ''

        if ';' in str(student_name):
            parts = str(student_name).split(';')
            last_name = parts[0].strip()
            first_name = parts[1].strip() if len(parts) > 1 else ""
            first_name_simple = first_name.split()[0] if first_name else ""
        else:
            last_name = str(student_name)
            first_name = ""
            first_name_simple = ""

        lookup_key = f"{first_name_simple} {last_name}".lower().strip()

        roster[lookup_key] = {
            "student_id": str(student_id),
            "student_name": f"{first_name} {last_name}".strip(),
            "first_name": first_name,
            "last_name": last_name,
            "email": email or "",
            "period": str(period) if period else ""
        }

        reverse_key = f"{last_name} {first_name_simple}".lower().strip()
        roster[reverse_key] = roster[lookup_key]

    print(f"📋 Loaded {len(roster)//2} students from Excel roster")
    return roster


# =============================================================================
# FILE PARSING - Extract student name from filename
# =============================================================================

def parse_filename(filename: str) -> dict:
    """
    Parse student info from filename.
    
    Expected format: FirstName_LastName_AssignmentName.docx
    Examples:
        A'kareah_West_Cornell Notes_ Political Parties.docx
        Eli_Long_Hamilton_Jefferson_Graphic_Organizer.docx
        Gabriella_Bueno_Cornell Notes – The Louisiana Purchase.docx
    
    Returns: {"first_name": ..., "last_name": ..., "assignment_part": ...}
    """
    # Remove extension
    name = Path(filename).stem
    
    # Split by underscore
    parts = name.split('_')
    
    if len(parts) >= 2:
        first_name = parts[0].strip()
        last_name = parts[1].strip()
        assignment_part = '_'.join(parts[2:]) if len(parts) > 2 else ""
        
        return {
            "first_name": first_name,
            "last_name": last_name,
            "assignment_part": assignment_part,
            "lookup_key": f"{first_name} {last_name}".lower()
        }
    else:
        # Can't parse - return filename as-is
        return {
            "first_name": name,
            "last_name": "",
            "assignment_part": "",
            "lookup_key": name.lower()
        }


def read_docx_file(filepath: str) -> str:
    """
    Read text content from a Word document (.docx)
    """
    try:
        from docx import Document
    except ImportError:
        print("❌ python-docx not installed. Run: pip install python-docx")
        return None
    
    try:
        doc = Document(filepath)
        full_text = []
        
        # Get text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)
        
        # Also get text from tables (Cornell notes often use tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text)
        
        return '\n'.join(full_text)
    except Exception as e:
        print(f"  ⚠️  Error reading file: {e}")
        return None


def read_image_file(filepath: str) -> dict:
    """
    Read an image file and return it as base64 for GPT-4o vision.
    
    Returns dict with:
    - type: "image"
    - data: base64 encoded image
    - media_type: image MIME type
    """
    import base64
    
    filepath = Path(filepath)
    extension = filepath.suffix.lower()
    
    # Map extensions to MIME types
    mime_types = {
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.webp': 'image/webp'
    }
    
    if extension not in mime_types:
        print(f"  ⚠️  Unsupported image type: {extension}")
        return None
    
    try:
        with open(filepath, 'rb') as f:
            image_data = base64.b64encode(f.read()).decode('utf-8')
        
        return {
            "type": "image",
            "data": image_data,
            "media_type": mime_types[extension]
        }
    except Exception as e:
        print(f"  ⚠️  Error reading image: {e}")
        return None


def read_assignment_file(filepath: str) -> dict:
    """
    Read assignment file based on its extension.
    Supports: .docx, .txt, .jpg, .jpeg, .png, .gif, .webp
    
    Returns dict with:
    - type: "text" or "image"
    - content: text content or base64 image data
    """
    filepath = Path(filepath)
    extension = filepath.suffix.lower()
    
    # Text-based files
    if extension == '.docx':
        content = read_docx_file(filepath)
        if content:
            return {"type": "text", "content": content}
        return None
    
    elif extension == '.txt':
        try:
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return {"type": "text", "content": f.read()}
        except Exception as e:
            print(f"  ⚠️  Error reading text file: {e}")
            return None
    
    # Image files
    elif extension in ['.jpg', '.jpeg', '.png', '.gif', '.webp']:
        image_data = read_image_file(filepath)
        if image_data:
            return {
                "type": "image",
                "content": image_data["data"],
                "media_type": image_data["media_type"]
            }
        return None
    
    else:
        print(f"  ⚠️  Unsupported file type: {extension}")
        return None


def extract_student_work(content: str) -> tuple:
    """
    Extract only the student work portions of the document.
    
    Looks for marker phrases and extracts content from each marked section
    until the next section header (like "Vocabulary" or "Notes").
    
    This ensures we only grade student responses, not teacher-provided content.
    
    Returns: (student_work, markers_found)
    - student_work: The extracted student content (or full content if no marker)
    - markers_found: List of markers that were found
    """
    content_lower = content.lower()
    lines = content.split('\n')
    
    # Find all marker positions
    found_markers = []
    marker_positions = []
    
    for marker in STUDENT_WORK_MARKERS:
        marker_lower = marker.lower()
        pos = content_lower.find(marker_lower)
        if pos != -1:
            marker_positions.append((pos, marker, len(marker)))
            if marker not in found_markers:
                found_markers.append(marker)
    
    if not marker_positions:
        # No markers found - return full content
        return content, []
    
    # Sort by position
    marker_positions.sort(key=lambda x: x[0])
    
    # Find all section header positions (these END a student section)
    header_positions = []
    for header in SECTION_HEADERS:
        header_lower = header.lower()
        pos = 0
        while True:
            pos = content_lower.find(header_lower, pos)
            if pos == -1:
                break
            header_positions.append(pos)
            pos += 1
    
    # Also treat markers as potential section boundaries
    all_boundaries = sorted(set(header_positions + [m[0] for m in marker_positions]))
    
    # Extract content from each marker section
    student_sections = []
    
    for marker_pos, marker_name, marker_len in marker_positions:
        # Start after the marker
        start_pos = marker_pos + marker_len
        
        # Find the next boundary (header or another marker) after this marker
        end_pos = len(content)
        for boundary in all_boundaries:
            if boundary > start_pos + 10:  # At least 10 chars after marker start
                # Check if this boundary is a non-student header
                boundary_text = content_lower[max(0, boundary-5):boundary+50]
                is_student_marker = any(m.lower() in boundary_text for m in STUDENT_WORK_MARKERS)
                is_header = any(h.lower() in boundary_text for h in SECTION_HEADERS)
                
                if is_header and not is_student_marker:
                    end_pos = boundary
                    break
        
        # Extract this section
        section_content = content[start_pos:end_pos].strip()
        if section_content and len(section_content) > 10:
            student_sections.append(section_content)
    
    # Combine all student sections
    if student_sections:
        combined = '\n\n---\n\n'.join(student_sections)
        return combined, found_markers
    
    # Fallback: return everything after first marker
    first_pos = marker_positions[0][0] + marker_positions[0][2]
    return content[first_pos:].strip(), found_markers


# =============================================================================
# AI GRADING - MULTI-PROVIDER SUPPORT
# =============================================================================

def _grade_with_openai(prompt: str, assignment_data: dict, model: str) -> str:
    """Grade using OpenAI API."""
    from openai import OpenAI
    client = OpenAI(api_key=OPENAI_API_KEY)

    if assignment_data["type"] == "text":
        messages = [{"role": "user", "content": prompt}]
    else:
        messages = [{
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{assignment_data['media_type']};base64,{assignment_data['content']}"
                    }
                }
            ]
        }]

    # Map model aliases to actual model names
    model_map = {
        'gpt-4o': 'gpt-4o',
        'gpt-4o-mini': 'gpt-4o-mini',
        'gpt-4-turbo': 'gpt-4-turbo',
        'gpt-4': 'gpt-4',
        'gpt-3.5-turbo': 'gpt-3.5-turbo'
    }
    actual_model = model_map.get(model, 'gpt-4o')

    response = client.chat.completions.create(
        model=actual_model,
        messages=messages,
        max_tokens=1500,
        temperature=0.3
    )
    return response.choices[0].message.content.strip()


def _grade_with_anthropic(prompt: str, assignment_data: dict, model: str) -> str:
    """Grade using Anthropic Claude API."""
    import anthropic
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)

    # Map model aliases to actual model names
    model_map = {
        'claude-sonnet': 'claude-sonnet-4-20250514',
        'claude-haiku': 'claude-3-5-haiku-20241022',
        'claude-opus': 'claude-3-opus-20240229'
    }
    actual_model = model_map.get(model, 'claude-sonnet-4-20250514')

    if assignment_data["type"] == "text":
        content = [{"type": "text", "text": prompt}]
    else:
        content = [
            {"type": "text", "text": prompt},
            {
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": assignment_data['media_type'],
                    "data": assignment_data['content']
                }
            }
        ]

    response = client.messages.create(
        model=actual_model,
        max_tokens=1500,
        messages=[{"role": "user", "content": content}]
    )
    return response.content[0].text.strip()


def _grade_with_gemini(prompt: str, assignment_data: dict, model: str) -> str:
    """Grade using Google Gemini API."""
    import google.generativeai as genai
    genai.configure(api_key=GEMINI_API_KEY)

    # Map model aliases to actual model names
    model_map = {
        'gemini-flash': 'gemini-2.0-flash',
        'gemini-pro': 'gemini-1.5-pro'
    }
    actual_model = model_map.get(model, 'gemini-2.0-flash')

    gen_model = genai.GenerativeModel(actual_model)

    if assignment_data["type"] == "text":
        response = gen_model.generate_content(prompt)
    else:
        # For images, we need to include the image data
        import base64
        image_data = base64.b64decode(assignment_data['content'])
        image_part = {
            "mime_type": assignment_data['media_type'],
            "data": image_data
        }
        response = gen_model.generate_content([prompt, image_part])

    return response.text.strip()


def grade_assignment(student_name: str, assignment_data: dict, custom_ai_instructions: str = '', grade_level: str = '6', subject: str = 'Social Studies', model: str = 'gpt-4o') -> dict:
    """
    Use AI (OpenAI, Anthropic, or Gemini) to grade a student assignment.

    FERPA COMPLIANCE: Student name is NOT sent to AI APIs.
    We use "Student" as a placeholder to protect privacy.

    Supports both text and image inputs.

    Parameters:
    - student_name: Name of the student (kept local, not sent to API)
    - assignment_data: dict with "type" ("text" or "image") and "content"
    - custom_ai_instructions: Additional grading instructions from the teacher
    - grade_level: The student's grade level (e.g., '6', '7', '8')
    - subject: The subject being graded (e.g., 'Social Studies', 'English/ELA')
    - model: AI model to use (gpt-4o, claude-sonnet, gemini-flash, etc.)

    Returns dict with:
    - score: numeric grade (0-100)
    - letter_grade: A, B, C, D, or F
    - feedback: detailed feedback for the student
    - breakdown: points for each rubric category
    - authenticity_flag: 'clean', 'review', or 'flagged'
    - authenticity_reason: Explanation for flagged or review status
    """

    # Check for empty/blank student submissions before sending to API
    content = assignment_data.get("content", "")
    if assignment_data.get("type") == "text" and content:
        import re

        # Method 1: Check for filled-in blanks (text between underscores like ___answer___)
        filled_blanks = re.findall(r'_{2,}([^_\n]+)_{2,}', content)
        filled_blanks = [b.strip() for b in filled_blanks if b.strip() and len(b.strip()) > 1]

        # Method 2: Check for content after colons that isn't just blanks
        # e.g., "Nationalism: the belief that..." vs "Nationalism: ___"
        after_colons = re.findall(r':\s*([^_\n:]{10,})', content)
        after_colons = [a.strip() for a in after_colons if a.strip() and not a.strip().startswith('_')]

        # Method 3: Look for paragraph-length responses (likely written answers)
        paragraphs = [p.strip() for p in content.split('\n\n') if len(p.strip()) > 50]
        # Filter out paragraphs that are mostly underscores
        real_paragraphs = [p for p in paragraphs if p.count('_') < len(p) * 0.3]

        # Method 4: Count lines that are JUST underscores (blank response lines)
        blank_lines = len(re.findall(r'^[\s_]+$', content, re.MULTILINE))
        total_lines = len([l for l in content.split('\n') if l.strip()])
        blank_ratio = blank_lines / max(total_lines, 1)

        # Method 5: Check for questions followed by no response
        lines = content.split('\n')
        unanswered_questions = []
        question_indices = []

        for i, line in enumerate(lines):
            line_stripped = line.strip()
            # Check if line is a question (ends with ? or starts with number/bullet or is a vocab term with colon)
            is_question = (
                line_stripped.endswith('?') or
                re.match(r'^\d+[\.\)]\s*\w', line_stripped) or
                re.match(r'^[a-zA-Z][\.\)]\s*\w', line_stripped) or
                (line_stripped.endswith(':') and len(line_stripped) > 5 and '_' not in line_stripped)
            )
            if is_question and len(line_stripped) > 5:
                question_indices.append(i)

        # Check content between consecutive questions
        for idx, q_idx in enumerate(question_indices):
            line_stripped = lines[q_idx].strip()
            next_q_idx = question_indices[idx + 1] if idx + 1 < len(question_indices) else len(lines)

            content_between = []
            for j in range(q_idx + 1, min(next_q_idx, q_idx + 6)):
                if j < len(lines):
                    between_line = lines[j].strip()
                    if between_line and not re.match(r'^[_\s\-\.]+$', between_line):
                        if len(between_line) > 3 and between_line.count('_') < len(between_line) * 0.5:
                            content_between.append(between_line)

            if not content_between:
                unanswered_questions.append(line_stripped[:60] + "..." if len(line_stripped) > 60 else line_stripped)

        # Determine if submission is blank
        has_filled_blanks = len(filled_blanks) >= 2
        has_written_responses = len(after_colons) >= 2 or len(real_paragraphs) >= 1
        mostly_blank_lines = blank_ratio > 0.4
        many_unanswered = len(unanswered_questions) >= 3

        is_blank = (not has_filled_blanks and not has_written_responses and mostly_blank_lines) or \
                   (many_unanswered and not has_filled_blanks and not has_written_responses)

        if is_blank:
            print(f"  ⚠️  BLANK/EMPTY SUBMISSION DETECTED")
            print(f"      Filled blanks: {len(filled_blanks)}, Written responses: {len(after_colons)}")
            print(f"      Blank line ratio: {blank_ratio:.1%}, Unanswered questions: {len(unanswered_questions)}")
            return {
                "score": 0,
                "letter_grade": "INCOMPLETE",
                "breakdown": {
                    "content_accuracy": 0,
                    "completeness": 0,
                    "critical_thinking": 0,
                    "communication": 0
                },
                "feedback": f"This assignment appears to be incomplete or blank. {len(unanswered_questions)} question(s) were found without responses. Please complete the assignment and resubmit, or see your teacher if you need help.",
                "student_responses": [],
                "unanswered_questions": unanswered_questions[:10],
                "authenticity_flag": "clean",
                "authenticity_reason": "",
                "skills_demonstrated": {}
            }

    # FERPA: Use anonymous placeholder instead of real student name
    anonymous_name = "Student"

    # Build custom instructions section if provided
    custom_section = ''
    if custom_ai_instructions:
        custom_section = f"""
---
TEACHER'S GRADING INSTRUCTIONS (FOLLOW THESE CAREFULLY):
{custom_ai_instructions}
---
"""

    # Map grade level to age range for context
    grade_age_map = {
        'K': '5-6', '1': '6-7', '2': '7-8', '3': '8-9', '4': '9-10', '5': '10-11',
        '6': '11-12', '7': '12-13', '8': '13-14', '9': '14-15', '10': '15-16',
        '11': '16-17', '12': '17-18'
    }
    age_range = grade_age_map.get(str(grade_level), '11-12')

    prompt_text = f"""
{GRADING_RUBRIC}

{ASSIGNMENT_INSTRUCTIONS}
{custom_section}
---

STUDENT CONTEXT:
- Grade Level: {grade_level}
- Subject: {subject}
- Expected Age Range: {age_range} years old

Please grade this student's work.
- Assess EVERY answer the student provided.
- For fill-in-the-blank: check if the answer is factually correct or close enough.
- Accept multiple valid answers and synonyms.
- DO NOT penalize spelling mistakes if the meaning is clear.
- Be GENEROUS - these are grade {grade_level} students ({age_range} years old)! A student who completes all work and gets most answers right deserves an A or B.
- IMPORTANT: If the teacher provided custom grading instructions above, follow them carefully.

IMPORTANT - AI/PLAGIARISM DETECTION:
Analyze the writing for signs that it may NOT be the student's original work.
Consider what is typical for a grade {grade_level} student ({age_range} years old) in {subject}:
- Is the vocabulary and sentence structure appropriate for grade {grade_level}?
- Are there sudden shifts in writing quality or voice?
- Does the writing sound overly formal, polished, or "adult" for a {age_range} year old?
- Are there unusual phrases or technical jargon that a grade {grade_level} student wouldn't typically use?
- Does the writing lack the natural errors and simplicity expected from this age group?
- Is the content suspiciously comprehensive or professionally structured?

Set "authenticity_flag" to one of:
- "clean" - Writing appears authentic and appropriate for grade {grade_level}
- "review" - Some concerns about writing level, teacher should review
- "flagged" - Writing significantly above expected grade {grade_level} level, likely AI-generated or copied

Provide your response in the following JSON format ONLY (no other text):
{{
    "score": <number 0-100>,
    "letter_grade": "<A, B, C, D, or F>",
    "breakdown": {{
        "content_accuracy": <points out of 40>,
        "completeness": <points out of 25>,
        "writing_quality": <points out of 20>,
        "effort_engagement": <points out of 15>
    }},
    "authenticity_flag": "<clean, review, or flagged>",
    "authenticity_reason": "<Brief explanation if review or flagged, otherwise empty string>",
    "feedback": "<2-3 paragraphs of encouraging feedback written directly to the student. Use simple, friendly language a grade {grade_level} student can understand. Start with specific praise for what they did well, then gently mention 1-2 areas to improve (if any), then end with encouragement. Be positive and supportive! Do NOT use the student's name - just say 'you' or 'your'.>"
}}
"""

    # Determine provider from model name
    if model.startswith('claude'):
        provider = 'anthropic'
    elif model.startswith('gemini'):
        provider = 'gemini'
    else:
        provider = 'openai'

    print(f"  🤖 Grading with {provider.upper()} ({model})...")

    try:
        # Build the message content based on input type
        if assignment_data["type"] == "text":
            full_prompt = prompt_text + f"\n\nSTUDENT'S RESPONSES/WORK:\n{assignment_data['content']}"
        elif assignment_data["type"] == "image":
            full_prompt = prompt_text + "\n\nSTUDENT'S WORK (see attached image):"
        else:
            return {"score": 0, "letter_grade": "ERROR", "breakdown": {}, "feedback": "Unknown content type"}

        # Route to appropriate provider
        if provider == 'anthropic':
            response_text = _grade_with_anthropic(full_prompt, assignment_data, model)
        elif provider == 'gemini':
            response_text = _grade_with_gemini(full_prompt, assignment_data, model)
        else:
            response_text = _grade_with_openai(full_prompt, assignment_data, model)
        
        # Clean up response (remove markdown code blocks if present)
        if response_text.startswith("```"):
            lines = response_text.split('\n')
            start = 1 if lines[0].startswith("```") else 0
            end = len(lines)
            for i in range(len(lines)-1, -1, -1):
                if lines[i].strip() == "```":
                    end = i
                    break
            response_text = '\n'.join(lines[start:end])
        response_text = response_text.strip()
        
        result = json.loads(response_text)
        return result
        
    except json.JSONDecodeError as e:
        print(f"  ⚠️  Error parsing AI response: {e}")
        return {
            "score": 0,
            "letter_grade": "ERROR",
            "breakdown": {},
            "feedback": "Error grading - please review manually."
        }
    except Exception as e:
        print(f"  ⚠️  API error: {e}")
        return {
            "score": 0,
            "letter_grade": "ERROR", 
            "breakdown": {},
            "feedback": f"API error: {e}"
        }


# =============================================================================
# EMAIL GENERATION
# =============================================================================

def generate_email_content(student_info: dict, grade_result: dict, assignment_name: str, teacher_name: str = '', subject_taught: str = '', school_name: str = '') -> tuple:
    """
    Generate email subject and body for a student.

    Returns: (subject, body)
    """
    first_name = student_info.get('first_name', 'Student').split()[0]  # Just first name

    subject = f"Grade for {assignment_name}: {grade_result['letter_grade']}"

    # Build signature
    signature = teacher_name if teacher_name else "Your Teacher"
    if subject_taught:
        signature += f"\n{subject_taught}"
    if school_name:
        signature += f"\n{school_name}"

    body = f"""Hi {first_name},

Here is your grade and feedback for {assignment_name}:

GRADE: {grade_result['score']}/100 ({grade_result['letter_grade']})

FEEDBACK:
{grade_result['feedback']}

If you have any questions about your grade, please see me during class.

- {signature}
"""
    return subject, body


def save_emails_to_folder(grades: list, output_folder: str, teacher_name: str = '', subject: str = '', school_name: str = ''):
    """
    Save emails as individual text files - ONE EMAIL PER STUDENT
    with feedback for ALL their assignments combined.
    """
    email_folder = Path(output_folder) / "emails"
    email_folder.mkdir(parents=True, exist_ok=True)
    
    # Group grades by student
    students = {}
    for grade in grades:
        student_name = grade.get('student_name', 'Unknown')
        if student_name not in students:
            # Get only first name (no middle initial)
            full_first = grade.get('first_name', student_name.split()[0])
            first_only = full_first.split()[0] if full_first else student_name.split()[0]
            students[student_name] = {
                'email': grade.get('email', ''),
                'first_name': first_only,
                'assignments': []
            }
        students[student_name]['assignments'].append(grade)
    
    email_count = 0
    for student_name, data in students.items():
        if not data['email']:
            continue
        
        # Build combined email
        assignments = data['assignments']
        first_name = data['first_name']
        
        # Subject line
        if len(assignments) == 1:
            subject = f"Grade for {assignments[0].get('assignment', 'Assignment')}: {assignments[0]['letter_grade']}"
        else:
            subject = f"Grades for {len(assignments)} Assignments"
        
        # Body
        body = f"Hi {first_name},\n\n"
        
        if len(assignments) == 1:
            a = assignments[0]
            body += f"Here is your grade and feedback for {a.get('assignment', 'your assignment')}:\n\n"
            body += f"GRADE: {a['score']}/100 ({a['letter_grade']})\n\n"
            body += f"FEEDBACK:\n{a.get('feedback', 'No feedback available.')}\n"
        else:
            body += "Here are your grades and feedback:\n\n"
            for a in assignments:
                assignment_name = a.get('assignment', 'Assignment')
                body += f"{'='*50}\n"
                body += f"📝 {assignment_name}\n"
                body += f"{'='*50}\n"
                body += f"GRADE: {a['score']}/100 ({a['letter_grade']})\n\n"
                body += f"FEEDBACK:\n{a.get('feedback', 'No feedback available.')}\n\n"
        
        body += "\nIf you have any questions about your grades, please see me during class.\n\n"
        # Build signature from teacher info
        signature = teacher_name if teacher_name else "Your Teacher"
        if subject:
            signature += f"\n{subject}"
        if school_name:
            signature += f"\n{school_name}"
        body += f"- {signature}\n"
        
        # Save file
        safe_name = re.sub(r'[^\w\s-]', '', student_name).replace(' ', '_')
        filepath = email_folder / f"{safe_name}_email.txt"
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"TO: {data['email']}\n")
            f.write(f"SUBJECT: {subject}\n")
            f.write(f"{'='*50}\n\n")
            f.write(body)
        
        email_count += 1
    
    print(f"📧 Saved {email_count} email files to: {email_folder}")


def create_outlook_drafts(grades: list):
    """
    Create draft emails in Outlook desktop app (Windows only).
    This lets you review each email before sending.
    """
    try:
        import win32com.client
    except ImportError:
        print("⚠️  pywin32 not installed. Run: pip install pywin32")
        print("   Falling back to saving emails as files.")
        return False
    
    try:
        outlook = win32com.client.Dispatch("Outlook.Application")
        count = 0
        
        for grade in grades:
            if not grade.get('email'):
                continue
                
            subject, body = generate_email_content(grade, grade, ASSIGNMENT_NAME)
            
            mail = outlook.CreateItem(0)  # 0 = mail item
            mail.To = grade['email']
            mail.Subject = subject
            mail.Body = body
            mail.Save()  # Save as draft
            count += 1
        
        print(f"📧 Created {count} draft emails in Outlook")
        return True
        
    except Exception as e:
        print(f"⚠️  Outlook error: {e}")
        return False


# =============================================================================
# CSV EXPORT FOR FOCUS
# =============================================================================

def export_focus_csv(grades: list, output_folder: str, assignment_name: str) -> list:
    """
    Create CSV files formatted for Focus import, SEPARATED BY ASSIGNMENT.
    
    Groups students by the assignment extracted from their filename,
    creates one CSV per assignment type.
    
    Format:
    Student ID,Score,Comment
    1950304,85,"Great job, Jackson! You correctly identified..."
    1956701,92,"Excellent work, Maria!..."
    
    Returns list of created file paths.
    """
    Path(output_folder).mkdir(parents=True, exist_ok=True)
    
    # Group grades by assignment
    assignments = {}
    for grade in grades:
        # Extract assignment name from filename (everything after FirstName_LastName_)
        filename = grade.get('filename', '')
        parts = Path(filename).stem.split('_')
        if len(parts) >= 3:
            # Assignment is everything after first two parts (first_last)
            assignment = '_'.join(parts[2:])
        else:
            assignment = "Unknown_Assignment"
        
        # Clean up assignment name for display (replace underscores with spaces)
        assignment = assignment.strip().replace('_', ' ')
        
        if assignment not in assignments:
            assignments[assignment] = []
        assignments[assignment].append(grade)
    
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    created_files = []
    
    print(f"\n  Creating {len(assignments)} separate Focus CSVs by assignment:")
    
    for assignment, assignment_grades in assignments.items():
        # Clean assignment name for filename
        safe_name = re.sub(r'[^\w\s-]', '', assignment).replace(' ', '_')[:50]
        filepath = Path(output_folder) / f"{safe_name}_{timestamp}.csv"
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            # Column headers: Student ID, Score, Comment
            writer.writerow(['Student ID', 'Score', 'Comment'])
            
            for grade in assignment_grades:
                if grade['student_id'] != "UNKNOWN":
                    # Get student's first name only (no middle initial)
                    full_first = grade.get('first_name', '')
                    first_name = full_first.split()[0] if full_first else ''
                    
                    # Clean up feedback
                    feedback = grade.get('feedback', '')
                    feedback_clean = ' '.join(feedback.split())
                    
                    # Varied natural phrases based on grade
                    score = grade.get('score', 0)
                    
                    if first_name:
                        if score >= 90:
                            openers = [
                                f"Great job, {first_name}!",
                                f"Excellent work, {first_name}!",
                                f"{first_name}, this was fantastic!",
                                f"Well done, {first_name}!",
                                f"{first_name}, you nailed this!"
                            ]
                        elif score >= 80:
                            openers = [
                                f"Nice work, {first_name}!",
                                f"Good job, {first_name}!",
                                f"{first_name}, solid effort here!",
                                f"Well done, {first_name}!",
                                f"{first_name}, this was good work!"
                            ]
                        elif score >= 70:
                            openers = [
                                f"{first_name}, decent effort!",
                                f"Good start, {first_name}!",
                                f"{first_name}, you're on the right track!",
                                f"Keep it up, {first_name}!",
                                f"{first_name}, nice try!"
                            ]
                        else:
                            openers = [
                                f"{first_name}, let's work on this together.",
                                f"{first_name}, I know you can do better!",
                                f"Keep trying, {first_name}!",
                                f"{first_name}, don't give up!",
                                f"{first_name}, let's review this."
                            ]
                        
                        opener = random.choice(openers)
                        comment = f"{opener} {feedback_clean}"
                    else:
                        comment = feedback_clean
                    
                    writer.writerow([
                        grade['student_id'], 
                        grade['score'],
                        comment
                    ])
        
        print(f"    📊 {assignment}: {len(assignment_grades)} students → {filepath.name}")
        created_files.append(str(filepath))
    
    return created_files


def save_to_master_csv(grades: list, output_folder: str):
    """
    Append grades to a master CSV file that tracks ALL grades over time.
    This enables progress tracking across the entire school year.
    
    Columns:
    - Date, Student ID, Student Name, Period, Assignment, Unit, Quarter
    - Overall Score, Letter Grade
    - Content Accuracy, Completeness, Writing Quality, Effort & Engagement
    """
    master_file = Path(output_folder) / "master_grades.csv"
    
    # Check if file exists to determine if we need headers
    file_exists = master_file.exists()
    
    # Determine current quarter based on date
    today = datetime.now()
    month = today.month
    if month >= 8 and month <= 10:
        quarter = "Q1"
    elif month >= 11 or month <= 1:
        quarter = "Q2"
    elif month >= 2 and month <= 4:
        quarter = "Q3"
    else:
        quarter = "Q4"
    
    with open(master_file, 'a', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        
        # Write header if new file
        if not file_exists:
            writer.writerow([
                'Date', 'Student ID', 'Student Name', 'First Name', 'Last Name',
                'Period', 'Assignment', 'Unit', 'Quarter',
                'Overall Score', 'Letter Grade',
                'Content Accuracy', 'Completeness', 'Writing Quality', 'Effort Engagement',
                'AI Flag', 'AI Reason', 'Feedback'
            ])
        
        # Write each grade
        for grade in grades:
            if grade.get('student_id') == "UNKNOWN":
                continue
            
            breakdown = grade.get('breakdown', {})
            
            writer.writerow([
                today.strftime('%Y-%m-%d'),
                grade.get('student_id', ''),
                grade.get('student_name', ''),
                grade.get('first_name', ''),
                grade.get('last_name', ''),
                grade.get('period', ''),
                grade.get('assignment', ''),
                grade.get('unit', ''),
                grade.get('grading_period', quarter),  # Use grading_period from settings, fallback to calculated
                grade.get('score', 0),
                grade.get('letter_grade', ''),
                breakdown.get('content_accuracy', 0),
                breakdown.get('completeness', 0),
                breakdown.get('writing_quality', 0),
                breakdown.get('effort_engagement', 0),
                grade.get('authenticity_flag', 'clean'),
                grade.get('authenticity_reason', ''),
                grade.get('feedback', '')[:500]  # Truncate feedback for CSV
            ])
    
    print(f"📊 Updated master grades file: {master_file}")


def export_detailed_report(grades: list, output_folder: str, assignment_name: str) -> str:
    """
    Create detailed CSV with all grading information for your records.
    """
    safe_name = re.sub(r'[^\w\s-]', '', assignment_name).replace(' ', '_')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filepath = Path(output_folder) / f"Detailed_Report_{safe_name}_{timestamp}.csv"
    
    with open(filepath, 'w', newline='', encoding='utf-8') as f:
        writer = csv.writer(f)
        writer.writerow([
            'Student ID', 'Student Name', 'Email', 'Assignment', 'Score', 'Letter Grade',
            'Content (40)', 'Completeness (25)', 'Writing (20)', 'Effort (15)',
            'AI Flag', 'AI Reason', 'Feedback', 'Filename'
        ])

        for grade in grades:
            breakdown = grade.get('breakdown', {})
            writer.writerow([
                grade.get('student_id', ''),
                grade.get('student_name', ''),
                grade.get('email', ''),
                grade.get('assignment', ''),
                grade.get('score', ''),
                grade.get('letter_grade', ''),
                breakdown.get('content_accuracy', ''),
                breakdown.get('completeness', ''),
                breakdown.get('writing_quality', ''),
                breakdown.get('effort_engagement', breakdown.get('critical_thinking', '')),
                grade.get('authenticity_flag', 'clean'),
                grade.get('authenticity_reason', ''),
                grade.get('feedback', ''),
                grade.get('filename', '')
            ])
    
    print(f"📋 Detailed report saved: {filepath}")
    return str(filepath)


# =============================================================================
# MAIN GRADING WORKFLOW
# =============================================================================

def run_grading(
    assignment_folder: str = ASSIGNMENT_FOLDER,
    output_folder: str = OUTPUT_FOLDER,
    roster_file: str = ROSTER_FILE,
    assignment_name: str = ASSIGNMENT_NAME,
    create_outlook_emails: bool = False  # Set True if you have Outlook on Windows
):
    """
    Main function - runs the complete grading workflow.
    
    1. Loads student roster
    2. Reads each .docx file from assignment folder
    3. Grades each assignment with AI
    4. Generates emails (saves to files or creates Outlook drafts)
    5. Creates Focus CSV for grade import
    6. Creates detailed report for your records
    """
    print("=" * 60)
    print("📚 ASSIGNMENT GRADER - 6th Grade Social Studies")
    print("=" * 60)
    print(f"📁 Assignment folder: {assignment_folder}")
    print(f"💾 Output folder: {output_folder}")
    print(f"📝 Assignment: {assignment_name}")
    print()
    
    # Load roster
    roster = load_roster(roster_file)
    
    # Get all .docx files
    assignment_path = Path(assignment_folder)
    if not assignment_path.exists():
        print(f"❌ Assignment folder not found: {assignment_folder}")
        return []
    
    # Find all supported files (docx, txt, and images)
    supported_extensions = ['*.docx', '*.txt', '*.jpg', '*.jpeg', '*.png', '*.gif', '*.webp']
    all_files = []
    for ext in supported_extensions:
        all_files.extend(assignment_path.glob(ext))
    
    print(f"📄 Found {len(all_files)} files ({', '.join(supported_extensions)})")
    print()
    
    # Process each file
    all_grades = []
    
    for i, filepath in enumerate(all_files, 1):
        print(f"[{i}/{len(all_files)}] {filepath.name}")
        
        # Parse filename to get student name
        parsed = parse_filename(filepath.name)
        student_name = f"{parsed['first_name']} {parsed['last_name']}"
        lookup_key = parsed['lookup_key']
        
        # Look up student in roster
        if lookup_key in roster:
            student_info = roster[lookup_key].copy()
            print(f"  👤 {student_info['student_name']} (ID: {student_info['student_id']})")
        else:
            # Not found in roster - use parsed name
            student_info = {
                "student_id": "UNKNOWN",
                "student_name": student_name,
                "first_name": parsed['first_name'],
                "last_name": parsed['last_name'],
                "email": ""
            }
            print(f"  👤 {student_name} (⚠️ NOT IN ROSTER)")
        
        # Read file content
        file_data = read_assignment_file(filepath)
        if not file_data:
            print(f"  ❌ Could not read file")
            continue
        
        # Handle based on file type
        markers_found = []
        
        if file_data["type"] == "text":
            # Text-based file - check for markers
            content = file_data["content"]
            
            if len(content.strip()) < 20:
                print(f"  ⚠️  File appears empty ({len(content)} chars)")
                continue
            
            # Extract only the student work portion
            student_work, markers_found = extract_student_work(content)
            
            if markers_found:
                print(f"  📝 Found marker(s): {', '.join(markers_found[:2])}{'...' if len(markers_found) > 2 else ''}")
            else:
                print(f"  ⚠️  NO MARKERS FOUND - Check if student uploaded wrong document!")
                print(f"      → Review manually: {filepath.name}")
            
            if len(student_work.strip()) < 10:
                print(f"  ⚠️  No student work found after marker")
                continue
            
            # Prepare data for grading
            grade_data = {"type": "text", "content": student_work}
        
        elif file_data["type"] == "image":
            # Image file - send entire image to AI for grading
            print(f"  🖼️  Image file - sending to AI for visual grading")
            grade_data = file_data
            markers_found = ["image"]  # Mark as having content
        
        else:
            print(f"  ❌ Unknown file type")
            continue
        
        # Grade with AI
        grade_result = grade_assignment(student_info['student_name'], grade_data)
        
        # Combine all info
        # Use assignment_name parameter (from config) instead of extracting from filename
        # Filename extraction can truncate long names
        grade_record = {
            **student_info,
            **grade_result,
            "filename": filepath.name,
            "assignment": assignment_name,
            "has_markers": len(markers_found) > 0
        }
        all_grades.append(grade_record)
        
        print(f"  ✅ Score: {grade_result['score']} ({grade_result['letter_grade']})")
        print()
    
    # Export results
    print("=" * 60)
    print("📊 EXPORTING RESULTS")
    print("=" * 60)
    
    # Focus CSVs (separated by assignment)
    focus_files = export_focus_csv(all_grades, output_folder, assignment_name)
    
    # Detailed report (one file with all grades)
    export_detailed_report(all_grades, output_folder, assignment_name)
    
    # Emails
    if create_outlook_emails:
        if not create_outlook_drafts(all_grades):
            save_emails_to_folder(all_grades, output_folder)
    else:
        save_emails_to_folder(all_grades, output_folder)
    
    # Summary
    print()
    print("=" * 60)
    print("📈 GRADING SUMMARY")
    print("=" * 60)
    
    if all_grades:
        scores = [g['score'] for g in all_grades if g['score'] > 0]
        if scores:
            print(f"Total graded: {len(all_grades)}")
            print(f"Average score: {sum(scores)/len(scores):.1f}")
            print(f"Highest: {max(scores)}")
            print(f"Lowest: {min(scores)}")
            
            # Grade distribution
            grade_dist = {}
            for g in all_grades:
                letter = g['letter_grade']
                grade_dist[letter] = grade_dist.get(letter, 0) + 1
            print(f"Distribution: {dict(sorted(grade_dist.items()))}")
            
            # Per-assignment breakdown
            print(f"\n📚 By Assignment:")
            assignments = {}
            for g in all_grades:
                a = g.get('assignment', 'Unknown')
                if a not in assignments:
                    assignments[a] = []
                assignments[a].append(g['score'])
            
            for assignment, scores_list in sorted(assignments.items()):
                valid_scores = [s for s in scores_list if s > 0]
                if valid_scores:
                    avg = sum(valid_scores) / len(valid_scores)
                    print(f"   • {assignment[:40]}: {len(scores_list)} students, avg {avg:.1f}")
        
        # List students not in roster
        unknown = [g for g in all_grades if g['student_id'] == 'UNKNOWN']
        if unknown:
            print(f"\n⚠️  {len(unknown)} students NOT FOUND in roster:")
            for g in unknown:
                print(f"   - {g['student_name']} ({g['filename']})")
        
        # List documents with no markers (possible wrong uploads)
        no_markers = [g for g in all_grades if not g.get('has_markers', True)]
        if no_markers:
            print(f"\n🚨 {len(no_markers)} documents had NO MARKERS - review for possible wrong uploads:")
            for g in no_markers:
                print(f"   - {g['student_name']}: {g['filename']}")
    
    print()
    print("✅ GRADING COMPLETE!")
    return all_grades


# =============================================================================
# RUN THE SCRIPT
# =============================================================================

if __name__ == "__main__":
    # Update the paths at the top of this file, then run!
    results = run_grading(
        create_outlook_emails=False  # Set True if you have Outlook desktop on Windows
    )
