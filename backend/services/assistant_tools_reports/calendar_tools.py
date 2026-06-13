"""Teaching-calendar tools: view, schedule, unschedule, holidays, plus the
curriculum-map date parser they rely on.

Pure-move of whole functions out of the former single-file module; bodies
are byte-identical.
"""
import os
import re
import json
import logging
from datetime import datetime, timedelta

from backend.services.assistant_tools import (
    _load_calendar, _save_calendar, DOCUMENTS_DIR,
)
from backend.utils.compliance import require_teacher_id
import sentry_sdk

_logger = logging.getLogger(__name__)


def _find_curriculum_map_unit_candidates(doc, req_start, req_end, _parse_map_date):
    """Pass 1: scan all tables for unit+date-range rows that overlap [req_start, req_end].

    Returns a list of (span_days, unit_name, dates_str, weeks, table_index) tuples.
    Extracted from _parse_curriculum_map_for_dates to keep the parent ≤200 LOC.
    """
    import re
    candidates = []  # (span_days, unit_name, dates_str, weeks, table_index)

    # First pass: scan detail table headers + benchmark rows
    current_unit_header = None
    current_table_idx = None
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            seen_cells = set()
            cells = []
            for c in row.cells:
                t = c.text.strip()
                if t and t not in seen_cells:
                    seen_cells.add(t)
                    cells.append(t)
            row_text = ' '.join(cells)

            # Detect unit header rows like "Unit 7: Manifest Destiny (4 weeks)"
            unit_header = re.search(r'(Unit\s+\d+:\s*[^\(]+)', row_text)
            if unit_header and ('weeks' in row_text.lower() or 'Standard' in row_text):
                current_unit_header = unit_header.group(1).strip()
                current_table_idx = ti
                continue

            # Look for date ranges in benchmark rows
            row_lower = row_text.lower()
            # Pattern 1: "February 17th - March 12th" (both have month)
            date_ranges = re.findall(
                r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z.]*\s+\d{1,2}(?:st|nd|rd|th)?)\s*[-–—]\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z.]*\s+\d{1,2}(?:st|nd|rd|th)?)',
                row_lower, re.IGNORECASE
            )
            # Pattern 2: "Feb. 2nd -6th" (end date has no month — inherit from start)
            short_ranges = re.findall(
                r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z.]*\s+\d{1,2}(?:st|nd|rd|th)?)\s*[-–—]\s*(\d{1,2}(?:st|nd|rd|th)?)\b',
                row_lower, re.IGNORECASE
            )
            for d_start_str, day_only in short_ranges:
                # Infer month from start date
                month_match = re.match(r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z.]*)', d_start_str, re.IGNORECASE)
                if month_match:
                    d_end_str = f"{month_match.group(1)} {day_only}"
                    date_ranges.append((d_start_str, d_end_str))

            for d_start_str, d_end_str in date_ranges:
                d_start = _parse_map_date(d_start_str)
                d_end = _parse_map_date(d_end_str)
                if d_start and d_end and d_end >= req_start and d_start <= req_end:
                    span = (d_end - d_start).days
                    # Use current_unit_header if in a detail table, else look in same row
                    unit_name = current_unit_header
                    if not unit_name:
                        for cell_text in cells:
                            um = re.search(r'(Unit\s+\d+:\s*[^|]+)', cell_text)
                            if um:
                                unit_name = um.group(1).strip()
                                break
                    if unit_name:
                        # Extract week numbers from range patterns like "25-28"
                        joined = ' '.join(cells)
                        week_ranges = re.findall(r'\b(\d{1,2})-(\d{1,2})\b', joined)
                        weeks = None
                        if week_ranges:
                            s, e = int(week_ranges[0][0]), int(week_ranges[0][1])
                            if 1 <= s <= 52 and 1 <= e <= 52:
                                weeks = list(range(s, e + 1))
                        candidates.append((span, unit_name, f"{d_start_str.strip()} – {d_end_str.strip()}", weeks, ti))
    return candidates


def _parse_curriculum_map_for_dates(start_date, end_date):
    """Parse curriculum map DOCX and return structured data for a date range.

    Looks for unit/week entries whose date ranges overlap with the requested dates.
    Returns dict with unit name, benchmarks, vocabulary, textbook, resources, or None.
    """
    import re
    from datetime import datetime as _dt

    # Find curriculum map file
    if not os.path.isdir(DOCUMENTS_DIR):
        return None
    curriculum_file = None
    for fname in os.listdir(DOCUMENTS_DIR):
        if not fname.lower().endswith('.docx') or fname.endswith('.meta.json'):
            continue  # Only parse DOCX files (python-docx can't read PDFs)
        meta_path = os.path.join(DOCUMENTS_DIR, fname + ".meta.json")
        if os.path.exists(meta_path):
            try:
                with open(meta_path, 'r') as f:
                    meta = json.load(f)
                if meta.get("doc_type") == "curriculum":
                    curriculum_file = os.path.join(DOCUMENTS_DIR, fname)
                    break
            except Exception:  # noqa: BLE001  # broad catch: error is logged
                _logger.debug("curriculum document metadata read failed", exc_info=True)
                continue
        elif 'curriculum' in fname.lower() or 'pacing' in fname.lower():
            curriculum_file = os.path.join(DOCUMENTS_DIR, fname)
            break
    if not curriculum_file or not os.path.exists(curriculum_file):
        return None

    try:
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.table import Table
    except ImportError:
        return None

    # Parse target date range
    try:
        req_start = _dt.strptime(start_date, '%Y-%m-%d')
        req_end = _dt.strptime(end_date, '%Y-%m-%d')
    except Exception:  # noqa: BLE001  # broad catch: returns fallback
        return None

    # Helper to parse flexible date strings like "Feb. 17th", "March 12th", "January 6th"
    def _parse_map_date(s, year=None):
        if not year:
            year = req_start.year
        s = re.sub(r'(st|nd|rd|th)\b', '', s.strip().rstrip('.'))
        for fmt in ('%B %d', '%b %d', '%b. %d'):
            try:
                d = _dt.strptime(s.strip(), fmt)
                return d.replace(year=year)
            except ValueError:
                continue
        return None

    doc = Document(curriculum_file)

    # Pass 1: Find unit that covers these dates (delegated to module-level helper).
    candidates = _find_curriculum_map_unit_candidates(doc, req_start, req_end, _parse_map_date)

    if not candidates:
        return None

    # Pick the narrowest matching date range
    candidates.sort(key=lambda x: x[0])
    _, matching_unit, matching_dates, matching_weeks, _ = candidates[0]

    # Pass 2: Find the unit's detail section — benchmarks, vocabulary, textbook, resources
    unit_keyword = matching_unit.split(':')[0].strip()  # e.g., "Unit 7"
    benchmarks = []
    vocabulary = []
    textbook = []
    resources = {"nearpod_activities": [], "nearpod_lessons": [], "videos": [], "dbqs": []}
    in_unit_section = False

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            row_text = ' '.join(cells)

            # Detect start of our unit's detail section
            if unit_keyword in row_text and ('Standard' in row_text or 'weeks' in row_text.lower()):
                in_unit_section = True
                continue

            # Detect end — next unit header or resources section end
            if in_unit_section:
                other_unit = re.search(r'Unit\s+\d+:', row_text)
                if other_unit and unit_keyword not in row_text:
                    in_unit_section = False
                    continue

            if not in_unit_section:
                continue

            # Extract benchmarks (format: "Quarter 3 | 25-28 | dates | SS.8.A.4.1 | description")
            for cell_text in cells:
                bm_match = re.match(r'(SS\.\d+\.[A-Z]+\.\d+(?:\.\d+)?)\s*$', cell_text.strip())
                if bm_match:
                    code = bm_match.group(1)
                    # Next non-empty cell after the code is the description
                    idx = cells.index(cell_text)
                    desc = cells[idx + 1] if idx + 1 < len(cells) else ""
                    if desc and not desc.startswith('SS.') and not desc.startswith('Quarter'):
                        benchmarks.append({"code": code, "description": desc})
                    elif not any(b["code"] == code for b in benchmarks):
                        benchmarks.append({"code": code, "description": ""})

            # Extract vocabulary
            if 'academic vocabulary' in row_text.lower():
                for cell_text in cells:
                    if 'academic vocabulary' in cell_text.lower():
                        vocab_text = re.sub(r'^academic vocabulary:?\s*', '', cell_text, flags=re.IGNORECASE)
                        # Handle typos: period after word (3+ chars) followed by space+uppercase → comma
                        # Skip initials like "F." in "Stephen F. Austin"
                        vocab_text = re.sub(r'(?<=[a-z]{3})\.\s+(?=[A-Z\u201c"])', ', ', vocab_text)
                        terms = [t.strip().strip('"').strip('\u201c').strip('\u201d') for t in vocab_text.split(',')]
                        vocabulary = [t for t in terms if t and len(t) > 1]
                        break

            # Extract textbook references
            if 'chapter' in row_text.lower() and ('pgs' in row_text.lower() or 'teacher textbook' in row_text.lower()):
                for cell_text in cells:
                    ch_match = re.search(r'Chapter\s+\d+\s*\(pgs?\.\s*[\w\-\u2013]+\)', cell_text)
                    if ch_match and ch_match.group() not in textbook:
                        textbook.append(ch_match.group())

            # Extract Nearpod/video resources
            if 'nearpod' in row_text.lower() or 'video' in row_text.lower() or 'dbq' in row_text.lower():
                for cell_text in cells:
                    cl = cell_text.lower()
                    if 'activities:' in cl:
                        items = re.sub(r'^activities:\s*', '', cell_text, flags=re.IGNORECASE)
                        resources["nearpod_activities"] = [i.strip() for i in items.split('\n') if i.strip()]
                    elif 'lessons:' in cl and 'nearpod' in ' '.join(cells).lower():
                        items = re.sub(r'^lessons:\s*', '', cell_text, flags=re.IGNORECASE)
                        resources["nearpod_lessons"] = [i.strip() for i in items.split('\n') if i.strip()]
                    elif 'videos:' in cl:
                        items = re.sub(r'^videos:\s*', '', cell_text, flags=re.IGNORECASE)
                        resources["videos"] = [i.strip() for i in items.split('\n') if i.strip()]
                    elif 'dbq' in cl:
                        dbq_items = re.findall(r'(?:DBQ\s*\|?\s*)?(.+?\(SS\.[^)]+\))', cell_text)
                        resources["dbqs"] = [d.strip() for d in dbq_items if d.strip()]

    # Deduplicate benchmarks
    seen_codes = set()
    unique_benchmarks = []
    for b in benchmarks:
        if b["code"] not in seen_codes:
            seen_codes.add(b["code"])
            unique_benchmarks.append(b)

    # Clean empty resource lists
    resources = {k: v for k, v in resources.items() if v}

    return {
        "unit": matching_unit,
        "weeks": matching_weeks,
        "dates": matching_dates,
        "benchmarks": unique_benchmarks,
        "vocabulary": vocabulary,
        "textbook": textbook,
        "resources": resources,
    }


def get_calendar(start_date=None, end_date=None, teacher_id='local-dev'):
    """Read the teaching calendar for a date range."""
    require_teacher_id(teacher_id)
    cal = _load_calendar(teacher_id)

    if not start_date:
        start_date = datetime.now().strftime('%Y-%m-%d')
    if not end_date:
        end_dt = datetime.strptime(start_date, '%Y-%m-%d') + timedelta(days=7)
        end_date = end_dt.strftime('%Y-%m-%d')

    # Filter lessons in range
    lessons = [s for s in cal.get("scheduled_lessons", [])
               if start_date <= s.get("date", "") <= end_date]
    lessons.sort(key=lambda s: s.get("date", ""))

    # Filter holidays in range (including multi-day overlaps)
    holidays = []
    for h in cal.get("holidays", []):
        h_start = h.get("date", "")
        h_end = h.get("end_date", h_start)
        if h_end >= start_date and h_start <= end_date:
            holidays.append(h)

    result = {
        "start_date": start_date,
        "end_date": end_date,
        "scheduled_lessons": lessons,
        "holidays": holidays,
        "total_lessons": len(lessons),
        "total_holidays": len(holidays),
    }

    # Always include curriculum map data for curricular context
    try:
        curriculum_data = _parse_curriculum_map_for_dates(start_date, end_date)
        if curriculum_data:
            result["curriculum_map"] = curriculum_data
            if not lessons:
                result["note"] = ("No lessons scheduled for this period. Curriculum map data shows what should be covered. "
                                  "Also check uploaded reference documents in your system context for additional details.")
    except Exception as e:  # noqa: BLE001  # broad catch: error is logged
        sentry_sdk.capture_exception(e)

    return result


def schedule_lesson_tool(date, lesson_title, unit=None, day_number=None, lesson_file=None, teacher_id='local-dev'):
    """Schedule a lesson on the teaching calendar."""
    require_teacher_id(teacher_id)
    import uuid as _uuid
    if not date or not lesson_title:
        return {"error": "date and lesson_title are required"}

    cal = _load_calendar(teacher_id)

    # Pick a color based on unit name
    unit_colors = ['#6366f1', '#8b5cf6', '#ec4899', '#f59e0b', '#22c55e', '#06b6d4', '#ef4444']
    color_idx = hash(unit or '') % len(unit_colors)

    entry = {
        "id": str(_uuid.uuid4()),
        "date": date,
        "unit": unit or "",
        "lesson_title": lesson_title,
        "day_number": day_number,
        "lesson_file": lesson_file or "",
        "color": unit_colors[color_idx],
    }

    # Remove existing lesson on same date with same day_number or title to avoid duplicates
    cal["scheduled_lessons"] = [
        s for s in cal["scheduled_lessons"]
        if not (s["date"] == date and (
            (day_number is not None and s.get("day_number") == day_number) or
            s.get("lesson_title") == lesson_title
        ))
    ]
    cal["scheduled_lessons"].append(entry)
    _save_calendar(cal, teacher_id)

    return {"status": "scheduled", "entry": entry}


def unschedule_lesson_tool(date, lesson_title=None, teacher_id='local-dev'):
    """Remove a lesson from the teaching calendar by date and optional title."""
    require_teacher_id(teacher_id)
    if not date:
        return {"error": "date is required"}

    cal = _load_calendar(teacher_id)
    before = len(cal["scheduled_lessons"])

    if lesson_title:
        cal["scheduled_lessons"] = [
            s for s in cal["scheduled_lessons"]
            if not (s["date"] == date and s.get("lesson_title") == lesson_title)
        ]
    else:
        cal["scheduled_lessons"] = [
            s for s in cal["scheduled_lessons"]
            if s["date"] != date
        ]

    removed = before - len(cal["scheduled_lessons"])
    if removed == 0:
        return {"status": "not_found", "message": f"No lessons found on {date}"}

    _save_calendar(cal, teacher_id)
    return {"status": "removed", "removed_count": removed}


def add_calendar_holiday(date, name, end_date=None, teacher_id='local-dev'):
    """Add a holiday or break to the teaching calendar."""
    require_teacher_id(teacher_id)
    if not date or not name:
        return {"error": "date and name are required"}

    cal = _load_calendar(teacher_id)

    holiday = {"date": date, "name": name}
    if end_date:
        holiday["end_date"] = end_date

    # Remove existing holiday on same date to avoid duplicates
    cal["holidays"] = [h for h in cal["holidays"] if h["date"] != date]
    cal["holidays"].append(holiday)
    cal["holidays"].sort(key=lambda h: h["date"])

    _save_calendar(cal, teacher_id)
    return {"status": "added", "holiday": holiday}
