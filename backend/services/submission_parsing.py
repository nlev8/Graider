"""Submission parsing for the grading pipeline: extract student info from a
filename + read submission file contents. Flask-free (pathlib / base64 / file I/O —
no LLM / network) extracted from assignment_grader.py. Wave 7 (grading-engine
decomposition). Diagnostic output uses the module logger (the grader's debug prints
became _logger calls on extraction — return values are unchanged).
"""
import logging
import re
from pathlib import Path

_logger = logging.getLogger(__name__)


def parse_filename(filename: str) -> dict:
    """
    Parse student info from filename.

    Expected formats:
        FirstName_LastName_AssignmentName.docx
        Last, First M._AssignmentName.docx
    Examples:
        A'kareah_West_Cornell Notes_ Political Parties.docx
        Eli_Long_Hamilton_Jefferson_Graphic_Organizer.docx
        Deloach, Rylee M._Washington_Stations_Handout.docx

    Returns: {"first_name": ..., "last_name": ..., "assignment_part": ...}
    """
    # Remove extension
    name = Path(filename).stem

    # Handle "Last, First M._Assignment" format (comma before first underscore)
    first_underscore = name.find('_')
    if first_underscore > 0 and ',' in name[:first_underscore]:
        name_part = name[:first_underscore]
        assignment_part = name[first_underscore + 1:] if first_underscore < len(name) - 1 else ""
        comma_parts = name_part.split(',')
        last_name = comma_parts[0].strip()
        # First name may include middle initial like "Rylee M."
        first_full = comma_parts[1].strip() if len(comma_parts) > 1 else ""
        first_name = first_full.split()[0] if first_full else ""

        # Strip apostrophes/curly quotes so Da'Juan matches Dajuan
        key = f"{first_name} {last_name}".lower()
        key = key.replace("'", "").replace("\u2019", "")
        return {
            "first_name": first_name,
            "last_name": last_name,
            "assignment_part": assignment_part,
            "lookup_key": key
        }

    # Standard format: FirstName_LastName_AssignmentName
    parts = name.split('_')

    if len(parts) >= 2:
        first_name = parts[0].strip()
        last_name = parts[1].strip()
        assignment_part = '_'.join(parts[2:]) if len(parts) > 2 else ""

        # Strip apostrophes/curly quotes so Da'Juan matches Dajuan
        key = f"{first_name} {last_name}".lower()
        key = key.replace("'", "").replace("\u2019", "")
        return {
            "first_name": first_name,
            "last_name": last_name,
            "assignment_part": assignment_part,
            "lookup_key": key
        }
    else:
        # Can't parse - return filename as-is
        return {
            "first_name": name,
            "last_name": "",
            "assignment_part": "",
            "lookup_key": name.lower().replace("'", "").replace("\u2019", "")
        }


def read_image_file(filepath: str) -> dict:
    """
    Read an image file and return it as base64 for GPT-4o vision.

    Returns dict with:
    - type: "image"
    - data: base64 encoded image
    - media_type: image MIME type
    """
    import base64

    filepath = Path(filepath)
    extension = filepath.suffix.lower()

    # Map extensions to MIME types
    mime_types = {
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.png': 'image/png',
        '.gif': 'image/gif',
        '.webp': 'image/webp'
    }

    if extension not in mime_types:
        _logger.warning("Unsupported image type: %s", extension)
        return None

    try:
        with open(filepath, 'rb') as f:
            image_data = base64.b64encode(f.read()).decode('utf-8')

        return {
            "type": "image",
            "data": image_data,
            "media_type": mime_types[extension]
        }
    except Exception as e:
        _logger.warning("Error reading image: %s", e)
        return None


def read_docx_file(filepath: str) -> str:
    """
    Read text content from a Word document (.docx) in document order.
    This properly interleaves paragraphs and tables as they appear.
    """
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        _logger.error("python-docx not installed. Run: pip install python-docx")
        return None

    try:
        doc = Document(filepath)
        full_text = []

        # Iterate through document body elements in order
        # This ensures tables and paragraphs appear in their actual document order
        for element in doc.element.body:
            # Check if it's a paragraph
            if element.tag.endswith('p'):
                para = Paragraph(element, doc)
                if para.text.strip():
                    full_text.append(para.text)
            # Check if it's a table
            elif element.tag.endswith('tbl'):
                table = Table(element, doc)
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))

        return '\n'.join(full_text)
    except Exception as e:
        _logger.warning("Error reading file: %s", e)
        return None


def extract_from_graider_text(document_text, exclude_markers=None):
    """Extract student responses from plain text containing [GRAIDER:TYPE:ID] markers.

    Fallback for when structured table reading fails (e.g., tables were flattened
    by Google Docs, copy-paste, or format conversion).  Parses the text between
    consecutive GRAIDER markers to capture student answers.

    Args:
        document_text: Plain text that may contain [GRAIDER:...] tags.
        exclude_markers: List of section names to skip.

    Returns:
        Same shape as extract_from_tables(), or None if no GRAIDER tags found.
    """
    tag_pattern = re.compile(r'\[GRAIDER:(VOCAB|QUESTION|SUMMARY):([^\]]+)\]')
    matches = list(tag_pattern.finditer(document_text))

    if not matches:
        return None

    _logger.info("Graider text fallback: Found %d GRAIDER markers in plain text", len(matches))

    extracted = []
    blank_questions = []
    excluded_sections = []
    exclude_lower = [em.lower().strip() for em in exclude_markers] if exclude_markers else []

    type_map = {
        "VOCAB": "vocab_term",
        "QUESTION": "numbered_question",
        "SUMMARY": "summary"
    }

    for i, match in enumerate(matches):
        tag_type = match.group(1)
        tag_id = match.group(2)

        # Text after the tag up to the next tag (or GRAIDER_TABLE_V1 marker or end of doc)
        start = match.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(document_text)
        raw_block = document_text[start:end]
        # Truncate at GRAIDER_TABLE_V1 marker if present (end-of-worksheet sentinel)
        marker_pos = raw_block.find('GRAIDER_TABLE_V1')
        if marker_pos != -1:
            raw_block = raw_block[:marker_pos]
        raw_block = raw_block.strip()

        # The block contains: "  visible_header  (N pts)\nstudent answer\n..."
        # Split into lines, skip the header line (contains the term/question + pts),
        # and grab everything else as the student response.
        lines = [ln.strip() for ln in raw_block.split('\n') if ln.strip()]

        # First non-empty line is typically the visible header (term + pts)
        header = lines[0] if lines else ""
        # Student response is everything after the header line,
        # filtering out section headers and metadata that appear between markers
        section_headers = {'vocabulary', 'questions', 'summary', 'question', 'vocab'}
        response_lines = []
        for ln in lines[1:]:
            # Skip section headers and metadata
            if ln.lower() in section_headers:
                continue
            if 'GRAIDER_TABLE_V1' in ln:
                continue
            response_lines.append(ln)
        response = '\n'.join(response_lines).strip()
        # Strip placeholder text from response cell
        response = response.replace("Type your answer here...", "").replace("Your Answer:", "").strip()

        # Fix: Student typed answer on same line as header (no newline separation)
        # Common when student types in the header cell of a Graider table
        if not response or len(re.sub(r'[_\s]', '', response)) < 2:
            pts_match = re.search(r'\(\d+\s*pts?\)', header)
            if pts_match:
                after_pts = header[pts_match.end():].strip()
                if after_pts and len(re.sub(r'[_\s]', '', after_pts)) >= 2:
                    response = after_pts
                    header = header[:pts_match.end()].strip()
                    _logger.debug("Text fallback: recovered same-line answer for %s:%s", tag_type, tag_id)

            # Summary fallback: no (N pts) marker — check if prompt is followed by student text
            if (not response or len(re.sub(r'[_\s]', '', response)) < 2) and tag_type == "SUMMARY":
                prompt_keywords = {'write', 'explain', 'summarize', 'describe', 'sentence', 'summary', 'paragraph'}
                best_split = None
                for m in re.finditer(r'[.?!]\s+', header):
                    before = header[:m.start() + 1].lower()
                    after = header[m.end():].strip()
                    after_clean = re.sub(r'[_\s]', '', after)
                    has_prompt_word = any(kw in before for kw in prompt_keywords)
                    if has_prompt_word and len(after_clean) >= 20:
                        best_split = m
                        break
                if best_split:
                    response = header[best_split.end():].strip()
                    header = header[:best_split.start() + 1].strip()
                    _logger.debug("Text fallback: recovered same-line answer for SUMMARY")

        # Check exclusion
        header_lower = header.lower()
        is_excluded = any(em in header_lower for em in exclude_lower)
        if is_excluded:
            excluded_sections.append(header)
            continue

        # Build question label
        if tag_type == "VOCAB":
            question = tag_id
        elif tag_type == "QUESTION":
            question = header
        elif tag_type == "SUMMARY":
            question = "Summary"
        else:
            question = header

        # Check if blank
        response_cleaned = re.sub(r'[_\s]', '', response)
        if len(response_cleaned) < 2:
            # Deduplicate
            if question not in blank_questions:
                blank_questions.append(question)
            continue

        # Deduplicate: skip if already extracted for this tag
        already_extracted = any(
            e.get("tag_id") == tag_id and e.get("section") == tag_type
            for e in extracted
        )
        if already_extracted:
            continue

        extracted.append({
            "question": question,
            "answer": response,
            "type": type_map.get(tag_type, "numbered_question"),
            "section": tag_type,
            "tag_id": tag_id
        })

    # Cross-reference: remove blank entries that were successfully extracted
    extracted_questions = {e.get("question", "").lower().strip() for e in extracted}
    blank_questions = [
        bq for bq in blank_questions
        if bq.lower().strip() not in extracted_questions
    ]

    total_q = len(extracted) + len(blank_questions)
    answered_q = len(extracted)
    summary = f"Graider text fallback: Found {answered_q} responses out of {total_q} sections."
    if blank_questions:
        summary += f" {len(blank_questions)} left blank."

    _logger.info("Graider text extraction: %d/%d answered", answered_q, total_q)

    return {
        "extracted_responses": extracted,
        "blank_questions": blank_questions,
        "total_questions": max(total_q, 1),
        "answered_questions": answered_q,
        "extraction_summary": summary,
        "excluded_sections": excluded_sections,
        "missing_sections": []  # Text fallback has no separate missing sections (avoids double-counting with blank_questions)
    }


def extract_from_tables(table_data, exclude_markers=None):
    """Extract student responses from Graider structured table data.

    Maps table entries to the same dict format as extract_student_responses().

    Args:
        table_data: List of dicts from read_docx_file_structured()["tables"]
        exclude_markers: List of section names to skip

    Returns:
        Same shape as extract_student_responses():
        {
            "extracted_responses": [{"question": ..., "answer": ..., "type": ...}],
            "blank_questions": [...],
            "total_questions": int,
            "answered_questions": int,
            "extraction_summary": str
        }
    """
    extracted = []
    blank_questions = []
    excluded_sections = []
    exclude_lower = [em.lower().strip() for em in exclude_markers] if exclude_markers else []

    type_map = {
        "VOCAB": "vocab_term",
        "QUESTION": "numbered_question",
        "SUMMARY": "summary"
    }

    for entry in table_data:
        tag_type = entry.get("tag_type", "")
        tag_id = entry.get("tag_id", "")
        header = entry.get("header_text", "")
        response = entry.get("response", "")
        # Strip placeholder text from response
        response = response.replace("Type your answer here...", "").replace("Your Answer:", "").strip()

        # Check exclusion
        header_lower = header.lower()
        is_excluded = any(em in header_lower for em in exclude_lower)
        if is_excluded:
            excluded_sections.append(header)
            continue

        # Build question label
        if tag_type == "VOCAB":
            question = tag_id  # The term name
        elif tag_type == "QUESTION":
            question = header  # e.g. "1) What is photosynthesis?"
        elif tag_type == "SUMMARY":
            question = "Summary"
        else:
            question = header

        # Check if blank
        response_cleaned = re.sub(r'[_\s]', '', response)
        if len(response_cleaned) < 2:
            # Backup: Student may have typed answer in the header cell after (N pts)
            # read_docx_file_structured has Cases 1-3 for this, but they can fail
            # when Word/Google Docs alters the cell structure
            pts_match = re.search(r'\(\d+\s*pts?\)', header)
            if pts_match:
                after_pts = header[pts_match.end():].strip()
                if after_pts and len(re.sub(r'[_\s]', '', after_pts)) >= 2:
                    response = after_pts
                    header = header[:pts_match.end()].strip()
                    # Rebuild question label with trimmed header
                    if tag_type == "QUESTION":
                        question = header
                    _logger.debug("Recovered same-line answer from header for %s:%s", tag_type, tag_id)

            # Summary fallback: no (N pts) marker — check if prompt is followed by student text
            # Find where the prompt ends and the student answer begins.
            # Strategy: find ALL sentence boundaries and pick the split where
            # the text before contains prompt keywords and text after is long enough.
            if len(re.sub(r'[_\s]', '', response)) < 2 and tag_type == "SUMMARY":
                prompt_keywords = {'write', 'explain', 'summarize', 'describe', 'sentence', 'summary', 'paragraph'}
                best_split = None
                for m in re.finditer(r'[.?!]\s+', header):
                    before = header[:m.start() + 1].lower()
                    after = header[m.end():].strip()
                    after_clean = re.sub(r'[_\s]', '', after)
                    # The text before the split should contain prompt language
                    has_prompt_word = any(kw in before for kw in prompt_keywords)
                    # The text after should be substantial (student wrote something real)
                    if has_prompt_word and len(after_clean) >= 20:
                        best_split = m
                        break  # Take the first valid split after prompt keywords
                if best_split:
                    response = header[best_split.end():].strip()
                    header = header[:best_split.start() + 1].strip()
                    question = "Summary"
                    _logger.debug("Recovered same-line answer from summary header")

        # After recovery attempts, final blank check
        response_cleaned = re.sub(r'[_\s]', '', response)
        if len(response_cleaned) < 2:
            # Deduplicate: don't add if this question is already in blank_questions
            if question not in blank_questions:
                blank_questions.append(question)
            continue

        # Deduplicate: don't add if we already have an extracted response for this tag
        already_extracted = any(
            e.get("tag_id") == tag_id and e.get("section") == tag_type
            for e in extracted
        )
        if already_extracted:
            continue

        extracted.append({
            "question": question,
            "answer": response,
            "type": type_map.get(tag_type, "numbered_question"),
            "section": tag_type,
            "tag_id": tag_id
        })

    # Build a set of extracted question labels to remove them from blank_questions
    extracted_questions = {e.get("question", "").lower().strip() for e in extracted}
    blank_questions = [
        bq for bq in blank_questions
        if bq.lower().strip() not in extracted_questions
    ]

    total_q = len(extracted) + len(blank_questions)
    answered_q = len(extracted)
    summary = "Table extraction: Found " + str(answered_q) + " responses out of " + str(total_q) + " sections."
    if blank_questions:
        summary += " " + str(len(blank_questions)) + " left blank."
    if excluded_sections:
        summary += " Excluded " + str(len(excluded_sections)) + " section(s)."

    _logger.info("Table extraction: %d/%d answered", answered_q, total_q)

    return {
        "extracted_responses": extracted,
        "blank_questions": blank_questions,
        "total_questions": max(total_q, 1),
        "answered_questions": answered_q,
        "extraction_summary": summary,
        "excluded_sections": excluded_sections,
        "missing_sections": []  # Table format has no separate missing sections (avoids double-counting with blank_questions)
    }


def read_docx_file_structured(filepath: str) -> dict:
    """Read a .docx file and detect Graider structured tables.

    Iterates through doc.element.body looking for 2-row tables whose first cell
    contains a [GRAIDER:TYPE:ID] tag. Also checks for the GRAIDER_TABLE_V1 marker.

    Returns:
        {
            "is_graider_table": bool,
            "plain_text": str (full document text for fallback),
            "tables": [
                {"tag_type": "VOCAB"|"QUESTION"|"SUMMARY",
                 "tag_id": str, "header_text": str, "response": str},
                ...
            ]
        }
    """
    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
        from docx.shared import Pt
    except ImportError:
        return {"is_graider_table": False, "plain_text": None, "tables": []}

    try:
        doc = Document(filepath)
        tables_found = []
        full_text = []
        has_marker = False
        tag_pattern = re.compile(r'\[GRAIDER:(VOCAB|QUESTION|SUMMARY):([^\]]+)\]')

        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = Paragraph(element, doc)
                text = para.text.strip()
                if text:
                    full_text.append(text)
                    if text == "GRAIDER_TABLE_V1":
                        has_marker = True

            elif element.tag.endswith('tbl'):
                table = Table(element, doc)
                rows = table.rows
                if len(rows) >= 2:
                    header_cell_text = rows[0].cells[0].text
                    match = tag_pattern.search(header_cell_text)
                    if match:
                        tag_type = match.group(1)
                        tag_id = match.group(2)
                        # Strip the hidden tag from visible text
                        visible_header = tag_pattern.sub('', header_cell_text).strip()
                        # Collect response text from ALL rows after the header
                        # (handles cases where Enter key or word processors add extra rows)
                        response_parts = []
                        for r_idx in range(1, len(rows)):
                            cell_text = rows[r_idx].cells[0].text.strip()
                            if cell_text:
                                response_parts.append(cell_text)
                        response = '\n'.join(response_parts)
                        # Strip placeholder text from response cell
                        response = response.replace("Type your answer here...", "").replace("Your Answer:", "").strip()

                        # Fix: Detect student text typed in header cell (row 0)
                        # Students sometimes type in the blue header cell instead of the white response cell
                        response_cleaned = re.sub(r'[_\s]', '', response)
                        if len(response_cleaned) < 2:
                            header_cell_obj = rows[0].cells[0]
                            header_paras = [p.text.strip() for p in header_cell_obj.paragraphs]

                            # Case 1: Student pressed Enter in header cell — extra paragraphs
                            if len(header_paras) > 1:
                                extra_text = '\n'.join(p for p in header_paras[1:] if p)
                                if extra_text and len(re.sub(r'[_\s]', '', extra_text)) >= 2:
                                    response = extra_text
                                    visible_header = tag_pattern.sub('', header_paras[0]).strip()

                            # Case 2: Student typed on same line after (N pts) — no Enter
                            if len(re.sub(r'[_\s]', '', response)) < 2:
                                pts_match = re.search(r'\(\d+\s*pts?\)', header_cell_text)
                                if pts_match:
                                    after_pts = header_cell_text[pts_match.end():].strip()
                                    if after_pts and len(re.sub(r'[_\s]', '', after_pts)) >= 2:
                                        response = after_pts
                                        visible_header = tag_pattern.sub('', header_cell_text[:pts_match.end()]).strip()

                            # Case 3: SUMMARY (no pts marker) — check for non-bold student runs
                            if len(re.sub(r'[_\s]', '', response)) < 2 and tag_type == "SUMMARY":
                                student_text_parts = []
                                for para in header_cell_obj.paragraphs:
                                    for run in para.runs:
                                        # Header runs are bold or tiny (1pt tag) or 9pt (pts)
                                        is_tag_run = run.font.size is not None and run.font.size <= Pt(2)
                                        is_pts_run = run.font.size is not None and run.font.size == Pt(9)
                                        is_header_run = run.bold is True
                                        if not is_tag_run and not is_pts_run and not is_header_run:
                                            txt = run.text.strip()
                                            if txt:
                                                student_text_parts.append(txt)
                                if student_text_parts:
                                    candidate = ' '.join(student_text_parts).strip()
                                    if len(re.sub(r'[_\s]', '', candidate)) >= 2:
                                        response = candidate

                        # Deduplicate: skip if we already have this tag (duplicate tables in doc)
                        existing = next(
                            (t for t in tables_found
                             if t["tag_type"] == tag_type and t["tag_id"] == tag_id),
                            None
                        )
                        if existing:
                            # Keep the version with the longer response (more likely to be correct)
                            if len(response) > len(existing.get("response", "")):
                                existing["response"] = response
                                existing["header_text"] = visible_header
                        else:
                            tables_found.append({
                                "tag_type": tag_type,
                                "tag_id": tag_id,
                                "header_text": visible_header,
                                "response": response
                            })
                        # Also add to plain text for content reference
                        full_text.append(visible_header)
                        if response:
                            full_text.append(response)
                        continue

                # Non-Graider table — add as plain text
                for row in rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(' | '.join(row_text))

        is_graider = has_marker or len(tables_found) > 0
        plain_text = '\n'.join(full_text)

        if is_graider:
            _logger.info("Detected Graider table format: %d structured tables", len(tables_found))

        return {
            "is_graider_table": is_graider,
            "plain_text": plain_text,
            "tables": tables_found
        }
    except Exception as e:
        _logger.warning("Error in structured read: %s", e)
        return {"is_graider_table": False, "plain_text": None, "tables": []}
