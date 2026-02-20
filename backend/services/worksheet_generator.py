"""
Worksheet Generator Service
============================
Creates structured .docx worksheets for students.
Expected answers are stored in the assignment config's gradingNotes field,
not embedded in the document itself.
"""

import os
import re
import json
from urllib.parse import quote

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from backend.services.document_generator import (
    _parse_markdown_runs, load_style, DEFAULT_STYLE, _apply_style_to_heading
)
from backend.services.visualization import add_image_to_docx


WORKSHEETS_DIR = os.path.expanduser("~/Downloads/Graider/Worksheets")
ASSIGNMENTS_DIR = os.path.expanduser("~/.graider_assignments")


def _embed_visual(doc, visual):
    """Embed a visual element (math, graph, shape, number_line, coordinate_plane) into the doc.

    Args:
        doc: python-docx Document object
        visual: dict with 'type' and type-specific params. Supported types:
            - math: {type, latex, font_size?}
            - number_line: {type, min, max, points?, labels?, title?, blank?}
            - coordinate_plane: {type, x_range?, y_range?, points?, labels?, title?, blank?}
            - graph: {type, graph_type (bar/line/scatter), + data params}
            - box_plot: {type, data, labels?, title?, blank?}
            - shape: {type, shape_type (triangle/rectangle), + dimension params}
    """
    vtype = visual.get('type', '')
    try:
        if vtype == 'math':
            from backend.services.visualization import render_latex
            img_data = render_latex(
                visual.get('latex', ''),
                font_size=visual.get('font_size', 20)
            )
            add_image_to_docx(doc, img_data, width_inches=4)
        elif vtype == 'number_line':
            from backend.services.visualization import create_number_line
            img_data = create_number_line(
                min_val=visual.get('min', 0), max_val=visual.get('max', 10),
                points=visual.get('points', []), labels=visual.get('labels', {}),
                title=visual.get('title', ''), blank=visual.get('blank', False)
            )
            add_image_to_docx(doc, img_data, width_inches=5)
        elif vtype == 'coordinate_plane':
            from backend.services.visualization import create_coordinate_plane
            img_data = create_coordinate_plane(
                x_range=visual.get('x_range', [-10, 10]),
                y_range=visual.get('y_range', [-10, 10]),
                points=visual.get('points', []),
                labels=visual.get('labels', {}),
                title=visual.get('title', ''),
                blank=visual.get('blank', False)
            )
            add_image_to_docx(doc, img_data, width_inches=4.5)
        elif vtype == 'graph':
            gt = visual.get('graph_type', 'bar')
            if gt == 'bar':
                from backend.services.visualization import create_bar_chart
                img_data = create_bar_chart(
                    categories=visual.get('categories', []),
                    values=visual.get('values', []),
                    title=visual.get('title', ''),
                    x_label=visual.get('x_label', ''),
                    y_label=visual.get('y_label', ''),
                    blank=visual.get('blank', False)
                )
            elif gt == 'line':
                from backend.services.visualization import create_line_graph
                img_data = create_line_graph(
                    x_data=visual.get('x_data', []),
                    y_data=visual.get('y_data', []),
                    title=visual.get('title', ''),
                    x_label=visual.get('x_label', ''),
                    y_label=visual.get('y_label', ''),
                    blank=visual.get('blank', False)
                )
            elif gt == 'scatter':
                from backend.services.visualization import create_scatter_plot
                img_data = create_scatter_plot(
                    x_data=visual.get('x_data', []),
                    y_data=visual.get('y_data', []),
                    title=visual.get('title', ''),
                    x_label=visual.get('x_label', ''),
                    y_label=visual.get('y_label', ''),
                    show_trend=visual.get('show_trend', False),
                    blank=visual.get('blank', False)
                )
            else:
                return
            add_image_to_docx(doc, img_data, width_inches=4.5)
        elif vtype == 'box_plot':
            from backend.services.visualization import create_box_plot
            img_data = create_box_plot(
                data=visual.get('data', []),
                labels=visual.get('labels', []),
                title=visual.get('title', ''),
                blank=visual.get('blank', False)
            )
            add_image_to_docx(doc, img_data, width_inches=4.5)
        elif vtype == 'shape':
            st = visual.get('shape_type', 'triangle')
            if st == 'triangle':
                from backend.services.visualization import create_triangle
                img_data = create_triangle(
                    base=visual.get('base', 6),
                    height=visual.get('height', 4),
                    title=visual.get('title', ''),
                    blank=visual.get('blank', False)
                )
            elif st == 'rectangle':
                from backend.services.visualization import create_rectangle
                img_data = create_rectangle(
                    width=visual.get('width', 8),
                    height=visual.get('height', 5),
                    title=visual.get('title', ''),
                    blank=visual.get('blank', False)
                )
            else:
                return
            add_image_to_docx(doc, img_data, width_inches=3.5)
        elif vtype == 'function_graph':
            from backend.services.visualization import create_function_graph
            img_data = create_function_graph(
                expressions=visual.get('expressions', []),
                x_range=tuple(visual.get('x_range', (-10, 10))),
                y_range=tuple(visual['y_range']) if visual.get('y_range') else None,
                title=visual.get('title', ''),
                show_grid=visual.get('show_grid', True),
                labels=visual.get('labels'),
                blank=visual.get('blank', False)
            )
            add_image_to_docx(doc, img_data, width_inches=5)
        elif vtype == 'circle':
            from backend.services.visualization import create_circle
            img_data = create_circle(
                radius=visual.get('radius', 5),
                center=tuple(visual.get('center', (0, 0))),
                show_radius=visual.get('show_radius', True),
                show_diameter=visual.get('show_diameter', False),
                show_area=visual.get('show_area', False),
                title=visual.get('title', ''),
                blank=visual.get('blank', False)
            )
            add_image_to_docx(doc, img_data, width_inches=3.5)
        elif vtype == 'polygon':
            from backend.services.visualization import create_polygon
            img_data = create_polygon(
                sides=visual.get('sides', 5),
                side_length=visual.get('side_length', 4),
                show_labels=visual.get('show_labels', True),
                show_dimensions=visual.get('show_dimensions', True),
                title=visual.get('title', ''),
                blank=visual.get('blank', False)
            )
            add_image_to_docx(doc, img_data, width_inches=3.5)
        elif vtype == 'histogram':
            from backend.services.visualization import create_histogram
            img_data = create_histogram(
                data=visual.get('data', []),
                bins=visual.get('bins', 10),
                title=visual.get('title', ''),
                x_label=visual.get('x_label', ''),
                y_label=visual.get('y_label', 'Frequency'),
                show_values=visual.get('show_values', True),
                blank=visual.get('blank', False)
            )
            add_image_to_docx(doc, img_data, width_inches=4.5)
        elif vtype == 'pie_chart':
            from backend.services.visualization import create_pie_chart
            img_data = create_pie_chart(
                categories=visual.get('categories', []),
                values=visual.get('values', []),
                title=visual.get('title', ''),
                show_percentages=visual.get('show_percentages', True),
                explode=visual.get('explode'),
                blank=visual.get('blank', False)
            )
            add_image_to_docx(doc, img_data, width_inches=4)
    except Exception:
        doc.add_paragraph("[Visual failed to render]")


def _add_graider_table(doc, header_text, graider_tag, points, style, height_twips,
                       body_font="Calibri", body_size=11):
    """Add a structured 2-row Graider table for a question/vocab/summary section.

    Row 0 (header): Hidden [GRAIDER:TYPE:ID] tag + visible question text + points.
                    Blue shading, bold white text.
    Row 1 (response): Empty cell where student types. White, minimum height.

    Args:
        doc: python-docx Document object
        header_text: Visible question/term/prompt text
        graider_tag: e.g. "GRAIDER:VOCAB:Osmosis" (without brackets)
        points: Point value to display
        style: Style dict for colors/fonts
        height_twips: Minimum height of response cell in twips (1440=1in, 2160=1.5in)
        body_font: Font name for body text
        body_size: Font size for body text
    """
    header_bg = style.get("table_header_bg", "#4472C4").lstrip('#')
    header_text_color = style.get("table_header_text_color", "#FFFFFF").lstrip('#')

    table = doc.add_table(rows=2, cols=1)
    table.style = 'Table Grid'

    # --- Row 0: Header with hidden tag + visible text ---
    header_cell = table.rows[0].cells[0]
    header_cell.text = ""  # Clear default
    hp = header_cell.paragraphs[0]

    # Hidden tag run: font color matches background so it's invisible in print/screen
    tag_run = hp.add_run("[" + graider_tag + "]")
    tag_run.font.size = Pt(1)
    r, g, b = int(header_bg[0:2], 16), int(header_bg[2:4], 16), int(header_bg[4:6], 16)
    tag_run.font.color.rgb = RGBColor(r, g, b)

    # Visible header text
    text_run = hp.add_run("  " + header_text)
    text_run.bold = True
    text_run.font.size = Pt(body_size)
    text_run.font.name = body_font
    tr, tg, tb = int(header_text_color[0:2], 16), int(header_text_color[2:4], 16), int(header_text_color[4:6], 16)
    text_run.font.color.rgb = RGBColor(tr, tg, tb)

    # Points indicator
    if points is not None:
        pts_run = hp.add_run("  (" + str(points) + " pts)")
        pts_run.font.size = Pt(9)
        pts_run.font.name = body_font
        pts_run.font.color.rgb = RGBColor(tr, tg, tb)

    # Apply shading to header cell
    shading = parse_xml(
        '<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), header_bg)
    )
    header_cell._tc.get_or_add_tcPr().append(shading)

    # --- Row 1: Empty response cell with minimum height ---
    response_cell = table.rows[1].cells[0]
    response_cell.text = ""
    # Set minimum row height
    tr_elem = table.rows[1]._tr
    trPr = tr_elem.get_or_add_trPr()
    trHeight = parse_xml(
        '<w:trHeight {} w:val="{}" w:hRule="atLeast"/>'.format(
            nsdecls('w'), str(height_twips)
        )
    )
    trPr.append(trHeight)

    # Set font and add placeholder to guide students to type here
    rp = response_cell.paragraphs[0]
    placeholder_run = rp.add_run("Type your answer here...")
    placeholder_run.font.name = body_font
    placeholder_run.font.size = Pt(body_size)
    placeholder_run.font.color.rgb = RGBColor(180, 180, 180)
    placeholder_run.font.italic = True

    # Add spacing after table
    spacer = doc.add_paragraph()
    spacer_fmt = spacer.paragraph_format
    spacer_fmt.space_before = Pt(4)
    spacer_fmt.space_after = Pt(4)

    return table


def _add_graider_marker(doc):
    """Add hidden end-of-document marker for table format detection.

    Adds a paragraph with 'GRAIDER_TABLE_V1' in 1pt white font.
    """
    marker_para = doc.add_paragraph()
    marker_run = marker_para.add_run("GRAIDER_TABLE_V1")
    marker_run.font.size = Pt(1)
    marker_run.font.color.rgb = RGBColor(255, 255, 255)


def create_worksheet_docx(filepath, title, worksheet_type, vocab_terms,
                          questions, summary_prompt, summary_key_points,
                          total_points, style=None):
    """Create a .docx worksheet with visible student layer and invisible answer key.

    Args:
        filepath: Output path for the .docx file
        title: Worksheet title
        worksheet_type: One of 'cornell-notes', 'fill-in-blank', 'short-answer', 'vocabulary'
        vocab_terms: List of dicts with 'term' and 'definition' keys
        questions: List of dicts with 'question', 'expected_answer', 'points' keys
        summary_prompt: Instruction text for the summary section
        summary_key_points: List of strings - key points for a good summary
        total_points: Total point value of the worksheet
        style: Optional style dict (from load_style). Uses DEFAULT_STYLE if None.
    """
    if style is None:
        style = dict(DEFAULT_STYLE)

    body_font = style.get("body_font_name", "Calibri")
    body_size = style.get("body_font_size", 11)

    doc = Document()

    # === VISIBLE LAYER ===

    # Title
    heading = doc.add_heading(title, level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = style.get("title_font_name", "Georgia")
        run.font.size = Pt(style.get("title_font_size", 24))
        run.bold = style.get("title_bold", True)

    # Header fields (Name / Date / Period)
    for field in ['Name', 'Date', 'Period']:
        p = doc.add_paragraph()
        run = p.add_run(field + ': ')
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = body_font
        fill_run = p.add_run('_' * 50)
        fill_run.font.name = body_font

    doc.add_paragraph()  # spacing

    # Vocabulary Section — structured tables
    if vocab_terms:
        vh = doc.add_heading('VOCABULARY', level=2)
        _apply_style_to_heading(vh, 2, style)
        for item in vocab_terms:
            term = item.get('term', '')
            pts = item.get('points', 5)
            _add_graider_table(
                doc, term + ":", "GRAIDER:VOCAB:" + term,
                pts, style, 1440, body_font, body_size  # 1 inch response height
            )

    # Questions Section — structured tables (with optional visuals)
    if questions:
        qh = doc.add_heading('QUESTIONS', level=2)
        _apply_style_to_heading(qh, 2, style)
        for i, q in enumerate(questions, 1):
            pts = q.get('points', 10)
            question_text = q.get('question', '')
            # Embed visual above the answer table if provided
            visual = q.get('visual')
            if visual and isinstance(visual, dict):
                _embed_visual(doc, visual)
            header = str(i) + ") " + question_text
            _add_graider_table(
                doc, header, "GRAIDER:QUESTION:" + str(i),
                pts, style, 2160, body_font, body_size  # 1.5 inch response height
            )

    # Summary Section — structured table
    if summary_prompt:
        sh = doc.add_heading('SUMMARY', level=2)
        _apply_style_to_heading(sh, 2, style)
        _add_graider_table(
            doc, summary_prompt, "GRAIDER:SUMMARY:main",
            None, style, 4320, body_font, body_size  # 3 inch response height
        )

    # Hidden detection marker at end of document
    _add_graider_marker(doc)

    doc.save(filepath)


def _build_document_text(title, vocab_terms, questions, summary_prompt,
                         summary_key_points, total_points):
    """Build plain-text version of the worksheet matching what read_docx_file extracts.

    This is stored in importedDoc.text so Grading Setup can display the template
    and the grading engine can use it for template comparison.
    """
    lines = [title, '', 'Name: ' + '_' * 50, 'Date: ' + '_' * 50,
             'Period: ' + '_' * 50, '']

    if vocab_terms:
        lines.append('VOCABULARY')
        for item in vocab_terms:
            lines.append(item.get('term', '') + ': ' + '_' * 60)
        lines.append('')

    if questions:
        lines.append('QUESTIONS')
        for i, q in enumerate(questions, 1):
            pts = q.get('points', 10)
            lines.append(str(i) + ') ' + q.get('question', '') + '  (' + str(pts) + ' pts)')
            lines.append('Response: ' + '_' * 55)
            lines.append('_' * 65)
            lines.append('')

    if summary_prompt:
        lines.append('SUMMARY')
        lines.append(summary_prompt)
        for _ in range(5):
            lines.append('_' * 70)

    return '\n'.join(lines)


def generate_worksheet(title, worksheet_type, vocab_terms=None, questions=None,
                       summary_prompt=None, summary_key_points=None,
                       total_points=100, subject=None, style_name=None):
    """Generate a .docx worksheet and save its config to Grading Setup.

    Returns dict with status, filepath, download_url, etc.
    """
    os.makedirs(WORKSHEETS_DIR, exist_ok=True)

    # Load visual style
    style = load_style(style_name)

    # Match naming convention used by assignment_routes.py (spaces, not underscores)
    safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()
    filename = safe_title + '.docx'
    filepath = os.path.join(WORKSHEETS_DIR, filename)

    vterms = vocab_terms or []
    qs = questions or []
    skp = summary_key_points or []

    create_worksheet_docx(
        filepath=filepath,
        title=title,
        worksheet_type=worksheet_type,
        vocab_terms=vterms,
        questions=qs,
        summary_prompt=summary_prompt,
        summary_key_points=skp,
        total_points=total_points,
        style=style
    )

    # Build plain-text representation matching what read_docx_file extracts
    doc_text = _build_document_text(title, vterms, qs, summary_prompt, skp, total_points)

    # Save assignment config to Grading Setup
    download_url = "/api/download-worksheet/" + quote(filename)
    config = _build_assignment_config(
        title, worksheet_type, vterms, qs,
        summary_prompt, skp, total_points, subject
    )
    config["worksheetDownloadUrl"] = download_url
    config["tableStructured"] = True
    config["tableVersion"] = "v1"
    config["importedDoc"] = {
        "text": doc_text,
        "html": "",
        "filename": filename,
        "loading": False
    }
    os.makedirs(ASSIGNMENTS_DIR, exist_ok=True)
    config_path = os.path.join(ASSIGNMENTS_DIR, safe_title + '.json')
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config, f, indent=2)

    return {
        "status": "created",
        "filepath": filepath,
        "filename": filename,
        "download_url": "/api/download-worksheet/" + quote(filename),
        "saved_to_builder": True,
        "config_name": safe_title,
        "style_used": style_name or "default",
    }


def _build_assignment_config(title, worksheet_type, vocab_terms, questions,
                             summary_prompt, summary_key_points, total_points,
                             subject=None):
    """Build a Grading Setup-compatible assignment config from worksheet data.

    Generates both customMarkers (as objects with start/points/type) and
    questions[] (individual question objects) so the editor can render them.
    """
    import time

    rubric_map = {
        "cornell-notes": "cornell-notes",
        "fill-in-blank": "fill-in-blank",
        "short-answer": "standard",
        "vocabulary": "fill-in-blank"
    }

    effort_points = 15
    content_points = total_points - effort_points

    # Count total items to distribute points across
    num_items = len(vocab_terms or []) + len(questions or [])
    if summary_prompt:
        num_items += 1
    per_item = round(content_points / num_items) if num_items > 0 else 0

    # Build customMarkers as objects and questions[] for the editor
    marker_objects = []
    question_objects = []
    ts = int(time.time() * 1000)

    if vocab_terms:
        vocab_points = per_item * len(vocab_terms)
        marker_objects.append({
            "start": "VOCABULARY",
            "points": vocab_points,
            "type": "vocab_term"
        })
        for i, v in enumerate(vocab_terms):
            pts = round(vocab_points / len(vocab_terms))
            question_objects.append({
                "id": ts + i,
                "type": "vocab_term",
                "prompt": v.get('term', ''),
                "points": pts,
                "marker": "VOCABULARY",
                "expected_answer": v.get('definition', ''),
            })

    if questions:
        q_points = per_item * len(questions)
        marker_objects.append({
            "start": "QUESTIONS",
            "points": q_points,
            "type": "short_answer"
        })
        for i, q in enumerate(questions):
            pts = round(q_points / len(questions))
            question_objects.append({
                "id": ts + 1000 + i,
                "type": "short_answer",
                "prompt": q.get('question', ''),
                "points": pts,
                "marker": "QUESTIONS",
                "expected_answer": q.get('expected_answer', ''),
            })

    if summary_prompt:
        s_points = per_item
        marker_objects.append({
            "start": "SUMMARY",
            "points": s_points,
            "type": "written"
        })
        question_objects.append({
            "id": ts + 2000,
            "type": "written",
            "prompt": summary_prompt,
            "points": s_points,
            "marker": "SUMMARY",
        })

    # Build grading notes with answer key
    grading_notes_parts = []
    if vocab_terms:
        grading_notes_parts.append("VOCABULARY EXPECTED DEFINITIONS:")
        for v in vocab_terms:
            grading_notes_parts.append("- " + v.get('term', '') + ": " +
                                       v.get('definition', 'Accept reasonable definition'))
    if questions:
        grading_notes_parts.append("")
        grading_notes_parts.append("EXPECTED ANSWERS:")
        for i, q in enumerate(questions, 1):
            grading_notes_parts.append("- Q" + str(i) + ": " +
                                       q.get('expected_answer', ''))
    if summary_key_points:
        grading_notes_parts.append("")
        grading_notes_parts.append("SUMMARY SHOULD INCLUDE:")
        for kp in summary_key_points:
            grading_notes_parts.append("- " + kp)

    return {
        "title": title,
        "subject": subject or "",
        "totalPoints": total_points,
        "instructions": "",
        "aliases": [],
        "customMarkers": marker_objects,
        "excludeMarkers": [],
        "gradingNotes": "\n".join(grading_notes_parts),
        "questions": question_objects,
        "responseSections": [],
        "rubricType": rubric_map.get(worksheet_type, "standard"),
        "customRubric": None,
        "useSectionPoints": True,
        "sectionTemplate": "Custom",
        "effortPoints": effort_points,
        "completionOnly": False,
        "countsTowardsGrade": True,
        "importedDoc": None,  # Populated by generate_worksheet() after file creation
        "worksheetDownloadUrl": None
    }
