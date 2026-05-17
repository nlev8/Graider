"""Document / visual / platform-export rendering for the planner. Pure logic extracted from planner_routes.py (no Flask)."""
import os
import json
import logging

_logger = logging.getLogger(__name__)


def _question_to_visual_dict(q):
    """Convert a planner question dict to the visual dict format expected by _embed_visual().

    Maps planner question fields (visual_type, question_type, etc.) to the
    {type, ...params} format used by worksheet_generator._embed_visual().
    Returns None if the question has no visual element.
    """
    q_type = q.get('question_type', q.get('visual_type', ''))
    if not q_type:
        return None

    # Direct type mappings — these map to _embed_visual 'type' values
    type_map = {
        'number_line': 'number_line',
        'coordinate_plane': 'coordinate_plane',
        'box_plot': 'box_plot',
        'function_graph': 'function_graph',
        'dot_plot': 'dot_plot',
        'stem_and_leaf': 'stem_and_leaf',
        'venn_diagram': 'venn_diagram',
        'histogram': 'histogram',
        'pie_chart': 'pie_chart',
        'bar_chart': 'graph',
        'protractor': 'protractor',
        'angle_protractor': 'protractor',
    }

    # Shape types map to 'shape' with shape_type sub-key
    shape_types = {'triangle', 'rectangle', 'geometry', 'pythagorean',
                   'trig', 'angles', 'similarity'}

    # Circle maps to 'circle'
    circle_types = {'circle'}

    # Polygon maps to 'polygon'
    polygon_types = {'regular_polygon'}

    if q_type in type_map:
        visual = dict(q)  # copy all params
        visual['type'] = type_map[q_type]
        if q_type == 'bar_chart':
            visual['graph_type'] = 'bar'
        visual['blank'] = True  # Student version — don't show answers
        return visual
    elif q_type in shape_types:
        visual = dict(q)
        visual['type'] = 'shape'
        visual['shape_type'] = 'triangle' if q_type in ('triangle', 'geometry', 'pythagorean', 'trig', 'angles', 'similarity') else 'rectangle'
        visual['blank'] = True
        return visual
    elif q_type in circle_types:
        visual = dict(q)
        visual['type'] = 'circle'
        visual['blank'] = True
        return visual
    elif q_type in polygon_types:
        visual = dict(q)
        visual['type'] = 'polygon'
        visual['blank'] = True
        return visual

    return None


def _export_assignment_docx_graider(assignment, output_folder, safe_title):
    """Export a generated assignment as a .docx with Graider table extraction tags.

    Creates a Word document with the same structure as the PDF export but using
    Graider tables for structured student response extraction.

    Returns the file path of the saved .docx.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from backend.services.worksheet_generator import _add_graider_table, _add_graider_marker, _embed_visual
    from backend.services.worksheet_generator import _add_options_with_bubbles, _create_answer_key_doc

    doc = Document()

    graider_style = {
        "table_header_bg": "#4472C4",
        "table_header_text_color": "#FFFFFF",
    }

    title = assignment.get('title', 'Assignment')
    instructions = assignment.get('instructions', '')
    sections = assignment.get('sections', [])
    total_points = assignment.get('total_points', 100)

    # Teacher name and subject header
    _teacher = assignment.get('teacher_name', '')
    _subject = assignment.get('subject', '')
    if _teacher or _subject:
        header_para = doc.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_parts = []
        if _teacher:
            header_parts.append(_teacher)
        if _subject:
            header_parts.append(_subject)
        header_run = header_para.add_run("  |  ".join(header_parts))
        header_run.font.size = Pt(11)
        header_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Title — 24pt, centered
    heading = doc.add_heading(title, 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(24)
        run.font.color.rgb = RGBColor(0, 0, 0)

    doc.add_paragraph()  # Skip line after title

    # Student info line
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run("Name: _______________________  Date: _______________  Period: _____")

    doc.add_paragraph()  # Skip line after name/date

    # Total points (no time limit on assignments)
    if total_points:
        meta_para = doc.add_paragraph()
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_para.add_run(f"Total Points: {total_points}")

    # Instructions
    if instructions:
        doc.add_paragraph()
        inst_para = doc.add_paragraph()
        inst_para.add_run("Instructions: ").bold = True
        inst_para.add_run(instructions)

    doc.add_paragraph()  # Space before questions

    question_num = 1
    answer_key_questions = []

    for section in sections:
        section_name = section.get('name', 'Section')
        section_points = section.get('points', 0)
        section_type = section.get('type', 'short_answer')
        questions = section.get('questions', [])

        # Section header
        pts_text = f" ({section_points} points)" if section_points else ""
        doc.add_heading(f"{section_name}{pts_text}", level=1)
        doc.add_paragraph()  # Space between section header and questions

        for q in questions:
            q_number = q.get('number', question_num)
            q_text = q.get('question', '')
            q_points = q.get('points', 0)
            q_options = q.get('options', [])
            q_type = q.get('question_type', section_type)
            q_visual = q.get('visual_type', None)

            # Embed visual element if present
            visual_dict = _question_to_visual_dict(q)
            if visual_dict:
                try:
                    # Render question text above the visual — black font
                    q_para = doc.add_paragraph()
                    _nr = q_para.add_run(f"{q_number}. ")
                    _nr.bold = True
                    _nr.font.color.rgb = RGBColor(0, 0, 0)
                    _tr = q_para.add_run(q_text)
                    _tr.font.color.rgb = RGBColor(0, 0, 0)
                    if q_points:
                        _pr = q_para.add_run(f" ({q_points} pts)")
                        _pr.italic = True
                        _pr.font.color.rgb = RGBColor(0, 0, 0)

                    _embed_visual(doc, visual_dict)
                except Exception as e:
                    _logger.warning("Could not embed visual for Q%s: %s", q_number, e)

            # Inject True/False options if missing for TF questions
            if not q_options and q_type in ('true_false', 'tf'):
                q_options = ['True', 'False']

            if q_options:
                # MC/TF: question + options with bubbles (no separate answer table)
                is_tf = q_type in ('true_false', 'tf')
                # Filter out metadata fields that AI sometimes includes as options
                q_options = [opt for opt in q_options if not str(opt).startswith('answer_key_field')]
                if not visual_dict:
                    q_para = doc.add_paragraph()
                    _nr = q_para.add_run(f"{q_number}. ")
                    _nr.bold = True
                    _nr.font.color.rgb = RGBColor(0, 0, 0)
                    _tr = q_para.add_run(q_text)
                    _tr.font.color.rgb = RGBColor(0, 0, 0)
                    if q_points:
                        _pr = q_para.add_run(f" ({q_points} pts)")
                        _pr.italic = True
                        _pr.font.color.rgb = RGBColor(0, 0, 0)

                # Options with empty bubbles — bubbles ARE the answer
                _add_options_with_bubbles(doc, q_options, is_tf=is_tf)
                doc.add_paragraph()  # Space after MC/TF question

                # Track for separate answer key file
                answer_key_questions.append({
                    "number": q_number,
                    "option_texts": q_options,
                    "correct_answer": q.get('answer', ''),
                    "is_tf": is_tf,
                    "question_text": q_text,
                })

            elif q_type == 'data_table':
                # Data table: render table headers above, then Graider table for answers
                if not visual_dict:
                    q_para = doc.add_paragraph()
                    _nr = q_para.add_run(f"{q_number}. ")
                    _nr.bold = True
                    _nr.font.color.rgb = RGBColor(0, 0, 0)
                    _tr = q_para.add_run(q_text)
                    _tr.font.color.rgb = RGBColor(0, 0, 0)
                    if q_points:
                        _pr = q_para.add_run(f" ({q_points} pts)")
                        _pr.italic = True
                        _pr.font.color.rgb = RGBColor(0, 0, 0)

                # Render visible data table if headers present
                headers = q.get('headers', q.get('column_headers', []))
                row_labels = q.get('row_labels', [])
                expected = q.get('expected_data', [])
                num_rows = q.get('num_rows', len(expected) if expected else 5)
                if headers:
                    from docx.shared import Inches as DocxInches
                    if row_labels:
                        table_data_headers = [''] + headers
                    else:
                        table_data_headers = headers
                    col_count = len(table_data_headers)
                    tbl = doc.add_table(rows=1 + num_rows, cols=col_count)
                    tbl.style = 'Table Grid'
                    for ci, h in enumerate(table_data_headers):
                        tbl.rows[0].cells[ci].text = str(h)
                    for ri in range(num_rows):
                        if row_labels and ri < len(row_labels):
                            tbl.rows[ri + 1].cells[0].text = str(row_labels[ri])
                    doc.add_paragraph()  # space after table

                _add_graider_table(doc, f"Data Analysis for Question {q_number}",
                                   f"GRAIDER:QUESTION:{q_number}", q_points,
                                   graider_style, 2160)  # 1.5 inch
                doc.add_paragraph()  # Space after question

            elif q_type == 'math_equation':
                _add_graider_table(doc, f"{q_number}. {q_text} ({q_points} pts) — Show your work",
                                   f"GRAIDER:QUESTION:{q_number}", q_points,
                                   graider_style, 3600)  # 2.5 inches
                doc.add_paragraph()  # Space after question

            elif q.get('terms') and q.get('definitions'):
                # Matching: two-column layout with draw-a-line instruction
                if not visual_dict:
                    q_para = doc.add_paragraph()
                    _nr = q_para.add_run(f"{q_number}. ")
                    _nr.bold = True
                    _nr.font.color.rgb = RGBColor(0, 0, 0)
                    _tr = q_para.add_run(q_text)
                    _tr.font.color.rgb = RGBColor(0, 0, 0)
                    if q_points:
                        _pr = q_para.add_run(f" ({q_points} pts)")
                        _pr.italic = True
                        _pr.font.color.rgb = RGBColor(0, 0, 0)

                inst_para = doc.add_paragraph()
                inst_para.add_run("Directions: ").bold = True
                inst_para.add_run("Draw a line from each term to its matching definition.")
                inst_para.italic = True

                terms = q.get('terms', [])
                definitions = q.get('definitions', [])
                max_len = max(len(terms), len(definitions))

                tbl = doc.add_table(rows=max_len + 1, cols=3)
                tbl.style = 'Table Grid'

                # Header row
                tbl.rows[0].cells[0].text = "Term"
                tbl.rows[0].cells[1].text = ""
                tbl.rows[0].cells[2].text = "Definition"
                for cell in tbl.rows[0].cells:
                    for paragraph in cell.paragraphs:
                        for run_obj in paragraph.runs:
                            run_obj.bold = True

                # Set narrow middle column
                for row in tbl.rows:
                    row.cells[1].width = Inches(0.5)

                for i in range(max_len):
                    if i < len(terms):
                        tbl.rows[i + 1].cells[0].text = f"{i + 1}. {terms[i]}"
                    if i < len(definitions):
                        letter_char = chr(65 + i)
                        tbl.rows[i + 1].cells[2].text = f"{letter_char}. {definitions[i]}"

                doc.add_paragraph()  # Space after table

            elif q_type in ('essay', 'extended_response'):
                # Render question text as black paragraph ABOVE the response box
                if not visual_dict:
                    q_para = doc.add_paragraph()
                    _nr = q_para.add_run(f"{q_number}. ")
                    _nr.bold = True
                    _nr.font.color.rgb = RGBColor(0, 0, 0)
                    _tr = q_para.add_run(q_text)
                    _tr.font.color.rgb = RGBColor(0, 0, 0)
                    if q_points:
                        _pr = q_para.add_run(f" ({q_points} pts)")
                        _pr.italic = True
                        _pr.font.color.rgb = RGBColor(0, 0, 0)
                _add_graider_table(doc, f"Response for Question {q_number}",
                                   f"GRAIDER:QUESTION:{q_number}", q_points,
                                   graider_style, 4320)  # 3 inches
                doc.add_paragraph()  # Space after question

            elif q_type == 'coordinates':
                if not visual_dict:
                    q_para = doc.add_paragraph()
                    _nr = q_para.add_run(f"{q_number}. ")
                    _nr.bold = True
                    _nr.font.color.rgb = RGBColor(0, 0, 0)
                    _tr = q_para.add_run(q_text)
                    _tr.font.color.rgb = RGBColor(0, 0, 0)
                    if q_points:
                        _pr = q_para.add_run(f" ({q_points} pts)")
                        _pr.italic = True
                        _pr.font.color.rgb = RGBColor(0, 0, 0)

                _add_graider_table(doc, f"Coordinates for Question {q_number}",
                                   f"GRAIDER:QUESTION:{q_number}", q_points,
                                   graider_style, 720)  # 0.5 inch
                doc.add_paragraph()  # Space after question

            else:
                # Short answer / default — render question text as black paragraph
                height = 2160  # 1.5 inches
                if not visual_dict:
                    q_para = doc.add_paragraph()
                    _nr = q_para.add_run(f"{q_number}. ")
                    _nr.bold = True
                    _nr.font.color.rgb = RGBColor(0, 0, 0)
                    _tr = q_para.add_run(q_text)
                    _tr.font.color.rgb = RGBColor(0, 0, 0)
                    if q_points:
                        _pr = q_para.add_run(f" ({q_points} pts)")
                        _pr.italic = True
                        _pr.font.color.rgb = RGBColor(0, 0, 0)
                _add_graider_table(doc, f"Response for Question {q_number}",
                                   f"GRAIDER:QUESTION:{q_number}", q_points,
                                   graider_style, height)
                doc.add_paragraph()  # Space after question

            question_num += 1

    # Add Graider marker at end
    _add_graider_marker(doc)

    # Save student worksheet
    filename = f"{safe_title}_Student.docx"
    filepath = os.path.join(output_folder, filename)
    doc.save(filepath)

    # Save separate answer key file
    if answer_key_questions:
        answer_key_doc = _create_answer_key_doc(title, answer_key_questions)
        if answer_key_doc:
            key_filepath = os.path.join(output_folder, f"{safe_title}_Answer_Key.docx")
            answer_key_doc.save(key_filepath)

    return filepath


def _create_visual_for_question(question: dict, show_answer: bool = False):
    """Create a visual element (graph, number line, etc.) for a question.

    Returns a ReportLab Image with both width and height set to preserve aspect ratio.
    Supports all geometry types: triangle, rectangle, circle, trapezoid, parallelogram,
    regular_polygon, rectangular_prism, cylinder.
    """
    try:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        import numpy as np
        import math as _math
        from reportlab.lib.units import inch
        from reportlab.platypus import Image
        import io

        q_type = question.get('question_type', question.get('visual_type', ''))
        fig = None

        if q_type == 'number_line':
            fig, ax = plt.subplots(figsize=(7, 1.5))
            min_val = question.get('min_val', -10)
            max_val = question.get('max_val', 10)

            ax.axhline(y=0, color='black', linewidth=2)
            ax.set_xlim(min_val - 0.5, max_val + 0.5)
            ax.set_ylim(-0.5, 0.5)

            for i in range(int(min_val), int(max_val) + 1):
                ax.plot([i, i], [-0.1, 0.1], 'k-', linewidth=1.5)
                ax.text(i, -0.25, str(i), ha='center', fontsize=10)

            ax.annotate('', xy=(max_val + 0.3, 0), xytext=(max_val, 0),
                       arrowprops=dict(arrowstyle='->', color='black', lw=2))
            ax.annotate('', xy=(min_val - 0.3, 0), xytext=(min_val, 0),
                       arrowprops=dict(arrowstyle='->', color='black', lw=2))

            if show_answer and question.get('points_to_plot'):
                for pt in question['points_to_plot']:
                    ax.plot(pt, 0, 'ro', markersize=10)
            ax.axis('off')

        elif q_type == 'coordinate_plane':
            fig, ax = plt.subplots(figsize=(4.5, 4.5))
            x_range = question.get('x_range', (-6, 6))
            y_range = question.get('y_range', (-6, 6))

            ax.axhline(y=0, color='black', linewidth=1.5)
            ax.axvline(x=0, color='black', linewidth=1.5)
            ax.grid(True, linestyle='--', alpha=0.5)
            ax.set_xlim(x_range[0] - 0.5, x_range[1] + 0.5)
            ax.set_ylim(y_range[0] - 0.5, y_range[1] + 0.5)
            ax.set_xticks(range(x_range[0], x_range[1] + 1))
            ax.set_yticks(range(y_range[0], y_range[1] + 1))

            offset = (x_range[1] - x_range[0]) * 0.35
            ax.text(offset, offset, 'I', fontsize=14, color='gray', alpha=0.5)
            ax.text(-offset, offset, 'II', fontsize=14, color='gray', alpha=0.5)
            ax.text(-offset, -offset, 'III', fontsize=14, color='gray', alpha=0.5)
            ax.text(offset, -offset, 'IV', fontsize=14, color='gray', alpha=0.5)

            ax.set_xlabel('x')
            ax.set_ylabel('y')
            ax.set_aspect('equal')

            if show_answer and question.get('points_to_plot'):
                labels = question.get('point_labels', [])
                for i, pt in enumerate(question['points_to_plot']):
                    ax.plot(pt[0], pt[1], 'ro', markersize=10)
                    label = labels[i] if i < len(labels) else f"({pt[0]}, {pt[1]})"
                    ax.annotate(label, xy=pt, xytext=(5, 5), textcoords='offset points', fontsize=10)

        elif q_type in ('geometry', 'triangle', 'pythagorean', 'trig', 'angles', 'similarity'):
            fig, ax = plt.subplots(figsize=(4, 3))
            base = question.get('base', 6)
            height = question.get('height', 4)

            vertices = [(0, 0), (base, 0), (base/2, height)]
            triangle = plt.Polygon(vertices, fill=True, facecolor='lightblue',
                                  edgecolor='black', linewidth=2)
            ax.add_patch(triangle)
            ax.plot([base/2, base/2], [0, height], 'r--', linewidth=1.5)
            ax.text(base/2, -0.4, f'b = {base}', ha='center', fontsize=11)
            ax.text(base/2 + 0.3, height/2, f'h = {height}', ha='left', fontsize=11)
            ax.set_xlim(-1, base + 1)
            ax.set_ylim(-1, height + 1)
            ax.set_aspect('equal')
            ax.axis('off')

        elif q_type == 'rectangle':
            fig, ax = plt.subplots(figsize=(4, 3))
            base = question.get('base', question.get('width', 6))
            height = question.get('height', 4)

            rect = plt.Rectangle((0, 0), base, height, fill=True,
                                facecolor='#bbf7d0', edgecolor='#22c55e', linewidth=2)
            ax.add_patch(rect)
            ax.text(base / 2, -0.4, f'w = {base}', ha='center', fontsize=11)
            ax.text(base + 0.3, height / 2, f'h = {height}', ha='left', fontsize=11)
            ax.set_xlim(-1, base + 1.5)
            ax.set_ylim(-1, height + 1)
            ax.set_aspect('equal')
            ax.axis('off')

        elif q_type == 'circle':
            fig, ax = plt.subplots(figsize=(3.5, 3.5))
            radius = question.get('radius', 5)
            circle_patch = plt.Circle((0, 0), radius, fill=True,
                                      facecolor='#dbeafe', edgecolor='#2563eb', linewidth=2)
            ax.add_patch(circle_patch)
            ax.plot(0, 0, 'ko', markersize=4)
            ax.plot([0, radius], [0, 0], 'r-', linewidth=1.5)
            ax.text(radius / 2, radius * 0.08, f'r = {radius}', ha='center', fontsize=11, color='red')
            margin = radius * 0.3
            ax.set_xlim(-radius - margin, radius + margin)
            ax.set_ylim(-radius - margin, radius + margin)
            ax.set_aspect('equal')
            ax.axis('off')

        elif q_type == 'trapezoid':
            fig, ax = plt.subplots(figsize=(4, 3))
            top_base = question.get('top_base', 4)
            base = question.get('base', 8)
            height = question.get('height', 5)
            offset_x = (base - top_base) / 2
            vertices = [(offset_x, height), (offset_x + top_base, height), (base, 0), (0, 0)]
            trap = plt.Polygon(vertices, fill=True, facecolor='#e9d5ff', edgecolor='#a855f7', linewidth=2)
            ax.add_patch(trap)
            ax.text((offset_x + offset_x + top_base) / 2, height + 0.3, f'a = {top_base}', ha='center', fontsize=11, color='#a855f7')
            ax.text(base / 2, -0.4, f'b = {base}', ha='center', fontsize=11, color='#a855f7')
            ax.plot([base / 2, base / 2], [0, height], 'r--', linewidth=1.5)
            ax.text(base / 2 + 0.3, height / 2, f'h = {height}', ha='left', fontsize=11, color='red')
            ax.set_xlim(-1, base + 1)
            ax.set_ylim(-1, height + 1.5)
            ax.set_aspect('equal')
            ax.axis('off')

        elif q_type == 'parallelogram':
            fig, ax = plt.subplots(figsize=(4, 3))
            base = question.get('base', 7)
            height = question.get('height', 4)
            slant = 1.5
            vertices = [(slant, height), (base + slant, height), (base, 0), (0, 0)]
            para = plt.Polygon(vertices, fill=True, facecolor='#fbcfe8', edgecolor='#ec4899', linewidth=2)
            ax.add_patch(para)
            ax.text(base / 2, -0.4, f'b = {base}', ha='center', fontsize=11, color='#ec4899')
            ax.plot([slant, slant], [0, height], 'r--', linewidth=1.5)
            ax.text(slant + 0.3, height / 2, f'h = {height}', ha='left', fontsize=11, color='red')
            ax.set_xlim(-1, base + slant + 1)
            ax.set_ylim(-1, height + 1)
            ax.set_aspect('equal')
            ax.axis('off')

        elif q_type == 'regular_polygon':
            fig, ax = plt.subplots(figsize=(3.5, 3.5))
            n = max(3, min(12, question.get('sides', 6)))
            side_length = question.get('side_length', 4)
            R = 2.5
            vertices = []
            for i in range(n):
                angle = (2 * _math.pi * i) / n - _math.pi / 2
                vertices.append((R * _math.cos(angle), R * _math.sin(angle)))
            vertices.append(vertices[0])
            xs, ys = zip(*vertices)
            ax.fill(xs, ys, facecolor='#dbeafe', edgecolor='#3b82f6', linewidth=2)
            mid_x = (vertices[0][0] + vertices[1][0]) / 2
            mid_y = (vertices[0][1] + vertices[1][1]) / 2
            ax.plot([0, mid_x], [0, mid_y], 'r--', linewidth=1.5)
            ax.text(mid_x, mid_y - 0.4, f's = {side_length}', ha='center', fontsize=11, color='#3b82f6')
            ax.text(0, 0.2, f'n = {n}', ha='center', fontsize=10, color='gray')
            ax.set_xlim(-3.5, 3.5)
            ax.set_ylim(-3.5, 3.5)
            ax.set_aspect('equal')
            ax.axis('off')

        elif q_type == 'rectangular_prism':
            fig, ax = plt.subplots(figsize=(4, 3.5))
            l = question.get('base', 5)
            w = question.get('width', 3)
            h = question.get('height', 4)
            # Front face
            ax.add_patch(plt.Rectangle((0, 0), l, h, fill=True, facecolor='#bfdbfe', edgecolor='#3b82f6', linewidth=2))
            # Top face
            top_verts = [(0, h), (w * 0.5, h + w * 0.4), (l + w * 0.5, h + w * 0.4), (l, h)]
            ax.add_patch(plt.Polygon(top_verts, fill=True, facecolor='#93c5fd', edgecolor='#3b82f6', linewidth=2))
            # Right face
            right_verts = [(l, 0), (l + w * 0.5, w * 0.4), (l + w * 0.5, h + w * 0.4), (l, h)]
            ax.add_patch(plt.Polygon(right_verts, fill=True, facecolor='#60a5fa', edgecolor='#3b82f6', linewidth=2))
            ax.text(l / 2, -0.5, f'l = {l}', ha='center', fontsize=10)
            ax.text(-0.5, h / 2, f'h = {h}', ha='right', fontsize=10)
            ax.text(l + w * 0.3, h + w * 0.25, f'w = {w}', fontsize=10, color='#1d4ed8')
            ax.set_xlim(-1, l + w * 0.5 + 1)
            ax.set_ylim(-1, h + w * 0.4 + 1)
            ax.set_aspect('equal')
            ax.axis('off')

        elif q_type == 'cylinder':
            fig, ax = plt.subplots(figsize=(3.5, 4))
            radius = question.get('radius', 3)
            h = question.get('height', 7)
            from matplotlib.patches import Ellipse
            # Body
            ax.add_patch(plt.Rectangle((-radius, 0), 2 * radius, h, facecolor='#bbf7d0', edgecolor='none'))
            ax.plot([-radius, -radius], [0, h], color='#22c55e', linewidth=2)
            ax.plot([radius, radius], [0, h], color='#22c55e', linewidth=2)
            # Bottom ellipse
            ax.add_patch(Ellipse((0, 0), 2 * radius, radius * 0.6, facecolor='#86efac', edgecolor='#22c55e', linewidth=2))
            # Top ellipse
            ax.add_patch(Ellipse((0, h), 2 * radius, radius * 0.6, facecolor='#bbf7d0', edgecolor='#22c55e', linewidth=2))
            # Radius line
            ax.plot([0, radius], [h, h], 'r-', linewidth=2)
            ax.plot(0, h, 'ro', markersize=4)
            ax.text(radius / 2, h + radius * 0.2, f'r = {radius}', ha='center', fontsize=10, color='red')
            ax.text(radius + 0.3, h / 2, f'h = {h}', ha='left', fontsize=10, color='#1d4ed8')
            margin = radius * 0.5
            ax.set_xlim(-radius - margin, radius + margin + 1)
            ax.set_ylim(-radius * 0.3 - 0.5, h + radius * 0.3 + 0.5)
            ax.set_aspect('equal')
            ax.axis('off')

        elif q_type == 'box_plot':
            fig, ax = plt.subplots(figsize=(6, 2.5))
            data = question.get('data', [[50, 60, 70, 75, 80, 85, 90]])
            labels = question.get('data_labels', [f'Set {i+1}' for i in range(len(data))])
            bp = ax.boxplot(data, patch_artist=True, tick_labels=labels)
            colors = plt.cm.Pastel1(np.linspace(0, 1, len(data)))
            for patch, color in zip(bp['boxes'], colors):
                patch.set_facecolor(color)
            ax.set_ylabel('Value')
            ax.grid(True, axis='y', linestyle='--', alpha=0.5)

        elif q_type == 'bar_chart':
            fig, ax = plt.subplots(figsize=(5, 3))
            categories = question.get('categories', ['A', 'B', 'C', 'D'])
            chart_data = question.get('chart_data', {})
            values = chart_data.get('values', question.get('values', [0, 0, 0, 0]))
            if not show_answer:
                values = [0] * len(categories)
            ax.bar(categories, values, color='steelblue', edgecolor='black')
            ax.set_ylabel(chart_data.get('y_label', question.get('y_label', 'Value')))
            ax.grid(True, axis='y', linestyle='--', alpha=0.5)

        elif q_type == 'function_graph':
            fig, ax = plt.subplots(figsize=(4.5, 3.5))
            x_range = question.get('x_range', (-10, 10))
            y_range = question.get('y_range', (-10, 10))
            ax.axhline(y=0, color='black', linewidth=0.8)
            ax.axvline(x=0, color='black', linewidth=0.8)
            ax.grid(True, linestyle='--', alpha=0.4)
            ax.set_xlim(x_range)
            ax.set_ylim(y_range)
            ax.set_xlabel('x')
            ax.set_ylabel('y')

            if show_answer and question.get('correct_expressions'):
                from sympy import sympify, lambdify, Symbol
                x_sym = Symbol('x')
                x_vals = np.linspace(x_range[0], x_range[1], 300)
                colors = ['#2563eb', '#dc2626', '#16a34a', '#9333ea']
                for idx, expr_str in enumerate(question['correct_expressions']):
                    try:
                        clean = expr_str.strip().replace('^', '**')
                        if clean.lower().startswith('y'):
                            clean = clean.split('=', 1)[-1].strip()
                        sym_expr = sympify(clean)
                        f = lambdify(x_sym, sym_expr, modules=['numpy'])
                        y_vals = f(x_vals)
                        if not hasattr(y_vals, '__len__'):
                            y_vals = np.full_like(x_vals, float(y_vals))
                        ax.plot(x_vals, y_vals, color=colors[idx % len(colors)], linewidth=2)
                    except Exception:
                        continue

        elif q_type == 'dot_plot':
            categories = question.get('categories', [])
            correct_dots = question.get('correct_dots', {})
            min_val = question.get('min_val', 0)
            max_val = question.get('max_val', 10)
            step = question.get('step', 1)
            items = categories if categories else [str(min_val + i * step) for i in range(int((max_val - min_val) / step) + 1)]
            fig, ax = plt.subplots(figsize=(6, 2.5))
            ax.axhline(y=0, color='black', linewidth=1.5)
            for idx, item in enumerate(items):
                ax.text(idx, -0.3, str(item), ha='center', fontsize=9)
                ax.plot([idx, idx], [-0.05, 0.05], 'k-', linewidth=1)
                if show_answer and correct_dots:
                    count = int(correct_dots.get(str(item), 0))
                    for d in range(count):
                        ax.plot(idx, 0.3 + d * 0.35, 'o', color='#6366f1', markersize=8)
            ax.set_xlim(-0.5, len(items) - 0.5)
            max_count = max(int(v) for v in correct_dots.values()) if correct_dots else 3
            ax.set_ylim(-0.7, 0.3 + max_count * 0.35 + 0.5)
            ax.axis('off')

        elif q_type == 'stem_and_leaf':
            data = question.get('data', [])
            correct_leaves = question.get('correct_leaves', {})
            fig, ax = plt.subplots(figsize=(4, 3))
            if data:
                stems = sorted(set(v // 10 for v in data))
                rows = []
                for s in stems:
                    leaves = sorted(v % 10 for v in data if v // 10 == s)
                    rows.append(f"  {s} | {' '.join(str(l) for l in leaves)}")
                text = "Stem | Leaf\n" + "-" * 20 + "\n" + "\n".join(rows)
                if show_answer:
                    ax.text(0.1, 0.5, text, transform=ax.transAxes, fontsize=12, fontfamily='monospace', va='center')
                else:
                    blank_rows = [f"  {s} |" for s in stems]
                    blank_text = "Stem | Leaf\n" + "-" * 20 + "\n" + "\n".join(blank_rows)
                    ax.text(0.1, 0.5, blank_text, transform=ax.transAxes, fontsize=12, fontfamily='monospace', va='center')
            ax.axis('off')

        elif q_type == 'unit_circle':
            fig, ax = plt.subplots(figsize=(5, 5))
            circle = plt.Circle((0, 0), 1, fill=False, color='#6366f1', linewidth=2)
            ax.add_patch(circle)
            ax.axhline(y=0, color='black', linewidth=1)
            ax.axvline(x=0, color='black', linewidth=1)
            key_degs = [0, 30, 45, 60, 90, 120, 135, 150, 180, 210, 225, 240, 270, 300, 315, 330]
            for d in key_degs:
                rad = _math.radians(d)
                px, py = _math.cos(rad), _math.sin(rad)
                ax.plot(px, py, 'o', color='#6366f1', markersize=4)
                ax.plot([0, px], [0, py], '--', color='#6366f1', alpha=0.2, linewidth=0.5)
                if show_answer:
                    ax.annotate(f'{d}\u00b0', xy=(px, py), xytext=(px * 1.15, py * 1.15), fontsize=7, ha='center')
            ax.set_xlim(-1.4, 1.4)
            ax.set_ylim(-1.4, 1.4)
            ax.set_aspect('equal')
            ax.grid(True, linestyle='--', alpha=0.2)
            ax.set_title('Unit Circle', fontsize=12)

        elif q_type == 'transformations':
            fig, ax = plt.subplots(figsize=(5, 5))
            grid_range = question.get('grid_range', [-8, 8])
            ax.axhline(y=0, color='black', linewidth=1)
            ax.axvline(x=0, color='black', linewidth=1)
            ax.grid(True, linestyle='--', alpha=0.3)
            ax.set_xlim(grid_range[0] - 0.5, grid_range[1] + 0.5)
            ax.set_ylim(grid_range[0] - 0.5, grid_range[1] + 0.5)
            ax.set_aspect('equal')
            # Draw original shape
            orig = question.get('original_vertices', [[1, 1], [4, 1], [4, 3]])
            if orig:
                xs = [v[0] for v in orig] + [orig[0][0]]
                ys = [v[1] for v in orig] + [orig[0][1]]
                ax.fill(xs, ys, alpha=0.3, color='#6366f1')
                ax.plot(xs, ys, '-', color='#6366f1', linewidth=2)
                for i, (x, y) in enumerate(orig):
                    ax.annotate(chr(65 + i), (x, y), xytext=(5, 5), textcoords='offset points', fontsize=10, fontweight='bold', color='#6366f1')
            # Draw correct if showing answers
            if show_answer:
                correct_v = question.get('correct_vertices', [])
                if correct_v:
                    xs = [v[0] for v in correct_v] + [correct_v[0][0]]
                    ys = [v[1] for v in correct_v] + [correct_v[0][1]]
                    ax.fill(xs, ys, alpha=0.2, color='#ec4899')
                    ax.plot(xs, ys, '-', color='#ec4899', linewidth=2)

        elif q_type == 'fraction_model':
            denom = question.get('denominator', 4)
            model_type = question.get('model_type', 'area')
            correct_num = question.get('correct_numerator', 0)
            fig, ax = plt.subplots(figsize=(4, 2.5))
            if model_type == 'circle':
                for i in range(denom):
                    start = i * 360 / denom
                    end = (i + 1) * 360 / denom
                    from matplotlib.patches import Wedge
                    filled = show_answer and i < correct_num
                    color = '#6366f1' if filled else 'white'
                    wedge = Wedge((0, 0), 1, start - 90, end - 90, facecolor=color, edgecolor='black', linewidth=1.5, alpha=0.4 if filled else 1)
                    ax.add_patch(wedge)
                ax.set_xlim(-1.3, 1.3)
                ax.set_ylim(-1.3, 1.3)
            else:  # area/strip
                cell_w = 4 / denom
                for i in range(denom):
                    filled = show_answer and i < correct_num
                    color = '#6366f1' if filled else 'white'
                    ax.add_patch(plt.Rectangle((i * cell_w, 0), cell_w, 1.5, facecolor=color, edgecolor='black', linewidth=1.5, alpha=0.4 if filled else 1))
                ax.set_xlim(-0.3, 4.3)
                ax.set_ylim(-0.5, 2.2)
            ax.set_aspect('equal')
            ax.axis('off')
            if show_answer:
                ax.set_title(f'{correct_num}/{denom}', fontsize=14, color='#6366f1', fontweight='bold')

        elif q_type == 'probability_tree':
            fig, ax = plt.subplots(figsize=(6, 4))
            tree = question.get('tree', {})
            def draw_tree(node, x, y, dx, level=0):
                branches = node.get('branches', [])
                n = len(branches)
                if n == 0:
                    return
                dy_step = 2.0 / (2 ** level) if level < 3 else 0.5
                for i, branch in enumerate(branches):
                    offset = (i - (n - 1) / 2) * dy_step
                    nx, ny = x + dx, y + offset
                    ax.annotate('', xy=(nx, ny), xytext=(x, y),
                               arrowprops=dict(arrowstyle='->', color='#6366f1', lw=1.5))
                    prob_text = branch.get('probability', '')
                    if not branch.get('hidden', False) or show_answer:
                        ax.text((x + nx) / 2, (y + ny) / 2 + 0.15, prob_text, ha='center', fontsize=9, color='#f59e0b')
                    ax.text(nx + 0.05, ny, branch.get('label', ''), fontsize=10, fontweight='bold',
                           bbox=dict(boxstyle='round,pad=0.2', facecolor='#e0e7ff', edgecolor='#6366f1'))
                    draw_tree(branch, nx, ny, dx * 0.8, level + 1)
            if tree:
                ax.text(0.5, 0, tree.get('label', 'Start'), fontsize=11, fontweight='bold',
                       bbox=dict(boxstyle='round,pad=0.3', facecolor='#6366f1', edgecolor='none', alpha=0.8),
                       color='white', ha='center')
                draw_tree(tree, 0.5, 0, 2.0)
            ax.set_xlim(-0.5, 7)
            ax.set_ylim(-3, 3)
            ax.axis('off')

        elif q_type == 'tape_diagram':
            tapes = question.get('tapes', [])
            fig, ax = plt.subplots(figsize=(6, 1.5 + len(tapes) * 1.2))
            y_pos = 0
            max_total = max((t.get('total', sum(s.get('value', 1) for s in t.get('segments', []))) for t in tapes), default=10)
            for tape in tapes:
                label = tape.get('label', '')
                segments = tape.get('segments', [])
                total = tape.get('total', sum(s.get('value', 1) for s in segments))
                bar_width = 5 * (total / max_total)
                ax.text(-0.3, y_pos + 0.3, label, ha='right', fontsize=10, fontweight='600')
                x_pos = 0
                for seg in segments:
                    seg_w = bar_width * (seg.get('value', 1) / total)
                    color = seg.get('color', '#6366f1')
                    ax.add_patch(plt.Rectangle((x_pos, y_pos), seg_w, 0.6, facecolor=color, edgecolor='black', linewidth=1, alpha=0.3))
                    if show_answer or not seg.get('hidden', False):
                        ax.text(x_pos + seg_w / 2, y_pos + 0.3, str(seg.get('value', '')), ha='center', fontsize=10, fontweight='bold')
                    else:
                        ax.text(x_pos + seg_w / 2, y_pos + 0.3, '?', ha='center', fontsize=10, fontweight='bold', color='red')
                    x_pos += seg_w
                ax.text(x_pos + 0.2, y_pos + 0.3, f'= {total}' if not tape.get('totalHidden') or show_answer else '= ?', fontsize=10)
                y_pos -= 1.2
            ax.set_xlim(-1.5, 7)
            ax.set_ylim(y_pos - 0.5, 1)
            ax.axis('off')

        elif q_type == 'venn_diagram':
            fig, ax = plt.subplots(figsize=(5, 4))
            sets = question.get('sets', 2)
            labels = question.get('set_labels', ['Set A', 'Set B', 'Set C'])
            from matplotlib.patches import Circle as MplCircle
            if sets >= 3:
                circles = [
                    MplCircle((-0.5, 0.3), 1.2, alpha=0.15, color='#6366f1', linewidth=2),
                    MplCircle((0.5, 0.3), 1.2, alpha=0.15, color='#ec4899', linewidth=2),
                    MplCircle((0, -0.5), 1.2, alpha=0.15, color='#10b981', linewidth=2),
                ]
                label_pos = [(-1.2, 1.2), (1.2, 1.2), (0, -1.8)]
            else:
                circles = [
                    MplCircle((-0.5, 0), 1.2, alpha=0.15, color='#6366f1', linewidth=2),
                    MplCircle((0.5, 0), 1.2, alpha=0.15, color='#ec4899', linewidth=2),
                ]
                label_pos = [(-1.2, 1.2), (1.2, 1.2)]
            for i, c in enumerate(circles):
                ax.add_patch(c)
                c_edge = MplCircle(c.center, c.radius, fill=False, edgecolor=c.get_facecolor(), linewidth=2)
                ax.add_patch(c_edge)
                if i < len(labels):
                    ax.text(label_pos[i][0], label_pos[i][1], labels[i], ha='center', fontsize=11, fontweight='bold')
            # Show region values if answers visible
            if show_answer:
                regions = question.get('correct_values', question.get('regions', {}))
                if sets == 2:
                    pos_map = {'only_a': (-1, 0), 'a_and_b': (0, 0), 'only_b': (1, 0)}
                else:
                    pos_map = {'only_a': (-1, 0.4), 'only_b': (1, 0.4), 'only_c': (0, -1),
                              'a_and_b': (0, 0.4), 'a_and_c': (-0.5, -0.3), 'b_and_c': (0.5, -0.3), 'all': (0, 0)}
                for key, pos in pos_map.items():
                    val = regions.get(key, '')
                    if val:
                        ax.text(pos[0], pos[1], str(val), ha='center', fontsize=12, fontweight='bold')
            ax.set_xlim(-2.5, 2.5)
            ax.set_ylim(-2.5, 2.5)
            ax.set_aspect('equal')
            ax.axis('off')

        elif q_type in ['protractor', 'angle_protractor']:
            fig, ax = plt.subplots(figsize=(4, 3))
            given_angle = question.get('given_angle', 45)
            from matplotlib.patches import Arc, FancyArrowPatch
            # Draw protractor arc
            arc = Arc((0, 0), 3, 3, angle=0, theta1=0, theta2=180, color='#6366f1', linewidth=2)
            ax.add_patch(arc)
            # Tick marks
            for d in range(0, 181, 10):
                rad = _math.radians(d)
                r1, r2 = 1.4, 1.55
                ax.plot([r1 * _math.cos(rad), r2 * _math.cos(rad)], [r1 * _math.sin(rad), r2 * _math.sin(rad)], 'k-', linewidth=0.8)
                if d % 30 == 0:
                    ax.text(1.25 * _math.cos(rad), 1.25 * _math.sin(rad), f'{d}\u00b0', ha='center', fontsize=7)
            # Base ray
            ax.plot([0, 1.6], [0, 0], 'k-', linewidth=2)
            # Angle ray
            angle_rad = _math.radians(given_angle)
            ax.plot([0, 1.6 * _math.cos(angle_rad)], [0, 1.6 * _math.sin(angle_rad)], '-', color='#ec4899', linewidth=2.5)
            # Angle arc
            angle_arc = Arc((0, 0), 0.8, 0.8, angle=0, theta1=0, theta2=given_angle, color='#ec4899', linewidth=2)
            ax.add_patch(angle_arc)
            if show_answer:
                mid_rad = _math.radians(given_angle / 2)
                ax.text(0.5 * _math.cos(mid_rad), 0.5 * _math.sin(mid_rad), f'{given_angle}\u00b0', ha='center', fontsize=11, color='#ec4899', fontweight='bold')
            else:
                mid_rad = _math.radians(given_angle / 2)
                ax.text(0.5 * _math.cos(mid_rad), 0.5 * _math.sin(mid_rad), '?', ha='center', fontsize=14, color='#ec4899', fontweight='bold')
            ax.set_xlim(-1.8, 1.8)
            ax.set_ylim(-0.3, 1.8)
            ax.set_aspect('equal')
            ax.axis('off')

        elif q_type == 'histogram':
            fig, ax = plt.subplots(figsize=(5, 3))
            data = question.get('data', [])
            bins = question.get('bins', 10)
            if data:
                ax.hist(data, bins=bins, color='steelblue', edgecolor='black')
            ax.set_ylabel('Frequency')
            ax.grid(True, axis='y', linestyle='--', alpha=0.5)

        elif q_type == 'pie_chart':
            fig, ax = plt.subplots(figsize=(4, 4))
            chart_data = question.get('chart_data', {})
            categories = chart_data.get('categories', question.get('categories', ['A', 'B', 'C']))
            values = chart_data.get('values', question.get('values', [1, 1, 1]))
            colors = ['#6366f1', '#ec4899', '#10b981', '#f59e0b', '#3b82f6', '#ef4444']
            ax.pie(values, labels=categories, autopct='%1.0f%%', colors=colors[:len(values)])

        else:
            return None

        if fig is None:
            return None

        # Save figure and calculate proper dimensions for PDF
        plt.tight_layout()
        buf = io.BytesIO()
        fig.savefig(buf, format='png', dpi=150, bbox_inches='tight', facecolor='white')
        buf.seek(0)

        # Get the figure's aspect ratio to calculate proper height
        fig_w, fig_h = fig.get_size_inches()
        aspect_ratio = fig_h / fig_w
        plt.close(fig)

        # Determine target width based on type, capped to page width (7 inches usable)
        if q_type in ['coordinate_plane', 'unit_circle', 'transformations']:
            target_width = 3.5 * inch
        elif q_type in ['number_line', 'dot_plot']:
            target_width = 5.5 * inch
        elif q_type in ['box_plot', 'bar_chart', 'function_graph', 'probability_tree',
                         'tape_diagram', 'histogram', 'venn_diagram']:
            target_width = 4.5 * inch
        elif q_type in ['circle', 'regular_polygon', 'cylinder', 'pie_chart']:
            target_width = 3 * inch
        elif q_type in ['fraction_model', 'protractor', 'angle_protractor', 'stem_and_leaf']:
            target_width = 3.5 * inch
        else:
            target_width = 3.5 * inch

        # Calculate height preserving aspect ratio
        target_height = target_width * aspect_ratio

        # Cap height to avoid overflow (max ~5 inches)
        max_height = 5 * inch
        if target_height > max_height:
            target_height = max_height
            target_width = target_height / aspect_ratio

        return Image(buf, width=target_width, height=target_height)

    except Exception as e:
        _logger.error("Error creating visual: %s", e)
        return None


def parse_template_structure(filepath, ext):
    """Parse a template file to understand its structure for export."""
    structure = {
        "columns": [],
        "format": ext,
        "sample_rows": []
    }

    try:
        if ext in ['.csv', '.txt']:
            import csv
            with open(filepath, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                rows = list(reader)
                if rows:
                    structure["columns"] = rows[0]
                    structure["sample_rows"] = rows[1:4]  # First 3 data rows

        elif ext in ['.xlsx', '.xls']:
            from openpyxl import load_workbook
            wb = load_workbook(filepath)
            ws = wb.active
            rows = list(ws.iter_rows(values_only=True))
            if rows:
                structure["columns"] = [str(c) if c else '' for c in rows[0]]
                structure["sample_rows"] = [[str(c) if c else '' for c in row] for row in rows[1:4]]

        elif ext == '.json':
            with open(filepath, 'r') as f:
                data = json.load(f)
                if isinstance(data, list) and data:
                    structure["columns"] = list(data[0].keys()) if isinstance(data[0], dict) else []
                    structure["sample_rows"] = data[:3]
                elif isinstance(data, dict):
                    structure["columns"] = list(data.keys())

    except Exception as e:
        structure["error"] = str(e)

    return structure


def generate_qti_xml(assessment, questions):
    """Generate QTI 1.2 XML for Canvas/LMS import."""
    title = assessment.get('title', 'Assessment')

    xml = f'''<?xml version="1.0" encoding="UTF-8"?>
<questestinterop xmlns="http://www.imsglobal.org/xsd/ims_qtiasiv1p2">
  <assessment ident="{title.replace(' ', '_')}" title="{title}">
    <section ident="root_section">
'''

    for q in questions:
        q_id = f"q_{q.get('number', 1)}"
        q_text = q.get('question', '')
        q_type = q.get('type', 'multiple_choice')
        points = q.get('points', 1)

        if q_type == 'multiple_choice' and q.get('options'):
            xml += f'''      <item ident="{q_id}" title="Question {q.get('number', '')}">
        <itemmetadata>
          <qtimetadata>
            <qtimetadatafield>
              <fieldlabel>question_type</fieldlabel>
              <fieldentry>multiple_choice_question</fieldentry>
            </qtimetadatafield>
            <qtimetadatafield>
              <fieldlabel>points_possible</fieldlabel>
              <fieldentry>{points}</fieldentry>
            </qtimetadatafield>
          </qtimetadata>
        </itemmetadata>
        <presentation>
          <material>
            <mattext texttype="text/html">{q_text}</mattext>
          </material>
          <response_lid ident="response1" rcardinality="Single">
            <render_choice>
'''
            correct_answer = q.get('answer', 'A')
            for i, opt in enumerate(q.get('options', [])):
                opt_id = chr(65 + i)  # A, B, C, D
                opt_text = opt
                if len(opt) > 2 and opt[1] == ')':
                    opt_text = opt[2:].strip()
                xml += f'''              <response_label ident="{opt_id}">
                <material>
                  <mattext texttype="text/html">{opt_text}</mattext>
                </material>
              </response_label>
'''
            xml += f'''            </render_choice>
          </response_lid>
        </presentation>
        <resprocessing>
          <outcomes>
            <decvar maxvalue="100" minvalue="0" varname="SCORE" vartype="Decimal"/>
          </outcomes>
          <respcondition continue="No">
            <conditionvar>
              <varequal respident="response1">{correct_answer}</varequal>
            </conditionvar>
            <setvar action="Set" varname="SCORE">100</setvar>
          </respcondition>
        </resprocessing>
      </item>
'''

    xml += '''    </section>
  </assessment>
</questestinterop>'''

    return xml


def _get_export_dir():
    """Get temp directory for exports (study guides, flashcards, etc.)."""
    import tempfile
    d = os.path.join(tempfile.gettempdir(), "graider_exports")
    os.makedirs(d, exist_ok=True)
    return d
