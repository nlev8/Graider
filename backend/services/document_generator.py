"""
Document Generator Service
===========================
Creates rich .docx documents with configurable visual styles.
Supports headings, paragraphs with markdown bold/italic, bullet lists,
numbered lists, and tables. Styles can be saved and reused.
"""

import os
import re
import json
from urllib.parse import quote

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


DOCUMENTS_DIR = os.path.expanduser("~/Downloads/Graider/Documents")
STYLES_DIR = os.path.expanduser("~/.graider_data/doc_styles")


DEFAULT_STYLE = {
    "title_font_name": "Georgia",
    "title_font_size": 24,
    "title_bold": True,
    "title_color": None,
    "heading_font_name": "Georgia",
    "heading_sizes": {"1": 18, "2": 14, "3": 12},
    "heading_bold": True,
    "heading_color": "#2F5496",
    "body_font_name": "Calibri",
    "body_font_size": 11,
    "line_spacing": 1.15,
    "table_header_bg": "#4472C4",
    "table_header_text_color": "#FFFFFF",
    "accent_color": "#808080",
}


def _hex_to_rgb(hex_str):
    """Convert '#RRGGBB' to RGBColor. Returns None if invalid."""
    if not hex_str or not isinstance(hex_str, str):
        return None
    hex_str = hex_str.lstrip('#')
    if len(hex_str) != 6:
        return None
    try:
        r, g, b = int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16)
        return RGBColor(r, g, b)
    except (ValueError, TypeError):
        return None


def load_style(style_name):
    """Load a saved style from disk, merged with defaults.

    Args:
        style_name: Name of the saved style (e.g., 'cornell-notes').

    Returns:
        Dict of style properties (defaults filled in for missing keys).
    """
    style = dict(DEFAULT_STYLE)
    if not style_name:
        return style

    filepath = os.path.join(STYLES_DIR, style_name + ".json")
    if not os.path.exists(filepath):
        return style

    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            saved = json.load(f)
        # Merge saved values on top of defaults
        for key, val in saved.items():
            if key in style and val is not None:
                style[key] = val
    except Exception:
        pass

    # Coerce numeric fields to proper types (JSON or AI tools may pass strings)
    _int_keys = {"title_font_size", "body_font_size"}
    _float_keys = {"line_spacing"}
    for k in _int_keys:
        if k in style:
            try:
                style[k] = int(style[k])
            except (ValueError, TypeError):
                style[k] = DEFAULT_STYLE[k]
    for k in _float_keys:
        if k in style:
            try:
                style[k] = float(style[k])
            except (ValueError, TypeError):
                style[k] = DEFAULT_STYLE[k]
    # Coerce heading_sizes values to int
    hs = style.get("heading_sizes")
    if isinstance(hs, dict):
        for level, size in hs.items():
            try:
                hs[level] = int(size)
            except (ValueError, TypeError):
                hs[level] = DEFAULT_STYLE["heading_sizes"].get(level, 14)

    return style


def save_style(name, style_dict):
    """Save a named style to disk.

    Args:
        name: Style name (e.g., 'cornell-notes', 'parent-letter').
        style_dict: Dict of visual properties to save.

    Returns:
        Dict with status and filepath.
    """
    os.makedirs(STYLES_DIR, exist_ok=True)
    filepath = os.path.join(STYLES_DIR, name + ".json")

    # Only keep recognized keys
    allowed_keys = set(DEFAULT_STYLE.keys()) | {"name"}
    filtered = {k: v for k, v in style_dict.items() if k in allowed_keys}
    filtered["name"] = name

    with open(filepath, 'w', encoding='utf-8') as f:
        json.dump(filtered, f, indent=2)

    return {"status": "saved", "style_name": name, "filepath": filepath}


def list_styles():
    """Return list of saved style names.

    Returns:
        Dict with styles list.
    """
    if not os.path.exists(STYLES_DIR):
        return {"styles": []}

    styles = []
    for f in sorted(os.listdir(STYLES_DIR)):
        if f.endswith('.json'):
            styles.append(f.replace('.json', ''))

    return {"styles": styles}


def _parse_markdown_runs(paragraph, text, font_name=None, font_size=None):
    """Parse **bold**, *italic*, ***bold+italic*** in text into formatted runs.

    Args:
        paragraph: python-docx Paragraph object to add runs to.
        text: String that may contain markdown bold/italic markers.
        font_name: Optional font name override.
        font_size: Optional font size in points.
    """
    # Pattern matches: ***bold+italic***, **bold**, *italic*, or plain text
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*)')

    last_end = 0
    for match in pattern.finditer(text):
        # Add plain text before this match
        if match.start() > last_end:
            plain = text[last_end:match.start()]
            run = paragraph.add_run(plain)
            if font_name:
                run.font.name = font_name
            if font_size:
                run.font.size = Pt(font_size)

        # Determine which group matched
        if match.group(2) is not None:
            # ***bold+italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(3) is not None:
            # **bold**
            run = paragraph.add_run(match.group(3))
            run.bold = True
        elif match.group(4) is not None:
            # *italic*
            run = paragraph.add_run(match.group(4))
            run.italic = True

        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = Pt(font_size)

        last_end = match.end()

    # Add remaining plain text
    if last_end < len(text):
        run = paragraph.add_run(text[last_end:])
        if font_name:
            run.font.name = font_name
        if font_size:
            run.font.size = Pt(font_size)


def _apply_style_to_heading(heading, level, style):
    """Apply font name, size, color, bold from style to a heading paragraph.

    Args:
        heading: python-docx heading Paragraph object.
        level: Heading level (1, 2, or 3).
        style: Style dict with heading properties.
    """
    heading_sizes = style.get("heading_sizes", {})
    size = heading_sizes.get(str(level), DEFAULT_STYLE["heading_sizes"].get(str(level), 14))
    font_name = style.get("heading_font_name", "Georgia")
    is_bold = style.get("heading_bold", True)
    color = _hex_to_rgb(style.get("heading_color"))

    for run in heading.runs:
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = is_bold
        if color:
            run.font.color.rgb = color


def _apply_table_header_style(row, style):
    """Apply background color and text styling to a table header row.

    Args:
        row: python-docx table Row object.
        style: Style dict with table_header_bg and table_header_text_color.
    """
    bg_hex = style.get("table_header_bg", "#4472C4")
    text_color = _hex_to_rgb(style.get("table_header_text_color", "#FFFFFF"))
    bg_clean = bg_hex.lstrip('#') if bg_hex else "4472C4"

    for cell in row.cells:
        # Set cell background
        shading = parse_xml(
            '<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_clean)
        )
        cell._tc.get_or_add_tcPr().append(shading)

        # Style the text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = style.get("heading_font_name", "Georgia")
                if text_color:
                    run.font.color.rgb = text_color


def create_document_docx(filepath, title, content_blocks, style):
    """Build a .docx file from structured content blocks with style applied.

    Args:
        filepath: Output path for the .docx file.
        title: Document title (main heading).
        content_blocks: List of dicts with 'type' and content fields.
            Supported types: heading, paragraph, bullet_list, numbered_list, table.
        style: Style dict (from load_style or DEFAULT_STYLE).
    """
    doc = Document()

    # Title
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = style.get("title_font_name", "Georgia")
        run.font.size = Pt(style.get("title_font_size", 24))
        run.bold = style.get("title_bold", True)
        title_color = _hex_to_rgb(style.get("title_color"))
        if title_color:
            run.font.color.rgb = title_color

    body_font = style.get("body_font_name", "Calibri")
    body_size = style.get("body_font_size", 11)

    for block in content_blocks:
        block_type = block.get("type", "paragraph")

        if block_type == "heading":
            level = block.get("level", 2)
            level = max(1, min(3, level))
            h = doc.add_heading(block.get("text", ""), level=level)
            _apply_style_to_heading(h, level, style)

        elif block_type == "paragraph":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = style.get("line_spacing", 1.15)
            _parse_markdown_runs(p, block.get("text", ""), body_font, body_size)

        elif block_type == "bullet_list":
            items = block.get("items", [])
            for item_text in items:
                p = doc.add_paragraph(style='List Bullet')
                p.clear()
                _parse_markdown_runs(p, item_text, body_font, body_size)

        elif block_type == "numbered_list":
            items = block.get("items", [])
            for item_text in items:
                p = doc.add_paragraph(style='List Number')
                p.clear()
                _parse_markdown_runs(p, item_text, body_font, body_size)

        elif block_type == "table":
            rows_data = block.get("rows", [])
            if not rows_data:
                continue
            num_cols = max(len(row) for row in rows_data)
            table = doc.add_table(rows=len(rows_data), cols=num_cols)
            table.style = 'Table Grid'

            for r_idx, row_data in enumerate(rows_data):
                for c_idx, cell_text in enumerate(row_data):
                    if c_idx < num_cols:
                        cell = table.rows[r_idx].cells[c_idx]
                        cell.text = ""
                        p = cell.paragraphs[0]
                        _parse_markdown_runs(p, cell_text, body_font, body_size)

            # Style header row (first row)
            if len(rows_data) > 0:
                _apply_table_header_style(table.rows[0], style)

            doc.add_paragraph()  # spacing after table

        elif block_type == "math":
            try:
                from backend.services.visualization import render_latex, add_image_to_docx
                latex = block.get("latex", "")
                font_size = block.get("font_size", 20)
                img = render_latex(latex, font_size=font_size)
                add_image_to_docx(doc, img, width_inches=4)
            except Exception:
                doc.add_paragraph("[Math image generation failed]")

        elif block_type == "number_line":
            try:
                from backend.services.visualization import create_number_line, add_image_to_docx
                img = create_number_line(
                    min_val=block.get("min", -10),
                    max_val=block.get("max", 10),
                    points=block.get("points"),
                    labels=block.get("labels"),
                    title=block.get("title"),
                    blank=block.get("blank", False),
                )
                add_image_to_docx(doc, img, width_inches=6)
            except Exception:
                doc.add_paragraph("[Number line image generation failed]")

        elif block_type == "coordinate_plane":
            try:
                from backend.services.visualization import create_coordinate_plane, add_image_to_docx
                img = create_coordinate_plane(
                    x_range=tuple(block.get("x_range", [-10, 10])),
                    y_range=tuple(block.get("y_range", [-10, 10])),
                    points=[tuple(p) for p in block.get("points", [])] if block.get("points") else None,
                    labels=block.get("labels"),
                    title=block.get("title"),
                    blank=block.get("blank", False),
                )
                add_image_to_docx(doc, img, width_inches=5)
            except Exception:
                doc.add_paragraph("[Coordinate plane image generation failed]")

        elif block_type == "graph":
            try:
                from backend.services.visualization import (
                    create_bar_chart, create_line_graph, create_scatter_plot, add_image_to_docx
                )
                graph_type = block.get("graph_type", "bar")
                if graph_type == "bar":
                    img = create_bar_chart(
                        categories=block.get("categories", []),
                        values=block.get("values", []),
                        title=block.get("title"),
                        x_label=block.get("x_label"),
                        y_label=block.get("y_label"),
                        blank=block.get("blank", False),
                    )
                elif graph_type == "line":
                    img = create_line_graph(
                        x_data=block.get("x_data", []),
                        y_data=block.get("y_data", []),
                        title=block.get("title"),
                        x_label=block.get("x_label"),
                        y_label=block.get("y_label"),
                        blank=block.get("blank", False),
                    )
                elif graph_type == "scatter":
                    img = create_scatter_plot(
                        x_data=block.get("x_data", []),
                        y_data=block.get("y_data", []),
                        title=block.get("title"),
                        x_label=block.get("x_label"),
                        y_label=block.get("y_label"),
                        show_trend=block.get("show_trend", False),
                        blank=block.get("blank", False),
                    )
                else:
                    doc.add_paragraph(f"[Unknown graph type: {graph_type}]")
                    continue
                add_image_to_docx(doc, img, width_inches=5)
            except Exception:
                doc.add_paragraph("[Graph image generation failed]")

        elif block_type == "box_plot":
            try:
                from backend.services.visualization import create_box_plot, add_image_to_docx
                img = create_box_plot(
                    data=block.get("data", []),
                    labels=block.get("labels"),
                    title=block.get("title"),
                    blank=block.get("blank", False),
                )
                add_image_to_docx(doc, img, width_inches=5)
            except Exception:
                doc.add_paragraph("[Box plot image generation failed]")

        elif block_type == "shape":
            try:
                from backend.services.visualization import create_triangle, create_rectangle, add_image_to_docx
                shape_type = block.get("shape_type", "triangle")
                if shape_type == "triangle":
                    img = create_triangle(
                        base=block.get("base", 6),
                        height=block.get("height", 4),
                        title=block.get("title"),
                        blank=block.get("blank", False),
                    )
                elif shape_type == "rectangle":
                    img = create_rectangle(
                        width=block.get("width", 6),
                        height=block.get("height", 4),
                        title=block.get("title"),
                        blank=block.get("blank", False),
                    )
                else:
                    doc.add_paragraph(f"[Unknown shape type: {shape_type}]")
                    continue
                add_image_to_docx(doc, img, width_inches=4)
            except Exception:
                doc.add_paragraph("[Shape image generation failed]")

    doc.save(filepath)


ASSIGNMENTS_DIR = os.path.expanduser("~/.graider_assignments")


def _build_document_text(title, content_blocks):
    """Build plain-text representation of the document for Grading Setup display."""
    lines = [title, '']
    for block in content_blocks:
        block_type = block.get("type", "paragraph")
        if block_type == "heading":
            lines.append(block.get("text", ""))
            lines.append('')
        elif block_type == "paragraph":
            lines.append(block.get("text", ""))
            lines.append('')
        elif block_type in ("bullet_list", "numbered_list"):
            for i, item in enumerate(block.get("items", []), 1):
                prefix = "- " if block_type == "bullet_list" else str(i) + ". "
                lines.append(prefix + item)
            lines.append('')
        elif block_type == "table":
            for row in block.get("rows", []):
                lines.append(' | '.join(row))
            lines.append('')
    return '\n'.join(lines)


def generate_document(title, content, style_name=None, save_to_builder=False):
    """Entry point: generate a formatted Word document.

    Args:
        title: Document title.
        content: List of content block dicts.
        style_name: Optional name of saved style to apply.
        save_to_builder: If True, also save as an assignment config in Grading Setup.

    Returns:
        Dict with status, filepath, download_url, and style_used.
    """
    os.makedirs(DOCUMENTS_DIR, exist_ok=True)

    style = load_style(style_name)

    safe_title = "".join(c for c in title if c.isalnum() or c in ' -_').strip()
    filename = safe_title + '.docx'
    filepath = os.path.join(DOCUMENTS_DIR, filename)

    content_blocks = content or []
    create_document_docx(filepath, title, content_blocks, style)

    download_url = "/api/download-document/" + quote(filename)

    result = {
        "status": "created",
        "filepath": filepath,
        "filename": filename,
        "download_url": download_url,
        "style_used": style_name or "default",
    }

    if save_to_builder:
        doc_text = _build_document_text(title, content_blocks)
        config = {
            "title": title,
            "subject": "",
            "totalPoints": 100,
            "instructions": "",
            "aliases": [],
            "customMarkers": [],
            "excludeMarkers": [],
            "gradingNotes": "",
            "questions": [],
            "responseSections": [],
            "rubricType": "standard",
            "customRubric": None,
            "useSectionPoints": False,
            "sectionTemplate": "Custom",
            "effortPoints": 15,
            "completionOnly": False,
            "countsTowardsGrade": True,
            "importedDoc": {
                "text": doc_text,
                "html": "",
                "filename": filename,
                "loading": False,
            },
            "worksheetDownloadUrl": download_url,
        }
        os.makedirs(ASSIGNMENTS_DIR, exist_ok=True)
        config_path = os.path.join(ASSIGNMENTS_DIR, safe_title + '.json')
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f, indent=2)
        result["saved_to_builder"] = True
        result["config_name"] = safe_title

    return result
