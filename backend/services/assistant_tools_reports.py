"""
Report / Export / Calendar / Resource Assistant Tools
=====================================================
Tools for exporting grades, generating documents & worksheets,
curriculum standards lookup, calendar management, resource browsing,
parent communication, and assignment config saving.
"""
import os
import csv
import json
import subprocess
import re
import uuid
from datetime import datetime, timedelta
from collections import defaultdict
from urllib.parse import quote

from backend.services.assistant_tools import (
    _load_master_csv, _load_results, _load_roster, _load_settings,
    _load_accommodations, _load_standards, _load_saved_lessons,
    _load_period_class_levels, _load_parent_contacts,
    _load_saved_assignments, _load_calendar, _save_calendar, _load_email_config,
    _fuzzy_name_match, _safe_int_score, _normalize_period,
    _normalize_assignment_name, _get_output_folder,
    ASSIGNMENTS_DIR, EXPORTS_DIR, STANDARDS_DIR, DOCUMENTS_DIR, LESSONS_DIR,
)
from backend.services.assistant_tools_grading import get_missing_assignments

# Constants
CREDS_FILE = os.path.expanduser("~/.graider_data/portal_credentials.json")
PROJECT_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
PARENT_CONTACTS_FILE = os.path.expanduser("~/.graider_data/parent_contacts.json")
CALENDAR_FILE = os.path.expanduser("~/.graider_data/teaching_calendar.json")


# ═══════════════════════════════════════════════════════
# TOOL DEFINITIONS
# ═══════════════════════════════════════════════════════

REPORT_TOOL_DEFINITIONS = [
    {
        "name": "create_focus_assignment",
        "description": "Create an assignment in Focus gradebook via browser automation. Requires VPortal credentials to be configured in Settings. The browser will open for the teacher to complete 2FA and verify before saving.",
        "input_schema": {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "Assignment name"
                },
                "category": {
                    "type": "string",
                    "description": "Category (e.g., 'Assessments', 'Classwork')"
                },
                "points": {
                    "type": "integer",
                    "description": "Point value for the assignment"
                },
                "date": {
                    "type": "string",
                    "description": "Due date in MM/DD/YYYY format"
                },
                "description": {
                    "type": "string",
                    "description": "Assignment description"
                }
            },
            "required": ["name"]
        }
    },
    {
        "name": "export_grades_csv",
        "description": "Export grades as a Focus SIS-compatible CSV file. Generates per-period CSVs with Student ID and Score columns.",
        "input_schema": {
            "type": "object",
            "properties": {
                "assignment": {
                    "type": "string",
                    "description": "Assignment name to filter export (omit for all)"
                },
                "period": {
                    "type": "string",
                    "description": "Period to filter export (omit for all periods)"
                }
            }
        }
    },
    {
        "name": "lookup_student_info",
        "description": "Look up student contact and roster information. Returns student ID, local ID, grade level, period, course codes, student email, parent emails, parent phone numbers, 504 plan status, detailed contacts (up to 3 guardians with names, relationships, roles), and full student schedule (all periods with teachers and courses). Search by student name, ID, or period. Supports BATCH lookup via student_ids array — use this to get parent contacts for multiple students in one call (e.g., after querying grades to find failing students).",
        "input_schema": {
            "type": "object",
            "properties": {
                "student_name": {
                    "type": "string",
                    "description": "Student name to search for (partial match, case-insensitive)"
                },
                "student_id": {
                    "type": "string",
                    "description": "Single student ID number to look up directly"
                },
                "student_ids": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of student ID numbers for batch lookup. Use this to get contacts for multiple students at once."
                },
                "period": {
                    "type": "string",
                    "description": "List all students in this period (e.g., 'Period 1', '1')"
                }
            }
        }
    },
    {
        "name": "generate_worksheet",
        "description": "Generate a structured worksheet document (Cornell Notes, Fill-in-the-Blank, short-answer, vocabulary) from a reading or topic. Creates a downloadable Word document with an embedded invisible answer key for consistent AI grading. Automatically saved to Grading Setup. Use when the teacher asks to create a worksheet, assignment, or activity from a reading, textbook page, or topic.",
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Worksheet title (e.g., 'Cornell Notes - Expanding into Native American Lands')"
                },
                "worksheet_type": {
                    "type": "string",
                    "enum": ["cornell-notes", "fill-in-blank", "short-answer", "vocabulary"],
                    "description": "Type of worksheet to generate"
                },
                "vocab_terms": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "term": {"type": "string"},
                            "definition": {"type": "string", "description": "Expected definition (for answer key)"}
                        },
                        "required": ["term"]
                    },
                    "description": "Vocabulary terms with expected definitions"
                },
                "questions": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "question": {"type": "string"},
                            "expected_answer": {"type": "string", "description": "Expected answer (for answer key)"},
                            "points": {"type": "integer", "default": 10},
                            "visual": {
                                "type": "object",
                                "description": "Optional visual element embedded above the answer box. Supports: math (LaTeX expression), number_line, coordinate_plane, graph (bar/line/scatter), box_plot, shape (triangle/rectangle). Example: {\"type\":\"math\",\"latex\":\"\\\\frac{1}{2}+\\\\frac{1}{3}\"} or {\"type\":\"shape\",\"shape_type\":\"triangle\",\"base\":8,\"height\":6} or {\"type\":\"graph\",\"graph_type\":\"bar\",\"categories\":[\"A\",\"B\"],\"values\":[3,7],\"title\":\"Sales\"}",
                                "properties": {
                                    "type": {"type": "string", "enum": ["math", "number_line", "coordinate_plane", "graph", "box_plot", "shape"]},
                                    "latex": {"type": "string", "description": "LaTeX expression (for math type)"},
                                    "font_size": {"type": "integer", "description": "Font size for math rendering"},
                                    "min": {"type": "number"}, "max": {"type": "number"},
                                    "points": {"type": "array", "items": {"type": "number"}},
                                    "labels": {"type": "object"}, "title": {"type": "string"},
                                    "blank": {"type": "boolean", "description": "If true, hide data for student to fill in"},
                                    "x_range": {"type": "array", "items": {"type": "number"}}, "y_range": {"type": "array", "items": {"type": "number"}},
                                    "graph_type": {"type": "string", "enum": ["bar", "line", "scatter"]},
                                    "categories": {"type": "array", "items": {"type": "string"}}, "values": {"type": "array", "items": {"type": "number"}},
                                    "x_data": {"type": "array", "items": {"type": "number"}}, "y_data": {"type": "array", "items": {"type": "number"}},
                                    "x_label": {"type": "string"}, "y_label": {"type": "string"},
                                    "data": {"type": "array", "items": {"type": "number"}}, "shape_type": {"type": "string", "enum": ["triangle", "rectangle"]},
                                    "base": {"type": "number"}, "height": {"type": "number"}, "width": {"type": "number"}
                                },
                                "required": ["type"]
                            }
                        },
                        "required": ["question"]
                    },
                    "description": "Questions with expected answers. Each question can include an optional 'visual' object to embed a math expression, graph, shape, or other visual above the answer box."
                },
                "summary_prompt": {
                    "type": "string",
                    "description": "Instruction for the summary section (e.g., 'Summarize the reading in 3-5 sentences...')"
                },
                "summary_key_points": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Key points that should appear in a good summary"
                },
                "total_points": {
                    "type": "integer",
                    "default": 100,
                    "description": "Total point value for the worksheet"
                },
                "style_name": {
                    "type": "string",
                    "description": "Name of a saved visual style to apply. Omit to use defaults."
                }
            },
            "required": ["title", "worksheet_type"]
        }
    },
    {
        "name": "generate_document",
        "description": "Generate a formatted Word document (.docx) with rich typography. Use for study guides, reference sheets, parent letters, lesson outlines, rubrics, or any document the teacher requests. NOT for gradeable worksheets (use generate_worksheet for those).",
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Document title (main heading)"
                },
                "content": {
                    "type": "array",
                    "description": "Ordered content blocks. Text supports **bold** and *italic* markdown. Visual blocks (math, number_line, coordinate_plane, graph, box_plot, shape) embed images.",
                    "items": {
                        "type": "object",
                        "properties": {
                            "type": {"type": "string", "enum": ["heading", "paragraph", "bullet_list", "numbered_list", "table", "math", "number_line", "coordinate_plane", "graph", "box_plot", "shape"]},
                            "text": {"type": "string", "description": "Text content (heading/paragraph)"},
                            "level": {"type": "integer", "description": "Heading level 1-3"},
                            "items": {"type": "array", "items": {"type": "string"}, "description": "List items"},
                            "rows": {"type": "array", "items": {"type": "array", "items": {"type": "string"}}, "description": "Table rows (first = header)"},
                            "latex": {"type": "string", "description": "LaTeX math expression for type=math (e.g., '\\\\frac{-b \\\\pm \\\\sqrt{b^2-4ac}}{2a}')"},
                            "font_size": {"type": "integer", "description": "Font size for math rendering (default 20)"},
                            "min": {"type": "number", "description": "Min value for number_line"},
                            "max": {"type": "number", "description": "Max value for number_line"},
                            "points": {"type": "array", "items": {"type": "number"}, "description": "Points to plot. number_line: array of numbers. coordinate_plane: array of [x,y] pairs."},
                            "labels": {"type": "array", "items": {"type": "string"}, "description": "Labels for points or data sets"},
                            "title": {"type": "string", "description": "Title for visual blocks"},
                            "blank": {"type": "boolean", "description": "If true, create blank template for students to fill in"},
                            "x_range": {"type": "array", "items": {"type": "number"}, "description": "coordinate_plane: [min, max] for x-axis"},
                            "y_range": {"type": "array", "items": {"type": "number"}, "description": "coordinate_plane: [min, max] for y-axis"},
                            "graph_type": {"type": "string", "enum": ["bar", "line", "scatter"], "description": "Type of graph for type=graph"},
                            "categories": {"type": "array", "items": {"type": "string"}, "description": "Bar chart category labels"},
                            "values": {"type": "array", "items": {"type": "number"}, "description": "Bar chart values"},
                            "x_data": {"type": "array", "items": {"type": "number"}, "description": "X values for line/scatter graphs"},
                            "y_data": {"type": "array", "items": {"type": "number"}, "description": "Y values for line/scatter graphs"},
                            "x_label": {"type": "string", "description": "X-axis label for graphs"},
                            "y_label": {"type": "string", "description": "Y-axis label for graphs"},
                            "show_trend": {"type": "boolean", "description": "Show trend line on scatter plot"},
                            "data": {"type": "array", "items": {"type": "array", "items": {"type": "number"}}, "description": "box_plot: array of number arrays (one per data set)"},
                            "shape_type": {"type": "string", "enum": ["triangle", "rectangle"], "description": "Shape type for type=shape"},
                            "base": {"type": "number", "description": "Triangle base length"},
                            "height": {"type": "number", "description": "Shape height (triangle or rectangle)"},
                            "width": {"type": "number", "description": "Rectangle width"}
                        },
                        "required": ["type"]
                    }
                },
                "style_name": {
                    "type": "string",
                    "description": "Name of a saved visual style to apply. Omit to use defaults."
                },
                "save_to_builder": {
                    "type": "boolean",
                    "description": "If true, also save this document to Grading Setup for grading. Only set true if the teacher confirms they want it saved."
                }
            },
            "required": ["title", "content"]
        }
    },
    {
        "name": "save_document_style",
        "description": "Save a visual style (fonts, sizes, colors, spacing) so future documents of this type always look the same. Use when the teacher says they like how a document looks and want to reuse that formatting.",
        "input_schema": {
            "type": "object",
            "properties": {
                "name": {"type": "string", "description": "Style name (e.g., 'cornell-notes', 'parent-letter')"},
                "style": {
                    "type": "object",
                    "description": "Visual properties to save",
                    "properties": {
                        "title_font_name": {"type": "string"},
                        "title_font_size": {"type": "integer"},
                        "title_bold": {"type": "boolean"},
                        "heading_font_name": {"type": "string"},
                        "heading_sizes": {"type": "object"},
                        "body_font_name": {"type": "string"},
                        "body_font_size": {"type": "integer"},
                        "line_spacing": {"type": "number"},
                        "table_header_bg": {"type": "string"},
                        "accent_color": {"type": "string"}
                    }
                }
            },
            "required": ["name", "style"]
        }
    },
    {
        "name": "list_document_styles",
        "description": "List saved document visual styles. Use before generating a document to check if a matching style exists.",
        "input_schema": {"type": "object", "properties": {}}
    },
    {
        "name": "generate_csv",
        "description": "Generate a downloadable spreadsheet file (XLSX or CSV). Use .xlsx for Wayground quizzes (required format) and any polished spreadsheet. Use .csv for simple data exports. When generating a Wayground quiz, use .xlsx and match the EXACT column structure from the Assessment Templates in your context (Question Text, Question Type, Option 1-5, Correct Answer, Time in seconds, Image Link, Answer explanation).",
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Name for the file. Use .xlsx for Wayground and polished spreadsheets, .csv for simple exports. Defaults to .xlsx if no extension given."
                },
                "headers": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Column headers for the CSV"
                },
                "rows": {
                    "type": "array",
                    "items": {
                        "type": "array",
                        "items": {"type": "string"}
                    },
                    "description": "Data rows. Each row is an array of string values matching the headers."
                }
            },
            "required": ["filename", "headers", "rows"]
        }
    },
    {
        "name": "recommend_next_lesson",
        "description": "Analyze student performance and recommend what the next lesson should focus on. Provides DIFFERENTIATED recommendations by class level (advanced/standard/support periods) with DOK-appropriate standards. Also identifies IEP/504 accommodation patterns — whether accommodated students struggled differently and what modifications may help. Cross-references rubric breakdowns, unanswered questions, developing skills, and curriculum standards. Use when teacher asks 'what should I teach next?', 'what do students need to work on?', or 'how should I differentiate the next lesson?'.",
        "input_schema": {
            "type": "object",
            "properties": {
                "assignment_name": {
                    "type": "string",
                    "description": "Assignment to base recommendations on (partial match). Omit to use the most recent assignment."
                },
                "period": {
                    "type": "string",
                    "description": "Focus on a specific period (optional). Different periods may have different class levels and weaknesses."
                },
                "num_assignments": {
                    "type": "integer",
                    "description": "Number of recent assignments to analyze (default 1, max 5). More assignments = broader trend analysis."
                }
            }
        }
    },
    {
        "name": "get_standards",
        "description": "Look up curriculum standards for the teacher's state and subject. Returns ALL standards when no topic is specified, or filter by keyword. Use this to get full details (vocabulary, learning targets, essential questions) for standards. For a quick overview of all standards, use list_all_standards first.",
        "input_schema": {
            "type": "object",
            "properties": {
                "topic": {
                    "type": "string",
                    "description": "Topic keyword to filter standards (e.g., 'fractions', 'civil war', 'photosynthesis'). Matches against benchmark text, topics, and vocabulary."
                },
                "dok_max": {
                    "type": "integer",
                    "description": "Maximum DOK level to include (1-4). Use 2 for support classes, 3 for standard, 4 for advanced."
                }
            }
        }
    },
    {
        "name": "get_recent_lessons",
        "description": "List saved lesson plans, optionally filtered by unit name. Shows what has been taught recently — topics, standards covered, vocabulary, and objectives per day. Use this when the teacher asks to create a quiz or worksheet 'for this unit', 'for what we've been doing', or references past lessons.",
        "input_schema": {
            "type": "object",
            "properties": {
                "unit_name": {
                    "type": "string",
                    "description": "Filter by unit name (partial match, case-insensitive). Omit to show all units and recent lessons."
                }
            }
        }
    },
    {
        "name": "get_calendar",
        "description": "Read the teaching calendar — scheduled lessons and holidays for a date range. AUTHORITATIVE source for what the teacher is teaching. If scheduled_lessons is non-empty, those lessons ARE scheduled — always reference them by title and topic. Use when the teacher asks 'what am I teaching this week?', 'what's on my calendar?', 'generate a worksheet for Tuesday', or anything about their schedule. When asked about a specific day, set both start_date and end_date to that exact date.",
        "input_schema": {
            "type": "object",
            "properties": {
                "start_date": {
                    "type": "string",
                    "description": "Start date in YYYY-MM-DD format. Defaults to today. For a specific day like 'Tuesday Feb 17', use '2026-02-17'."
                },
                "end_date": {
                    "type": "string",
                    "description": "End date in YYYY-MM-DD format. Defaults to 7 days from start. For a specific day, set equal to start_date."
                }
            }
        }
    },
    {
        "name": "schedule_lesson",
        "description": "Schedule a saved lesson plan onto the teaching calendar on a specific date. Use when the teacher says 'schedule the Revolution unit starting Monday', 'put Unit 3 on the calendar', or asks to plan out their week/month. For multi-day lessons, call this once per day with incrementing day_number and dates.",
        "input_schema": {
            "type": "object",
            "properties": {
                "date": {
                    "type": "string",
                    "description": "Date to schedule in YYYY-MM-DD format"
                },
                "unit": {
                    "type": "string",
                    "description": "Unit name the lesson belongs to"
                },
                "lesson_title": {
                    "type": "string",
                    "description": "Title of the lesson"
                },
                "day_number": {
                    "type": "integer",
                    "description": "Day number within the lesson (e.g., 1 for Day 1). Optional."
                },
                "lesson_file": {
                    "type": "string",
                    "description": "Path to lesson file like 'Unit Name/Lesson Title.json'. Optional."
                }
            },
            "required": ["date", "lesson_title"]
        }
    },
    {
        "name": "unschedule_lesson",
        "description": "Remove a lesson from the teaching calendar. Use when the teacher says 'remove the lesson on Friday', 'clear March 5th', or when correcting a schedule by removing the old entry before adding the new one.",
        "input_schema": {
            "type": "object",
            "properties": {
                "date": {
                    "type": "string",
                    "description": "Date to remove lesson(s) from in YYYY-MM-DD format"
                },
                "lesson_title": {
                    "type": "string",
                    "description": "Title of the specific lesson to remove. If omitted, removes ALL lessons on that date."
                }
            },
            "required": ["date"]
        }
    },
    {
        "name": "add_calendar_holiday",
        "description": "Add a holiday or break to the teaching calendar. Use when the teacher says 'add Spring Break', 'mark Monday as a holiday', or 'we're off next Friday'.",
        "input_schema": {
            "type": "object",
            "properties": {
                "date": {
                    "type": "string",
                    "description": "Start date in YYYY-MM-DD format"
                },
                "name": {
                    "type": "string",
                    "description": "Name of the holiday or break (e.g., 'Spring Break', 'Teacher Workday')"
                },
                "end_date": {
                    "type": "string",
                    "description": "End date for multi-day breaks in YYYY-MM-DD format. Omit for single-day holidays."
                }
            },
            "required": ["date", "name"]
        }
    },
    {
        "name": "list_all_standards",
        "description": "Get a compact index of ALL curriculum standards for the teacher's subject. Returns every standard code with a short benchmark summary. Use this first to see what standards exist, then use get_standards with a topic keyword to get full details (vocabulary, learning targets, essential questions) for specific standards.",
        "input_schema": {
            "type": "object",
            "properties": {}
        }
    },
    {
        "name": "list_resources",
        "description": "List all supporting documents uploaded by the teacher (curriculum guides, pacing calendars, rubrics, etc.). Shows filename, type, description, and size. Use this to discover what reference materials are available before reading specific ones with read_resource.",
        "input_schema": {
            "type": "object",
            "properties": {}
        }
    },
    {
        "name": "read_resource",
        "description": "Read the full text content of an uploaded supporting document. The filenames are listed in the UPLOADED REFERENCE DOCUMENTS section of your context — pass the exact filename. Supports PDF, DOCX, DOC, TXT, and MD files. ALWAYS use this when the teacher mentions 'curriculum map', 'pacing guide', 'resources', or any uploaded document. Call this BEFORE answering any question about curriculum content, pacing, or uploaded materials.",
        "input_schema": {
            "type": "object",
            "properties": {
                "filename": {
                    "type": "string",
                    "description": "Exact filename of the document to read (from list_resources results)"
                }
            },
            "required": ["filename"]
        }
    },
    {
        "name": "send_parent_emails",
        "description": "Generate a preview of personalized emails to parents/guardians via Outlook. Returns a preview with a confirmation action — messages are NOT sent by this tool. The teacher confirms sending via the UI. Supports template placeholders: {student_first_name}, {student_last_name}, {student_name}, {parent_name}, {period}, {teacher_name}, {subject_area}. Can target specific students or auto-find students with zero submissions.",
        "input_schema": {
            "type": "object",
            "properties": {
                "email_subject": {
                    "type": "string",
                    "description": "Email subject line. Supports {student_first_name} etc."
                },
                "email_body": {
                    "type": "string",
                    "description": "Email body text. Use placeholders and newlines for paragraphs."
                },
                "student_names": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of student names to email parents of."
                },
                "period": {
                    "type": "string",
                    "description": "Email all parents in this period."
                },
                "zero_submissions": {
                    "type": "boolean",
                    "description": "If true, auto-target students with zero assignment submissions."
                }
            },
            "required": ["email_subject", "email_body"]
        }
    },
    {
        "name": "send_focus_comms",
        "description": "Generate a preview of email and/or SMS messages to parents via Focus SIS Communications. This is the DEFAULT and PREFERRED method for contacting parents — use this instead of send_parent_emails unless the teacher specifically asks for Outlook. Returns a preview with a confirmation action — messages are NOT sent by this tool. The teacher confirms sending via the UI. Supports template placeholders: {student_first_name}, {student_last_name}, {student_name}. When sending both email AND SMS, the SMS should be a short notification like 'Please check your email for a message regarding {subject}.' When sending SMS-only, omit email_body.",
        "input_schema": {
            "type": "object",
            "properties": {
                "email_subject": {
                    "type": "string",
                    "description": "Subject line for the email. Also used as the topic reference in SMS-only messages. Supports {student_first_name}, {student_last_name}, {student_name} placeholders."
                },
                "email_body": {
                    "type": "string",
                    "description": "Email body text. Omit to send SMS only. Supports template placeholders. Use newlines for paragraphs."
                },
                "sms_body": {
                    "type": "string",
                    "description": "SMS text (max 500 chars). If sending both email and SMS, use a short notification like 'Please check your email for a message about [topic] from [teacher].' If omitted, only email is sent."
                },
                "student_names": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "List of student names (Last, First format) to message parents of."
                },
                "period": {
                    "type": "string",
                    "description": "Send to all parents in this period."
                },
                "recipient_type": {
                    "type": "string",
                    "enum": ["Primary Contacts", "Students"],
                    "description": "Who receives the message. 'Primary Contacts' sends to parents (default). 'Students' sends directly to the student."
                }
            },
            "required": ["email_subject"]
        }
    },
    {
        "name": "confirm_and_send",
        "description": "Execute a pending email/SMS send after the teacher has confirmed the preview. Call this ONLY after showing a preview from send_focus_comms or send_parent_emails AND the teacher has confirmed they want to send it (e.g., 'yes', 'send it', 'looks good'). This triggers the actual Playwright automation to send the messages. Takes no parameters — it reads the pending payload automatically.",
        "input_schema": {
            "type": "object",
            "properties": {},
        }
    },
    {
        "name": "save_assignment_config",
        "description": "Save or update an assignment config in Grading Setup. Use when teacher asks to save an assignment, update point values, rename, change rubric type, or modify grading notes for an existing assignment. Can update individual fields without overwriting the entire config. IMPORTANT: If the teacher uploaded a document, always include document_text so the editor can display and mark it.",
        "input_schema": {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Assignment title (used as filename). If updating existing, must match exactly."
                },
                "document_text": {
                    "type": "string",
                    "description": "Full text content of the assignment document. Include this when saving a new assignment from an uploaded file so the editor can display it."
                },
                "questions": {
                    "type": "array",
                    "description": "Array of question objects: {id, type, prompt, points, marker, expected_answer}",
                    "items": {
                        "type": "object",
                        "properties": {
                            "id": {"type": "integer"},
                            "type": {"type": "string", "enum": ["short_answer", "vocab_term", "written", "fill_in_blank"]},
                            "prompt": {"type": "string"},
                            "points": {"type": "integer"},
                            "marker": {"type": "string"},
                            "expected_answer": {"type": "string"}
                        },
                        "required": ["type", "prompt", "points"]
                    }
                },
                "totalPoints": {
                    "type": "integer",
                    "description": "Total points for the assignment"
                },
                "effortPoints": {
                    "type": "integer",
                    "description": "Points allocated for effort (default 15)"
                },
                "gradingNotes": {
                    "type": "string",
                    "description": "Grading notes / expected answers for the AI grader"
                },
                "rubricType": {
                    "type": "string",
                    "enum": ["standard", "cornell-notes", "fill-in-blank"],
                    "description": "Rubric type"
                },
                "customMarkers": {
                    "type": "array",
                    "description": "Section markers: [{start, points, type}]",
                    "items": {
                        "type": "object",
                        "properties": {
                            "start": {"type": "string"},
                            "points": {"type": "integer"},
                            "type": {"type": "string"}
                        }
                    }
                }
            },
            "required": ["title"]
        }
    },
]


# ═══════════════════════════════════════════════════════
# HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════

def _analyze_group_weaknesses(group_results):
    """Shared weakness analysis for a group of student results.
    Returns dict with category_weaknesses, content_gaps, developing_skills, strengths, scores."""
    total = len(group_results)
    if not total:
        return None

    # Category breakdown
    cat_totals = defaultdict(list)
    for r in group_results:
        bd = r.get('breakdown', {})
        if bd:
            for cat, val in bd.items():
                try:
                    cat_totals[cat].append(int(float(val)) if val else 0)
                except (ValueError, TypeError):
                    pass

    category_weaknesses = []
    for cat, vals in cat_totals.items():
        avg = round(sum(vals) / len(vals), 1) if vals else 0
        max_possible = max(vals) if vals else 0
        zeros = sum(1 for v in vals if v == 0)
        category_weaknesses.append({
            "category": cat,
            "average": avg,
            "max_seen": max_possible,
            "zero_count": zeros,
            "zero_pct": round(zeros / len(vals) * 100, 1) if vals else 0,
        })
    category_weaknesses.sort(key=lambda x: x["average"])

    # Unanswered questions
    unanswered_counts = defaultdict(int)
    for r in group_results:
        uq = r.get('unanswered_questions')
        if uq and isinstance(uq, list):
            for q in uq:
                unanswered_counts[q] += 1
    content_gaps = sorted(unanswered_counts.items(), key=lambda x: x[1], reverse=True)[:8]

    # Skills
    developing_freq = defaultdict(int)
    strength_freq = defaultdict(int)
    for r in group_results:
        skills = r.get('skills_demonstrated', {})
        if isinstance(skills, dict):
            for s in (skills.get('developing', []) or []):
                developing_freq[s.strip()] += 1
            for s in (skills.get('strengths', []) or []):
                strength_freq[s.strip()] += 1

    top_developing = sorted(developing_freq.items(), key=lambda x: x[1], reverse=True)[:6]
    top_strengths = sorted(strength_freq.items(), key=lambda x: x[1], reverse=True)[:4]

    # Scores
    scores = [_safe_int_score(r.get('score')) for r in group_results]
    failing = [s for s in scores if s < 70]
    avg_score = round(sum(scores) / len(scores), 1) if scores else 0

    # Omissions
    omission_count = sum(1 for r in group_results
                         if r.get('unanswered_questions') and len(r.get('unanswered_questions', [])) > 0)

    return {
        "total_students": total,
        "average_score": avg_score,
        "failing_count": len(failing),
        "failing_pct": round(len(failing) / total * 100, 1) if total else 0,
        "omission_rate": round(omission_count / total * 100, 1) if total else 0,
        "category_weaknesses": category_weaknesses,
        "content_gaps": [{"topic": q, "students_missed": c,
                          "pct": round(c / total * 100, 1)}
                         for q, c in content_gaps],
        "developing_skills": [{"skill": s, "count": c} for s, c in top_developing],
        "student_strengths": [{"skill": s, "count": c} for s, c in top_strengths],
    }


def _match_standards(weakness_data, standards, target_dok=None):
    """Match curriculum standards to identified weaknesses, optionally filtering by DOK level."""
    if not standards or not weakness_data:
        return []

    weakness_keywords = set()
    for gap in weakness_data.get("content_gaps", []):
        for word in gap["topic"].lower().split():
            if len(word) > 3:
                weakness_keywords.add(word)
    for dev in weakness_data.get("developing_skills", []):
        for word in dev["skill"].lower().split():
            if len(word) > 3:
                weakness_keywords.add(word)

    relevant = []
    for std in standards:
        topics = [t.lower() for t in std.get('topics', [])]
        benchmark = std.get('benchmark', '').lower()
        vocab = [v.lower() for v in std.get('vocabulary', [])]
        all_text = ' '.join(topics + [benchmark] + vocab)

        match_count = sum(1 for kw in weakness_keywords if kw in all_text)
        if match_count > 0:
            std_dok = std.get('dok', '')
            dok_match = True
            if target_dok is not None and std_dok:
                try:
                    dok_val = int(std_dok) if str(std_dok).isdigit() else 0
                    dok_match = dok_val <= target_dok
                except (ValueError, TypeError):
                    dok_match = True

            relevant.append({
                "code": std.get('code', ''),
                "benchmark": std.get('benchmark', '')[:200],
                "topics": std.get('topics', []),
                "dok": std_dok,
                "essential_questions": std.get('essential_questions', [])[:2],
                "learning_targets": std.get('learning_targets', [])[:2],
                "relevance_score": match_count,
                "dok_appropriate": dok_match,
            })

    relevant.sort(key=lambda x: (-int(x["dok_appropriate"]), -x["relevance_score"]))
    return relevant[:5]


def _parse_curriculum_map_for_dates(start_date, end_date):
    """Parse curriculum map DOCX and return structured data for a date range.

    Looks for unit/week entries whose date ranges overlap with the requested dates.
    Returns dict with unit name, benchmarks, vocabulary, textbook, resources, or None.
    """
    import re
    from datetime import datetime as _dt

    # Find curriculum map file
    if not os.path.isdir(DOCUMENTS_DIR):
        return None
    curriculum_file = None
    for fname in os.listdir(DOCUMENTS_DIR):
        if not fname.lower().endswith('.docx') or fname.endswith('.meta.json'):
            continue  # Only parse DOCX files (python-docx can't read PDFs)
        meta_path = os.path.join(DOCUMENTS_DIR, fname + ".meta.json")
        if os.path.exists(meta_path):
            try:
                with open(meta_path, 'r') as f:
                    meta = json.load(f)
                if meta.get("doc_type") == "curriculum":
                    curriculum_file = os.path.join(DOCUMENTS_DIR, fname)
                    break
            except Exception:
                continue
        elif 'curriculum' in fname.lower() or 'pacing' in fname.lower():
            curriculum_file = os.path.join(DOCUMENTS_DIR, fname)
            break
    if not curriculum_file or not os.path.exists(curriculum_file):
        return None

    try:
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.table import Table
    except ImportError:
        return None

    # Parse target date range
    try:
        req_start = _dt.strptime(start_date, '%Y-%m-%d')
        req_end = _dt.strptime(end_date, '%Y-%m-%d')
    except Exception:
        return None

    # Helper to parse flexible date strings like "Feb. 17th", "March 12th", "January 6th"
    def _parse_map_date(s, year=None):
        if not year:
            year = req_start.year
        s = re.sub(r'(st|nd|rd|th)\b', '', s.strip().rstrip('.'))
        for fmt in ('%B %d', '%b %d', '%b. %d'):
            try:
                d = _dt.strptime(s.strip(), fmt)
                return d.replace(year=year)
            except ValueError:
                continue
        return None

    doc = Document(curriculum_file)

    # Pass 1: Find unit that covers these dates.
    # Strategy: scan detail table headers (e.g., "Unit 7: Manifest Destiny (4 weeks)")
    # then check benchmark rows for matching date ranges. Pick the narrowest match.
    candidates = []  # (span_days, unit_name, dates_str, weeks, table_index)

    # First pass: scan detail table headers + benchmark rows
    current_unit_header = None
    current_table_idx = None
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            seen_cells = set()
            cells = []
            for c in row.cells:
                t = c.text.strip()
                if t and t not in seen_cells:
                    seen_cells.add(t)
                    cells.append(t)
            row_text = ' '.join(cells)

            # Detect unit header rows like "Unit 7: Manifest Destiny (4 weeks)"
            unit_header = re.search(r'(Unit\s+\d+:\s*[^\(]+)', row_text)
            if unit_header and ('weeks' in row_text.lower() or 'Standard' in row_text):
                current_unit_header = unit_header.group(1).strip()
                current_table_idx = ti
                continue

            # Look for date ranges in benchmark rows
            row_lower = row_text.lower()
            # Pattern 1: "February 17th - March 12th" (both have month)
            date_ranges = re.findall(
                r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z.]*\s+\d{1,2}(?:st|nd|rd|th)?)\s*[-\u2013\u2014]\s*((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z.]*\s+\d{1,2}(?:st|nd|rd|th)?)',
                row_lower, re.IGNORECASE
            )
            # Pattern 2: "Feb. 2nd -6th" (end date has no month — inherit from start)
            short_ranges = re.findall(
                r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z.]*\s+\d{1,2}(?:st|nd|rd|th)?)\s*[-\u2013\u2014]\s*(\d{1,2}(?:st|nd|rd|th)?)\b',
                row_lower, re.IGNORECASE
            )
            for d_start_str, day_only in short_ranges:
                # Infer month from start date
                month_match = re.match(r'((?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)[a-z.]*)', d_start_str, re.IGNORECASE)
                if month_match:
                    d_end_str = f"{month_match.group(1)} {day_only}"
                    date_ranges.append((d_start_str, d_end_str))

            for d_start_str, d_end_str in date_ranges:
                d_start = _parse_map_date(d_start_str)
                d_end = _parse_map_date(d_end_str)
                if d_start and d_end and d_end >= req_start and d_start <= req_end:
                    span = (d_end - d_start).days
                    # Use current_unit_header if in a detail table, else look in same row
                    unit_name = current_unit_header
                    if not unit_name:
                        for cell_text in cells:
                            um = re.search(r'(Unit\s+\d+:\s*[^|]+)', cell_text)
                            if um:
                                unit_name = um.group(1).strip()
                                break
                    if unit_name:
                        # Extract week numbers from range patterns like "25-28"
                        joined = ' '.join(cells)
                        week_ranges = re.findall(r'\b(\d{1,2})-(\d{1,2})\b', joined)
                        weeks = None
                        if week_ranges:
                            s, e = int(week_ranges[0][0]), int(week_ranges[0][1])
                            if 1 <= s <= 52 and 1 <= e <= 52:
                                weeks = list(range(s, e + 1))
                        candidates.append((span, unit_name, f"{d_start_str.strip()} \u2013 {d_end_str.strip()}", weeks, ti))

    if not candidates:
        return None

    # Pick the narrowest matching date range
    candidates.sort(key=lambda x: x[0])
    _, matching_unit, matching_dates, matching_weeks, _ = candidates[0]

    # Pass 2: Find the unit's detail section — benchmarks, vocabulary, textbook, resources
    unit_keyword = matching_unit.split(':')[0].strip()  # e.g., "Unit 7"
    benchmarks = []
    vocabulary = []
    textbook = []
    resources = {"nearpod_activities": [], "nearpod_lessons": [], "videos": [], "dbqs": []}
    in_unit_section = False

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            row_text = ' '.join(cells)

            # Detect start of our unit's detail section
            if unit_keyword in row_text and ('Standard' in row_text or 'weeks' in row_text.lower()):
                in_unit_section = True
                continue

            # Detect end — next unit header or resources section end
            if in_unit_section:
                other_unit = re.search(r'Unit\s+\d+:', row_text)
                if other_unit and unit_keyword not in row_text:
                    in_unit_section = False
                    continue

            if not in_unit_section:
                continue

            # Extract benchmarks (format: "Quarter 3 | 25-28 | dates | SS.8.A.4.1 | description")
            for cell_text in cells:
                bm_match = re.match(r'(SS\.\d+\.[A-Z]+\.\d+(?:\.\d+)?)\s*$', cell_text.strip())
                if bm_match:
                    code = bm_match.group(1)
                    # Next non-empty cell after the code is the description
                    idx = cells.index(cell_text)
                    desc = cells[idx + 1] if idx + 1 < len(cells) else ""
                    if desc and not desc.startswith('SS.') and not desc.startswith('Quarter'):
                        benchmarks.append({"code": code, "description": desc})
                    elif not any(b["code"] == code for b in benchmarks):
                        benchmarks.append({"code": code, "description": ""})

            # Extract vocabulary
            if 'academic vocabulary' in row_text.lower():
                for cell_text in cells:
                    if 'academic vocabulary' in cell_text.lower():
                        vocab_text = re.sub(r'^academic vocabulary:?\s*', '', cell_text, flags=re.IGNORECASE)
                        # Handle typos: period after word (3+ chars) followed by space+uppercase → comma
                        # Skip initials like "F." in "Stephen F. Austin"
                        vocab_text = re.sub(r'(?<=[a-z]{3})\.\s+(?=[A-Z\u201c"])', ', ', vocab_text)
                        terms = [t.strip().strip('"').strip('\u201c').strip('\u201d') for t in vocab_text.split(',')]
                        vocabulary = [t for t in terms if t and len(t) > 1]
                        break

            # Extract textbook references
            if 'chapter' in row_text.lower() and ('pgs' in row_text.lower() or 'teacher textbook' in row_text.lower()):
                for cell_text in cells:
                    ch_match = re.search(r'Chapter\s+\d+\s*\(pgs?\.\s*[\w\-\u2013]+\)', cell_text)
                    if ch_match and ch_match.group() not in textbook:
                        textbook.append(ch_match.group())

            # Extract Nearpod/video resources
            if 'nearpod' in row_text.lower() or 'video' in row_text.lower() or 'dbq' in row_text.lower():
                for cell_text in cells:
                    cl = cell_text.lower()
                    if 'activities:' in cl:
                        items = re.sub(r'^activities:\s*', '', cell_text, flags=re.IGNORECASE)
                        resources["nearpod_activities"] = [i.strip() for i in items.split('\n') if i.strip()]
                    elif 'lessons:' in cl and 'nearpod' in ' '.join(cells).lower():
                        items = re.sub(r'^lessons:\s*', '', cell_text, flags=re.IGNORECASE)
                        resources["nearpod_lessons"] = [i.strip() for i in items.split('\n') if i.strip()]
                    elif 'videos:' in cl:
                        items = re.sub(r'^videos:\s*', '', cell_text, flags=re.IGNORECASE)
                        resources["videos"] = [i.strip() for i in items.split('\n') if i.strip()]
                    elif 'dbq' in cl:
                        dbq_items = re.findall(r'(?:DBQ\s*\|?\s*)?(.+?\(SS\.[^)]+\))', cell_text)
                        resources["dbqs"] = [d.strip() for d in dbq_items if d.strip()]

    # Deduplicate benchmarks
    seen_codes = set()
    unique_benchmarks = []
    for b in benchmarks:
        if b["code"] not in seen_codes:
            seen_codes.add(b["code"])
            unique_benchmarks.append(b)

    # Clean empty resource lists
    resources = {k: v for k, v in resources.items() if v}

    return {
        "unit": matching_unit,
        "weeks": matching_weeks,
        "dates": matching_dates,
        "benchmarks": unique_benchmarks,
        "vocabulary": vocabulary,
        "textbook": textbook,
        "resources": resources,
    }


def _extract_pdf_text(filepath):
    """Extract text from a PDF file path using PyMuPDF."""
    try:
        import fitz
        doc = fitz.open(filepath)
        pages = []
        for page in doc:
            pages.append(page.get_text())
        page_count = len(pages)
        doc.close()
        return "\n\n".join(pages), page_count
    except ImportError:
        return "[PDF extraction requires PyMuPDF: pip install pymupdf]", 0
    except Exception as e:
        return f"[Error extracting PDF: {e}]", 0


def _extract_docx_text(filepath):
    """Extract text from a DOCX file path using python-docx."""
    try:
        from docx import Document
        from docx.text.paragraph import Paragraph
        from docx.table import Table

        doc = Document(filepath)
        full_text = []
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = Paragraph(element, doc)
                if para.text.strip():
                    full_text.append(para.text)
            elif element.tag.endswith('tbl'):
                table = Table(element, doc)
                for row in table.rows:
                    # Deduplicate merged cells (same text repeated across merged columns)
                    seen = set()
                    row_text = []
                    for cell in row.cells:
                        txt = cell.text.strip()
                        if txt and txt not in seen:
                            seen.add(txt)
                            row_text.append(txt)
                    if row_text:
                        full_text.append(' | '.join(row_text))
        return '\n'.join(full_text)
    except ImportError:
        return "[DOCX extraction requires python-docx: pip install python-docx]"
    except Exception as e:
        return f"[Error extracting DOCX: {e}]"


def _parse_student_name(name):
    """Parse 'Last, First Middle' or 'First Last' into components.

    Returns dict with first_name, last_name, full_name (as 'First Last').
    """
    if not name:
        return {"first_name": "Student", "last_name": "", "full_name": "Student"}

    name = name.strip()
    for sep in [',', ';']:
        if sep in name:
            parts = name.split(sep, 1)
            last = parts[0].strip()
            after = parts[1].strip()
            first = after.split()[0] if after else last
            return {
                "first_name": first,
                "last_name": last,
                "full_name": first + " " + last,
            }
    words = name.split()
    return {
        "first_name": words[0],
        "last_name": words[-1] if len(words) > 1 else "",
        "full_name": name,
    }


def _fill_email_template(template, replacements):
    """Replace {placeholders} in a template string."""
    result = template
    for key, value in replacements.items():
        result = result.replace("{" + key + "}", str(value))
    return result


# ═══════════════════════════════════════════════════════
# TOOL HANDLER FUNCTIONS
# ═══════════════════════════════════════════════════════

def create_focus_assignment(name, category=None, points=None, date=None, description=None):
    """Launch Focus automation to create an assignment."""
    # Check credentials
    if not os.path.exists(CREDS_FILE):
        return {"error": "VPortal credentials not configured. Go to Settings > Tools to set them up."}

    script_path = os.path.join(PROJECT_ROOT, "focus-automation.js")
    if not os.path.exists(script_path):
        return {"error": "focus-automation.js not found in project root."}

    cmd = ["node", script_path, "assignment", "--name", name]
    if category:
        cmd.extend(["--category", category])
    if points:
        cmd.extend(["--points", str(points)])
    if date:
        cmd.extend(["--date", date])
    if description:
        cmd.extend(["--description", description])

    try:
        process = subprocess.Popen(
            cmd,
            cwd=PROJECT_ROOT,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        return {
            "status": "launched",
            "message": f"Browser automation started for '{name}'. Check your phone for 2FA approval, then review the form before saving.",
            "pid": process.pid
        }
    except FileNotFoundError:
        return {"error": "Node.js not found. Make sure Node.js is installed."}
    except Exception as e:
        return {"error": f"Failed to launch automation: {str(e)}"}


def export_grades_csv(assignment=None, period=None):
    """Export grades as Focus-compatible CSV."""
    results = _load_results()
    if not results:
        return {"error": "No grading results to export"}

    os.makedirs(EXPORTS_DIR, exist_ok=True)

    # Filter
    filtered = results
    if assignment:
        filtered = [r for r in filtered
                    if assignment.lower() in r.get('assignment', '').lower()]
    if period:
        filtered = [r for r in filtered
                    if r.get('period', '') == period or r.get('quarter', '') == period]

    if not filtered:
        return {"error": "No results match the specified filters"}

    # Group by period
    by_period = defaultdict(list)
    for r in filtered:
        p = r.get('period', r.get('quarter', 'All'))
        by_period[p].append(r)

    safe_name = assignment or 'grades'
    safe_name = ''.join(c if c.isalnum() or c in ' -_' else '' for c in safe_name).strip().replace(' ', '_')

    exported_files = []
    total_rows = 0
    for p, items in by_period.items():
        safe_period = p.replace(' ', '_').replace('/', '-')
        filename = f"{safe_name}_{safe_period}.csv"
        filepath = os.path.join(EXPORTS_DIR, filename)

        csv_lines = ['Student ID,Score']
        matched = 0
        for r in items:
            student_id = r.get('student_id', '')
            score = r.get('score', 0)
            if student_id:
                csv_lines.append(f"{student_id},{score}")
                matched += 1
            else:
                name = r.get('student_name', 'Unknown')
                csv_lines.append(f"# {name},{score}")

        with open(filepath, 'w', encoding='utf-8') as f:
            f.write('\n'.join(csv_lines))

        exported_files.append({
            "file": filename,
            "period": p,
            "rows": matched,
            "download_url": "/api/download-export/" + filename,
        })
        total_rows += matched

    download_urls = [{"url": f["download_url"], "filename": f["file"]} for f in exported_files]

    result = {
        "status": "exported",
        "export_dir": EXPORTS_DIR,
        "files": exported_files,
        "total_rows": total_rows,
        "download_urls": download_urls,
    }

    # Single-file convenience: set top-level download_url for the simple download button path
    if len(exported_files) == 1:
        result["download_url"] = exported_files[0]["download_url"]
        result["filename"] = exported_files[0]["file"]

    return result


def recommend_next_lesson(assignment_name=None, period=None, num_assignments=1):
    """Analyze performance and recommend next lesson focus with period differentiation and IEP awareness."""
    results = _load_results()
    if not results:
        return {"error": "No grading results available"}

    num_assignments = min(num_assignments or 1, 5)

    # Determine which assignment(s) to analyze
    if assignment_name:
        target_results = [r for r in results
                          if assignment_name.lower() in r.get('assignment', '').lower()]
    else:
        sorted_results = sorted(results, key=lambda r: r.get('graded_at', ''), reverse=True)
        recent_assignments = []
        seen = set()
        for r in sorted_results:
            a = r.get('assignment', '')
            if a and a not in seen:
                seen.add(a)
                recent_assignments.append(a)
            if len(recent_assignments) >= num_assignments:
                break
        target_results = [r for r in results if r.get('assignment', '') in recent_assignments]

    if period:
        target_results = [r for r in target_results if r.get('period', '') == period]

    if not target_results:
        return {"error": "No matching results found for analysis"}

    analyzed_assignments = list(set(r.get('assignment', '') for r in target_results))

    # Load supplementary data
    period_levels = _load_period_class_levels()
    accommodations = _load_accommodations()
    standards = _load_standards()
    settings = _load_settings()
    config = settings.get('config', {})
    global_notes = settings.get('globalAINotes', '')
    saved_lessons = _load_saved_lessons()

    # === Overall weakness analysis ===
    overall = _analyze_group_weaknesses(target_results)

    # === Per-class-level analysis ===
    # Group results by class level (advanced/standard/support)
    level_groups = defaultdict(list)
    periods_by_level = defaultdict(list)

    for r in target_results:
        p = r.get('period', '')
        level = period_levels.get(p, 'standard')
        level_groups[level].append(r)
        if p not in periods_by_level[level]:
            periods_by_level[level].append(p)

    # DOK targets by class level
    dok_targets = {"advanced": 4, "standard": 3, "support": 2}

    class_level_analysis = {}
    for level in ["advanced", "standard", "support"]:
        group = level_groups.get(level, [])
        if not group:
            continue
        weakness = _analyze_group_weaknesses(group)
        if not weakness:
            continue

        target_dok = dok_targets.get(level, 3)
        matched_standards = _match_standards(weakness, standards, target_dok=target_dok)

        # Build skill needs for this level
        skill_needs = []
        cw = weakness["category_weaknesses"]
        if cw:
            weakest = cw[0]
            skill_needs.append(f"Weakest area: {weakest['category']} (avg {weakest['average']})")
        cg = weakness["content_gaps"]
        if cg:
            skill_needs.append(f"Most skipped: '{cg[0]['topic']}' ({cg[0]['pct']}%)")
        ds = weakness["developing_skills"]
        if ds:
            skill_needs.append(f"Top developing: '{ds[0]['skill']}' ({ds[0]['count']} students)")
        if weakness["omission_rate"] > 40:
            skill_needs.append(f"High omission rate: {weakness['omission_rate']}%")

        class_level_analysis[level] = {
            "class_level": level,
            "periods": periods_by_level[level],
            "target_dok": target_dok,
            "identified_needs": skill_needs,
            "recommended_standards": matched_standards,
            **weakness,
        }

    # === IEP/504 Accommodation Analysis ===
    accommodation_analysis = None
    if accommodations:
        # Match accommodated students to results by student_id or name
        accom_ids = set(accommodations.keys())
        accom_results = []
        non_accom_results = []
        matched_students = {}

        for r in target_results:
            sid = r.get('student_id', '')
            sname = r.get('student_name', '').lower().replace(' ', '_')
            if sid in accom_ids:
                accom_results.append(r)
                matched_students[sid] = accommodations[sid]
            elif sname in accom_ids:
                accom_results.append(r)
                matched_students[sname] = accommodations[sname]
            else:
                non_accom_results.append(r)

        if accom_results:
            accom_weakness = _analyze_group_weaknesses(accom_results)
            non_accom_weakness = _analyze_group_weaknesses(non_accom_results) if non_accom_results else None

            # Collect unique presets across matched students
            all_presets = set()
            for info in matched_students.values():
                for p in info.get('presets', []):
                    all_presets.add(p)

            score_gap = None
            if accom_weakness and non_accom_weakness:
                score_gap = round(non_accom_weakness["average_score"] - accom_weakness["average_score"], 1)

            accommodation_analysis = {
                "accommodated_student_count": len(accom_results),
                "total_student_count": len(target_results),
                "accommodation_presets_in_use": sorted(all_presets),
                "accommodated_avg_score": accom_weakness["average_score"] if accom_weakness else None,
                "non_accommodated_avg_score": non_accom_weakness["average_score"] if non_accom_weakness else None,
                "score_gap": score_gap,
                "accommodated_failing_pct": accom_weakness["failing_pct"] if accom_weakness else None,
                "accommodated_omission_rate": accom_weakness["omission_rate"] if accom_weakness else None,
                "accommodated_weaknesses": accom_weakness["category_weaknesses"][:3] if accom_weakness else [],
                "note": "IEP/504 students may need modified lesson pacing, scaffolded activities, or alternative assessments.",
            }

    # === Overall standards (no DOK filter) for the combined view ===
    overall_standards = _match_standards(overall, standards, target_dok=None)

    # === Build top-level identified needs ===
    skill_needs = []
    if overall["category_weaknesses"]:
        weakest = overall["category_weaknesses"][0]
        skill_needs.append(f"Weakest rubric area: {weakest['category']} (avg {weakest['average']})")
    if overall["content_gaps"]:
        top_gap = overall["content_gaps"][0]
        skill_needs.append(f"Most skipped section: '{top_gap['topic']}' ({top_gap['students_missed']} students, "
                          f"{top_gap['pct']}%)")
    if overall["developing_skills"]:
        skill_needs.append(f"Top developing skill: '{overall['developing_skills'][0]['skill']}' "
                          f"({overall['developing_skills'][0]['count']} students)")
    if overall["omission_rate"] > 40:
        skill_needs.append(f"High omission rate: {overall['omission_rate']}% of students "
                          f"left questions blank — consider assignment completion mini-lesson")

    # Note if class levels diverge significantly
    if len(class_level_analysis) > 1:
        level_avgs = {lv: d["average_score"] for lv, d in class_level_analysis.items()}
        max_avg = max(level_avgs.values())
        min_avg = min(level_avgs.values())
        if max_avg - min_avg > 10:
            skill_needs.append(
                f"Large gap between class levels: {max_avg} (highest) vs {min_avg} (lowest) — "
                f"differentiated lesson planning recommended"
            )

    return {
        "analyzed_assignments": analyzed_assignments,
        "period_filter": period or "all periods",
        "total_students": overall["total_students"],
        "average_score": overall["average_score"],
        "failing_count": overall["failing_count"],
        "failing_pct": overall["failing_pct"],
        "identified_needs": skill_needs,
        "category_weaknesses": overall["category_weaknesses"],
        "content_gaps": overall["content_gaps"],
        "developing_skills": overall["developing_skills"],
        "student_strengths": overall["student_strengths"],
        "relevant_standards": overall_standards,
        "existing_lessons": [l["title"] for l in saved_lessons],
        "teacher_subject": config.get('subject', ''),
        "teacher_grade": config.get('grade_level', ''),
        "period_differentiation_notes": global_notes[:500] if global_notes else "",
        "class_level_breakdown": class_level_analysis if class_level_analysis else None,
        "period_class_levels": period_levels if period_levels else None,
        "accommodation_analysis": accommodation_analysis,
    }


def lookup_student_info(student_name=None, student_id=None, student_ids=None, period=None):
    """Look up student roster and contact information.
    Supports batch lookup via student_ids (list of IDs)."""
    roster = _load_roster()
    parent_contacts = _load_parent_contacts()
    results_json = _load_results()

    # Build email lookup from grading results (student_id -> email)
    email_lookup = {}
    for r in results_json:
        sid = r.get('student_id', '')
        email = r.get('student_email', '')
        if sid and email:
            email_lookup[sid] = email

    if not roster and not parent_contacts:
        return {"error": "No student roster data found. Import from Focus SIS or upload class period CSVs in Settings."}

    # Batch lookup by student_ids list
    if student_ids and isinstance(student_ids, list):
        id_set = set(str(sid) for sid in student_ids)
        matches = [s for s in roster if s["student_id"] in id_set]
        # Also check parent contacts for IDs not found in roster
        found_ids = set(s["student_id"] for s in matches)
        for sid in id_set - found_ids:
            contact = parent_contacts.get(sid)
            if contact:
                matches.append({
                    "name": contact.get("student_name", "Unknown"),
                    "student_id": sid,
                    "local_id": "",
                    "grade": "",
                    "period": contact.get("period", ""),
                })
    else:
        # Single lookup mode
        matches = roster
        if student_id:
            matches = [s for s in matches if s["student_id"] == str(student_id)]
        if student_name:
            matches = [s for s in matches if _fuzzy_name_match(student_name, s["name"])]
        if period:
            # Normalize period input: "1" -> matches "Period 1", "Period 1" -> matches "Period 1"
            period_lower = period.lower().strip()
            matches = [s for s in matches
                       if period_lower in s["period"].lower()
                       or (period_lower.isdigit() and f"period {period_lower}" in s["period"].lower())]

        if not matches and student_id:
            # Try parent contacts as fallback (has student_id keys even without roster)
            contact = parent_contacts.get(str(student_id))
            if contact:
                matches = [{
                    "name": contact.get("student_name", "Unknown"),
                    "student_id": str(student_id),
                    "local_id": "",
                    "grade": "",
                    "period": contact.get("period", ""),
                }]

        if not matches and student_name:
            # Try parent contacts by name (fuzzy word match)
            for sid, contact in parent_contacts.items():
                if _fuzzy_name_match(student_name, contact.get("student_name", "")):
                    matches.append({
                        "name": contact.get("student_name", "Unknown"),
                        "student_id": sid,
                        "local_id": "",
                        "grade": "",
                        "period": contact.get("period", ""),
                    })

    if not matches:
        return {"error": "No students found matching the search criteria.", "searched": {
            "name": student_name, "id": student_id, "ids": student_ids, "period": period
        }}

    # Deduplicate by student_id (batch mode may have overlap)
    seen_ids = set()
    unique_matches = []
    for s in matches:
        sid = s["student_id"]
        if sid not in seen_ids:
            seen_ids.add(sid)
            unique_matches.append(s)

    # Enrich each match with contact info, email, 504 status, schedule, and course codes
    students = []
    for s in unique_matches:
        sid = s["student_id"]
        contact = parent_contacts.get(sid, {})
        student_email = email_lookup.get(sid, "")

        entry = {
            "name": s["name"],
            "student_id": sid,
            "local_id": s.get("local_id", ""),
            "grade_level": s.get("grade", ""),
            "period": s["period"],
            "course_codes": s.get("course_codes", []),
            "student_email": student_email,
            "parent_emails": contact.get("parent_emails", []),
            "parent_phones": contact.get("parent_phones", []),
            "has_504": contact.get("has_504", False),
            "contacts": contact.get("contacts", []),
            "schedule": contact.get("schedule", []),
        }
        students.append(entry)

    return {
        "students": students,
        "total_found": len(students),
    }


def generate_worksheet_tool(title, worksheet_type, vocab_terms=None, questions=None,
                            summary_prompt=None, summary_key_points=None,
                            total_points=100, style_name=None):
    """Generate a .docx worksheet and save to Grading Setup."""
    try:
        from backend.services.worksheet_generator import generate_worksheet
        # Load subject from teacher settings
        settings = _load_settings()
        config = settings.get('config', {})
        subject = config.get('subject', '')
        return generate_worksheet(
            title=title,
            worksheet_type=worksheet_type,
            vocab_terms=vocab_terms,
            questions=questions,
            summary_prompt=summary_prompt,
            summary_key_points=summary_key_points,
            total_points=total_points,
            subject=subject,
            style_name=style_name
        )
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    except Exception as e:
        return {"error": "Failed to generate worksheet: " + str(e)}


def generate_document_tool(title, content, style_name=None, save_to_builder=False):
    """Generate a formatted Word document with rich typography."""
    try:
        from backend.services.document_generator import generate_document
        return generate_document(
            title=title, content=content,
            style_name=style_name, save_to_builder=save_to_builder
        )
    except ImportError:
        return {"error": "python-docx not installed. Run: pip install python-docx"}
    except Exception as e:
        return {"error": "Failed to generate document: " + str(e)}


def generate_csv_tool(filename, headers, rows):
    """Generate a downloadable CSV or XLSX file based on filename extension."""
    from urllib.parse import quote

    EXPORT_DIR = os.path.expanduser("~/Downloads/Graider/Exports")
    os.makedirs(EXPORT_DIR, exist_ok=True)

    # Sanitize filename
    safe_name = "".join(c for c in filename if c.isalnum() or c in ' -_.').strip()
    is_xlsx = safe_name.lower().endswith('.xlsx')

    # Default to .xlsx if no recognized extension
    if not safe_name.lower().endswith(('.csv', '.xlsx')):
        safe_name += '.xlsx'
        is_xlsx = True

    filepath = os.path.join(EXPORT_DIR, safe_name)

    if is_xlsx:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Quiz"

        # Write header row with styling
        header_font = Font(bold=True, color="FFFFFF", size=11)
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin')
        )
        for col_idx, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_idx, value=header)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', wrap_text=True)
            cell.border = thin_border

        # Write data rows
        for row_idx, row in enumerate(rows, 2):
            for col_idx, value in enumerate(row, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True, vertical='top')

        # Auto-fit column widths (approximate)
        for col_idx, header in enumerate(headers, 1):
            max_len = len(str(header))
            for row in rows:
                if col_idx - 1 < len(row):
                    max_len = max(max_len, len(str(row[col_idx - 1])))
            ws.column_dimensions[openpyxl.utils.get_column_letter(col_idx)].width = min(max_len + 4, 50)

        wb.save(filepath)
    else:
        import csv
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            for row in rows:
                writer.writerow(row)

    download_url = "/api/download-csv/" + quote(safe_name)

    return {
        "status": "created",
        "filename": safe_name,
        "filepath": filepath,
        "download_url": download_url,
        "row_count": len(rows),
        "columns": headers,
    }


def save_document_style_tool(name, style):
    """Save a named visual style for documents."""
    try:
        from backend.services.document_generator import save_style
        return save_style(name=name, style_dict=style)
    except Exception as e:
        return {"error": "Failed to save style: " + str(e)}


def list_document_styles_tool():
    """List saved document visual styles."""
    try:
        from backend.services.document_generator import list_styles
        return list_styles()
    except Exception as e:
        return {"error": "Failed to list styles: " + str(e)}


def get_standards_tool(topic=None, dok_max=None):
    """Look up curriculum standards filtered by topic and DOK level."""
    all_standards = _load_standards()
    if not all_standards:
        settings = _load_settings()
        config = settings.get('config', {})
        subj = config.get('subject', 'unknown')
        st = config.get('state', 'unknown')
        return {"error": f"No standards found for {subj} in {st}. Check Settings > Subject and State."}

    results = all_standards

    # Filter by DOK level
    if dok_max is not None:
        results = [s for s in results if s.get('dok', 99) <= dok_max]

    # Filter by topic keyword
    if topic:
        topic_lower = topic.lower()
        filtered = []
        for s in results:
            searchable = " ".join([
                s.get('benchmark', ''),
                " ".join(s.get('topics', [])),
                " ".join(s.get('vocabulary', [])),
                s.get('item_specs', ''),
            ]).lower()
            if topic_lower in searchable:
                filtered.append(s)
        results = filtered

    if not results:
        return {"error": f"No standards found matching '{topic or 'all'}' (DOK <= {dok_max or 'any'})"}

    return {
        "count": len(results),
        "standards": [
            {
                "code": s.get("code", ""),
                "benchmark": s.get("benchmark", ""),
                "dok": s.get("dok"),
                "topics": s.get("topics", []),
                "vocabulary": s.get("vocabulary", []),
                "learning_targets": s.get("learning_targets", []),
                "essential_questions": s.get("essential_questions", []),
                "sample_assessment": s.get("sample_assessment", ""),
            }
            for s in results
        ]
    }


def list_all_standards_tool():
    """Return a compact index of ALL curriculum standards for the teacher's subject."""
    all_standards = _load_standards()
    if not all_standards:
        settings = _load_settings()
        config = settings.get('config', {})
        subj = config.get('subject', 'unknown')
        st = config.get('state', 'unknown')
        return {"error": f"No standards found for {subj} in {st}. Check Settings > Subject and State."}

    settings = _load_settings()
    config = settings.get('config', {})

    compact = []
    for s in all_standards:
        benchmark = s.get("benchmark", "")
        compact.append({
            "code": s.get("code", ""),
            "benchmark": benchmark[:120] + "..." if len(benchmark) > 120 else benchmark,
            "dok": s.get("dok"),
            "topics": s.get("topics", []),
        })

    return {
        "subject": config.get('subject', 'unknown'),
        "state": config.get('state', 'unknown'),
        "grade_level": config.get('grade_level', 'unknown'),
        "total_count": len(compact),
        "standards": compact,
    }


def get_recent_lessons(unit_name=None):
    """List saved lesson plans with full detail for document generation context."""
    if not os.path.exists(LESSONS_DIR):
        return {"error": "No saved lessons found. Generate and save lesson plans in the Planner tab first."}

    lessons = []
    for unit_dir in os.listdir(LESSONS_DIR):
        unit_path = os.path.join(LESSONS_DIR, unit_dir)
        if not os.path.isdir(unit_path):
            continue

        # Filter by unit name if provided
        if unit_name and unit_name.lower() not in unit_dir.lower():
            continue

        for fname in os.listdir(unit_path):
            if not fname.endswith('.json'):
                continue
            try:
                with open(os.path.join(unit_path, fname), 'r', encoding='utf-8') as fh:
                    data = json.load(fh)

                # Extract day-level details
                days_summary = []
                all_vocab = []
                all_standards = []
                for day in data.get('days', []):
                    day_info = {
                        "day": day.get("day"),
                        "topic": day.get("topic", ""),
                        "objective": day.get("objective", ""),
                    }
                    # Collect standards addressed
                    stds = day.get("standards_addressed", [])
                    if stds:
                        day_info["standards"] = stds
                        all_standards.extend(stds)
                    # Collect vocabulary
                    vocab = day.get("vocabulary", [])
                    for v in vocab:
                        term = v.get("term", v) if isinstance(v, dict) else str(v)
                        if term and term not in all_vocab:
                            all_vocab.append(term)
                    days_summary.append(day_info)

                lessons.append({
                    "unit": unit_dir,
                    "title": data.get("title", fname.replace(".json", "")),
                    "overview": data.get("overview", ""),
                    "essential_questions": data.get("essential_questions", []),
                    "num_days": len(data.get("days", [])),
                    "days": days_summary,
                    "vocabulary": all_vocab,
                    "standards_covered": list(set(all_standards)),
                    "saved_at": data.get("_saved_at", ""),
                })
            except Exception:
                continue

    if not lessons:
        msg = f"No lessons found for unit '{unit_name}'." if unit_name else "No saved lessons found."
        return {"error": msg}

    # Sort by saved_at (most recent first), cap at 10
    lessons.sort(key=lambda x: x.get("saved_at", ""), reverse=True)
    lessons = lessons[:10]

    # Group by unit for readability
    units = {}
    for lesson in lessons:
        unit = lesson["unit"]
        if unit not in units:
            units[unit] = []
        units[unit].append(lesson)

    return {
        "total_lessons": len(lessons),
        "units": list(units.keys()),
        "lessons": lessons,
    }


def get_calendar(start_date=None, end_date=None):
    """Read the teaching calendar for a date range."""
    cal = _load_calendar()

    if not start_date:
        start_date = datetime.now().strftime('%Y-%m-%d')
    if not end_date:
        end_dt = datetime.strptime(start_date, '%Y-%m-%d') + timedelta(days=7)
        end_date = end_dt.strftime('%Y-%m-%d')

    # Filter lessons in range
    lessons = [s for s in cal.get("scheduled_lessons", [])
               if start_date <= s.get("date", "") <= end_date]
    lessons.sort(key=lambda s: s.get("date", ""))

    # Filter holidays in range (including multi-day overlaps)
    holidays = []
    for h in cal.get("holidays", []):
        h_start = h.get("date", "")
        h_end = h.get("end_date", h_start)
        if h_end >= start_date and h_start <= end_date:
            holidays.append(h)

    result = {
        "start_date": start_date,
        "end_date": end_date,
        "scheduled_lessons": lessons,
        "holidays": holidays,
        "total_lessons": len(lessons),
        "total_holidays": len(holidays),
    }

    # Always include curriculum map data for curricular context
    try:
        curriculum_data = _parse_curriculum_map_for_dates(start_date, end_date)
        if curriculum_data:
            result["curriculum_map"] = curriculum_data
            if not lessons:
                result["note"] = ("No lessons scheduled for this period. Curriculum map data shows what should be covered. "
                                  "Also check uploaded reference documents in your system context for additional details.")
    except Exception:
        pass

    return result


def schedule_lesson_tool(date, lesson_title, unit=None, day_number=None, lesson_file=None):
    """Schedule a lesson on the teaching calendar."""
    import uuid as _uuid
    if not date or not lesson_title:
        return {"error": "date and lesson_title are required"}

    cal = _load_calendar()

    # Pick a color based on unit name
    unit_colors = ['#6366f1', '#8b5cf6', '#ec4899', '#f59e0b', '#22c55e', '#06b6d4', '#ef4444']
    color_idx = hash(unit or '') % len(unit_colors)

    entry = {
        "id": str(_uuid.uuid4()),
        "date": date,
        "unit": unit or "",
        "lesson_title": lesson_title,
        "day_number": day_number,
        "lesson_file": lesson_file or "",
        "color": unit_colors[color_idx],
    }

    # Remove existing lesson on same date with same day_number or title to avoid duplicates
    cal["scheduled_lessons"] = [
        s for s in cal["scheduled_lessons"]
        if not (s["date"] == date and (
            (day_number is not None and s.get("day_number") == day_number) or
            s.get("lesson_title") == lesson_title
        ))
    ]
    cal["scheduled_lessons"].append(entry)
    _save_calendar(cal)

    return {"status": "scheduled", "entry": entry}


def unschedule_lesson_tool(date, lesson_title=None):
    """Remove a lesson from the teaching calendar by date and optional title."""
    if not date:
        return {"error": "date is required"}

    cal = _load_calendar()
    before = len(cal["scheduled_lessons"])

    if lesson_title:
        cal["scheduled_lessons"] = [
            s for s in cal["scheduled_lessons"]
            if not (s["date"] == date and s.get("lesson_title") == lesson_title)
        ]
    else:
        cal["scheduled_lessons"] = [
            s for s in cal["scheduled_lessons"]
            if s["date"] != date
        ]

    removed = before - len(cal["scheduled_lessons"])
    if removed == 0:
        return {"status": "not_found", "message": f"No lessons found on {date}"}

    _save_calendar(cal)
    return {"status": "removed", "removed_count": removed}


def add_calendar_holiday(date, name, end_date=None):
    """Add a holiday or break to the teaching calendar."""
    if not date or not name:
        return {"error": "date and name are required"}

    cal = _load_calendar()

    holiday = {"date": date, "name": name}
    if end_date:
        holiday["end_date"] = end_date

    # Remove existing holiday on same date to avoid duplicates
    cal["holidays"] = [h for h in cal["holidays"] if h["date"] != date]
    cal["holidays"].append(holiday)
    cal["holidays"].sort(key=lambda h: h["date"])

    _save_calendar(cal)
    return {"status": "added", "holiday": holiday}


MAX_RESOURCE_TEXT = 120000


def list_resources_tool():
    """List all uploaded supporting documents from the documents directory."""
    if not os.path.isdir(DOCUMENTS_DIR):
        return {"documents": [], "message": "No documents directory found. Upload documents in Settings > Resources."}

    documents = []
    try:
        for fname in sorted(os.listdir(DOCUMENTS_DIR)):
            if fname.endswith('.meta.json'):
                continue
            fpath = os.path.join(DOCUMENTS_DIR, fname)
            if not os.path.isfile(fpath):
                continue

            # Try to load metadata
            meta_path = fpath + ".meta.json"
            meta = {}
            if os.path.exists(meta_path):
                try:
                    with open(meta_path, 'r', encoding='utf-8') as f:
                        meta = json.load(f)
                except Exception:
                    pass

            size_kb = round(os.path.getsize(fpath) / 1024, 1)
            documents.append({
                "filename": fname,
                "doc_type": meta.get("doc_type", "unknown"),
                "description": meta.get("description", ""),
                "size_kb": size_kb,
            })
    except Exception as e:
        return {"error": f"Error reading documents directory: {e}"}

    return {"documents": documents, "total": len(documents)}


def read_resource_tool(filename):
    """Read and return the text content of an uploaded document."""
    if not filename or not filename.strip():
        return {"error": "No filename provided"}

    # Prevent path traversal
    safe_name = os.path.basename(filename)
    filepath = os.path.join(DOCUMENTS_DIR, safe_name)

    if not os.path.exists(filepath):
        # Also check project root for built-in docs like User_Manual.md
        project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
        alt_path = os.path.join(project_root, safe_name)
        if os.path.exists(alt_path):
            filepath = alt_path
        else:
            return {"error": f"Document not found: {safe_name}. Use list_resources to see available files."}

    ext = os.path.splitext(safe_name)[1].lower()
    pages = None

    try:
        if ext == '.pdf':
            content, pages = _extract_pdf_text(filepath)
        elif ext in ('.docx', '.doc'):
            content = _extract_docx_text(filepath)
        elif ext in ('.txt', '.md'):
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
        else:
            return {"error": f"Unsupported file type: {ext}. Supported: PDF, DOCX, DOC, TXT, MD."}
    except Exception as e:
        return {"error": f"Error reading {safe_name}: {e}"}

    truncated = False
    if len(content) > MAX_RESOURCE_TEXT:
        content = content[:MAX_RESOURCE_TEXT]
        truncated = True

    result = {
        "filename": safe_name,
        "content": content,
    }
    if pages:
        result["pages"] = pages
    if truncated:
        result["warning"] = f"Content truncated to {MAX_RESOURCE_TEXT} characters. Full document is larger."

    # Include metadata if available
    meta_path = filepath + ".meta.json"
    if os.path.exists(meta_path):
        try:
            with open(meta_path, 'r', encoding='utf-8') as f:
                meta = json.load(f)
            result["doc_type"] = meta.get("doc_type", "unknown")
            result["description"] = meta.get("description", "")
        except Exception:
            pass

    return result


def save_assignment_config(title, document_text=None, questions=None, totalPoints=None,
                           effortPoints=None, gradingNotes=None, rubricType=None,
                           customMarkers=None):
    """Save or update an assignment config in Grading Setup.

    Merge-updates: loads existing config if present, applies only the
    provided fields, and writes back. This lets the assistant update
    point values or questions without wiping the rest of the config.
    """
    import time

    if not title or not title.strip():
        return {"error": "Title is required."}

    safe_title = ''.join(c for c in title if c.isalnum() or c in ' -_').strip()
    if not safe_title:
        return {"error": "Title contains no valid characters."}

    os.makedirs(ASSIGNMENTS_DIR, exist_ok=True)
    config_path = os.path.join(ASSIGNMENTS_DIR, safe_title + '.json')

    # Load existing config if present
    existing = {}
    if os.path.exists(config_path):
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                existing = json.load(f)
        except Exception:
            pass

    # Merge updates
    existing["title"] = title
    if document_text is not None:
        existing["importedDoc"] = {
            "text": document_text,
            "html": "",
            "filename": safe_title + ".docx",
            "loading": False,
        }
    if questions is not None:
        # Ensure each question has an id
        ts = int(time.time() * 1000)
        for i, q in enumerate(questions):
            if not q.get("id"):
                q["id"] = ts + i
        existing["questions"] = questions
    if totalPoints is not None:
        existing["totalPoints"] = totalPoints
    if effortPoints is not None:
        existing["effortPoints"] = effortPoints
    if gradingNotes is not None:
        existing["gradingNotes"] = gradingNotes
    if rubricType is not None:
        existing["rubricType"] = rubricType
    if customMarkers is not None:
        existing["customMarkers"] = customMarkers

    # Ensure required fields exist with defaults
    defaults = {
        "subject": "", "totalPoints": 100, "instructions": "",
        "aliases": [], "customMarkers": [], "excludeMarkers": [],
        "gradingNotes": "", "questions": [], "responseSections": [],
        "rubricType": "standard", "customRubric": None,
        "useSectionPoints": False, "sectionTemplate": "Custom",
        "effortPoints": 15, "completionOnly": False,
        "countsTowardsGrade": True,
    }
    for key, default_val in defaults.items():
        if key not in existing:
            existing[key] = default_val

    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(existing, f, indent=2)

    q_count = len(existing.get("questions", []))
    q_total = sum(q.get("points", 0) for q in existing.get("questions", []))

    return {
        "status": "saved",
        "config_name": safe_title,
        "title": title,
        "questions_count": q_count,
        "questions_points": q_total,
        "total_points": existing.get("totalPoints", 100),
        "message": "Assignment config saved with " + str(q_count) + " questions (" + str(q_total) + " pts + " + str(existing.get("effortPoints", 15)) + " effort).",
    }


def send_parent_emails(email_subject, email_body, student_names=None, period=None,
                       zero_submissions=False, dry_run=True):
    """Generate email preview for parents/guardians via Outlook automation.

    When called from the AI assistant, this ALWAYS returns a preview.
    Actual sending is triggered by the frontend confirm action via /api/confirm-send.
    """
    # Programmatic guard: AI assistant must never send directly
    dry_run = True
    # Load parent contacts
    if not os.path.exists(PARENT_CONTACTS_FILE):
        return {"error": "No parent contacts imported. Upload class list in Settings first."}

    try:
        with open(PARENT_CONTACTS_FILE, 'r', encoding='utf-8') as f:
            contacts = json.load(f)
    except Exception as e:
        return {"error": "Failed to load parent contacts: " + str(e)}

    if not contacts:
        return {"error": "Parent contacts file is empty."}

    # Load teacher settings
    settings = _load_settings()
    config = settings.get('config', {})
    email_config = _load_email_config()
    teacher_name = email_config.get('teacher_name', '') or config.get('teacher_name', 'Your Teacher')
    subject_area = config.get('subject', '')
    email_signature = email_config.get('email_signature', '')

    # Build reverse map: normalized student name -> (student_id, contact data)
    name_to_contact = {}
    for student_id, contact in contacts.items():
        sname = contact.get('student_name', '')
        if sname:
            name_to_contact[sname] = (student_id, contact)

    # Resolve target students
    target_students = []  # list of (student_id, contact_data, period)

    if zero_submissions:
        missing_data = get_missing_assignments(period="all")
        if "error" in missing_data:
            return {"error": "Could not check missing assignments: " + missing_data["error"]}

        zero_list = missing_data.get("zero_submission_students", [])
        if not zero_list:
            return {"message": "No students with zero submissions found.", "total_emails": 0}

        for entry in zero_list:
            sname = entry.get("student_name", "")
            speriod = entry.get("period", "")

            # Try exact match first, then fuzzy
            matched = False
            for cname, (sid, cdata) in name_to_contact.items():
                if _fuzzy_name_match(sname, cname):
                    target_students.append((sid, cdata, speriod))
                    matched = True
                    break
            if not matched:
                target_students.append((None, {"student_name": sname}, speriod))

    elif student_names:
        for search_name in student_names:
            matched = False
            for cname, (sid, cdata) in name_to_contact.items():
                if _fuzzy_name_match(search_name, cname):
                    speriod = cdata.get('period', '')
                    target_students.append((sid, cdata, speriod))
                    matched = True
                    break
            if not matched:
                target_students.append((None, {"student_name": search_name}, ""))

    elif period:
        target_period = _normalize_period(period)
        for sid, cdata in contacts.items():
            cperiod = cdata.get('period', '')
            if _normalize_period(cperiod) == target_period:
                target_students.append((sid, cdata, cperiod))

    else:
        return {"error": "Provide student_names, period, or set zero_submissions=true to target students."}

    # Build email payloads
    emails = []
    skipped = []

    for sid, cdata, speriod in target_students:
        parent_emails = cdata.get('parent_emails', [])
        sname = cdata.get('student_name', 'Student')

        if not parent_emails:
            skipped.append(sname)
            continue

        parsed = _parse_student_name(sname)
        parent_name = cdata.get('primary_contact_name', '')
        if not parent_name:
            parent_name = "Parent/Guardian"

        replacements = {
            "student_first_name": parsed["first_name"],
            "student_last_name": parsed["last_name"],
            "student_name": parsed["full_name"],
            "parent_name": parent_name,
            "period": speriod or cdata.get('period', ''),
            "teacher_name": teacher_name,
            "subject_area": subject_area,
        }

        filled_subject = _fill_email_template(email_subject, replacements)
        filled_body = _fill_email_template(email_body, replacements)

        # Append teacher signature
        if email_signature:
            filled_body += "\n\n" + email_signature
        elif teacher_name and teacher_name != 'Your Teacher':
            filled_body += "\n\n" + teacher_name

        to_email = parent_emails[0]
        cc_emails = parent_emails[1:] if len(parent_emails) > 1 else []

        emails.append({
            "to": to_email,
            "cc": ', '.join(cc_emails) if cc_emails else '',
            "subject": filled_subject,
            "body": filled_body,
            "student_name": sname,
        })

    if not emails and skipped:
        return {
            "error": "No parent emails found for any targeted students.",
            "skipped_students": skipped,
        }
    if not emails:
        return {"error": "No matching students found."}

    # Dry run: return preview
    if dry_run:
        previews = []
        for e in emails[:3]:
            previews.append({
                "to": e["to"],
                "cc": e.get("cc", ""),
                "subject": e["subject"],
                "body": e["body"][:500] + ("..." if len(e["body"]) > 500 else ""),
                "student_name": e["student_name"],
            })
        # Store pending payload for confirm_and_send tool
        pending_path = os.path.join(os.path.expanduser("~/.graider_data"), "pending_send.json")
        os.makedirs(os.path.dirname(pending_path), exist_ok=True)
        with open(pending_path, 'w') as pf:
            json.dump({"action": "send_parent_emails", "emails": emails}, pf)

        return {
            "dry_run": True,
            "NOT_SENT": True,
            "preview_count": len(previews),
            "total_emails": len(emails),
            "previews": previews,
            "skipped_students": skipped,
            "message": "PREVIEW ONLY — emails have NOT been sent yet. Show this preview to the teacher and ask if they want to send it. If they confirm, call confirm_and_send.",
        }

    # Actually send via Outlook/Playwright
    try:
        from backend.routes.email_routes import launch_outlook_sender
        result = launch_outlook_sender(emails)
        result["skipped_students"] = skipped
        result["total_emails"] = len(emails)
        return result
    except ImportError:
        return {"error": "Outlook sender not available. Check backend installation."}
    except Exception as e:
        return {"error": "Failed to launch Outlook sender: " + str(e)}


def send_focus_comms(email_subject, email_body=None, sms_body=None, student_names=None,
                     period=None, dry_run=True, recipient_type=None):
    """Generate email/SMS preview for parents via Focus SIS Communications.

    When called from the AI assistant, this ALWAYS returns a preview.
    Actual sending is triggered by the frontend confirm action via /api/confirm-send.
    Supports email-only, SMS-only, or both.
    """
    # Programmatic guard: AI assistant must never send directly
    dry_run = True

    if not email_body and not sms_body:
        return {"error": "Provide email_body, sms_body, or both."}

    # Auto-generate SMS notification if email is provided but SMS is not.
    # Default behavior: always send both email + SMS unless SMS-only was requested.
    if email_body and not sms_body:
        settings = _load_settings()
        config = settings.get('config', {})
        email_config = _load_email_config()
        teacher_name = email_config.get('teacher_name', '') or config.get('teacher_name', 'Your Teacher')
        sms_body = "Please check your email for a message regarding " + email_subject + ". -" + teacher_name
    FOCUS_ROSTER_FILE = os.path.expanduser("~/.graider_data/focus_roster_import.json")

    if not os.path.exists(FOCUS_ROSTER_FILE):
        return {"error": "No Focus roster imported. Import roster from Focus in Settings first."}

    try:
        with open(FOCUS_ROSTER_FILE, 'r', encoding='utf-8') as f:
            roster = json.load(f)
    except Exception as e:
        return {"error": "Failed to load Focus roster: " + str(e)}

    # Build flat list of all students from roster periods
    all_students = []  # list of (student_name, period_name)
    periods = roster.get("periods", {})
    for period_name, period_data in periods.items():
        for student in period_data.get("students", []):
            all_students.append((student.get("name", ""), period_name))

    if not all_students:
        return {"error": "Focus roster is empty."}

    # Resolve target students
    target_students = []  # list of (student_name, period_name)

    if student_names:
        for search_name in student_names:
            matched = False
            for sname, pname in all_students:
                if _fuzzy_name_match(search_name, sname):
                    target_students.append((sname, pname))
                    matched = True
                    break
            if not matched:
                target_students.append((search_name, ""))
    elif period:
        target_period = _normalize_period(period)
        for sname, pname in all_students:
            if _normalize_period(pname) == target_period:
                target_students.append((sname, pname))
    else:
        return {"error": "Provide student_names or period to target students."}

    if not target_students:
        return {"error": "No matching students found in Focus roster."}

    # Build messages in focus-comms.js format
    messages = []
    skipped = []

    for sname, pname in target_students:
        parsed = _parse_student_name(sname)

        replacements = {
            "student_first_name": parsed["first_name"],
            "student_last_name": parsed["last_name"],
            "student_name": parsed["full_name"],
        }

        filled_subject = _fill_email_template(email_subject, replacements)
        filled_body = _fill_email_template(email_body, replacements) if email_body else ""
        filled_sms = _fill_email_template(sms_body, replacements) if sms_body else ""

        msg_entry = {
            "student_name": sname,
            "subject": filled_subject,
            "email_body": filled_body,
            "sms_body": filled_sms,
            "cc_emails": [],
        }
        if recipient_type and recipient_type != "Primary Contacts":
            msg_entry["recipient_type"] = recipient_type
        messages.append(msg_entry)

    if not messages:
        return {"error": "No messages to send."}

    # Dry run: return preview
    if dry_run:
        previews = []
        for m in messages[:3]:
            previews.append({
                "student_name": m["student_name"],
                "subject": m["subject"],
                "email_body": m["email_body"][:500] + ("..." if len(m["email_body"]) > 500 else ""),
                "sms_body": m["sms_body"][:200] if m["sms_body"] else "(no SMS)",
            })
        # Store pending payload for confirm_and_send tool
        pending_path = os.path.join(os.path.expanduser("~/.graider_data"), "pending_send.json")
        os.makedirs(os.path.dirname(pending_path), exist_ok=True)
        with open(pending_path, 'w') as pf:
            json.dump({"action": "send_focus_comms", "messages": messages}, pf)

        return {
            "dry_run": True,
            "NOT_SENT": True,
            "preview_count": len(previews),
            "total_messages": len(messages),
            "previews": previews,
            "message": "PREVIEW ONLY — messages have NOT been sent yet. Show this preview to the teacher and ask if they want to send it. If they confirm, call confirm_and_send.",
        }

    # Actually send via focus-comms.js
    try:
        from backend.routes.email_routes import launch_focus_comms
        result = launch_focus_comms(messages)
        result["total_messages"] = len(messages)
        return result
    except ImportError:
        return {"error": "Focus Comms route not available. Check backend installation."}
    except Exception as e:
        return {"error": "Failed to launch Focus Comms: " + str(e)}


def confirm_and_send():
    """Execute the pending send action after teacher confirmation.

    Reads the pending payload saved by send_focus_comms or send_parent_emails,
    then triggers the actual Playwright automation.
    """
    pending_path = os.path.join(os.path.expanduser("~/.graider_data"), "pending_send.json")
    if not os.path.exists(pending_path):
        return {"error": "No pending send action. Generate a preview first using send_focus_comms or send_parent_emails."}

    try:
        with open(pending_path, 'r') as f:
            pending = json.load(f)
    except Exception as e:
        return {"error": "Failed to read pending send: " + str(e)}

    # Remove pending file so it can't be sent twice
    os.remove(pending_path)

    action = pending.get("action")

    try:
        if action == "send_focus_comms":
            from backend.routes.email_routes import launch_focus_comms
            messages = pending.get("messages", [])
            if not messages:
                return {"error": "No messages in pending payload."}
            result = launch_focus_comms(messages)
            result["total_messages"] = len(messages)
            return result
        elif action == "send_parent_emails":
            from backend.routes.email_routes import launch_outlook_sender
            emails = pending.get("emails", [])
            if not emails:
                return {"error": "No emails in pending payload."}
            result = launch_outlook_sender(emails)
            result["total_emails"] = len(emails)
            return result
        else:
            return {"error": f"Unknown pending action: {action}"}
    except Exception as e:
        return {"error": f"Failed to launch send: {str(e)}"}


# ═══════════════════════════════════════════════════════
# EXPORT MAP
# ═══════════════════════════════════════════════════════

REPORT_TOOL_DEFINITIONS = REPORT_TOOL_DEFINITIONS  # re-export for clarity

REPORT_TOOL_HANDLERS = {
    "create_focus_assignment": create_focus_assignment,
    "export_grades_csv": export_grades_csv,
    "lookup_student_info": lookup_student_info,
    "generate_worksheet": generate_worksheet_tool,
    "generate_document": generate_document_tool,
    "generate_csv": generate_csv_tool,
    "save_document_style": save_document_style_tool,
    "list_document_styles": list_document_styles_tool,
    "recommend_next_lesson": recommend_next_lesson,
    "get_standards": get_standards_tool,
    "list_all_standards": list_all_standards_tool,
    "get_recent_lessons": get_recent_lessons,
    "get_calendar": get_calendar,
    "schedule_lesson": schedule_lesson_tool,
    "unschedule_lesson": unschedule_lesson_tool,
    "add_calendar_holiday": add_calendar_holiday,
    "list_resources": list_resources_tool,
    "read_resource": read_resource_tool,
    "save_assignment_config": save_assignment_config,
    "send_parent_emails": send_parent_emails,
    "send_focus_comms": send_focus_comms,
    "confirm_and_send": confirm_and_send,
}
