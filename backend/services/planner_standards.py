"""Standards loading + matching for the planner. Pure logic extracted from planner_routes.py (no Flask)."""
import os
import json
import logging
from pathlib import Path

from backend.services.assignment_post_processing import _extract_usage, _record_planner_cost

_logger = logging.getLogger(__name__)

# Path to standards data
DATA_DIR = Path(__file__).parent.parent / 'data'
DOCUMENTS_DIR = os.path.expanduser("~/.graider_data/documents")

# Image extensions routed to GPT-4o vision in extract_text_from_upload. Single
# source of truth: the route imports this for its lazy-key check so the two
# branches can never diverge (a divergence would let an image reach the service
# with no key and fail in the 500 catch-all instead of the proper 400).
IMAGE_EXTENSIONS = ('.png', '.jpg', '.jpeg', '.gif', '.webp', '.bmp')


def load_support_documents_for_planning() -> str:
    """Load curriculum guides, standards, and other planning documents."""
    if not os.path.exists(DOCUMENTS_DIR):
        return ""

    docs_content = []
    total_chars = 0
    max_chars = 12000  # Increased limit for richer planning context

    # Document types useful for lesson planning (prioritized)
    planning_doc_types = ['curriculum', 'standards', 'pacing_guide', 'textbook', 'assessment', 'general']

    for f in os.listdir(DOCUMENTS_DIR):
        if f.endswith('.meta.json'):
            try:
                with open(os.path.join(DOCUMENTS_DIR, f), 'r') as mf:
                    metadata = json.load(mf)

                doc_type = metadata.get('doc_type', 'general')
                filepath = metadata.get('filepath', '')
                description = metadata.get('description', '')

                # Include all planning-relevant document types
                if doc_type not in planning_doc_types:
                    continue

                if not os.path.exists(filepath):
                    continue

                content = ""
                if filepath.endswith('.txt') or filepath.endswith('.md'):
                    with open(filepath, 'r', encoding='utf-8') as df:
                        content = df.read()
                elif filepath.endswith('.docx'):
                    try:
                        from docx import Document
                        doc = Document(filepath)
                        content = '\n'.join([p.text for p in doc.paragraphs])
                    except Exception:  # noqa: BLE001  # broad catch: error is logged
                        _logger.debug("standards docx extraction failed", exc_info=True)
                        continue
                elif filepath.endswith('.pdf'):
                    try:
                        import fitz
                        pdf = fitz.open(filepath)
                        content = '\n'.join([page.get_text() for page in pdf])
                        pdf.close()
                    except Exception:  # noqa: BLE001  # broad catch: error is logged
                        _logger.debug("standards pdf extraction failed", exc_info=True)
                        continue

                if content and total_chars + len(content) < max_chars:
                    doc_label = doc_type.upper()
                    if description:
                        doc_label += f" - {description}"
                    # Use more content per document (up to 4000 chars)
                    chunk = content[:4000]
                    docs_content.append(f"[{doc_label}]\n{chunk}")
                    total_chars += len(chunk)

            except Exception:  # noqa: BLE001  # broad catch: error is logged
                _logger.debug("support document load failed", exc_info=True)
                continue

    if not docs_content:
        return ""

    return "\n\nREFERENCE DOCUMENTS:\n" + "\n\n".join(docs_content)


# Module-level cache for standards map
_standards_map_cache = None


def _extract_grade_from_code(code):
    """Extract grade level from a standards code across all frameworks."""
    if not code:
        return None
    parts = code.split('.')

    # NGSS: prefix-based (MS-PS1-1, HS-LS1-1)
    if code.startswith('MS-'):
        return 'MS'
    if code.startswith('HS-'):
        return 'HS'

    # CCSS Math: CCSS.MATH.CONTENT.{G}.{DOMAIN}...
    if code.startswith('CCSS.MATH') and len(parts) >= 4:
        return parts[3]

    # CCSS ELA: CCSS.ELA-LITERACY.{STRAND}.{G}...
    if code.startswith('CCSS.ELA') and len(parts) >= 4:
        return parts[3]

    # C3 Social Studies: D2.His.1.6-8
    if code.startswith('D') and len(code) > 1 and code[1:2].isdigit() and len(parts) >= 4:
        return parts[3]

    # FL B.E.S.T., TX TEKS, VA SOL: {SUBJ}.{G}.{DOMAIN}...
    if len(parts) >= 2:
        candidate = parts[1]
        if candidate == 'K12':
            return 'K12'
        if candidate.isdigit() or candidate == 'K':
            return candidate
        if '-' in candidate:
            return candidate

    return None


def _grade_matches(code_grade, requested_grade):
    """Check if extracted grade matches the requested grade."""
    if code_grade is None:
        return False
    if code_grade == 'K12':
        return True
    if code_grade == requested_grade:
        return True
    if code_grade == 'MS' and requested_grade in ('6', '7', '8'):
        return True
    if code_grade == 'HS' and requested_grade in ('9', '10', '11', '12'):
        return True
    if '-' in str(code_grade):
        try:
            lo, hi = code_grade.split('-')
            req = int(requested_grade) if requested_grade.isdigit() else 0
            return int(lo) <= req <= int(hi)
        except (ValueError, IndexError):
            _logger.debug("Unparseable grade range %r in standards filter", code_grade)
    if code_grade == '912' and requested_grade in ('9', '10', '11', '12'):
        return True
    return False


def _get_standards_map():
    """Load and cache standards_map.json."""
    global _standards_map_cache
    if _standards_map_cache is None:
        map_path = DATA_DIR / 'standards' / 'standards_map.json'
        if map_path.exists():
            with open(map_path, 'r') as f:
                _standards_map_cache = json.load(f)
        else:
            _standards_map_cache = {}
    return _standards_map_cache


def _load_standards_file(filepath):
    """Load standards from a JSON file. Returns list or empty list."""
    if not filepath.exists():
        return []
    try:
        with open(filepath, 'r', encoding='utf-8') as f:
            data = json.load(f)
        if isinstance(data, list):
            return data
        return data.get('standards', [])
    except Exception:  # noqa: BLE001  # broad catch: returns fallback
        return []


def load_standards(state, subject, grade=None):
    """Load standards with mapping-based resolution and fallback.

    Returns dict: {standards, fallback_used, fallback_framework, no_framework, state_note}
    """
    result = {
        'standards': [],
        'fallback_used': False,
        'fallback_framework': None,
        'no_framework': False,
        'state_note': None,
    }

    smap = _get_standards_map()
    states = smap.get('states', {})
    subject_fallbacks = smap.get('subject_fallbacks', {})
    subject_to_filename = smap.get('subject_to_filename', {})

    # Look up state config
    state_config = states.get(state.upper(), {}) if state else {}
    framework = state_config.get('framework', 'ccss')
    result['state_note'] = state_config.get('note')

    # Map subject to filename
    filename = subject_to_filename.get(subject)
    if not filename:
        filename = subject.lower().replace(' ', '_').replace('/', '-')

    # Try primary path: standards/{framework}/{filename}.json
    primary_path = DATA_DIR / 'standards' / framework / (filename + '.json')
    standards = _load_standards_file(primary_path)

    if not standards:
        # Primary file not found — try subject-specific fallback
        # This handles: CCSS states needing NGSS for science, state-specific
        # frameworks missing certain subjects, etc.
        fallback_fw = subject_fallbacks.get(subject)
        if fallback_fw is None:
            # No fallback defined for this subject (e.g., Spanish, World Languages)
            if framework not in ('ccss', 'ngss'):
                result['no_framework'] = True
        elif fallback_fw and fallback_fw != framework:
            fallback_path = DATA_DIR / 'standards' / fallback_fw / (filename + '.json')
            standards = _load_standards_file(fallback_path)
            if standards:
                result['fallback_used'] = True
                result['fallback_framework'] = fallback_fw

    if not standards:
        # Legacy fallback: standards_{state}_{subject}.json
        subject_clean = subject.lower().replace(' ', '_').replace('/', '-')
        legacy_path = DATA_DIR / ('standards_' + state.lower() + '_' + subject_clean + '.json')
        standards = _load_standards_file(legacy_path)

    # Filter by grade
    if grade and standards:
        filtered = [s for s in standards if _grade_matches(_extract_grade_from_code(s.get('code', '')), str(grade))]
        if filtered:
            standards = filtered
        elif not filtered:
            # Preserve existing high school course mapping for FL
            GRADE_TO_COURSE = {
                'math': {'9': 'Algebra 1', '10': 'Geometry', '11': 'Algebra 2', '12': 'Pre-Calculus'},
                'science': {'9': 'Biology', '10': 'Chemistry', '11': 'Physics', '12': 'Earth/Space Science'},
            }
            subject_clean = subject.lower().replace(' ', '_').replace('/', '-')
            courses = GRADE_TO_COURSE.get(subject_clean, {})
            if str(grade) in courses:
                course_filtered = [s for s in standards if s.get('course') == courses[str(grade)]]
                if course_filtered:
                    standards = course_filtered

    result['standards'] = standards
    return result


def rewrite_for_alignment_content(*, enriched_questions, doc_text, grade, subject, api_key):
    """Call the LLM to rewrite questions for standards alignment; return the
    AI result dict + usage, or {'error': ...} on a non-JSON response (the route
    jsonifies either). Wave 6 Slice 8 - extracted from planner_routes.
    """
    from backend.services.llm_adapter import LLMRequest, Message, OpenAIAdapter, ResponseFormat, TextPart
    adapter = OpenAIAdapter(api_key=api_key)

    system_prompt = (
        f"You are an expert curriculum specialist for grade {grade} {subject}. "
        "Rewrite the given questions to better align with the target standards. "
        "Preserve the general topic and difficulty level but adjust the focus, "
        "vocabulary, and cognitive demand to match the standard's benchmark. "
        "Keep the question appropriate for the grade level."
    )

    user_prompt = json.dumps({
        "task": "Rewrite each question to better align with its target standard.",
        "document_context": doc_text[:3000],
        "questions": enriched_questions,
        "return_format": {
            "rewrites": [{"original_text": "str", "rewritten_text": "str", "standard_code": "str", "change_explanation": "brief explanation of what changed and why"}]
        }
    })

    completion = adapter.chat(LLMRequest(
        model="gpt-4o-mini",
        system_prompt=system_prompt,
        messages=[Message(role="user", content=[TextPart(text=user_prompt)])],
        response_format=ResponseFormat(type="json_object"),
        temperature=0.3,
        metadata={"feature_label": "rewrite_for_alignment"},
    ))

    raw_content = completion.content_parts[0].text if completion.content_parts else "{}"
    try:
        result = json.loads(raw_content)
    except json.JSONDecodeError:
        _logger.warning("[rewrite-alignment] Non-JSON response: %s", raw_content[:500])
        return {"error": "AI returned non-JSON response. Possibly rate limited."}

    usage = _extract_usage(completion, "gpt-4o-mini")
    _record_planner_cost(usage)

    return {**result, "usage": usage}


def align_document_to_standards_content(*, doc_text, standards_ref, api_key):
    """Call the LLM to analyze a document against the given standards reference;
    return the alignment result dict + usage, or {'error': ...} on a non-JSON
    response (the route jsonifies either). Wave 6 Slice 9 - extracted from
    planner_routes.
    """
    from backend.services.llm_adapter import LLMRequest, Message, OpenAIAdapter, ResponseFormat, TextPart
    adapter = OpenAIAdapter(api_key=api_key)

    # Truncate document to fit in context with standards
    truncated_doc = doc_text[:8000]

    system_prompt = (
        "You are an expert curriculum alignment specialist. "
        "Analyze the educational document against the provided standards and return a detailed alignment analysis in JSON format. "
        "Be specific about what content in the document maps to each standard. "
        "Only include standards with at least some relevance (confidence > 0.2). "
        "Sort matched_standards by confidence descending."
    )

    user_prompt = json.dumps({
        "task": "Analyze this educational document and identify which standards it aligns to.",
        "document_text": truncated_doc,
        "available_standards": standards_ref,
        "return_format": {
            "matched_standards": [{"code": "str", "benchmark": "str", "confidence": "float 0.0-1.0", "evidence": "brief quote or description from document", "alignment_notes": "what is well-covered vs missing"}],
            "unmatched_standards": ["standard codes not covered"],
            "overall_alignment_score": "float 0.0-1.0",
            "suggestions": ["improvement suggestion strings"],
            "question_analysis": [{"question_text": "truncated question", "aligned_standard": "code or null", "alignment_quality": "strong|partial|weak|none", "rewrite_suggestion": "optional string or null"}]
        }
    })

    completion = adapter.chat(LLMRequest(
        model="gpt-4o",
        system_prompt=system_prompt,
        messages=[Message(role="user", content=[TextPart(text=user_prompt)])],
        response_format=ResponseFormat(type="json_object"),
        temperature=0.3,
        metadata={"feature_label": "align_document_to_standards"},
    ))

    raw_content = completion.content_parts[0].text if completion.content_parts else "{}"
    try:
        result = json.loads(raw_content)
    except json.JSONDecodeError:
        _logger.warning("[align-standards] Non-JSON response: %s", raw_content[:500])
        return {"error": "AI returned non-JSON response. Possibly rate limited."}

    usage = _extract_usage(completion, "gpt-4o")
    _record_planner_cost(usage)

    return {**result, "usage": usage}


class TextExtractionError(Exception):
    """Raised by extract_text_from_upload for a file type that passed the route's
    ALLOWED_DOC_EXTENSIONS gate but has no extraction branch (.doc/.rtf). The route
    maps this to a 400 with this message. Wave 6 Slice 10.
    """


def extract_text_from_upload(*, file_data, filename, api_key):
    """Extract plain text from an uploaded document (docx/pdf/txt) or image
    (png/jpg/...). Pure dispatch (Flask-free): the route validates the upload,
    resolves the OpenAI key for images (and owns the missing-key 400), and passes
    the bytes + filename + key in. Returns the extracted text string; raises
    TextExtractionError for an allowed-but-unhandled type. Wave 6 Slice 10 -
    extracted from planner_routes.
    """
    # Documents — extract text directly
    if filename.endswith('.docx'):
        import io
        from docx import Document
        doc = Document(io.BytesIO(file_data))
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([c.text.strip() for c in row.cells if c.text.strip()])
                if row_text:
                    text_parts.append(row_text)
        return '\n'.join(text_parts)

    elif filename.endswith('.pdf'):
        import io
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_data)) as pdf:
                text_parts = [page.extract_text() or '' for page in pdf.pages]
            return '\n'.join(text_parts).strip()
        except ImportError:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(file_data))
            text_parts = [page.extract_text() or '' for page in reader.pages]
            return '\n'.join(text_parts).strip()

    elif filename.endswith('.txt'):
        return file_data.decode('utf-8', errors='replace')

    # Images — use GPT-4o vision to extract text
    elif filename.endswith(IMAGE_EXTENSIONS):
        import base64

        ext = filename.rsplit('.', 1)[-1]
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "gif": "image/gif", "webp": "image/webp", "bmp": "image/bmp"}.get(ext, "image/png")
        b64 = base64.b64encode(file_data).decode('utf-8')

        from backend.services.llm_adapter import ImagePart, LLMRequest, Message, OpenAIAdapter, TextPart
        adapter = OpenAIAdapter(api_key=api_key)
        completion = adapter.chat(LLMRequest(
            model="gpt-4o",
            system_prompt="Extract ALL text from this image. Return only the extracted text, preserving paragraphs and structure. Do not add commentary.",
            messages=[Message(role="user", content=[
                ImagePart(url=None, base64=b64, mime_type=mime),
                TextPart(text="Extract all text from this image."),
            ])],
            max_tokens=4000,
            temperature=0,
            metadata={"feature_label": "extract_text_image"},
        ))

        _record_planner_cost(_extract_usage(completion, "gpt-4o"))

        return (completion.content_parts[0].text if completion.content_parts else "").strip()

    else:
        raise TextExtractionError("Unsupported file type. Use .docx, .pdf, .txt, .png, .jpg, or .jpeg")
